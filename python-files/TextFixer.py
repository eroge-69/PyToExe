# textfixer_pro.py
import sys
import re
import os
import difflib
import time
import torch
from PySide6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
                            QSplitter, QPlainTextEdit, QPushButton, QLabel, QStatusBar,
                            QFileDialog, QDialog, QListWidget, QLineEdit, QMessageBox,
                            QProgressBar, QSizePolicy, QMenu)
from PySide6.QtCore import Qt, QThread, Signal, QPropertyAnimation, QEasingCurve, QSize
from PySide6.QtGui import QTextCharFormat, QColor, QFont, QSyntaxHighlighter, QTextDocument, QPalette, QIcon, QAction
from transformers import T5ForConditionalGeneration, T5Tokenizer
from pymorphy2 import MorphAnalyzer
from natasha import Segmenter, MorphVocab, NewsEmbedding, NewsMorphTagger, Doc
from docx import Document

# ======================
# СИСТЕМНЫЕ НАСТРОЙКИ
# ======================
DARK_THEME = True
DICTIONARY_FILE = "user_dictionary.txt"
MAX_TEXT_LENGTH = 10000  # Макс. длина текста для обработки

# ======================
# ВСПОМОГАТЕЛЬНЫЕ КЛАССЫ
# ======================
class TextHighlighter(QSyntaxHighlighter):
    def __init__(self, document):
        super().__init__(document)
        self.errors = []
        
    def set_errors(self, errors):
        self.errors = errors
        self.rehighlight()
        
    def highlightBlock(self, text):
        error_format = QTextCharFormat()
        error_format.setUnderlineColor(QColor("#ff6b6b"))
        error_format.setUnderlineStyle(QTextCharFormat.WaveUnderline)
        error_format.setFontUnderline(True)
        
        for start, end in self.errors:
            if start <= self.currentBlock().position() + len(text) and end >= self.currentBlock().position():
                start_pos = max(0, start - self.currentBlock().position())
                end_pos = min(len(text), end - self.currentBlock().position())
                length = end_pos - start_pos
                if length > 0:
                    self.setFormat(start_pos, length, error_format)

class CheckThread(QThread):
    progress = Signal(int)
    result_ready = Signal(str, list, dict)
    error_occurred = Signal(str)
    
    def __init__(self, text, model, tokenizer, device, user_dictionary, segmenter, morph_vocab, morph_tagger):
        super().__init__()
        self.text = text
        self.model = model
        self.tokenizer = tokenizer
        self.device = device
        self.user_dictionary = user_dictionary
        self.segmenter = segmenter
        self.morph_vocab = morph_vocab
        self.morph_tagger = morph_tagger
        
    def run(self):
        try:
            # Шаг 1: Предварительная обработка
            self.progress.emit(10)
            doc = Doc(self.text)
            doc.segment(self.segmenter)
            doc.tag_morph(self.morph_tagger)
            
            # Шаг 2: Разбивка на предложения
            sentences = [sent.text for sent in doc.sents]
            if not sentences:
                self.result_ready.emit(self.text, [], {"words": 0, "unique": 0, "sentences": 0, "errors": 0})
                return
                
            # Шаг 3: Обработка каждого предложения
            corrected_sentences = []
            all_errors = []
            total_sentences = len(sentences)
            
            for i, sent in enumerate(sentences):
                # Пропускаем предложения из пользовательского словаря
                if any(word.lower() in self.user_dictionary for word in re.findall(r'\w+', sent)):
                    corrected_sentences.append(sent)
                    continue
                    
                # Обработка модели
                input_ids = self.tokenizer.encode("ocr: " + sent, return_tensors="pt", 
                                                max_length=256, truncation=True).to(self.device)
                outputs = self.model.generate(input_ids, max_length=256)
                corrected = self.tokenizer.decode(outputs[0], skip_special_tokens=True)
                corrected_sentences.append(corrected)
                
                # Поиск различий
                matcher = difflib.SequenceMatcher(None, sent, corrected)
                for tag, i1, i2, j1, j2 in matcher.get_opcodes():
                    if tag != 'equal':
                        start = sum(len(s) for s in sentences[:i]) + i1
                        end = start + (i2 - i1)
                        all_errors.append((start, end))
                
                # Прогресс
                self.progress.emit(10 + int(80 * (i + 1) / total_sentences))
            
            # Шаг 4: Сборка результата
            corrected_text = " ".join(corrected_sentences)
            
            # Шаг 5: Статистика
            words = re.findall(r'\w+', self.text.lower())
            stats = {
                "words": len(words),
                "unique": len(set(words)),
                "sentences": len(sentences),
                "errors": len(all_errors)
            }
            
            self.progress.emit(95)
            self.result_ready.emit(corrected_text, all_errors, stats)
            self.progress.emit(100)
            
        except Exception as e:
            self.error_occurred.emit(f"Ошибка обработки: {str(e)}")

# ======================
# ГЛАВНОЕ ОКНО ПРИЛОЖЕНИЯ
# ======================
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("TextFixer Pro")
        self.setGeometry(100, 100, 1200, 700)
        self.user_dictionary = self.load_dictionary()
        self.model = None
        self.tokenizer = None
        self.device = self.get_device()
        self.current_errors = []
        self.init_ui()
        self.init_nlp()
        self.load_ai_model()
        self.apply_theme(DARK_THEME)
        
    def get_device(self):
        return torch.device("cuda" if torch.cuda.is_available() else "cpu")
    
    def init_ui(self):
        # Центральный виджет
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        main_layout.setContentsMargins(15, 15, 15, 10)
        main_layout.setSpacing(15)
        
        # Панель текстовых редакторов
        text_panel = QSplitter(Qt.Horizontal)
        text_panel.setHandleWidth(8)
        text_panel.setStyleSheet("""
            QSplitter::handle {
                background: #5c5c5c;
                border: 1px solid #444;
                border-radius: 3px;
            }
        """)
        
        # Левая панель (ввод)
        self.input_edit = QPlainTextEdit()
        self.input_edit.setPlaceholderText("Введите текст для проверки...")
        self.input_edit.setStyleSheet("""
            QPlainTextEdit {
                background-color: #ffffff;
                color: #333;
                border-radius: 10px;
                padding: 15px;
                font-size: 14px;
                border: 1px solid #ccc;
            }

            QPushButton {
                background-color: #0078d7;
                color: white;
                border-radius: 8px;
                padding: 10px 18px;
                font-size: 14px;
                font-weight: bold;
            }

            QPushButton:hover {
                background-color: #0063b1;
            }

            QPushButton:pressed {
                background-color: #005ba1;
            }

        """)
        self.input_highlighter = TextHighlighter(self.input_edit.document())
        text_panel.addWidget(self.input_edit)
        
        # Правая панель (результат)
        self.output_edit = QPlainTextEdit()
        self.output_edit.setReadOnly(True)
        self.output_edit.setPlaceholderText("Здесь появится исправленный текст...")
        self.output_edit.setStyleSheet("""
            QPlainTextEdit {
                background: #252526;
                color: #dcdcdc;
                border: 1px solid #3c3c3c;
                border-radius: 5px;
                padding: 10px;
                font-size: 14px;
            }
        """)
        text_panel.addWidget(self.output_edit)
        
        text_panel.setSizes([600, 600])
        main_layout.addWidget(text_panel)
        
        # Панель инструментов
        toolbar = QHBoxLayout()
        toolbar.setSpacing(12)
        
        # Кнопки с иконками
        self.create_button("Проверить", self.check_text, "#4caf50", "check-circle")
        self.create_button("Автозамена", self.auto_correct, "#2196f3", "edit")
        self.create_button("Заменить все", self.replace_all, "#ff9800", "copy")
        self.create_button("Сохранить", self.save_result, "#9c27b0", "save")
        self.create_button("Загрузить", self.load_file, "#607d8b", "folder")
        self.create_button("Словарь", self.manage_dictionary, "#e91e63", "book")
        
        # Добавляем кнопки в панель
        for btn in [self.check_btn, self.autocorrect_btn, self.replace_btn, 
                    self.save_btn, self.load_btn, self.dict_btn]:
            toolbar.addWidget(btn)
            
        main_layout.addLayout(toolbar)
        
        # Статус бар
        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setFixedWidth(200)
        self.progress_bar.setVisible(False)
        self.status_bar.addPermanentWidget(self.progress_bar)
        
        self.stats_label = QLabel("Ошибок: 0 | Слов: 0 | Уникальных: 0 | Предложений: 0")
        self.status_bar.addPermanentWidget(self.stats_label)
        
        # Меню
        self.init_menu()
        
    def create_button(self, text, handler, color, icon_name):
        btn = QPushButton(text)
        btn.setMinimumHeight(40)
        btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {color};
                color: white;
                font-weight: bold;
                border-radius: 8px;
                padding: 8px 16px;
                font-size: 14px;
                border: none;
            }}
            QPushButton:hover {{
                background-color: {self.adjust_color(color, 20)};
            }}
            QPushButton:pressed {{
                background-color: {self.adjust_color(color, -20)};
            }}
            QPushButton:disabled {{
                background-color: #aaaaaa;
            }}
        """)
        btn.clicked.connect(handler)
        
        # Анимация при наведении
        anim = QPropertyAnimation(btn, b"geometry")
        anim.setDuration(200)
        anim.setEasingCurve(QEasingCurve.OutQuad)
        
        def enter_event(event):
            anim.setStartValue(btn.geometry())
            anim.setEndValue(btn.geometry().adjusted(-2, -2, 2, 2))
            anim.start()
            
        def leave_event(event):
            anim.setStartValue(btn.geometry())
            anim.setEndValue(btn.geometry().adjusted(2, 2, -2, -2))
            anim.start()
            
        btn.enterEvent = enter_event
        btn.leaveEvent = leave_event
        
        # Сохраняем ссылки на кнопки
        if text == "Проверить":
            self.check_btn = btn
        elif text == "Автозамена":
            self.autocorrect_btn = btn
        elif text == "Заменить все":
            self.replace_btn = btn
        elif text == "Сохранить":
            self.save_btn = btn
        elif text == "Загрузить":
            self.load_btn = btn
        elif text == "Словарь":
            self.dict_btn = btn
            
        return btn
        
    def init_menu(self):
        menu_bar = self.menuBar()
        
        # Меню Файл
        file_menu = menu_bar.addMenu("Файл")
        
        new_action = QAction("Новый", self)
        new_action.triggered.connect(self.new_document)
        file_menu.addAction(new_action)
        
        load_action = QAction("Загрузить текст", self)
        load_action.triggered.connect(self.load_file)
        file_menu.addAction(load_action)
        
        save_action = QAction("Сохранить результат", self)
        save_action.triggered.connect(self.save_result)
        file_menu.addAction(save_action)
        
        file_menu.addSeparator()
        
        exit_action = QAction("Выход", self)
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)
        
        # Меню Вид
        view_menu = menu_bar.addMenu("Вид")
        
        theme_menu = QMenu("Тема", self)
        
        light_action = QAction("Светлая", self)
        light_action.triggered.connect(lambda: self.apply_theme(False))
        
        dark_action = QAction("Темная", self)
        dark_action.triggered.connect(lambda: self.apply_theme(True))
        
        theme_menu.addAction(light_action)
        theme_menu.addAction(dark_action)
        view_menu.addMenu(theme_menu)
        
    def init_nlp(self):
        self.segmenter = Segmenter()
        self.morph_vocab = MorphVocab()
        emb = NewsEmbedding()
        self.morph_tagger = NewsMorphTagger(emb)
        self.morph = MorphAnalyzer()
        
    def load_ai_model(self):
        try:
            self.status_bar.showMessage("Загрузка модели исправления...")
            model_name = model_name = "UrukHan/t5-russian-spell" 
            self.tokenizer = T5Tokenizer.from_pretrained(model_name)
            self.model = T5ForConditionalGeneration.from_pretrained(model_name).to(self.device)
            self.status_bar.showMessage(f"Модель загружена ({'GPU' if torch.cuda.is_available() else 'CPU'})", 3000)
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось загрузить модель: {str(e)}")
            
    def load_dictionary(self):
        dictionary = set()
        if os.path.exists(DICTIONARY_FILE):
            try:
                with open(DICTIONARY_FILE, "r", encoding="utf-8") as f:
                    for line in f:
                        word = line.strip().lower()
                        if word:
                            dictionary.add(word)
            except Exception as e:
                print(f"Ошибка загрузки словаря: {e}")
        return dictionary
        
    def save_dictionary(self):
        try:
            with open(DICTIONARY_FILE, "w", encoding="utf-8") as f:
                for word in sorted(self.user_dictionary):
                    f.write(word + "\n")
            return True
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить словарь: {str(e)}")
            return False
            
    def apply_theme(self, dark=True):
        palette = QPalette()
    
        if dark:
            # Темная тема
            palette.setColor(QPalette.Window, QColor("#1e1e1e"))
            palette.setColor(QPalette.WindowText, QColor("#dcdcdc"))
            palette.setColor(QPalette.Base, QColor("#2d2d30"))
            palette.setColor(QPalette.AlternateBase, QColor("#252526"))
            palette.setColor(QPalette.ToolTipBase, QColor("#1e1e1e"))
            palette.setColor(QPalette.ToolTipText, Qt.white)
            palette.setColor(QPalette.Text, QColor("#dcdcdc"))
            palette.setColor(QPalette.Button, QColor("#3c3c3c"))
            palette.setColor(QPalette.ButtonText, QColor("#dcdcdc"))
            palette.setColor(QPalette.BrightText, Qt.red)
            palette.setColor(QPalette.Highlight, QColor("#0078d7"))
            palette.setColor(QPalette.HighlightedText, Qt.white)

            self.input_edit.setStyleSheet("QPlainTextEdit { background-color: #2d2d30; color: #dcdcdc; }")
            self.output_edit.setStyleSheet("QPlainTextEdit { background-color: #252526; color: #dcdcdc; }")
        else:
            # Светлая тема
            palette.setColor(QPalette.Window, QColor("#f0f0f0"))
            palette.setColor(QPalette.WindowText, Qt.black)
            palette.setColor(QPalette.Base, Qt.white)
            palette.setColor(QPalette.AlternateBase, QColor("#f0f0f0"))
            palette.setColor(QPalette.ToolTipBase, Qt.white)
            palette.setColor(QPalette.ToolTipText, Qt.black)
            palette.setColor(QPalette.Text, Qt.black)
            palette.setColor(QPalette.Button, QColor("#e0e0e0"))
            palette.setColor(QPalette.ButtonText, Qt.black)
            palette.setColor(QPalette.BrightText, Qt.red)
            palette.setColor(QPalette.Highlight, QColor("#0078d7"))
            palette.setColor(QPalette.HighlightedText, Qt.white)

            self.input_edit.setStyleSheet("QPlainTextEdit { background-color: white; color: black; }")
            self.output_edit.setStyleSheet("QPlainTextEdit { background-color: white; color: black; }")

        QApplication.setPalette(palette)
        self.repaint()

        
    def adjust_color(self, hex_color, amount):
        """Осветлить или затемнить цвет"""
        hex_color = hex_color.lstrip('#')
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)

        
        r = max(0, min(255, r + amount))
        g = max(0, min(255, g + amount))
        b = max(0, min(255, b + amount))
        
        return f"#{r:02x}{g:02x}{b:02x}"
        
    # ======================
    # ОСНОВНЫЕ ФУНКЦИИ
    # ======================
    def check_text(self):
        text = self.input_edit.toPlainText().strip()
        if not text:
            QMessageBox.warning(self, "Внимание", "Введите текст для проверки")
            return
            
        if len(text) > MAX_TEXT_LENGTH:
            QMessageBox.warning(self, "Внимание", f"Текст слишком длинный (макс. {MAX_TEXT_LENGTH} символов)")
            return
            
        # Блокируем кнопки во время обработки
        self.set_buttons_enabled(False)
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        self.status_bar.showMessage("Проверка текста...")
        
        # Запуск фонового потока
        self.thread = CheckThread(
            text, 
            self.model, 
            self.tokenizer, 
            self.device, 
            self.user_dictionary,
            self.segmenter,
            self.morph_vocab,
            self.morph_tagger
        )
        self.thread.progress.connect(self.progress_bar.setValue)
        self.thread.result_ready.connect(self.handle_check_result)
        self.thread.error_occurred.connect(self.handle_check_error)
        self.thread.finished.connect(self.thread.deleteLater)
        self.thread.start()
        
    def handle_check_result(self, corrected_text, errors, stats):
        self.output_edit.setPlainText(corrected_text)
        self.current_errors = errors
        self.input_highlighter.set_errors(errors)
        
        # Обновляем статистику
        self.stats_label.setText(
            f"Ошибок: {stats['errors']} | "
            f"Слов: {stats['words']} | "
            f"Уникальных: {stats['unique']} | "
            f"Предложений: {stats['sentences']}"
        )
        
        # Разблокируем кнопки
        self.set_buttons_enabled(True)
        self.progress_bar.setVisible(False)
        self.status_bar.showMessage("Проверка завершена", 3000)
        
    def handle_check_error(self, message):
        QMessageBox.critical(self, "Ошибка", message)
        self.set_buttons_enabled(True)
        self.progress_bar.setVisible(False)
        
    def auto_correct(self):
        if not self.current_errors:
            QMessageBox.information(self, "Автозамена", "Ошибки не найдены")
            return
            
        # Заменяем текст в поле ввода на исправленный
        self.input_edit.setPlainText(self.output_edit.toPlainText())
        self.input_highlighter.set_errors([])
        self.current_errors = []
        self.status_bar.showMessage("Автозамена применена", 2000)
        
    def replace_all(self):
        if not self.output_edit.toPlainText():
            QMessageBox.warning(self, "Внимание", "Нет исправленного текста")
            return
            
        self.input_edit.setPlainText(self.output_edit.toPlainText())
        self.status_bar.showMessage("Текст заменен", 2000)
        
    def save_result(self):
        text = self.output_edit.toPlainText()
        if not text:
            QMessageBox.warning(self, "Внимание", "Нет текста для сохранения")
            return
            
        path, _ = QFileDialog.getSaveFileName(
            self,
            "Сохранить результат",
            "",
            "Текстовые файлы (*.txt);;Документы Word (*.docx)"
        )
        
        if not path:
            return
            
        try:
            if path.endswith('.docx'):
                doc = Document()
                doc.add_paragraph(text)
                doc.save(path)
            else:
                with open(path, 'w', encoding='utf-8') as f:
                    f.write(text)
                    
            self.status_bar.showMessage(f"Файл сохранен: {os.path.basename(path)}", 3000)
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить файл: {str(e)}")
            
    def load_file(self):
        path, _ = QFileDialog.getOpenFileName(
            self,
            "Загрузить файл",
            "",
            "Текстовые файлы (*.txt);;Документы Word (*.docx)"
        )
        
        if not path:
            return
            
        try:
            if path.endswith('.docx'):
                doc = Document(path)
                text = "\n".join([para.text for para in doc.paragraphs])
            else:
                with open(path, 'r', encoding='utf-8') as f:
                    text = f.read()
                    
            self.input_edit.setPlainText(text)
            self.status_bar.showMessage(f"Файл загружен: {os.path.basename(path)}", 3000)
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось загрузить файл: {str(e)}")
            
    def manage_dictionary(self):
        dialog = DictionaryDialog(self.user_dictionary, self)
        if dialog.exec() == QDialog.Accepted:
            self.user_dictionary = dialog.get_dictionary()
            self.save_dictionary()
            
    def new_document(self):
        self.input_edit.clear()
        self.output_edit.clear()
        self.input_highlighter.set_errors([])
        self.stats_label.setText("Ошибок: 0 | Слов: 0 | Уникальных: 0 | Предложений: 0")
        self.status_bar.showMessage("Новый документ создан", 2000)
        
    def set_buttons_enabled(self, enabled):
        self.check_btn.setEnabled(enabled)
        self.autocorrect_btn.setEnabled(enabled)
        self.replace_btn.setEnabled(enabled)
        self.save_btn.setEnabled(enabled)
        self.dict_btn.setEnabled(enabled)
        
# ======================
# ДИАЛОГ СЛОВАРЯ
# ======================
class DictionaryDialog(QDialog):
    def __init__(self, dictionary, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Управление словарем")
        self.setGeometry(300, 300, 500, 400)
        self.dictionary = dictionary.copy()
        self.init_ui()
        
    def init_ui(self):
        layout = QVBoxLayout()
        
        # Панель добавления
        add_layout = QHBoxLayout()
        self.word_input = QLineEdit()
        self.word_input.setPlaceholderText("Введите новое слово...")
        add_layout.addWidget(self.word_input)
        
        add_btn = QPushButton("Добавить")
        add_btn.clicked.connect(self.add_word)
        add_btn.setStyleSheet("""
            QPushButton {
                background-color: #4caf50;
                color: white;
                font-weight: bold;
                padding: 6px 12px;
                border-radius: 4px;
            }
        """)
        add_layout.addWidget(add_btn)
        layout.addLayout(add_layout)
        
        # Список слов
        self.word_list = QListWidget()
        self.word_list.addItems(sorted(self.dictionary))
        self.word_list.setStyleSheet("""
            QListWidget {
                background: #2d2d30;
                color: #dcdcdc;
                border: 1px solid #3c3c3c;
                border-radius: 5px;
                padding: 5px;
            }
        """)
        layout.addWidget(self.word_list)
        
        # Панель удаления
        del_btn = QPushButton("Удалить выбранное")
        del_btn.clicked.connect(self.delete_word)
        del_btn.setStyleSheet("""
            QPushButton {
                background-color: #f44336;
                color: white;
                font-weight: bold;
                padding: 8px;
                border-radius: 4px;
            }
        """)
        layout.addWidget(del_btn)
        
        # Кнопки диалога
        button_layout = QHBoxLayout()
        ok_btn = QPushButton("Сохранить")
        ok_btn.clicked.connect(self.accept)
        button_layout.addWidget(ok_btn)
        
        cancel_btn = QPushButton("Отмена")
        cancel_btn.clicked.connect(self.reject)
        button_layout.addWidget(cancel_btn)
        
        layout.addLayout(button_layout)
        self.setLayout(layout)
        
    def add_word(self):
        word = self.word_input.text().strip().lower()
        if not word:
            QMessageBox.warning(self, "Ошибка", "Введите слово")
            return
            
        if word in self.dictionary:
            QMessageBox.warning(self, "Ошибка", "Слово уже в словаре")
            return
            
        self.dictionary.add(word)
        self.word_list.addItem(word)
        self.word_list.sortItems()
        self.word_input.clear()
        
    def delete_word(self):
        selected = self.word_list.currentItem()
        if not selected:
            return
            
        word = selected.text()
        self.dictionary.remove(word)
        self.word_list.takeItem(self.word_list.row(selected))
        
    def get_dictionary(self):
        return self.dictionary

# ======================
# ЗАПУСК ПРИЛОЖЕНИЯ
# ======================
if __name__ == "__main__":
    app = QApplication(sys.argv)
    
    # Установка стиля
    app.setStyle("Fusion")
    
    # Проверка GPU
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    print(f"Используемое устройство: {device}")
    
    # Создание главного окна
    window = MainWindow()
    window.show()
    
    sys.exit(app.exec())