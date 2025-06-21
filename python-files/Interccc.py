import tkinter as tk
from tkinter import ttk, messagebox, simpledialog
import pandas as pd
import os
from datetime import datetime
import csv
import json
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

# Global variables
OB_COUNTER_FILE = "ob_counters.json"
EXCEL_FILE = "testlist.xlsx"
CALL_JOBS_FILE = "call_jobs.csv"
ADMIN_UI = "52651"
ADMIN_PASSWORD = "password"
APPROVAL_COUNTER_FILE = "approval_counter.txt"

# Column names for call job details
CALL_JOB_COLUMNS = [
    "Date", "Time", "Officer", "Contact Phone", "OC Case", "Unit", 
    "Case Reference", "Called Time", "Interpreter Reg. No.", 
    "Called by", "Result", "Team", "Approval No.", "Remarks"
]

class EditableTreeview(ttk.Treeview):
    """Treeview with editable cells and zebra striping"""
    def __init__(self, parent, columns, **kwargs):
        super().__init__(parent, columns=columns, show="headings", **kwargs)
        self.columns = columns
        self.editing_cell = None
        self.entry = None
        
        # Configure column headings
        for col in columns:
            self.heading(col, text=col)
            self.column(col, width=100, anchor="w", minwidth=100, stretch=True)
        
        # Bind double click for editing
        self.bind("<Double-1>", self.on_double_click)
        self.bind("<ButtonRelease-1>", self.on_single_click)
        
        # Configure zebra striping
        self.tag_configure('evenrow', background='#f0f0f0')
        self.tag_configure('oddrow', background='#ffffff')
        self.zebra_stripe()
        
        # Bind events to maintain striping
        self.bind("<<TreeviewSelect>>", self.zebra_stripe)
        self.bind("<Delete>", self.zebra_stripe)
        
    def zebra_stripe(self, event=None):
        """Apply zebra striping to all rows"""
        for i, item in enumerate(self.get_children()):
            if i % 2 == 0:
                self.item(item, tags=('evenrow',))
            else:
                self.item(item, tags=('oddrow',))
    
    def insert(self, parent, index, **kwargs):
        """Override insert to maintain striping"""
        item_id = super().insert(parent, index, **kwargs)
        self.zebra_stripe()
        return item_id
        
    def delete(self, *items):
        """Override delete to maintain striping"""
        super().delete(*items)
        self.zebra_stripe()
        
    def on_double_click(self, event):
        """Handle double-click to edit cell"""
        region = self.identify_region(event.x, event.y)
        if region != "cell":
            return
            
        column = self.identify_column(event.x)
        column_index = int(column[1:]) - 1
        item = self.identify_row(event.y)
        
        if not item:
            return
            
        # Get current value
        values = self.item(item, "values")
        current_value = values[column_index] if column_index < len(values) else ""
        
        # Get cell coordinates
        x, y, width, height = self.bbox(item, column)
        
        # Create entry widget
        self.entry = ttk.Entry(self, font=("Arial", 10))
        self.entry.insert(0, current_value)
        self.entry.select_range(0, tk.END)
        self.entry.focus()
        self.entry.place(x=x, y=y, width=width, height=height)
        
        # Bind events
        self.entry.bind("<Return>", lambda e: self.save_edit(item, column_index))
        self.entry.bind("<Escape>", self.cancel_edit)
        self.entry.bind("<FocusOut>", lambda e: self.save_edit(item, column_index))
        
        # Store editing state
        self.editing_cell = (item, column_index)
        
    def on_single_click(self, event):
        """Handle single click to select row"""
        region = self.identify_region(event.x, event.y)
        if region == "cell":
            item = self.identify_row(event.y)
            self.selection_set(item)
        
    def save_edit(self, item, column_index):
        """Save edited value to treeview"""
        if not self.entry:
            return
            
        new_value = self.entry.get()
        values = list(self.item(item, "values"))
        
        # Ensure we have enough columns
        if len(values) < len(self.columns):
            values.extend([""] * (len(self.columns) - len(values)))
            
        values[column_index] = new_value
        self.item(item, values=values)
        
        self.cancel_edit()
        self.zebra_stripe()  # Maintain striping after edit
        
    def cancel_edit(self, event=None):
        """Cancel editing"""
        if self.entry:
            self.entry.place_forget()
            self.entry = None
            self.editing_cell = None

class InterpreterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Interpreter Management System")
        self.root.geometry("1200x900")
        self.root.configure(bg="#f0f0f0")
        
        # Configure treeview styles
        self.configure_styles()
        
        # Set application icon
        try:
            self.root.iconbitmap("app_icon.ico")
        except:
            pass
        
        self.ob_counters = self.load_ob_counters()
        self.languages = self.load_languages()
        self.call_job_data = []
        self.current_language = None
        self.current_ob = None
        self.original_counter = None  # For OB counter reversion
        self.pending_ob = False  # Track if OB needs reversion
        self.show_login()
        
    def configure_styles(self):
        """Configure styles for treeview widgets"""
        style = ttk.Style()
        style.configure("Treeview", 
                        background="#ffffff",
                        foreground="#000000",
                        rowheight=30,
                        fieldbackground="#ffffff")
        style.configure("Treeview.Heading", 
                        background="#e0e0e0",
                        foreground="#000000",
                        font=('Arial', 10, 'bold'))
        style.map('Treeview', background=[('selected', '#347083')])
        
        # Custom styles for remarks column
        style.configure("OldRemark.Treeview", foreground="red")
        style.configure("NextRemark.Treeview", foreground="yellow")

    def load_languages(self):
        """Load available languages from Excel sheet names"""
        try:
            if os.path.exists(EXCEL_FILE):
                xl = pd.ExcelFile(EXCEL_FILE)
                return xl.sheet_names
            return []
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load Excel file:\n{str(e)}")
            return []

    def load_ob_counters(self):
        """Load OB counters from JSON file"""
        try:
            if os.path.exists(OB_COUNTER_FILE):
                with open(OB_COUNTER_FILE, "r") as f:
                    return json.load(f)
            return {lang: 1 for lang in self.load_languages()}
        except:
            return {}

    def save_ob_counters(self):
        """Save OB counters to JSON file"""
        try:
            with open(OB_COUNTER_FILE, "w") as f:
                json.dump(self.ob_counters, f, indent=2)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save OB counters:\n{str(e)}")
            
    def generate_ob_number(self, language):
        """Generate OB number with language prefix"""
        if language not in self.ob_counters:
            self.ob_counters[language] = 1
            
        prefix = language[:3].upper()
        ob_num = f"{prefix}{self.ob_counters[language]:03d}"
        self.ob_counters[language] += 1
        self.save_ob_counters()
        return ob_num

    def revert_ob_counter(self, language):
        """Revert OB counter if OB wasn't saved"""
        if language in self.ob_counters:
            self.ob_counters[language] -= 1
            self.save_ob_counters()

    def load_call_jobs(self):
        """Load existing call jobs from CSV file"""
        if not os.path.exists(CALL_JOBS_FILE):
            return []
            
        try:
            with open(CALL_JOBS_FILE, "r", newline="") as f:
                reader = csv.DictReader(f)
                jobs = list(reader)
                return sorted(jobs, key=lambda x: (x['OB'], x['Time']))
        except:
            return []

    def save_call_jobs(self):
        """Save call jobs to CSV file"""
        try:
            sorted_jobs = sorted(self.call_job_data, key=lambda x: (x['OB'], x['Time']))
            
            with open(CALL_JOBS_FILE, "w", newline="") as f:
                writer = csv.DictWriter(f, fieldnames=["OB"] + CALL_JOB_COLUMNS)
                writer.writeheader()
                for job in sorted_jobs:
                    writer.writerow(job)
            return True
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save call jobs:\n{str(e)}")
            return False

    def show_login(self):
        self.clear_window()
        
        login_frame = tk.Frame(self.root, bg="#f0f0f0")
        login_frame.place(relx=0.5, rely=0.5, anchor="center")
        
        tk.Label(login_frame, text="Interpreter Management System", 
                font=("Arial", 16, "bold"), bg="#f0f0f0").grid(row=0, column=0, columnspan=2, pady=20)
        
        tk.Label(login_frame, text="UI:", bg="#f0f0f0", font=("Arial", 12)).grid(row=1, column=0, sticky="e", padx=5, pady=5)
        tk.Label(login_frame, text="Password:", bg="#f0f0f0", font=("Arial", 12)).grid(row=2, column=0, sticky="e", padx=5, pady=5)
        
        self.ui_entry = tk.Entry(login_frame, font=("Arial", 12))
        self.password_entry = tk.Entry(login_frame, show="*", font=("Arial", 12))
        
        self.ui_entry.grid(row=1, column=1, sticky="w", padx=5, pady=5)
        self.password_entry.grid(row=2, column=1, sticky="w", padx=5, pady=5)
        
        login_btn = tk.Button(login_frame, text="Login", command=self.authenticate,
                             bg="#4CAF50", fg="white", font=("Arial", 12), padx=20)
        login_btn.grid(row=3, column=0, columnspan=2, pady=20)

    def authenticate(self):
        ui = self.ui_entry.get()
        password = self.password_entry.get()
        
        if ui == ADMIN_UI and password == ADMIN_PASSWORD:
            self.call_job_data = self.load_call_jobs()
            self.show_main_menu()
        else:
            messagebox.showerror("Error", "Invalid UI or Password")

    def show_main_menu(self):
        self.clear_window()
        
        tk.Label(self.root, text="Main Menu", font=("Arial", 16, "bold"), bg="#f0f0f0").pack(pady=20)
        
        button_frame = tk.Frame(self.root, bg="#f0f0f0")
        button_frame.pack(pady=20)
        
        buttons = [
            ("New Call Job (OB)", self.new_call_job),
            ("Edit Existing OB", self.edit_ob),
            ("Generate Approval Memo", self.generate_memo),
            ("Admin/Report", self.admin_report)
        ]
        
        for i, (text, command) in enumerate(buttons):
            btn = tk.Button(button_frame, text=text, command=command,
                           bg="#2196F3", fg="white", font=("Arial", 12), 
                           height=2, width=25, wraplength=180)
            btn.grid(row=i//2, column=i%2, padx=15, pady=15)

    def new_call_job(self):
        self.clear_window()
        self.current_ob = None
        self.current_language = None
        self.pending_ob = False  # Reset pending OB status
        
        # Header frame
        header_frame = tk.Frame(self.root, bg="#f0f0f0")
        header_frame.pack(fill="x", pady=10)
        
        self.ob_var = tk.StringVar()
        self.ob_var.set("OB Number: Not generated")
        tk.Label(header_frame, textvariable=self.ob_var, 
                font=("Arial", 14, "bold"), bg="#f0f0f0").pack(side="left", padx=20)
        
        # Language selection
        lang_frame = tk.Frame(self.root, bg="#f0f0f0")
        lang_frame.pack(fill="x", pady=10)
        
        tk.Label(lang_frame, text="Select Language:", 
                font=("Arial", 12), bg="#f0f0f0").pack(side="left", padx=20)
        
        self.lang_var = tk.StringVar()
        lang_dropdown = ttk.Combobox(lang_frame, textvariable=self.lang_var, 
                                    values=self.languages, state="readonly", width=20)
        lang_dropdown.pack(side="left", padx=5)
        if self.languages:
            lang_dropdown.current(0)
        
        # Buttons
        btn_frame = tk.Frame(self.root, bg="#f0f0f0")
        btn_frame.pack(pady=10)
        
        tk.Button(btn_frame, text="Load Interpreters", 
                 command=self.load_interpreters_handler,
                 bg="#FF9800", fg="white", font=("Arial", 12)).pack(side="left", padx=10)
        
        # Modified to handle OB reversion
        tk.Button(btn_frame, text="Back to Menu", command=self.cleanup_new_call_job,
                 bg="#9E9E9E", fg="white", font=("Arial", 12)).pack(side="left", padx=10)
        
        # Create paned window for the three sections
        paned = tk.PanedWindow(self.root, orient=tk.VERTICAL, bg="#f0f0f0", sashrelief=tk.RAISED)
        paned.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Top frame for interpreter list (increased height)
        self.top_frame = tk.Frame(paned, bg="#f0f0f0", height=240)
        paned.add(self.top_frame)
        
        # Middle frame for previous OBs
        self.middle_frame = tk.Frame(paned, bg="#f0f0f0", height=200)
        paned.add(self.middle_frame)
        
        # Bottom frame for call job details (increased height)
        bottom_frame = tk.Frame(paned, bg="#f0f0f0", height=240)
        paned.add(bottom_frame)
        
        # Call job details section
        tk.Label(bottom_frame, text="Call Job Details", 
                font=("Arial", 12, "bold"), bg="#f0f0f0").pack(pady=5)
        
        # Create scrollable frame for call job table
        table_frame = tk.Frame(bottom_frame, bg="#f0f0f0")
        table_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Create treeview for call job details
        self.call_job_tree = EditableTreeview(table_frame, columns=CALL_JOB_COLUMNS)
        
        # Add scrollbars
        vsb = ttk.Scrollbar(table_frame, orient="vertical", command=self.call_job_tree.yview)
        hsb = ttk.Scrollbar(table_frame, orient="horizontal", command=self.call_job_tree.xview)
        self.call_job_tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        
        # Layout
        self.call_job_tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")
        
        # Configure grid weights
        table_frame.grid_rowconfigure(0, weight=1)
        table_frame.grid_columnconfigure(0, weight=1)
        
        # Add default row
        self.add_new_row()
        
        # Add button frame
        button_frame = tk.Frame(bottom_frame, bg="#f0f0f0")
        button_frame.pack(pady=5)
        
        tk.Button(button_frame, text="Add Row", command=self.add_new_row,
                 bg="#2196F3", fg="white", font=("Arial", 10)).pack(side="left", padx=5)
        
        tk.Button(button_frame, text="Delete Selected", command=self.delete_selected_row,
                 bg="#f44336", fg="white", font=("Arial", 10)).pack(side="left", padx=5)
        
        # Add save button
        save_frame = tk.Frame(bottom_frame, bg="#f0f0f0")
        save_frame.pack(pady=10)
        
        tk.Button(save_frame, text="Save Call Job", command=self.save_current_job,
                 bg="#4CAF50", fg="white", font=("Arial", 12)).pack(padx=10)

    def cleanup_new_call_job(self):
        """Handle cleanup when leaving new call job screen"""
        if self.pending_ob and self.current_language:
            # Revert OB counter if not saved
            self.revert_ob_counter(self.current_language)
            self.pending_ob = False
        self.show_main_menu()

    def load_interpreters_handler(self):
        """Handle interpreter loading with OB generation"""
        language = self.lang_var.get()
        if not language:
            messagebox.showwarning("Warning", "Please select a language")
            return
            
        # Generate OB number if not already generated
        if not self.current_ob:
            # Store current counter state
            if language in self.ob_counters:
                self.original_counter = self.ob_counters[language]
            else:
                self.original_counter = 1
                
            self.current_ob = self.generate_ob_number(language)
            self.current_language = language
            self.ob_var.set(f"OB Number: {self.current_ob}")
            self.pending_ob = True  # Mark as pending save
            
        self.show_interpreters(language)
        self.show_previous_obs(language)

    def show_previous_obs(self, language):
        """Display previous two OBs for the selected language"""
        # Clear middle frame
        for widget in self.middle_frame.winfo_children():
            widget.destroy()
            
        # Get language prefix
        prefix = language[:3].upper()
        
        # Find all OBs for this language (excluding current OB)
        lang_obs = [job for job in self.call_job_data 
                   if job["OB"].startswith(prefix) and job["OB"] != self.current_ob]
        
        if not lang_obs:
            tk.Label(self.middle_frame, text="No previous OBs found for this language",
                    font=("Arial", 12), bg="#f0f0f0").pack(pady=50)
            return
            
        # Get unique OBs and sort descending
        unique_obs = sorted(set(job["OB"] for job in lang_obs), reverse=True)
        
        # Take last two OBs
        last_two_obs = unique_obs[:2]
        
        # Filter records for these OBs
        records_to_show = [job for job in lang_obs if job["OB"] in last_two_obs]
        
        # Sort by OB and Time
        records_to_show = sorted(records_to_show, key=lambda x: (x['OB'], x['Time']))
        
        # Create frame for title
        title_frame = tk.Frame(self.middle_frame, bg="#f0f0f0")
        title_frame.pack(fill="x", pady=(5, 0))
        
        tk.Label(title_frame, text=f"Previous OBs for {language}",
                font=("Arial", 12, "bold"), bg="#f0f0f0").pack()
        
        # Create scrollable frame
        table_frame = tk.Frame(self.middle_frame, bg="#f0f0f0")
        table_frame.pack(fill="both", expand=True, padx=10, pady=5)
        
        # Create treeview
        columns = ["OB"] + CALL_JOB_COLUMNS
        tree = ttk.Treeview(table_frame, columns=columns, show="headings")
        
        # Configure tags for zebra striping
        tree.tag_configure('evenrow', background='#f0f0f0')
        tree.tag_configure('oddrow', background='#ffffff')
        
        # Configure column headings and properties
        for col in columns:
            tree.heading(col, text=col)
            tree.column(col, width=100, anchor="w", minwidth=100, stretch=True)
        
        # Add scrollbars
        vsb = ttk.Scrollbar(table_frame, orient="vertical", command=tree.yview)
        hsb = ttk.Scrollbar(table_frame, orient="horizontal", command=tree.xview)
        tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        
        # Layout
        tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")
        
        # Configure grid weights
        table_frame.grid_rowconfigure(0, weight=1)
        table_frame.grid_columnconfigure(0, weight=1)
        
        # Insert data with alternating row colors and remark styling
        remarks_index = len(columns) - 1  # Last column is Remarks
        for idx, job in enumerate(records_to_show):
            values = [job.get(col, "") for col in columns]
            
            # Determine tags based on remarks
            tags = []
            if idx % 2 == 0:
                tags.append('evenrow')
            else:
                tags.append('oddrow')
                
            # Check remarks for special styling
            remarks = job.get("Remarks", "")
            if remarks.startswith("Old"):
                tags.append('old_remark')
            elif remarks.startswith("Next") or remarks.startswith("1st"):
                tags.append('next_remark')
            
            tree.insert("", "end", values=values, tags=tuple(tags))
        
        # Configure custom remark styles - UPDATED FOR YELLOW TEXT
        tree.tag_configure('old_remark', foreground='red')
        tree.tag_configure('next_remark', foreground='blue', background='yellow')

    def add_new_row(self):
        """Add a new row to the call job table"""
        today = datetime.now().strftime("%Y-%m-%d")
        current_time = datetime.now().strftime("%H:%M")
        self.call_job_tree.insert("", "end", values=(
            today, current_time, "", "", "", "", "", "", "", "", "", "", "", ""
        ))

    def delete_selected_row(self):
        """Delete selected row from call job table"""
        selected = self.call_job_tree.selection()
        if not selected:
            messagebox.showwarning("Warning", "Please select a row to delete")
            return
            
        for item in selected:
            self.call_job_tree.delete(item)

    def save_current_job(self):
        """Save the current call job details"""
        if not self.call_job_tree.get_children():
            messagebox.showwarning("Warning", "No call job data to save")
            return
            
        # Get data from treeview
        job_data = []
        for item in self.call_job_tree.get_children():
            values = self.call_job_tree.item(item, "values")
            
            # Create dictionary with column names and values
            job_dict = {}
            for i, col in enumerate(CALL_JOB_COLUMNS):
                job_dict[col] = values[i] if i < len(values) else ""
                
            # Add OB number
            job_dict["OB"] = self.current_ob
            job_data.append(job_dict)
            
        # Add to call job data
        self.call_job_data.extend(job_data)
        
        # Save to file
        if self.save_call_jobs():
            messagebox.showinfo("Success", f"{len(job_data)} call job records saved successfully")
            
            # Clear current rows and add one empty row
            self.call_job_tree.delete(*self.call_job_tree.get_children())
            self.add_new_row()
            
            # Reset pending OB status
            self.pending_ob = False

    def show_interpreters(self, language):
        try:
            df = pd.read_excel(EXCEL_FILE, sheet_name=language)
            if df.empty:
                messagebox.showinfo("Info", f"No interpreters found for {language}")
                return
                
            # Clear top frame
            for widget in self.top_frame.winfo_children():
                widget.destroy()
                
            # Create interpreter list in top frame
            tk.Label(self.top_frame, text=f"{language} Interpreters", 
                    font=("Arial", 12, "bold"), bg="#f0f0f0").pack(pady=5)
            
            # Create a frame to hold treeview
            tree_frame = tk.Frame(self.top_frame, bg="#f0f0f0")
            tree_frame.pack(fill="both", expand=True, padx=10, pady=10)
            
            # Get first 6 columns
            columns = df.columns.tolist()[:6]
            
            # Create treeview
            tree = ttk.Treeview(tree_frame, columns=columns, show="headings")
            
            # Configure tags for zebra striping
            tree.tag_configure('evenrow', background='#f0f0f0')
            tree.tag_configure('oddrow', background='#ffffff')
            
            # Configure column headings and properties
            for col in columns:
                tree.heading(col, text=col)
                tree.column(col, width=120, anchor="w", minwidth=100, stretch=True)
            
            # Vertical scrollbar
            vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
            tree.configure(yscrollcommand=vsb.set)
            
            # Horizontal scrollbar
            hsb = ttk.Scrollbar(tree_frame, orient="horizontal", command=tree.xview)
            tree.configure(xscrollcommand=hsb.set)
            
            # Layout
            tree.grid(row=0, column=0, sticky="nsew")
            vsb.grid(row=0, column=1, sticky="ns")
            hsb.grid(row=1, column=0, sticky="ew")
            
            # Configure grid weights
            tree_frame.grid_rowconfigure(0, weight=1)
            tree_frame.grid_columnconfigure(0, weight=1)
            
            # Configure text wrapping
            style = ttk.Style()
            style.configure("Treeview", rowheight=30)
            
            # Insert data with alternating row colors
            for idx, row in enumerate(df.iterrows()):
                _, data = row
                values = []
                for val in data.tolist()[:6]:
                    if pd.isna(val) or val == "":
                        values.append("/")
                    else:
                        values.append(str(val))
                
                if idx % 2 == 0:
                    tree.insert("", "end", values=values, tags=('evenrow',))
                else:
                    tree.insert("", "end", values=values, tags=('oddrow',))
            
            # Add "Select" button
            tk.Button(self.top_frame, text="Select Interpreter", 
                     command=lambda: self.select_interpreter(tree),
                     bg="#2196F3", fg="white", font=("Arial", 12)).pack(pady=10)
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load interpreters:\n{str(e)}")

    def select_interpreter(self, tree):
        """Handle interpreter selection"""
        selected = tree.selection()
        if not selected:
            messagebox.showwarning("Warning", "Please select an interpreter")
            return
            
        values = tree.item(selected[0], "values")
        if values:
            # Set interpreter registration number
            if not self.call_job_tree.get_children():
                self.add_new_row()
                
            # Find first empty row
            selected_item = None
            for item in self.call_job_tree.get_children():
                row_values = self.call_job_tree.item(item, "values")
                reg_index = CALL_JOB_COLUMNS.index("Interpreter Reg. No.")
                if reg_index < len(row_values) and row_values[reg_index] == "":
                    selected_item = item
                    break
                    
            # Add new row if none found
            if not selected_item:
                self.add_new_row()
                selected_item = self.call_job_tree.get_children()[-1]
                
            # Update values
            current_values = list(self.call_job_tree.item(selected_item, "values"))
            if len(current_values) < len(CALL_JOB_COLUMNS):
                current_values.extend([""] * (len(CALL_JOB_COLUMNS) - len(current_values)))
            
            reg_no_index = CALL_JOB_COLUMNS.index("Interpreter Reg. No.")
            current_values[reg_no_index] = values[1]  # Second column is reg no (index 1)
            
            self.call_job_tree.item(selected_item, values=current_values)
            
            messagebox.showinfo("Success", f"Selected interpreter: {values[1]} - {values[2]}")
            self.call_job_tree.see(selected_item)
            self.call_job_tree.selection_set(selected_item)

    def edit_ob(self):
        edit_window = tk.Toplevel(self.root)
        edit_window.title("Edit Existing OB")
        edit_window.geometry("800x600")
        
        tk.Label(edit_window, text="Enter OB Number to Edit:", 
                font=("Arial", 12)).pack(pady=20)
        
        ob_entry = tk.Entry(edit_window, font=("Arial", 12), width=15)
        ob_entry.pack(pady=5)
        
        tk.Button(edit_window, text="Load OB", 
                 command=lambda: self.load_ob_details(edit_window, ob_entry.get().upper()),
                 bg="#2196F3", fg="white", font=("Arial", 12)).pack(pady=20)

    def load_ob_details(self, window, ob_number):
        if not ob_number:
            messagebox.showwarning("Warning", "Please enter an OB number")
            return
            
        # Find the OB in our data
        jobs = [job for job in self.call_job_data if job["OB"] == ob_number]
        if not jobs:
            messagebox.showinfo("Info", f"OB {ob_number} not found")
            return
            
        # Clear window and display job details
        for widget in window.winfo_children():
            widget.destroy()
            
        tk.Label(window, text=f"Editing OB: {ob_number}", 
                font=("Arial", 14, "bold")).pack(pady=10)
        
        # Create editable treeview
        tree_frame = tk.Frame(window)
        tree_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        tree = EditableTreeview(tree_frame, columns=CALL_JOB_COLUMNS)
        
        # Add scrollbars
        vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
        hsb = ttk.Scrollbar(tree_frame, orient="horizontal", command=tree.xview)
        tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        
        # Layout
        tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")
        
        # Configure grid weights
        tree_frame.grid_rowconfigure(0, weight=1)
        tree_frame.grid_columnconfigure(0, weight=1)
        
        # Insert job data sorted by time
        jobs_sorted = sorted(jobs, key=lambda job: job.get('Time', ''))
        for job in jobs_sorted:
            values = [job.get(col, "") for col in CALL_JOB_COLUMNS]
            tree.insert("", "end", values=values)
        
        # Button frame
        button_frame = tk.Frame(window)
        button_frame.pack(pady=5)
        
        tk.Button(button_frame, text="Add Row", command=lambda: self.add_row_to_ob(tree),
                 bg="#2196F3", fg="white", font=("Arial", 10)).pack(side="left", padx=5)
        
        tk.Button(button_frame, text="Delete Selected", command=lambda: self.delete_ob_row(tree),
                 bg="#f44336", fg="white", font=("Arial", 10)).pack(side="left", padx=5)
        
        # Save button
        save_frame = tk.Frame(window)
        save_frame.pack(pady=10)
        
        tk.Button(save_frame, text="Save Changes", 
                 command=lambda: self.save_ob_changes(tree, ob_number, window),
                 bg="#4CAF50", fg="white", font=("Arial", 12)).pack(side="left", padx=10)
        
        tk.Button(save_frame, text="Cancel", command=window.destroy,
                 bg="#f44336", fg="white", font=("Arial", 12)).pack(side="left", padx=10)

    def add_row_to_ob(self, tree):
        """Add a new row to the OB edit table"""
        today = datetime.now().strftime("%Y-%m-%d")
        current_time = datetime.now().strftime("%H:%M")
        tree.insert("", "end", values=(
            today, current_time, "", "", "", "", "", "", "", "", "", "", "", ""
        ))

    def delete_ob_row(self, tree):
        """Delete selected row from OB edit table"""
        selected = tree.selection()
        if not selected:
            messagebox.showwarning("Warning", "Please select a row to delete")
            return
            
        for item in selected:
            tree.delete(item)

    def save_ob_changes(self, tree, ob_number, window):
        """Save changes to existing OB"""
        if not tree.get_children():
            messagebox.showwarning("Warning", "No data to save")
            return
            
        # Remove existing records for this OB
        self.call_job_data = [job for job in self.call_job_data if job["OB"] != ob_number]
        
        # Add updated records
        for item in tree.get_children():
            values = tree.item(item, "values")
            
            # Create dictionary with column names and values
            job_dict = {"OB": ob_number}
            for i, col in enumerate(CALL_JOB_COLUMNS):
                job_dict[col] = values[i] if i < len(values) else ""
                
            self.call_job_data.append(job_dict)
            
        # Save to file
        if self.save_call_jobs():
            messagebox.showinfo("Success", "Changes saved successfully")
            window.destroy()

    def generate_memo(self):
        """Generate approval memo for selected OB"""
        # Prompt for OB number
        ob_number = simpledialog.askstring("Generate Approval Memo", 
                                          "Enter OB Number:", 
                                          parent=self.root)
        if not ob_number:
            return
            
        # Filter jobs for this OB with Result="Yes"
        jobs = [job for job in self.call_job_data 
               if job["OB"] == ob_number and job.get("Result", "").strip().lower() == "yes"]
        
        if not jobs:
            messagebox.showinfo("Info", f"No jobs with 'Yes' result found for OB {ob_number}")
            return
            
        # Generate approval numbers and process each job
        for job in jobs:
            # Skip if already has approval number
            if job.get("Approval No."):
                continue
                
            # Generate new approval number
            approval_number = self.get_next_approval_number()
            job["Approval No."] = approval_number
            
            # Create memo document
            self.create_memo_document(job, approval_number)
            
            # Show success message
            messagebox.showinfo("Success", 
                              f"Approval memo generated: {approval_number}.docx\n"
                              f"Interpreter: {job.get('Interpreter Reg. No.', '')}")
        
        # Save updated jobs
        self.save_call_jobs()
        
    def get_next_approval_number(self):
        """Get next approval number from counter file"""
        try:
            if os.path.exists(APPROVAL_COUNTER_FILE):
                with open(APPROVAL_COUNTER_FILE, "r") as f:
                    counter = int(f.read().strip())
            else:
                counter = 1
        except:
            counter = 1
            
        # Save next counter value
        with open(APPROVAL_COUNTER_FILE, "w") as f:
            f.write(str(counter + 1))
            
        return counter

    def get_dc_value(self, case_ref):
        """Determine DC value based on case reference"""
        if not case_ref:
            return ""
            
        case_ref = case_ref.strip().upper()
        
        # Define mapping rules
        if case_ref.startswith("W RN") or case_ref.startswith("ABD") or case_ref.startswith("STY"):
            return "DC WDIST"
        elif case_ref.startswith("C RN"):
            return "DC CDIST"
        elif case_ref.startswith("WCH") or case_ref.startswith("HV"):
            return "DC WCHDIST"
        elif case_ref.startswith("NP") or case_ref.startswith("CW"):
            return "DC EDIST"
        elif case_ref.startswith("HKI"):
            return "SSP CRM HKI"
        elif case_ref.startswith("OCTB"):
            return "CSP OCTB"
        elif case_ref.startswith("NB"):
            return "CSP NB"
        elif case_ref.startswith("T HKI"):
            return "SSP T HKI"
        elif case_ref.startswith("CAPO"):
            return "SP CAPO"
        else:
            return ""

    def get_language_name_from_prefix(self, prefix):
        """Get full language name from OB prefix"""
        for lang in self.languages:
            if lang[:3].upper() == prefix:
                return lang
        return prefix  # Return prefix if no match found

    def get_interpreter_info(self, language, reg_no):
        """Lookup interpreter details from Excel sheet with correct column indices"""
        if not language or not reg_no:
            return ("", "")
            
        try:
            df = pd.read_excel(EXCEL_FILE, sheet_name=language)
            if df.empty:
                return ("", "")
                
            # Find interpreter by registration number
            for _, row in df.iterrows():
                # Column indices:
                # 0: S/N (ignore)
                # 1: Registration Number
                # 2: Name of part-time interpreter
                # 3: Other Languages (ignore)
                # 4: Telephone/Pager No.
                if str(row.iloc[1]) == reg_no:  # Registration Number is at index 1
                    name = str(row.iloc[2]) if pd.notna(row.iloc[2]) else ""
                    tel = str(row.iloc[4]) if pd.notna(row.iloc[4]) else ""
                    return (name, tel)
            return ("", "")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load interpreter info:\n{str(e)}")
            return ("", "")

    def create_memo_document(self, job_data, approval_number):
        """Create memo document from template with enhanced field replacement"""
        try:
            # Load template document
            doc = Document("memo.docx")
            
            # Get language from OB prefix
            ob_prefix = job_data["OB"][:3].upper()
            lang_name = self.get_language_name_from_prefix(ob_prefix)
            
            # Get interpreter details
            interpreter_info = self.get_interpreter_info(
                lang_name, 
                job_data.get("Interpreter Reg. No.", "")
            )
            
            # Determine DC field based on Case Reference
            case_ref = job_data.get("Case Reference", "")
            dc_value = self.get_dc_value(case_ref)
            
            # Create replacement dictionary
            replacements = {
                "<Officer>": job_data.get("Officer", ""),
                "<OC Case>": job_data.get("OC Case", ""),
                "<Unit>": job_data.get("Unit", ""),
                "<Case Reference>": job_data.get("Case Reference", ""),
                "<Interpreter Reg. No.>": job_data.get("Interpreter Reg. No.", ""),
                "<Called>": job_data.get("Called by", ""),
                "<Team>": job_data.get("Team", ""),
                "<app#>": str(approval_number),
                "<dc>": dc_value,
                "<lang>": lang_name.upper(),
                "<intname>": interpreter_info[0],
                "<intel>": interpreter_info[1]
            }
            
            # Replace in paragraphs
            for para in doc.paragraphs:
                for key, value in replacements.items():
                    if key in para.text:
                        para.text = para.text.replace(key, value)
            
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in replacements.items():
                            if key in cell.text:
                                cell.text = cell.text.replace(key, value)
            
            # Save the document
            filename = f"{approval_number}.docx"
            doc.save(filename)
            return True
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate memo:\n{str(e)}")
            return False

    def admin_report(self):
        report_window = tk.Toplevel(self.root)
        report_window.title("Admin Report")
        report_window.geometry("1200x800")
        
        tk.Label(report_window, text="Call Job Report", 
                font=("Arial", 16, "bold")).pack(pady=20)
        
        # Create treeview
        tree_frame = tk.Frame(report_window)
        tree_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        columns = ["OB"] + CALL_JOB_COLUMNS
        tree = ttk.Treeview(tree_frame, columns=columns, show="headings")
        
        # Configure tags for zebra striping
        tree.tag_configure('evenrow', background='#f0f0f0')
        tree.tag_configure('oddrow', background='#ffffff')
        
        # Configure columns
        for col in columns:
            tree.heading(col, text=col)
            tree.column(col, width=100, anchor="w", minwidth=100)
        
        # Add scrollbars
        vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
        hsb = ttk.Scrollbar(tree_frame, orient="horizontal", command=tree.xview)
        tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        
        # Layout
        tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")
        
        # Configure grid weights
        tree_frame.grid_rowconfigure(0, weight=1)
        tree_frame.grid_columnconfigure(0, weight=1)
        
        # Insert data sorted by OB and Time with alternating row colors
        sorted_jobs = sorted(self.call_job_data, key=lambda job: (job['OB'], job.get('Time', '')))
        for idx, job in enumerate(sorted_jobs):
            values = [job.get(col, "") for col in columns]
            if idx % 2 == 0:
                tree.insert("", "end", values=values, tags=('evenrow',))
            else:
                tree.insert("", "end", values=values, tags=('oddrow',))
        
        # Export button
        tk.Button(report_window, text="Export to CSV", 
                 command=self.export_report,
                 bg="#4CAF50", fg="white", font=("Arial", 12)).pack(pady=10)

    def export_report(self):
        """Export report to CSV file"""
        try:
            filename = f"call_job_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
            with open(filename, "w", newline="") as f:
                writer = csv.DictWriter(f, fieldnames=["OB"] + CALL_JOB_COLUMNS)
                writer.writeheader()
                
                # Sort before writing
                sorted_jobs = sorted(self.call_job_data, key=lambda job: (job['OB'], job.get('Time', '')))
                writer.writerows(sorted_jobs)
                
            messagebox.showinfo("Success", f"Report exported to:\n{filename}")
        except Exception as e:
            messagebox.showerror("Error", f"Export failed:\n{str(e)}")

    def clear_window(self):
        for widget in self.root.winfo_children():
            widget.destroy()

if __name__ == "__main__":
    root = tk.Tk()
    app = InterpreterApp(root)
    root.mainloop()