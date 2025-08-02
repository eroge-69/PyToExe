import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import os
import re

def detect_and_replace(run, target, replacement):
    """Replace all case-insensitive occurrences of target with replacement, preserving text style."""
    if not target or not replacement:
        return False

    style = {
        'bold': run.bold,
        'italic': run.italic,
        'underline': run.underline
    }

    # Compile regex for case-insensitive matching, escaping special characters
    pattern = re.compile(re.escape(target), re.IGNORECASE)
    new_text, count = pattern.subn(replacement, run.text)

    if count > 0:
        run.text = new_text
        run.bold = style['bold']
        run.italic = style['italic']
        run.underline = style['underline']
        return True
    return False

def process_documents():
    files = filedialog.askopenfilenames(filetypes=[("Word Documents", "*.docx")])
    if not files:
        messagebox.showerror("Error", "No documents selected.")
        return

    left_from, left_to = l_from.get().strip(), l_to.get().strip()
    right_from, right_to = r_from.get().strip(), r_to.get().strip()
    subj_from, subj_to = s_from.get().strip(), s_to.get().strip()

    save_path = filedialog.askdirectory(title="Select Save Location")
    if not save_path:
        messagebox.showerror("Error", "No save location selected.")
        return

    modified_count = 0
    report = []

    for path in files:
        doc = Document(path)
        left_done = right_done = subject_done = False
        doc_name = os.path.basename(path)
        feedback = [f"\n📝 {doc_name}"]

        # HEADER replacement
        for section in doc.sections:
            header = section.header
            for para in header.paragraphs:
                for run in para.runs:
                    # Left Header
                    if left_from and left_to:
                        if detect_and_replace(run, left_from, left_to):
                            left_done = True
                    # Right Header
                    if right_from and right_to:
                        if detect_and_replace(run, right_from, right_to):
                            right_done = True

        # SUBJECT replacement in document body
        if subj_from and subj_to:
            for para in doc.paragraphs:
                for run in para.runs:
                    if detect_and_replace(run, subj_from, subj_to):
                        subject_done = True

        # Feedback
        feedback.append("✅ Left Header updated." if left_done else "⚠️ Left Header not changed.")
        feedback.append("✅ Right Header updated." if right_done else "⚠️ Right Header not changed.")
        feedback.append("✅ Subject updated in body." if subject_done else "⚠️ Subject not changed.")

        # Save file
        if left_done or right_done or subject_done:
            new_file = os.path.join(save_path, doc_name.replace(".docx", "_modified.docx"))
            doc.save(new_file)
            feedback.append(f"📂 Saved to: {new_file}")
            modified_count += 1
        else:
            feedback.append("🚫 No changes made to this document.")

        report.extend(feedback)

    # Final report
    report.append(f"\n📊 Total documents: {len(files)}")
    report.append(f"✅ Modified documents: {modified_count}")
    messagebox.showinfo("Process Summary", "\n".join(report))

# GUI Setup
root = tk.Tk()
root.title("Word Header & Subject Modifier")

labels = [
    ("Left Header - Change From:", 0), ("Left Header - Change To:", 1),
    ("Right Header - Change From:", 2), ("Right Header - Change To:", 3),
    ("Subject - Change From:", 4), ("Subject - Change To:", 5),
]
entries = {}
for text, row in labels:
    tk.Label(root, text=text).grid(row=row, column=0, sticky="e", padx=5, pady=3)
    entry = tk.Entry(root, width=45)
    entry.grid(row=row, column=1)
    entries[text] = entry

l_from, l_to = entries["Left Header - Change From:"], entries["Left Header - Change To:"]
r_from, r_to = entries["Right Header - Change From:"], entries["Right Header - Change To:"]
s_from, s_to = entries["Subject - Change From:"], entries["Subject - Change To:"]

tk.Button(root, text="Select Word Documents and Apply Changes", command=process_documents).grid(row=6, columnspan=2, pady=10)
root.mainloop()
