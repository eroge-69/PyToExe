import argparse
import json
from docx import Document
from docx.shared import Pt

def parse_ssplist_to_docx(json_file, output_docx):
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    ssp_list = data.get('sspList', [])
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    for ssp in ssp_list:
        doc.add_paragraph(f"Name: {ssp.get('name', '')}")
        doc.add_paragraph(f"Description / purpose: {ssp.get('description', '')}")
        doc.add_paragraph(f"Type: {ssp.get('type', '')}")
        doc.add_paragraph(f"Size: {ssp.get('size', '')}")
        doc.add_paragraph(f"Security strength: {ssp.get('strength', '')}")
        gen = (', '.join(ssp.get('generatedByList', [])) or ', '.join(ssp.get('establishedByList', [])) or 'No')
        doc.add_paragraph(f"Generated / Established: {gen}")
        inp = ', '.join(ssp.get('inputOutputList', []))
        doc.add_paragraph(f"Input / Imported: {inp}")
        out = ', '.join(ssp.get('inputOutputList', []))
        doc.add_paragraph(f"Output / Exported: {out}")
        doc.add_paragraph(f"Pre-loaded by module manufacturer: No")
        doc.add_paragraph(f"Stored persistently across reboots: No")
        doc.add_paragraph(f"Stored temporarily in memory: Yes")
        zeros = ', '.join(ssp.get('zeroizationList', []))
        doc.add_paragraph(f"Zeroisation: {zeros}")
        used = ', '.join(ssp.get('usedByList', []))
        doc.add_paragraph(f"SSP used with functions: {used}")
        rel = ', '.join([f"{relssp.get('sspName','')}: {relssp.get('relationship','')}" for relssp in ssp.get('relatedSspList', [])])
        doc.add_paragraph(f"Related SSPs: {rel}")
        doc.add_paragraph("")  # Empty line for separation
    doc.save(output_docx)

def main():
    parser = argparse.ArgumentParser(description='Parse sspList from JSON and output a formatted .docx file.')
    parser.add_argument('-i', '--input', required=True, help='Input module.json filename')
    parser.add_argument('-o', '--output', required=True, help='Output .docx filename')
    args = parser.parse_args()
    parse_ssplist_to_docx(args.input, args.output)

if __name__ == '__main__':
    main()