import tkinter
from tkinter import messagebox
from tkinter import ttk
from tkinter import filedialog
import openpyxl
from docx import Document
from docx.shared import Cm
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time

fname=''
cl=''
section=''
month=''
wfile=''
wfname=''
al=False

def exit():
    messagebox.showinfo("Close","App is getting closed")
    w.destroy()

def getFile():
    global fname
    fname=filedialog.askopenfilenames(title="Select files to Analyze",filetypes=[("Excel Files","*.xlsx")])
    txtFile.insert(0,str(fname).replace("(","").replace(")","").replace("'",""))
    fname=fname[0]

def getStudentDetails(wbs,wss):
    global month
    stu={}
    slist=[]
    alist=[]
    tlist=[]
    plist=[]
    try:
        txtProgress.insert(1.0,"Getting Student Details\n")
        txtProgress.update()
        time.sleep(0.1)
        mmatch=-1
            
        for col in wss.iter_cols(min_row=4,min_col=2):
            if mmatch==0 or mmatch==1:
                mmatch+=1
            if mmatch>=2 and mmatch<=4:
                L=[]
                for i in col:
                    if i.value==None:
                        continue
                    L.append(i.value)
                L=L[0:stucount+1]
                if mmatch==2:
                    alist=list(L[1:])
                    txtProgress.insert(1.0,"Student Attendance retrieved.\n")
                    txtProgress.update()
                    time.sleep(0.1)
                elif mmatch==3:
                    tlist=list(L[1:])
                    txtProgress.insert(1.0,"Total Attendance retrieved.\n")
                    txtProgress.update()
                    time.sleep(0.1)
                elif mmatch==4:
                    plist=list(L[1:])
                    txtProgress.insert(1.0,"Student Attendance Percentage retrieved.\n")
                    txtProgress.update()
                    time.sleep(0.5)
                    txtProgress.insert(1.0,"Preparing Data of students with less than 75% attendance.\n")
                    txtProgress.update()
                    time.sleep(1)
                    
                    for i in range(len(slist)):
                        if plist[i]>=75:
                            continue
                        stu[slist[i]]=[alist[i],tlist[i],plist[i]]
                    txtProgress.insert(1.0,"Data of students with less than 75% attendance collected successfully.\n")
                    txtProgress.update()
                    time.sleep(1)
                    return stu
                mmatch+=1
            elif mmatch>4:
                break
            
            if col[0].value==None:
                continue
            elif 'MONTH' in col[0].value:
                txtProgress.insert(1.0,"Getting Student Names\n")
                txtProgress.update()
                time.sleep(0.1)
                
                for i in col:
                    if i.value==None:
                        break
                    elif 'MONTH' in i.value or 'NAME' in i.value:
                        continue
                    else:
                        slist.append(i.value)
                stucount=len(slist)
                txtProgress.insert(1.0,"Student Names successfully retrieved.\n")
                txtProgress.update()
                time.sleep(0.1)
            elif col[0].value[:-5]==month:
                txtProgress.insert(1.0,"Month-wise details found.\n")
                txtProgress.update()
                time.sleep(0.1)
                mmatch+=1          
    except:
        messagebox.showinfo("ERROR"," Some Error occurred. Student Details cannot be retrieved and compiled.")
        return None

def getBlankFile():
    global wfile,wfname
    txtProgress.insert(1.0,"Creating Word file.\n")
    txtProgress.update()
    time.sleep(0.5)
            
    messagebox.showinfo("SAVE","Save file for warning letters")
    wfname = filedialog.asksaveasfilename(defaultextension='.docx',filetypes=[("Word Files","*.docx")])
            
    wfile = Document()
    sec = wfile.sections
    for s in sec:
        s.top_margin = Cm(0.5)
        s.bottom_margin = Cm(0.5)
        s.left_margin = Cm(2.0)
        s.right_margin = Cm(1.0)
            

def generate():
    global fname,cl,section,month,wfile,wfname,al
    #Get Values
    if al==False:
        cl=cbClass.get()
        txtProgress.insert(1.0,"Checking Class\n")
        txtProgress.update()
        time.sleep(0.1)
        if cl=='':
            messagebox.showinfo("ERROR","Select Class")
            return
    
        section=cbSection.get()
        txtProgress.insert(1.0,"Checking Section\n")
        txtProgress.update()
        time.sleep(0.1)
        if section=='':
            messagebox.showinfo("ERROR","Select Section")
            return
    
    month=cbMonth.get()
    txtProgress.insert(1.0,"Checking Month\n")
    txtProgress.update()
    time.sleep(0.1)
    if month=='':
        messagebox.showinfo("ERROR","Select Month")
        return
    if fname=='':
        messagebox.showinfo("ERROR","Load File")
        return
    
    #Opening file and worksheet
    try:
        txtProgress.insert(1.0,"Checking File and Sheet it may take some time\n")
        txtProgress.update()
        time.sleep(0.1)
        wb = openpyxl.load_workbook(fname,data_only=True)
        ws = wb[str(cl)+section]
        txtProgress.insert(1.0,"File and worksheet "+str(cl)+section+" found.\n")
        txtProgress.update()
        time.sleep(0.1)
        
    #Calling function to provide student details in dictionary
        try:
            student=getStudentDetails(wb,ws)
            
            if student==None:
                return
            
            if al==False:
                getBlankFile()
            
            txtProgress.insert(1.0,"Adding data to file.\n")
            txtProgress.update()
            time.sleep(0.1)
            
            for i in student:
                thead=wfile.add_table(rows=1,cols=3)
                thead.style = "TableGrid"
                thead.cell(0,0).width=Cm(1)
                thead.cell(0,1).width=Cm(17)
                thead.cell(0,2).width=Cm(1)
                c1=thead.rows[0].cells[0]
                p1 = c1.paragraphs[0]
                run = p1.add_run()
                run.add_picture('KVS_Logo.png',width=Cm(1),height=Cm(1))
                
                p2 = thead.cell(0,1).paragraphs[0]
                run1=p2.add_run("पीएम श्री केन्द्रीय विद्यालय ओईएफ कानपुर\nPM SHRI KENDRIYA VIDYALAYA OEF KANPUR")
                run1.bold=True
                run1.font.size = Pt(12)
                p2.paragraph_format.alignment = 1
                
                p3 = thead.cell(0,2).paragraphs[0]
                run = p3.add_run()
                run.add_picture('PMSHRI_Logo.png',width=Cm(1),height=Cm(1))
                #h=wfile.add_heading("पीएम श्री केन्द्रीय विद्यालय ओईएफ कानपुर\nPM SHRI KENDRIYA VIDYALAYA OEF KANPUR",level=1)
                #h.alignment=1
                
                
                msg="दिनांक: ____________"
                p1=wfile.add_paragraph(msg)
                p1.alignment=2
                msg='''प्रिय अभिभावक,
'छात्र/छात्रा का नाम: {}'
कक्षा एवं विभाग: {}-'{}'
महोदय/महोदया,'''.format(i,str(cl),str(section))
                p1=wfile.add_paragraph(msg)
                p1.alignment=0
                
                msg='''आपको यह सूचित किया जाता है कि आपके पाल्य/पाल्या की माह {} 2025 तक की उपस्थिति केवल {} बैठकों में से केवल {} बैठक अर्थात {}% है | इस विषय में कई बार मौखिक रूप से एवं व्हाट्सएप द्वारा सूचित किया गया है कि छात्रों की विद्यालय से अनुपस्थिति विद्यार्थियों के अकादमीय विकास में बाधक हैं |
केन्द्रीय माध्यमिक शिक्षा बोर्ड एवं केन्द्रीय विद्यालय संगठन के नियमानुसार प्रत्येक छात्र की 75% उपस्थिति होनी अनिवार्य है | ऐसा न होने की स्थिति में छात्र/छात्रा को परीक्षा देने से वंचित किया जा सकता है | अत: आपसे अनुरोध है कि अपने पाल्य की 75% उपस्थिति सुनिश्चित करने का कष्ट करें |
सधन्यवाद'''.format(month,str(student[i][1]),str(student[i][0]),str(student[i][2]))
                p1=wfile.add_paragraph(msg)
                p1.alignment=0
                
                wfile.add_picture('signature.jpg',width=Cm(1.5),height=Cm(0.5))
                
                msg="प्राचार्य"
                p1=wfile.add_paragraph(msg)
                p1.alignment=0
                
                msg="प्राप्ति रसीद:"
                p1=wfile.add_paragraph(msg)
                p1.alignment=1
                
                msg='''पत्र प्राप्त किया |
हस्ताक्षर __________________ दिनांक: ______________ अभिभावक का नाम: ______________________________\n'''
                p1=wfile.add_paragraph(msg)
                p1.alignment=0
                
                thead=wfile.add_table(rows=1,cols=3)
                thead.style = "TableGrid"
                thead.cell(0,0).width=Cm(1)
                thead.cell(0,1).width=Cm(17)
                thead.cell(0,2).width=Cm(1)
                c1=thead.rows[0].cells[0]
                p1 = c1.paragraphs[0]
                run = p1.add_run()
                run.add_picture('KVS_Logo.png',width=Cm(1),height=Cm(1))
                
                p2 = thead.cell(0,1).paragraphs[0]
                run1=p2.add_run("पीएम श्री केन्द्रीय विद्यालय ओईएफ कानपुर\nPM SHRI KENDRIYA VIDYALAYA OEF KANPUR")
                run1.bold=True
                run1.font.size = Pt(12)
                p2.paragraph_format.alignment = 1
                
                p3 = thead.cell(0,2).paragraphs[0]
                run = p3.add_run()
                run.add_picture('PMSHRI_Logo.png',width=Cm(1),height=Cm(1))
                '''h=wfile.add_heading("पीएम श्री केन्द्रीय विद्यालय ओईएफ कानपुर\nPM SHRI KENDRIYA VIDYALAYA OEF KANPUR",level=1)
                h.alignment=1'''
                msg="दिनांक: ____________"
                p1=wfile.add_paragraph(msg)
                p1.alignment=2
                msg='''प्रिय अभिभावक,
'छात्र/छात्रा का नाम: {}'
कक्षा एवं विभाग: {}-'{}'
महोदय/महोदया,'''.format(i,str(cl),str(section))
                p1=wfile.add_paragraph(msg)
                p1.alignment=0
                
                msg='''आपको यह सूचित किया जाता है कि आपके पाल्य/पाल्या की माह {} 2025 तक की उपस्थिति केवल {} बैठकों में से केवल {} बैठक अर्थात {}% है | इस विषय में कई बार मौखिक रूप से एवं व्हाट्सएप द्वारा सूचित किया गया है कि छात्रों की विद्यालय से अनुपस्थिति विद्यार्थियों के अकादमीय विकास में बाधक हैं |
केन्द्रीय माध्यमिक शिक्षा बोर्ड एवं केन्द्रीय विद्यालय संगठन के नियमानुसार प्रत्येक छात्र की 75% उपस्थिति होनी अनिवार्य है | ऐसा न होने की स्थिति में छात्र/छात्रा को परीक्षा देने से वंचित किया जा सकता है | अत: आपसे अनुरोध है कि अपने पाल्य की 75% उपस्थिति सुनिश्चित करने का कष्ट करें |
सधन्यवाद'''.format(month,str(student[i][1]),str(student[i][0]),str(student[i][2]))
                p1=wfile.add_paragraph(msg)
                p1.alignment=0
                
                wfile.add_picture('signature.jpg',width=Cm(1.5),height=Cm(0.5))
                
                msg="प्राचार्य"
                p1=wfile.add_paragraph(msg)
                p1.alignment=0
                
                wfile.add_page_break()
                wfile.save(wfname)
                txtProgress.insert(1.0,"Warning for "+str(i)+" saved succesfully.\n")
                txtProgress.update()
                time.sleep(0.1)
            
            wfile.save(wfname)
            
            txtProgress.insert(1.0,"File successfully saved for class "+str(cl)+"-"+section+"\n")
            txtProgress.update()
            time.sleep(0.1)
            if al==False:
                messagebox.showinfo("COMPLETION","File successfully created.")
        except:
            messagebox.showinfo('ERROR',"Some error occured. File cannot be opened or saved.")
    except:
        messagebox.showinfo("ERROR","File or worksheet "+str(cl)+section+" cannot be found")
        return

def generateAll():
    global fname,cl,section,month,wfile,wfname,al
    al=True
    month=cbMonth.get()
    txtProgress.insert(1.0,"Checking Month\n")
    txtProgress.update()
    time.sleep(0.1)
    if month=='':
        messagebox.showinfo("ERROR","Select Month")
        return
    if fname=='':
        messagebox.showinfo("ERROR","Load File")
        return
    getBlankFile()
    for cl in range(1,13):
        for section in ['A','B','C']:
            if (cl==11 or cl==12) and section=='C':
                continue
            else:
                generate()
    al=False
    messagebox.showinfo("Complete","File Successfully Created")

w=tkinter.Tk()
w.title("ATTENDANCE WARNING GENERATOR")
w.geometry("400x400")
w.configure(background="maroon")

lblMain=tkinter.Label(w,text="Attendance Warning Generator", foreground="white",background="maroon",width="35",font=32)
lblMain.place(x=5,y=5)

lblClass=tkinter.Label(w,text="Choose Class", font=12, foreground="white",background="maroon")
lblClass.place(x=10,y=40)

cbClass=ttk.Combobox(w,values=[1,2,3,4,5,6,7,8,9,10,11,12],font=12,width=10)
cbClass.place(x=250,y=40)

lblSection=tkinter.Label(w,text="Choose Section",font=12, foreground="white",background="maroon")
lblSection.place(x=10,y=80)

cbSection=ttk.Combobox(w,values=['A','B','C'],font=12,width=10)
cbSection.place(x=250,y=80)

lblMonth=tkinter.Label(w,text="Choose Month",font=12, foreground="white",background="maroon")
lblMonth.place(x=10,y=120)

cbMonth=ttk.Combobox(w,values=['APRIL','MAY','JUNE','JULY','AUGUST','SEPTEMBER','OCTOBER','NOVEMBER','DECEMBER','JANUARY'],width=12,font=12)
cbMonth.place(x=250,y=120)

txtFile=tkinter.Entry(w,width=30,font=8)
txtFile.place(x=5,y=150)

btnLoadFile = tkinter.Button(w,text="Load File", font=16,command=getFile)
btnLoadFile.place(x=300,y=150)

txtProgress=tkinter.Text(w,width=48,height=8)
txtProgress.place(x=2,y=250)
txtProgress.insert(1.0,"Select Class, Section, Month; then load file then click Generate.\n")
txtProgress.insert(1.0,"You will see the progress here\n")

btnGenerate = tkinter.Button(w,text="GENERATE", font=16,command=generate)
btnGenerate.place(x=20,y=200)

btnGenerateAll = tkinter.Button(w,text="GENERATE ALL", font=16,command=generateAll)
btnGenerateAll.place(x=145,y=200)

btnClose = tkinter.Button(w,text="CLOSE", font=16,command=exit)
btnClose.place(x=300,y=200)

w.mainloop()