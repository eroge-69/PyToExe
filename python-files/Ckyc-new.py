import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
from docx.shared import Pt
from datetime import datetime
import os

def update_fields(*args):
    entity = entity_type_var.get()

    # Hide all optional fields first
    poa_label.grid_remove()
    poa_holder_entry.grid_remove()
    signatory_label.grid_remove()
    signatory_entry.grid_remove()

    # Then show only as needed
    if entity == "Non-Corporate with POA":
        poa_label.grid(row=2, column=0, sticky="e", padx=8, pady=2)
        poa_holder_entry.grid(row=2, column=1, padx=10)
    elif entity == "Corporate":
        signatory_label.grid(row=2, column=0, sticky="e", padx=8, pady=2)
        signatory_entry.grid(row=2, column=1, padx=10)

def generate_ckyc_letter():
    name = name_var.get().strip()
    poa_holder = poa_holder_var.get().strip()
    signatory = signatory_var.get().strip()
    entity_type = entity_type_var.get()

    if not name:
        messagebox.showerror("Missing Input", "Please enter the name.")
        return

    doc = Document()

    # Set font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Garamond'
    font.size = Pt(11)

    # Header and Address
    doc.add_paragraph("Date: __________________")
    address = (
        "To;\n"
        "Prachay Capital Limited \n"
        "(Formerly known as Prachay Capital Private Limited);\n"
        "Address: Office No. 1401/1402, 14th Floor,\n"
        "Next Gen Avenue, Wing B, CTS No. 2850,\n"
        "Survey No. 103, Bahiratwadi, Near ICC Tower,\n"
        "Senapati Bapat Road, Pune-411016"
    )
    doc.add_paragraph(address)
    doc.add_paragraph("\nSubject: Consent Letter for CKYC (Central Know Your Customer) Compliance as required under the laws\n")

    # Body and Footer
    if entity_type == "Corporate":
        doc.add_paragraph(f"We, {name}, hereby provide our consent for the processing of our KYC (Know Your Customer) details under the Central KYC (CKYC) system.")
        doc.add_paragraph(
            "We understand that CKYC is a centralized repository of KYC records of customers in the financial sector, maintained by "
            "the Central Registry of Securitization Asset Reconstruction and Security Interest of India (CERSAI). By providing our consent, "
            "we authorize Prachay Capital Limited (Prachay Capital Private Limited) to share our KYC information with CERSAI for the purpose of "
            "CKYC compliance and other legal authorities as required under the law."
        )
        doc.add_paragraph(
            "We acknowledge that this consent is given voluntarily and that we have been duly informed of the purpose and implications of CKYC compliance. "
            "We understand that the information provided by us will be used in accordance with applicable laws and regulations governing KYC norms."
        )
        doc.add_paragraph(
            "We hereby declare that the information provided by us is true, accurate, and complete to the best of our knowledge and belief. "
            "We undertake to promptly inform Prachay Capital Limited (Formerly known as Prachay Capital Private Limited) of any changes or updates to our KYC details in the future."
        )
        doc.add_paragraph("\n\nFor " + name + ".\n\n\n")
        doc.add_paragraph(signatory)
        suffix = "_Corporate"

    elif entity_type == "Non-Corporate with POA":
        doc.add_paragraph(f"I, {name}, hereby provide my consent for the processing of my KYC (Know Your Customer) details under the Central KYC (CKYC) system.")
        doc.add_paragraph(
            "I understand that CKYC is a centralized repository of KYC records of customers in the financial sector, maintained by "
            "the Central Registry of Securitization Asset Reconstruction and Security Interest of India (CERSAI). By providing my consent, "
            "I authorize Prachay Capital Limited (formerly known as Prachay Capital Private Limited) to share my KYC information with CERSAI "
            "for the purpose of CKYC compliance and other legal authorities as required under the law."
        )
        doc.add_paragraph(
            "I acknowledge that this consent is given voluntarily and that we have been duly informed of the purpose and implications of CKYC compliance. "
            "I understand that the information provided by me will be used in accordance with applicable laws and regulations governing KYC norms."
        )
        doc.add_paragraph(
            "I hereby declare that the information provided by me is true, accurate, and complete to the best of my knowledge and belief. "
            "I undertake to promptly inform Prachay Capital Limited (Formerly known as Prachay Capital Private Limited) of any changes or updates to my KYC details in the future."
        )
        doc.add_paragraph("\n\n_________________________\n")
        doc.add_paragraph(f"{name}  (Through the hands of his power of attorney holder {poa_holder})")
        suffix = "_POA"

    else:  # Non-Corporate
        doc.add_paragraph(f"I, {name}, hereby provide my consent for the processing of my KYC (Know Your Customer) details under the Central KYC (CKYC) system.")
        doc.add_paragraph(
            "I understand that CKYC is a centralized repository of KYC records of customers in the financial sector, maintained by "
            "the Central Registry of Securitization Asset Reconstruction and Security Interest of India (CERSAI). By providing my consent, "
            "I authorize Prachay Capital Limited (formerly known as Prachay Capital Private Limited) to share my KYC information with CERSAI "
            "for the purpose of CKYC compliance and other legal authorities as required under the law."
        )
        doc.add_paragraph(
            "I acknowledge that this consent is given voluntarily and that we have been duly informed of the purpose and implications of CKYC compliance. "
            "I understand that the information provided by me will be used in accordance with applicable laws and regulations governing KYC norms."
        )
        doc.add_paragraph(
            "I hereby declare that the information provided by me is true, accurate, and complete to the best of my knowledge and belief. "
            "I undertake to promptly inform Prachay Capital Limited (Formerly known as Prachay Capital Private Limited) of any changes or updates to my KYC details in the future."
        )
        doc.add_paragraph("\n\n_________________________\n")
        doc.add_paragraph(name)
        suffix = "_CKYC"

    # Save to Desktop
    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
    filename = os.path.join(desktop_path, f"{name.replace(' ', '_')}{suffix}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")

    try:
        doc.save(filename)
        messagebox.showinfo("Success", f"Letter saved to:\n{filename}")
        os.startfile(filename)
    except PermissionError:
        messagebox.showerror("Permission Denied", f"Close the file if open:\n{filename}")
    except Exception as e:
        messagebox.showerror("Error", str(e))

# GUI setup
root = tk.Tk()
root.title("CKYC Consent Letter Generator")

# Variables
entity_type_var = ttk.Combobox(root, values=["Non-Corporate", "Non-Corporate with POA", "Corporate"], state="readonly")
entity_type_var.set("Non-Corporate")
entity_type_var.bind("<<ComboboxSelected>>", update_fields)

name_var = tk.StringVar()
poa_holder_var = tk.StringVar()
signatory_var = tk.StringVar()

# Widgets
tk.Label(root, text="Entity Type").grid(row=0, column=0, sticky="e", padx=8, pady=5)
entity_type_var.grid(row=0, column=1, padx=10, pady=5)

tk.Label(root, text="Name").grid(row=1, column=0, sticky="e", padx=8)
tk.Entry(root, textvariable=name_var).grid(row=1, column=1, padx=10)

poa_label = tk.Label(root, text="POA Holder Name")
poa_holder_entry = tk.Entry(root, textvariable=poa_holder_var)

signatory_label = tk.Label(root, text="Authorised Signatory")
signatory_entry = tk.Entry(root, textvariable=signatory_var)

tk.Button(root, text="Generate CKYC Letter", command=generate_ckyc_letter, bg="#c6f7c6", width=30).grid(row=5, columnspan=2, pady=15)

update_fields()
root.mainloop()
