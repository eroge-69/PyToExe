#!/usr/bin/env python
# coding: utf-8

# In[8]:


from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime

# Add this at the top or before main()
def get_security_classification():
    options = [
        "SECRET",
        "PERSONAL AND CONFIDENTIAL",
        "RESTRICTED"
    ]
    print("Choose security classification:")
    for idx, val in enumerate(options, 1):
        print(f"{idx}. {val}")
    while True:
        try:
            choice = int(input("Enter number [1-3]: "))
            if 1 <= choice <= len(options):
                return options[choice - 1]
            else:
                print("Invalid choice. Try again.")
        except ValueError:
            print("Please enter a number from 1 to 3.")


def input_lines(prompt):
    print(f"Enter {prompt} (enter blank line to finish):")
    lines = []
    while True:
        line = input()
        if not line.strip():
            break
        lines.append(line)
    return lines

def input_paragraphs(prefix='', indent=0):
    paras = []
    idx = 1
    while True:
        num = f"{prefix}{idx}."
        text = input(f"Enter text for para {num} (enter blank to finish this level): ").strip()
        if not text:
            break
        has_sub = input(f"Does para {num} have sub-paras? (y/n): ").strip().lower()
        para = {'num': num, 'text': text, 'subs': []}
        if has_sub == 'y':
            para['subs'] = input_paragraphs(prefix=num, indent=indent+1)
        paras.append(para)
        idx += 1
    return paras

def set_font_all(doc):
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    for section in doc.sections:
        for container in [section.header, section.footer]:
            for para in container.paragraphs:
                for run in para.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

def add_para(doc, num, text, indent=0):
    """
    Adds a para/sub-para with correct JSSD-style recursive indent:
    - 'indent' is number of spaces before the para number,
    - then para number, then 8 spaces, then text.
    """
    spaces_before_number = ' ' * indent         # align number under parent's body
    spaces_after_number = ' ' * 8              # always 8 spaces after number
    para_line = f"{spaces_before_number}{num}{spaces_after_number}{text}"
    p = doc.add_paragraph(para_line, style='Normal')
    p.paragraph_format.left_indent = Inches(0)
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_after = Pt(12)
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    return p

def add_paragraphs(doc, paras, indent=0, parent_num=''):
    for para in paras:
        add_para(doc, para['num'], para['text'], indent)
        if para.get('subs'):
            # Compute indent for child:
            # Current indent + len(this para's number) + 8 (spaces after number)
            next_indent = indent + len(para['num']) + 8
            add_paragraphs(doc, para['subs'], next_indent, para['num'])


def add_list_section(doc, heading, items):
    if not items:
        return
    
    # Create the first line: heading (bold) + " 1." + 8 spaces + first item text
    first_line = doc.add_paragraph()
    run_heading = first_line.add_run(heading.capitalize() + ":- ")
    run_heading.bold = True
    run_heading.font.name = 'Arial'
    run_heading.font.size = Pt(12)
    run_heading._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # First serial, NOT bold, after heading, on same line
    run_num = first_line.add_run("1." + " " * 8)
    run_num.bold = False
    run_num.font.name = 'Arial'
    run_num.font.size = Pt(12)
    run_num._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run_item = first_line.add_run(items[0])
    run_item.font.name = 'Arial'
    run_item.font.size = Pt(12)
    run_item._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # Subsequent serials: start on new paragraphs, same indent as after heading
    for i, text in enumerate(items[1:], 2):
        # Spaces to align with the start of first serial after heading
        spaces = " " * (len(heading) + 4 + 2)  # len of heading + ":- " (3 or 4) + "1." (2)
        p = doc.add_paragraph()
        run_num = p.add_run(f"{spaces}{i}." + " " * 8)
        run_num.bold = False
        run_num.font.name = 'Arial'
        run_num.font.size = Pt(12)
        run_num._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        run_item = p.add_run(text)
        run_item.font.name = 'Arial'
        run_item.font.size = Pt(12)
        run_item._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')



        
def main():
    print("\n=== JSSD Service Letter Generator ===\n")
    doc = Document()
    sect = doc.sections[0]
    sect.left_margin   = Inches(1.3)
    sect.right_margin  = Inches(0.5)
    sect.top_margin    = Inches(0.5)
    sect.bottom_margin = Inches(0.5)

    # Security classification dropdown
    sec_class = get_security_classification()      # <- ADDED

    # HEADER / FOOTER: Security classification (centered, bold, Arial 12)
    for part in [sect.header.paragraphs[0], sect.footer.paragraphs[0]]:
        run = part.add_run(sec_class)
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        part.alignment = WD_ALIGN_PARAGRAPH.CENTER


    tele = input("Telephone: ")
    fax  = input("Fax: ")
    email = input("E-mail: ")
    doc.add_paragraph(f"Tele    : {tele}")
    doc.add_paragraph(f"Fax     : {fax}")
    doc.add_paragraph(f"E-mail  : {email}")
    doc.add_paragraph()

    org_block = input_lines("originator office address block lines")
    for line in org_block:
        doc.add_paragraph(line)
    doc.add_paragraph()

    file_ref = input("File reference number: ")
    today = datetime.now().strftime("%d-%m-%Y")
    date_str = input(f"Date [{today}]: ").strip() or today
    doc.add_paragraph(f"{file_ref}  {date_str}")
    doc.add_paragraph()

    # Addressee blocks (any number)
    addresses = []
    while True:
        addr = input_lines("Addressee block lines")
        if not addr:
            break
        addresses.append(addr)
        if input("Add another addressee block? (y/n): ").strip().lower() != 'y':
            break
    for addr in addresses:
        for line in addr:
            doc.add_paragraph(line)
        doc.add_paragraph()

    # Subject
    subject = input("Subject (will be in all caps and centered, no 'subject:'): ").strip().upper()
    subj_para = doc.add_paragraph()
    run = subj_para.add_run(subject)
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    subj_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Main and sub-paragraphs
    print("\n== Paragraphs (enter as many as needed, you can add sub-paras to each) ==")
    paras = input_paragraphs()
    add_paragraphs(doc, paras)

    # Spacing before signatory (4 blank)
    for _ in range(4):
        doc.add_paragraph()
    # Signature block, left-aligned
    # Signature block, left-aligned
    sign_name = input("Signatory's Name: ")
    sign_rank = input("Signatory's Rank: ")
    sign_appoint = input("Signatory's Appointment: ")
    p = doc.add_paragraph(f"({sign_name})")
    doc.add_paragraph(sign_rank)
    doc.add_paragraph(sign_appoint)

    doc.add_paragraph()    # <-- single blank line after signatory block

    # Annexures/enclosures/copy to input
    annexures = input_lines("annexures (each line is one; blank to skip)")
    enclosures = input_lines("enclosures (each line is one; blank to skip)")

    def input_copy_to_blocks():
        blocks = []
        print('Enter "Copy to" address blocks (multi-line). At the end of each address, input an empty line. Enter an extra empty line to finish.')
        while True:
            lines = []
            while True:
                line = input()
                if line == "":
                    break
                lines.append(line)
            if not lines:
                break
            blocks.append(lines)
        return blocks

    copy_blocks = input_copy_to_blocks()

    add_list_section(doc, "Annexures", annexures)

    doc.add_paragraph()    # <-- single blank line after annexures

    add_list_section(doc, "Enclosure", enclosures)

    doc.add_paragraph()    # <-- single blank line after enclosures

    # add_copy_to(doc, copy_blocks)


    set_font_all(doc)

    doc.save("JSSD_Service_Letter.docx")
    print("Saved: JSSD_Service_Letter.docx\n")

if __name__ == "__main__":
    main()


# In[9]:


pwd


# In[ ]:




