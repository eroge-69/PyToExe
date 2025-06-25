import os
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from docx import Document
from tkinter import messagebox
import tkinter as tk


# === Load Document ===
DOC_PATH = r"D:\TAFE\TAFE Notes.docx"
if not os.path.exists(DOC_PATH):
    raise FileNotFoundError(f"File not found: {DOC_PATH}")

doc = Document(DOC_PATH)

# === Extract Headings from the DOCX ===
all_headings = []
for para in doc.paragraphs:
    if para.style.name.startswith("Heading") and para.text.strip():
        all_headings.append(para.text.strip())

search_history = []

# === GUI ===
app = ttk.Window(themename="darkly")
app.title("TAFE Notes Search Tool")
app.geometry("1000x650")

ttk.Label(app, text=f"Loaded: {os.path.basename(DOC_PATH)}", font=("Segoe UI", 10, "italic")).pack(pady=5)

main_frame = ttk.Frame(app, padding=10)
main_frame.pack(fill=BOTH, expand=True)

left_frame = ttk.Frame(main_frame)
left_frame.pack(side=LEFT, fill=BOTH, expand=True, padx=(0, 10))

# === Heading Search Area ===
heading_var = ttk.StringVar()
keyword_var = ttk.StringVar()

ttk.Label(left_frame, text="Search by Heading:", font=("Segoe UI", 10, "bold")).pack(anchor="w")
heading_entry = ttk.Entry(left_frame, textvariable=heading_var, width=50)
heading_entry.pack(pady=(0, 5))

# Suggestion box (listbox shown below heading_entry)
suggestion_box = tk.Listbox(left_frame, height=5, font=("Segoe UI", 10))

suggestion_box.pack_forget()

def update_suggestions(event=None):
    query = heading_var.get().strip().lower()
    matches = [h for h in all_headings if h.lower().startswith(query)] if query else []
    
    suggestion_box.delete(0, "end")
    if matches:
        for match in matches:
            suggestion_box.insert("end", match)
        suggestion_box.place(x=heading_entry.winfo_x(), 
                             y=heading_entry.winfo_y() + heading_entry.winfo_height() + 5)
        suggestion_box.lift()
    else:
        suggestion_box.pack_forget()
        suggestion_box.place_forget()

def select_suggestion(event):
    if suggestion_box.curselection():
        selected = suggestion_box.get(suggestion_box.curselection())
        heading_var.set(selected)
        suggestion_box.place_forget()
        search_heading()

heading_entry.bind("<KeyRelease>", update_suggestions)
suggestion_box.bind("<<ListboxSelect>>", select_suggestion)

def update_history(entry, mode):
    label = f"{mode}: {entry}"
    if label not in search_history:
        search_history.append(label)
        history_listbox.insert(END, label)

def run_history_search(event):
    selected = history_listbox.curselection()
    if not selected:
        return
    label = history_listbox.get(selected[0])
    mode, term = label.split(": ", 1)
    if mode == "Heading":
        heading_var.set(term)
        search_heading(from_history=True)
    elif mode == "Keyword":
        keyword_var.set(term)
        search_keyword(from_history=True)

def reset_all():
    heading_var.set("")
    keyword_var.set("")
    result_text.delete("1.0", "end")
    history_listbox.delete(0, "end")
    search_history.clear()
    suggestion_box.place_forget()

def search_heading(from_history=False):
    heading = heading_var.get().strip().lower()
    if not heading:
        messagebox.showwarning("No Heading", "Please enter a heading.")
        return

    if not from_history:
        update_history(heading, "Heading")

    result_text.delete("1.0", "end")
    found = False
    collecting = False

    for para in doc.paragraphs:
        style = para.style.name
        text = para.text.strip()

        if style.startswith("Heading") and text.lower() == heading:
            found = True
            collecting = True
            continue
        elif style.startswith("Heading") and collecting:
            break

        if collecting:
            result_text.insert("end", text + "\n")

    if not found:
        result_text.insert("end", "Heading not found.\n")

def search_keyword(from_history=False):
    keyword = keyword_var.get().strip().lower()
    if not keyword:
        messagebox.showwarning("No Keyword", "Please enter a keyword.")
        return

    if not from_history:
        update_history(keyword, "Keyword")

    result_text.delete("1.0", "end")
    results = []

    for para in doc.paragraphs:
        if keyword in para.text.lower():
            results.append(para.text.strip())

    if results:
        for r in results:
            result_text.insert("end", r + "\n\n")
    else:
        result_text.insert("end", "No results found.\n")

# === Search Buttons ===
ttk.Button(left_frame, text="Search Heading", bootstyle="primary-outline", command=search_heading).pack(pady=5)

ttk.Label(left_frame, text="Search by Keyword:", font=("Segoe UI", 10, "bold")).pack(anchor="w", pady=(10, 0))
ttk.Entry(left_frame, textvariable=keyword_var, width=50).pack(pady=3)
ttk.Button(left_frame, text="Search Keyword", bootstyle="success-outline", command=search_keyword).pack(pady=5)

# === Results Box ===
ttk.Label(left_frame, text="Results:", font=("Segoe UI", 10, "bold")).pack(anchor="w", pady=(10, 0))
result_text = ttk.ScrolledText(left_frame, height=20, wrap="word", font=("Segoe UI", 10))
result_text.pack(fill=BOTH, expand=True, pady=5)

# === Reset Button ===
ttk.Button(left_frame, text="Reset All", bootstyle="danger", command=reset_all).pack(pady=(5, 0))

# === Right Panel: Search History ===
right_frame = ttk.Frame(main_frame)
right_frame.pack(side=RIGHT, fill=Y)

ttk.Label(right_frame, text="Search History:", font=("Segoe UI", 10, "bold")).pack(pady=(0, 5))
history_listbox = tk.Listbox(right_frame, width=30, height=30, font=("Segoe UI", 9))

history_listbox.pack(fill=Y, expand=True)
history_listbox.bind("<<ListboxSelect>>", run_history_search)

# Run the app
app.mainloop()
