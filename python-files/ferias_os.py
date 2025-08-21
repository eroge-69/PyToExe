# -*- coding: utf-8 -*-
"""
Gera DOCX de Férias por Subunidade lendo:
- CONFIG + cabeçalhos/ordem a partir de "FERIAS PARA ORDEM.xlsx" (mesma pasta do exe)
- Plano de Férias (PDF) -> linhas com |data|dias|nome|matricula|
- Efetivo (XLSX) -> folhas: "Nominal", "NOMINAL CIVIS", "SETUP"

Saída: FERIAS_OFICIAIS.docx / FERIAS_SARGENTOS.docx / FERIAS_GUARDAS.docx / FERIAS_CIVIS.docx
"""

import os, sys, re
from collections import defaultdict, namedtuple
from dataclasses import dataclass
from typing import Dict, List, Tuple, Any

import pdfplumber
from openpyxl import load_workbook
from unidecode import unidecode
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

# ---------- ficheiro de configuração ----------
CONFIG_XLSX_NAME = "FERIAS PARA ORDEM.xlsx"

# ---------- util paths ----------
def base_dir() -> str:
    if getattr(sys, "frozen", False):  # PyInstaller
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))

# ---------- tipos ----------
@dataclass
class EfetivoItem:
    categoria_final: str
    nome_titulo: str
    posto_curto: str
    orgao_code: str
    subunidade_norm: str
    antig_cter: Any

@dataclass
class EfetivoData:
    by_mat: Dict[str, "EfetivoItem"]
    orgao_to_ordem: Dict[str, str]

Row = namedtuple("Row", "secao categoria_macro posto_sub inicio dias nome matricula")

# ---------- helpers ----------
MESES = ["JAN","FEV","MAR","ABR","MAI","JUN","JUL","AGO","SET","OUT","NOV","DEZ"]

def norm(s: str) -> str:
    return unidecode((s or "").strip().lower())

def norm_key(s: str) -> str:
    """Chave canónica para órgãos: UPPER + sem acentos + espaços colapsados."""
    t = unidecode((s or "").replace("\xa0", " ")).upper().strip()
    return " ".join(t.split())

def only_digits(s: Any) -> str:
    return "".join(ch for ch in str(s) if ch.isdigit())

def to_title(s: str) -> str:
    return " ".join(w.capitalize() for w in (s or "").strip().lower().split())

def norm_sub(s: str) -> str:
    n = (s or "").strip().upper()
    return "CEM" if n in ("", "PAS") else n

def format_pt_date(ddmmyyyy: str) -> str:
    # aceita "d/m/aaaa" ou "dd/mm/aaaa"
    try:
        dd, mm, yy = ddmmyyyy.strip().split("/")
        dd_i = int(dd); mm_i = int(mm); yy_i = int(yy)
        return f"{dd_i:02d}{MESES[mm_i-1]}{str(yy_i)[-2:]}"
    except Exception:
        return ddmmyyyy  # fallback

def civis_rank(posto: str) -> int:
    s = (posto or "").upper().replace("  ", " ")
    if s == "MFP": return 1
    if s == "MFL": return 2
    if s == "GFL": return 3
    if s == "TECSUP": return 4
    if s in ("A. OP.", "A.OP.", "A OP"): return 5
    # heurística
    ns = unidecode(s)
    if "MESTRE" in s and "PRINC" in s: return 1
    if "MESTRE FLORESTAL" in s: return 2
    if "GUARDA FLORESTAL" in s: return 3
    if "TECNICO" in ns: return 4
    if "ASSISTENTE OPERACIONAL" in s: return 5
    return 99

# ---------- ler CONFIG no XLSX ----------
def read_config_and_headers(xlsx_path: str) -> Tuple[str,str,Dict[str,str],Dict[str,int]]:
    """
    Lê:
      - B2: caminho do EFETIVO_XLSX
      - B3: caminho do PDF
      - Desde linha 6: A=CODIGO, B=POR_EXTENSO, C=ORDEM
    """
    wb = load_workbook(xlsx_path, data_only=True, read_only=True)
    ws = wb["CONFIG"] if "CONFIG" in wb.sheetnames else wb.active

    try:
        efetivo_path = str(ws["B2"].value or "").strip()
        pdf_path     = str(ws["B3"].value or "").strip()
    except Exception:
        wb.close()
        raise RuntimeError("CONFIG inválida: confirme B2 (efetivo) e B3 (PDF).")

    headers_map: Dict[str,str] = {}
    headers_order: Dict[str,int] = {}

    # Desde a linha 6: A=CODIGO, B=POR_EXTENSO, C=ORDEM
    for row in ws.iter_rows(min_row=6, values_only=True):
        code = (row[0] or "").strip()
        if not code:
            continue
        por_ext = (row[1] or "").strip() or code
        ord_val = row[2]
        codeU = code.upper()
        headers_map[codeU] = por_ext
        if ord_val is not None:
            try:
                headers_order[codeU] = int(ord_val)
            except Exception:
                pass

    wb.close()
    return efetivo_path, pdf_path, headers_map, headers_order

# ---------- ler EFETIVO XLSX ----------
def header_index(ws) -> Dict[str,int]:
    header = next(ws.iter_rows(min_row=1, max_row=1, values_only=True))
    idx = {}
    for j, val in enumerate(header):
        name = norm(str(val or ""))
        if name:
            idx[name] = j
    return idx

def safe_ws(wb, name: str):
    if name in wb.sheetnames:
        return wb[name]
    raise RuntimeError(f"Folha em falta: {name}")

def load_efetivo(efetivo_xlsx: str) -> EfetivoData:
    wb = load_workbook(efetivo_xlsx, data_only=True, read_only=True)

    by_mat: Dict[str, EfetivoItem] = {}
    orgao_to_ordem: Dict[str,str] = {}

    # Nominal
    ws_nom = safe_ws(wb, "Nominal")
    h = header_index(ws_nom)
    for r in ws_nom.iter_rows(min_row=2, values_only=True):
        try:
            mat = only_digits(r[h.get("matricula","")])
            if not mat:
                continue
            org_raw = r[h.get("orgao", h.get("órgão", ""))]
            org_code = norm_key(str(org_raw or ""))
            item = EfetivoItem(
                categoria_final=(str(r[h.get("categoria","")]) or "").upper(),
                nome_titulo=to_title(str(r[h.get("nome","")] or "")),
                posto_curto=str(r[h.get("posto","")] or ""),
                orgao_code=org_code,
                subunidade_norm=norm_sub(str(r[h.get("subunidade","")] or "")),
                antig_cter=r[h.get("antiguidade_no_cter","")]
            )
            by_mat[mat] = item
        except Exception:
            continue

    # NOMINAL CIVIS
    if "NOMINAL CIVIS" in wb.sheetnames:
        ws_civ = wb["NOMINAL CIVIS"]
        hc = header_index(ws_civ)
        for r in ws_civ.iter_rows(min_row=2, values_only=True):
            try:
                mat = only_digits(r[hc.get("matricula","")])
                if not mat:
                    continue
                org_raw = r[hc.get("orgao", hc.get("órgão",""))]
                org_code = norm_key(str(org_raw or ""))
                item = EfetivoItem(
                    categoria_final="CIVIS",
                    nome_titulo=to_title(str(r[hc.get("nome","")] or "")),
                    posto_curto=str(r[hc.get("posto","")] or ""),
                    orgao_code=org_code,
                    subunidade_norm=norm_sub(str(r[hc.get("subunidade","")] or "")),
                    antig_cter=None
                )
                by_mat[mat] = item
            except Exception:
                continue

    # SETUP: ORGAO -> ORDEM (detetar a linha de cabeçalho)
    if "SETUP" in wb.sheetnames:
        ws_setup = wb["SETUP"]
        hdr_row_idx = None
        org_idx = ord_idx = None

        for i, row in enumerate(ws_setup.iter_rows(min_row=1, max_row=ws_setup.max_row, values_only=True), start=1):
            vals = [unidecode(str(v or "")).upper().strip() for v in row]
            if not any(vals):
                continue
            if "ORGAO" in vals and "ORDEM" in vals:
                org_idx = vals.index("ORGAO")
                ord_idx = vals.index("ORDEM")
                hdr_row_idx = i
                break

        if hdr_row_idx and org_idx is not None and ord_idx is not None:
            for row in ws_setup.iter_rows(min_row=hdr_row_idx+1, max_row=ws_setup.max_row, values_only=True):
                try:
                    org = norm_key(str((row[org_idx] or "")))
                    ordem = str(row[ord_idx] or "").strip()
                    if org:
                        orgao_to_ordem[org] = ordem
                except Exception:
                    continue

    wb.close()
    return EfetivoData(by_mat=by_mat, orgao_to_ordem=orgao_to_ordem)

# ---------- ler PDF ----------
LINE_RE = re.compile(
    r"\|\s*(\d{1,2}/\d{1,2}/\d{4})\s*\|\s*(\d+)\s*\|\s*([^|]+?)\s*\|\s*(\d{6,7})\s*\|"
)

def read_pdf_rows(pdf_path: str) -> List[Row]:
    txt_lines: List[str] = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            txt_lines.extend([ln.rstrip() for ln in t.splitlines()])

    rows: List[Row] = []
    secao = ""
    cat = ""
    posto_sub = ""

    for ln in txt_lines:
        s = ln.strip()
        if not s:
            continue

        if ("Concessão de férias do ano anterior" in s and "F02" in s):
            secao = "F02"; cat = ""; posto_sub = ""; continue
        if ("Concessão de férias do corrente ano" in s and "F01" in s):
            secao = "F01"; cat = ""; posto_sub = ""; continue

        if s in ("OFICIAIS","SARGENTOS","GUARDAS","CIVIS"):
            cat = s; posto_sub = ""; continue

        # cabeçalho de bloco (subunidade)
        if s == s.upper() and not s.startswith("|") and not s.startswith("+"):
            if all(k not in s for k in ("Início","Matrícula","GBMAPFR","AUSÊNCIAS","Pag.","** FIM")) and len(s) <= 40:
                posto_sub = s
                continue

        m = LINE_RE.search(s)
        if m and secao and cat:
            inicio, dias, nome, matricula = m.groups()
            rows.append(Row(secao=secao, categoria_macro=cat, posto_sub=posto_sub,
                            inicio=inicio, dias=int(dias), nome=nome.strip(), matricula=matricula))
    return rows

# ---------- geração de DOCX ----------
def header_extenso(code: str, headers_map: Dict[str,str]) -> str:
    code = norm_sub(code)
    return headers_map.get(code, code)

def sort_subs(subs: List[str], headers_map: Dict[str,str], headers_order: Dict[str,int]) -> List[str]:
    def key_fn(c):
        cN = norm_sub(c)
        order = headers_order.get(cN, 10**9)
        label = headers_map.get(cN, cN).upper()
        return (order, label)
    return sorted(subs, key=key_fn)

def sort_militares(arr: List[dict]) -> List[dict]:
    def key_fn(x):
        antig = x.get("antig")
        if antig is None or antig == "" or (isinstance(antig, str) and not antig.strip()):
            return (1, float("inf"))
        try:
            return (0, float(antig))
        except Exception:
            return (0, float("inf"))
    return sorted(arr, key=key_fn)

def sort_civis(arr: List[dict]) -> List[dict]:
    def key_fn(x):
        return (civis_rank(x.get("posto_curto","")), x.get("nome_titulo","").upper())
    return sorted(arr, key=key_fn)

def add_paragraph(doc: Document, text: str, make_red: bool=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if make_red:
        run.font.color.rgb = RGBColor(255,0,0)
    run.font.size = Pt(11)

def generate_docx(categoria: str, rows: List[Row], dados: EfetivoData,
                  headers_map: Dict[str,str], headers_order: Dict[str,int],
                  out_dir: str):
    doc = Document()
    h = doc.add_paragraph()
    h_run = h.add_run(f"Concessão de férias ({categoria})")
    h_run.bold = True
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    for secao in ("F02","F01"):
        items = [r for r in rows if r.secao == secao and r.categoria_macro.upper() == categoria.upper()]
        if not items:
            continue

        sec_p = doc.add_paragraph()
        sec_r = sec_p.add_run("Concessão de férias do ano anterior" if secao=="F02" else "Concessão de férias do corrente ano")
        sec_r.bold = True

        by_sub: Dict[str, List[dict]] = defaultdict(list)

        for t in items:
            mat = t.matricula
            if mat in dados.by_mat:
                it = dados.by_mat[mat]
                not_found = False
            else:
                it = EfetivoItem(
                    categoria_final=t.categoria_macro,
                    nome_titulo=t.nome,
                    posto_curto=t.posto_sub,
                    orgao_code="",  # sem órgão
                    subunidade_norm="CEM",
                    antig_cter=None
                )
                not_found = True

            if it.categoria_final.upper() != categoria.upper():
                continue

            sub_key = it.subunidade_norm
            ordem_txt = dados.orgao_to_ordem.get(norm_key(it.orgao_code), it.orgao_code)

            rec = dict(
                matricula=mat,
                inicio=t.inicio,
                dias=t.dias,
                nome_titulo=it.nome_titulo,
                posto_curto=it.posto_curto,
                ordem_txt=ordem_txt,
                antig=it.antig_cter,
                not_found=not_found
            )
            by_sub[sub_key].append(rec)

        if not by_sub:
            continue

        sub_list = sort_subs(list(by_sub.keys()), headers_map, headers_order)

        for sub_code in sub_list:
            doc.add_paragraph("")  # espaço
            hh = doc.add_paragraph()
            rr = hh.add_run(header_extenso(sub_code, headers_map))
            rr.bold = True

            arr = by_sub[sub_code]
            if categoria.upper() != "CIVIS":
                arr = sort_militares(arr)
            else:
                arr = sort_civis(arr)

            for r in arr:
                linha = f'{r["posto_curto"]} N.º {r["matricula"]} - {r["nome_titulo"]}, {r["ordem_txt"]}, desde {format_pt_date(r["inicio"])}, por {r["dias"]} dias;'
                add_paragraph(doc, linha, make_red=bool(r.get("not_found", False)))

    out_path = os.path.join(out_dir, f'FERIAS_{categoria.upper()}.docx')
    doc.save(out_path)
    print(f"[OK] {out_path}")

# ---------- main ----------
def main():
    folder = base_dir()
    cfg_path = os.path.join(folder, CONFIG_XLSX_NAME)
    if not os.path.exists(cfg_path):
        raise SystemExit(f"Ficheiro de configuração não encontrado: {cfg_path}")

    efetivo_path, pdf_path, headers_map, headers_order = read_config_and_headers(cfg_path)

    # caminhos relativos no XLSX → resolver para a pasta do exe
    if efetivo_path and not os.path.isabs(efetivo_path):
        efetivo_path = os.path.join(folder, efetivo_path)
    if pdf_path and not os.path.isabs(pdf_path):
        pdf_path = os.path.join(folder, pdf_path)

    if not os.path.exists(efetivo_path):
        raise SystemExit(f"Efetivo não encontrado: {efetivo_path}")
    if not os.path.exists(pdf_path):
        raise SystemExit(f"PDF não encontrado: {pdf_path}")

    dados = load_efetivo(efetivo_path)
    rows = read_pdf_rows(pdf_path)

    out_dir = folder
    for cat in ("OFICIAIS","SARGENTOS","GUARDAS","CIVIS"):
        generate_docx(cat, rows, dados, headers_map, headers_order, out_dir)

    print("\nConcluído.")

if __name__ == "__main__":
    main()
