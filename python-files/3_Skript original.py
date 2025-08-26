from docx import Document
import re

def lade_platzhalter(pfad):
    mapping = {}
    with open(pfad, encoding="utf-16") as f:
        for zeile in f:
            match = re.match(r"\[\[(.*?)\]\]\s+(.+)", zeile.strip())
            if match:
                key, value = match.groups()
                mapping[f"[[{key}]]"] = value
    return mapping

def ersetze_runs(paragraph, mapping):
    # Alle Runs zu einem Text zusammenfügen
    full_text = ''.join(run.text for run in paragraph.runs)
    pattern = r"\[\[.*?\]\]"

    if not re.search(pattern, full_text):
        return

    for placeholder, value in mapping.items():
        full_text = full_text.replace(placeholder, value)

    paragraph.clear()
    paragraph.add_run(full_text)

def entferne_leere_platzhalter(doc):
    absätze_zum_entfernen = []
    geloeschte_platzhalter = []

    for para in doc.paragraphs:
        full_text = para.text.strip()

        # Nur Platzhalter mit optionalem Bullet am Anfang
        match = re.fullmatch(r"[-•*]?\s*(\[\[.*?\]\])", full_text)
        if match:
            absätze_zum_entfernen.append(para)
            geloeschte_platzhalter.append(match.group(1))  # [[...]] speichern

    # Absätze entfernen
    for para in absätze_zum_entfernen:
        p = para._element
        p.getparent().remove(p)
        p._p = p._element = None

    return geloeschte_platzhalter

def main():
    mapping = lade_platzhalter("2_master.txt")
    doc = Document("4_Datei.docx")

    for para in doc.paragraphs:
        ersetze_runs(para, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    ersetze_runs(para, mapping)

    # Leere Platzhalter entfernen & Namen zurückgeben
    geloeschte_platzhalter = entferne_leere_platzhalter(doc)

    doc.save("5_Datei_modifiziert.docx")

    if geloeschte_platzhalter:
        anzahl = len(geloeschte_platzhalter)
        liste = ', '.join(geloeschte_platzhalter)
        print(f"Dokument aktualisiert – {anzahl} leere Platzhalter entfernt ({liste}).")
    else:
        print("Dokument aktualisiert – keine leeren Platzhalter gefunden.")

if __name__ == "__main__":
    main()
