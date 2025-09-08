import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, filedialog, font
import queue
import json
import os
import pyaudio
from vosk import Model, KaldiRecognizer
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import sys
import platform
import time
from datetime import datetime

class VoiceToDocApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Голос в Документ - Vosk")
        self.root.geometry("1000x800")
        self.root.resizable(True, True)
        
        # Создаем стиль для интерфейса
        self.style = ttk.Style()
        self.style.theme_use("clam")
        self.style.configure("Accent.TButton", foreground="white", background="#4CAF50")
        self.style.map("Accent.TButton", background=[("active", "#45a049")])
        self.style.configure("Stop.TButton", foreground="white", background="#f44336")
        self.style.map("Stop.TButton", background=[("active", "#da190b")])
        
        # Определяем путь к папке model
        if getattr(sys, 'frozen', False):
            self.app_dir = os.path.dirname(sys.executable)
        else:
            self.app_dir = os.path.dirname(os.path.abspath(__file__))
        
        # Устанавливаем рабочую директорию в папку программы
        os.chdir(self.app_dir)
        
        self.model_dir = os.path.join(self.app_dir, "model")
        
        # Отладочная информация
        print(f"📁 Рабочая директория: {os.getcwd()}")
        print(f"📁 Папка программы: {self.app_dir}")
        print(f"📁 Папка модели: {self.model_dir}")
        print(f"📁 Модель существует: {os.path.exists(self.model_dir)}")
        
        # Переменные состояния
        self.recording = False
        self.quotes_open = False
        self.font_size = 12
        self.font_style = "normal"
        self.font_family = "Arial"
        self.alignment = "left"
        self.model = None
        self.recognizer = None
        self.audio = None
        self.stream = None
        self.queue = queue.Queue()
        self.update_id = None
        self.recording_start_time = None
        self.words_count = 0
        self.session_text = ""
        self.microphone_sensitivity = 50  # Чувствительность микрофона (0-100)
        self.selected_microphone = None
        self.available_microphones = []
        self.capitalize_next_word = True  # Делать заглавной первую букву следующего слова
        self.uppercase_mode = False  # Режим полного ввода ЗАГЛАВНЫМИ
        
        # Создание интерфейса
        self.create_widgets()
        
        # Загрузка модели
        self.load_model()

    def load_model(self):
        """Загружает модель распознавания речи из папки model"""
        try:
            if not os.path.exists(self.model_dir):
                self.show_model_error("Папка 'model' не найдена!")
                return False
            
            # Сначала проверяем, является ли сама папка model валидной моделью
            if self.is_valid_model(self.model_dir):
                model_path = self.model_dir
            else:
                # Ищем подпапки с моделями
                model_path = None
                for entry in os.listdir(self.model_dir):
                    full_path = os.path.join(self.model_dir, entry)
                    if os.path.isdir(full_path) and self.is_valid_model(full_path):
                        model_path = full_path
                        break
            
            if model_path:
                print(f"Загружаем модель из: {model_path}")
                self.model = Model(model_path)
                model_name = os.path.basename(model_path)
                self.status_var.set(f"✅ Модель загружена: {model_name}")
                self.record_btn.config(state=tk.NORMAL)
                return True
            
            self.show_model_error("В папке 'model' не найдена подходящая модель Vosk!")
            return False
            
        except Exception as e:
            self.show_model_error(f"Ошибка загрузки модели: {str(e)}")
            return False

    def is_valid_model(self, model_path):
        """Проверяет наличие необходимых файлов модели"""
        required_files = [
            "am/final.mdl",
            "graph/HCLG.fst",
            "ivector/final.dubm"
        ]
        
        print(f"Проверяем модель: {model_path}")
        for file in required_files:
            full_path = os.path.join(model_path, file)
            print(f"  Проверяем файл: {full_path}")
            if not os.path.exists(full_path):
                print(f"    ❌ Файл не найден: {file}")
                return False
            else:
                print(f"    ✅ Файл найден: {file}")
        print(f"✅ Модель валидна: {model_path}")
        return True

    def show_model_error(self, message):
        """Показывает сообщение об ошибке загрузки модели"""
        self.status_var.set(f"❌ {message}")
        error_msg = (
            f"{message}\n\n"
            "Пожалуйста:\n"
            "1. Создайте папку 'model' рядом с программой\n"
            "2. Скачайте модель с сайта: https://alphacephei.com/vosk/models\n"
            "3. Распакуйте скачанный архив в папку 'model'\n\n"
            "Рекомендуемые модели:\n"
            "- vosk-model-small-ru-0.22 (40 МБ) - быстрая\n"
            "- vosk-model-ru-0.42 (67 МБ) - точная\n"
            "- vosk-model-ru-0.42-new (1.8 ГБ) - максимальная точность"
        )
        messagebox.showerror("Ошибка модели", error_msg)

    def create_widgets(self):
        """Создает элементы интерфейса"""
        # Основные фреймы
        main_frame = ttk.Frame(self.root, padding=10)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Панель редактирования текста (сверху)
        self.create_text_formatting_toolbar(main_frame)
        
        # Панель статистики
        stats_frame = ttk.LabelFrame(main_frame, text="Статистика сессии", padding=5)
        stats_frame.pack(fill=tk.X, pady=5)
        
        stats_inner = ttk.Frame(stats_frame)
        stats_inner.pack(fill=tk.X)
        
        self.words_label = ttk.Label(stats_inner, text="Слов: 0")
        self.words_label.pack(side=tk.LEFT, padx=10)
        
        self.time_label = ttk.Label(stats_inner, text="Время: 00:00")
        self.time_label.pack(side=tk.LEFT, padx=10)
        
        self.speed_label = ttk.Label(stats_inner, text="Скорость: 0 слов/мин")
        self.speed_label.pack(side=tk.LEFT, padx=10)
        
        # Текстовое поле с прокруткой
        text_frame = ttk.LabelFrame(main_frame, text="Текст документа", padding=10)
        text_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        self.text_area = scrolledtext.ScrolledText(
            text_frame, wrap=tk.WORD, font=("Arial", 11), padx=10, pady=10
        )
        self.text_area.pack(fill=tk.BOTH, expand=True)
        
        # Панель управления
        control_frame = ttk.Frame(main_frame)
        control_frame.pack(fill=tk.X, pady=10)
        
        # Кнопки управления
        self.record_btn = ttk.Button(
            control_frame, 
            text="▶ Начать запись", 
            command=self.toggle_recording,
            style="Accent.TButton",
            state=tk.DISABLED
        )
        self.record_btn.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        
        ttk.Button(
            control_frame, 
            text="🗑 Очистить текст", 
            command=self.clear_text,
            width=15
        ).pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        
        ttk.Button(
            control_frame, 
            text="📄 Открыть документ", 
            command=self.open_document,
            width=18
        ).pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        
        ttk.Button(
            control_frame, 
            text="💾 Сохранить документ", 
            command=self.save_document,
            width=20
        ).pack(side=tk.RIGHT, padx=5, fill=tk.X, expand=True)
        
        # Панель форматирования
        format_frame = ttk.LabelFrame(main_frame, text="Голосовые команды форматирования", padding=10)
        format_frame.pack(fill=tk.X, pady=5)
        
        # Подсказки по командам
        commands_info = ttk.Label(
            format_frame,
            text=(
                "🎤 Доступные команды:\n"
                "• Знаки препинания: 'точка', 'запятая', 'восклицательный знак', 'вопросительный знак'\n"
                "• Форматирование: 'жирный', 'курсив', 'подчеркивание', 'обычный'\n"
                "• Размер текста: 'мелкий шрифт', 'обычный шрифт', 'крупный шрифт', 'заголовок'\n"
                "• Выравнивание: 'выравнивание по левому краю', 'по центру', 'по правому краю'\n"
                "• Структура: 'новая строка', 'абзац', 'кавычки', 'тире', 'двоеточие'\n"
                "• Специальные: 'точка с запятой', 'многоточие', 'скобки', 'круглые скобки'"
            ),
            wraplength=800,
            justify=tk.LEFT
        )
        commands_info.pack(fill=tk.X)
        
        # Панель управления микрофоном (снизу)
        self.create_microphone_controls(main_frame)
        
        # Статус бар
        self.status_var = tk.StringVar(value="Инициализация приложения...")
        status_bar = ttk.Label(
            self.root, 
            textvariable=self.status_var, 
            relief=tk.SUNKEN, 
            anchor=tk.W, 
            padding=5
        )
        status_bar.pack(side=tk.BOTTOM, fill=tk.X)
        
        # Индикатор записи
        self.recording_indicator = ttk.Label(
            self.root,
            text="🔴",
            font=("Arial", 16, "bold")
        )
        self.recording_indicator.pack(side=tk.BOTTOM, anchor=tk.E, padx=10)
        self.recording_indicator.pack_forget()
        
        # Инициализируем микрофоны после создания всех элементов интерфейса
        self.refresh_microphones()


    def init_audio(self):
        """Инициализирует аудиосистему"""
        try:
            if self.audio is None:
                self.audio = pyaudio.PyAudio()
                
            if self.stream is None:
                # Используем выбранный микрофон или устройство по умолчанию
                input_device_index = self.selected_microphone if self.selected_microphone is not None else None
                
                self.stream = self.audio.open(
                    format=pyaudio.paInt16,
                    channels=1,
                    rate=16000,
                    input=True,
                    input_device_index=input_device_index,
                    frames_per_buffer=8000,
                    stream_callback=self.audio_callback
                )
                self.stream.stop_stream()
                
            if self.model and not self.recognizer:
                self.recognizer = KaldiRecognizer(self.model, 16000)
            
            return True
        except Exception as e:
            self.status_var.set(f"❌ Аудио ошибка: {str(e)}")
            return False

    def audio_callback(self, in_data, frame_count, time_info, status):
        """Обратный вызов для обработки аудиоданных"""
        if self.recording and self.recognizer:
            try:
                if self.recognizer.AcceptWaveform(in_data):
                    result = json.loads(self.recognizer.Result())
                    text = result.get("text", "")
                    if text:
                        self.queue.put(("text", text))
                else:
                    partial = json.loads(self.recognizer.PartialResult())
                    partial_text = partial.get("partial", "")
                    if partial_text:
                        self.queue.put(("partial", partial_text))
            except Exception:
                pass
        return (in_data, pyaudio.paContinue)
    
    def start_recording(self):
        """Начинает запись голоса"""
        if not self.model:
            messagebox.showerror("Ошибка", "Модель распознавания не загружена!")
            return
            
        try:
            # Закрываем существующий поток если он есть
            if self.stream:
                if self.stream.is_active():
                    self.stream.stop_stream()
                self.stream.close()
                self.stream = None
            
            if not self.init_audio():
                return
                
            self.recording = True
            self.recording_start_time = time.time()
            self.record_btn.config(text="⏹ Остановить запись", style="Stop.TButton")
            self.status_var.set("🎤 Запись... Говорите сейчас!")
            self.recording_indicator.pack(side=tk.BOTTOM, anchor=tk.E, padx=10)
            
            # Очищаем очередь
            while not self.queue.empty():
                self.queue.get_nowait()
                
            self.stream.start_stream()
            
            # Запуск обработки очереди и обновления статистики
            self.process_queue()
            self.update_stats()
        except Exception as e:
            self.status_var.set(f"❌ Ошибка записи: {str(e)}")
    
    def stop_recording(self):
        """Останавливает запись голоса"""
        self.recording = False
        self.record_btn.config(text="▶ Начать запись", style="Accent.TButton")
        self.status_var.set("⏹ Запись остановлена")
        self.recording_indicator.pack_forget()
        
        if self.stream and self.stream.is_active():
            self.stream.stop_stream()
        
        if self.update_id:
            self.root.after_cancel(self.update_id)
            self.update_id = None
    
    def toggle_recording(self):
        """Переключает режим записи"""
        if self.recording:
            self.stop_recording()
        else:
            self.start_recording()
    
    def update_stats(self):
        """Обновляет статистику записи"""
        if self.recording and self.recording_start_time:
            elapsed_time = time.time() - self.recording_start_time
            minutes = elapsed_time / 60
            
            # Обновляем время
            time_str = f"{int(elapsed_time // 60):02d}:{int(elapsed_time % 60):02d}"
            self.time_label.config(text=f"Время: {time_str}")
            
            # Обновляем скорость
            if minutes > 0:
                speed = self.words_count / minutes
                self.speed_label.config(text=f"Скорость: {speed:.1f} слов/мин")
            
            self.root.after(1000, self.update_stats)
    
    def process_queue(self):
        """Обрабатывает данные из очереди распознавания"""
        try:
            while not self.queue.empty():
                data_type, data = self.queue.get_nowait()
                
                if data_type == "text":
                    # Обрабатываем команды и добавляем текст только один раз
                    self.process_voice_command(data)
                    self.session_text += data + " "
                    self.words_count += len(data.split())
                    self.words_label.config(text=f"Слов: {self.words_count}")
                    self.text_area.see(tk.END)
                elif data_type == "partial":
                    self.status_var.set(f"🎯 Распознаётся: {data}")
                
                self.root.update_idletasks()
        
        except queue.Empty:
            pass
        
        if self.recording:
            self.update_id = self.root.after(100, self.process_queue)
    
    def process_voice_command(self, text):
        """Обрабатывает голосовые команды для форматирования"""
        original_text = text
        text = text.lower()
        
        # Расширенные знаки препинания
        punctuation_map = {
            "точка": ".", 
            "запятая": ",", 
            "восклицательный знак": "!",
            "вопросительный знак": "?",
            "точка с запятой": ";",
            "многоточие": "...",
            "новая строка": "\n",
            "абзац": "\n\n",
            "тире": " - ",
            "двоеточие": ":",
            "скобки": "[]",
            "круглые скобки": "()",
        }
        
        for cmd, symbol in punctuation_map.items():
            if cmd in text:
                self.text_area.insert(tk.END, symbol)
                # После окончания предложения или начала новой строки/абзаца следующая буква должна быть заглавной
                if symbol in (".", "!", "?", "...", "\n", "\n\n"):
                    self.capitalize_next_word = True
                self.status_var.set(f"✅ Добавлено: {symbol}")
                return
        
        # Кавычки
        if "кавычки" in text:
            self.toggle_quotes()
            return
        
        # Форматирование текста
        format_map = {
            "жирный": ("bold", "🔴 Жирный текст"),
            "курсив": ("italic", "🔴 Курсивный текст"),
            "подчеркивание": ("underline", "🔴 Подчеркнутый текст"),
            "обычный": ("normal", "🔴 Обычный текст"),
        }
        
        for cmd, (style, msg) in format_map.items():
            if cmd in text:
                self.font_style = style
                self.status_var.set(msg)
                return
        
        # Размер шрифта
        size_commands = {
            "мелкий шрифт": 10,
            "обычный шрифт": 12,
            "крупный шрифт": 14,
            "заголовок": 18,
        }
        
        for cmd, size in size_commands.items():
            if cmd in text:
                self.font_size = size
                self.status_var.set(f"🔴 Размер шрифта: {size} pt")
                return
        
        # Выравнивание
        alignment_map = {
            "выравнивание по левому краю": ("left", "🔴 Выравнивание по левому краю"),
            "выравнивание по центру": ("center", "🔴 Выравнивание по центру"),
            "выравнивание по правому краю": ("right", "🔴 Выравнивание по правому краю"),
        }
        
        for cmd, (align, msg) in alignment_map.items():
            if cmd in text:
                self.alignment = align
                self.status_var.set(msg)
                return
        
        # Если не найдена команда, добавляем текст с учётом режима заглавных и автокапитализации
        self.insert_text_with_capitalization(original_text)

    def insert_text_with_capitalization(self, text):
        """Вставляет текст, применяя режим ЗАГЛАВНЫХ и автокапитализацию первого слова."""
        if not text:
            return
        insert_text = text
        if self.uppercase_mode:
            insert_text = insert_text.upper()
        elif self.capitalize_next_word:
            insert_text = self._capitalize_first_letter(insert_text)
            self.capitalize_next_word = False
        self.text_area.insert(tk.END, insert_text + " ")

    def _capitalize_first_letter(self, s: str) -> str:
        """Делает заглавной первую буквенно-цифровую букву в строке, не меняя остальные символы."""
        for idx, ch in enumerate(s):
            if ch.isalpha() or ch.isdigit():
                return s[:idx] + ch.upper() + s[idx+1:]
        return s
    
    def toggle_quotes(self):
        """Переключает открытие/закрытие кавычек"""
        symbol = "«" if not self.quotes_open else "»"
        self.text_area.insert(tk.END, symbol)
        self.quotes_open = not self.quotes_open
        self.status_var.set(f"✅ Вставлены {symbol} кавычки")
    
    def clear_text(self):
        """Очищает текстовое поле"""
        if messagebox.askyesno("Подтверждение", "Очистить весь текст?"):
            self.text_area.delete(1.0, tk.END)
            self.session_text = ""
            self.words_count = 0
            self.words_label.config(text="Слов: 0")
            self.status_var.set("🗑 Текст очищен")
    
    def open_document(self):
        """Открывает существующий документ Word"""
        file_path = filedialog.askopenfilename(
            filetypes=[("Документы Word", "*.docx"), ("Все файлы", "*.*")],
            title="Открыть документ"
        )
        
        if not file_path:
            return
        
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            self.text_area.delete(1.0, tk.END)
            self.text_area.insert(1.0, text)
            self.session_text = text
            self.words_count = len(text.split())
            self.words_label.config(text=f"Слов: {self.words_count}")
            self.status_var.set(f"📄 Документ открыт: {os.path.basename(file_path)}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть документ: {str(e)}")
    
    def save_document(self):
        """Сохраняет документ в формате Word"""
        text = self.text_area.get(1.0, tk.END).strip()
        if not text:
            messagebox.showwarning("Предупреждение", "Нет текста для сохранения!")
            return
        
        # Предлагаем имя файла с датой
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        default_name = f"голосовой_документ_{timestamp}.docx"
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Документы Word", "*.docx"), ("Все файлы", "*.*")],
            title="Сохранить документ",
            initialfile=default_name
        )
        
        if not file_path:
            return
        
        try:
            doc = Document()
            
            # Настройка стилей документа
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(12)
            
            # Добавляем заголовок
            title = doc.add_heading('Голосовой документ', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Добавляем информацию о создании
            info_para = doc.add_paragraph()
            info_para.add_run(f"Создан: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}\n")
            info_para.add_run(f"Количество слов: {self.words_count}\n")
            info_para.add_run(f"Время записи: {self.time_label.cget('text').replace('Время: ', '')}\n")
            info_para.add_run(f"Скорость: {self.speed_label.cget('text').replace('Скорость: ', '')}\n")
            
            # Добавляем разделитель
            doc.add_paragraph("=" * 50)
            
            # Добавляем основной текст
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            run.font.size = Pt(self.font_size)
            
            # Применение стиля шрифта
            if self.font_style == "bold":
                run.bold = True
            elif self.font_style == "italic":
                run.italic = True
            elif self.font_style == "underline":
                run.underline = True
            
            # Применение выравнивания
            if self.alignment == "left":
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif self.alignment == "center":
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif self.alignment == "right":
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            doc.save(file_path)
            self.status_var.set(f"💾 Документ сохранён: {os.path.basename(file_path)}")
            messagebox.showinfo("Успех", f"Документ успешно сохранён!\n{file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить документ: {str(e)}")
    
    def create_text_formatting_toolbar(self, parent):
        """Создает панель редактирования текста"""
        toolbar_frame = ttk.LabelFrame(parent, text="Редактирование текста", padding=5)
        toolbar_frame.pack(fill=tk.X, pady=(0, 5))
        
        # Основная панель инструментов
        toolbar = ttk.Frame(toolbar_frame)
        toolbar.pack(fill=tk.X)
        
        # Кнопки форматирования
        ttk.Button(toolbar, text="B", command=self.toggle_bold, width=3).pack(side=tk.LEFT, padx=2)
        ttk.Button(toolbar, text="I", command=self.toggle_italic, width=3).pack(side=tk.LEFT, padx=2)
        ttk.Button(toolbar, text="U", command=self.toggle_underline, width=3).pack(side=tk.LEFT, padx=2)
        
        # Разделитель
        ttk.Separator(toolbar, orient=tk.VERTICAL).pack(side=tk.LEFT, fill=tk.Y, padx=5)
        
        # Выбор шрифта
        ttk.Label(toolbar, text="Шрифт:").pack(side=tk.LEFT, padx=(5, 2))
        self.font_var = tk.StringVar(value=self.font_family)
        self.font_combo = ttk.Combobox(toolbar, textvariable=self.font_var, width=12, state="readonly")
        self.font_combo.pack(side=tk.LEFT, padx=2)
        self.font_combo.bind("<<ComboboxSelected>>", self.on_font_change)
        
        # Размер шрифта
        ttk.Label(toolbar, text="Размер:").pack(side=tk.LEFT, padx=(10, 2))
        self.size_var = tk.IntVar(value=self.font_size)
        self.size_scale = ttk.Scale(toolbar, from_=8, to=24, variable=self.size_var, 
                                   orient=tk.HORIZONTAL, length=100, command=self.on_size_change)
        self.size_scale.pack(side=tk.LEFT, padx=2)
        
        self.size_label = ttk.Label(toolbar, text=f"{self.font_size}pt")
        self.size_label.pack(side=tk.LEFT, padx=2)
        
        # Разделитель
        ttk.Separator(toolbar, orient=tk.VERTICAL).pack(side=tk.LEFT, fill=tk.Y, padx=5)
        
        # Выравнивание
        ttk.Button(toolbar, text="◀", command=lambda: self.set_alignment("left"), width=3).pack(side=tk.LEFT, padx=2)
        ttk.Button(toolbar, text="▣", command=lambda: self.set_alignment("center"), width=3).pack(side=tk.LEFT, padx=2)
        ttk.Button(toolbar, text="▶", command=lambda: self.set_alignment("right"), width=3).pack(side=tk.LEFT, padx=2)
        
        # Заполнитель
        ttk.Frame(toolbar).pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # Кнопка сброса форматирования
        ttk.Button(toolbar, text="Сброс", command=self.reset_formatting).pack(side=tk.RIGHT, padx=2)
        # Кнопка режима ВЕРХНЕГО РЕГИСТРА
        self.uppercase_btn = ttk.Button(toolbar, text="ABC", command=self.toggle_uppercase_mode, width=4)
        self.uppercase_btn.pack(side=tk.RIGHT, padx=2)
        
        # Заполняем список шрифтов
        self.populate_fonts()

    def create_microphone_controls(self, parent):
        """Создает панель управления микрофоном"""
        mic_frame = ttk.LabelFrame(parent, text="Настройки микрофона", padding=5)
        mic_frame.pack(fill=tk.X, pady=(5, 0))
        
        # Основная панель
        mic_panel = ttk.Frame(mic_frame)
        mic_panel.pack(fill=tk.X)
        
        # Выбор микрофона
        ttk.Label(mic_panel, text="Микрофон:").pack(side=tk.LEFT, padx=(0, 5))
        self.mic_var = tk.StringVar()
        self.mic_combo = ttk.Combobox(mic_panel, textvariable=self.mic_var, width=30, state="readonly")
        self.mic_combo.pack(side=tk.LEFT, padx=(0, 10))
        self.mic_combo.bind("<<ComboboxSelected>>", self.on_microphone_change)
        
        # Кнопка обновления списка микрофонов
        ttk.Button(mic_panel, text="🔄", command=self.refresh_microphones, width=3).pack(side=tk.LEFT, padx=2)
        
        # Чувствительность
        ttk.Label(mic_panel, text="Чувствительность:").pack(side=tk.LEFT, padx=(20, 5))
        self.sensitivity_var = tk.IntVar(value=self.microphone_sensitivity)
        self.sensitivity_scale = ttk.Scale(mic_panel, from_=0, to=100, variable=self.sensitivity_var,
                                          orient=tk.HORIZONTAL, length=150, command=self.on_sensitivity_change)
        self.sensitivity_scale.pack(side=tk.LEFT, padx=2)
        
        self.sensitivity_label = ttk.Label(mic_panel, text=f"{self.microphone_sensitivity}%")
        self.sensitivity_label.pack(side=tk.LEFT, padx=2)
        
        # Заполнитель
        ttk.Frame(mic_panel).pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # Кнопка тестирования микрофона
        self.test_mic_btn = ttk.Button(mic_panel, text="🎤 Тест", command=self.test_microphone)
        self.test_mic_btn.pack(side=tk.RIGHT, padx=2)
        
        # Инициализируем микрофоны после создания status_var
        # self.refresh_microphones()  # Будет вызван после создания status_var

    def populate_fonts(self):
        """Заполняет список доступных шрифтов"""
        try:
            fonts = list(font.families())
            fonts.sort()
            self.font_combo['values'] = fonts
        except Exception:
            # Fallback шрифты
            self.font_combo['values'] = ["Arial", "Times New Roman", "Calibri", "Verdana", "Tahoma"]

    def toggle_bold(self):
        """Переключает жирный шрифт"""
        if self.font_style == "bold":
            self.font_style = "normal"
        else:
            self.font_style = "bold"
        self.update_text_formatting()
        self.status_var.set("🔴 Жирный текст" if self.font_style == "bold" else "🔴 Обычный текст")

    def toggle_italic(self):
        """Переключает курсив"""
        if self.font_style == "italic":
            self.font_style = "normal"
        else:
            self.font_style = "italic"
        self.update_text_formatting()
        self.status_var.set("🔴 Курсивный текст" if self.font_style == "italic" else "🔴 Обычный текст")

    def toggle_underline(self):
        """Переключает подчеркивание"""
        if self.font_style == "underline":
            self.font_style = "normal"
        else:
            self.font_style = "underline"
        self.update_text_formatting()
        self.status_var.set("🔴 Подчеркнутый текст" if self.font_style == "underline" else "🔴 Обычный текст")

    def on_font_change(self, event=None):
        """Обработчик изменения шрифта"""
        self.font_family = self.font_var.get()
        self.update_text_formatting()
        self.status_var.set(f"🔴 Шрифт: {self.font_family}")

    def on_size_change(self, value):
        """Обработчик изменения размера шрифта"""
        self.font_size = int(float(value))
        self.size_label.config(text=f"{self.font_size}pt")
        self.update_text_formatting()
        self.status_var.set(f"🔴 Размер шрифта: {self.font_size} pt")

    def set_alignment(self, alignment):
        """Устанавливает выравнивание текста"""
        self.alignment = alignment
        self.update_text_formatting()
        alignment_names = {"left": "по левому краю", "center": "по центру", "right": "по правому краю"}
        self.status_var.set(f"🔴 Выравнивание {alignment_names[alignment]}")

    def reset_formatting(self):
        """Сбрасывает форматирование к значениям по умолчанию"""
        self.font_style = "normal"
        self.font_size = 12
        self.font_family = "Arial"
        self.alignment = "left"
        
        # Обновляем интерфейс
        self.size_var.set(self.font_size)
        self.size_label.config(text=f"{self.font_size}pt")
        self.font_var.set(self.font_family)
        
        self.update_text_formatting()
        self.status_var.set("🔴 Форматирование сброшено")

    def toggle_uppercase_mode(self):
        """Переключает режим полного ввода ЗАГЛАВНЫМИ."""
        self.uppercase_mode = not self.uppercase_mode
        btn_text = "ABC" if not self.uppercase_mode else "abc"
        self.uppercase_btn.config(text=btn_text)
        self.status_var.set("🔠 Режим ВЕРХНИЙ РЕГИСТР: ВКЛ" if self.uppercase_mode else "🔠 Режим ВЕРХНИЙ РЕГИСТР: ВЫКЛ")

    def update_text_formatting(self):
        """Обновляет форматирование текста в области ввода"""
        # Применяем форматирование к новому тексту
        current_font = (self.font_family, self.font_size)
        if self.font_style == "bold":
            current_font = (self.font_family, self.font_size, "bold")
        elif self.font_style == "italic":
            current_font = (self.font_family, self.font_size, "italic")
        elif self.font_style == "underline":
            current_font = (self.font_family, self.font_size, "underline")
        
        # Обновляем шрифт для новых вставок
        self.text_area.config(font=current_font)

    def refresh_microphones(self):
        """Обновляет список доступных микрофонов"""
        try:
            if self.audio is None:
                self.audio = pyaudio.PyAudio()
            
            self.available_microphones = []
            device_count = self.audio.get_device_count()
            
            for i in range(device_count):
                device_info = self.audio.get_device_info_by_index(i)
                if device_info['maxInputChannels'] > 0:  # Устройство ввода
                    name = device_info['name']
                    self.available_microphones.append((i, name))
            
            # Обновляем список в комбобоксе
            mic_names = [f"{name} (ID: {idx})" for idx, name in self.available_microphones]
            self.mic_combo['values'] = mic_names
            
            if mic_names and not self.mic_var.get():
                self.mic_combo.current(0)
                self.selected_microphone = self.available_microphones[0][0]
            
            self.status_var.set(f"🎤 Найдено микрофонов: {len(self.available_microphones)}")
            
        except Exception as e:
            self.status_var.set(f"❌ Ошибка получения микрофонов: {str(e)}")

    def on_microphone_change(self, event=None):
        """Обработчик изменения выбранного микрофона"""
        selection = self.mic_combo.current()
        if selection >= 0 and selection < len(self.available_microphones):
            self.selected_microphone = self.available_microphones[selection][0]
            mic_name = self.available_microphones[selection][1]
            self.status_var.set(f"🎤 Выбран микрофон: {mic_name}")

    def on_sensitivity_change(self, value):
        """Обработчик изменения чувствительности микрофона"""
        self.microphone_sensitivity = int(float(value))
        self.sensitivity_label.config(text=f"{self.microphone_sensitivity}%")
        self.status_var.set(f"🎤 Чувствительность: {self.microphone_sensitivity}%")

    def test_microphone(self):
        """Тестирует выбранный микрофон"""
        if not self.selected_microphone is None:
            try:
                # Создаем тестовый поток
                test_stream = self.audio.open(
                    format=pyaudio.paInt16,
                    channels=1,
                    rate=16000,
                    input=True,
                    input_device_index=self.selected_microphone,
                    frames_per_buffer=1024
                )
                
                # Читаем небольшой фрагмент
                data = test_stream.read(1024)
                test_stream.stop_stream()
                test_stream.close()
                
                # Проверяем уровень сигнала
                import struct
                audio_data = struct.unpack('1024h', data)
                max_amplitude = max(abs(sample) for sample in audio_data)
                level = min(100, (max_amplitude / 32768) * 100)
                
                self.status_var.set(f"🎤 Тест микрофона: уровень {level:.1f}%")
                messagebox.showinfo("Тест микрофона", f"Микрофон работает!\nУровень сигнала: {level:.1f}%")
                
            except Exception as e:
                self.status_var.set(f"❌ Ошибка тестирования: {str(e)}")
                messagebox.showerror("Ошибка", f"Не удалось протестировать микрофон:\n{str(e)}")
        else:
            messagebox.showwarning("Предупреждение", "Выберите микрофон для тестирования!")
    
    def on_closing(self):
        """Обрабатывает закрытие приложения"""
        self.recording = False
        
        # Останавливаем и закрываем аудиопоток
        if hasattr(self, 'stream') and self.stream:
            if self.stream.is_active():
                self.stream.stop_stream()
            self.stream.close()
        
        # Завершаем аудиосистему
        if hasattr(self, 'audio') and self.audio:
            self.audio.terminate()
        
        self.root.destroy()

if __name__ == "__main__":
    root = tk.Tk()
    
    # Проверяем платформу для корректного отображения иконок
    if platform.system() == "Windows":
        try:
            root.iconbitmap(default="icon.ico")
        except:
            pass
    
    app = VoiceToDocApp(root)
    root.protocol("WM_DELETE_WINDOW", app.on_closing)
    root.mainloop()



