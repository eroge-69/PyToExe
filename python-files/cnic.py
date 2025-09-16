import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image
from docx import Document
from docx.shared import Inches
import os

# --- Constants ---
IMAGES_PER_PAGE = 8  # 4 rows x 2 columns (like in your image)
IMAGE_WIDTH_INCHES = 2.0
IMAGE_HEIGHT_INCHES = 1.25

# --- Function to create DOCX ---
def create_word_doc(front_path, back_path):
    try:
        document = Document()
        
        # Add front page
        document.add_paragraph("CNIC Front Side")
        table = document.add_table(rows=4, cols=2)
        for row in table.rows:
            for cell in row.cells:
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(front_path, width=Inches(IMAGE_WIDTH_INCHES), height=Inches(IMAGE_HEIGHT_INCHES))

        document.add_page_break()

        # Add back page
        document.add_paragraph("CNIC Back Side")
        table = document.add_table(rows=4, cols=2)
        for row in table.rows:
            for cell in row.cells:
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(back_path, width=Inches(IMAGE_WIDTH_INCHES), height=Inches(IMAGE_HEIGHT_INCHES))

        output_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if output_path:
            document.save(output_path)
            messagebox.showinfo("Success", f"Word file saved:\n{output_path}")
        else:
            messagebox.showwarning("Canceled", "Save canceled.")

    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {e}")

# --- GUI Setup ---
def select_front():
    path = filedialog.askopenfilename(filetypes=[("Image files", "*.png;*.jpg;*.jpeg")])
    if path:
        front_entry.delete(0, tk.END)
        front_entry.insert(0, path)

def select_back():
    path = filedialog.askopenfilename(filetypes=[("Image files", "*.png;*.jpg;*.jpeg")])
    if path:
        back_entry.delete(0, tk.END)
        back_entry.insert(0, path)

def generate():
    front_path = front_entry.get()
    back_path = back_entry.get()
    if not os.path.exists(front_path) or not os.path.exists(back_path):
        messagebox.showerror("Error", "Please select valid front and back images.")
        return
    create_word_doc(front_path, back_path)

# --- Build GUI ---
root = tk.Tk()
root.title("CNIC Batch Print Generator")

tk.Label(root, text="Front Image:").grid(row=0, column=0, padx=10, pady=10, sticky='e')
front_entry = tk.Entry(root, width=40)
front_entry.grid(row=0, column=1, padx=10)
tk.Button(root, text="Browse", command=select_front).grid(row=0, column=2)

tk.Label(root, text="Back Image:").grid(row=1, column=0, padx=10, pady=10, sticky='e')
back_entry = tk.Entry(root, width=40)
back_entry.grid(row=1, column=1, padx=10)
tk.Button(root, text="Browse", command=select_back).grid(row=1, column=2)

tk.Button(root, text="Generate Word File", command=generate, bg="green", fg="white").grid(row=2, column=1, pady=20)

root.mainloop()
