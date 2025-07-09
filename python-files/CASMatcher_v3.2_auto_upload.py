import os, re, pymupdf, shutil, asyncio, time, warnings
import numpy as np
from docx import Document
import pandas as pd
from pathlib import Path
import ttkbootstrap as tb
from datetime import datetime
from openpyxl.styles import PatternFill
from tkinter import filedialog, Listbox
from playwright.async_api import async_playwright, TimeoutError as PlaywrightTimeoutError
from threading import Thread
warnings.simplefilter(action='ignore', category=FutureWarning)

# ------------------------------------------------------------------
# Folder that permanently holds the master copies of your standards
# ------------------------------------------------------------------
STANDARD_SOURCE_DIR = "./Standard_storage"

# create essential folders
def check_essential_folders():
    folders = [
     "./1_Standard",                             # 存放標準文件(如化學物質限制清單)
     "./2_IMDS_Lists_to_be_Downloaded",          # 存放要下載的 IMDS 報告 ID 清單                        
      "./3_Downloaded_IMDS_Report",              # 存放下載完成的 IMDS 報告文件
      "./4_Output_Result",                       # 存放處理結果的 Excel 文件
      "./5_Summary",                             # 存放摘要報告文件
      "./6_DownloadFail",                        # 存放下載失敗的報告 ID 清單
      "./Match_Temp_Files"                       # 存放臨時文件
      ]
    for folder in folders:
        if not os.path.exists(folder):
            os.makedirs(folder)
    return None

# get the default edge path for playwright to browse
# 尋找預設的 Edge 瀏覽器路徑,若找不到則返回 None
def find_edge_path():
    paths = [
        "C:\\Program Files (x86)\\Microsoft\\Edge\\Application\\msedge.exe",
        "C:\\Program Files\\Microsoft\\Edge\\Application\\msedge.exe",
        "/Applications/Microsoft Edge.app/Contents/MacOS/Microsoft Edge",
        "/usr/bin/microsoft-edge",
        "/usr/local/bin/microsoft-edge"
    ]
    for path in paths:
        if Path(path).exists():
            return path
    return None

# fit the excel column widths to the contents
# 自動調整 Excel 檔案每一欄的寬度
# writer: ExcelWriter 的物件, 把 pandas 的 DataFrame 輸出(寫入)到 Excel 檔案的工具
# 把多個 DataFrame 寫到同一個 Excel 檔案的不同工作表(sheet)
# Excel
def fit_excel(writer):
    # for every sheet in a excel file
    for sheet_name in writer.sheets:
        worksheet = writer.sheets[sheet_name]
        # col -> 直行
        for col in worksheet.columns:
            max_length = 0
            # Get the column name of this column
            col_name = col[0].column_letter
            for cell in col:
                # 如果這個儲存格的內容長度比目前記錄的最大長度還長,就更新 max_length
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            worksheet.column_dimensions[col_name].width = max_length + 2

# highlight the HomoPPM where the substance exceed the threshold limit
def highlight_excel(writer):
    # 設定一個黃色的儲存格填滿樣式
    #start_color="FFFF00" 代表黃色(Excel 的顏色代碼)。
    #fill_type="solid" 代表是實心填滿。
    highlight_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
    # 取得 Excel 檔案中名為 "Summary" 的工作表。
    worksheet = writer.sheets['Summary']
    # min_row=2:從第2列開始(通常第1列是標題)
    # .min_col=6 : 第6欄開始。
    # 從第二行開始,到最後一行,只看第六到最後一欄
    # 如果這一欄的值是 1,就把它塗成黃色
    """
        假設你的 Excel "Summary" 工作表長這樣:
        A	B	C	D	E	F (第6欄)	G	H	...
        標題1	標題2	...	...	...	homoPPM	...	exceed	...
        ...	...	...	...	...	1234	...	1	...
        ...	...	...	...	...	5678	...	0	...
        ...	...	...	...	...	9999	...	1	...

        從第2列、第6欄開始, 逐列檢查每一列的第6欄到最後一欄的所有儲存格
    """
    # iter_rows 是 openpyxl 套件中 Worksheet 物件的方法,用來逐列(row by row)遍歷 Excel 工作表的儲存格。
    # 每次迴圈,row 會是一個 tuple,裡面是這一列你指定範圍的所有 cell 物件。
    # 如果這一列的最後一欄(exceed)是 1,就把這一列的第6欄(homoPPM)塗成黃色。
    for row in worksheet.iter_rows(min_row=2, max_row=worksheet.max_row, min_col=6, max_col=worksheet.max_column):
        if row[-1].value == 1:
            row[0].fill = highlight_fill
    # 把目前工作表的最後一欄整欄刪除
    worksheet.delete_cols(worksheet.max_column)

# a class for autodownloading of the IMDS reports
class AutoDownload:
    # 建立 AutoDownload 物件時會自動執行
    def __init__(self, headless, edge_path):
        # 布林值,決定瀏覽器要不要「無頭模式」(不顯示視窗,背景執行)
        # True = 無頭模式(背景執行,不顯示瀏覽器視窗)
        # False = 顯示瀏覽器視窗(方便調試和觀察操作過程)
        self.headless = headless
        
        # Edge 瀏覽器的執行檔路徑
        # 例如: "C:\\Program Files (x86)\\Microsoft\\Edge\\Application\\msedge.exe"
        # 如果為 None,則使用系統預設的 Chromium
        self.edge_path = edge_path
        
        # 用來啟動 Playwright 的物件
        # Playwright 是一個瀏覽器自動化框架
        # 初始值為 None,在 start_browser() 方法中會被初始化
        self.playwright = None
        
        # 用來啟動瀏覽器的物件
        # 代表一個瀏覽器實例(如 Chrome, Edge 等)
        # 初始值為 None,在 start_browser() 方法中會被初始化
        self.browser = None
        
        # 用來建立瀏覽器上下文(Context)的物件
        # 瀏覽器上下文類似於隱身模式,每個上下文都是獨立的
        # 可以同時運行多個上下文而不互相影響
        # 初始值為 None,在 start_browser() 方法中會被初始化
        self.context = None
        
        # 用來操作瀏覽器頁面的物件
        # 代表瀏覽器中的一個標籤頁(Tab)
        # 所有的頁面操作(點擊、輸入、導航等)都通過這個物件進行
        # 初始值為 None,在 start_browser() 方法中會被初始化
        self.page = None
        
        # 布林值,決定要使用哪一種搜尋方式(Motor 或 Inbox)
        # True = 使用 MotorIMDS (汽車產業專用介面)
        # False = 使用一般 IMDS (標準收發件管理介面)
        self.isMotor = True
        
        """
        一般 IMDS(Inbox/Outbox):標準的收發件管理介面,所有用戶都會用到。
        MotorIMDS:汽車產業專用的進階管理介面,功能更強大,操作流程不同。
        """
        
        # 判斷目前排序是否為遞減(用於下載時排序控制)
        # True = 已按遞減順序排序
        # False = 未排序或按遞增順序排序
        # 這個變數用於避免重複進行排序操作,提高效率
        self.descending = False

    def _check_page_initialized(self, message=None):
        """
        檢查頁面是否已初始化
        
        參數:
        - message: UI 元件，用來顯示錯誤訊息。如果為 None, 則不顯示錯誤訊息
        
        返回值:
        - True: 頁面已初始化 (self.page 不是 None)
        - False: 頁面未初始化 (self.page 是 None)
        
        邏輯流程:
        1. 檢查 self.page 是否為 None
        2. 如果為 None (未初始化):
           - 如果有傳入 message 參數，在 UI 上顯示錯誤訊息
           - 返回 False
        3. 如果不為 None (已初始化):
           - 直接返回 True
        """
        # 步驟1: 檢查頁面是否為 None (未初始化)
        if self.page is None:
            # 步驟2: 檢查是否有傳入 message 參數
            if message:
                # 步驟3: 如果有 message，在 UI 上顯示錯誤訊息
                message.config(text="Browser not initialized!")
            # 步驟4: 無論是否有 message，都返回 False (表示未初始化)
            return False
        
        # 步驟5: 如果頁面不是 None，返回 True (表示已初始化)
        return True

    # start the edge browser and navigate to the login page
    """
    async (異步)
    """
    async def start_browser(self, message):
        # 使用 try-except 來捕獲瀏覽器初始化過程中可能發生的錯誤
        try:
            # 啟動 Playwright 引擎
            # Playwright 是一個跨瀏覽器的自動化測試框架
            # 支援 Chrome, Firefox, Safari 等多種瀏覽器
            '''
            # await 告訴程序："等待這個操作完成後，再繼續下一行"
            '''
            self.playwright = await async_playwright().start()
            
            # 啟動 Chromium 瀏覽器
            # executable_path: 指定 Edge 瀏覽器的執行檔路徑
            # headless: 決定是否顯示瀏覽器視窗
            #   - True: 無頭模式,不顯示視窗(適合伺服器環境)
            #   - False: 顯示視窗(適合調試和觀察)
            self.browser = await self.playwright.chromium.launch(executable_path=self.edge_path, headless=self.headless)
            
            # 創建新的瀏覽器上下文
            # 瀏覽器上下文類似於隱身模式,每個上下文都是獨立的
            # 可以同時運行多個上下文而不互相影響
            # 這對於並行處理多個任務很有用
            self.context = await self.browser.new_context()
            
            # 在瀏覽器上下文中創建新頁面
            # 頁面代表瀏覽器中的一個標籤頁(Tab)
            # 所有的頁面操作(點擊、輸入、導航等)都通過這個物件進行
            self.page = await self.context.new_page()
            
            # 導航到 IMDS 系統的登入頁面
            # 這是 IMDS (International Material Data System) 的官方登入頁面
            # 用戶需要在這個頁面手動登入
            await self.page.goto('https://www.mdsystem.com/imdsnt/faces/login')
            
            # 在 UI 上顯示成功消息
            # 告訴用戶瀏覽器已經成功初始化並導航到登入頁面
            message.config(text="Browser and page initialized successfully.")
            
        # 捕獲並處理所有可能發生的異常
        except Exception as e:
            # 在 UI 上顯示錯誤消息
            # 包含具體的錯誤信息,幫助用戶診斷問題
            message.config(text=f"Error initializing browser: {e}")

    # navigate to inbox 收件匣 search
    # 查看其他公司發給您的報告
    async def go_inbox(self, message):
        '''
        # 情況2:如果 self.page 不是 None
        self.page = some_page_object
        result = self._check_page_initialized(message)  # result = True
        if not result:  # if not True = if False
            return      # 不執行 return, 繼續執行後面的程式碼
        '''
        if not self._check_page_initialized(message):
            return
        self.isMotor = False  # 設定為一般 IMDS 模式（非 Motor 模式）

        # 模擬按下 Alt+Shift+I 快捷鍵導航到 Inbox 頁面
        await self.page.keyboard.down('Alt')
        await self.page.keyboard.down('Shift')
        await self.page.keyboard.press('I')
        await self.page.keyboard.up('Shift')
        await self.page.keyboard.up('Alt')
        
        # 等待 Inbox 頁面的搜尋欄位出現（最多等待10秒）
        await self.page.wait_for_selector('#pt1\\:dcCmds\\:sfIbLU\\:itModuleId\\:\\:content', timeout=10000)
        await asyncio.sleep(1)  # 等待頁面穩定
        
        # 在下拉選單中選擇版本 "0"
        await self.page.select_option('#pt1\\:dcCmds\\:sfIbLU\\:socVersion\\:\\:content', value="0")
        await asyncio.sleep(0.75)  # 等待選擇生效
        
        message.config(text="Ready to Download!")  # 通知用戶準備就緒
        return
 
    # navigate to outbox 發件匣 search
    # 查看	您發給其他公司的報告
    async def go_outbox(self, message):
        # check if the page is initialized
        if not self._check_page_initialized(message):
            return
        self.isMotor = True  # 設定為 MotorIMDS 模式（汽車產業專用介面）
        self.descending = False  # 重置排序狀態
        
        # 模擬按下 Alt+Shift+M 快捷鍵導航到 MotorIMDS Outbox 頁面
        await self.page.keyboard.down('Alt')
        await self.page.keyboard.down('Shift')
        await self.page.keyboard.press('M')
        await self.page.keyboard.up('Shift')
        await self.page.keyboard.up('Alt')
        
        # 等待 MotorIMDS Outbox 頁面的搜尋欄位出現（最多等待10秒）
        await self.page.wait_for_selector('#pt1\\:dcSear\\:sfMdsLU\\:dcSearP\\:itModuleId\\:\\:content', timeout=10000)
        await asyncio.sleep(0.75)  # 等待頁面穩定
        
        message.config(text="Ready to Download!")  # 通知用戶準備就緒
        return

    # download one report according to the report ID
    # check if the page is initialized
    async def run(self, report_ID, num):
        """
        下載單一 IMDS 報告
        
        參數:
        - report_ID: 要下載的報告 ID (字串)
        - num: 報告的序號 (用於顯示進度)
        
        返回值:
        - download_path: 下載檔案的完整路徑
        
        流程:
        1. 檢查瀏覽器是否已初始化
        2. 記錄開始時間
        3. 根據模式 (MotorIMDS 或一般 IMDS) 執行不同的下載流程
        4. 等待下載完成並儲存檔案
        5. 記錄下載時間並返回檔案路徑
        """
        """
        locator()	        找到頁面元素	                self.page.locator('#id')
        fill()	            填入文字                       self.page.fill('#input', 'text')
        press()	            按鍵                           self.page.press('#input', 'Enter')
        click()	            點擊	                       self.page.click('#button')
        hover()	            懸停	                       element.hover()
        dispatch_event()	觸發事件	                   element.dispatch_event('click')
        expect_download()	等待下載	                   async with page.expect_download()
        keyboard.press()	鍵盤按鍵	                   page.keyboard.press('Tab')
        query_selector()	查詢元素	                   page.query_selector('#img')
        get_attribute()	    獲取屬性	                   element.get_attribute('src')
        wait_for_selector()	等待元素出現	               await page.wait_for_selector('#id')
        select_option()  	選擇選項	                   await page.select_option('#select', 'option')
        wait_for_event()	等待事件	                   async with page.expect_event('click')
        wait_for_timeout()	等待時間	                   await page.wait_for_timeout(1000)
        wait_for_function()	等待函數	                    wait page.wait_for_function('() => document.querySelector("#id").textContent === "text"')
        wait_for_load_state()	等待載入狀態	           await page.wait_for_load_state('networkidle')
        """

        if not self._check_page_initialized():
            raise Exception("Browser not initialized!")

        start = time.time()
        await self.wait_for_loading()
        await asyncio.sleep(1.5)
        if self.isMotor:
            # 步驟4a: 搜尋報告
            await self.wait_for_loading()
            await asyncio.sleep(1.5)
            await self.page.fill('#pt1\\:dcSear\\:sfMdsLU\\:dcSearP\\:itModuleId\\:\\:content', report_ID)
            await self.page.press('#pt1\\:dcSear\\:sfMdsLU\\:dcSearP\\:itModuleId\\:\\:content', 'Enter')
            await self.wait_for_loading()
            await asyncio.sleep(2)
            ##################################
            # 步驟4b: 表格排序 (確保最新版本在前)
            if not self.descending:
                try:
                    parentLocator = self.page.locator('#pt1\\:dcSear\\:sfMdsLU\\:pc2\\:tMdsRes\\:c7')
                    await parentLocator.hover(timeout=1000)
                    tableLocator = self.page.locator('#pt1\\:dcSear\\:sfMdsLU\\:pc2\\:tMdsRes\\:c7\\:\\:afrSI')
                    await tableLocator.hover(timeout=1000)
                    # 點擊排序按鈕(遞減)
                    await self.page.click('#pt1\\:dcSear\\:sfMdsLU\\:pc2\\:tMdsRes\\:c7\\:\\:afrSI tbody tr td[_afrsortdesc="1"] a.x17t', timeout=1000)
                    self.descending = True
                except PlaywrightTimeoutError:
                    pass
            await self.wait_for_loading()
            ##################################
            # 步驟4c: 選擇第一行結果
            firstRowLocator = self.page.locator('#pt1\\:dcSear\\:sfMdsLU\\:pc2\\:tMdsRes\\:\\:db table.x17f.x184 tbody tr:first-child')
            await firstRowLocator.dispatch_event("click", timeout=5000)
            ##################################
            # 步驟4d: 點擊右鍵選單
            button = self.page.locator('#pt1\\:dcSear\\:sfMdsLU\\:pc2\\:ctbMenT')
            await button.dispatch_event("click", timeout=3000)
            button = self.page.locator('tr#pt1\\:dcSear\\:sfMdsLU\\:pc2\\:tMdsRes\\:pt_mReports')
            await button.dispatch_event("click", timeout=3000)
            ##################################
            # 步驟4e: 選擇 Report 選項
             # 點擊右鍵選單 → Reports → Report
            button = self.page.locator('#pt1\\:dcSear\\:sfMdsLU\\:pc2\\:tMdsRes\\:pt_cmiReport')
            await self.wait_for_loading()
            try:
                async with self.page.expect_download() as download_info:
                    await button.dispatch_event("click", timeout=3000)
                    await self.wait_for_loading()
                    await asyncio.sleep(0.5)
                    await self.page.keyboard.press('Tab')
                    await self.wait_for_loading()
                    await asyncio.sleep(0.5)
                    await self.page.keyboard.press('Enter')
            except PlaywrightTimeoutError:
                async with self.page.expect_download() as download_info:
                    await self.page.keyboard.press('Enter')
        else:
             # 步驟4a: 搜尋報告
            await self.wait_for_loading()
            await asyncio.sleep(2)
            await self.page.fill('#pt1\\:dcCmds\\:sfIbLU\\:itModuleId\\:\\:content', report_ID)
            await self.page.press('#pt1\\:dcCmds\\:sfIbLU\\:itModuleId\\:\\:content', 'Enter')         
            await self.wait_for_loading()
            await asyncio.sleep(4)
             # 步驟4b: 選單操作
            button = self.page.locator('#pt1\\:dcCmds\\:sfIbLU\\:pc2\\:ctbMenT')
            await button.dispatch_event("click", timeout=5000)
            button = self.page.locator('tr#pt1\\:dcCmds\\:sfIbLU\\:pc2\\:tResult\\:pt_mReports')
            await button.dispatch_event("click", timeout=3000)
            button = self.page.locator('#pt1\\:dcCmds\\:sfIbLU\\:pc2\\:tResult\\:pt_cmiReport')
            await self.wait_for_loading()
             # 步驟4c: 等待下載
            async with self.page.expect_download() as download_info:
                await button.dispatch_event("click", timeout=3000)
        download = await download_info.value
        download_path = "./3_Downloaded_IMDS_Report/" + f"MDSReport_{report_ID}.pdf"
        await download.save_as(download_path)
        print(f"{num} - Download Time for Report {report_ID}: {time.time()-start}s")
        return download_path

    async def wait_for_loading(self):
        # 檢查頁面是否已初始化
        if not self._check_page_initialized():
            return
        # 持續檢查頁面載入狀態
        while True:
            # 尋找載入狀態指示器圖片
            img_src = await self.page.query_selector('#pt1\\:pt_si1 img')
            # 如果找不到圖片，等待1秒後繼續檢查
            if img_src is None:
                await asyncio.sleep(1)
                continue
            # 獲取圖片來源，判斷載入狀態
            src_value = await img_src.get_attribute('src')
            # 如果圖片是 status_idle.gif，表示載入完成
            if src_value == '/imdsnt/images/headerbar/status_idle.gif':
                break
            # 如果還沒載入完成，等待1秒後繼續檢查
            await asyncio.sleep(1)

    # logout and close the browser
    async def shut_browser(self):
        # 步驟1: 如果頁面存在，先登出並關閉頁面
        if self.page:
            await self.page.click('#pt1\\:pt_ctbLogoff')  # 點擊登出按鈕
            await self.page.close()                       # 關閉頁面
        # 步驟2: 如果瀏覽器存在，關閉瀏覽器
        if self.browser:
            await self.browser.close()
        # 步驟3: 如果 Playwright 存在，停止 Playwright 引擎
        if self.playwright:
            await self.playwright.stop()

# a class to generate required outputs
class CASMatcher:
    def __init__(self, paths_compareList, preprocesseds, names_compareList, message):
        self.compareLists = []
        self.error_standard = False
        self.error_reports = False
        self.error_output = False
        self.message = message
        for i, preprocessed in enumerate(preprocesseds):
            if preprocessed:
                self.compareLists.append(pd.read_excel(paths_compareList[i]))
            else:
                self.compareLists.append(self.preprocess_compareList(paths_compareList[i], names_compareList[i]))
        self.names_compareList = names_compareList

    # uses for recognizing the "Tree Level" and helps get the required tables from IMDS report
    def is_single_digit(self, string):
        return len(string) == 1 and string.isdigit()

    # uses for recognizing the "CAS Number" and helps to get the CAS Number rows only
    def find_numeric_hyphen_strings(self, string):
        pattern_1 = r'^\d+-\d+$'
        pattern_2 = r'^\d+-\d+-\d+$'
        if re.match(pattern_1, string) or re.match(pattern_2, string):
            return True
        return False

    # uses for renaming the columns and modifies the types after extracting the required tables from the IMDS report
    def change_type(self, df):
        df = df.rename(
            columns={
                df.columns[0]: 'Level', df.columns[1]: 'Substance Name',
                df.columns[2]: 'CAS Number' ,df.columns[3]: 'IMDS ID',
                df.columns[4]: 'Quantity', df.columns[5]: 'Weight',
                df.columns[6]: 'Portion', df.columns[7]: 'maxPortion'
                }
            )
        df['maxPortion'] = df['maxPortion'].map(lambda x: x.split('-')[-1].strip())
        df['maxPortion'] = pd.to_numeric(df['maxPortion'], errors='coerce')
        df['Level'] = pd.to_numeric(df['Level'], errors='coerce')
        df['Portion'] = pd.to_numeric(df['Portion'], errors='coerce')
        df['maxPortion'] = pd.to_numeric(df['maxPortion'], errors='coerce')
        df['Weight'] = pd.to_numeric(df['Weight'], errors='coerce')
        df['homoWW%'] = df['maxPortion'].combine_first(df['Portion'])
        df['homoWW%'] = pd.to_numeric(df['homoWW%'], errors='coerce')
        df['homoPPM'] = df['homoWW%']*10000
        df['homoPPM'] = pd.to_numeric(df['homoPPM'], errors='coerce')
        df = df.drop(columns=['IMDS ID', 'Quantity', 'Portion', 'maxPortion'])
        return df

    # *****extracts the required tables from IMDS report in docx format and output as Pandas DataFrame*****
    def get_MDSReport_docx(self, path_MDS):
        doc = Document(path_MDS)
        data = []
        for table in doc.tables:
            for row in table.rows:
                temp = row.cells[0].text.strip(" |-,")
                if self.is_single_digit(temp):
                    data.append([cell.text for cell in row.cells])
        data = np.array(data)[:,0:8]
        data = np.char.strip(data.astype(str), " |-,")
        df = pd.DataFrame(data)
        return self.change_type(df)

    # *****extracts the required tables from IMDS report in pdf format and output as Pandas DataFrame*****
    def get_MDSReport_pdf(self, path_MDS):
        doc = pymupdf.open(path_MDS)
        dfs = []
        for page in doc:
            tabs = page.find_tables()
            if tabs and len(tabs.tables):
                for tab in tabs:
                    data = np.array(tab.extract())[:,0:8]
                    data = np.char.strip(data.astype(str), " |-,")
                    df = pd.DataFrame(data)
                    df = df[df.iloc[:, 0].apply(self.is_single_digit)]
                    dfs.append(self.change_type(df))
        merged_df = pd.concat(dfs, ignore_index=True)
        return merged_df

    # *****processes the standards to desired structure and gets the required columns*****
    def preprocess_compareList(self, filename, name):
        try:
            df = pd.read_excel(filename)
            if "Threshold Limit (ppm)" not in df.columns:
                df['Threshold Limit (ppm)'] = None
            if "Characteristic" not in df.columns:
                df['Characteristic'] = None
            if "Group" not in df.columns:
                df = df.loc[:, ['CAS Number', 'Chemical Name', 'Threshold Limit (ppm)', 'Characteristic']]
                df['CAS Number'] = df['CAS Number'].str.strip(" ,")
                df['CAS Number'] = df['CAS Number'].str.split(',')
                df = df.explode('CAS Number').reset_index(drop=True)
                df['Chemical Name'] = df.index
            else:
                df = df.loc[:, ['CAS Number', 'Group', 'Threshold Limit (ppm)', 'Characteristic']]
                df.rename(columns={"Group": "Chemical Name"}, inplace=True)
            df = df.dropna(subset='CAS Number').drop_duplicates(subset='CAS Number')
            df = df[df['CAS Number'].apply(self.find_numeric_hyphen_strings)]
            df.to_excel('./Match_Temp_Files/' + 'processed_' + filename.split('/')[-1], index=False)
            return df
        except:
            self.message.config(text=f"Invalid Standard Format: {name}")
            self.error_standard = True
            return None

    # *****create the Pandas DataFrame to store the data for the excel outputs in "Total List" sheet*****
    def get_total(self, result):
        result = result[result.iloc[:, 2].apply(self.find_numeric_hyphen_strings)].loc[:,["Level", "Substance Name", "CAS Number", "Chemical Name"]]
        result.rename(columns={"Chemical Name": "Found"}, inplace=True)
        result['Found'] = result['Found'].fillna('no')
        result.loc[result['Found'] != 'no', 'Found'] = "yes"
        return result

    # *****create the Pandas DataFrame to store the data for the excel outputs in "Summary" sheet*****
    def get_summary(self, result):
        level = 0
        yielded_rows = []
        result.loc[result['Chemical Name'].isna(), 'homoPPM'] = None
        for _, row in result[::-1].iterrows():
            currL = row['Level']
            if pd.notna(row['Chemical Name']) or level > currL:
                yielded_rows.append(row)
                level = currL
        yielded_rows = yielded_rows[::-1]
        ppm = {'Sub Level Weight': [], 'exceed': [], 'Parents': []}
        current_weight = 0
        group = {}
        parent = []
        for row in yielded_rows:
            if pd.isna(row['Chemical Name']):
                if pd.notna(row["Weight"]):
                    current_weight = row["Weight"]
                else:
                    current_weight = 0
                ppm['Sub Level Weight'].append(None)
                group = {}
                ppm['exceed'].append(0)
                while len(parent) > 0 and parent[-1][0] >= row['Level']:
                    parent.pop()
                parent.append((row['Level'], row['Substance Name']))
                ppm['Parents'].append(None)
            else:
                portion = row['homoPPM']/1000000
                ppm['Sub Level Weight'].append(portion*current_weight)
                group.setdefault(row['Chemical Name'], 0)
                group[row['Chemical Name']] += row['homoPPM']
                if pd.notna(row["Threshold Limit (ppm)"]) and group[row['Chemical Name']] > row['Threshold Limit (ppm)']:
                    ppm['exceed'].append(1)
                    ppm['Parents'].append(list(parent))
                else:
                    ppm['exceed'].append(0)
                    ppm['Parents'].append(None)
        summary = pd.DataFrame(yielded_rows).assign(**ppm)
        return summary

    # *****combine the DataFrame from standards to IMDS report on "CAS Number" and get the output result by calling get_total() and get_summary()*****
    # *****return the report summary DataFrame (reportLevel) if and only if existing substance exceed threshold limit*****
    def get_result(self, path_MDS):
        try:
            file_type = path_MDS.split('.')[-1]
            if file_type == 'docx':
                data_MDS = self.get_MDSReport_docx(path_MDS)
            else:
                data_MDS = self.get_MDSReport_pdf(path_MDS)
        except:
            self.message.config(text=f"Invalid Report Format: {path_MDS.split('/')[-1]}")
            self.error_reports = True
            return None
        reportLevel = []
        for i, compareList in enumerate(self.compareLists):
            result = pd.merge(
                data_MDS.copy(),
                compareList.copy(), on='CAS Number',
                how='left')
            count = result['Chemical Name'].notnull().sum()
            highlight = 'no'
            total = self.get_total(result.copy())
            if count:
                summary = self.get_summary(result.copy())
                if summary['exceed'].sum():
                    df = summary.loc[(summary['exceed'] == 1), ['Substance Name', 'CAS Number', 'homoPPM', 'Threshold Limit (ppm)', 'Parents', 'Characteristic']].reset_index(drop=True)
                    df.insert(0, 'Standard', None)
                    df.loc[0, 'Standard'] = self.names_compareList[i]
                    reportLevel.append(df)
                    highlight = 'yes'
                summary = summary.loc[:, ['Level', 'Substance Name', 'CAS Number', 'homoWW%', 'Sub Level Weight', 'homoPPM', 'Threshold Limit (ppm)', 'Characteristic', 'exceed']]
            filename = path_MDS.split('/')[-1].split('.')[0] + '.xlsx'
            filepath = './4_Output_Result/' + self.names_compareList[i] + '%' + str(count) + "%" + highlight + '%' + filename
            try:
                with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
                    total.to_excel(writer, sheet_name="Total List", index=False)
                    if count:
                        summary.to_excel(writer, sheet_name="Summary", index=False)
                        highlight_excel(writer)
                    fit_excel(writer)
            except:
                self.message.config(text=f"Error on Create Output: Check if {filename} is closed!")
                self.error_output = True
                return None
        if len(reportLevel):
            return pd.concat(reportLevel, ignore_index=True)
        else:
            return None

# a class to create a multiple selection combo box, used for standards selection
class MultiSelectComboBox(tb.Frame):
    def __init__(self, master, options, max_height=5, **kwargs):
        super().__init__(master, **kwargs)
        self.options = None
        self.selected_options = []
        self.max_height = max_height
        self.combo_var = tb.StringVar()
        self.combo = tb.Entry(self, textvariable=self.combo_var, state="readonly")
        self.combo.pack(fill="x", padx=5, pady=5)
        self.combo.bind("<Button-1>", self.toggle_dropdown)
        self.dropdown_frame = tb.Frame()
        self.listbox = Listbox(self.dropdown_frame, selectmode="multiple")
        self.update_listbox(options)
        self.listbox.pack(fill="both", expand=True)
        self.listbox.bind("<<ListboxSelect>>", self.set_entry_text)
        self.dropdown_visible = False

    # update the choices in the combo box
    def update_listbox(self, options):
        self.options = options
        self.listbox.delete(0, "end")
        for option in self.options:
            self.listbox.insert("end", option)
        listbox_height = min(len(options), self.max_height)
        self.listbox.config(height=listbox_height)

    # show/hide the listbox when onclick on the combo box
    def toggle_dropdown(self, event):
        if self.dropdown_visible:
            self.dropdown_frame.place_forget()
        else:
            self.dropdown_frame.place(in_=self.combo, relx=0, rely=1, relwidth=1)
        self.dropdown_visible = not self.dropdown_visible

    # get the selected options of the listbox
    def get_selected_options(self):
        selected_indices = self.listbox.curselection()
        self.selected_options = [self.listbox.get(i) for i in selected_indices]
        return self.selected_options

    # set the combo box to display the number of chosen standards
    def set_entry_text(self, event):
        indices = self.listbox.curselection()
        if len(indices) > 0:
            self.combo_var.set(f"{len(indices)} Standards Selected!")
        else:
            self.combo_var.set("")

# a class to rendar the UI and contain the controller functions for the UI elements
class CASMatcherApp:
    def __init__(self):
        self.root = tb.Window(themename="superhero")
        self.root.title("CAS Matcher v3.1.3")
        self.root.geometry("720x850")
        self.autoDownloader = None
        self.browser_loop = asyncio.new_event_loop()
        self.browser_thread = Thread(target=self.__start_event_loop__)
        self.browser_thread.start()
        self.matcher = None
        self.tree_data = []

        self.notebook = tb.Notebook(self.root)
        self.notebook.pack(pady=10)
        
        self.tab_autoDownload = tb.Frame(self.notebook)
        self.edge_path = find_edge_path()
        if self.edge_path is None:
            self.__create_browser_selection_frame__()
        self.__create_start_close_browser__()
        self.__create_page_setting__()
        self.__create_standard_selection_frame__()
        self.__create_list_upload_frame__()
        self.__create_progress_frame__()
        self.__create_message_frame__()
        self.__create_treeview_frame__()
        self.__create_clear_frame__()
        self.notebook.add(self.tab_autoDownload, text="AutoDownload")

        self.tab_casmatcher = tb.Frame(self.notebook)
        self.__create_standard_selection_frame_cas__()
        self.__create_report_upload_frame_cas__()
        self.__create_progress_frame_cas__()
        self.__create_message_frame_cas__()
        self.__create_treeview_frame_cas__()
        self.__create_clear_frame_cas__()
        self.notebook.add(self.tab_casmatcher, text="CasMatcher")
        
        self.update_tree()
        self.select_List.dropdown_frame.lift()
        self.select_List_cas.dropdown_frame.lift()
        
        self.root.mainloop()
        if self.autoDownloader is not None:
            asyncio.run_coroutine_threadsafe(self.autoDownloader.shut_browser(), self.browser_loop)
        self.browser_loop.call_soon_threadsafe(self.browser_loop.stop)
        self.browser_thread.join()

    def __create_browser_selection_frame__(self):
        frame = tb.Labelframe(self.tab_autoDownload, text="Locate Edge Browser", padding=(10, 5))
        frame.pack(pady=5)
        self.selectBrowser = tb.Button(frame, text="Select", command=self.get_edge_path)
        self.selectBrowser.pack(pady=5)

    def __create_standard_selection_frame__(self):
        frame = tb.Labelframe(self.tab_autoDownload, text="Choose/Upload the Standard", padding=(10, 5))
        frame.pack(pady=10)
        self.List_Names = [path for path in os.listdir("./1_Standard") if "~$" not in path]
        self.select_List = MultiSelectComboBox(frame, options=self.List_Names)
        self.select_List.grid(row=1, column=0, padx=5, pady=5)
        self.upload_standard_button = tb.Button(frame, text="Upload Standards", command=self.upload_standards)
        self.upload_standard_button.grid(row=1, column=1, padx=5, pady=5)

    def __create_standard_selection_frame_cas__(self):
        frame = tb.Labelframe(self.tab_casmatcher, text="Choose/Upload the Standard", padding=(10, 5))
        frame.pack(pady=10)
        self.select_List_cas = MultiSelectComboBox(frame, options=self.List_Names)
        self.select_List_cas.grid(row=1, column=0, padx=5, pady=5)
        self.upload_standard_button_cas = tb.Button(frame, text="Upload Standards", command=self.upload_standards)
        self.upload_standard_button_cas.grid(row=1, column=1, padx=5, pady=5)

    def __create_report_upload_frame_cas__(self):
        frame = tb.Labelframe(self.tab_casmatcher, text="Upload the Reports", padding=(10, 5))
        frame.pack(pady=10)
        self.upload_report_button_cas = tb.Button(frame, text="Reports", command=self.upload_reports_cas)
        self.upload_report_button_cas.pack(pady=5)

    def __create_list_upload_frame__(self):
        frame = tb.Labelframe(self.tab_autoDownload, text="Upload the List to Download & Process", padding=(10, 5))
        frame.pack(pady=10)
        self.upload_list_button = tb.Button(frame, text="Upload", command=self.upload_lists)
        self.upload_list_button.grid(row=0, column=0, padx=66, pady=5, sticky="w")
        self.process_button = tb.Button(frame, text="Process", command=self.perform_process)
        self.process_button.grid(row=0, column=1, padx=66, pady=5, sticky="e")

    def __create_start_close_browser__(self):
        frame = tb.Labelframe(self.tab_autoDownload, text="Open/Close Browser", padding=(10, 2))
        frame.pack(pady=5)
        self.open_button = tb.Button(frame, text="Open", command=self.init_browser)
        self.open_button.grid(row=0, column=0, padx=60, pady=2, sticky="w")
        self.close_button = tb.Button(frame, text="Close", command=self.close_browser)
        self.close_button.grid(row=0, column=1, padx=60, pady=2, sticky="e")

    def __create_page_setting__(self):
        frame = tb.Labelframe(self.tab_autoDownload, text="Navigate", padding=(10, 5))
        frame.pack(pady=5)
        self.headless_var = tb.BooleanVar(value=True)
        self.headless_checkbox = tb.Checkbutton(frame, text="Show Browser", variable=self.headless_var)
        self.headless_checkbox.grid(row=2, column=0, columnspan=1, padx=5, pady=5)
        self.isMotor_var = tb.BooleanVar(value=True)
        self.isMotor_checkbox = tb.Checkbutton(frame, text="MotorIMDS", variable=self.isMotor_var)
        self.isMotor_checkbox.grid(row=2, column=1, columnspan=1, padx=5, pady=5)
        self.go_button = tb.Button(frame, text="Go", command=self.navigate)
        self.go_button.grid(row=3, column=0, columnspan=2, padx=5, pady=5)
        
    def __create_progress_frame__(self):
        frame = tb.Frame(self.tab_autoDownload)
        frame.pack(pady=2)
        self.progress_bar = tb.Progressbar(frame, length=550, value=0)
        self.progress_bar.pack(pady=3)

    def __create_progress_frame_cas__(self):
        frame = tb.Frame(self.tab_casmatcher)
        frame.pack(pady=2)
        self.progress_bar_cas = tb.Progressbar(frame, length=550, value=0)
        self.progress_bar_cas.pack(pady=3)

    def __create_clear_frame__(self):
        frame = tb.Frame(self.tab_autoDownload)
        frame.pack(pady=2)
        self.clear_button = tb.Button(frame, text="Clear Outputs", command=self.clear_folder)
        self.clear_button.pack(pady=3)

    def __create_clear_frame_cas__(self):
        frame = tb.Frame(self.tab_casmatcher)
        frame.pack(pady=2)
        self.clear_button_cas = tb.Button(frame, text="Clear Outputs", command=self.clear_folder)
        self.clear_button_cas.pack(pady=3)

    def __create_message_frame__(self):
        frame = tb.Frame(self.tab_autoDownload)
        frame.pack(pady=5)
        self.message = tb.Label(frame, text="", font=("Helvetica", 12), wraplength=700, justify='center', anchor='center')
        self.message.pack(pady=5)

    def __create_message_frame_cas__(self):
        frame = tb.Frame(self.tab_casmatcher)
        frame.pack(pady=5)
        self.message_cas = tb.Label(frame, text="", font=("Helvetica", 12), wraplength=700, justify='center', anchor='center')
        self.message_cas.pack(pady=5)

    def __create_treeview_frame__(self):
        frame = tb.Labelframe(self.tab_autoDownload, text="Output Results", padding=(10, 5))
        frame.pack(pady=5)
        self.tree_columns = ("list_name", "matching_number", "found_exceed", "filename")
        self.found_list = tb.Treeview(frame, columns=self.tree_columns, show="headings")
        self.found_list.pack(pady=3)
        self.found_list.heading('list_name', text='List Name')
        self.found_list.heading('matching_number', text='Matching Number')
        self.found_list.heading('found_exceed', text='Exceed')
        self.found_list.heading('filename', text='Filename')
        self.found_list.column('list_name', width=180)
        self.found_list.column('matching_number', anchor='center', width=75)
        self.found_list.column('found_exceed', anchor='center', width=75)
        self.found_list.column('filename', width=350)
        self.found_list.bind("<Double-1>", self.open_file)

    def __create_treeview_frame_cas__(self):
        frame = tb.Labelframe(self.tab_casmatcher, text="Output Results", padding=(10, 5))
        frame.pack(pady=5)
        self.found_list_cas = tb.Treeview(frame, columns=self.tree_columns, show="headings")
        self.found_list_cas.pack(pady=3)
        self.found_list_cas.heading('list_name', text='List Name')
        self.found_list_cas.heading('matching_number', text='Matching Number')
        self.found_list_cas.heading('found_exceed', text='Exceed')
        self.found_list_cas.heading('filename', text='Filename')
        self.found_list_cas.column('list_name', width=180)
        self.found_list_cas.column('matching_number', anchor='center', width=75)
        self.found_list_cas.column('found_exceed', anchor='center', width=75)
        self.found_list_cas.column('filename', width=350)   
        self.found_list_cas.bind("<Double-1>", self.open_file_cas)

    # create async loop for the auto browsing and downloading purpose
    def __start_event_loop__(self):
        asyncio.set_event_loop(self.browser_loop)
        self.browser_loop.run_forever()

    # allow double clicks on the tree view items to open the files (AutoDownload)
    def open_file(self, event):
        selected_item = self.found_list.focus()
        if not selected_item:
            return
        x, y, z, k = self.found_list.item(selected_item, 'value')
        file_path = f"4_Output_Result/{x}%{y}%{z}%{k}"
        absolute_path = os.path.abspath(file_path)
        if os.path.isfile(absolute_path):
            os.startfile(absolute_path) 

    # allow double clicks on the treeview items to open the files (CasMatcher)
    def open_file_cas(self, event):
        selected_item = self.found_list_cas.focus()
        if not selected_item:
            return
        x, y, z, k = self.found_list_cas.item(selected_item, 'value')
        file_path = f"4_Output_Result/{x}%{y}%{z}%{k}"
        absolute_path = os.path.abspath(file_path)
        if os.path.isfile(absolute_path):
            os.startfile(absolute_path)

    # get the output filenames and structure them as data for the treeview
    def get_tree_data(self):
        file_paths = os.listdir('./4_Output_Result')
        self.tree_data.clear()
        if len(file_paths) != 0:
            for file_path in file_paths:
                splits = file_path.split('%')
                if "~$" not in splits[0]:
                    self.tree_data.append((splits[0], splits[1], splits[2], splits[3]))

    # update the treeview according to the tree data
    def update_tree(self):
        for item in self.found_list.get_children():
            self.found_list.delete(item)
        for item in self.found_list_cas.get_children():
            self.found_list_cas.delete(item)
        self.get_tree_data()
        for data in self.tree_data:
            self.found_list.insert('', 'end', values=data)
            self.found_list_cas.insert('', 'end', values=data)

    # upload the standards to the "1_Standard" folder and allow for selection in the combo box
# ------------------------------------------------------------------
# Pull every standard from the designated folder – no file dialog
# ------------------------------------------------------------------
    def upload_standards(self):
        src_dir = STANDARD_SOURCE_DIR
        if not os.path.isdir(src_dir):
            msg = f"Standard folder not found: {src_dir}"
            self.message.config(text=msg)
            self.message_cas.config(text=msg)
            return

        excel_files = [f for f in os.listdir(src_dir)
                    if f.lower().endswith(".xlsx") and "~$" not in f]

        if not excel_files:
            msg = "No *.xlsx standards found in the designated folder!"
            self.message.config(text=msg)
            self.message_cas.config(text=msg)
            return

        for fname in excel_files:
            src_path = os.path.join(src_dir, fname)
            dst_path = os.path.join("./1_Standard", fname)
            # copy only if the file is new or has changed
            if not os.path.isfile(dst_path) or \
            os.path.getmtime(src_path) > os.path.getmtime(dst_path):
                shutil.copy2(src_path, dst_path)

        self.update_combobox()
        self.message.config(text="Standards synced from designated folder.")
        self.message_cas.config(text="Standards synced from designated folder.")

    # upload the IMDS reports and calling the reports process function to generate outputs (CasMatcher)
    def upload_reports_cas(self):
        if not self.get_matcher_cas():
            return
        file_paths = filedialog.askopenfilenames(filetypes=[("Word Documents", "*.docx"), ("PDF Files", "*.pdf")])
        if file_paths:
            Thread(target=self.perform_process_thread_cas, args=(file_paths, )).start()

    # upload the "to download list" to the specific folders
    def upload_lists(self):
        file_paths = filedialog.askopenfilenames(filetypes=[("Excel files", "*.xlsx"), ("CSV files", "*.csv")])
        if len(file_paths) != 0:
            for file_path in file_paths:
                shutil.copy(file_path, "./2_IMDS_Lists_to_be_Downloaded")
            self.message.config(text="Lists Upload Successfully!")

    # update the standard selection combo box
    def update_combobox(self):
        self.List_Names = [path for path in os.listdir("./1_Standard") if "~$" not in path]
        self.select_List.update_listbox(self.List_Names)
        self.select_List_cas.update_listbox(self.List_Names)

    # clear the output result folder
    def clear_folder(self):
        self.message.config(text="Clearing ...")
        self.message_cas.config(text="Clearing ...")
        files = os.listdir('./4_Output_Result')
        for file in files:
            try:
                os.remove(os.path.join('./4_Output_Result', file))
            except:
                self.message.config(text=f"Error on Clear Output: Check if {file} is closed!")
                self.message_cas.config(text=f"Error on Clear Output: Check if {file} is closed!")
                return        
        self.update_tree()
        self.message.config(text="All files are cleared!")
        self.message_cas.config(text="All files are cleared!")

    # get the CASMatcher objects for generate outputs (AutoDownload)
    def get_matcher(self):
        list_names = self.select_List.get_selected_options()
        if len(list_names) == 0:
            self.message.config(text="Choose a standard!")
            return False
        else:
            self.matcher = None
            path = []
            preprocesseds = []
            name = []
            for list_name in list_names:
                path_processed = './Match_Temp_Files/' + 'processed_' + list_name
                path_non_processed = './1_Standard/' + list_name
                if os.path.exists(path_processed):
                    path.append(path_processed)
                    preprocesseds.append(True)
                else:
                    path.append(path_non_processed)
                    preprocesseds.append(False)
                name.append(list_name.split('.')[0])
            self.matcher = CASMatcher(path, preprocesseds, name, self.message)
            if self.matcher.error_standard:
                self.matcher = None
                return False
            return True

    # get the CASMatcher objects for generate outputs (CasMatcher)
    def get_matcher_cas(self):
        list_names = self.select_List_cas.get_selected_options()
        if len(list_names) == 0:
            self.message_cas.config(text="Choose a standard!")
            return False
        else:
            self.matcher = None
            path = []
            preprocesseds = []
            name = []
            for list_name in list_names:
                path_processed = './Match_Temp_Files/' + 'processed_' + list_name
                path_non_processed = './1_Standard/' + list_name
                if os.path.exists(path_processed):
                    path.append(path_processed)
                    preprocesseds.append(True)
                else:
                    path.append(path_non_processed)
                    preprocesseds.append(False)
                name.append(list_name.split('.')[0])
            self.matcher = CASMatcher(path, preprocesseds, name, self.message_cas)
            if self.matcher.error_standard:
                self.matcher = None
                return False
            return True

    # get the all the paths of the "to download list"
    def get_list_to_download(self):
        file_paths = os.listdir("./2_IMDS_Lists_to_be_Downloaded")
        if len(file_paths) == 0:
            return None
        return file_paths

    # process a IMDS report according to the file path and get the output result
    def process_file(self, file_path):
        if self.matcher is None:
            return None
        reportLevel = self.matcher.get_result(file_path)
        self.update_tree()
        return reportLevel

    # uses for selecting the edge broswer if default edge path is not found
    def get_edge_path(self):
        file_path = filedialog.askopenfilename(
            title="Select Edge Browser",
            filetypes=[("Executable Files", "*.exe"), ("All Files", "*.*")]
        )
        if file_path:
            if file_path.split('/')[-1] == "msedge.exe":
                self.edge_path = file_path
                self.message.config(text="Edge Browser is Located!")
            else:
                self.message.config(text="Please Choose Edge Browser!")

    # call the reports download (according to the report IDs in "to download list") and process function (AutoDownload)
    def perform_process(self):
        if self.autoDownloader is None:
            self.message.config(text="Please Open Browser and Login First!")
            return
        if not self.get_matcher():
            return
        if self.get_list_to_download() is None:
            self.message.config(text="Upload MDS ID list to be download!")
            return
        Thread(target=self.perform_process_thread, args=(True, )).start()

    # reports download and process function (AutoDownload)
    def perform_process_thread(self, allowRepeat):
        # First Trial: allowRepeat == True
        # Second Trial (repeat): allowRepeat == False
        report_lists = self.get_list_to_download()
        if report_lists is None:
            return
        for report_list in report_lists:
            failPath = f'./2_IMDS_Lists_to_be_Downloaded/{report_list.split(".")[0]}.xlsx'
            if not allowRepeat:
                failPath = f'./6_DownloadFail/{report_list.split(".")[0]}.xlsx'
                print(f"{report_list} - Re-Downloading Failed Report IDs...")
            report_IDs = self.get_report_IDs(report_list)
            if report_IDs is None:
                return
            num = len(report_IDs)
            downloadCount = 0
            failedDownload = []
            reportLevels = []
            self.progress_bar['value'] = 0
            for i, report_ID in enumerate(report_IDs, 1):
                try:
                    self.message.config(text="Downloading...")
                    if self.autoDownloader is None:
                        self.message.config(text="Browser not initialized!")
                        return
                    future = asyncio.run_coroutine_threadsafe(self.autoDownloader.run(report_ID, i), self.browser_loop)
                    file_path = future.result()
                    self.message.config(text="Processing...")
                    reportLevel = self.process_file(file_path)
                    if self.matcher and (self.matcher.error_reports or self.matcher.error_output):
                        return
                    if reportLevel is not None:
                        reportLevel.insert(0, 'Report ID', '')
                        reportLevel.loc[0, 'Report ID'] = 'MDSReport_' + report_ID
                        reportLevels.append(reportLevel)
                    downloadCount += 1
                except PlaywrightTimeoutError:
                    print(f"{i} - Download Failed for Report {report_ID}.")
                    failedDownload.append(report_ID)
                    pass
                self.progress_bar['value'] += 100/num
            if len(reportLevels):
                reportLevels = pd.concat(reportLevels, ignore_index=True)
                filename = f"auto_{report_list.split('.')[0]}.xlsx"
                try:
                    path = f'./5_Summary/{filename}'
                    if not allowRepeat:
                        path = f'./5_Summary/reDownload_{filename}'
                    with pd.ExcelWriter(path, engine='openpyxl') as writer:
                        reportLevels.to_excel(writer, sheet_name="Report Summary", index=False)
                        fit_excel(writer)
                except:
                    self.message.config(text=f"Error on Create Summary: Check if {filename} is closed!")
                    return
            print(f"Number of downloaded reports: {downloadCount}")
            print(f"Proportion of successfully downloaded reports: {downloadCount}/{num}")
            print(f"Failed to Download Reports: {failedDownload}")
            print("------------------------------------------------------------------------")
            if (len(failedDownload)):
                df = pd.DataFrame({'ITEM IMDS ID NO.': failedDownload})
                df.to_excel(failPath, index=False)
            self.message.config(text="All files are downloaded and processed!")
        if allowRepeat:
            self.perform_process_thread(False)

    # reports process function (CasMatcher)
    def perform_process_thread_cas(self, file_paths):
        num = len(file_paths)
        reportLevels = []
        self.message_cas.config(text="Processing...")
        self.progress_bar_cas['value'] = 0
        for file_path in file_paths:
            reportLevel = self.process_file(file_path)
            if self.matcher.error_reports or self.matcher.error_output:
                return
            if reportLevel is not None:
                reportLevel.insert(0, 'Report ID', '')
                report_ID = file_path.split('/')[-1].split('.')[0]
                reportLevel.loc[0, 'Report ID'] = report_ID
                reportLevels.append(reportLevel)
            self.progress_bar_cas['value'] += 100/num
        if len(reportLevels):
            reportLevels = pd.concat(reportLevels, ignore_index=True)
            filename = f"cas_{datetime.now().strftime("%d-%m-%Y_%H-%M-%S")}.xlsx"
            try:
                with pd.ExcelWriter(f'./5_Summary/{filename}', engine='openpyxl') as writer:
                    reportLevels.to_excel(writer, sheet_name="Report Summary", index=False)
                    fit_excel(writer)
            except:
                self.message_cas.config(text=f"Error on Create Summary: Check if {filename} is closed!")
                return
        self.message_cas.config(text="All files are downloaded and processed!")

    # start the browser
    def init_browser(self):
        if self.autoDownloader is not None:
            self.message.config(text="Browser is Already On!")
        else:
            headless = not self.headless_var.get()
            self.autoDownloader = AutoDownload(headless, self.edge_path)
            asyncio.run_coroutine_threadsafe(self.autoDownloader.start_browser(self.message), self.browser_loop)

    # close the browser
    def close_browser(self):
        if self.autoDownloader is not None:
            asyncio.run_coroutine_threadsafe(self.autoDownloader.shut_browser(), self.browser_loop)
            self.autoDownloader = None
            self.message.config(text="Browser is Closed!")
        else:
            self.message.config(text="No Browser is Opened!")

    # navigate to different search page
    def navigate(self):
        if self.autoDownloader is None:
            self.message.config(text="Open browser first.")
            return
        self.message.config(text="Navigating ...")
        isMotor = self.isMotor_var.get()
        if isMotor:
            asyncio.run_coroutine_threadsafe(self.autoDownloader.go_outbox(self.message), self.browser_loop)
        else:
            asyncio.run_coroutine_threadsafe(self.autoDownloader.go_inbox(self.message), self.browser_loop)
        

    # get the report IDs from the "to download list"
    def get_report_IDs(self, report_list):
        result = []
        if report_list:
            report_type = report_list.split('.')[-1]
            if (report_type == "csv"):
                df = pd.read_csv("./2_IMDS_Lists_to_be_Downloaded/" + report_list)
            else:
                df = pd.read_excel("./2_IMDS_Lists_to_be_Downloaded/" + report_list)
            try:
                df = df.loc[:,'ITEM IMDS ID NO.'].dropna().drop_duplicates().astype(str)
            except:
                self.message.config(text=f'Invalid List Format: {report_list}')
                os.remove("./2_IMDS_Lists_to_be_Downloaded/" + report_list)
                return None
            df = df.map(lambda x: x.split('/')[0].strip() if "." not in x.split('/')[-1].strip() else None).dropna().drop_duplicates()
            os.remove("./2_IMDS_Lists_to_be_Downloaded/" + report_list)
            print(f"{report_list} - Total Number of Report to Download: {len(df.to_list())}")
            print("------------------------------------------------------------------------")
            return df.to_list()

# first function to call
def main():
    check_essential_folders()
    CASMatcherApp()

# start point of the program
if __name__ == '__main__':
    main()
