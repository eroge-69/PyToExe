import os
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
import pandas as pd
from docx2pdf import convert
import tempfile
import threading
import re
from collections import defaultdict

class ContractGeneratorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Insurance Contract Generator")
        self.root.geometry("700x500")
        
        # Variables
        self.word_path = tk.StringVar()
        self.excel_path = tk.StringVar()
        self.output_folder = tk.StringVar()
        self.progress = tk.DoubleVar()
        self.status = tk.StringVar(value="Ready")
        
        # Create GUI
        self.create_widgets()
        
    def create_widgets(self):
        # Main frame
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Title
        ttk.Label(main_frame, text="Insurance Contract Generator", 
                 font=('Helvetica', 16, 'bold')).grid(row=0, column=0, columnspan=3, pady=10)
        
        # File selection
        ttk.Label(main_frame, text="Word Template:").grid(row=1, column=0, sticky=tk.W, pady=5)
        ttk.Entry(main_frame, textvariable=self.word_path, width=50).grid(row=1, column=1, padx=5)
        ttk.Button(main_frame, text="Browse", command=self.browse_word).grid(row=1, column=2)
        
        ttk.Label(main_frame, text="Excel Data File:").grid(row=2, column=0, sticky=tk.W, pady=5)
        ttk.Entry(main_frame, textvariable=self.excel_path, width=50).grid(row=2, column=1, padx=5)
        ttk.Button(main_frame, text="Browse", command=self.browse_excel).grid(row=2, column=2)
        
        ttk.Label(main_frame, text="Output Folder:").grid(row=3, column=0, sticky=tk.W, pady=5)
        ttk.Entry(main_frame, textvariable=self.output_folder, width=50).grid(row=3, column=1, padx=5)
        ttk.Button(main_frame, text="Browse", command=self.browse_output).grid(row=3, column=2)
        
        # Progress
        ttk.Label(main_frame, text="Progress:").grid(row=4, column=0, sticky=tk.W, pady=10)
        ttk.Progressbar(main_frame, variable=self.progress, length=450).grid(
            row=4, column=1, columnspan=2, sticky=tk.W)
        
        # Status and log
        ttk.Label(main_frame, text="Status:").grid(row=5, column=0, sticky=tk.W, pady=5)
        ttk.Label(main_frame, textvariable=self.status).grid(row=5, column=1, sticky=tk.W)
        
        # Log output
        log_frame = ttk.LabelFrame(main_frame, text="Processing Log", padding=10)
        log_frame.grid(row=6, column=0, columnspan=3, sticky="ew", pady=10)
        self.log_text = tk.Text(log_frame, height=8, wrap=tk.WORD)
        scrollbar = ttk.Scrollbar(log_frame, command=self.log_text.yview)
        self.log_text.configure(yscrollcommand=scrollbar.set)
        self.log_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Generate button
        ttk.Button(main_frame, text="Generate Contracts", command=self.start_generation).grid(
            row=7, column=1, pady=10)
    
    def browse_word(self):
        file_path = filedialog.askopenfilename(title="Select Word Template", 
                                             filetypes=[("Word Files", "*.docx")])
        if file_path:
            self.word_path.set(file_path)
            self.log(f"Selected template: {file_path}")
    
    def browse_excel(self):
        file_path = filedialog.askopenfilename(title="Select Excel Data File", 
                                             filetypes=[("Excel Files", "*.xlsx *.xls")])
        if file_path:
            self.excel_path.set(file_path)
            self.log(f"Selected data file: {file_path}")
    
    def browse_output(self):
        folder_path = filedialog.askdirectory(title="Select Output Folder")
        if folder_path:
            self.output_folder.set(folder_path)
            self.log(f"Selected output folder: {folder_path}")
    
    def start_generation(self):
        if not all([self.word_path.get(), self.excel_path.get(), self.output_folder.get()]):
            messagebox.showerror("Error", "Please select all required files and folders")
            return
        
        # Clear previous logs
        self.log_text.delete(1.0, tk.END)
        
        # Disable UI during generation
        self.status.set("Processing...")
        self.root.config(cursor="watch")
        
        # Run generation in separate thread
        threading.Thread(
            target=self.generate_contracts,
            args=(self.word_path.get(), self.excel_path.get(), self.output_folder.get()),
            daemon=True
        ).start()
    
    def generate_contracts(self, word_path, excel_path, output_folder):
        try:
            # Read Excel data
            df = pd.read_excel(excel_path, sheet_name='Sheet1')
            total_contracts = len(df)
            self.log(f"Found {total_contracts} contracts to generate")
            
            # Verify required columns
            required_columns = ['Contract No', 'Name_1', 'IDNUM', 'Date', 'Month', 
                              'Year', 'Year word', 'Address', 'Region', 'Branch']
            missing_cols = [col for col in required_columns if col not in df.columns]
            if missing_cols:
                raise ValueError(f"Missing required columns: {', '.join(missing_cols)}")
            
            for index, row in df.iterrows():
                contract_no = str(row['Contract No'])
                self.log(f"\nProcessing contract {contract_no}...")
                
                # Update progress
                self.progress.set((index + 1) / total_contracts * 100)
                self.status.set(f"Processing {index + 1}/{total_contracts}")
                self.root.update_idletasks()
                
                # Prepare contract data
                contract_data = {
                    '«Contract_No»': contract_no,
                    '«Name_1»': str(row['Name_1']).strip(),
                    '«IDNUM»': str(row['IDNUM']).strip(),
                    '«Date»': self.add_ordinal_suffix(str(row['Date'])),
                    '«Month»': str(row['Month']).strip(),
                    '«Year»': str(row['Year']).strip(),
                    '«Year_word»': str(row['Year word']).strip(),
                    '«Address»': str(row['Address']).strip()
                }
                
                self.log(f"Data to insert: {contract_data}")
                
                # Load document
                doc = Document(word_path)
                
                # Replace all placeholders
                replacements = self.replace_all_placeholders(doc, contract_data)
                
                # Verify all placeholders were replaced
                missing = [k for k in contract_data if replacements.get(k, 0) == 0]
                if missing:
                    self.log(f"Warning: These placeholders were not found: {', '.join(missing)}")
                
                # Save output
                region = str(row['Region']).strip()
                branch = str(row['Branch']).strip()
                output_path = self.save_contract(doc, contract_no, region, branch, output_folder)
                
                self.log(f"Successfully saved: {output_path}")
            
            # Completion
            self.status.set("Generation complete")
            messagebox.showinfo("Success", f"Generated {total_contracts} contracts")
            
        except Exception as e:
            self.log(f"ERROR: {str(e)}")
            self.status.set("Error occurred")
            messagebox.showerror("Error", f"Failed to generate contracts: {str(e)}")
        finally:
            self.root.config(cursor="")
            self.progress.set(0)
    
    def replace_all_placeholders(self, doc, contract_data):
        """Replace all placeholders in document while preserving formatting"""
        replacements = defaultdict(int)
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in contract_data.items():
                if self.replace_in_paragraph(paragraph, key, value):
                    replacements[key] += 1
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in contract_data.items():
                            if self.replace_in_paragraph(paragraph, key, value):
                                replacements[key] += 1
                                
        return replacements
    
    def replace_in_paragraph(self, paragraph, key, value):
        """Replace placeholder in a single paragraph while preserving formatting"""
        if key not in paragraph.text:
            return False
            
        # Store original runs
        original_runs = [run for run in paragraph.runs]
        paragraph.clear()
        
        # Rebuild paragraph with replacements
        for run in original_runs:
            if key in run.text:
                # Split the text around the placeholder
                parts = run.text.split(key)
                for i, part in enumerate(parts):
                    if part:
                        new_run = paragraph.add_run(part)
                        self.copy_run_formatting(new_run, run)
                    if i < len(parts) - 1:
                        # Insert the replacement value
                        new_run = paragraph.add_run(value)
                        self.copy_run_formatting(new_run, run)
            else:
                # Keep runs without the placeholder
                new_run = paragraph.add_run(run.text)
                self.copy_run_formatting(new_run, run)
        
        return True
    
    def copy_run_formatting(self, new_run, original_run):
        """Copy all formatting from original run to new run"""
        new_run.bold = original_run.bold
        new_run.italic = original_run.italic
        new_run.underline = original_run.underline
        new_run.font.name = original_run.font.name
        new_run.font.size = original_run.font.size
        if original_run.font.color.rgb:
            new_run.font.color.rgb = original_run.font.color.rgb
    
    def save_contract(self, doc, contract_no, region, branch, output_folder):
        """Save the document to appropriate folder structure"""
        region_folder = os.path.join(output_folder, region)
        branch_folder = os.path.join(region_folder, branch)
        os.makedirs(branch_folder, exist_ok=True)
        
        with tempfile.TemporaryDirectory() as temp_dir:
            # Save temporary DOCX
            temp_docx = os.path.join(temp_dir, f"{contract_no}.docx")
            doc.save(temp_docx)
            
            # Convert to PDF
            pdf_path = os.path.join(branch_folder, f"{contract_no}.pdf")
            convert(temp_docx, pdf_path)
            
        return pdf_path
    
    def add_ordinal_suffix(self, day):
        """Convert day number to ordinal (1 → 1st, 2 → 2nd, etc.)"""
        try:
            day = str(day).strip()
            day_num = int(re.sub(r'[^\d]', '', day))  # Extract numbers only
            if 11 <= (day_num % 100) <= 13:
                suffix = 'th'
            else:
                suffix = {1: 'st', 2: 'nd', 3: 'rd'}.get(day_num % 10, 'th')
            return f"{day_num}{suffix}"
        except ValueError:
            return day
    
    def log(self, message):
        """Add timestamped message to log"""
        timestamp = pd.Timestamp.now().strftime('%H:%M:%S')
        self.log_text.insert(tk.END, f"[{timestamp}] {message}\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()

if __name__ == "__main__":
    root = tk.Tk()
    app = ContractGeneratorApp(root)
    root.mainloop()