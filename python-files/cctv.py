import tkinter as tk
from tkinter import ttk, messagebox
import pandas as pd
import json
import os
import stat
import bcrypt
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tkcalendar import Calendar

# File paths
TASKS_FILE = "tasks.json"
USERS_FILE = "users.json"

def init_data_files():
    """מאתחל קבצי JSON אם אינם קיימים."""
    if not os.path.exists(TASKS_FILE):
        with open(TASKS_FILE, 'w', encoding='utf-8') as f:
            json.dump([], f, ensure_ascii=False, indent=2)
    if not os.path.exists(USERS_FILE):
        hashed = bcrypt.hashpw("admin123".encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
        with open(USERS_FILE, 'w', encoding='utf-8') as f:
            json.dump({"users": [{"username": "admin", "password": hashed, "role": "admin"}]}, f, ensure_ascii=False, indent=2)

class TaskManagerApp:
    """אפליקציית ניהול משימות עם ממשק גרפי."""
    
    def __init__(self, root):
        """מאתחל את חלון האפליקציה הראשי.
        
        Args:
            root (tk.Tk): חלון Tkinter הראשי.
        """
        self.root = root
        self.root.title("ניהול משימות")
        self.current_user = None
        self.entries = {}
        
        init_data_files()
        self._setup_login_ui() # התחלה עם מסך ההתחברות

    def _setup_login_ui(self):
        """מאתחל את ממשק המשתמש של חלון ההתחברות."""
        if hasattr(self, 'app_frame') and self.app_frame.winfo_exists():
            self.app_frame.destroy()
        if hasattr(self, 'footer') and self.footer.winfo_exists():
            self.footer.destroy()

        login_width = 300
        login_height = 200
        self.root.geometry(f"{login_width}x{login_height}")
        self.root.resizable(False, False)
        self._center_window_on_screen(self.root, login_width, login_height)

        self.auth_frame = tk.Frame(self.root)
        self.auth_frame.pack(expand=True)

        tk.Label(self.auth_frame, text="שם משתמש").pack()
        self.username_entry = tk.Entry(self.auth_frame, width=20)
        self.username_entry.pack()
        tk.Label(self.auth_frame, text="סיסמה").pack()
        self.password_entry = tk.Entry(self.auth_frame, show="*", width=20)
        self.password_entry.pack()
        tk.Button(self.auth_frame, text="התחבר", command=self.login).pack(pady=10)

    def _setup_main_ui(self):
        """מאתחל את ממשק המשתמש של האפליקציה המלאה לאחר התחברות."""
        if hasattr(self, 'auth_frame') and self.auth_frame.winfo_exists():
            self.auth_frame.destroy()

        main_width = 850
        main_height = 700
        self._center_window_on_screen(self.root, main_width, main_height)
        self.root.geometry(f"{main_width}x{main_height}")
        self.root.resizable(True, True)


        # Menu bar
        self.menu_bar = tk.Menu(self.root)
        self.root.config(menu=self.menu_bar)
        self.admin_menu = tk.Menu(self.menu_bar, tearoff=0)
        self.menu_bar.add_cascade(label="ניהול", menu=self.admin_menu)
        self.admin_menu.add_command(label="ניהול משתמשים", command=self.show_admin_panel)

        self.app_frame = tk.Frame(self.root)
        self.app_frame.pack(expand=True, fill='both')
        
        self.app_frame.grid_columnconfigure(0, weight=1)
        self.app_frame.grid_columnconfigure(1, weight=0)
        self.app_frame.grid_columnconfigure(2, weight=1)

        self.app_frame.grid_rowconfigure(0, weight=0)
        self.app_frame.grid_rowconfigure(1, weight=0)
        self.app_frame.grid_rowconfigure(2, weight=1)

        self.logged_in_label = tk.Label(self.app_frame, text=f"משתמש מחובר: {self.current_user['username']}", font=("Arial", 10), fg="blue")
        self.logged_in_label.grid(row=0, column=1, pady=5, sticky='n')
        
        tk.Label(self.app_frame, text="הזנת פרטים", font=("Arial", 14)).grid(row=1, column=1, pady=10, sticky='n')

        self._create_input_frame()

        self.footer = tk.Label(self.root, text="Created By Arthur Wasserstein", anchor='w', font=("Arial", 9), fg="gray")
        self.footer.pack(side='bottom', anchor='sw', padx=10, pady=5)

    def _create_input_frame(self):
        """יוצר מסגרת לשדות הקלט, ממורכזת, וכן ממקם את כפתור 'נקה נתונים'."""
        right_frame = tk.Frame(self.app_frame)
        right_frame.grid(row=2, column=1, sticky='') 

        right_frame.grid_columnconfigure(0, weight=1)
        right_frame.grid_columnconfigure(1, weight=0)
        right_frame.grid_columnconfigure(2, weight=1)

        right_frame.grid_rowconfigure(0, weight=0)
        right_frame.grid_rowconfigure(1, weight=1)

        inner_content_frame = tk.Frame(right_frame)
        inner_content_frame.grid(row=1, column=1, sticky='') 

        inner_content_frame.grid_columnconfigure(0, weight=1)
        inner_content_frame.grid_columnconfigure(1, weight=0)
        inner_content_frame.grid_columnconfigure(2, weight=1)


        # Input fields configuration
        fields = [
            ("מס' תיק", lambda parent: tk.Entry(parent, width=15, justify='right')), 
            ("שם היחידה", lambda parent: tk.Entry(parent, width=15, justify='right')), 
            ("שם השוטר המעתיק", lambda parent: tk.Entry(parent, width=15, justify='right')), 
            ("מ.א", lambda parent: tk.Entry(parent, width=15, justify='right')), 
            ("פרטי אירוע", lambda parent: tk.Text(parent, height=4, width=20, relief="solid", borderwidth=1)),
            ("מיקום המצלמה", lambda parent: tk.Entry(parent, width=15, justify='right')), 
            ("מספר מצלמות", lambda parent: ttk.Combobox(parent, values=list(range(1, 11)), state="readonly", width=3)), 
        ]

        current_row = 0

        tk.Button(right_frame, text="נקה נתונים", command=self.clear_form).grid(
            row=0, column=0, columnspan=3, sticky='w', padx=20, pady=5)

        for idx, (label_text, widget_creator) in enumerate(fields):
            tk.Label(inner_content_frame, text=label_text).grid(row=current_row, column=1, sticky='e', padx=(0, 5), pady=0)
            current_row += 1 

            widget = widget_creator(inner_content_frame)
            if label_text == "מספר מצלמות":
                widget.grid(row=current_row, column=1, sticky='', padx=(0, 5), pady=(0, 5))
            else:
                widget.grid(row=current_row, column=1, sticky='ew', padx=(0, 5), pady=(0, 5))
            self.entries[label_text] = widget
            
            current_row += 1 

        # Apply right justification for the Text widgets after creation
        if "פרטי אירוע" in self.entries and isinstance(self.entries["פרטי אירוע"], tk.Text):
            self.entries["פרטי אירוע"].tag_configure("right", justify='right')
            self.entries["פרטי אירוע"].tag_add("right", "1.0", "end")


        # Download frame - centered within inner_content_frame
        download_frame = tk.Frame(inner_content_frame)
        download_frame.grid(row=current_row, column=0, columnspan=3, pady=5) 
        
        download_frame.grid_columnconfigure(0, weight=1)
        download_frame.grid_columnconfigure(1, weight=0)
        download_frame.grid_columnconfigure(2, weight=1)

        date_hour_widgets_frame = tk.Frame(download_frame)
        date_hour_widgets_frame.grid(row=0, column=1, sticky='')

        tk.Label(date_hour_widgets_frame, text="מתאריך").pack(side=tk.RIGHT, padx=(5,0))
        from_date_subframe = tk.Frame(date_hour_widgets_frame)
        from_date_subframe.pack(side=tk.RIGHT, padx=2)
        self.from_date_entry = tk.Entry(from_date_subframe, width=10, justify='right')
        self.from_date_entry.pack(side=tk.RIGHT)
        tk.Button(from_date_subframe, text="📅", command=lambda: self.show_calendar(self.from_date_entry), width=3).pack(side=tk.RIGHT, padx=2)

        tk.Label(date_hour_widgets_frame, text="עד תאריך").pack(side=tk.RIGHT, padx=(5,0))
        to_date_subframe = tk.Frame(date_hour_widgets_frame)
        to_date_subframe.pack(side=tk.RIGHT, padx=2)
        self.to_date_entry = tk.Entry(to_date_subframe, width=10, justify='right')
        self.to_date_entry.pack(side=tk.RIGHT)
        tk.Button(to_date_subframe, text="📅", command=lambda: self.show_calendar(self.to_date_entry), width=3).pack(side=tk.RIGHT, padx=2)

        tk.Label(date_hour_widgets_frame, text="משעה").pack(side=tk.RIGHT, padx=(5,0))
        self.from_hour_entry = tk.Entry(date_hour_widgets_frame, width=10, justify='right')
        self.from_hour_entry.pack(side=tk.RIGHT, padx=2)

        tk.Label(date_hour_widgets_frame, text="עד שעה").pack(side=tk.RIGHT, padx=(5,0))
        self.to_hour_entry = tk.Entry(date_hour_widgets_frame, width=10, justify='right')
        self.to_hour_entry.pack(side=tk.RIGHT, padx=2)


        current_row += 1 

        # Notes - centered within inner_content_frame
        notes_frame = tk.Frame(inner_content_frame)
        notes_frame.grid(row=current_row, column=0, columnspan=3, pady=2) 
        
        notes_frame.grid_columnconfigure(0, weight=1)
        notes_frame.grid_columnconfigure(1, weight=0)
        notes_frame.grid_columnconfigure(2, weight=1)

        tk.Label(notes_frame, text="הערות").grid(row=0, column=1, sticky='s')
        self.notes_entry = tk.Text(notes_frame, height=2, width=20, relief="solid", borderwidth=1)
        self.notes_entry.grid(row=1, column=1, sticky='ew')
        # Apply right justification for the notes Text widget
        self.notes_entry.tag_configure("right", justify='right')
        self.notes_entry.tag_add("right", "1.0", "end")

        current_row += 1 

        # Buttons - centered within inner_content_frame
        button_frame = tk.Frame(inner_content_frame)
        button_frame.grid(row=current_row, column=0, columnspan=3, pady=10)
        tk.Button(button_frame, text="שמור וצור Word&Excel", command=self.save_task).pack()
        tk.Button(button_frame, text="התנתק", command=self.logout).pack(pady=5)

    def login(self):
        """מבצע התחברות למערכת עם בדיקת שם משתמש וסיסמה."""
        username = self.username_entry.get()
        password = self.password_entry.get().encode('utf-8')
        try:
            with open(USERS_FILE, 'r', encoding='utf-8') as f:
                users = json.load(f)["users"]
            for user in users:
                if user["username"] == username and bcrypt.checkpw(password, user["password"].encode('utf-8')):
                    self.current_user = user
                    self._setup_main_ui()
                    return
            messagebox.showerror("שגיאה", "שם משתמש או סיסמה שגויים", parent=self.root)
        except FileNotFoundError:
            messagebox.showerror("שגיאה", f"קובץ משתמשים ({USERS_FILE}) לא נמצא. אנא ודא שהקובץ קיים.", parent=self.root)
        except json.JSONDecodeError:
            messagebox.showerror("שגיאה", f"קובץ משתמשים ({USERS_FILE}) פגום. אנא מחק אותו והפעל מחדש.", parent=self.root)
        except Exception as e:
            messagebox.showerror("שגיאה", f"שגיאה בהתחברות: {str(e)}", parent=self.root)

    def logout(self):
        """מבצע התנתקות ומנקה את הממשק."""
        self.current_user = None
        self.clear_form()
        self._setup_login_ui()

    def clear_form(self):
        """מנקה את כל שדות הקלט."""
        for label, widget in self.entries.items():
            if isinstance(widget, tk.Entry) or isinstance(widget, ttk.Combobox):
                widget.delete(0, tk.END)
            elif isinstance(widget, tk.Text):
                widget.delete("1.0", tk.END)
        self.from_hour_entry.delete(0, tk.END)
        self.to_hour_entry.delete(0, tk.END)
        self.from_date_entry.delete(0, tk.END)
        self.to_date_entry.delete(0, tk.END)
        self.notes_entry.delete("1.0", tk.END)

    def show_admin_panel(self):
        """מציג חלון ניהול משתמשים למנהלים."""
        if self.current_user and self.current_user["role"] == "admin":
            self.admin_window = tk.Toplevel(self.root)
            self.admin_window.title("ניהול משתמשים")
            
            admin_width = 350
            admin_height = 400
            
            # וודא שחלון האב מעודכן לפני המרכוז היחסי (ובתוך הפונקציה window.update_idletasks())
            self._center_window_relative_to_parent(self.admin_window, admin_width, admin_height)
            self.admin_window.geometry(f"{admin_width}x{admin_height}") # הגדר את גודל החלון לאחר המרכוז
            self.admin_window.resizable(False, False)


            menu_bar = tk.Menu(self.admin_window)
            self.admin_window.config(menu=menu_bar)
            users_menu = tk.Menu(menu_bar, tearoff=0)
            menu_bar.add_cascade(label="ניהול משתמשים", menu=users_menu)
            users_menu.add_command(label="הצג משתמשים פעילים", command=self.show_active_users)

            tk.Label(self.admin_window, text="ניהול משתמשים", font=("Arial", 14)).pack(pady=10)

            try:
                with open(USERS_FILE, 'r', encoding='utf-8') as f:
                    users_data = json.load(f)
                    usernames = [f"{user['username']} ({user['role']})" for user in users_data["users"]]
            except Exception as e:
                messagebox.showerror("שגיאה", f"שגיאה בטעינת משתמשים: {str(e)}", parent=self.admin_window)
                return

            tk.Label(self.admin_window, text="בחר משתמש").pack(pady=5)
            self.user_combo = ttk.Combobox(self.admin_window, values=usernames)
            self.user_combo.pack()
            self.user_combo.bind("<<ComboboxSelected>>", self.on_user_select)

            self.new_username_entry_label = tk.Label(self.admin_window, text="שם משתמש חדש")
            self.new_username_entry_label.pack(pady=5)
            self.new_username_entry = tk.Entry(self.admin_window)
            self.new_username_entry.pack()
            self.new_password_entry_label = tk.Label(self.admin_window, text="סיסמה חדשה")
            self.new_password_entry_label.pack(pady=5)
            self.new_password_entry = tk.Entry(self.admin_window, show="*")
            self.new_password_entry.pack()
            self.new_role_combo_label = tk.Label(self.admin_window, text="תפקיד")
            self.new_role_combo_label.pack(pady=5)
            self.new_role_combo = ttk.Combobox(self.admin_window, values=["admin", "user"])
            self.new_role_combo.pack()

            self.create_user_button = tk.Button(self.admin_window, text="צור משתמש", command=self.create_user)
            self.create_user_button.pack(pady=5)
            self.remove_button = tk.Button(self.admin_window, text="הסר משתמש", command=self.remove_user, state='disabled')
            self.update_button = tk.Button(self.admin_window, text="עדכן סיסמה", command=self.show_update_password_window, state='disabled')
            self.make_admin_button = tk.Button(self.admin_window, text="תן אדמין", command=self.make_admin, state='disabled')
            self.remove_admin_button = tk.Button(self.admin_window, text="הסר אדמין", command=self.remove_admin, state='disabled')

            self.admin_window.transient(self.root)
            self.admin_window.grab_set()
        else:
            messagebox.showerror("שגיאה", "רק אדמין יכול לגשת לניהול משתמשים", parent=self.root)

    def show_active_users(self):
        """מציג רשימת משתמשים פעילים."""
        active_window = tk.Toplevel(self.root)
        active_window.title("משתמשים פעילים")
        
        active_width = 300
        active_height = 200
        self._center_window_relative_to_parent(active_window, active_width, active_height)
        active_window.geometry(f"{active_width}x{active_height}") # הגדר את גודל החלון לאחר המרכוז

        try:
            with open(USERS_FILE, 'r', encoding='utf-8') as f:
                users_data = json.load(f)
                users = [f"{user['username']} ({user['role']})" for user in users_data["users"]]
                active_count = len(users)
                for i, user in enumerate(users, 1):
                    tk.Label(active_window, text=f"{i}. {user}", justify="right").pack()
                tk.Label(active_window, text=f"סה\"כ משתמשים פעילים: {active_count}", justify="right").pack()
        except Exception as e:
            messagebox.showerror("שגיאה", f"שגיאה בטעינת משתמשים: {str(e)}", parent=active_window)

        active_window.transient(self.root)
        active_window.grab_set()

    def on_user_select(self, event):
        """מטפל בבחירת משתמש בחלון ניהול משתמשים."""
        if self.user_combo.get():
            self.create_user_button.pack_forget()
            self.remove_button.config(state='normal')
            self.update_button.config(state='normal')
            self.make_admin_button.config(state='normal')
            self.remove_admin_button.config(state='disabled')
            self.remove_button.pack(pady=5)
            self.update_button.pack(pady=5)
            self.make_admin_button.pack(pady=5)
            selected_username = self.user_combo.get().split(' (')[0]
            try:
                with open(USERS_FILE, 'r', encoding='utf-8') as f:
                    users_data = json.load(f)
                    for user in users_data["users"]:
                        if user["username"] == selected_username and user["role"] == "admin" and user["username"] != self.current_user["username"]:
                            self.remove_admin_button.config(state='normal')
                            self.remove_admin_button.pack(pady=5)
                            break
            except Exception as e:
                messagebox.showerror("שגיאה", f"שגיאה בטעינת משתמשים: {str(e)}", parent=self.admin_window)
            self.new_username_entry_label.pack_forget()
            self.new_username_entry.pack_forget()
            self.new_password_entry_label.pack_forget()
            self.new_password_entry.pack_forget()
            self.new_role_combo_label.pack_forget()
            self.new_role_combo.pack_forget()
        else:
            self.remove_button.pack_forget()
            self.update_button.pack_forget()
            self.make_admin_button.pack_forget()
            self.remove_admin_button.pack_forget()
            self.remove_button.config(state='disabled')
            self.update_button.config(state='disabled')
            self.make_admin_button.config(state='disabled')
            self.remove_admin_button.config(state='disabled')
            self.new_username_entry_label.pack(pady=5)
            self.new_username_entry.pack()
            self.new_password_entry_label.pack(pady=5)
            self.new_password_entry.pack()
            self.new_role_combo_label.pack(pady=5)
            self.new_role_combo.pack()
            self.create_user_button.pack(pady=5)

    def show_update_password_window(self):
        """מציג חלון לעדכון סיסמה."""
        username = self.user_combo.get().split(' (')[0]
        update_window = tk.Toplevel(self.root)
        update_window.title("עדכן סיסמה")
        
        update_width = 250
        update_height = 150
        self._center_window_relative_to_parent(update_window, update_width, update_height)
        update_window.geometry(f"{update_width}x{update_height}") # הגדר את גודל החלון לאחר המרכוז

        tk.Label(update_window, text="סיסמה חדשה").pack(pady=10)
        new_password_entry = tk.Entry(update_window, show="*", justify='right')
        new_password_entry.pack(pady=5)

        def save_new_password():
            new_password = new_password_entry.get()
            if not new_password:
                messagebox.showerror("שגיאה", "הזן סיסמה חדשה", parent=update_window)
                return
            try:
                hashed = bcrypt.hashpw(new_password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
                with open(USERS_FILE, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                for user in data["users"]:
                    if user["username"] == username:
                        user["password"] = hashed
                        break
                with open(USERS_FILE, 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
                messagebox.showinfo("הצלחה", f"סיסמת המשתמש {username} עודכנה בהצלחה!", parent=update_window)
                update_window.destroy()
                self.user_combo.set("")
                self.on_user_select(None)
            except Exception as e:
                messagebox.showerror("שגיאה", f"שגיאה בעדכון סיסמה: {str(e)}", parent=self.admin_window)

        tk.Button(update_window, text="שמור", command=save_new_password).pack(pady=10)
        tk.Button(update_window, text="ביטול", command=update_window.destroy).pack(pady=5)
        update_window.transient(self.root)
        update_window.grab_set()

    def show_calendar(self, entry):
        """מציג לוח שנה לבחירת תאריך."""
        cal_window = tk.Toplevel(self.root)
        cal_window.title("בחר תאריך")
        
        cal_width = 300
        cal_height = 250
        self._center_window_relative_to_parent(cal_window, cal_width, cal_height)
        cal_window.geometry(f"{cal_width}x{cal_height}") # הגדר את גודל החלון לאחר המרכוז

        cal = Calendar(cal_window, selectmode="day", date_pattern="dd/mm/yyyy")
        cal.pack(pady=10)

        def set_date():
            date = cal.get_date()
            entry.delete(0, tk.END)
            entry.insert(0, date)
            cal_window.destroy()

        tk.Button(cal_window, text="אישור", command=set_date).pack(pady=5)
        tk.Button(cal_window, text="ביטול", command=cal_window.destroy).pack(pady=5)
        cal_window.transient(self.root)
        cal_window.grab_set()

    def create_user(self):
        """יוצר משתמש חדש עם שם, סיסמה ותפקיד."""
        username = self.new_username_entry.get()
        password = self.new_password_entry.get()
        role = self.new_role_combo.get()
        if not (username and password and role):
            messagebox.showerror("שגיאה", "מלא את כל השדות", parent=self.admin_window)
            return
        try:
            with open(USERS_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
            if any(u["username"] == username for u in data["users"]):
                messagebox.showerror("שגיאה", "שם משתמש כבר קיים", parent=self.admin_window)
                return
            hashed = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
            data["users"].append({"username": username, "password": hashed, "role": role})
            with open(USERS_FILE, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            messagebox.showinfo("הצלחה", "משתמש נוצר!", parent=self.admin_window)
            self.new_username_entry.delete(0, tk.END)
            self.new_password_entry.delete(0, tk.END)
            self.new_role_combo.set("")
            self.admin_window.focus_set()
        except Exception as e:
            messagebox.showerror("שגיאה", f"שגיאה ביצירת משתמש: {str(e)}", parent=self.admin_window)

    def remove_user(self):
        """מסיר משתמש קיים."""
        username = self.user_combo.get().split(' (')[0]
        if not username:
            messagebox.showerror("שגיאה", "בחר משתמש להסרה", parent=self.admin_window)
            return
        if username == self.current_user["username"]:
            messagebox.showerror("שגיאה", "לא ניתן להסיר את המשתמש המחובר", parent=self.admin_window)
            return
        if username == "admin":
            messagebox.showerror("שגיאה", "לא ניתן להסיר את המשתמש admin", parent=self.admin_window)
            return
        try:
            with open(USERS_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
            data["users"] = [user for user in data["users"] if user["username"] != username]
            with open(USERS_FILE, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            messagebox.showinfo("הצלחה", f"המשתמש {username} הוסר בהצלחה!", parent=self.admin_window)
            self.user_combo.set("")
            self.on_user_select(None)
        except Exception as e:
            messagebox.showerror("שגיאה", f"שגיאה בהסרת משתמש: {str(e)}", parent=self.admin_window)

    def make_admin(self):
        """מעניק הרשאות אדמין למשתמש."""
        username = self.user_combo.get().split(' (')[0]
        if not username:
            messagebox.showerror("שגיאה", "בחר משתמש לשינוי תפקיד", parent=self.admin_window)
            return
        if username == self.current_user["username"]:
            messagebox.showerror("שגיאה", "לא ניתן לשנות את התפקיד של המשתמש המחובר", parent=self.admin_window)
            return
        try:
            with open(USERS_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
            for user in data["users"]:
                if user["username"] == username and user["role"] != "admin":
                    user["role"] = "admin"
                    break
            with open(USERS_FILE, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            messagebox.showinfo("הצלחה", f"המשתמש {username} הפך ל-admin בהצלחה!", parent=self.admin_window)
            self.user_combo.set("")
            self.on_user_select(None)
        except Exception as e:
            messagebox.showerror("שגיאה", f"שגיאה בהפיכת משתמש לאדמין: {str(e)}", parent=self.admin_window)

    def remove_admin(self):
        """מסיר הרשאות אדמין ממשתמש."""
        username = self.user_combo.get().split(' (')[0]
        if not username:
            messagebox.showerror("שגיאה", "בחר משתמש לשינוי תפקיד", parent=self.admin_window)
            return
        if username == self.current_user["username"]:
            messagebox.showerror("שגיאה", "לא ניתן לשנות את התפקיד של המשתמש המחובר", parent=self.admin_window)
            return
        try:
            with open(USERS_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
            for user in data["users"]:
                if user["username"] == username and user["role"] == "admin":
                    user["role"] = "user"
                    break
            with open(USERS_FILE, 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
            messagebox.showinfo("הצלחה", f"המשתמש {username} הוסר מתפקיד admin בהצלחה!", parent=self.admin_window)
            self.user_combo.set("")
            self.on_user_select(None)
        except Exception as e:
            messagebox.showerror("שגיאה", f"שגיאה בהסרת הרשאות אדמין: {str(e)}", parent=self.admin_window)

    def validate_dates(self, from_date, to_date):
        """בודק תקינות תאריכים."""
        if not from_date or not to_date:
            return True
        try:
            start = datetime.strptime(from_date, "%d/%m/%Y")
            end = datetime.strptime(to_date, "%d/%m/%Y")
            if start > end:
                raise ValueError("תאריך התחלה לא יכול להיות מאוחר מתאריך סיום")
            return True
        except ValueError as e:
            messagebox.showerror("שגיאה", f"תאריכים לא תקינים: {str(e)}", parent=self.root)
            return False

    def validate_hours(self, from_hour, to_hour):
        """בודק תקינות שעות."""
        if not from_hour or not to_hour:
            return True
        try:
            start_hour = datetime.strptime(from_hour, "%H:%M")
            end_hour = datetime.strptime(to_hour, "%H:%M")
            if start_hour > end_hour:
                raise ValueError("שעת התחלה לא יכולה להיות מאוחרת משעת סיום")
            return True
        except ValueError as e:
            messagebox.showerror("שגיאה", f"שעות לא תקינות: {str(e)}", parent=self.root)
            return False

    def get_next_doc_path(self, folder):
        """מחזיר נתיב פנוי לקובץ Word."""
        base_path = os.path.join(folder, "detailed_report")
        i = 1
        while True:
            path = f"{base_path}{'' if i == 1 else i}.docx"
            if not os.path.exists(path):
                return path
            i += 1

    def save_task(self):
        """שומר משימה ל-JSON, Excel ו-Word."""
        if not self.current_user:
            messagebox.showerror("שגיאה", "התחבר תחילה", parent=self.root)
            return

        task = {
            "מס' תיק": self.entries["מס' תיק"].get(),
            "שם היחידה": self.entries["שם היחידה"].get(),
            "שם השוטר המעתיק": self.entries["שם השוטר המעתיק"].get(),
            "מ.א": self.entries["מ.א"].get(),
            "פרטי אירוע": self.entries["פרטי אירוע"].get("1.0", tk.END).strip(),
            "מיקום המצלמה": self.entries["מיקום המצלמה"].get(),
            "מספר מצלמות": self.entries["מספר מצלמות"].get(),
            "משעה": self.from_hour_entry.get(),
            "עד שעה": self.to_hour_entry.get(),
            "מתאריך": self.from_date_entry.get(),
            "עד תאריך": self.to_date_entry.get(),
            "הערות": self.notes_entry.get("1.0", tk.END).strip(),
            "משתמש": self.current_user["username"]
        }

        required_fields = ["מס' תיק", "שם היחידה", "שם השוטר המעתיק", "מ.א", "פרטי אירוע", "מיקום המצלמה", "מספר מצלמות", "משעה", "עד שעה"]
        missing = [field for field in required_fields if not task[field]]
        if missing:
            messagebox.showerror("שגיאה", f"מלא את השדות החובה: {', '.join(missing)}", parent=self.root)
            return

        if not self.validate_dates(task["מתאריך"], task["עד תאריך"]):
            return
        if not self.validate_hours(task["משעה"], task["עד שעה"]):
            return

        try:
            # Save to JSON
            try:
                with open(TASKS_FILE, 'r', encoding='utf-8') as f:
                    tasks = json.load(f)
            except (json.JSONDecodeError, FileNotFoundError):
                tasks = []
            tasks.append(task)
            with open(TASKS_FILE, 'w', encoding='utf-8') as f:
                json.dump(tasks, f, ensure_ascii=False, indent=2)

            # Save to Excel
            if os.path.exists("tasks.xlsx"):
                os.chmod("tasks.xlsx", stat.S_IWRITE)
            df_tasks = pd.DataFrame(tasks)
            df_tasks = df_tasks.drop(columns=["טל'", "זמן התחלה", "זמן סיום", "מספר סידורי"], errors='ignore')
            with pd.ExcelWriter("tasks.xlsx", engine="openpyxl") as writer:
                df_tasks.to_excel(writer, index=False)
                for column_cells in writer.sheets['Sheet1'].columns:
                    length = max(len(str(cell.value)) for cell in column_cells)
                    writer.sheets['Sheet1'].column_dimensions[column_cells[0].column_letter].width = length + 2
            os.chmod("tasks.xlsx", stat.S_IREAD)

            # Save to Word
            officer_folder = task["שם השוטר המעתיק"]
            if not os.path.exists(officer_folder):
                os.makedirs(officer_folder)
            doc_path = self.get_next_doc_path(officer_folder)
            doc = Document()

            table = doc.add_table(rows=1, cols=8)  # Updated to 8 columns to include "מתאריך"
            table.autofit = True
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            headers = ["הערות", "מס' תיק", "מיקום המצלמה", "מספר מצלמות", "משעה", "עד שעה", "מתאריך", "עד תאריך"]
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                run = hdr_cells[i].paragraphs[0].runs[0] if hdr_cells[i].paragraphs[0].runs else hdr_cells[i].paragraphs[0].add_run()
                run.bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_cells = table.add_row().cells
            row_cells[0].text = task["הערות"] if task["הערות"] else "ללא הערות"
            row_cells[1].text = task["מס' תיק"]
            row_cells[2].text = task["מיקום המצלמה"]
            row_cells[3].text = task["מספר מצלמות"]
            row_cells[4].text = task["משעה"]
            row_cells[5].text = task["עד שעה"]
            row_cells[6].text = task["מתאריך"] if task["מתאריך"] else ""  # Added "מתאריך" column
            row_cells[7].text = task["עד תאריך"] if task["עד תאריך"] else ""
            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            doc.add_paragraph().add_run('\n\n')
            confidentiality_paragraph = doc.add_paragraph()
            confidentiality_paragraph.add_run("הצהרת סודיות ושימוש בחומר מצולם").bold = True
            confidentiality_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            doc.add_paragraph("הנני מצהיר כי ידוע לי באופן חד משמעי, כי חל איסור חמור על העברת הצילומים או כל חלק מהם לגורמים שאינם קשורים לארגון שהציב את המצלמות, חל איסור על מסירת חומר מצולם לגורמי תקשורת מכל סוג.").alignment = WD_ALIGN_PARAGRAPH.RIGHT

            doc.add_paragraph().add_run('\n\n')
            security_paragraph = doc.add_paragraph()
            security_paragraph.add_run("הנחיות אבטחת מידע").bold = True
            security_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            instructions = [
                "יש לשמור את ההעתקה אך ורק על התקן אחסון חיצוני.",
                "חל איסור מוחלט על שמירה מקומית במחשב העמדה.",
                "חל איסור מוחלט לצלם מסך.",
                "כל חומר שנשמר באופן זמני בעמדה, ימחק באופן סופי תוך 48 שעות (באחריות השוטר)."
            ]
            for text in instructions:
                p = doc.add_paragraph(text)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            signature_paragraph = doc.add_paragraph()
            signature_paragraph.add_run(f"חתימת השוטר המעתיק: {task['שם השוטר המעתיק']} מ.א: {task['מ.א']}")
            signature_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            doc.save(doc_path)

            messagebox.showinfo("הצלחה", "נשמרו קבצי Excel ודו\"ח Word!", parent=self.root)
        except Exception as e:
            messagebox.showerror("שגיאה", f"שגיאה בשמירת המשימה: {str(e)}", parent=self.root)
            
    def _center_window_on_screen(self, window, width, height):
        """מרכז חלון Tkinter במרכז המסך."""
        window.update_idletasks() # וודא שמימדי החלון מעודכנים
        screen_width = window.winfo_screenwidth()
        screen_height = window.winfo_screenheight()
        x = (screen_width / 2) - (width / 2)
        y = (screen_height / 2) - (height / 2)
        window.geometry(f'{width}x{height}+{int(x)}+{int(y)}')

    def _center_window_relative_to_parent(self, window, width, height):
        """מרכז חלון Tkinter יחסית לחלון האב (החלון הראשי)."""
        parent = window.master
        parent.update_idletasks() # וודא שמימדי חלון האב מעודכנים
        window.update_idletasks() # **חשוב!** וודא שמימדי חלון הילד מעודכנים לפני החישוב
        
        parent_x = parent.winfo_x()
        parent_y = parent.winfo_y()
        parent_width = parent.winfo_width()
        parent_height = parent.winfo_height()

        center_x = parent_x + (parent_width - width) // 2
        center_y = parent_y + (parent_height - height) // 2
        window.geometry(f"{width}x{height}+{center_x}+{center_y}")


if __name__ == "__main__":
    root = tk.Tk()
    app = TaskManagerApp(root)
    root.mainloop()