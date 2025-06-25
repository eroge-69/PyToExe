from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openpyxl import load_workbook
#import docx2pdf  # optionally used
import win32com.client
import os
import re
import random
import sys

# === Resource loader (for PyInstaller compatibility) ===
def resource_path(relative_path):
    base_path = getattr(sys, '_MEIPASS', os.path.abspath("."))
    return os.path.join(base_path, relative_path)

# === Create unique filenames ===
def get_unique_filename(base_name, extension):
    filename = f"{base_name}.{extension}"
    counter = 1
    while os.path.exists(filename):
        filename = f"{base_name}_{counter}.{extension}"
        counter += 1
    return filename

# === Replace placeholders in runs (for both paragraphs and table cells) ===
def replace_placeholders_in_paragraph(paragraph, replacements):
    full_text = "".join(run.text for run in paragraph.runs)
    replaced = full_text
    for key, val in replacements.items():
        replaced = replaced.replace(key, str(val))

    if replaced != full_text:
        # Clear all runs and replace with one updated run
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = replaced

# === Paths ===
excel_path = resource_path("Macro_Exam Generator.xlsm")
word_template_path = resource_path("Exam Template.docx")
section_b_image_dir = resource_path("Section B Images")
section_c_image_dir = resource_path("Section C Images")

# === Step 1: Run Excel macro to generate new random values ===
excel_app = win32com.client.Dispatch("Excel.Application")
excel_app.Visible = False
workbook = excel_app.Workbooks.Open(os.path.abspath("Macro_Exam Generator.xlsm"))
excel_app.Application.Run("GenerateVCEPracticeExam")
workbook.Save()
workbook.Close()
excel_app.Quit()

# === Step 2: Load updated Excel workbook ===
wb = load_workbook(excel_path, data_only=True)
ws = wb["Dashboard"]

# === Step 3: Prepare file output ===
output_docx_path = get_unique_filename("Generated_Exam", "docx")
output_pdf_path = output_docx_path.replace(".docx", ".pdf")

# === Step 4: Read values from Dashboard ===
replacements = {
    "{A7}": ws["A7"].value,
    "{A8}": ws["A8"].value,
    "{A10}": ws["A10"].value,
    "{A11}": ws["A11"].value,
    "{G9}": ws["G9"].value,
    "{G11}": ws["G11"].value,
    "{G15}": ws["G15"].value,
    "{G13}": "[IMAGE_B]",
    "{N13}": "[IMAGE_1]",
    "{N16}": "[IMAGE_2]",
    "{N17}": "[IMAGE_3]"
}

# === Step 5: Load Word template ===
doc = Document(word_template_path)

# === Step 6: Replace placeholders in text and tables ===
for paragraph in doc.paragraphs:
    replace_placeholders_in_paragraph(paragraph, replacements)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_placeholders_in_paragraph(paragraph, replacements)

# === Step 7: Insert Section B Image ===
section_b_prompt_number = random.randint(1, 11)  # Adjust range if needed
section_b_image_path = os.path.join(section_b_image_dir, f"{section_b_prompt_number}.png")

for paragraph in doc.paragraphs:
    if "[IMAGE_B]" in paragraph.text:
        paragraph.clear()
        run = paragraph.add_run()
        if os.path.exists(section_b_image_path):
            run.add_picture(section_b_image_path, width=Inches(4.5))
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            paragraph.add_run("[IMAGE NOT FOUND]")

# === Step 8: Insert Section C Images ===
section_c_prompt = random.randint(1, 11)

if section_c_prompt < 10:
    img1 = os.path.join(section_c_image_dir, f"{section_c_prompt}1.png")
    img2 = os.path.join(section_c_image_dir, f"{section_c_prompt}2.png")
    img3 = os.path.join(section_c_image_dir, f"{section_c_prompt}3.png")
else:
    img1 = os.path.join(section_c_image_dir, f"{section_c_prompt}-1.png")
    img2 = os.path.join(section_c_image_dir, f"{section_c_prompt}-2.png")
    img3 = os.path.join(section_c_image_dir, f"{section_c_prompt}-3.png")

def insert_image_at_placeholder(doc, placeholder, image_path):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.clear()
            if os.path.exists(image_path):
                run = paragraph.add_run()
                run.add_picture(image_path, width=Inches(5.5))
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                paragraph.add_run("[IMAGE NOT FOUND]")
            break

insert_image_at_placeholder(doc, "[IMAGE_1]", img1)
insert_image_at_placeholder(doc, "[IMAGE_2]", img2)
insert_image_at_placeholder(doc, "[IMAGE_3]", img3)

# === Step 9: Save and convert ===
doc.save(output_docx_path)
# convert(output_docx_path, output_pdf_path)  # Optional, depends on Word

print(f"✅ Exam generated: {output_docx_path} → {output_pdf_path}")
