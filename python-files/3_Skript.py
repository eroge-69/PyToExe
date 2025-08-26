import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import re
import os

# Deine bestehenden Funktionen
def lade_platzhalter(pfad):
    mapping = {}
    with open(pfad, encoding="utf-16") as f:
        for zeile in f:
            match = re.match(r"\[\[(.*?)\]\]\s+(.+)", zeile.strip())
            if match:
                key, value = match.groups()
                mapping[f"[[{key}]]"] = value
    return mapping

def ersetze_runs(paragraph, mapping):
    full_text = ''.join(run.text for run in paragraph.runs)
    pattern = r"\[\[.*?\]\]"
    if not re.search(pattern, full_text):
        return
    for placeholder, value in mapping.items():
        full_text = full_text.replace(placeholder, value)
    paragraph.clear()
    paragraph.add_run(full_text)

def entferne_leere_platzhalter(doc):
    absätze_zum_entfernen = []
    geloeschte_platzhalter = []
    for para in doc.paragraphs:
        full_text = para.text.strip()
        match = re.fullmatch(r"[-•*]?\s*(\[\[.*?\]\])", full_text)
        if match:
            absätze_zum_entfernen.append(para)
            geloeschte_platzhalter.append(match.group(1))
    for para in absätze_zum_entfernen:
        p = para._element
        p.getparent().remove(p)
        para._p = para._element = None
    return geloeschte_platzhalter

# --- UI Funktionen ---
def waehle_datei():
    datei_path = filedialog.askopenfilename(
        title="Wähle die Word-Datei aus",
        filetypes=[("Word-Dateien", "*.docx")]
    )
    if datei_path:
        datei_label.config(text=datei_path)
        start_button.config(state="normal")

def starte_verarbeitung():
    datei_path = datei_label.cget("text")
    if not datei_path or not os.path.exists(datei_path):
        messagebox.showerror("Fehler", "Keine Datei ausgewählt oder Datei existiert nicht.")
        return

    mapping = lade_platzhalter("2_master.txt")
    doc = Document(datei_path)

    for para in doc.paragraphs:
        ersetze_runs(para, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    ersetze_runs(para, mapping)

    geloeschte_platzhalter = entferne_leere_platzhalter(doc)

    # Speicherpfad automatisch generieren
    ordner, dateiname = os.path.split(datei_path)
    name, ext = os.path.splitext(dateiname)
    output_path = os.path.join(ordner, f"{name}_modifiziert{ext}")
    doc.save(output_path)

    message = f"Dokument gespeichert:\n{output_path}"
    if geloeschte_platzhalter:
        message += f"\n{len(geloeschte_platzhalter)} leere Platzhalter entfernt."
    messagebox.showinfo("Fertig", message)

# --- Tkinter UI ---
root = tk.Tk()
root.title("Platzhalter ersetzen")

tk.Label(root, text="1. Word-Datei auswählen:").pack(pady=5)
datei_label = tk.Label(root, text="", fg="blue")
datei_label.pack(pady=5)
tk.Button(root, text="Datei auswählen", command=waehle_datei).pack(pady=5)

tk.Label(root, text="2. Verarbeitung starten:").pack(pady=10)
start_button = tk.Button(root, text="Start", command=starte_verarbeitung, state="disabled")
start_button.pack(pady=5)

root.mainloop()
