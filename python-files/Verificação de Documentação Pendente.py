import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from docx import Document
from PIL import Image, ImageTk
import os
import subprocess
import platform

# ----------------------
# Documentos por seções
# ----------------------
docs_pessoais = [
    "CPF – Autenticado em Cartório",
    "RG – Autenticado em Cartório",
    "CPF (Cônjuge) – Autenticado em Cartório",
    "RG (Cônjuge) – Autenticado em Cartório",
    "Certidão de Nascimento ou Casamento (ATUALIZADA)",  # checkbox interativa
    "Comprovante de Endereço (ATUALIZADO)",
    "Comprovante de Rendimentos Mensais",
    "Petição Inicial – Reconhecido Firma em Cartório",
    "Cadastro Socioeconômico"
]

docs_imovel = [
    "Comprovante de Endereço (ANTERIOR À 22/12/2016)",
    "Declaração/Comprovante de Posse – Reconhecido Firma",
    "Declaração de Reconhecimento de Limite – Reconhecido Firma",
    "Matrícula do Imóvel ou Certidão Negativa de Matrícula"
]

docs_situacao = [
    "Cadastro Imobiliário – CIB",
    "Certidão Negativa de Débitos Municipais (IPTU)",
    "Cadastro Ambiental Rural – CAR",
    "CCIR – Certificado de Cadastro de Imóvel Rural",
    "Certidão de Regularidade Fiscal (ITR)"
]

docs_engenheiro = [
    "Estudo de Georreferenciamento",
    "Memorial Descritivo do Georeferenciamento",
    "ART - Anotação de Responsabilidade Técnica"
]

todos_documentos = docs_pessoais + docs_imovel + docs_situacao + docs_engenheiro

# ----------------------
# Variáveis de estado
# ----------------------
status_docs = {}
opcoes_docs = {}
checkbox_widgets = {}

# Bloqueios por tipo de imóvel
docs_urbano_block = ["Cadastro Ambiental Rural – CAR", "CCIR – Certificado de Cadastro de Imóvel Rural", "Certidão de Regularidade Fiscal (ITR)"]
docs_rural_block = ["Cadastro Imobiliário – CIB", "Certidão Negativa de Débitos Municipais (IPTU)"]

# Bloqueio por estado civil
docs_conjuge = ["CPF (Cônjuge) – Autenticado em Cartório", "RG (Cônjuge) – Autenticado em Cartório"]

# Nome da checkbox interativa
certidao_doc_nome = "Certidão de Nascimento ou Casamento (ATUALIZADA)"

# ----------------------
# Funções
# ----------------------
def atualizar_tipo_imovel(*args):
    tipo = tipo_imovel_var.get()
    for doc in docs_urbano_block + docs_rural_block:
        checkbox_widgets[doc].config(state="normal")
        if doc in opcoes_docs and hasattr(opcoes_docs[doc], "config"):
            opcoes_docs[doc].config(state="readonly")
    if tipo == "Urbano":
        for doc in docs_urbano_block:
            checkbox_widgets[doc].config(state="disabled")
            status_docs[doc].set(0)
            if doc in opcoes_docs and hasattr(opcoes_docs[doc], "config"):
                opcoes_docs[doc].config(state="disabled")
    elif tipo == "Rural":
        for doc in docs_rural_block:
            checkbox_widgets[doc].config(state="disabled")
            status_docs[doc].set(0)
            if doc in opcoes_docs and hasattr(opcoes_docs[doc], "config"):
                opcoes_docs[doc].config(state="disabled")

def atualizar_estado_civil(*args):
    estado = estado_civil_var.get()
    if estado not in ["Casado(a)", "União Estável"]:
        for doc in docs_conjuge:
            checkbox_widgets[doc].config(state="disabled")
            status_docs[doc].set(0)
            if doc in opcoes_docs and hasattr(opcoes_docs[doc], "config"):
                opcoes_docs[doc].config(state="disabled")
    else:
        for doc in docs_conjuge:
            checkbox_widgets[doc].config(state="normal")
            if doc in opcoes_docs and hasattr(opcoes_docs[doc], "config"):
                opcoes_docs[doc].config(state="readonly")
    # Atualiza a checkbox da certidão
    if certidao_doc_nome in checkbox_widgets:
        if estado == "Solteiro(a)":
            checkbox_widgets[certidao_doc_nome].config(text="Certidão de Nascimento Atualizada")
        elif estado == "Casado(a)":
            checkbox_widgets[certidao_doc_nome].config(text="Certidão de Casamento Atualizada")
        elif estado == "Divorciado(a)":
            checkbox_widgets[certidao_doc_nome].config(text="Certidão de Divórcio Atualizada")
        elif estado == "Viúvo(a)":
            checkbox_widgets[certidao_doc_nome].config(text="Certidão de Casamento + Certidão de Óbito do Cônjuge")
        elif estado == "União Estável":
            checkbox_widgets[certidao_doc_nome].config(text="Certidão de União Estável Atualizada")

def criar_bloco(frame_pai, titulo, lista_docs):
    bloco = tk.LabelFrame(frame_pai, text=titulo, font=("Arial", 12, "bold"), padx=10, pady=10)
    bloco.pack(side="left", padx=10, pady=10, fill="both", expand=True)
    
    for doc in lista_docs:
        row = tk.Frame(bloco)
        row.pack(anchor="w", pady=2, fill="x")
        
        var_check = tk.IntVar()
        chk = tk.Checkbutton(row, text=doc, variable=var_check)
        chk.pack(side="left")
        status_docs[doc] = var_check
        checkbox_widgets[doc] = chk
        
        if "Autenticado" in doc or "Reconhecido" in doc:
            cb = ttk.Combobox(
                row,
                values=["Autenticado/Reconhecido", "Falta autenticar/Reconhecer firma", "Não se aplica"],
                state="readonly",
                width=25
            )
            cb.current(0)
            cb.pack(side="left", padx=5)
            opcoes_docs[doc] = cb
        else:
            opcoes_docs[doc] = "Não se aplica"
    
    return bloco

def escolher_pasta():
    pasta = filedialog.askdirectory()
    if pasta:
        output_dir.set(pasta)

def gerar_relatorio():
    requerente = entry_nome.get().strip()
    processo = entry_processo.get().strip()
    tipo_imovel = tipo_imovel_var.get()
    estado_civil = estado_civil_var.get()

    doc = Document()
    doc.add_heading("RELATÓRIO DE DOCUMENTOS - REURB", level=1)
    doc.add_paragraph(f"Requerente: {requerente if requerente else '---'}")
    doc.add_paragraph(f"Nº do Processo: {processo if processo else '---'}")
    doc.add_paragraph(f"Tipo de Imóvel: {tipo_imovel if tipo_imovel else '---'}")
    doc.add_paragraph(f"Estado Civil: {estado_civil if estado_civil else '---'}")
    doc.add_paragraph("")

    faltando = []
    pendentes_autenticacao = []

    for nome in todos_documentos:
        if tipo_imovel == "Urbano" and nome in docs_urbano_block:
            continue
        if tipo_imovel == "Rural" and nome in docs_rural_block:
            continue
        if estado_civil not in ["Casado(a)", "União Estável"] and nome in docs_conjuge:
            continue

        entregue = status_docs[nome].get()
        estado = opcoes_docs[nome].get() if hasattr(opcoes_docs[nome], "get") else opcoes_docs[nome]
        nome_relatorio = checkbox_widgets[nome].cget("text")
        if not entregue:
            faltando.append(nome_relatorio)
        elif estado == "Falta autenticar/Reconhecer firma":
            pendentes_autenticacao.append(nome_relatorio)

    if pendentes_autenticacao:
        doc.add_heading("Pendentes de Autenticação/Firma", level=2)
        tabela = doc.add_table(rows=1, cols=2)
        hdr_cells = tabela.rows[0].cells
        hdr_cells[0].text = "Documento"
        hdr_cells[1].text = "Situação"
        for pa in pendentes_autenticacao:
            row_cells = tabela.add_row().cells
            row_cells[0].text = pa
            row_cells[1].text = "⚠ Falta autenticar/reconhecer firma"

    if faltando:
        doc.add_heading("Documentos Faltantes", level=2)
        tabela = doc.add_table(rows=1, cols=2)
        hdr_cells = tabela.rows[0].cells
        hdr_cells[0].text = "Documento"
        hdr_cells[1].text = "Situação"
        for f in faltando:
            row_cells = tabela.add_row().cells
            row_cells[0].text = f
            row_cells[1].text = "✘ Não entregue"

    if not pendentes_autenticacao and not faltando:
        doc.add_paragraph("✔ Todos os documentos entregues e regularizados.")

    # salva arquivo
    nome_arquivo = f"{requerente.replace(' ','_')}_{processo.replace('/','_')}_Docs_Pendentes.docx"
    caminho_final = os.path.join(output_dir.get(), nome_arquivo) if output_dir.get() else nome_arquivo
    doc.save(caminho_final)

    # Abrir automaticamente no Word / visualizador
    try:
        sistema = platform.system()
        if sistema == "Windows":
            os.startfile(caminho_final)
        elif sistema == "Darwin":  # macOS
            subprocess.Popen(["open", caminho_final])
        else:  # Linux
            subprocess.Popen(["xdg-open", caminho_final])
    except Exception as e:
        print(f"Não foi possível abrir automaticamente o arquivo: {e}")

    # Janela de opções
    janela_opcoes = tk.Toplevel(root)
    janela_opcoes.title("Relatório gerado")
    janela_opcoes.geometry("400x220")
    tk.Label(
        janela_opcoes, 
        text=f"Relatório de documentos faltantes do requerente {requerente}\nProcesso nº {processo}\nfoi gerado.",
        justify="left"
    ).pack(pady=15, padx=15)

    def abrir_pasta():
        if output_dir.get():
            subprocess.Popen(["open", output_dir.get()])  # MacOS
        else:
            subprocess.Popen(["open", os.getcwd()])
        janela_opcoes.destroy()

    def imprimir_relatorio():
        subprocess.Popen(["open", caminho_final])
        janela_opcoes.destroy()

    tk.Button(janela_opcoes, text="Abrir Pasta", command=abrir_pasta, width=20).pack(pady=5)
    tk.Button(janela_opcoes, text="Imprimir", command=imprimir_relatorio, width=20).pack(pady=5)
    tk.Button(janela_opcoes, text="Fechar", command=janela_opcoes.destroy, width=20).pack(pady=5)

# ----------------------
# Janela principal
# ----------------------
root = tk.Tk()
root.title("SIVAD – Sistema Integrado de Verificação de Documentos")
root.geometry("1366x768")

# Variáveis
output_dir = tk.StringVar()
tipo_imovel_var = tk.StringVar()
estado_civil_var = tk.StringVar()

# ----------------------
# Campos topo
# ----------------------
frame_top = tk.Frame(root)
frame_top.pack(pady=10, padx=20, fill="x")

tk.Label(frame_top, text="Nome do Requerente:").grid(row=0, column=0, sticky="w")
entry_nome = tk.Entry(frame_top, width=50)
entry_nome.grid(row=0, column=1, padx=5)

tk.Label(frame_top, text="Nº do Processo:").grid(row=0, column=2, sticky="w")
entry_processo = tk.Entry(frame_top, width=30)
entry_processo.grid(row=0, column=3, sticky="w", padx=5)

# Tipo de Imóvel
tk.Label(frame_top, text="Tipo de Imóvel:").grid(row=1, column=0, sticky="w")
tipo_imovel_cb = ttk.Combobox(frame_top, textvariable=tipo_imovel_var,
                               values=["", "Urbano", "Rural"], state="readonly", width=20)
tipo_imovel_cb.grid(row=1, column=1, sticky="w")
tipo_imovel_var.trace("w", atualizar_tipo_imovel)

# Estado civil
tk.Label(frame_top, text="Estado Civil:").grid(row=1, column=2, sticky="w")
estado_civil_cb = ttk.Combobox(frame_top, textvariable=estado_civil_var,
                               values=["", "Solteiro(a)", "Casado(a)", "Divorciado(a)", "Viúvo(a)", "União Estável"],
                               state="readonly", width=20)
estado_civil_cb.grid(row=1, column=3, sticky="w")
estado_civil_var.trace("w", atualizar_estado_civil)

# ----------------------
# Blocos 2x2
# ----------------------
frame_blocos = tk.Frame(root)
frame_blocos.pack(padx=10, pady=10, fill="both", expand=True)

frame_blocos_top = tk.Frame(frame_blocos)
frame_blocos_top.pack(fill="both", expand=True)
bloco1 = criar_bloco(frame_blocos_top, "DOS DOCUMENTOS PESSOAIS", docs_pessoais)
bloco2 = criar_bloco(frame_blocos_top, "DOCUMENTOS DO IMÓVEL", docs_imovel)

frame_blocos_bottom = tk.Frame(frame_blocos)
frame_blocos_bottom.pack(fill="both", expand=True)
bloco3 = criar_bloco(frame_blocos_bottom, "SITUAÇÃO FISCAL DO IMÓVEL", docs_situacao)
bloco4 = criar_bloco(frame_blocos_bottom, "ATRIBUIÇÕES DO ENGENHEIRO", docs_engenheiro)

# ----------------------
# Botões
# ----------------------
frame_botoes = tk.Frame(root)
frame_botoes.pack(pady=5)
tk.Button(frame_botoes, text="Escolher pasta de saída", command=escolher_pasta, width=20).pack(side="left", padx=5)
tk.Button(frame_botoes, text="Gerar Relatório", command=gerar_relatorio, width=20).pack(side="left", padx=5)

# ----------------------
# Rodapé
# ----------------------
tk.Label(root, text="© SITARF v1.0 - Desenvolvido por Espedito Neiva de Sousa Lima Filho.",
         font=("Arial", 7), justify="left").pack(side="bottom", anchor="w", padx=5, pady=(0, 5))

# ----------------------
# Loop principal
# ----------------------
root.mainloop()
