import os
import xml.etree.ElementTree as ET
import requests
from docx import Document
from docx.shared import Inches, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import tkinter as tk
from tkinter import filedialog
from PIL import Image

IMAGE_COMPRESSION_QUALITY = 85
IMAGE_DOWNLOAD_FOLDER = 'downloaded_screenshots'
COMPRESSED_IMAGE_FOLDER = os.path.join(IMAGE_DOWNLOAD_FOLDER, 'compressed')

def select_xliff_files():
    root = tk.Tk()
    root.withdraw()
    print("Opening file dialog to select XLIFF file(s)...")
    filepaths = filedialog.askopenfilenames(
        title="Select XLIFF Files",
        filetypes=(("XLIFF files", "*.xliff"), ("XML files", "*.xml"), ("All files", "*.*"))
    )
    return filepaths

def create_output_folders():
    if not os.path.exists(IMAGE_DOWNLOAD_FOLDER):
        os.makedirs(IMAGE_DOWNLOAD_FOLDER)
    if not os.path.exists(COMPRESSED_IMAGE_FOLDER):
        os.makedirs(COMPRESSED_IMAGE_FOLDER)
    print(f"Image folders are ready.")

def download_image(url):
    try:
        filename = os.path.join(IMAGE_DOWNLOAD_FOLDER, url.split('/')[-1])
        if not os.path.exists(filename):
            print(f"Downloading {url}...")
            response = requests.get(url, stream=True, timeout=15)
            response.raise_for_status()
            with open(filename, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)
        return filename
    except requests.exceptions.RequestException as e:
        print(f"Error downloading {url}: {e}")
        return None

def compress_and_save_image(original_path):
    if not original_path:
        return None
    base_name = os.path.basename(original_path)
    compressed_name = os.path.splitext(base_name)[0] + '.jpg'
    compressed_path = os.path.join(COMPRESSED_IMAGE_FOLDER, compressed_name)
    if os.path.exists(compressed_path):
        return compressed_path
    try:
        with Image.open(original_path) as img:
            if img.mode in ('RGBA', 'P'):
                img = img.convert('RGB')
            print(f"Compressing {base_name}...")
            img.save(compressed_path, 'JPEG', quality=IMAGE_COMPRESSION_QUALITY, optimize=True)
        return compressed_path
    except Exception as e:
        print(f"Could not compress image {original_path}: {e}. Using original.")
        return original_path

def parse_xliff(file_path):
    print(f"\n--- Parsing {os.path.basename(file_path)} ---")
    ns = {'xliff': 'urn:oasis:names:tc:xliff:document:1.2'}
    try:
        tree = ET.parse(file_path)
        root = tree.getroot()
    except ET.ParseError:
        print(f"ERROR: '{os.path.basename(file_path)}' is not a valid XML/XLIFF. Skipping.")
        return []
    units_with_images = []
    segment_counter = 0
    all_trans_units = root.findall('.//xliff:trans-unit', ns)
    for trans_unit in all_trans_units:
        segment_counter += 1
        image_urls = [ctx.text.strip() for ctx in trans_unit.findall('.//xliff:context[@context-type="x-screenshots"]', ns) if ctx.text]
        if image_urls:
            source_element = trans_unit.find('xliff:source', ns)
            source_text = source_element.text.strip() if source_element is not None and source_element.text else "NO SOURCE TEXT"
            units_with_images.append({'number': segment_counter, 'source': source_text, 'urls': image_urls})
    print(f"Found {len(units_with_images)} segments with screenshots (out of {segment_counter} total segments).")
    return units_with_images

def add_bookmark(paragraph, bookmark_name):
    # Add a Word bookmark to a paragraph (for TOC linking)
    run = paragraph.add_run()
    tag = run._r
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), '0')
    start.set(qn('w:name'), bookmark_name)
    tag.addprevious(start)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), '0')
    tag.addnext(end)

def create_docx_with_toc(file_units_dict, output_filename):
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    doc.add_heading('Translation References', level=1)

    # Table of Contents placeholder
    toc_paragraph = doc.add_paragraph('Table of Contents\n', style='Intense Quote')
    toc_entries = []

    # Add a blank paragraph for spacing
    doc.add_paragraph("")

    # For each file, add a heading (with bookmark), table, and collect TOC info
    for idx, (file_path, units) in enumerate(file_units_dict.items()):
        file_name = os.path.basename(file_path)
        # Heading with bookmark for TOC
        heading = doc.add_heading(file_name, level=2)
        bookmark_name = f"file_{idx}"
        add_bookmark(heading, bookmark_name)
        toc_entries.append((file_name, bookmark_name))
        # Table for this file
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        table.columns[0].width = Cm(2.4)
        table.columns[1].width = Inches(3.0)
        table.columns[2].width = Inches(3.56)
        hdr_cells = table.rows[0].cells
        header_texts = ["Segment #", "Source Text", "Reference Screenshot(s)"]
        for i, text in enumerate(header_texts):
            p = hdr_cells[i].paragraphs[0]
            p.text = ""
            run = p.add_run(text)
            run.bold = True
        for unit in units:
            row_cells = table.add_row().cells
            row_cells[0].text = f"#{unit['number']}"
            row_cells[1].text = unit['source']
            paragraph = row_cells[2].paragraphs[0]
            for img_path in unit['image_paths']:
                try:
                    run = paragraph.add_run()
                    run.add_picture(img_path, width=Inches(3.3))
                    paragraph.add_run("\n")
                except Exception as e:
                    print(f"Could not add image {img_path} to Word: {e}")
        doc.add_paragraph("")  # Spacing after each table

    # Insert TOC entries after placeholder
    # (Word does not support auto-hyperlinked TOC via python-docx, but you can add manual links)
    toc_paragraph = doc.paragraphs[1]  # The paragraph after the main heading
    for file_name, bookmark_name in toc_entries:
        p = doc.add_paragraph()
        run = p.add_run(file_name)
        # Insert field code for hyperlink to bookmark (works in Word, not in all viewers)
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = f'HYPERLINK \\l "{bookmark_name}"'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        r_element = run._r
        r_element.append(fldChar1)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar3)
        p.style = 'List Bullet'

    doc.save(output_filename)
    print(f"\n✅ Success! Final DOCX file created: {output_filename}")

def main():
    xliff_files = select_xliff_files()
    if not xliff_files:
        print("No files selected. Exiting.")
        return
    create_output_folders()
    file_units_dict = {}
    for file_path in xliff_files:
        units_from_file = parse_xliff(file_path)
        # Download and compress images for each file separately
        for unit in units_from_file:
            unit['image_paths'] = []
            for url in unit['urls']:
                original_path = download_image(url)
                compressed_path = compress_and_save_image(original_path)
                if compressed_path:
                    unit['image_paths'].append(compressed_path)
        if units_from_file:
            file_units_dict[file_path] = units_from_file
    if not file_units_dict:
        print("\nNo translation units with screenshots were found. Exiting.")
        return
    output_dir = os.path.dirname(xliff_files[0])
    output_filename = os.path.join(output_dir, "consolidated_references.docx")
    create_docx_with_toc(file_units_dict, output_filename)
    try:
        os.startfile(output_filename)
    except AttributeError:
        print(f"\nTo view the file, open it from its location:\n{output_dir}")
    except Exception as e:
        print(f"Could not auto-open the file: {e}")

if __name__ == '__main__':
    main()
