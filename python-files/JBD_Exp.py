import pandas as pd
from docx import Document
from docx.shared import Cm
from datetime import datetime
import time
import os
import sys
import argparse
import gc
from pathlib import Path
from typing import Optional, Tuple, Dict, List, Set
import warnings
warnings.filterwarnings('ignore')

class OptimizedExcelToWordConverter:
    def __init__(self):
        self.start_time = time.time()
        self.processed_dates = 0
        self.inserted_events = 0
        self.replaced_events = 0
        self.error_count = 0
        self.skipped_events = 0
        
    def clean_text(self, text) -> str:
        """Очистка текста с минимальным использованием памяти"""
        if text is None:
            return ""
        if not isinstance(text, str):
            text = str(text)
        if len(text) == 0:
            return text
            
        # Быстрая очистка
        result = text
        for char in ['\x07', '\r', '\n', '\t']:
            if char in result:
                result = result.replace(char, '')
        
        # Удаляем конечные пробелы и специальные символы
        result = result.strip()
        while len(result) > 0 and result[-1] in ['\r', '\x07']:
            result = result[:-1]
            
        return result.strip()

    def is_valid_time(self, time_str: str) -> bool:
        """Проверка валидности времени БЕЗ регулярных выражений"""
        if not time_str or not isinstance(time_str, str):
            return False
        
        # Заменяем точку на двоеточие для единообразия
        time_str = time_str.replace('.', ':').strip()
        
        # Разделяем на части
        parts = time_str.split(':')
        
        # Проверяем что получилось 2 части
        if len(parts) != 2:
            return False
        
        try:
            hour, minute = int(parts[0]), int(parts[1])
            return 0 <= hour <= 23 and 0 <= minute <= 59
        except (ValueError, TypeError):
            return False

    def format_time(self, time_value) -> str:
        """Форматирование времени из различных форматов"""
        if time_value is None or pd.isna(time_value):
            return ""
            
        # Если это уже строка в формате HH:MM
        if isinstance(time_value, str):
            if ':' in time_value or '.' in time_value:
                time_str = time_value.replace('.', ':').strip()
                if self.is_valid_time(time_str):
                    parts = time_str.split(':')
                    return f"{int(parts[0]):02d}:{int(parts[1]):02d}"
            return ""
            
        # Если это datetime объект
        try:
            if hasattr(time_value, 'strftime'):
                return time_value.strftime('%H:%M')
            # Если это время в виде float (Excel serial time)
            elif isinstance(time_value, (int, float)):
                hours = int(time_value * 24)
                minutes = int((time_value * 24 - hours) * 60)
                if 0 <= hours <= 23 and 0 <= minutes <= 59:
                    return f"{hours:02d}:{minutes:02d}"
        except:
            pass
            
        return ""

    def format_date(self, date_value) -> str:
        """Форматирование даты из различных форматов"""
        if date_value is None or pd.isna(date_value):
            return ""
            
        # Если это уже строка
        if isinstance(date_value, str):
            # Пробуем разные форматы
            for fmt in ['%d.%m.%Y', '%Y-%m-%d', '%d/%m/%Y', '%m/%d/%Y']:
                try:
                    dt = datetime.strptime(date_value.strip(), fmt)
                    return dt.strftime('%d.%m.%Y')
                except:
                    continue
            return date_value.strip()
            
        # Если это datetime объект
        try:
            if hasattr(date_value, 'strftime'):
                return date_value.strftime('%d.%m.%Y')
        except:
            pass
            
        return str(date_value)

    def time_to_minutes(self, time_str: str) -> int:
        """Быстрое преобразование времени в минуты для сортировки"""
        try:
            time_str = time_str.replace('.', ':').strip()
            parts = time_str.split(':')
            return int(parts[0]) * 60 + int(parts[1])
        except:
            return 0

    def build_date_index(self, doc) -> Dict[str, Tuple]:
        """Построение индекса дат для быстрого поиска"""
        date_index = {}
        
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                try:
                    # Проверяем, что это объединенная ячейка (дата)
                    if len(row.cells) == 1 or (len(row.cells) > 1 and 
                                               all(cell.text == row.cells[0].text for cell in row.cells)):
                        date_text = self.clean_text(row.cells[0].text)
                        if date_text and not self.is_valid_time(date_text):
                            date_index[date_text] = (table, row_idx)
                except Exception as e:
                    print(f"Предупреждение при индексации строки {row_idx}: {e}")
                    continue
                    
        return date_index

    def get_first_two_sentences(self, text: str) -> str:
        """Извлечение первых двух предложений"""
        if not text:
            return ""
            
        sentences = []
        sentence_start = 0
        
        for i, char in enumerate(text):
            if char in '.!?':
                # Проверяем, что это действительно конец предложения
                if i == len(text) - 1 or text[i+1] in ' \t\n\r':
                    sentences.append(text[sentence_start:i+1].strip())
                    sentence_start = i + 1
                    if len(sentences) == 2:
                        break
                        
        # Если нашли меньше двух предложений, добавляем остаток
        if len(sentences) < 2 and sentence_start < len(text):
            remaining = text[sentence_start:].strip()
            if remaining:
                sentences.append(remaining)
                
        return ' '.join(sentences)

    def compare_first_two_sentences(self, text1: str, text2: str) -> bool:
        """Сравнение первых двух предложений"""
        return self.get_first_two_sentences(text1) == self.get_first_two_sentences(text2)

    def insert_row_standard(self, table, row_index: int) -> object:
        """Стандартная вставка строки в таблицу"""
        try:
            # Используем стандартный API python-docx
            new_row = table.add_row()
            
            # Перемещаем строку на нужную позицию
            if row_index < len(table.rows) - 1:
                # Получаем XML элементы
                tbl = table._tbl
                new_tr = new_row._tr
                
                # Удаляем строку из конца
                tbl.remove(new_tr)
                
                # Вставляем на нужную позицию
                ref_row = table.rows[row_index]._tr
                ref_row.addprevious(new_tr)
                
            return new_row
        except Exception as e:
            print(f"Ошибка при вставке строки: {e}")
            return None

    def process_event(self, table, date_row_idx: int, event_time: str, event_text: str) -> None:
        """Обработка события с минимальным использованием памяти"""
        time_found = False
        event_minutes = self.time_to_minutes(event_time)
        
        # Ищем подходящее место для вставки
        for row_idx in range(date_row_idx + 1, len(table.rows)):
            try:
                row = table.rows[row_idx]
                
                # Проверяем, что это строка с событием (2 ячейки)
                if len(row.cells) < 2:
                    break
                    
                # Если ячейки объединены (новая дата), прекращаем
                if row.cells[0].text == row.cells[1].text:
                    break
                    
                current_time = self.clean_text(row.cells[0].text)
                current_event = self.clean_text(row.cells[1].text)
                
                # Если время невалидное, возможно это конец событий для данной даты
                if not self.is_valid_time(current_time):
                    if not time_found and row_idx == date_row_idx + 1:
                        # Это первая строка после даты и она пустая - вставляем сюда
                        row.cells[0].text = event_time
                        row.cells[1].text = event_text
                        self.inserted_events += 1
                        return
                    break
                    
                current_minutes = self.time_to_minutes(current_time)
                
                # Нашли место для вставки
                if current_minutes >= event_minutes:
                    time_found = True
                    
                    if current_time == event_time:
                        # Время совпадает - проверяем текст
                        if self.compare_first_two_sentences(current_event, event_text):
                            # Заменяем событие
                            row.cells[1].text = event_text
                            self.replaced_events += 1
                        else:
                            # Вставляем новую строку после текущей
                            new_row = self.insert_row_standard(table, row_idx + 1)
                            if new_row:
                                new_row.cells[0].text = event_time
                                new_row.cells[1].text = event_text
                                self.inserted_events += 1
                    else:
                        # Вставляем перед текущей строкой
                        new_row = self.insert_row_standard(table, row_idx)
                        if new_row:
                            new_row.cells[0].text = event_time
                            new_row.cells[1].text = event_text
                            self.inserted_events += 1
                    return
                    
            except Exception as e:
                print(f"Ошибка при обработке строки {row_idx}: {e}")
                continue
                
        # Если не нашли подходящее место, добавляем в конец
        if not time_found:
            try:
                new_row = table.add_row()
                new_row.cells[0].text = event_time
                new_row.cells[1].text = event_text
                self.inserted_events += 1
            except Exception as e:
                print(f"Ошибка при добавлении строки в конец: {e}")
                self.error_count += 1

    def read_excel_data(self, excel_path: str) -> Optional[List]:
        """Улучшенное чтение данных из Excel"""
        print(f"Чтение данных из: {excel_path}")
        
        try:
            # Пробуем разные способы чтения Excel
            
            # Способ 1: Чтение с автоопределением листа
            try:
                # Сначала получаем информацию о файле
                xl_file = pd.ExcelFile(excel_path)
                sheet_names = xl_file.sheet_names
                print(f"Найдены листы: {sheet_names}")
                
                # Используем первый лист
                sheet_name = sheet_names[0] if sheet_names else 0
                
                # Читаем данные
                df = pd.read_excel(
                    excel_path,
                    sheet_name=sheet_name,
                    header=None,  # Без заголовков
                    skiprows=1    # Пропускаем первую строку если это заголовок
                )
                
                print(f"Прочитано {len(df)} строк, {len(df.columns)} колонок")
                
                # Проверяем количество колонок
                if len(df.columns) < 4:
                    print(f"Недостаточно колонок. Найдено: {len(df.columns)}, требуется минимум 4")
                    return None
                    
                # Берем колонки B, C, D (индексы 1, 2, 3)
                df_subset = df.iloc[:, 1:4].copy()
                df_subset.columns = ['date', 'time', 'event']
                
            except Exception as e:
                print(f"Ошибка при чтении с автоопределением: {e}")
                
                # Способ 2: Прямое чтение с указанием колонок
                df_subset = pd.read_excel(
                    excel_path,
                    usecols="B:D",  # Используем буквенные обозначения
                    names=['date', 'time', 'event'],
                    skiprows=1
                )
                
            # Обработка данных
            excel_data = []
            
            for idx, row in df_subset.iterrows():
                try:
                    # Форматируем дату
                    date_str = self.format_date(row['date'])
                    if not date_str:
                        continue
                        
                    # Форматируем время
                    time_str = self.format_time(row['time'])
                    if not time_str:
                        continue
                        
                    # Очищаем текст события
                    event_str = self.clean_text(str(row['event']))
                    if not event_str:
                        continue
                        
                    excel_data.append([date_str, time_str, event_str])
                    
                except Exception as e:
                    print(f"Предупреждение: пропуск строки {idx}: {e}")
                    continue
                    
            print(f"Успешно обработано {len(excel_data)} записей")
            
            # Сортировка по дате и времени
            excel_data.sort(key=lambda x: (x[0], self.time_to_minutes(x[1])))
            
            return excel_data
            
        except Exception as e:
            print(f"Критическая ошибка чтения Excel: {e}")
            import traceback
            traceback.print_exc()
            return None

    def restore_column_widths(self, doc) -> None:
        """Восстановление ширины столбцов"""
        try:
            for table in doc.tables:
                if len(table.columns) >= 2:
                    # Устанавливаем ширину первого столбца
                    for row in table.rows:
                        if len(row.cells) >= 2:
                            row.cells[0].width = Cm(2)
        except Exception as e:
            print(f"Предупреждение при восстановлении ширины столбцов: {e}")

    def process_data(self, excel_path: str, word_path: str, output_path: Optional[str] = None) -> bool:
        """Оптимизированная обработка данных"""
        try:
            if output_path is None:
                output_path = word_path
                
            # Чтение данных из Excel
            excel_data = self.read_excel_data(excel_path)
            if not excel_data:
                print("Не удалось прочитать данные из Excel!")
                return False
                
            print(f"\nОткрытие Word документа: {word_path}")
            doc = Document(word_path)
            
            # Создаем индекс дат
            date_index = self.build_date_index(doc)
            print(f"Найдено {len(date_index)} дат в документе")
            
            if not date_index:
                print("В документе не найдены даты!")
                print("Убедитесь, что в документе есть таблицы с датами в объединенных ячейках")
                return False
                
            # Обработка данных
            processed_dates_set = set()
            
            for i, (date_text, event_time, event_text) in enumerate(excel_data):
                if i % 50 == 0 and i > 0:
                    print(f"Обработано {i}/{len(excel_data)} записей...")
                    
                # Поиск даты в индексе
                if date_text in date_index:
                    table, row_idx = date_index[date_text]
                    
                    if date_text not in processed_dates_set:
                        self.processed_dates += 1
                        processed_dates_set.add(date_text)
                        
                    self.process_event(table, row_idx, event_time, event_text)
                else:
                    self.skipped_events += 1
                    if self.skipped_events <= 5:  # Показываем только первые 5 пропущенных
                        print(f"  Пропущено: дата '{date_text}' не найдена в документе")
                        
            # Восстановление ширины столбцов
            print("\nВосстановление форматирования...")
            self.restore_column_widths(doc)
            
            # Сохранение документа
            print(f"Сохранение документа в: {output_path}")
            doc.save(output_path)
            
            return True
            
        except Exception as e:
            self.error_count += 1
            print(f"Критическая ошибка: {str(e)}")
            import traceback
            traceback.print_exc()
            return False

def gui_file_selector(title: str, extension: str) -> Optional[str]:
    """Выбор файла через GUI (если доступен tkinter)"""
    try:
        import tkinter as tk
        from tkinter import filedialog
        
        root = tk.Tk()
        root.withdraw()
        
        file_path = filedialog.askopenfilename(
            title=title,
            filetypes=[(f"{extension.upper()} files", f"*.{extension}"), ("All files", "*.*")]
        )
        
        root.destroy()
        return file_path if file_path else None
        
    except ImportError:
        return None

def cli_file_selector(extension: str, work_dir: Path) -> Optional[str]:
    """Выбор файла через командную строку"""
    files = list(work_dir.glob(f"*.{extension}")) + list(work_dir.glob(f"*.{extension.upper()}"))
    
    if not files:
        print(f"В директории {work_dir} не найдены файлы .{extension}")
        return None
        
    print(f"\nНайдены файлы .{extension}:")
    for i, file in enumerate(files, 1):
        print(f"  {i}. {file.name}")
        
    while True:
        try:
            choice = input(f"Выберите номер файла (1-{len(files)}) или 0 для отмены: ")
            if choice == '0':
                return None
            idx = int(choice) - 1
            if 0 <= idx < len(files):
                return str(files[idx])
            else:
                print("Неверный номер. Попробуйте еще раз.")
        except (ValueError, KeyboardInterrupt):
            return None

def run_interactive(work_dir: Path):
    """Интерактивный режим работы"""
    converter = OptimizedExcelToWordConverter()
    
    print("="*60)
    print("КОНВЕРТЕР EXCEL В WORD (Журнал Боевых Действий)")
    print("="*60)
    
    # Выбор Excel файла
    print("\n1. Выберите Excel файл с данными:")
    excel_path = gui_file_selector('Выберите Excel файл', 'xlsx')
    if not excel_path:
        excel_path = cli_file_selector('xlsx', work_dir)
    
    if not excel_path:
        print("Excel файл не выбран. Выход.")
        return
        
    # Выбор Word файла
    print("\n2. Выберите Word документ для обновления:")
    word_path = gui_file_selector('Выберите Word документ', 'docx')
    if not word_path:
        word_path = cli_file_selector('docx', work_dir)
        
    if not word_path:
        print("Word файл не выбран. Выход.")
        return
        
    # Создание имени выходного файла
    output_path = str(Path(word_path).with_suffix('')) + '_updated.docx'
    
    print(f"\n3. Результат будет сохранен в: {output_path}")
    
    # Обработка
    print("\n" + "="*60)
    print("НАЧАЛО ОБРАБОТКИ")
    print("="*60)
    
    success = converter.process_data(excel_path, word_path, output_path)
    
    # Вывод статистики
    elapsed_time = time.time() - converter.start_time
    
    print("\n" + "="*60)
    if success:
        print("ОБРАБОТКА ЗАВЕРШЕНА УСПЕШНО!")
    else:
        print("ОБРАБОТКА ЗАВЕРШЕНА С ОШИБКАМИ!")
    print("="*60)
    
    print(f"Время выполнения: {elapsed_time:.2f} секунд")
    print(f"Обработано дат: {converter.processed_dates}")
    print(f"Вставлено событий: {converter.inserted_events}")
    print(f"Заменено событий: {converter.replaced_events}")
    print(f"Пропущено событий: {converter.skipped_events}")
    print(f"Ошибок: {converter.error_count}")
    
    if success:
        print(f"\nРезультат сохранен в: {output_path}")

def main():
    """Главная точка входа"""
    parser = argparse.ArgumentParser(
        description='Конвертер Excel в Word для Журнала Боевых Действий',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Примеры использования:
  python jbd_exp.py                                    # Интерактивный режим
  python jbd_exp.py --excel data.xlsx --word doc.docx  # Прямое указание файлов
  python jbd_exp.py --dir /path/to/files              # Указание рабочей директории
        """
    )
    
    parser.add_argument('--excel', type=str, help='Путь к Excel файлу с данными')
    parser.add_argument('--word', type=str, help='Путь к Word документу')
    parser.add_argument('--output', type=str, help='Путь для сохранения результата')
    parser.add_argument('--dir', type=str, help='Рабочая директория с файлами')
    
    args = parser.parse_args()
    
    # Определение рабочей директории
    if args.dir:
        work_dir = Path(args.dir)
    elif args.excel:
        work_dir = Path(args.excel).parent
    else:
        work_dir = Path.cwd()
        
    # Проверка существования директории
    if not work_dir.exists():
        print(f"Ошибка: директория {work_dir} не существует!")
        sys.exit(1)
        
    # Режим работы
    if args.excel and args.word:
        # Прямой режим с указанными файлами
        converter = OptimizedExcelToWordConverter()
        output_path = args.output or str(Path(args.word).with_suffix('')) + '_updated.docx'
        
        print("="*60)
        print("КОНВЕРТЕР EXCEL В WORD")
        print("="*60)
        
        if converter.process_data(args.excel, args.word, output_path):
            elapsed_time = time.time() - converter.start_time
            print(f"\nОбработка завершена за {elapsed_time:.2f} секунд")
            print(f"Результат: {output_path}")
        else:
            print("\nОбработка завершена с ошибками")
            sys.exit(1)
    else:
        # Интерактивный режим
        run_interactive(work_dir)

if __name__ == "__main__":
    main()