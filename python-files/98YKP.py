import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re
import sys
import io
from datetime import datetime
import pymorphy3
import glob
import os
import shutil

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
morph = pymorphy3.MorphAnalyzer()


def decline_fio_to_genitive(fio: str) -> str:
    """Склоняет ФИО в родительный падеж"""
    try:
        return ' '.join([morph.parse(part)[0].inflect({'gent'}).word.title() 
                if morph.parse(part)[0].inflect({'gent'}) 
                else part 
                for part in fio.split()])
    except Exception:
        return fio + ' (ошибка склонения)'

def select_all(event):
    event.widget.select_range(0, 'end')
    return 'break'

class EnhancedEntry(ttk.Entry):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.bind('<Control-a>', lambda e: (e.widget.event_generate('<<SelectAll>>'), 'break'))
        self.bind('<Control-A>', lambda e: (e.widget.event_generate('<<SelectAll>>'), 'break'))
        self.bind('<<SelectAll>>', select_all)

FAMILY_TEMPLATES = {
    'Член семьи': {
        'fields': ['член_семьи', 'фио_член_семьи', 'дата_рожд_член_семьи', 'банк_член_семьи',
                'лиц_счет_член_семьи', 'бик_член_семьи'],
        'format': lambda data: (
            f"{data['член_семьи']} - {data['фио_член_семьи']}, {data['дата_рожд_член_семьи']} г.р., в размере 5 000 000 рублей.\n"
            f"\tБанк получателя - {data['банк_член_семьи']}, лицевой счет - {data['лиц_счет_член_семьи']}, БИК банка - {data['бик_член_семьи']}."
        )
    },
    'Ребенок': {
        'fields': ['член_семьи', 'фио_член_семьи','дата_рожд_член_семьи', 'банк_член_семьи',
                'лиц_счет_член_семьи', 'бик_член_семьи'],
        'format': lambda data: (
            f"{data['член_семьи']} - {data['фио_член_семьи']}, {data['дата_рожд_член_семьи']} г.р., в размере 5 000 000 рублей.\n"
            f"Банк получателя - {data['банк_член_семьи']}, лицевой счет - {data['лиц_счет_член_семьи']}, БИК банка - {data['бик_член_семьи']}."
        )
    },
    'Законный представитель ребенка': {
        'fields': [
            'нл_член_семьи', 'дата_рождения_нл_член_семьи', 'член_семьи',
            'дата_рождения_член_семьи', 'серия_номер_паспорт_член_семьи',
            'кем_выдан_паспорт_член_семьи', 'паспорт_член_семьи_выдан',
            'банк_член_семьи', 'лиц_счет_член_семьи', 'бик_член_семьи'
        ],
        'format': lambda data: (
            f"законному представителю-несовершеннолетнего ребенка {data['нл_член_семьи']}, "
            f"{data['дата_рождения_нл_член_семьи']} г.р.; {data['член_семьи']}, "
            f"{data['дата_рождения_член_семьи']} г.р., паспорт "
            f"{data['серия_номер_паспорт_член_семьи']}, "
            f"{data['кем_выдан_паспорт_член_семьи']}, выдан "
            f"{data['паспорт_член_семьи_выдан']}.\n"
            f"Банк получателя - {data['банк_член_семьи']},"
            f" лицевой счет - "
            f"{data['лиц_счет_член_семьи']},\nБИК банка - {data['бик_член_семьи']}."
        )
    },
    'Нет данных': {
        'fields': ['член_семьи', 'фио_член_семьи'],
        'format': lambda data: f"{data['член_семьи']} - {data['фио_член_семьи']}.\nДанные не предоставлены."
    }
}

STATIC_FIELDS = ['звание_погибшего', 'фио_погибшего', 'дата_рожд_погибшего', 
                 'лн_погибшего', 'гибель_погибшего', 'должность', 'банки_абзац', 'искл_из_части']

CERTIFICATE_TYPES = ['копия свидетельства о смерти', 'копия свидетельства о рождении', 
                     'копия свидетельства о заключении брака', 'копия свидетельства о расторжении брака','копия медицинского свидетельства о смерти',
                     'копия свидетельства об усыновлении', 'копия справки о смерти', 'копия свидетельства об установлении отцовства'
                     'копия медицинского свидетельства о смерти', 'копия справки о смерти гражданина',
                     'копия заявления на ЕДВ ', 'копия паспорта ','копия банковских реквизитов ','копия номинального счета', 'копия справки об отсутствии факта государственной регистрации акта гражданского состояния',
                     'другое']

TEMPLATE_PATTERNS = [
    "Проект-приказа 98 УК (*).docx",
    "Реестр 98 (*)"


]


TEMPLATES = {}
family_frames = []
certificate_frames = []
output_folder = Path()

def create_gui():
    root = tk.Tk()
    root.title("98 Приказ")
    root.geometry("1200x900")

    main_canvas = tk.Canvas(root)
    main_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    scrollbar = ttk.Scrollbar(root, orient=tk.VERTICAL, command=main_canvas.yview)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    main_canvas.configure(yscrollcommand=scrollbar.set)
    main_canvas.bind('<Configure>', lambda e: main_canvas.configure(scrollregion=main_canvas.bbox("all")))
    main_canvas.bind("<MouseWheel>", lambda e: main_canvas.yview_scroll(int(-1*(e.delta/120)), "units"))

    main_frame = ttk.Frame(main_canvas)
    main_canvas.create_window((0,0), window=main_frame, anchor="nw")

    static_frame = ttk.LabelFrame(main_frame, text="Основные данные")
    static_frame.pack(fill=tk.X, pady=10, padx=15)

    static_entries = {}
    for field in STATIC_FIELDS:
        row = ttk.Frame(static_frame)
        row.pack(fill=tk.X, padx=5, pady=3)
        ttk.Label(row, text=field.replace('_', ' ').title() + ":", width=25).pack(side=tk.LEFT)
        entry = EnhancedEntry(row)
        entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        static_entries[field] = entry

    control_frame = ttk.Frame(main_frame)
    control_frame.pack(fill=tk.X, pady=10, padx=15)

    ttk.Button(control_frame, text="Выбрать папку с шаблонами", 
              command=lambda: load_template_folder()).pack(side=tk.LEFT, padx=5)
    ttk.Button(control_frame, text="Добавить родственника", 
              command=lambda: add_family_member(members_frame)).pack(side=tk.LEFT, padx=5)
    ttk.Button(control_frame, text="Добавить документ", 
              command=lambda: add_certificate(certificates_frame)).pack(side=tk.LEFT, padx=5)

    members_frame = ttk.LabelFrame(main_frame, text="Члены семьи")
    members_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=10)

    certificates_frame = ttk.LabelFrame(main_frame, text="Документы-основания")
    certificates_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=10)

    ttk.Button(main_frame, text="Создать документы", 
              style='Accent.TButton',
              command=lambda: generate_documents(static_entries)).pack(pady=15)

    root.mainloop()

def add_family_member(parent):
    frame = ttk.Frame(parent)
    frame.pack(fill=tk.X, pady=5, expand=True)
    
    template_var = tk.StringVar()
    cb = ttk.Combobox(frame, textvariable=template_var, 
                     values=list(FAMILY_TEMPLATES.keys()), 
                     state='readonly', width=25)
    cb.pack(side=tk.LEFT, padx=5)
    
    fields_frame = ttk.Frame(frame)
    fields_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

    def update_fields(event=None):
        for widget in fields_frame.winfo_children():
            widget.destroy()
        
        template_name = template_var.get()
        if not template_name:
            return
            
        template = FAMILY_TEMPLATES[template_name]
        entries = {}
        
        for field in template['fields']:
            row = ttk.Frame(fields_frame)
            row.pack(fill=tk.X, padx=2, pady=2)
            ttk.Label(row, text=field.replace('_', ' ').title() + ":", width=20).pack(side=tk.LEFT)
            entry = EnhancedEntry(row, width=35)
            entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
            entries[field] = entry
        
        frame.entries = entries
        frame.template = template_name
    
    cb.bind('<<ComboboxSelected>>', update_fields)
    ttk.Button(frame, text="×", command=lambda: remove_frame(frame, family_frames)).pack(side=tk.RIGHT)
    family_frames.append(frame)

def add_certificate(parent):
    frame = ttk.Frame(parent)
    frame.pack(fill=tk.X, pady=5, expand=True)
    
    cert_type = tk.StringVar()
    cb = ttk.Combobox(frame, textvariable=cert_type, 
                     values=CERTIFICATE_TYPES, 
                     state='readonly', width=30)
    cb.pack(side=tk.LEFT, padx=5)
    
    frame.cert_type = cert_type
    frame.entries = {}
    
    fields = [('фио_чьё_свид', 30), ('серия', 15), ('номер_свид', 20), ('документ', 30)]
    
    def update_fields(event=None):
        for widget in frame.winfo_children()[1:]:
            widget.destroy()
            
        frame.entries.clear()
        fields_frame = ttk.Frame(frame)
        fields_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        for (field, width) in fields[:3]:
            row = ttk.Frame(fields_frame)
            row.pack(side=tk.LEFT, padx=2)
            ttk.Label(row, text=field.replace('_', ' ').title() + ":").pack()
            entry = EnhancedEntry(row, width=width)
            entry.pack()
            frame.entries[field] = entry
            
        if cert_type.get() == 'другое':
            row = ttk.Frame(fields_frame)
            row.pack(side=tk.LEFT, padx=2)
            ttk.Label(row, text="Название документа:").pack()
            entry = EnhancedEntry(row, width=30)
            entry.pack()
            frame.entries['документ'] = entry
    
    cb.bind('<<ComboboxSelected>>', update_fields)
    ttk.Button(frame, text="×", command=lambda: remove_frame(frame, certificate_frames)).pack(side=tk.RIGHT)
    certificate_frames.append(frame)

def remove_frame(frame, frame_list):
    if frame in frame_list:
        frame_list.remove(frame)
    frame.destroy()

def load_template_folder():
    folder_path = filedialog.askdirectory()
    if folder_path:
        global TEMPLATES, output_folder
        TEMPLATES = {}
        
        # Загрузка основных шаблонов
        for pattern in TEMPLATE_PATTERNS:
            files = glob.glob(str(Path(folder_path) / pattern.replace('(*)', '*')), recursive=True)
            for file in files:
                key = Path(file).name.split(' ')[0].lower().replace('(', '').replace(')', '')
                TEMPLATES[key] = file
        
        if not TEMPLATES:
            messagebox.showerror("Ошибка", "Не найдены необходимые шаблоны!")
        else:
            output_folder = Path(folder_path)
            messagebox.showinfo("Успех", f"Загружено шаблонов: {len(TEMPLATES)}")


def get_basis_items():
    """Получает список оснований для приказа и реестра с сериями и номерами"""
    basis_items = []
    
    for frame in family_frames:
        if hasattr(frame, 'entries') and hasattr(frame, 'template'):
            data = {f: e.get() for f, e in frame.entries.items()}
            if frame.template != 'Нет данных' and 'фио_член_семьи' in data:
                series_number = data.get('серия_номер_паспорт_член_семьи', '')
                parts = series_number.split()
                series = parts[0] if len(parts) >= 1 else ''
                number = parts[1] if len(parts) >= 2 else ''
                
    
    for frame in certificate_frames:
        if hasattr(frame, 'entries') and hasattr(frame, 'cert_type'):
            data = {f: e.get() for f, e in frame.entries.items()}
            cert_type = frame.cert_type.get()
            
            if cert_type and data.get('фио_чьё_свид'):
                doc_text = data.get('документ', '') if cert_type == 'другое' else cert_type
                basis_items.append((
                    f"{doc_text} {data.get('фио_чьё_свид', '')}",
                    data.get('серия', ''),
                    data.get('номер_свид', '')
                ))
    
    return basis_items

def process_family_block(doc):
    for paragraph in doc.paragraphs:
        if "(СЕМЬЯ)" in paragraph.text:
            family_content = []
            for frame in family_frames:
                if hasattr(frame, 'entries') and hasattr(frame, 'template'):
                    data = {f: e.get() for f, e in frame.entries.items()}
                    content = FAMILY_TEMPLATES[frame.template]['format'](data)
                    family_content.append(content)
            
            new_text = '\n\n'.join(family_content)
            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.text = new_text

def process_basis_block(doc):
    basis_items = get_basis_items()
    
    for paragraph in doc.paragraphs:
        if "(основание)" in paragraph.text:
            basis_text = ", ".join([item[0] for item in basis_items]) + "."
            if paragraph.runs:
                paragraph.runs[0].text = basis_text
            else:
                paragraph.text = basis_text


def process_registry_table(doc):
    basis_items = get_basis_items()
    
    for table in doc.tables:
        placeholder_found = False
        start_row = 0
        
        # Поиск строки с плейсхолдером
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                if "(основание_реестр)" in cell.text:
                    placeholder_found = True
                    start_row = row_idx
                    cell.text = ""
                    break
            if placeholder_found:
                break

        if placeholder_found:
            # Удаляем только строку с плейсхолдером
            if len(table.rows) > start_row:
                table._tbl.remove(table.rows[start_row]._tr)

            # Добавляем новые строки
            for idx, item in enumerate(basis_items):
                if len(item) != 3:
                    continue
                
                new_row = table.add_row()
                item_name, series, number = item
                
                # Форматирование названия документа
                formatted_name = '. '.join([part.strip().capitalize() 
                                          for part in item_name.split('.')])
                
                # Номер строки (центрирование и стиль)
                if len(new_row.cells) > 0:
                    num_cell = new_row.cells[0]
                    num_cell.text = ""
                    for paragraph in num_cell.paragraphs:
                        paragraph.alignment = 1  # Центрирование
                        run = paragraph.add_run()
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)

                # Наименование документа
                if len(new_row.cells) > 1:
                    name_cell = new_row.cells[1]
                    name_cell.text = formatted_name
                    for paragraph in name_cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(12)

                # Количество (центрирование)
                if len(new_row.cells) > 2:
                    count_cell = new_row.cells[2]
                    count_cell.text = "1"
                    for paragraph in count_cell.paragraphs:
                        paragraph.alignment = 1
                        for run in paragraph.runs:
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(12)

                # Серия/номер
                if len(new_row.cells) > 3:
                    num_cell = new_row.cells[3]
                    num_text = ""
                    if series:
                        num_text += f"{series.strip().upper()}"
                    if number:
                        num_text += f"{number.strip()}"
                    num_cell.text = num_text.strip()
                    for paragraph in num_cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(12)

            # Перенумерация и стиль для всех строк
            for row_idx, row in enumerate(table.rows):
                if row_idx < 1: 
                    continue  # Пропускаем заголовки
                
                # Номер строки
                if len(row.cells) > 0:
                    num_cell = row.cells[0]
                    num_cell.text = str(row_idx - 0)
                    for paragraph in num_cell.paragraphs:
                        paragraph.alignment = 1  # Центрирование
                        for run in paragraph.runs:
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(12)



def process_template(doc, static_data):
    static_data['дата_сегодня'] = datetime.now().strftime("%d.%m.%Y")
    static_data['фио_погибшего_кого'] = decline_fio_to_genitive(static_data['фио_погибшего'])
    
    for paragraph in doc.paragraphs:
        full_text = paragraph.text
        new_text = full_text
        
        for ph, value in static_data.items():
            new_text = new_text.replace(f"({ph})", str(value))
            new_text = new_text.replace(f"({ph}_кого)", str(static_data['фио_погибшего_кого']))
        
        if new_text != full_text:
            if paragraph.runs:
                first_run = paragraph.runs[0]
                font = first_run.font
                style = first_run.style
                paragraph.clear()
                new_run = paragraph.add_run(new_text, style)
                new_run.font.name = font.name
                new_run.font.size = font.size
            else:
                paragraph.text = new_text

    process_family_block(doc)
    process_basis_block(doc)
    process_registry_table(doc)

def process_general_template(doc, static_data):
    # Общая обработка для всех документов
    for paragraph in doc.paragraphs:
        full_text = paragraph.text
        new_text = full_text
        
        # Замена плейсхолдеров
        for ph, value in static_data.items():
            new_text = new_text.replace(f"({ph})", str(value))
            new_text = new_text.replace(f"({ph}_кого)", str(static_data['фио_погибшего_кого']))
        
        # Сохранение форматирования
        if new_text != full_text:
            if paragraph.runs:
                first_run = paragraph.runs[0]
                font = first_run.font
                style = first_run.style
                paragraph.clear()
                new_run = paragraph.add_run(new_text, style)
                new_run.font.name = "Times New Roman"
                new_run.font.size = Pt(14)
            else:
                paragraph.text = new_text
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(14)

    # Обработка таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(14)

def generate_documents(static_entries):
    try:
        if not TEMPLATES:
            raise ValueError("Сначала выберите папку с шаблонами!")
        
        static_data = {field: entry.get() for field, entry in static_entries.items()}
        fio = static_data.get('фио_погибшего', '').strip()
        
        if not fio:
            raise ValueError("Поле 'фио_погибшего' обязательно для заполнения!")
        
        # Подготовка данных
        safe_fio = re.sub(r'[\\/*?:"<>|]', '_', fio)
        static_data['safe_fio'] = safe_fio
        static_data['дата_сегодня'] = datetime.now().strftime("%d.%m.%Y")
        static_data['фио_погибшего_кого'] = decline_fio_to_genitive(fio)
        
        # Создание основной папки
        output_dir = output_folder / safe_fio
        os.makedirs(output_dir, exist_ok=True)

        # Генерация основных документов
        for template_name, template_path in TEMPLATES.items():
            if 'Прочее' not in str(template_path):
                doc = Document(template_path)
                process_template(doc, static_data)
                original_name = Path(template_path).name
                # Замена всех плейсхолдеров в имени файла
                new_name = original_name.replace('(*)', safe_fio)\
                                       .replace('(фио_погибшего)', safe_fio)
                doc.save(output_dir / new_name)

        # Генерация документов разбирательства
        investigation_dir = output_dir / "Прочее"
        os.makedirs(investigation_dir, exist_ok=True)

        # Обработка дополнительных документов
        additional_patterns = [
            ('1', '1 Титульный лист.docx'),
            ('2', '2 Опись.docx'),
            ('3', '3 Рапорт (фио_погибшего).docx'),
            ('4', '4 Объяснительная (фио_погибшего).docx'),
            ('5', '5 Заключение (фио_погибшего).docx')
        ]

        for key, pattern in additional_patterns:
            template_path = TEMPLATES.get(key)
            if template_path:
                doc = Document(template_path)
                process_general_template(doc, static_data)
                # Замена плейсхолдеров в имени файла
                new_name = Path(template_path).name\
                    .replace('(фио_погибшего)', safe_fio)\
                    .replace('(*)', safe_fio)
                output_path = investigation_dir / new_name
                doc.save(output_path)

        messagebox.showinfo("Готово!", f"Документы сохранены в папку:\n{output_dir}")
        
    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка генерации:\n{str(e)}")



if __name__ == "__main__":
    create_gui()
