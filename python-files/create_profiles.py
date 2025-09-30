"""
create_profiles.py
Baca file Excel "Data Gereja GSJPDI - (Otonom).xlsx" di folder yang sama,
hasilkan:
 - Profil_75_Gereja_GSJPDI_Jawa.docx
 - Profil_75_Gereja_GSJPDI_Jawa.pdf

Cara pakai sebelum dikompilasi: python create_profiles.py
"""
import sys
import os
from pathlib import Path
import pandas as pd

# Word
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# PDF (reportlab)
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm

# ----- Config -----
EXCEL_FILENAME = "Data Gereja GSJPDI - (Otonom).xlsx"   # ubah bila perlu
OUTPUT_DOCX = "Profil_75_Gereja_GSJPDI_Jawa.docx"
OUTPUT_PDF = "Profil_75_Gereja_GSJPDI_Jawa.pdf"

# ----- Helpers -----
def clean_colname(c):
    # normalisasi nama kolom dari Excel supaya gampang dipanggil
    if not isinstance(c, str):
        return str(c)
    s = c.strip()
    s = s.replace("\n", " ")
    s = " ".join(s.split())
    return s

def safe_get(row, cols, fallback=""):
    # mencoba beberapa kemungkinan nama kolom, kembalikan yang pertama ada
    for c in cols:
        if c in row and pd.notna(row[c]):
            return str(row[c]).strip()
    return fallback

# ----- Main -----
def main():
    cwd = Path.cwd()
    excel_path = cwd / EXCEL_FILENAME
    if not excel_path.exists():
        print(f"ERROR: File Excel tidak ditemukan di: {excel_path}")
        print("Letakkan file Excel di folder yang sama dengan program ini, atau ubah EXCEL_FILENAME di script.")
        input("Tekan Enter untuk keluar...")
        sys.exit(1)

    print("Membaca file Excel:", excel_path)
    try:
        xls = pd.ExcelFile(excel_path)
        sheet = xls.sheet_names[0]
        df = pd.read_excel(excel_path, sheet_name=sheet, dtype=str)
    except Exception as e:
        print("Gagal membaca Excel:", e)
        input("Enter untuk keluar...")
        sys.exit(1)

    # normalisasi kolom
    df.columns = [clean_colname(c) for c in df.columns]
    print(f"Ditemukan {len(df)} baris (entri gereja).")

    # Bikin DOCX
    print("Membuat file Word:", OUTPUT_DOCX)
    doc = Document()
    # set Normal style font ke Times New Roman 12
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    doc.add_heading("Profil 75 Gereja GSJPDI Jawa", level=0)

    # kolom-kolom kemungkinan (sesuaikan urutan nama kolom yang ada di Excel Anda)
    # ini daftar pilihan nama kolom yang script akan coba pakai (berurut)
    col_map = {
        "nama": ["Nama Gereja :", "Nama Gereja", "Nama Gereja : " , "Nama Gereja"],
        "alamat": ["Alamat : Kepemilikan :", "Alamat :", "Alamat", "Alamat :"],
        "kepemilikan": ["Alamat : Kepemilikan :", "Kepemilikan :", "Kepemilikan"],
        "bentuk": ["Bentuk Bangunan :", "Bentuk Bangunan"],
        "ukuran": ["Ukuran Gereja : (Lebar x Panjang)", "Ukuran Gereja : (Lebar x Panjang)", "Ukuran"],
        "tahun": ["Keberadaan Gereja Sejak Tahun :", "Tahun Berdiri", "Tahun"],
        "cabang": ["Jumlah Gereja Cabang : .... Gereja   0", "Jumlah Gereja Cabang :", "Jumlah Cabang"],
        "gembala": ["DATA PELAYAN GEREJA   Nama Gembala Sidang :", "Nama Gembala Sidang :", "Nama Gembala Sidang", "Gembala Sidang"],
        "lahir": ["Tempat,Tanggal Lahir", "Tempat, Tanggal Lahir", "Tempat, Tgl Lahir"],
        "alamat_gembala": ["Alamat :.1", "Alamat .1", "Alamat Pelayan"],
        "hp": ["No. HP :", "No HP :", "No. HP"],
        "email": ["Email : (jika ada)", "Email Address", "Email"],
        "pejabat": ["Nama Pejabat Grejawi yang membantu Pelayanan : 1. 2. dst.", "Nama Pejabat Grejawi", "Pejabat Gerejawi"],
        "statistik": ["Statistik Jemaat Gereja Stempat Jemaat Dewasa Pria :", "Statistik Jemaat", "Statistik Jemaat Gereja Stempat"],
        "ibadah": ["Hari-Hari Ibadah : Nama Ibadah, Hari Ibadah, Jam Ibadah (Contoh) Ibadah Raya, Minggu, pukul 09.00 WIB Ibadah Doa, selasa pukul18.00 WIB", "Hari-Hari Ibadah", "Hari Ibadah"]
    }

    # iterate rows
    for idx, row in df.iterrows():
        i = idx + 1
        nama = safe_get(row, col_map["nama"], fallback=f"Gereja #{i}")
        doc.add_page_break()
        doc.add_heading(f"{i}. {nama}", level=1)

        alamat_raw = safe_get(row, col_map["alamat"], "")
        # sometimes "Alamat : Kepemilikan :" contains both; try split heuristically
        kepemilikan_raw = safe_get(row, col_map["kepemilikan"], "")
        # if alamat_raw contains '       ' or multiple spaces, attempt to split
        alamat = alamat_raw
        kepemilikan = kepemilikan_raw
        if "       " in alamat_raw and not kepemilikan_raw:
            parts = [p.strip() for p in alamat_raw.split("       ") if p.strip()]
            if len(parts) >= 2:
                alamat = parts[0]
                kepemilikan = parts[-1]

        doc.add_paragraph(f"Alamat: {alamat}")
        if kepemilikan:
            doc.add_paragraph(f"Kepemilikan: {kepemilikan}")

        bentuk = safe_get(row, col_map["bentuk"], "")
        ukuran = safe_get(row, col_map["ukuran"], "")
        if bentuk or ukuran:
            doc.add_paragraph(f"Bentuk Bangunan: {bentuk} {('('+ukuran+')') if ukuran else ''}")

        tahun = safe_get(row, col_map["tahun"], "")
        if tahun:
            doc.add_paragraph(f"Tahun Berdiri: {tahun}")

        ca
