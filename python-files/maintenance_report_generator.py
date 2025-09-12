
import pandas as pd
from docx import Document
from datetime import datetime
import tkinter as tk
from tkinter import filedialog, simpledialog, messagebox

def generate_report(file_path, report_type="monthly"):
    # Load data
    df = pd.read_excel(file_path)

    # Basic KPIs
    total_wos = len(df)
    completed = df[df['Status'].str.lower() == "completed"]
    backlog = df[df['Status'].str.lower() != "completed"]

    # Create Word Report
    doc = Document()
    doc.add_heading(f"{report_type.capitalize()} Maintenance Report", 0)
    doc.add_paragraph(f"Generated on {datetime.now().strftime('%d %B %Y')}\n")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        f"This {report_type} report provides key insights from Maximo work orders. "
        f"Total WOs: {total_wos}, Completed: {len(completed)}, Backlog: {len(backlog)}."
    )

    doc.add_heading("KPI Dashboard", level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Light List Accent 1"
    table.cell(0,0).text, table.cell(0,1).text = "Total Work Orders", str(total_wos)
    table.cell(1,0).text, table.cell(1,1).text = "Completed WOs", str(len(completed))
    table.cell(2,0).text, table.cell(2,1).text = "Backlog WOs", str(len(backlog))
    table.cell(3,0).text, table.cell(3,1).text = "Completion Rate (%)", f"{len(completed)/total_wos*100:.1f}%"

    doc.add_heading("Recommendations", level=1)
    doc.add_paragraph("• Improve preventive maintenance compliance.")
    doc.add_paragraph("• Focus on reducing backlog.")
    doc.add_paragraph("• Monitor MTTR for high-priority failures.")

    # Save file
    output_file = f"Report_{report_type}_{datetime.now().strftime('%Y_%m_%d')}.docx"
    doc.save(output_file)
    return output_file

def main():
    root = tk.Tk()
    root.withdraw()  # hide main window

    # Step 1: Choose Excel file
    file_path = filedialog.askopenfilename(
        title="Select Maximo Excel File",
        filetypes=[("Excel files", "*.xlsx *.xls")]
    )
    if not file_path:
        messagebox.showwarning("No File", "No Excel file selected.")
        return

    # Step 2: Choose report type
    report_type = simpledialog.askstring(
        "Report Type", "Enter report type (weekly, monthly, quarterly):"
    )
    if report_type not in ["weekly", "monthly", "quarterly"]:
        messagebox.showwarning("Invalid Input", "Please enter weekly, monthly, or quarterly.")
        return

    # Step 3: Generate report
    try:
        output_file = generate_report(file_path, report_type)
        messagebox.showinfo("Success", f"✅ Report saved as {output_file}")
    except Exception as e:
        messagebox.showerror("Error", str(e))

if __name__ == "__main__":
    main()
