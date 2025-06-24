import os
import tkinter as tk
from tkinter import filedialog, messagebox
from openpyxl import load_workbook
from docx import Document
from docx.shared import RGBColor
from copy import deepcopy
from docx.enum.text import WD_BREAK

# Helper to replace runs colored red in a document
def replace_runs(runs, mapping):
    for run in runs:
        color = run.font.color
        if color and getattr(color, 'rgb', None) == RGBColor(0xEE, 0x00, 0x00):
            key = run.text.strip()
            if key in mapping:
                run.text = str(mapping[key])

# Main generation function
def generate_protocols(excel_path, template_path, output_path):
    try:
        wb = load_workbook(excel_path, data_only=True)
    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось открыть Excel-файл:\n{e}")
        return
    sheet = wb.active
    headers = [cell.value for cell in sheet[1]]
    header_index = {header: idx for idx, header in enumerate(headers)}

    protocols_created = 0
    # Load template to preserve styles, then clear content
    master_doc = Document(template_path)
    body = master_doc.element.body
    for child in list(body):
        body.remove(child)

    # Iterate and append each filled protocol
    for row in sheet.iter_rows(min_row=2, values_only=True):
        data = {header: row[idx] for header, idx in header_index.items()}
        flag = str(data.get('Красный/не красный', '')).strip().lower()
        diploma_text = 'с отличием' if flag == 'красный' else 'без отличия'
        mapping = {
            '1': data.get('№ протокола', ''),
            'ФИО Студента': data.get('ФИО Студента', ''),
            'Группа': data.get('Группа', ''),
            'Тема диплома': data.get('Тема диплома', ''),
            'Руководитель ВКР': data.get('Руководитель ВКР', ''),
            'Кол-во листов': data.get('Кол-во листов', ''),
            'Оценка за ВКР': data.get('Оценка за ВКР', ''),
            'Красный/не красный': diploma_text
        }
        # Fill a fresh template copy
        filled = Document(template_path)
        for para in filled.paragraphs:
            replace_runs(para.runs, mapping)
        for table in filled.tables:
            for row_cells in table.rows:
                for cell in row_cells.cells:
                    for para in cell.paragraphs:
                        replace_runs(para.runs, mapping)

        # Before appending, insert page break if not first
        if protocols_created > 0:
            p = master_doc.add_paragraph()
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)
        # Append all elements of filled to master_doc
        for element in filled.element.body:
            master_doc.element.body.append(deepcopy(element))
        protocols_created += 1

    if protocols_created == 0:
        messagebox.showwarning("Внимание", "Нет данных для генерации протоколов.")
        return

    try:
        master_doc.save(output_path)
        messagebox.showinfo(
            "Готово",
            f"Сохранён файл с {protocols_created} протоколами:\n{output_path}"
        )
    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось сохранить файл:\n{e}")

# GUI setup
root = tk.Tk()
root.title("Генератор протоколов")
root.geometry("500x230")

# Excel-файл
tk.Label(root, text="Excel-файл:").grid(row=0, column=0, sticky="e", padx=5, pady=5)
excel_var = tk.StringVar()
tk.Entry(root, textvariable=excel_var, width=45).grid(row=0, column=1)
tk.Button(root, text="Обзор...", command=lambda: excel_var.set(filedialog.askopenfilename(filetypes=[("Excel files","*.xlsx *.xls")]))).grid(row=0, column=2)

# Шаблон документа
tk.Label(root, text="Документ-шаблон:").grid(row=1, column=0, sticky="e", padx=5, pady=5)
template_var = tk.StringVar()
tk.Entry(root, textvariable=template_var, width=45).grid(row=1, column=1)
tk.Button(root, text="Обзор...", command=lambda: template_var.set(filedialog.askopenfilename(filetypes=[("Word files","*.docx")]))).grid(row=1, column=2)

# Итоговый файл
tk.Label(root, text="Сохранить в файл:").grid(row=2, column=0, sticky="e", padx=5, pady=5)
out_var = tk.StringVar()
tk.Entry(root, textvariable=out_var, width=45).grid(row=2, column=1)
tk.Button(root, text="Обзор...", command=lambda: out_var.set(filedialog.asksaveasfilename(defaultextension='.docx', filetypes=[("Word files","*.docx")], initialfile="Протоколы.docx"))).grid(row=2, column=2)

# Кнопка генерации
tk.Button(root, text="Сгенерировать файл", command=lambda: generate_protocols(excel_var.get(), template_var.get(), out_var.get())).grid(row=3, column=0, columnspan=3, pady=30)

root.mainloop()
