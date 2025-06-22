import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import pandas as pd
from docx import Document
import os
from pathlib import Path
import sys

class CustomerDataExporter:
    def __init__(self, root):
        self.root = root
        self.root.title("Customer Data Exporter v1.0")
        self.root.geometry("550x350")
        self.root.resizable(False, False)
        self.center_window()
        self.input_file_path = ""
        self.create_widgets()
        
    def center_window(self):
        self.root.update_idletasks()
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry(f'{width}x{height}+{x}+{y}')
        
    def create_widgets(self):
        main_frame = ttk.Frame(self.root, padding="25")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        title_label = ttk.Label(main_frame, text="Customer Data Exporter", font=("Arial", 18, "bold"))
        title_label.grid(row=0, column=0, columnspan=2, pady=(0, 25))
        subtitle_label = ttk.Label(main_frame, text="Xuất dữ liệu khách hàng từ CSV/Excel sang Word", font=("Arial", 10))
        subtitle_label.grid(row=1, column=0, columnspan=2, pady=(0, 20))
        ttk.Label(main_frame, text="Chọn file dữ liệu:", font=("Arial", 10, "bold")).grid(row=2, column=0, sticky=tk.W, pady=5)
        file_frame = ttk.Frame(main_frame)
        file_frame.grid(row=3, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        self.file_path_var = tk.StringVar()
        self.file_entry = ttk.Entry(file_frame, textvariable=self.file_path_var, width=55, state="readonly", font=("Arial", 9))
        self.file_entry.grid(row=0, column=0, padx=(0, 10))
        browse_btn = ttk.Button(file_frame, text="Chọn file", command=self.browse_file)
        browse_btn.grid(row=0, column=1)
        progress_frame = ttk.Frame(main_frame)
        progress_frame.grid(row=4, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=20)
        ttk.Label(progress_frame, text="Tiến trình:").grid(row=0, column=0, sticky=tk.W)
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(progress_frame, variable=self.progress_var, maximum=100, length=450)
        self.progress_bar.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=5)
        self.status_var = tk.StringVar()
        self.status_var.set("Sẵn sàng - Chọn file để bắt đầu")
        status_label = ttk.Label(main_frame, textvariable=self.status_var, font=("Arial", 9), foreground="blue")
        status_label.grid(row=5, column=0, columnspan=2, pady=5)
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=6, column=0, columnspan=2, pady=25)
        export_btn = ttk.Button(button_frame, text="🚀 Xuất dữ liệu", command=self.export_data)
        export_btn.grid(row=0, column=0, padx=10)
        exit_btn = ttk.Button(button_frame, text="❌ Thoát", command=self.root.quit)
        exit_btn.grid(row=0, column=1, padx=10)
        main_frame.columnconfigure(0, weight=1)
        progress_frame.columnconfigure(0, weight=1)
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        
    def browse_file(self):
        file_types = [("Tất cả file hỗ trợ", "*.csv;*.xlsx;*.xls"), ("File CSV", "*.csv"), ("File Excel", "*.xlsx;*.xls"), ("Tất cả file", "*.*")]
        filename = filedialog.askopenfilename(title="Chọn file dữ liệu khách hàng", filetypes=file_types)
        if filename:
            self.input_file_path = filename
            self.file_path_var.set(filename)
            self.status_var.set(f"✅ Đã chọn file: {os.path.basename(filename)}")
    
    def read_data_file(self, file_path):
        try:
            file_extension = Path(file_path).suffix.lower()
            if file_extension == '.csv':
                encodings = ['utf-8', 'utf-8-sig', 'cp1252', 'iso-8859-1', 'latin1']
                df = None
                for encoding in encodings:
                    try:
                        df = pd.read_csv(file_path, encoding=encoding)
                        break
                    except (UnicodeDecodeError, UnicodeError):
                        continue
                if df is None:
                    raise Exception("Không thể đọc file CSV với các encoding thông dụng")
            elif file_extension in ['.xlsx', '.xls']:
                df = pd.read_excel(file_path, engine='openpyxl' if file_extension == '.xlsx' else None)
            else:
                raise Exception("Định dạng file không được hỗ trợ")
            return df
        except Exception as e:
            raise Exception(f"Lỗi đọc file: {str(e)}")
    
    def create_customer_document(self, customer_data, customer_name, order_number):
        doc = Document()
        title = doc.add_heading(f'THÔNG TIN KHÁCH HÀNG', 0)
        title.alignment = 1
        name_para = doc.add_heading(f'{customer_name}', level=1)
        name_para.alignment = 1
        doc.add_paragraph()
        doc.add_heading('Chi tiết thông tin:', level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Thông tin'
        hdr_cells[1].text = 'Giá trị'
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].bold = True
        for column, value in customer_data.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(column)
            row_cells[1].text = str(value) if pd.notna(value) else "N/A"
        doc.add_paragraph()
        doc.add_paragraph(f"Số thứ tự: {order_number:03d}")
        doc.add_paragraph(f"Ngày tạo: {pd.Timestamp.now().strftime('%d/%m/%Y %H:%M:%S')}")
        doc.add_paragraph(f"Tạo bởi: Customer Data Exporter v1.0")
        return doc
    
    def export_data(self):
        if not self.input_file_path:
            messagebox.showerror("Lỗi", "Vui lòng chọn file dữ liệu trước khi xuất!")
            return
        if not messagebox.askyesno("Xác nhận", "Bạn có muốn xuất dữ liệu không?"):
            return
        try:
            self.status_var.set("📖 Đang đọc dữ liệu...")
            self.root.update()
            df = self.read_data_file(self.input_file_path)
            if df.empty:
                messagebox.showerror("Lỗi", "File dữ liệu trống!")
                return
            input_dir = os.path.dirname(self.input_file_path)
            output_dir = os.path.join(input_dir, "Customer Data")
            os.makedirs(output_dir, exist_ok=True)
            name_columns = ['name', 'customer_name', 'ten', 'ho_ten', 'fullname', 'customer', 'họ tên', 'tên']
            name_col = None
            for col in df.columns:
                col_lower = str(col).lower().strip()
                if any(keyword in col_lower for keyword in name_columns):
                    name_col = col
                    break
            if name_col is None:
                name_col = df.columns[0]
                messagebox.showinfo("Thông báo", f"Không tìm thấy cột tên khách hàng.\nSử dụng cột '{name_col}' làm tên khách hàng.")
            total_customers = len(df)
            success_count = 0
            for index, (_, customer_row) in enumerate(df.iterrows(), 1):
                try:
                    self.status_var.set(f"📝 Đang xuất khách hàng {index}/{total_customers}...")
                    self.progress_var.set((index / total_customers) * 100)
                    self.root.update()
                    customer_name = str(customer_row[name_col]) if pd.notna(customer_row[name_col]) else f"Customer_{index}"
                    clean_name = "".join(c for c in customer_name if c.isalnum() or c in (' ', '-', '_', '.')).strip()
                    clean_name = clean_name.replace(' ', '_')
                    if not clean_name or clean_name == 'nan':
                        clean_name = f"Customer_{index}"
                    filename = f"{index:03d}_{clean_name}.docx"
                    filepath = os.path.join(output_dir, filename)
                    doc = self.create_customer_document(customer_row, customer_name, index)
                    doc.save(filepath)
                    success_count += 1
                except Exception as e:
                    print(f"Lỗi khi xuất khách hàng {index}: {str(e)}")
                    continue
            self.progress_var.set(100)
            self.status_var.set(f"✅ Hoàn thành! Đã xuất {success_count}/{total_customers} file")
            messagebox.showinfo("Thành công", f"Đã xuất thành công {success_count}/{total_customers} file Word!\n\n📁 Vị trí: {output_dir}\n\nBạn có muốn mở thư mục không?")
            try:
                os.startfile(output_dir)
            except AttributeError:
                try:
                    os.system(f'open "{output_dir}"')
                except:
                    os.system(f'xdg-open "{output_dir}"')
        except Exception as e:
            self.progress_var.set(0)
            self.status_var.set("❌ Có lỗi xảy ra")
            messagebox.showerror("Lỗi", f"Có lỗi xảy ra:\n{str(e)}\n\nVui lòng kiểm tra file dữ liệu.")

def main():
    if getattr(sys, 'frozen', False):
        import ctypes
        ctypes.windll.user32.ShowWindow(ctypes.windll.kernel32.GetConsoleWindow(), 0)
    root = tk.Tk()
    app = CustomerDataExporter(root)
    def on_closing():
        if messagebox.askokcancel("Thoát", "Bạn có muốn thoát ứng dụng không?"):
            root.destroy()
    root.protocol("WM_DELETE_WINDOW", on_closing)
    root.mainloop()

if __name__ == "__main__":
    main()