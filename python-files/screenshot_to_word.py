import keyboard
import pyautogui
from docx import Document
from datetime import datetime
import os

# Word file name
doc_name = "Test_Report.docx"

# Create or open document
if os.path.exists(doc_name):
    doc = Document(doc_name)
else:
    doc = Document()

def take_screenshot():
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    img_name = f"screenshot_{timestamp}.png"
    
    # Take screenshot
    img = pyautogui.screenshot()
    img.save(img_name)

    # Add to Word
    doc.add_picture(img_name, width=None)
    doc.add_paragraph(f"Step captured at {timestamp}")
    doc.save(doc_name)
    print(f"Screenshot saved: {img_name} -> {doc_name}")

# Assign shortcut
keyboard.add_hotkey('ctrl+shift+s', take_screenshot)

print("Press Ctrl+Shift+S to take screenshot and save to Word.")
print("Press ESC to quit the tool.")
keyboard.wait('esc')
