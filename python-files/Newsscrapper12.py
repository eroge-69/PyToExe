import feedparser
import datetime
import os
import unicodedata
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from urllib.parse import urlparse
from dateutil.parser import parse as parse_date
import threading
import tkinter as tk
from tkinter import scrolledtext, messagebox

# Define keywords (English and Urdu, focused on Karachi, Sindh, and Sindh Government)
keywords = [
    # Existing English keywords
    'Karachi', 'Sindh', 'Crime', 'Hyderabad', 'Landhi', 'MQM-P', 'MQM-L', 'Korangi', 'CPLC',
    'Creek', 'DHA', 'Pakistan Rangers (Sindh)', 'Hub', 'Katcha Area', 'Sea View', 'SITE',
    'Super Highway', 'Northern Bypass', 'Hawks Bay', 'Clifton', 'Malir', 'Gulshan-e-Iqbal',
    'Lyari', 'Saddar', 'Defence', 'Quaid’s Mausoleum', 'Port Qasim', 'Karachi Port', 'robbery',
    'flood', 'protest', 'traffic', 'outage', 'strike', 'Sukkur', 'Tando Muhammad Khan',
    'Tando Allahyar', 'Matiari', 'Jamshoro', 'Thatta', 'Sujawal', 'Badin', 'Mirpur Khas',
    'Umerkot', 'Tharparkar', 'Sanghar', 'Shaheed Benazirabad', 'Naushero Feroze', 'Dadu',
    'Larkana', 'Qambar Shahdadkot', 'Shikarpur', 'Jacobabad', 'Kashmore', 'Ghotki', 'Khairpur',
    'Nazimabad', 'North Karachi', 'Orangi', 'Gulistan-e-Johar', 'Surjani', 'Baldia', 'Keamari',
    'Shahrah-e-Faisal', 'Jinnah International Airport', 'Karachi University', 'Frere Hall',
    'Empress Market', 'Mohatta Palace', 'Merewether Tower', 'Bagh-e-Ibn-e-Qasim', 'Teen Talwar',
    'Do Darya', 'Bin Qasim', 'Manora', 'Sandspit', 'Kemari', 'dacoity', 'street crime',
    'water shortage', 'load shedding', 'sewage', 'traffic jam', 'encroachment', 'BYC',
    'Karachi Metropolitan Corporation', 'KMC', 'Sindh Police', 'Karachi Water Board', 'KWSC', 'Water Board',
    'Lyari Expressway', 'Saddar Town', 'Gulberg', 'PIB Colony', 'PECHS', 'Nazimabad No. 2',
    'Shah Faisal Colony', 'Liaquatabad', 'Garden East', 'Garden West', 'Mehmoodabad', 'Mehmoodabad Nullah',
    'Karachi Circular Railway', 'Green Line Bus', 'Karachi Stock Exchange', 'Karachi Shipyard',
    'Karachi Zoo', 'Karachi Expo Center', 'Karachi Electric', 'K-Electric', 'Karachi Fire Brigade',
    'Karachi Traffic Police', 'Karachi Press Club', 'Karachi Arts Council', 'Karachi Fish Harbour',
    'Karachi Steel Mills', 'Karachi Airport Security Force', 'Karachi Port Trust', 'Karachi Development Authority',
    'Karachi Building Control Authority', 'Saddar Bazaar', 'Boat Basin', 'Burns Road', 'Tariq Road',
    'Korangi Industrial Area', 'Federal B Area', 'Gulshan-e-Maymar', 'Shershah', 'Mauripur', 'Steel Town',
    'Malir Cantt', 'Karachi Gymkhana', 'Pakistan Maritime Museum', 'Safari Park', 'Hill Park',
    'Clifton Bridge', 'NIPA Chowrangi', 'Civic Center', 'Expo Centre', 'Karachi Golf Club',
    'Lyari Gang War', 'Abdullah Shah Ghazi', 'Gizri', 'Shah Latif Town', 'Korangi Creek',
    'Gulshan-e-Hadeed', 'Safoora Goth', 'Sohrab Goth', 'Manghopir', 'Gadap Town',
    'Bin Qasim Town', 'Fishermen Cooperative Society', 'PIA Township', 'Pakistan Steel Mills',
    'Port Grand', 'Hussainabad', 'Karsaz', 'Shahrah-e-Quaideen', 'Bahadurabad', 'Mehran Town',
    'Chandni Chowk', 'Ranchor Line', 'Golimar', 'Nazimabad No. 4', 'Paposh Nagar',
    'Ayesha Manzil', 'Nagan Chowrangi', 'Power House Chowrangi', 'Shershah Bridge', 'Memon Goth',
    'Malir River', 'Gulshan-e-Iqbal Park', 'Hill Park', 'Safari Park', 'PIB Colony',
    'Liaquat National Hospital', 'Aga Khan Hospital', 'Civil Hospital Karachi', 'Jinnah Hospital',
    'Abbasi Shaheed Hospital', 'Edhi Foundation', 'Chhipa', 'Saylani Welfare', 'Indus Hospital',
    'Indus River', 'Indus Delta', 'Rann of Kutch', 'Manchar Lake', 'Keenjhar Lake',
    'Gorakh Hill', 'Kirthar National Park', 'Makli Necropolis', 'Chaukhandi Tombs',
    'Bhambore', 'Kot Diji Fort', 'Mohenjo-daro', 'Thatta', 'Sehwan Sharif',
    'Lal Shahbaz Qalandar', 'Sufi shrines Sindh', 'Sindhi culture', 'Sindhi language',
    'Sindhi Ajrak', 'Sindhi Topi', 'Sindhi literature', 'Sindh Assembly', 'Sindh Cabinet',
    'Sindh Chief Secretary', 'Sindh Revenue Department', 'Sindh Board of Revenue',
    'Sindh Public Health Engineering', 'Sindh Rural Support Organization',
    'Sindh Education Foundation', 'Sindh High Court', 'Sindh Bar Council',
    'Sindh Local Government Department', 'Sindh Environmental Protection Agency',
    'Sindh Food Authority', 'Sindh Irrigation Department', 'Sindh Wildlife Department',
    'Sindh Police', 'Sindh Rangers', 'Sindh Home Department', 'Sindh Disaster Management Authority',
    'Sindh Public Service Commission', 'Karachi Circular Railway', 'K-IV water project',
    'Thar Coal Project', 'Green Line Bus Karachi', 'Orange Line Bus Karachi', 'Lyari Expressway',
    'Malir Expressway', 'Karachi Transformation Plan', 'Sindh Solar Energy Project',
    'Benazir Income Support Programme', 'Sindh Education Reform Program',
    'Civil Hospital Karachi', 'Jinnah Postgraduate Medical Centre', 'Liaquat University Hospital',
    'Dow University of Health Sciences', 'Aga Khan University Hospital', 'Indus Hospital Karachi',
    'Sindh Institute of Urology and Transplantation', 'Sindh Medical College', 'Sindh University Jamshoro',
    'Mehran University of Engineering and Technology', 'Shah Abdul Latif University Khairpur',
    'IBA Sukkur', 'Sindh Industrial Trading Estate', 'Karachi Stock Exchange', 'Port Qasim Authority',
    'Karachi Port Trust', 'Pakistan Steel Mills', 'Sindh Board of Investment',
    'Sindh Small Industries Corporation', 'Sindh Agriculture Department', 'Sindh Fisheries Department',
    'Sindh floods', 'Sindh drought', 'Sindh heatwave', 'Sindh local body elections', 'Sindh budget',
    'Sindh law and order', 'Sindh education crisis', 'Sindh health crisis', 'Sindh infrastructure',
    'Sindh unemployment', 'Sindh poverty', 'Sindh women empowerment', 'Sindh minorities', 'Sindh child labor',
    'water crisis', 'energy crisis', 'power outage', 'Sindh Government', 'Sindh Chief Minister',
    'Murad Ali Shah', 'Sindh Secretariat', 'Sindh Governor', 'Kamran Tessori', 'Sindh Public Works Department',
    'Sindh Planning Commission', 'Sindh Development Plan', 'Karachi Water and Sewerage Board', 'KWSB',
    'Sindh Transport Department', 'Sindh Health Department', 'Sindh Agriculture University',
    'Sindh Tourism Department', 'Sindh Culture Department', 'Sindh Energy Department',
    'Sindh Information Department', 'Sindh Labour Department', 'Sindh Women Development Department',
    'Sindh Human Rights Commission', 'Sindh Anti-Corruption Establishment', 'Sindh Ombudsman',
    'Karachi Urban Transport', 'Sindh Solid Waste Management', 'Sindh Building Control Authority', 'SBCA',
    'Karachi Development Package', 'Sindh Coastal Development', 'Sindh Livestock Department',
    'Sindh Forest Department', 'Sindh Fisheries Cooperative', 'Karachi Public Transport',
    'Sindh Mass Transit', 'Sindh Economic Zones', 'Thar Engro Coal Power Project',
    'Sindh Public Private Partnership', 'Sindh Investment Department', 'Karachi Urban Flooding',
    'Sindh Rural Development', 'Sindh Climate Change', 'Sindh Youth Policy',
    'Sindh Education Policy', 'Sindh Health Policy', 'Sindh Water Policy', 'Sindh Energy Policy',
    'Nagarparkar', 'Nagar Parkar', 'Nagarparkar Granite', 'Thar Desert', 'Nagarparkar Temples'
]

# Urdu keywords
urdu_keywords = [
    'کراچی', 'سندھ', 'سکھر', 'حیدرآباد', 'ٹنڈو محمد خان', 'ٹنڈو الہ یار', 'مٹیاری', 'جامشورو', 'ٹھٹھہ', 'سجاول', 'بدین', 'میرپور خاص', 'عمرکوٹ', 'تھرپارکر', 'سانگھڑ', 'شہید بے نظیر آباد', 'نوشہرو فیروز', 'دادو', 'لاڑکانہ', 'قمبر شہدادکوٹ', 'شکارپور', 'جیکب آباد', 'کشمور', 'گھوٹکی', 'خیرپور',
    'کلفٹن', 'مالیر', 'گلشن اقبال', 'لیاری', 'صدر', 'ڈیفنس', 'قائد کا مزار', 'بندرگاہ', 'ناظم آباد', 'نارتھ کراچی', 'اورنگی', 'گلستان جوہر', 'سرجانی', 'بلدیہ', 'کیماڑی', 'شاہراہ فیصل', 'جناح ایئرپورٹ', 'کراچی یونیورسٹی', 'فریئر ہال', 'ایمپریس مارکیٹ', 'موہٹہ پیلس', 'میرویتھر ٹاور', 'باغ ابن قاسم', 'تین تلوار', 'دو دریا', 'بن قاسم', 'مانورہ',
    'ڈکیتی', 'اسٹریٹ کرائم', 'پانی کی قلت', 'لوڈ شیڈنگ', 'سیوریج', 'ٹریفک جام', 'تجاوزات', 'سیلاب', 'احتجاج', 'ہڑتال', 'بجلی بندش', 'جرائم', 'پولیس', 'رینجرز', 'سڑک', 'ٹریفک', 'پانی', 'بجلی', 'چوری', 'قتل', 'اغوا', 'دہشت گردی', 'فساد', 'آگ', 'حادثہ', 'بارش', 'کچرا', 'صفائی',
    'کراچی میٹروپولیٹن کارپوریشن', 'کے ایم سی', 'سندھ پولیس', 'کراچی واٹر بورڈ', 'واٹر بورڈ', 'سندھ اسمبلی', 'سندھ کابینہ', 'سندھ چیف سیکریٹری', 'سندھ ریونیو ڈیپارٹمنٹ', 'سندھ بورڈ آف ریونیو', 'سندھ پبلک ہیلتھ انجینئرنگ', 'سندھ رورل سپورٹ آرگنائزیشن', 'سندھ ایجوکیشن فاونڈیشن', 'سندھ ہائی کورٹ', 'سندھ بار کونسل', 'سندھ لوکل گورنمنٹ ڈیپارٹمنٹ', 'سندھ انوائرمنٹل پروٹیکشن ایجنسی', 'سندھ فوڈ اتھارٹی', 'سندھ ایریگیشن ڈیپارٹمنٹ', 'سندھ وائلڈ لائف ڈیپارٹمنٹ', 'سندھ ہوم ڈیپارٹمنٹ', 'سندھ ڈیزاسٹر مینجمنٹ اتھارٹی', 'سندھ پبلک سروس کمیشن',
    'کراچی سرکلر ریلوے', 'کے فور واٹر پروجیکٹ', 'تھر کول پروجیکٹ', 'گرین لائن بس', 'اورنج لائن بس', 'لیاری ایکسپریس وے', 'ملیر ایکسپریس وے', 'کراچی ٹرانسفارمیشن پلان', 'سندھ سولر انرجی پروجیکٹ', 'بینظیر انکم سپورٹ پروگرام', 'سندھ ایجوکیشن ریفارم پروگرام',
    'سول ہسپتال کراچی', 'جناح پوسٹ گریجویٹ میڈیکل سینٹر', 'لیاقت یونیورسٹی ہسپتال', 'ڈاؤ یونیورسٹی آف ہیلتھ سائنسز', 'آغا خان یونیورسٹی ہسپتال', 'انڈس ہسپتال کراچی', 'سندھ انسٹی ٹیوٹ آف یورولوجی اینڈ ٹرانسپلانٹیشن', 'سندھ میڈیکل کالج', 'سندھ یونیورسٹی جامشورو', 'مہران یونیورسٹی آف انجینئرنگ اینڈ ٹیکنالوجی', 'شاہ عبداللطیف یونیورسٹی خیرپور', 'آئی بی اے سکھر',
    'سندھ انڈسٹریل ٹریڈنگ اسٹیٹ', 'کراچی اسٹاک ایکسچینج', 'پورٹ قاسم اتھارٹی', 'کراچی پورٹ ٹرسٹ', 'پاکستان اسٹیل ملز', 'سندھ بورڈ آف انویسٹمنٹ', 'سندھ اسمال انڈسٹریز کارپوریشن', 'سندھ ایگریکلچر ڈیپارٹمنٹ', 'سندھ فشریز ڈیپارٹمنٹ',
    'سندھ سیلاب', 'سندھ خشک سالی', 'سندھ ہیٹ ویو', 'سندھ بلدیاتی انتخابات', 'سندھ بجٹ', 'سندھ لاء اینڈ آرڈر', 'سندھ ایجوکیشن کرائسز', 'سندھ ہیلتھ کرائسز', 'سندھ انفراسٹرکچر', 'سندھ بے روزگاری', 'سندھ غربت', 'سندھ خواتین', 'سندھ اقلیتیں', 'سندھ چائلڈ لیبر',
    'کراچی واٹر اینڈ سیوریج', 'سندھ حکومت', 'جرائم پیشہ', 'ٹریفک مسائل', 'پانی کی کمی', 'بجلی کی بندش', 'پانی کا بحران', 'توانائی بحران',
    'سندھ گورنمنٹ', 'سندھ وزیراعلی', 'مراد علی شاہ', 'سندھ سیکریٹریٹ', 'سندھ گورنر', 'کامران ٹیسوری',
    'سندھ پبلک ورکس ڈیپارٹمنٹ', 'سندھ پلاننگ کمیشن', 'سندھ ڈیولپمنٹ پلان', 'کراچی واٹر اینڈ سیوریج بورڈ',
    'کے ڈبلیو ایس بی', 'سندھ ٹرانسپورٹ ڈیپارٹمنٹ', 'سندھ ہیلتھ ڈیپارٹمنٹ', 'سندھ ایگریکلچر یونیورسٹی',
    'سندھ ٹورزم ڈیپارٹمنٹ', 'سندھ کلچر ڈیپارٹمنٹ', 'سندھ انرجی ڈیپارٹمنٹ', 'سندھ انفارمیشن ڈیپارٹمنٹ',
    'سندھ لیبر ڈیپارٹمنٹ', 'سندھ وومن ڈیولپمنٹ ڈیپارٹمنٹ', 'سندھ ہیومن رائٹس کمیشن',
    'سندھ اینٹی کرپشن اسٹیبلشمنٹ', 'سندھ محتسب', 'کراچی اربن ٹرانسپورٹ', 'سندھ سالڈ ویسٹ مینجمنٹ',
    'سندھ بلڈنگ کنٹرول اتھارٹی', 'ایس بی سی اے', 'کراچی ڈیولپمنٹ پیکج', 'سندھ کوسٹل ڈیولپمنٹ',
    'سندھ لائیو اسٹاک ڈیپارٹمنٹ', 'سندھ فاریسٹ ڈیپارٹمنٹ', 'سندھ فشریز کوآپریٹو', 'کراچی پبلک ٹرانسپورٹ',
    'سندھ ماس ٹرانزٹ', 'سندھ اکنامک زونز', 'تھر اینگرو کول پاور پروجیکٹ', 'سندھ پبلک پرائیویٹ پارٹنرشپ',
    'سندھ انویسٹمنٹ ڈیپارٹمنٹ', 'کراچی اربن فلڈنگ', 'سندھ رورل ڈیولپمنٹ', 'سندھ کلائمیٹ چینج',
    'سندھ یوتھ پالیسی', 'سندھ ایجوکیشن پالیسی', 'سندھ ہیلتھ پالیسی', 'سندھ واٹر پالیسی', 'سندھ انرجی پالیسی',
    'ناگرپارکر', 'تھر ڈیزرٹ', 'ناگرپارکر گرینائٹ', 'ناگرپارکر مندر'
]

# Define RSS feed lists
english_rss_urls = [
    # Existing English RSS feeds
    'https://arynews.tv/en/category/pakistan/sindh/karachi/feed/',
    'https://en.dailypakistan.com.pk/rss/category/karachi',
    'https://www.samaa.tv/city/karachi/feed/',
    'https://www.bolnews.com/feed/karachi/',
    'https://www.24newshd.tv/rss/karachi',
    'https://nation.com.pk/rss/national/karachi',
    'https://www.nation.com.pk/rss/karachi',
    'https://www.dawn.com/feeds/karachi',
    'https://tribune.com.pk/rss/category/karachi',
    'https://www.geo.tv/rss/1',
    'https://dunyanews.tv/en/rss/City-Karachi',
    'https://www.samaa.tv/city/sindh/feed/',
    'https://www.urdupoint.com/en/rss/rss.aspx?site=latest-news/pakistan/sindh',
    'https://www.geo.tv/rss/2',
    'https://dunyanews.tv/en/rss/City-Sindh',
    'https://www.samaa.tv/city/sukkur/feed/',
    'https://www.samaa.tv/city/larkana/feed/',
    'https://www.samaa.tv/city/hyderabad/feed/',
    'https://www.samaa.tv/city/thatta/feed/',
    'https://www.samaa.tv/city/badin/feed/',
    'https://www.urdupoint.com/en/rss/rss.aspx?site=latest-news/pakistan/sindh/thatta',
    'https://www.urdupoint.com/en/rss/rss.aspx?site=latest-news/pakistan/sindh/badin',
    'https://www.urdupoint.com/en/rss/rss.aspx?site=latest-news/pakistan/sindh/larkana',
    'https://www.urdupoint.com/en/rss/rss.aspx?site=latest-news/pakistan/sindh/sukkur',
    'https://www.urdupoint.com/en/rss/rss.aspx?site=latest-news/pakistan/sindh/hyderabad',
    'https://tribune.com.pk/rss/category/hyderabad',
    'https://dunyanews.tv/en/rss/City-Hyderabad',
    'https://www.brecorder.com/rss/pakistan/sindh',
    'https://www.thenews.com.pk/rss/1/1',
    'https://www.thenews.com.pk/rss/1/2',
    'https://www.thenews.com.pk/rss/1/3',
    'https://www.thenews.com.pk/rss/1/4',
    'https://www.thenews.com.pk/rss/1/5',
    'https://www.thenews.com.pk/rss/1/6',
    'https://www.geo.tv/rss/1/karachi',
    'https://www.geo.tv/rss/2/sindh',
    'https://www.geo.tv/rss/3/hyderabad',
    'https://www.geo.tv/rss/4/sukkur',
    'https://www.geo.tv/rss/5/larkana',
    'https://www.geo.tv/rss/6/thatta',
    'https://www.geo.tv/rss/7/badin',
    'https://www.geo.tv/rss/8/mirpurkhas',
    'https://www.geo.tv/rss/9/khairpur',
    'https://www.geo.tv/rss/10/dadu',
    'https://www.geo.tv/rss/11/jacobabad',
    'https://www.geo.tv/rss/12/shikarpur',
    'https://www.geo.tv/rss/13/naushero-feroze',
    'https://www.geo.tv/rss/14/qambar-shahdadkot',
    'https://www.geo.tv/rss/15/umarkot',
    'https://www.geo.tv/rss/16/tharparkar',
    'https://www.geo.tv/rss/17/sanghar',
    'https://www.geo.tv/rss/18/gotki',
    'https://www.geo.tv/rss/19/kashmore',
    'https://www.geo.tv/rss/20/sujawal',
    'https://www.geo.tv/rss/21/matiari',
    'https://www.geo.tv/rss/22/jamshoro',
    'https://www.geo.tv/rss/23/tando-allahyar',
    'https://www.geo.tv/rss/24/tando-muhammad-khan',
    'https://www.geo.tv/rss/25/karachi-crime',
    'https://www.geo.tv/rss/26/karachi-water',
    'https://www.geo.tv/rss/27/karachi-energy',
    'https://www.geo.tv/rss/28/karachi-traffic',
    'https://www.dawn.com/feeds/sindh',
    'https://www.dawn.com/feeds/hyderabad',
    'https://www.dawn.com/feeds/sukkur',
    'https://www.dawn.com/feeds/larkana',
    'https://www.dawn.com/feeds/thatta',
    'https://www.dawn.com/feeds/badin',
    'https://www.dawn.com/feeds/mirpurkhas',
    'https://www.dawn.com/feeds/khairpur',
    'https://www.dawn.com/feeds/dadu',
    'https://www.dawn.com/feeds/jacobabad',
    'https://www.dawn.com/feeds/shikarpur',
    'https://www.dawn.com/feeds/naushero-feroze',
    'https://www.dawn.com/feeds/qambar-shahdadkot',
    'https://www.dawn.com/feeds/umarkot',
    'https://www.dawn.com/feeds/tharparkar',
    'https://www.dawn.com/feeds/sanghar',
    'https://www.dawn.com/feeds/gotki',
    'https://www.dawn.com/feeds/kashmore',
    'https://www.dawn.com/feeds/sujawal',
    'https://www.dawn.com/feeds/matiari',
    'https://www.dawn.com/feeds/jamshoro',
    'https://www.dawn.com/feeds/tando-allahyar',
    'https://www.dawn.com/feeds/tando-muhammad-khan',
    'https://www.thenews.com.pk/rss/1/7',
    'https://www.thenews.com.pk/rss/1/8',
    'https://www.thenews.com.pk/rss/1/9',
    'https://www.thenews.com.pk/rss/1/10',
    'https://www.thenews.com.pk/rss/1/11',
    'https://www.thenews.com.pk/rss/1/12',
    'https://www.thenews.com.pk/rss/1/13',
    'https://www.thenews.com.pk/rss/1/14',
    'https://www.thenews.com.pk/rss/1/15',
    'https://www.thenews.com.pk/rss/1/16',
    'https://www.thenews.com.pk/rss/1/17',
    'https://www.thenews.com.pk/rss/1/18',
    'https://www.thenews.com.pk/rss/1/19',
    'https://www.thenews.com.pk/rss/1/20',
    'https://www.thenews.com.pk/rss/1/21',
    'https://www.thenews.com.pk/rss/1/22',
    'https://www.thenews.com.pk/rss/1/23',
    'https://www.thenews.com.pk/rss/1/24',
    'https://www.thenews.com.pk/rss/1/25',
    'https://tribune.com.pk/rss/category/karachi-crime',
    'https://www.samaa.tv/city/karachi/crime/feed/',
    'https://www.bolnews.com/feed/karachi-crime/',
    'https://www.dawn.com/feeds/karachi-crime',
    'https://www.urdupoint.com/en/rss/rss.aspx?site=latest-news/pakistan/sindh/karachi/crime',
    'https://tribune.com.pk/rss/category/karachi-water',
    'https://www.samaa.tv/city/karachi/water/feed/',
    'https://www.bolnews.com/feed/karachi-water/',
    'https://www.urdupoint.com/en/rss/rss.aspx?site=latest-news/pakistan/sindh/karachi/water',
    'https://tribune.com.pk/rss/category/karachi-energy',
    'https://www.samaa.tv/city/karachi/energy/feed/',
    'https://www.bolnews.com/feed/karachi-energy/',
    'https://www.urdupoint.com/en/rss/rss.aspx?site=latest-news/pakistan/sindh/karachi/energy',
    'https://tribune.com.pk/rss/category/karachi-traffic',
    'https://www.samaa.tv/city/karachi/traffic/feed/',
    'https://www.bolnews.com/feed/karachi-traffic/',
    'https://www.urdupoint.com/en/rss/rss.aspx?site=latest-news/pakistan/sindh/karachi/traffic',
    # Additional English RSS feeds
    'https://pakobserver.net/category/pakistan/sindh/feed/',
    'https://www.thefridaytimes.com/category/pakistan/sindh/feed/',
    'https://www.brecorder.com/rss/pakistan/karachi',
    'https://www.thenews.com.pk/rss/1/pakistan/sindh',
    'https://tribune.com.pk/rss/category/sindh',
    'https://www.pakistantoday.com.pk/category/national/sindh/feed/',
    'https://dailytimes.com.pk/feed/?s=sindh',
    'https://www.nation.com.pk/rss/pakistan/sindh',
    'https://www.app.com.pk/category/national/sindh/feed/',
    'https://en.dailypakistan.com.pk/rss/category/hyderabad',
    'https://en.dailypakistan.com.pk/rss/category/sukkur',
    'https://www.bolnews.com/feed/hyderabad/',
    'https://www.bolnews.com/feed/sukkur/',
    # Newly added feeds
    'https://www.dawn.com/feeds/rss',
    'https://www.dawn.com/pakistan/feed',
    'https://pakistangulfeconomist.com/feed',
    'https://dailynewspk.com/feed'
]

urdu_rss_urls = [
    # Existing Urdu RSS feeds
    'https://www.urdupoint.com/ur/rss/rss.aspx?site=latest-news/pakistan/sindh/karachi',
    'https://www.samaa.tv/urdu/city/karachi/feed/',
    'https://www.bolnews.com/urdu/feed/karachi/',
    'https://www.24newshd.tv/urdu/rss/karachi',
    'https://www.nawaiwaqt.com.pk/rss/karachi',
    'https://urdu.geo.tv/rss/1',
    'https://express.pk/rss/karachi',
    'https://jang.com.pk/rss/karachi',
    'https://dunya.com.pk/rss/karachi',
    'https://www.urdupoint.com/ur/rss/rss.aspx?site=latest-news/pakistan/sindh',
    'https://urdu.geo.tv/rss/2',
    'https://dunya.com.pk/rss/sindh',
    'https://www.urdupoint.com/ur/rss/rss.aspx?site=latest-news/pakistan/sindh/thatta',
    'https://www.urdupoint.com/ur/rss/rss.aspx?site=latest-news/pakistan/sindh/badin',
    'https://www.urdupoint.com/ur/rss/rss.aspx?site=latest-news/pakistan/sindh/larkana',
    'https://www.urdupoint.com/ur/rss/rss.aspx?site=latest-news/pakistan/sindh/sukkur',
    'https://www.urdupoint.com/ur/rss/rss.aspx?site=latest-news/pakistan/sindh/hyderabad',
    'https://express.pk/rss/hyderabad',
    'https://jang.com.pk/rss/hyderabad',
    'https://dunya.com.pk/rss/hyderabad',
    'https://www.nawaiwaqt.com.pk/rss/hyderabad',
    'https://urdu.geo.tv/rss/1/karachi',
    'https://urdu.geo.tv/rss/2/sindh',
    'https://urdu.geo.tv/rss/3/hyderabad',
    'https://urdu.geo.tv/rss/4/sukkur',
    'https://urdu.geo.tv/rss/5/larkana',
    'https://urdu.geo.tv/rss/6/thatta',
    'https://urdu.geo.tv/rss/7/badin',
    'https://urdu.geo.tv/rss/8/mirpurkhas',
    'https://urdu.geo.tv/rss/9/khairpur',
    'https://urdu.geo.tv/rss/10/dadu',
    'https://urdu.geo.tv/rss/11/jacobabad',
    'https://urdu.geo.tv/rss/12/shikarpur',
    'https://urdu.geo.tv/rss/13/naushero-feroze',
    'https://urdu.geo.tv/rss/14/qambar-shahdadkot',
    'https://urdu.geo.tv/rss/15/umarkot',
    'https://urdu.geo.tv/rss/16/tharparkar',
    'https://urdu.geo.tv/rss/17/sanghar',
    'https://urdu.geo.tv/rss/18/gotki',
    'https://urdu.geo.tv/rss/19/kashmore',
    'https://urdu.geo.tv/rss/20/sujawal',
    'https://urdu.geo.tv/rss/21/matiari',
    'https://urdu.geo.tv/rss/22/jamshoro',
    'https://urdu.geo.tv/rss/23/tando-allahyar',
    'https://urdu.geo.tv/rss/24/tando-muhammad-khan',
    'https://www.express.pk/rss/sindh',
    'https://www.express.pk/rss/karachi',
    'https://www.express.pk/rss/hyderabad',
    'https://www.express.pk/rss/sukkur',
    'https://www.express.pk/rss/larkana',
    'https://www.express.pk/rss/thatta',
    'https://www.express.pk/rss/badin',
    'https://www.express.pk/rss/mirpurkhas',
    'https://www.express.pk/rss/khairpur',
    'https://www.express.pk/rss/dadu',
    'https://www.express.pk/rss/jacobabad',
    'https://www.express.pk/rss/shikarpur',
    'https://www.express.pk/rss/naushero-feroze',
    'https://www.express.pk/rss/qambar-shahdadkot',
    'https://www.express.pk/rss/umarkot',
    'https://www.express.pk/rss/tharparkar',
    'https://www.express.pk/rss/sanghar',
    'https://www.express.pk/rss/gotki',
    'https://www.express.pk/rss/kashmore',
    'https://www.express.pk/rss/sujawal',
    'https://www.express.pk/rss/matiari',
    'https://www.express.pk/rss/jamshoro',
    'https://www.express.pk/rss/tando-allahyar',
    'https://www.express.pk/rss/tando-muhammad-khan',
    'https://urdu.arynews.tv/feed/',
    'https://www.bbc.com/urdu/topics/pakistan/rss',
    'https://www.dawnnews.tv/rss',
    # Additional Urdu RSS feeds
    'https://www.nawaiwaqt.com.pk/rss/sindh',
    'https://www.nawaiwaqt.com.pk/rss/sukkur',
    'https://www.nawaiwaqt.com.pk/rss/thatta',
    'https://www.nawaiwaqt.com.pk/rss/badin',
    'https://urdu.dailypakistan.com.pk/rss/category/sindh',
    'https://urdu.dailypakistan.com.pk/rss/category/karachi',
    'https://urdu.dailypakistan.com.pk/rss/category/hyderabad',
    'https://urdu.dailypakistan.com.pk/rss/category/sukkur',
    'https://jang.com.pk/rss/sindh',
    'https://jang.com.pk/rss/sukkur',
    'https://jang.com.pk/rss/thatta',
    'https://jang.com.pk/rss/badin',
    'https://dunya.com.pk/rss/sukkur',
    'https://dunya.com.pk/rss/thatta',
    'https://dunya.com.pk/rss/badin'
]

# Deduplicate RSS URLs
english_rss_urls = list(dict.fromkeys(english_rss_urls))
urdu_rss_urls = list(dict.fromkeys(urdu_rss_urls))

# Function to set RTL property for a paragraph
def set_rtl(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)

# Function to highlight keywords in bold and red
def add_highlighted_text(paragraph, text, keywords):
    words = text.split()
    for word in words:
        match = next((kw for kw in keywords if kw.lower() in word.lower() or kw in word), None)
        run = paragraph.add_run(f"{word} ")
        if match:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)

# Function to clean HTML content
def clean_html_content(html_content, log_callback):
    cleaned_text = ""
    if html_content:
        if isinstance(html_content, list):
            for item in html_content:
                if isinstance(item, dict) and 'value' in item:
                    cleaned_text += f"{item['value']} "
                elif isinstance(item, str):
                    cleaned_text += f"{item} "
        elif isinstance(html_content, dict) and 'value' in html_content:
            cleaned_text = html_content['value']
        elif isinstance(html_content, str):
            cleaned_text = html_content
        else:
            log_callback(f"[WARNING] Unhandled content type: {type(html_content)}")
            return ""

        if cleaned_text:
            try:
                cleaned_text = unicodedata.normalize('NFKC', cleaned_text)
                soup = BeautifulSoup(cleaned_text, 'html.parser')
                text_content = soup.get_text(separator=' ', strip=True)
                text_content = ' '.join(text_content.split())
                if text_content and '[]' not in text_content and text_content != '':
                    return text_content
                else:
                    log_callback(f"[WARNING] Cleaned content is empty or invalid: {text_content}")
            except Exception as e:
                log_callback(f"Error cleaning HTML content: {e}")
    return ""

# Function to normalize URLs
def normalize_url(url):
    parsed = urlparse(url)
    return f"{parsed.scheme}://{parsed.netloc}{parsed.path}".rstrip('/')

# Function to scrape news from RSS feeds
def scrape_news(all_keywords, rss_urls, urdu_rss_urls, log_callback):
    news_items = []
    now = datetime.datetime.now(datetime.timezone.utc)
    english_keywords = [k for k in all_keywords if k not in urdu_keywords]
    normalized_urdu_rss_urls = set(normalize_url(u) for u in urdu_rss_urls)
    
    # Check for overlapping URLs
    overlapping_urls = set(english_rss_urls) & set(urdu_rss_urls)
    if overlapping_urls:
        log_callback(f"[WARNING] Overlapping URLs found: {overlapping_urls}")

    total_feeds = len(rss_urls)
    urdu_skipped = 0
    for i, url in enumerate(rss_urls):
        normalized_url = normalize_url(url)
        is_urdu = normalized_url in normalized_urdu_rss_urls
        log_callback(f"Processing feed {i+1}/{total_feeds}: {url} | is_urdu={is_urdu}")
        try:
            feed = feedparser.parse(url)
            entries = getattr(feed, 'entries', [])
            log_callback(f"[DEBUG] Feed {url} returned {len(entries)} entries.")
            if is_urdu:
                for idx, entry in enumerate(entries[:5]):
                    title = entry.get('title', 'No title')
                    published = entry.get('published', 'No published')
                    log_callback(f"[DEBUG-URDU-ALL] Entry {idx+1} Title: {title} | Published: {published}")
            if len(entries) == 0:
                log_callback(f"[WARNING] Feed {url} returned ZERO entries!")
            for entry in entries:
                published_time = None
                if hasattr(entry, 'published_parsed'):
                    published_time = datetime.datetime(*entry.published_parsed[:6], tzinfo=datetime.timezone.utc)
                elif hasattr(entry, 'published'):
                    try:
                        published_time = parse_date(entry.published)
                        published_time = published_time.replace(tzinfo=datetime.timezone.utc)
                    except ValueError:
                        log_callback(f"[WARNING] Could not parse date: {entry.get('published', 'N/A')}")
                        published_time = None

                # Relax date filtering for Urdu
                if published_time is None or (now - published_time).total_seconds() <= 24 * 3600 or is_urdu:
                    title = entry.get('title', 'No title')
                    content_sources = [
                        entry.get('content', ''),
                        entry.get('summary', ''),
                        entry.get('description', ''),
                        entry.get('summary_detail', {}).get('value', ''),
                        entry.get('content', [{}])[0].get('value', '')
                    ]
                    cleaned_content = ""
                    for content_source in content_sources:
                        cleaned_content = clean_html_content(content_source, log_callback)
                        if cleaned_content:
                            break

                    if not cleaned_content and title.strip() and any(keyword in title for keyword in urdu_keywords if is_urdu):
                        log_callback(f"[INFO] Using title only for {url}: {title}")
                        cleaned_content = title

                    if not cleaned_content:
                        if is_urdu:
                            urdu_skipped += 1
                        log_callback(f"Skipping entry from {url} due to empty content: {title}")
                        continue

                    text_to_check = f"{title} {cleaned_content}"
                    matches_keywords = False
                    if is_urdu:
                        matches_keywords = any(keyword in text_to_check for keyword in urdu_keywords)
                        matched_keywords = [kw for kw in urdu_keywords if kw in text_to_check]
                        log_callback(f"[DEBUG-URDU] Title: {title}\nContent: {cleaned_content}\nMatched keywords: {matched_keywords}")
                    else:
                        matches_keywords = any(keyword.lower() in text_to_check.lower() for keyword in english_keywords)

                    general_keywords_check = 'karachi' in text_to_check.lower() or 'sindh' in text_to_check.lower()

                    if (matches_keywords or general_keywords_check) and title.strip():
                        news_items.append({
                            'source': url,
                            'title': title,
                            'link': entry.get('link', 'No link'),
                            'published': published_time or datetime.datetime.min.replace(tzinfo=datetime.timezone.utc),
                            'content': cleaned_content,
                            'is_urdu': is_urdu
                        })
        except Exception as e:
            log_callback(f"Error processing RSS feed {url}: {e}")
    
    log_callback(f"[INFO] Total Urdu entries skipped: {urdu_skipped}")
    return news_items

# Function to create the Word document
def create_word_document(news_data, filename="Karachi_Sindh_News.docx"):
    document = Document()
    style = document.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # Add Executive Summary
    heading = document.add_heading('Executive Summary', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centered
    heading.style.font.name = 'Arial'
    heading.paragraph_format.line_spacing = 1.5

    if news_data:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.style.font.name = 'Arial'
        add_highlighted_text(p, "The following news items were published related to Karachi and Sindh in the last 24 hours:", keywords)
        for i, news_item in enumerate(news_data):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if news_item['is_urdu'] else WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.5
            p.style.font.name = 'Arial'
            if news_item['is_urdu']:
                set_rtl(p)
            add_highlighted_text(p, f"{i+1}.\t{news_item['title']} ({news_item['source']}).", keywords)
    else:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.style.font.name = 'Arial'
        add_highlighted_text(p, "No news found for Karachi and Sindh from the provided RSS feeds in the last 24 hours.", keywords)

    document.add_page_break()

    # Add Detailed News Items
    heading = document.add_heading('Karachi and Sindh News (Last 24 Hours)', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    heading.style.font.name = 'Arial'
    heading.paragraph_format.line_spacing = 1.5

    if news_data:
        english_news = [item for item in news_data if not item['is_urdu']]
        urdu_news = [item for item in news_data if item['is_urdu']]

        # English News Section
        if english_news:
            heading = document.add_heading('English News', level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            heading.style.font.name = 'Arial'
            heading.paragraph_format.line_spacing = 1.5
            for news_item in english_news:
                heading = document.add_heading(news_item['title'], level=2)
                heading.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                heading.style.font.name = 'Arial'
                heading.paragraph_format.line_spacing = 1.5
                
                p = document.add_paragraph(f"Source: {news_item['source']}")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = 1.5
                p.style.font.name = 'Arial'
                
                published_str = news_item['published'].strftime('%Y-%m-%d %H:%M:%S UTC') if news_item['published'] > datetime.datetime.min.replace(tzinfo=datetime.timezone.utc) else "N/A"
                p = document.add_paragraph(f"Published: {published_str}")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = 1.5
                p.style.font.name = 'Arial'
                
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = 1.5
                p.style.font.name = 'Arial'
                add_highlighted_text(p, news_item['content'], keywords)
                
                p = document.add_paragraph()
                p.paragraph_format.line_spacing = 1.5
                p.style.font.name = 'Arial'
                
                p = document.add_paragraph(f"Link: {news_item['link']}")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = 1.5
                p.style.font.name = 'Arial'
                
                p = document.add_paragraph("-" * 50)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = 1.5
                p.style.font.name = 'Arial'

        if english_news and urdu_news:
            document.add_page_break()

        # Urdu News Section
        if urdu_news:
            heading = document.add_heading('Urdu News', level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_rtl(heading)
            heading.style.font.name = 'Arial'
            heading.paragraph_format.line_spacing = 1.5
            for news_item in urdu_news:
                heading = document.add_heading(news_item['title'], level=2)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_rtl(heading)
                heading.style.font.name = 'Arial'
                heading.paragraph_format.line_spacing = 1.5
                
                p = document.add_paragraph(f"Source: {news_item['source']}")
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_rtl(p)
                p.style.font.name = 'Arial'
                p.paragraph_format.line_spacing = 1.5
                
                published_str = news_item['published'].strftime('%Y-%m-%d %H:%M:%S UTC') if news_item['published'] > datetime.datetime.min.replace(tzinfo=datetime.timezone.utc) else "N/A"
                p = document.add_paragraph(f"Published: {published_str}")
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_rtl(p)
                p.style.font.name = 'Arial'
                p.paragraph_format.line_spacing = 1.5
                
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_rtl(p)
                p.style.font.name = 'Arial'
                p.paragraph_format.line_spacing = 1.5
                add_highlighted_text(p, news_item['content'], keywords)
                
                p = document.add_paragraph()
                p.paragraph_format.line_spacing = 1.5
                p.style.font.name = 'Arial'
                
                p = document.add_paragraph(f"Link: {news_item['link']}")
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_rtl(p)
                p.style.font.name = 'Arial'
                p.paragraph_format.line_spacing = 1.5
                
                p = document.add_paragraph("-" * 50)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_rtl(p)
                p.style.font.name = 'Arial'
                p.paragraph_format.line_spacing = 1.5
    else:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.style.font.name = 'Arial'
        add_highlighted_text(p, "No detailed news items available.", keywords)

    # Add page number and RESTD in header, and RESTD in footer
    for section in document.sections:
        # Header: page number (centered), then RESTD below (bold, underlined)
        header = section.header
        # Remove the default topmost (empty) paragraph if present
        if header.paragraphs and header.paragraphs[0].text.strip() == '':
            p = header.paragraphs[0]._element
            header._element.remove(p)
        header_para = header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_para.style.font.name = 'Arial'
        header_para.paragraph_format.line_spacing = 1.0
        # Add page number field
        run = header_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        run._r.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)
        # Add RESTD below page number (bold, underlined)
        restd_para = header.add_paragraph()
        restd_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        restd_para.style.font.name = 'Arial'
        restd_para.paragraph_format.line_spacing = 1.0
        restd_run = restd_para.add_run('RESTD')
        restd_run.bold = True
        restd_run.underline = True

        # Footer: Algorithmically Curated News (left, blue) above RESTD (bold, underlined)
        footer = section.footer
        # Add Algorithmically Curated News (left, blue)
        acn_para = footer.add_paragraph()
        acn_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        acn_para.style.font.name = 'Arial'
        acn_para.paragraph_format.line_spacing = 1.0
        acn_run = acn_para.add_run('Algorithmically Curated News')
        acn_run.italic = True
        acn_run.font.color.rgb = RGBColor(0, 0, 225)
        # Add RESTD (bold, underlined, centered)
        restd_para = footer.add_paragraph()
        restd_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        restd_para.style.font.name = 'Arial'
        restd_para.paragraph_format.line_spacing = 1.0
        restd_run = restd_para.add_run('RESTD')
        restd_run.bold = True
        restd_run.underline = True
        # Add line above RESTD in footer (bottom border)
        p = restd_para._element
        pPr = p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '8')
        top.set(qn('w:space'), '1')
        top.set(qn('w:color'), '000000')
        pbdr.append(top)
        pPr.append(pbdr)

    desktop = os.path.join(os.path.expanduser("~"), 'Desktop')
    output_path = os.path.join(desktop, filename)
    document.save(output_path)
    return output_path

# Function to remove duplicates
def remove_duplicates(news_items):
    seen = set()
    unique_news = []
    for item in news_items:
        key = (item['title'].strip(), item['content'].strip())
        if key not in seen:
            seen.add(key)
            unique_news.append(item)
    return unique_news

# Function to run the scraper in GUI mode
def run_scraper_gui(log_callback, button=None):
    try:
        log_callback("Starting news scraping...")
        log_callback(f"Total English RSS feeds: {len(english_rss_urls)}")
        log_callback(f"Total Urdu RSS feeds: {len(urdu_rss_urls)}")
        log_callback(f"Total combined (deduplicated) RSS feeds: {len(set(english_rss_urls + urdu_rss_urls))}")

        all_rss_urls = english_rss_urls + urdu_rss_urls
        news_items = scrape_news(keywords, all_rss_urls, urdu_rss_urls, log_callback)
        log_callback(f"Found {len(news_items)} relevant news items.")

        if news_items:
            news_items_sorted = sorted(news_items, key=lambda x: (not x['is_urdu'], x['published']), reverse=True)
            news_items_unique = remove_duplicates(news_items_sorted)
            timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            filename = f"Karachi_Sindh_News_{timestamp}.docx"
            output_path = create_word_document(news_items_unique, filename)
            log_callback(f"SUCCESS! Word document saved to: {output_path}")
            if button:
                messagebox.showinfo("Done", f"Word document saved to:\n{output_path}")
        else:
            log_callback("No relevant news found.")
            if button:
                messagebox.showinfo("Done", "No relevant news found.")
    except Exception as e:
        log_callback(f"[ERROR] {e}")
        if button:
            messagebox.showerror("Error", str(e))
    finally:
        if button:
            button.config(state=tk.NORMAL)

# Function to start the Tkinter GUI
def start_gui():
    root = tk.Tk()
    root.title("Karachi/Sindh News Scraper")
    root.geometry("700x500")

    label = tk.Label(root, text="Karachi/Sindh News Scraper", font=("Arial", 16, "bold"))
    label.pack(pady=10)

    log_area = scrolledtext.ScrolledText(root, width=80, height=20, font=("Arial", 10))
    log_area.pack(padx=10, pady=10)
    log_area.config(state=tk.DISABLED)

    def log_callback(msg):
        log_area.config(state=tk.NORMAL)
        log_area.insert(tk.END, msg + "\n")
        log_area.see(tk.END)
        log_area.config(state=tk.DISABLED)

    def on_run():
        run_button.config(state=tk.DISABLED)
        log_area.config(state=tk.NORMAL)
        log_area.delete(1.0, tk.END)
        log_area.config(state=tk.DISABLED)
        threading.Thread(target=run_scraper_gui, args=(log_callback, run_button), daemon=True).start()

    run_button = tk.Button(root, text="Run Scraper", font=("Arial", 12), command=on_run)
    run_button.pack(pady=10)

    root.mainloop()

# Main function
def run_scraper():
    def log(message):
        print(message)

    log("Starting news scraping...")
    log(f"Total English RSS feeds: {len(english_rss_urls)}")
    log(f"Total Urdu RSS feeds: {len(urdu_rss_urls)}")
    log(f"Total combined (deduplicated) RSS feeds: {len(set(english_rss_urls + urdu_rss_urls))}")

    all_rss_urls = english_rss_urls + urdu_rss_urls
    news_items = scrape_news(keywords, all_rss_urls, urdu_rss_urls, log)
    log(f"Found {len(news_items)} relevant news items.")

    if news_items:
        news_items_sorted = sorted(news_items, key=lambda x: (not x['is_urdu'], x['published']), reverse=True)
        news_items_unique = remove_duplicates(news_items_sorted)
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"Karachi_Sindh_News_{timestamp}.docx"
        output_path = create_word_document(news_items_unique, filename)
        log(f"SUCCESS! Word document saved to: {output_path}")
    else:
        log("No relevant news found.")

if __name__ == "__main__":
    try:
        import sys
        if hasattr(sys, 'frozen') or (len(sys.argv) > 1 and sys.argv[1] == '--gui'):
            start_gui()
        else:
            run_scraper()
    except Exception as e:
        print(f"[ERROR] {e}")