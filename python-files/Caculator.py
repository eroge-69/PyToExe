"""
BloxVest Group Investment Tools - Main Application
Core implementation for the BloxVest desktop app as per the final design specification.
This is a high-level, modular implementation covering:
- Core calculators (Payout, Loan, ROI, Valuation, Deal Finder)
- UI (dark/light theme, accessibility, tabbed navigation)
- Export (Word/PDF with branding)
- Local encrypted scenario save/load/autosave
- Guided onboarding/help
- Accessibility features

Dependencies:
- PySide6 (Qt for Python) for GUI
- python-docx, reportlab for export
- pycryptodome for encryption
- matplotlib for charts
- (other stdlib modules as needed)

To run: `python main.py`
"""

import sys
import os
import json
import datetime
import threading
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QTabWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QPushButton, QTableWidget, QTableWidgetItem, QHeaderView,
    QSpinBox, QDoubleSpinBox, QComboBox, QCheckBox, QTextEdit, QFileDialog,
    QMessageBox, QScrollArea, QSlider, QStyleFactory, QMenuBar, QDialog,
    QFormLayout, QDialogButtonBox, QGroupBox, QGridLayout, QStackedWidget
)
from PySide6.QtGui import QIcon, QFont, QActionGroup, QAction, QPixmap, QColor, QPalette
from PySide6.QtCore import Qt, QTimer, QSize

# Encryption
from Crypto.Cipher import AES
from Crypto.Random import get_random_bytes
from Crypto.Protocol.KDF import PBKDF2

# Export
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# Charts
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

# Constants
APP_NAME = "BloxVest Group Investment Tools"
APP_VERSION = "1.0"
APP_ICON_PATH = "resources/bloxvest_icon.ico"
BRAND_ORANGE = "#FF8800"
BRAND_DARK = "#1e1e1e"
BRAND_GRAY = "#333333"
BRAND_FONT = "Segoe UI"
ENCRYPTION_KEY = b"bloxvest_super_secret_key_32bytes!!"  # 32 bytes for AES-256

AUTOSAVE_INTERVAL_MS = 60 * 1000  # 60 seconds

# Utility functions for encryption
def encrypt_data(data: dict) -> bytes:
    raw = json.dumps(data).encode("utf-8")
    salt = get_random_bytes(16)
    key = PBKDF2(ENCRYPTION_KEY, salt, dkLen=32)
    iv = get_random_bytes(16)
    cipher = AES.new(key, AES.MODE_CBC, iv)
    # Pad to 16 bytes
    pad_len = 16 - (len(raw) % 16)
    raw += bytes([pad_len]) * pad_len
    ct = cipher.encrypt(raw)
    return salt + iv + ct

def decrypt_data(blob: bytes) -> dict:
    salt = blob[:16]
    iv = blob[16:32]
    ct = blob[32:]
    key = PBKDF2(ENCRYPTION_KEY, salt, dkLen=32)
    cipher = AES.new(key, AES.MODE_CBC, iv)
    raw = cipher.decrypt(ct)
    pad_len = raw[-1]
    raw = raw[:-pad_len]
    return json.loads(raw.decode("utf-8"))

# --- Core Calculator Logic ---

def calculate_payout(total, participants):
    # participants: list of dicts with 'name' and 'percent'
    total_percent = sum(p['percent'] for p in participants)
    payouts = []
    for p in participants:
        share = round(total * (p['percent'] / 100))
        payouts.append({'name': p['name'], 'percent': p['percent'], 'payout': share})
    # Adjust for rounding
    payout_sum = sum(p['payout'] for p in payouts)
    diff = total - payout_sum
    if diff != 0 and payouts:
        # Adjust largest share
        max_idx = max(range(len(payouts)), key=lambda i: payouts[i]['payout'])
        payouts[max_idx]['payout'] += diff
    return payouts, total_percent

def calculate_loan_schedule(principal, annual_rate, months):
    r = annual_rate / 12 / 100
    n = months
    if r == 0:
        payment = principal / n
    else:
        payment = principal * r * (1 + r) ** n / ((1 + r) ** n - 1)
    payment = round(payment)
    schedule = []
    balance = principal
    total_interest = 0
    for i in range(1, n + 1):
        interest = round(balance * r)
        principal_paid = payment - interest
        if i == n:
            # Last payment, adjust for rounding
            principal_paid = balance
            payment = principal_paid + interest
        balance -= principal_paid
        total_interest += interest
        schedule.append({
            'period': i,
            'payment': payment,
            'interest': interest,
            'principal': principal_paid,
            'balance': max(0, balance)
        })
    total_paid = sum(row['payment'] for row in schedule)
    return schedule, payment, total_paid, total_interest

def calculate_roi(initial, final, months=None):
    roi = ((final - initial) / initial) * 100 if initial else 0
    annualized = None
    if months and months > 0:
        annualized = ((final / initial) ** (12 / months) - 1) * 100
    return roi, annualized

def calculate_valuation(pre=None, post=None, invest=None, percent=None, mode="percent"):
    # mode: which value to calculate
    if mode == "percent":
        # Given pre, invest
        post = pre + invest
        percent = invest / post * 100 if post else 0
        return {'pre': pre, 'post': post, 'invest': invest, 'percent': percent}
    elif mode == "valuation":
        # Given invest, percent
        post = invest / (percent / 100) if percent else 0
        pre = post - invest
        return {'pre': pre, 'post': post, 'invest': invest, 'percent': percent}
    # Add more modes as needed
    return {}

def calculate_deal_finder(plans):
    # plans: list of dicts with 'name', 'invest', 'percent'
    results = []
    for plan in plans:
        post = plan['invest'] / (plan['percent'] / 100) if plan['percent'] else 0
        pre = post - plan['invest']
        results.append({
            'name': plan['name'],
            'invest': plan['invest'],
            'percent': plan['percent'],
            'post': post,
            'pre': pre
        })
    # Rank by post-money valuation (higher is better)
    sorted_results = sorted(results, key=lambda x: x['post'], reverse=True)
    for idx, r in enumerate(sorted_results):
        r['rank'] = idx + 1
    # Mark best (rank 1)
    best_post = sorted_results[0]['post'] if sorted_results else 0
    for r in results:
        r['rank'] = 1 if r['post'] == best_post else 2
    return results

# --- UI Components ---

class PayoutTab(QWidget):
    DEVEX_RATE = 285  # Robux per 1 USD (can be made user-editable)
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()
        form = QFormLayout()
        self.total_input = QSpinBox()
        self.total_input.setMaximum(10**9)
        self.total_input.setSuffix(" R$")
        form.addRow("Total Robux to Distribute:", self.total_input)
        self.devex_rate_input = QSpinBox()
        self.devex_rate_input.setMaximum(100000)
        self.devex_rate_input.setValue(self.DEVEX_RATE)
        self.devex_rate_input.setSuffix(" R$ per USD")
        self.devex_rate_input.setToolTip("Robux per 1 USD for DevEx. Default: 285")
        form.addRow("DevEx Rate:", self.devex_rate_input)
        self.after_payment_checkbox = QCheckBox("Show After Payment State (remaining group funds)")
        form.addRow(self.after_payment_checkbox)
        self.participants = []
        self.participant_layout = QVBoxLayout()
        self.add_participant_row()
        self.add_participant_row()
        add_btn = QPushButton("+ Add Participant")
        add_btn.clicked.connect(self.add_participant_row)
        form.addRow("Participants:", add_btn)
        form.addRow(self.participant_layout)
        layout.addLayout(form)
        self.calc_btn = QPushButton("Calculate")
        self.calc_btn.clicked.connect(self.calculate)
        layout.addWidget(self.calc_btn)
        self.result_table = QTableWidget(0, 5)
        self.result_table.setHorizontalHeaderLabels(["Name", "Payout (R$)", "USD", "Tax (R$)", "USD After Tax"])
        layout.addWidget(self.result_table)
        self.result_table2 = QTableWidget(0, 5)
        self.result_table2.setHorizontalHeaderLabels(["Name", "Payout (R$)", "USD", "Tax (R$)", "USD After Tax"])
        layout.addWidget(self.result_table2)
        self.result_table2.hide()
        self.summary_label = QLabel("")
        layout.addWidget(self.summary_label)
        self.summary_label2 = QLabel("")
        layout.addWidget(self.summary_label2)
        self.setLayout(layout)

    def add_participant_row(self):
        row = QHBoxLayout()
        name = QLineEdit()
        percent = QDoubleSpinBox()
        percent.setSuffix(" %")
        percent.setMaximum(100.0)
        percent.setMinimum(0.0)
        row.addWidget(name)
        row.addWidget(percent)
        self.participant_layout.addLayout(row)
        self.participants.append((name, percent))

    def calculate(self):
        total = self.total_input.value()
        devex_rate = self.devex_rate_input.value() or self.DEVEX_RATE
        show_after = self.after_payment_checkbox.isChecked()
        participants = []
        for name, percent in self.participants:
            pname = name.text() or f"Participant {self.participants.index((name, percent))+1}"
            pval = percent.value()
            if pval > 0:
                participants.append({'name': pname, 'percent': pval})
        payouts, total_percent = calculate_payout(total, participants)
        # Table 1: Before payment
        self.result_table.setRowCount(len(payouts))
        total_robux = total_usd = total_tax = total_usd_after = 0
        for i, p in enumerate(payouts):
            robux = p['payout']
            usd = robux / devex_rate
            tax = robux * 0.3
            usd_after = (robux - tax) / devex_rate
            total_robux += robux
            total_usd += usd
            total_tax += tax
            total_usd_after += usd_after
            self.result_table.setItem(i, 0, QTableWidgetItem(p['name']))
            self.result_table.setItem(i, 1, QTableWidgetItem(f"{robux} R$"))
            self.result_table.setItem(i, 2, QTableWidgetItem(f"{usd:.2f} $"))
            self.result_table.setItem(i, 3, QTableWidgetItem(f"{tax:.0f} R$"))
            self.result_table.setItem(i, 4, QTableWidgetItem(f"{usd_after:.2f} $"))
        self.summary_label.setText(
            f"Total = {total_percent:.1f}% allocated. Total payout: {total_robux} R$ ({total_usd:.2f} $), Taxes: {total_tax:.0f} R$, USD after tax: {total_usd_after:.2f} $"
        )
        # Table 2: After payment (remaining group funds)
        if show_after:
            self.result_table2.show()
            self.summary_label2.show()
            remaining = total - total_robux
            payouts2, total_percent2 = calculate_payout(remaining, participants)
            self.result_table2.setRowCount(len(payouts2))
            t2_robux = t2_usd = t2_tax = t2_usd_after = 0
            for i, p in enumerate(payouts2):
                robux = p['payout']
                usd = robux / devex_rate
                tax = robux * 0.3
                usd_after = (robux - tax) / devex_rate
                t2_robux += robux
                t2_usd += usd
                t2_tax += tax
                t2_usd_after += usd_after
                self.result_table2.setItem(i, 0, QTableWidgetItem(p['name']))
                self.result_table2.setItem(i, 1, QTableWidgetItem(f"{robux} R$"))
                self.result_table2.setItem(i, 2, QTableWidgetItem(f"{usd:.2f} $"))
                self.result_table2.setItem(i, 3, QTableWidgetItem(f"{tax:.0f} R$"))
                self.result_table2.setItem(i, 4, QTableWidgetItem(f"{usd_after:.2f} $"))
            self.summary_label2.setText(
                f"After payment: Remaining group funds: {remaining} R$. Total payout: {t2_robux} R$ ({t2_usd:.2f} $), Taxes: {t2_tax:.0f} R$, USD after tax: {t2_usd_after:.2f} $"
            )
        else:
            self.result_table2.hide()
            self.summary_label2.hide()

class LoanTab(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()
        form = QFormLayout()
        self.principal_input = QSpinBox()
        self.principal_input.setMaximum(10**9)
        self.principal_input.setSuffix(" R$")
        self.rate_input = QDoubleSpinBox()
        self.rate_input.setSuffix(" %")
        self.rate_input.setMaximum(100.0)
        self.term_input = QSpinBox()
        self.term_input.setMaximum(360)
        self.term_input.setSuffix(" months")
        form.addRow("Principal (R$):", self.principal_input)
        form.addRow("Interest Rate (%/yr):", self.rate_input)
        form.addRow("Term (months):", self.term_input)
        layout.addLayout(form)
        self.warning_label = QLabel(
            '<b style="color:#FF8800">Warning:</b> Roblox group loans are not protected by contracts. Repayment is based on trust or scripts. There is no legal recourse if the borrower defaults!'
        )
        self.warning_label.setWordWrap(True)
        layout.addWidget(self.warning_label)
        self.calc_btn = QPushButton("Calculate")
        self.calc_btn.clicked.connect(self.calculate)
        layout.addWidget(self.calc_btn)
        self.result_table = QTableWidget(0, 5)
        self.result_table.setHorizontalHeaderLabels(
            ["Period", "Payment", "Interest", "Principal", "Balance"]
        )
        layout.addWidget(self.result_table)
        self.summary_label = QLabel("")
        layout.addWidget(self.summary_label)
        self.setLayout(layout)

    def calculate(self):
        principal = self.principal_input.value()
        rate = self.rate_input.value()
        months = self.term_input.value()
        schedule, payment, total_paid, total_interest = calculate_loan_schedule(principal, rate, months)
        self.result_table.setRowCount(len(schedule))
        for i, row in enumerate(schedule):
            self.result_table.setItem(i, 0, QTableWidgetItem(str(row['period'])))
            self.result_table.setItem(i, 1, QTableWidgetItem(f"{row['payment']} R$"))
            self.result_table.setItem(i, 2, QTableWidgetItem(f"{row['interest']} R$"))
            self.result_table.setItem(i, 3, QTableWidgetItem(f"{row['principal']} R$"))
            self.result_table.setItem(i, 4, QTableWidgetItem(f"{row['balance']} R$"))
        self.summary_label.setText(
            f"Monthly Payment: {payment} R$; Total Paid: {total_paid} R$; Total Interest: {total_interest} R$"
        )

class ROITab(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()
        form = QFormLayout()
        self.initial_input = QSpinBox()
        self.initial_input.setMaximum(10**9)
        self.initial_input.setSuffix(" R$")
        self.final_input = QSpinBox()
        self.final_input.setMaximum(10**9)
        self.final_input.setSuffix(" R$")
        self.months_input = QSpinBox()
        self.months_input.setMaximum(120)
        self.months_input.setSuffix(" months")
        self.revenue_history = QTextEdit()
        self.revenue_history.setPlaceholderText("Enter monthly group revenue (Robux), one per line, most recent last.")
        self.revenue_history.setToolTip("Optional: Enter monthly group revenue history for more accurate ROI.")
        self.projected_revenue = QSpinBox()
        self.projected_revenue.setMaximum(10**9)
        self.projected_revenue.setSuffix(" R$")
        self.projected_revenue.setToolTip("Optional: Projected next month's revenue (Robux)")
        form.addRow("Initial Investment (R$):", self.initial_input)
        form.addRow("Final Value (R$):", self.final_input)
        form.addRow("Time Frame (months):", self.months_input)
        form.addRow("Monthly Revenue History:", self.revenue_history)
        form.addRow("Projected Next Month Revenue:", self.projected_revenue)
        layout.addLayout(form)
        self.calc_btn = QPushButton("Calculate ROI")
        self.calc_btn.clicked.connect(self.calculate)
        layout.addWidget(self.calc_btn)
        self.result_label = QLabel("")
        layout.addWidget(self.result_label)
        self.setLayout(layout)

    def calculate(self):
        initial = self.initial_input.value()
        final = self.final_input.value()
        months = self.months_input.value()
        # Use revenue history if provided
        revenue_lines = self.revenue_history.toPlainText().strip().splitlines()
        revenues = [int(line.strip()) for line in revenue_lines if line.strip().isdigit()]
        projected = self.projected_revenue.value()
        if revenues:
            total_revenue = sum(revenues)
            if projected:
                total_revenue += projected
                months_used = len(revenues) + 1
            else:
                months_used = len(revenues)
            roi = ((total_revenue - initial) / initial) * 100 if initial else 0
            annualized = ((total_revenue / initial) ** (12 / months_used) - 1) * 100 if initial and months_used > 0 else None
            text = f"ROI (using revenue): {roi:.2f}%"
            if annualized is not None:
                text += f"; Annualized ROI: {annualized:.2f}%"
            text += f"\nTotal Revenue: {total_revenue} R$ over {months_used} months"
        else:
            roi, annualized = calculate_roi(initial, final, months)
            text = f"ROI: {roi:.2f}%"
            if annualized is not None:
                text += f"; Annualized ROI: {annualized:.2f}%"
        self.result_label.setText(text)

class ValuationTab(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()
        form = QFormLayout()
        self.method_combo = QComboBox()
        self.method_combo.addItems([
            "Revenue Multiplier (Roblox)",
            "Asset Value (Funds + Limiteds)",
            "Member Value (per member)",
            "Custom (Manual)"
        ])
        self.method_combo.currentIndexChanged.connect(self.update_method_fields)
        self.pre_input = QSpinBox()
        self.pre_input.setMaximum(10**9)
        self.pre_input.setSuffix(" R$")
        self.invest_input = QSpinBox()
        self.invest_input.setMaximum(10**9)
        self.invest_input.setSuffix(" R$")
        self.percent_input = QDoubleSpinBox()
        self.percent_input.setMaximum(100.0)
        self.percent_input.setSuffix(" %")
        self.multiplier_input = QDoubleSpinBox()
        self.multiplier_input.setMaximum(100.0)
        self.multiplier_input.setValue(8.0)
        self.multiplier_input.setSuffix(" x")
        self.funds_input = QSpinBox()
        self.funds_input.setMaximum(10**9)
        self.funds_input.setSuffix(" R$")
        self.limiteds_input = QSpinBox()
        self.limiteds_input.setMaximum(10**9)
        self.limiteds_input.setSuffix(" R$")
        self.members_input = QSpinBox()
        self.members_input.setMaximum(10**7)
        self.members_input.setSuffix(" members")
        self.per_member_input = QSpinBox()
        self.per_member_input.setMaximum(10000)
        self.per_member_input.setValue(1)
        self.per_member_input.setSuffix(" R$")
        self.custom_input = QSpinBox()
        self.custom_input.setMaximum(2_000_000_000)
        self.custom_input.setSuffix(" R$")
        form.addRow("Valuation Method:", self.method_combo)
        form.addRow("Pre-Money Valuation (R$):", self.pre_input)
        form.addRow("Investment Amount (R$):", self.invest_input)
        form.addRow("Investor Equity %:", self.percent_input)
        form.addRow("Revenue Multiplier:", self.multiplier_input)
        form.addRow("Group Funds (R$):", self.funds_input)
        form.addRow("Limited Items Value (R$):", self.limiteds_input)
        form.addRow("Member Count:", self.members_input)
        form.addRow("Value per Member (R$):", self.per_member_input)
        form.addRow("Custom Valuation (R$):", self.custom_input)
        layout.addLayout(form)
        self.calc_btn = QPushButton("Calculate")
        self.calc_btn.clicked.connect(self.calculate)
        layout.addWidget(self.calc_btn)
        self.result_label = QLabel("")
        layout.addWidget(self.result_label)
        self.setLayout(layout)
        self.update_method_fields()

    def update_method_fields(self):
        idx = self.method_combo.currentIndex()
        self.multiplier_input.setVisible(idx == 0)
        self.funds_input.setVisible(idx == 1)
        self.limiteds_input.setVisible(idx == 1)
        self.members_input.setVisible(idx == 2)
        self.per_member_input.setVisible(idx == 2)
        self.custom_input.setVisible(idx == 3)

    def calculate(self):
        method = self.method_combo.currentIndex()
        if method == 0:  # Revenue Multiplier
            monthly_revenue = self.pre_input.value()
            multiplier = self.multiplier_input.value()
            valuation = int(round(monthly_revenue * multiplier))
            method_str = f"Revenue Multiplier: {monthly_revenue} R$ x {multiplier} = {valuation} R$"
        elif method == 1:  # Asset Value
            funds = self.funds_input.value()
            limiteds = self.limiteds_input.value()
            valuation = funds + limiteds
            method_str = f"Asset Value: Funds {funds} R$ + Limiteds {limiteds} R$ = {valuation} R$"
        elif method == 2:  # Member Value
            members = self.members_input.value()
            per_member = self.per_member_input.value()
            valuation = members * per_member
            method_str = f"Member Value: {members} x {per_member} R$ = {valuation} R$"
        else:  # Custom
            valuation = self.custom_input.value()
            method_str = f"Custom Valuation: {valuation} R$"
        invest = self.invest_input.value()
        percent = self.percent_input.value()
        investor_percent = (invest / valuation * 100) if valuation else 0
        self.result_label.setText(
            f"Valuation: {valuation:,.0f} R$\n{method_str}\nInvestment: {invest:,.0f} R$\nInvestor %: {investor_percent:.2f}%"
        )

class DealFinderTab(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()
        self.plan_table = QTableWidget(2, 6)
        self.plan_table.setHorizontalHeaderLabels([
            "Plan Name", "Investment (R$)", "Equity %", "Group Activity (1-10)", "Members", "Monthly Revenue (R$)"
        ])
        layout.addWidget(self.plan_table)
        self.add_btn = QPushButton("+ Add Plan")
        self.add_btn.clicked.connect(self.add_plan_row)
        layout.addWidget(self.add_btn)
        self.compare_btn = QPushButton("Compare Plans")
        self.compare_btn.clicked.connect(self.compare)
        layout.addWidget(self.compare_btn)
        self.result_table = QTableWidget(0, 7)
        self.result_table.setHorizontalHeaderLabels([
            "Plan Name", "Investment (R$)", "Equity %", "Implied Valuation (R$)", "Activity", "Members", "Score"
        ])
        layout.addWidget(self.result_table)
        self.narrative_label = QLabel("")
        layout.addWidget(self.narrative_label)
        self.setLayout(layout)

    def add_plan_row(self):
        row = self.plan_table.rowCount()
        self.plan_table.insertRow(row)

    def compare(self):
        plans = []
        for row in range(self.plan_table.rowCount()):
            name = self.plan_table.item(row, 0)
            invest = self.plan_table.item(row, 1)
            percent = self.plan_table.item(row, 2)
            activity = self.plan_table.item(row, 3)
            members = self.plan_table.item(row, 4)
            revenue = self.plan_table.item(row, 5)
            if name and invest and percent:
                try:
                    plan = {
                        'name': name.text(),
                        'invest': float(invest.text()),
                        'percent': float(percent.text()),
                        'activity': int(activity.text()) if activity and activity.text().isdigit() else 0,
                        'members': int(members.text()) if members and members.text().isdigit() else 0,
                        'revenue': int(revenue.text()) if revenue and revenue.text().isdigit() else 0
                    }
                    plans.append(plan)
                except Exception:
                    continue
        results = []
        for plan in plans:
            post = plan['invest'] / (plan['percent'] / 100) if plan['percent'] else 0
            # Score: weighted sum (activity*2 + members/1000 + revenue/1000 + post/10000)
            score = plan['activity']*2 + plan['members']/1000 + plan['revenue']/1000 + post/10000
            results.append({
                'name': plan['name'],
                'invest': plan['invest'],
                'percent': plan['percent'],
                'post': post,
                'activity': plan['activity'],
                'members': plan['members'],
                'score': score
            })
        sorted_results = sorted(results, key=lambda x: x['score'], reverse=True)
        self.result_table.setRowCount(len(sorted_results))
        for i, r in enumerate(sorted_results):
            self.result_table.setItem(i, 0, QTableWidgetItem(r['name']))
            self.result_table.setItem(i, 1, QTableWidgetItem(f"{r['invest']:,.0f}"))
            self.result_table.setItem(i, 2, QTableWidgetItem(f"{r['percent']:.2f}%"))
            self.result_table.setItem(i, 3, QTableWidgetItem(f"{r['post']:,.0f}"))
            self.result_table.setItem(i, 4, QTableWidgetItem(str(r['activity'])))
            self.result_table.setItem(i, 5, QTableWidgetItem(str(r['members'])))
            self.result_table.setItem(i, 6, QTableWidgetItem(f"{r['score']:.2f}"))
        if sorted_results:
            best = sorted_results[0]
            self.narrative_label.setText(
                f"Best Plan: {best['name']} (Score: {best['score']:.2f}, Implied Valuation: {best['post']:,.0f} R$)"
            )
        else:
            self.narrative_label.setText("No valid plans entered.")

# --- New Investment Strategy Tabs ---

class EquityTab(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()
        label = QLabel("Equity Investment Calculator (Coming Soon)\n\nAnalyze direct equity investments in Roblox groups. Enter group valuation, investment amount, and equity % to see ownership, returns, and dilution.\n\n[Further logic and UI to be implemented.]")
        label.setWordWrap(True)
        layout.addWidget(label)
        self.setLayout(layout)

class AssetFlippingTab(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()
        label = QLabel("Asset Flipping Calculator (Coming Soon)\n\nModel buying and selling limiteds, group assets, or collectibles for profit. Enter buy/sell prices, fees, and holding period.\n\n[Further logic and UI to be implemented.]")
        label.setWordWrap(True)
        layout.addWidget(label)
        self.setLayout(layout)

class ProductTab(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()
        label = QLabel("Product Investment Calculator (Coming Soon)\n\nEstimate returns from investing in Roblox products (games, clothing, assets). Enter dev costs, expected sales, and revenue share.\n\n[Further logic and UI to be implemented.]")
        label.setWordWrap(True)
        layout.addWidget(label)
        self.setLayout(layout)

class ServiceTab(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()
        label = QLabel("Service Investment Calculator (Coming Soon)\n\nAnalyze investments in services (commissions, dev-for-hire, marketing). Enter contract terms, rates, and expected ROI.\n\n[Further logic and UI to be implemented.]")
        label.setWordWrap(True)
        layout.addWidget(label)
        self.setLayout(layout)

class KnowledgeIPTab(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()
        label = QLabel("Knowledge/IP Investment Calculator (Coming Soon)\n\nModel investments in scripts, systems, or IP. Enter licensing terms, royalties, and projected usage.\n\n[Further logic and UI to be implemented.]")
        label.setWordWrap(True)
        layout.addWidget(label)
        self.setLayout(layout)

class MarketingTab(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()
        label = QLabel("Marketing Investment Calculator (Coming Soon)\n\nEstimate ROI from group ads, sponsorships, and influencer campaigns. Enter spend, reach, and conversion rates.\n\n[Further logic and UI to be implemented.]")
        label.setWordWrap(True)
        layout.addWidget(label)
        self.setLayout(layout)

class PortfolioTab(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()
        label = QLabel("Portfolio Analysis (Coming Soon)\n\nTrack and analyze a portfolio of Roblox investments. Enter multiple holdings, values, and returns.\n\n[Further logic and UI to be implemented.]")
        label.setWordWrap(True)
        layout.addWidget(label)
        self.setLayout(layout)

class PrivateSilentTab(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()
        label = QLabel("Private/Silent Investment Calculator (Coming Soon)\n\nModel silent partnerships or private deals. Enter terms, expected returns, and risk factors.\n\n[Further logic and UI to be implemented.]")
        label.setWordWrap(True)
        layout.addWidget(label)
        self.setLayout(layout)

class AnalyticsTab(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()
        label = QLabel("Advanced Analytics & Arbitrage (Coming Soon)\n\nRun advanced analytics, arbitrage, and scenario modeling for Roblox investments.\n\n[Further logic and UI to be implemented.]")
        label.setWordWrap(True)
        layout.addWidget(label)
        self.setLayout(layout)

class InvestorGuideTab(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()
        guide = QTextEdit()
        guide.setReadOnly(True)
        guide.setHtml("""
        <h2>BloxVest Investor Guide</h2>
        <p>Welcome to the BloxVest Roblox Group Investment Suite. This guide covers all major investment strategies and how to use each calculator tab:</p>
        <ul>
            <li><b>Payout:</b> Calculate group payouts and after-tax values.</li>
            <li><b>Loan:</b> Model group loans and repayment schedules.</li>
            <li><b>ROI:</b> Analyze return on investment for any scenario.</li>
            <li><b>Valuation:</b> Estimate group value using multiple methods.</li>
            <li><b>Deal Finder:</b> Compare multiple investment deals.</li>
            <li><b>Equity:</b> Analyze direct equity investments in groups.</li>
            <li><b>Asset Flipping:</b> Model profits from buying/selling assets.</li>
            <li><b>Product:</b> Estimate returns from Roblox products.</li>
            <li><b>Service:</b> Analyze service/commission investments.</li>
            <li><b>Knowledge/IP:</b> Model IP and licensing deals.</li>
            <li><b>Marketing:</b> Estimate ROI from marketing spend.</li>
            <li><b>Portfolio:</b> Track and analyze a portfolio of investments.</li>
            <li><b>Private/Silent:</b> Model silent/private investment deals.</li>
            <li><b>Analytics:</b> Run advanced analytics and arbitrage tools.</li>
        </ul>
        <p>Click each tab for a specialized calculator and narrative summary. Use the Help menu for tutorials and security info. For more details, visit <a href='https://bloxvest.com/guide'>bloxvest.com/guide</a>.</p>
        """)
        layout.addWidget(guide)
        self.setLayout(layout)

# --- Main Application Window ---

class MainWindow(QMainWindow):
    def init_menu(self):
        menubar = self.menuBar()
        file_menu = menubar.addMenu("&File")
        export_action = QAction("&Export Report", self)
        export_action.triggered.connect(self.export_report)
        file_menu.addAction(export_action)
        save_action = QAction("&Save Scenario", self)
        save_action.triggered.connect(self.save_scenario)
        file_menu.addAction(save_action)
        load_action = QAction("&Load Scenario", self)
        load_action.triggered.connect(self.load_scenario)
        file_menu.addAction(load_action)
        exit_action = QAction("&Exit", self)
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)
        view_menu = menubar.addMenu("&View")
        theme_action = QAction("&Toggle Theme", self)
        theme_action.triggered.connect(self.toggle_theme)
        view_menu.addAction(theme_action)
        help_menu = menubar.addMenu("&Help")
        about_action = QAction("&About", self)
        about_action.triggered.connect(self.show_about)
        help_menu.addAction(about_action)
        howto_action = QAction("&How to Use", self)
        howto_action.triggered.connect(self.show_howto)
        help_menu.addAction(howto_action)
        tutorial_action = QAction("&Tutorial", self)
        tutorial_action.triggered.connect(self.show_tutorial)
        help_menu.addAction(tutorial_action)
        security_action = QAction("&Security Policy", self)
        security_action.triggered.connect(self.show_security)
        help_menu.addAction(security_action)

    def show_howto(self):
        QMessageBox.information(self, "How to Use", (
            "Welcome to BloxVest Group Investment Tools!\n\n"
            "1. Select a tab for the tool you need (Payout, Loan, ROI, etc).\n"
            "2. Fill in the required fields. Hover over any label for a tooltip.\n"
            "3. Click Calculate to see results.\n"
            "4. Use Export to save a report, or Save Scenario to keep your work.\n"
            "5. For detailed help, click the Tutorial button on each tab.\n\n"
            "All calculations are local and private."
        ))

    def show_tutorial(self):
        tab_idx = self.tabs.currentIndex()
        tab_name = self.tabs.tabText(tab_idx)
        tutorial_texts = {
            "Payout": (
                "Payout Tab Tutorial:\n\n"
                "- Enter the total Robux to distribute and the DevEx rate.\n"
                "- Add participants and set their % share.\n"
                "- Click Calculate to see each participant's payout, USD value, and tax.\n"
                "- Use 'Show After Payment' to see remaining group funds."
            ),
            "Loan": (
                "Loan Tab Tutorial:\n\n"
                "- Enter the loan principal, interest rate, and term.\n"
                "- Click Calculate for a full amortization schedule.\n"
                "- Review the warning about Roblox loan risks."
            ),
            "ROI": (
                "ROI Tab Tutorial:\n\n"
                "- Enter your initial and final investment, and the time frame.\n"
                "- Optionally, enter monthly revenue history for more accurate ROI.\n"
                "- Click Calculate to see ROI and annualized return."
            ),
            "Valuation": (
                "Valuation Tab Tutorial:\n\n"
                "- Choose a valuation method.\n"
                "- Fill in the relevant fields.\n"
                "- Click Calculate to see pre-money, post-money, and investor % equity."
            ),
            "Deal Finder": (
                "Deal Finder Tab Tutorial:\n\n"
                "- Enter multiple deal scenarios (one per row).\n"
                "- Click Compare Plans to see scores and the best deal."
            ),
            # New investment strategy tabs
            "Equity": (
                "Equity Tab Tutorial:\n\n"
                "- Use this tab to model equity investments in Roblox groups or projects.\n"
                "- Enter group/project valuation, investment amount, and desired equity split.\n"
                "- Click Calculate to see post-money valuation, investor share, and after-tax values.\n"
                "- Narrative and data modes available for client-friendly summaries."
            ),
            "Asset Flipping": (
                "Asset Flipping Tab Tutorial:\n\n"
                "- Model the purchase and resale of Roblox assets (e.g., limiteds, collectibles).\n"
                "- Enter buy/sell prices, fees, and holding period.\n"
                "- Calculate profit, ROI, and after-tax returns in Robux and USD.\n"
                "- Charts and narrative summaries included."
            ),
            "Product": (
                "Product Investment Tab Tutorial:\n\n"
                "- Analyze investments in Roblox products (games, plugins, UGC items).\n"
                "- Enter development costs, expected sales, and revenue share.\n"
                "- See break-even, profit, and ROI in Robux/USD/after-tax.\n"
                "- Narrative mode for client reports."
            ),
            "Service": (
                "Service Investment Tab Tutorial:\n\n"
                "- Model investments in Roblox services (commissions, dev-for-hire, etc).\n"
                "- Enter service cost, expected revenue, and time frame.\n"
                "- Calculate ROI, payback period, and after-tax values."
            ),
            "Knowledge/IP": (
                "Knowledge/IP Investment Tab Tutorial:\n\n"
                "- Evaluate investments in knowledge, scripts, or intellectual property.\n"
                "- Enter acquisition cost, expected licensing/sales, and time frame.\n"
                "- See projected returns and narrative summary."
            ),
            "Marketing": (
                "Marketing Investment Tab Tutorial:\n\n"
                "- Model marketing spend for Roblox groups/products.\n"
                "- Enter campaign cost, expected reach, and conversion rates.\n"
                "- Calculate cost per acquisition, ROI, and after-tax impact."
            ),
            "Portfolio": (
                "Portfolio Tab Tutorial:\n\n"
                "- Track and analyze a portfolio of Roblox investments.\n"
                "- Add multiple investments across strategies.\n"
                "- See aggregate returns, risk, and narrative summary."
            ),
            "Private/Silent": (
                "Private/Silent Investment Tab Tutorial:\n\n"
                "- Model private or silent investments in Roblox groups.\n"
                "- Enter terms, expected returns, and confidentiality options.\n"
                "- See after-tax and narrative outputs."
            ),
            "Analytics": (
                "Analytics Tab Tutorial:\n\n"
                "- Analyze investment performance, trends, and risk.\n"
                "- Use built-in charts and data mode for deep dives.\n"
                "- Narrative mode for executive summaries."
            ),
            "Investor Guide": (
                "Investor Guide Tab Tutorial:\n\n"
                "- Access a comprehensive guide to Roblox group investing.\n"
                "- Click links to jump to each calculator tab.\n"
                "- Use this as a reference for best practices and strategy."
            ),
        }
        text = tutorial_texts.get(tab_name, "Select a tab for context-specific help.")
        QMessageBox.information(self, f"{tab_name} Tutorial", text)

    def show_security(self):
        QMessageBox.information(self, "Security Policy", (
            "• All data stays local on your device.\n"
            "• No external connections are made.\n"
            "• Scenario files are encrypted (AES-256).\n"
            "• Download only from trusted sources.\n"
            "• Keep Python and your OS updated."
        ))

    def show_about(self):
        QMessageBox.about(self, "About", f"{APP_NAME}\nVersion {APP_VERSION}\n© 2024 BloxVest, Inc.")

    def show_onboarding(self):
        if not os.path.exists(".bloxvest_onboarded"):
            QMessageBox.information(self, "Welcome to BloxVest!", (
                "Thank you for choosing BloxVest Group Investment Tools.\n\n"
                "• Use the Help menu for tutorials and security info.\n"
                "• Hover over any label for a quick explanation.\n"
                "• Your data is always private.\n\n"
                "Happy investing!"
            ))
            with open(".bloxvest_onboarded", "w") as f:
                f.write("onboarded")

    def __init__(self):
        super().__init__()
        self.setWindowTitle(f"{APP_NAME} v{APP_VERSION}")
        self.setWindowIcon(QIcon(APP_ICON_PATH))
        self.resize(1200, 800)
        self.tabs = QTabWidget()
        self.tabs.addTab(PayoutTab(), "Payout")
        self.tabs.addTab(LoanTab(), "Loan")
        self.tabs.addTab(ROITab(), "ROI")
        self.tabs.addTab(ValuationTab(), "Valuation")
        self.tabs.addTab(DealFinderTab(), "Deal Finder")
        # --- New Tabs ---
        self.tabs.addTab(EquityTab(), "Equity")
        self.tabs.addTab(AssetFlippingTab(), "Asset Flipping")
        self.tabs.addTab(ProductTab(), "Product")
        self.tabs.addTab(ServiceTab(), "Service")
        self.tabs.addTab(KnowledgeIPTab(), "Knowledge/IP")
        self.tabs.addTab(MarketingTab(), "Marketing")
        self.tabs.addTab(PortfolioTab(), "Portfolio")
        self.tabs.addTab(PrivateSilentTab(), "Private/Silent")
        self.tabs.addTab(AnalyticsTab(), "Analytics")
        self.tabs.addTab(InvestorGuideTab(), "Investor Guide")
        self.setCentralWidget(self.tabs)
        self.init_menu()
        self.theme = "dark"
        self.apply_theme(self.theme)
        self.autosave_timer = QTimer(self)
        self.autosave_timer.timeout.connect(self.autosave)
        self.autosave_timer.start(AUTOSAVE_INTERVAL_MS)
        self.show_onboarding()

    def apply_theme(self, theme):
        if theme == "dark":
            palette = QPalette()
            palette.setColor(QPalette.Window, QColor(BRAND_DARK))
            palette.setColor(QPalette.WindowText, Qt.white)
            palette.setColor(QPalette.Base, QColor("#232323"))
            palette.setColor(QPalette.AlternateBase, QColor("#2d2d2d"))
            palette.setColor(QPalette.ToolTipBase, Qt.white)
            palette.setColor(QPalette.ToolTipText, Qt.white)
            palette.setColor(QPalette.Text, Qt.white)
            palette.setColor(QPalette.Button, QColor(BRAND_ORANGE))
            palette.setColor(QPalette.ButtonText, Qt.white)
            palette.setColor(QPalette.Highlight, QColor(BRAND_ORANGE))
            palette.setColor(QPalette.HighlightedText, Qt.black)
            self.setPalette(palette)
        else:
            self.setPalette(QApplication.style().standardPalette())

    def toggle_theme(self):
        self.theme = "light" if self.theme == "dark" else "dark"
        self.apply_theme(self.theme)

    def export_report(self):
        # --- Export current tab's results to Word/PDF with branding, cover, summary, table, chart ---
        from PySide6.QtWidgets import QDialog, QFormLayout, QLineEdit, QComboBox, QDialogButtonBox
        # 1. Prompt for Client Name, Scenario Name, Format
        class ExportDialog(QDialog):
            def __init__(self, parent=None):
                super().__init__(parent)
                self.setWindowTitle("Export Report")
                layout = QFormLayout(self)
                self.client_input = QLineEdit()
                self.scenario_input = QLineEdit()
                self.format_combo = QComboBox()
                self.format_combo.addItems(["Word (.docx)", "PDF (.pdf)"])
                layout.addRow("Client Name:", self.client_input)
                layout.addRow("Scenario Name:", self.scenario_input)
                layout.addRow("Export Format:", self.format_combo)
                buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
                buttons.accepted.connect(self.accept)
                buttons.rejected.connect(self.reject)
                layout.addWidget(buttons)
        
        dlg = ExportDialog(self)
        if dlg.exec() != QDialog.Accepted:
            return
        client_name = dlg.client_input.text().strip() or "Client"
        scenario_name = dlg.scenario_input.text().strip() or "Scenario"
        fmt = dlg.format_combo.currentIndex()  # 0=Word, 1=PDF
        # 2. Gather data from current tab
        tab = self.tabs.currentWidget()
        tab_name = self.tabs.tabText(self.tabs.currentIndex())
        summary, table_data, chart_type, chart_data = self.get_tab_export_data(tab, tab_name)
        # 3. Ask for file path
        ext = ".docx" if fmt == 0 else ".pdf"
        file_path, _ = QFileDialog.getSaveFileName(self, "Export Report", f"{client_name}_{scenario_name}{ext}",
                                                   "Word Document (*.docx);;PDF File (*.pdf)")
        if not file_path:
            return
        # 4. Generate chart image
        chart_img = None
        if chart_type and chart_data:
            if chart_type == "pie":
                chart_img = self.create_pie_chart(chart_data, file_path + "_chart.png")
            elif chart_type == "bar":
                chart_img = self.create_bar_chart(chart_data, file_path + "_chart.png")
            elif chart_type == "line":
                chart_img = self.create_line_chart(chart_data, file_path + "_chart.png")
        # 5. Export
        try:
            disclaimer = "This report is for Roblox group investments. All values are in Robux (R$). Estimates are for informational purposes only and may not reflect real-world investment returns. Roblox group investments are not protected by legal contracts."
            if fmt == 0:
                self.generate_word_report(file_path, client_name, scenario_name, tab_name, summary, table_data, chart_img, disclaimer)
            else:
                self.generate_pdf_report(file_path, client_name, scenario_name, tab_name, summary, table_data, chart_img, disclaimer)
            QMessageBox.information(self, "Export", f"Report exported successfully to:\n{file_path}")
        except Exception as e:
            QMessageBox.critical(self, "Export Failed", f"Error: {e}")

    def generate_word_report(self, file_path, client, scenario, tab_name, summary, table_data, chart_img, disclaimer):
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        # Cover page
        doc.add_picture(APP_ICON_PATH, width=Inches(1.2))
        p = doc.add_paragraph()
        run = p.add_run(f"BloxVest Group Investment Tools\n{tab_name} Report")
        run.font.size = Pt(24)
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Client: {client}", style='Intense Quote')
        doc.add_paragraph(f"Scenario: {scenario}")
        doc.add_paragraph(f"Date: {datetime.datetime.now().strftime('%Y-%m-%d')}")
        doc.add_page_break()
        # Summary
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(summary)
        # Table
        if table_data:
            doc.add_heading("Details", level=2)
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            for i, row in enumerate(table_data):
                for j, val in enumerate(row):
                    cell = table.cell(i, j)
                    cell.text = val
                    if i == 0:
                        cell.paragraphs[0].runs[0].font.bold = True
            table.style = 'Table Grid'
        # Chart
        if chart_img and os.path.exists(chart_img):
            doc.add_heading("Chart", level=2)
            doc.add_picture(chart_img, width=Inches(4.5))
        # Branding/footer
        doc.add_paragraph("\nReport generated by BloxVest Group Investment Tools.", style='Quote')
        doc.add_paragraph(disclaimer, style='Intense Quote')
        doc.save(file_path)
        if chart_img and os.path.exists(chart_img):
            os.remove(chart_img)

    def generate_pdf_report(self, file_path, client, scenario, tab_name, summary, table_data, chart_img, disclaimer):
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib import colors
        from reportlab.platypus import Table, TableStyle, Paragraph, SimpleDocTemplate, Spacer, Image
        from reportlab.lib.styles import getSampleStyleSheet
        doc = SimpleDocTemplate(file_path, pagesize=letter)
        elements = []
        styles = getSampleStyleSheet()
        # Cover page
        if os.path.exists(APP_ICON_PATH):
            elements.append(Image(APP_ICON_PATH, width=60, height=60))
        elements.append(Paragraph(f"<b>BloxVest Group Investment Tools</b><br/>{tab_name} Report", styles['Title']))
        elements.append(Spacer(1, 12))
        elements.append(Paragraph(f"<b>Client:</b> {client}", styles['Normal']))
        elements.append(Paragraph(f"<b>Scenario:</b> {scenario}", styles['Normal']))
        elements.append(Paragraph(f"<b>Date:</b> {datetime.datetime.now().strftime('%Y-%m-%d')}", styles['Normal']))
        elements.append(Spacer(1, 24))
        # Summary
        elements.append(Paragraph("<b>Summary</b>", styles['Heading1']))
        elements.append(Paragraph(summary, styles['Normal']))
        elements.append(Spacer(1, 12))
        # Table
        if table_data:
            elements.append(Paragraph("<b>Details</b>", styles['Heading2']))
            t = Table(table_data)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(BRAND_ORANGE)),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            elements.append(t)
            elements.append(Spacer(1, 12))
        # Chart
        if chart_img and os.path.exists(chart_img):
            elements.append(Paragraph("<b>Chart</b>", styles['Heading2']))
            elements.append(Image(chart_img, width=300, height=200))
        # Branding/footer
        elements.append(Spacer(1, 24))
        elements.append(Paragraph("Report generated by BloxVest Group Investment Tools.", styles['Italic']))
        elements.append(Paragraph(disclaimer, styles['Italic']))
        doc.build(elements)
        if chart_img and os.path.exists(chart_img):
            os.remove(chart_img)

    def get_scenario_data(self):
        # Gather all tab input and result data for saving
        data = {}
        # Payout Tab
        payout_tab = self.tabs.widget(0)
        payout = {
            'total': payout_tab.total_input.value(),
            'devex_rate': payout_tab.devex_rate_input.value(),
            'show_after': payout_tab.after_payment_checkbox.isChecked(),
            'participants': [
                {'name': name.text(), 'percent': percent.value()}
                for name, percent in payout_tab.participants
            ],
            'results': [
                [payout_tab.result_table.item(row, col).text() if payout_tab.result_table.item(row, col) else ''
                 for col in range(payout_tab.result_table.columnCount())]
                for row in range(payout_tab.result_table.rowCount())
            ],
            'results2': [
                [payout_tab.result_table2.item(row, col).text() if payout_tab.result_table2.item(row, col) else ''
                 for col in range(payout_tab.result_table2.columnCount())]
                for row in range(payout_tab.result_table2.rowCount())
            ]
        }
        data['payout'] = payout
        # Loan Tab
        loan_tab = self.tabs.widget(1)
        loan = {
            'principal': loan_tab.principal_input.value(),
            'rate': loan_tab.rate_input.value(),
            'months': loan_tab.term_input.value(),
            'results': [
                [loan_tab.result_table.item(row, col).text() if loan_tab.result_table.item(row, col) else ''
                 for col in range(loan_tab.result_table.columnCount())]
                for row in range(loan_tab.result_table.rowCount())
            ]
        }
        data['loan'] = loan
        # ROI Tab
        roi_tab = self.tabs.widget(2)
        roi = {
            'initial': roi_tab.initial_input.value(),
            'final': roi_tab.final_input.value(),
            'months': roi_tab.months_input.value(),
            'revenue_history': roi_tab.revenue_history.toPlainText(),
            'projected_revenue': roi_tab.projected_revenue.value(),
            'result': roi_tab.result_label.text()
        }
        data['roi'] = roi
        # Valuation Tab
        val_tab = self.tabs.widget(3)
        val = {
            'method': val_tab.method_combo.currentIndex(),
            'pre': val_tab.pre_input.value(),
            'invest': val_tab.invest_input.value(),
            'percent': val_tab.percent_input.value(),
            'multiplier': val_tab.multiplier_input.value(),
            'funds': val_tab.funds_input.value(),
            'limiteds': val_tab.limiteds_input.value(),
            'members': val_tab.members_input.value(),
            'per_member': val_tab.per_member_input.value(),
            'custom': val_tab.custom_input.value(),
            'result': val_tab.result_label.text()
        }
        data['valuation'] = val
        # Deal Finder Tab
        deal_tab = self.tabs.widget(4)
        deal = {
            'plans': [
                [deal_tab.plan_table.item(row, col).text() if deal_tab.plan_table.item(row, col) else ''
                 for col in range(deal_tab.plan_table.columnCount())]
                for row in range(deal_tab.plan_table.rowCount())
            ],
            'results': [
                [deal_tab.result_table.item(row, col).text() if deal_tab.result_table.item(row, col) else ''
                 for col in range(deal_tab.result_table.columnCount())]
                for row in range(deal_tab.result_table.rowCount())
            ],
            'narrative': deal_tab.narrative_label.text()
        }
        data['deal_finder'] = deal
        return data

    def set_scenario_data(self, data):
        # Restore all tab input/result data from dict
        # Payout Tab
        payout_tab = self.tabs.widget(0)
        payout = data.get('payout', {})
        payout_tab.total_input.setValue(payout.get('total', 0))
        payout_tab.devex_rate_input.setValue(payout.get('devex_rate', PayoutTab.DEVEX_RATE))
        payout_tab.after_payment_checkbox.setChecked(payout.get('show_after', False))
        # Remove extra participant rows
        while len(payout_tab.participants) > len(payout.get('participants', [])):
            row = payout_tab.participant_layout.takeAt(len(payout_tab.participants)-1)
            for i in reversed(range(row.count())):
                widget = row.itemAt(i).widget()
                if widget:
                    widget.deleteLater()
            payout_tab.participants.pop()
        # Add missing participant rows
        while len(payout_tab.participants) < len(payout.get('participants', [])):
            payout_tab.add_participant_row()
        for i, p in enumerate(payout.get('participants', [])):
            name, percent = payout_tab.participants[i]
            name.setText(p.get('name', ''))
            percent.setValue(p.get('percent', 0))
        payout_tab.calculate()
        # Loan Tab
        loan_tab = self.tabs.widget(1)
        loan_tab.principal_input.setValue(data.get('loan', {}).get('principal', 0))
        loan_tab.rate_input.setValue(data.get('loan', {}).get('rate', 0))
        loan_tab.term_input.setValue(data.get('loan', {}).get('months', 0))
        loan_tab.calculate()
        # ROI Tab
        roi_tab = self.tabs.widget(2)
        roi_tab.initial_input.setValue(data.get('roi', {}).get('initial', 0))
        roi_tab.final_input.setValue(data.get('roi', {}).get('final', 0))
        roi_tab.months_input.setValue(data.get('roi', {}).get('months', 0))
        roi_tab.revenue_history.setPlainText(data.get('roi', {}).get('revenue_history', ''))
        roi_tab.projected_revenue.setValue(data.get('roi', {}).get('projected_revenue', 0))
        roi_tab.calculate()
        # Valuation Tab
        val_tab = self.tabs.widget(3)
        val = data.get('valuation', {})
        val_tab.method_combo.setCurrentIndex(val.get('method', 0))
        val_tab.pre_input.setValue(val.get('pre', 0))
        val_tab.invest_input.setValue(val.get('invest', 0))
        val_tab.percent_input.setValue(val.get('percent', 0))
        val_tab.multiplier_input.setValue(val.get('multiplier', 8.0))
        val_tab.funds_input.setValue(val.get('funds', 0))
        val_tab.limiteds_input.setValue(val.get('limiteds', 0))
        val_tab.members_input.setValue(val.get('members', 0))
        val_tab.per_member_input.setValue(val.get('per_member', 1))
        val_tab.custom_input.setValue(val.get('custom', 0))
        val_tab.calculate()
        # Deal Finder Tab
        deal_tab = self.tabs.widget(4)
        # Plans
        plans = data.get('deal_finder', {}).get('plans', [])
        while deal_tab.plan_table.rowCount() > len(plans):
            deal_tab.plan_table.removeRow(deal_tab.plan_table.rowCount()-1)
        while deal_tab.plan_table.rowCount() < len(plans):
            deal_tab.plan_table.insertRow(deal_tab.plan_table.rowCount())
        for row, plan in enumerate(plans):
            for col, val in enumerate(plan):
                item = QTableWidgetItem(val)
                deal_tab.plan_table.setItem(row, col, item)
        deal_tab.compare()

    def save_scenario(self, file_path=None):
        if not file_path:
            file_path, _ = QFileDialog.getSaveFileName(self, "Save Scenario", "bloxvest_scenario.bvs", "BloxVest Scenario (*.bvs)")
            if not file_path:
                return
        data = self.get_scenario_data()
        enc = encrypt_data(data)
        with open(file_path, 'wb') as f:
            f.write(enc)
        QMessageBox.information(self, "Save Scenario", f"Scenario saved to:\n{file_path}")

    def load_scenario(self, file_path=None):
        if not file_path:
            file_path, _ = QFileDialog.getOpenFileName(self, "Load Scenario", "", "BloxVest Scenario (*.bvs)")
            if not file_path:
                return
        with open(file_path, 'rb') as f:
            enc = f.read()
        try:
            data = decrypt_data(enc)
            self.set_scenario_data(data)
            QMessageBox.information(self, "Load Scenario", f"Scenario loaded from:\n{file_path}")
        except Exception as e:
            QMessageBox.critical(self, "Load Failed", f"Error: {e}")

    def autosave(self):
        # Autosave to a default file in the app directory
        try:
            data = self.get_scenario_data()
            enc = encrypt_data(data)
            with open("autosave.bvs", 'wb') as f:
                f.write(enc)
        except Exception:
            pass

    def get_tab_export_data(self, tab, tab_name):
        # Returns: summary(str), table_data(list of lists), chart_type(str), chart_data(dict)
        if tab_name == "Payout":
            summary = tab.summary_label.text()
            table = tab.result_table
            table_data = [[table.horizontalHeaderItem(i).text() for i in range(table.columnCount())]]
            for row in range(table.rowCount()):
                table_data.append([table.item(row, col).text() if table.item(row, col) else '' for col in range(table.columnCount())])
            chart_type = "pie"
            # Pie chart: payout in Robux per participant
            chart_data = {table.item(row, 0).text(): float(table.item(row, 1).text().replace(' R$', '').replace(',', '')) for row in range(table.rowCount()) if table.item(row, 1)}
            return summary, table_data, chart_type, chart_data
        elif tab_name == "Loan":
            summary = tab.summary_label.text()
            table = tab.result_table
            table_data = [[table.horizontalHeaderItem(i).text() for i in range(table.columnCount())]]
            for row in range(table.rowCount()):
                table_data.append([table.item(row, col).text() if table.item(row, col) else '' for col in range(table.columnCount())])
            chart_type = "line"
            # Chart: Balance over time
            chart_data = {str(table.item(row, 0).text()): float(table.item(row, 4).text().replace(' R$', '').replace(',', '')) for row in range(table.rowCount()) if table.item(row, 4)}
            return summary, table_data, chart_type, chart_data
        elif tab_name == "ROI":
            summary = tab.result_label.text()
            table_data = [["Initial", "Final", "Months"]]
            table_data.append([str(tab.initial_input.value()), str(tab.final_input.value()), str(tab.months_input.value())])
            chart_type = None
            chart_data = None
            return summary, table_data, chart_type, chart_data
        elif tab_name == "Valuation":
            summary = tab.result_label.text()
            table_data = [["Pre-Money", "Investment", "Equity %"]]
            table_data.append([str(tab.pre_input.value()), str(tab.invest_input.value()), str(tab.percent_input.value())])
            chart_type = None
            chart_data = None
            return summary, table_data, chart_type, chart_data
        elif tab_name == "Deal Finder":
            summary = tab.narrative_label.text()
            table = tab.result_table
            table_data = [[table.horizontalHeaderItem(i).text() for i in range(table.columnCount())]]
            for row in range(table.rowCount()):
                table_data.append([table.item(row, col).text() if table.item(row, col) else '' for col in range(table.columnCount())])
            chart_type = "bar"
            # Chart: Plan Name vs Implied Valuation
            chart_data = {table.item(row, 0).text(): float(table.item(row, 3).text().replace(',', '')) for row in range(table.rowCount()) if table.item(row, 3)}
            return summary, table_data, chart_type, chart_data
        return "", [], None, None

def main():
    app = QApplication(sys.argv)
    app.setStyle(QStyleFactory.create("Fusion"))
    window = MainWindow()
    window.show()
    sys.exit(app.exec())

if __name__ == "__main__":
    main()

