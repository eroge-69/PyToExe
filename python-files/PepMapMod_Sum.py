# Updated GUI-Enabled Script
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import os
import numpy as np
import tkinter as tk
from tkinter import filedialog
from docx import Document
from docx.shared import Inches
import io

# --- Report Generation & Analysis Functions ---

def generate_overall_summary(df, abundance_cols, report):
    """
    Generates a high-level text summary of modification levels across all samples
    and adds it directly to the report object with proper bolding.
    """
    unique_mod_types = df['Modification Type'].unique()
    
    for mod_type in unique_mod_types:
        mod_data = df[df['Modification Type'] == mod_type].copy()
        
        has_considerable_level = (mod_data[abundance_cols] > 5).any().any()
        
        p = report.add_paragraph()
        p.add_run('• ')
        
        if has_considerable_level:
            if mod_type == 'Oxidation':
                total_mod_abundance = mod_data[abundance_cols].sum()
                highest_sample = total_mod_abundance.idxmax().replace('% Abundance ', '')
                lowest_sample = total_mod_abundance.idxmin().replace('% Abundance ', '')
                
                p.add_run("We observed considerable ").bold = True
                p.add_run(f"{mod_type}").bold = True
                p.add_run(f" modifications in some of the samples. (Highest in {highest_sample}, lowest in {lowest_sample})").bold = False
            else:
                p.add_run("Considerable ").bold = True
                p.add_run(f"{mod_type}").bold = True
                p.add_run(" modifications are observed across all tested samples.").bold = True
        else:
            p.add_run("Relatively low abundance of ").bold = True
            p.add_run(f"{mod_type}").bold = True
            p.add_run(" modifications are observed.").bold = True


def analyze_per_plot_data(mod_data, abundance_cols, report):
    """
    Analyzes a specific modification type's data and generates a detailed description
    which is then added to the report, with improved natural phrasing.
    """
    p = report.add_paragraph()
    
    # Part 1: Analyze and compare different modifications within the same type
    mod_data['Average Abundance'] = mod_data[abundance_cols].mean(axis=1)
    has_high_average_mod = (mod_data['Average Abundance'] > 5).any()
    
    highest_mod_text_added = False
    
    if has_high_average_mod:
        mod_data_high_avg = mod_data[mod_data['Average Abundance'] > 5].sort_values(by='Average Abundance', ascending=False)
        
        if len(mod_data_high_avg) >= 2:
            highest_mod_row = mod_data_high_avg.iloc[0]
            second_highest_mod_row = mod_data_high_avg.iloc[1]
            
            highest_avg = highest_mod_row['Average Abundance']
            second_highest_avg = second_highest_mod_row['Average Abundance']
            
            if highest_avg > second_highest_avg * 1.20:
                # Use the Modification name instead of the Residue #
                modification_name = highest_mod_row['Modification']
                mod_type = highest_mod_row['Modification Type']
                
                p.add_run("The modification of ").bold = False
                p.add_run(f"{modification_name}").bold = True
                p.add_run(f" shows the highest {mod_type} modification level, with an average abundance of {highest_avg:.2f}%. ").bold = False
                highest_mod_text_added = True

    # Part 2: Compare samples to a reference
    reference_col = [col for col in abundance_cols if 'Reference' in col or 'reference' in col]
    test_samples_cols = [col for col in abundance_cols if col not in reference_col]
    
    if reference_col and test_samples_cols:
        reference_col = reference_col[0]
        mod_type = mod_data['Modification Type'].iloc[0]
        
        test_samples_abundance = mod_data[test_samples_cols]
        reference_abundance = mod_data[reference_col]
        
        if not test_samples_abundance.empty and not reference_abundance.empty:
            test_median_abundance = test_samples_abundance.median().median()
            reference_median = reference_abundance.median()
            
            # Add a transition if Part 1 already has a description
            if highest_mod_text_added:
                p.add_run("Furthermore, when compared to the reference sample, ").bold = False
            else:
                p.add_run(f"When compared to the reference sample, ").bold = False
            
            # Add the comparison text with bold keywords
            if test_median_abundance > 5 and test_median_abundance > reference_median * 1.20:
                p.add_run(f"the tested samples show a ").bold = False
                p.add_run("higher").bold = True
                p.add_run(f" level of {mod_type} modifications.").bold = False
            elif test_median_abundance > 5 and test_median_abundance < reference_median * 0.80:
                p.add_run(f"the tested samples show a ").bold = False
                p.add_run("lower").bold = True
                p.add_run(f" level of {mod_type} modifications.").bold = False
            else:
                p.add_run(f"the tested samples show a ").bold = False
                p.add_run("similar").bold = True
                p.add_run(f" level of {mod_type} modifications.").bold = False


# --- Main Application Logic ---

def run_analysis():
    """
    Main function to be called by the GUI. It handles file selection,
    data processing, and report generation.
    """
    status_label.config(text="Please select one or more Excel files...", fg="black")
    root.update_idletasks()
    
    file_paths = filedialog.askopenfilenames(
        title="Select one or more Excel files",
        filetypes=[("Excel files", "*.xlsx;*.xls")]
    )

    if not file_paths:
        status_label.config(text="No files selected. Process aborted.", fg="red")
        return

    status_label.config(text="Processing files...", fg="blue")
    root.update_idletasks()

    try:
        all_processed_data = []

        for file_path in file_paths:
            df = pd.read_excel(file_path, skiprows=6)
            abundance_column = [col for col in df.columns if "% Abundance" in col]
            
            if not abundance_column:
                print(f"Skipping {file_path}: No '% Abundance' column found.")
                continue
            
            abundance_column = abundance_column[0]
            
            df = df[~df['Comment'].str.contains("Low confidence|Poor recovery|Detection artifact", na=False)].copy()
            df = df[~df['Category'].str.contains("Unknown Modification|Artifact", na=False)].copy()
            df = df[~df['Modification'].str.contains("~", na=False)].copy()
            df = df[df['Confidence'] >= 95].copy()
            df = df[df[abundance_column] >= 0.05].copy()

            df['Modification Type'] = df['Category'].apply(lambda x: x.split('+')[-1].split('-')[-1])

            df_for_merge = df[[
                'Protein', 'Residue #', 'Modification', 'Category', 
                'Peptides', 'Peptide Sequence', 'Modification Type', 
                abundance_column
            ]].copy()
            
            all_processed_data.append(df_for_merge)

        if not all_processed_data:
            status_label.config(text="No valid data to merge. Process aborted.", fg="red")
            return

        status_label.config(text="Merging data...", fg="blue")
        root.update_idletasks()
        
        merged_df = all_processed_data[0]
        for i in range(1, len(all_processed_data)):
            merged_df = pd.merge(
                merged_df, 
                all_processed_data[i], 
                on=['Protein', 'Residue #', 'Modification', 'Category', 'Peptides', 'Peptide Sequence', 'Modification Type'],
                how='outer'
            )

        abundance_columns = [col for col in merged_df.columns if "% Abundance" in col]
        merged_df[abundance_columns] = merged_df[abundance_columns].fillna(0)
        merged_df = merged_df.sort_values(by='Modification Type').copy()

        output_folder = os.path.dirname(file_paths[0])
        output_path_excel = os.path.join(output_folder, "Modification_Summary.xlsx")
        merged_df.to_excel(output_path_excel, index=False)

        status_label.config(text="Generating report...", fg="blue")
        root.update_idletasks()

        report = Document()
        report.add_heading('Peptide Mapping Modification Summary Report', level=1)
        report.add_heading('Overall Modification Level Summary', level=2)
        generate_overall_summary(merged_df.copy(), abundance_columns, report)
        
        report.add_heading('Detailed Modification Abundance Plots', level=2)
        report.add_paragraph(f"The full processed data is available in the Excel file: {os.path.basename(output_path_excel)}")

        plot_folder = os.path.join(output_folder, "plots")
        os.makedirs(plot_folder, exist_ok=True)
        
        unique_mod_types = merged_df['Modification Type'].unique()
        for mod_type in unique_mod_types:
            mod_data = merged_df[merged_df['Modification Type'] == mod_type].copy()
            if mod_data.empty:
                continue

            modifications = mod_data['Modification']
            abundance_data = mod_data[abundance_columns]
            
            plt.figure(figsize=(16, 10))
            x = np.arange(len(modifications))
            bar_width = 0.8 / len(abundance_columns)
            
            for i, col in enumerate(abundance_columns):
                plt.bar(x + i * bar_width, mod_data[col], width=bar_width, label=col.replace('% Abundance ', ''))
            
            plt.xticks(
                x + bar_width * (len(abundance_columns) - 1) / 2, 
                modifications, 
                rotation=45, 
                ha='right', 
                fontsize=12
            )
            
            plt.xlabel("Modification", fontsize=14)
            plt.ylabel("% Abundance", fontsize=14)
            plt.title(f"{mod_type} Abundance Comparison", fontsize=16)
            plt.legend(
                title="Sample", 
                bbox_to_anchor=(1.05, 1), 
                loc='upper left', 
                fontsize=10, 
                title_fontsize=12
            )
            
            plot_path = os.path.join(plot_folder, f"{mod_type}_abundance_comparison.png")
            plt.tight_layout(rect=[0, 0, 0.95, 1])
            plt.savefig(plot_path, dpi=300)
            plt.close()
            
            report.add_heading(f"{mod_type} Abundance", level=3)
            analyze_per_plot_data(mod_data, abundance_columns, report)
            report.add_picture(plot_path, width=Inches(6))

        report_output_path = os.path.join(output_folder, "Peptide_Mapping_Modification_Summary.docx")
        report.save(report_output_path)
        
        status_label.config(text=f"Report saved to: {report_output_path}", fg="green")

    except Exception as e:
        status_label.config(text=f"An error occurred: {e}", fg="red")

# --- GUI Setup ---

root = tk.Tk()
root.title("Peptide Mapping Reporter")

frame = tk.Frame(root, padx=20, pady=20)
frame.pack(padx=10, pady=10)

title_label = tk.Label(frame, text="Peptide Mapping Report Generator", font=("Arial", 16, "bold"))
title_label.pack(pady=10)

info_label = tk.Label(frame, text="Click the button to select your Excel files and generate the report.")
info_label.pack(pady=5)

run_button = tk.Button(frame, text="Select Files & Generate Report", command=run_analysis)
run_button.pack(pady=20)

status_label = tk.Label(frame, text="Ready", fg="black")
status_label.pack(pady=5)

# Start the GUI event loop
root.mainloop()