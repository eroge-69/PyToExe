import sys
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from bs4 import BeautifulSoup

def convert_html_to_docx(html_path, logo_path):
    # Load HTML content
    with open(html_path, 'r', encoding='utf-8') as f:
        soup = BeautifulSoup(f, 'html.parser')

    # Create Word document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # Insert logo
    if os.path.exists(logo_path):
        doc.add_picture(logo_path, width=Inches(1.5))

    # Add "FOR INTERNAL USE ONLY" aligned right
    p = doc.add_paragraph("FOR INTERNAL USE ONLY")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p.runs[0].bold = True

    # Add instruction table
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.autofit = True
    table.allow_autofit = True

    # Instruction row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].merge(hdr_cells[1])
    hdr_cells[0].text = "Please fill in the following fields:"
    hdr_cells[0].paragraphs[0].runs[0].bold = True

    # Sample metadata row
    row_cells = table.rows[1].cells
    row_cells[0].text = "Organization:"
    row_cells[1].text = "Bühler Group"

    # Add title
    doc.add_paragraph("Product documentation Models, drawings, bills of materials and associated documents", style='Heading 1')

    # Add contents
    headings = soup.find_all(['h1', 'h2', 'h3'])
    if headings:
        doc.add_paragraph("Contents", style='Heading 2')
        for h in headings:
            doc.add_paragraph(h.get_text(), style='Normal')

    # Add purpose
    doc.add_paragraph("Purpose", style='Heading 2')
    doc.add_paragraph("This document provides product documentation including models, drawings, bills of materials, and associated documents.", style='Normal')

    # Add HTML content
    for tag in soup.body.descendants:
        if tag.name == 'p':
            doc.add_paragraph(tag.get_text(), style='Normal')
        elif tag.name in ['h1', 'h2', 'h3']:
            doc.add_paragraph(tag.get_text(), style='Heading 2')

    # Save document
    base = os.path.splitext(os.path.basename(html_path))[0]
    output_path = os.path.join(os.path.dirname(html_path), f"{base}_converted.docx")
    doc.save(output_path)
    print(f"Document saved to: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Drag and drop an HTML file onto this executable.")
        sys.exit(1)
    html_file = sys.argv[1]
    logo_file = os.path.join(os.path.dirname(sys.argv[0]), "UploadedImage1.jpg")
    convert_html_to_docx(html_file, logo_file)
