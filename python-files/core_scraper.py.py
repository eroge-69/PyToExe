import requests
from bs4 import BeautifulSoup
import datetime
import feedparser
from docx import Document
from docx.shared import Pt
from dateutil.parser import parse as parse_date
import unicodedata
import os # We'll use this to get the user's desktop path

# --- All your global variables and functions remain the same ---

# Define keywords (mix of English and Urdu) globally
keywords = ['Karachi', 'Sindh','Crime','Hyderabad','Landhi','MQM-P', 'MQM-L', 'Landhi','Korangi','CPLC',
            'Creek', 'DHA', 'Pakistan Rangers (Sindh)', 'Hub', 'Katcha Area', 'Sea View', 'Korangi', 'SITE',
            'Super Highway','Northern Bypass', 'Hawks Bay',
            'کراچی', 'سندھ', 'کچھ', 'سکھر', 'حیدرآباد', 'ٹنڈو محمد خان', 'ٹنڈو الہ یار',
            'مٹیاری', 'جامشورو', 'ٹھٹھہ', 'سجاول', 'بدین', 'میرپور خاص', 'عمرکوٹ',
            'تھرپارکر', 'سانگھڑ', 'شہید بے نظیر آباد', 'نوشہرو فیروز', 'دادو',
            'لاڑکانہ', 'قمبر شہدادکوٹ', 'شکارپور', 'جیکب آباد', 'کشمور', 'گھوٹکی', 'خیرپور']

urdu_keywords = ['کراچی', 'سندھ', 'کچھ', 'سکھر', 'حیدرآباد', 'ٹنڈو محمد خان', 'ٹنڈو الہ یار',
                 'مٹیاری', 'جامشورو', 'ٹھٹھہ', 'سجاول', 'بدin', 'میرپور خاص', 'عمرکوٹ',
                 'تھرپارکر', 'سانگھڑ', 'شہید بے نظیر آباد', 'نوشہرو فیروز', 'دادو',
                 'لاڑکانہ', 'قمبر شہدادکوٹ', 'شکارپور', 'جیکب آباد', 'کشمور', 'گھوٹکی', 'خیرپور']

english_rss_urls = [
    'https://www.dawn.com/feed/sindh',
    'https://www.pakistantoday.com.pk/feeds/city/karachi/',
    'https://www.samaaenglish.tv/rss-feed.xml',
    'https://arynews.tv/feed/sindh',
    'https://arynews.tv/feed/Karachi',
    'https://www.dawn.com/feeds/tag/karachi.rss',
    'https://www.thenews.com.pk/rss/2/11',
    'https://tribune.com.pk/feed/tag/karachi',
    'https://dailytimes.com.pk/feed/?s=karachi',
    'https://www.pakistantoday.com.pk/feed/?s=karachi',
    'https://www.brecorder.com/rss?section=karachi',
    'https://www.dawn.com/feeds/tag/sindh-police.rss',
    'https://www.dawn.com/feeds/tag/sindh-government.rss',
    'https://tribune.com.pk/feed/tag/sindh-assembly',
    'https://tribune.com.pk/feed/tag/interior-sindh',
    'https://www.thenews.com.pk/rss/2/12',
    'https://www.thenews.com.pk/rss/2/27',
    'https://www.dawn.com/feeds/tag/sindh.rss',
    'https://tribune.com.pk/feed/tag/sindh',
    'https://www.thenews.com.pk/rss/2/22',
    'https://dailytimes.com.pk/feed/?s=sindh',
    'https://www.pakistantoday.com.pk/feed/?s=sindh',
    'https://www.brecorder.com/rss?section=sindh',
    'https://www.urdupoint.com/en/rss/rss.aspx?site=latest-news/sindh',
    'https://www.urdupoint.com/en/rss/rss.aspx?site=latest-news/karachi',
    'https://www.urdupoint.com/en/rss/rss.aspx?site=latest-news&q=katcha',
    'https://en.dailypakistan.com.pk/rss/national',
    'https://www.geo.tv/rss/1/6',
    'https://www.geo.tv/rss/1/13',
    'https://92newshd.tv/feed/?s=sindh',
    'https://92newshd.tv/feed/?s=karachi',
    'https://www.bolnews.com/feed/?s=karachi',
    'https://www.bolnews.com/feed/?s=sindh',
    'https://nation.com.pk/rss/national',
    'https://www.urdupoint.com/en/rss/rss.aspx?site=latest-news/crime',
    'https://www.urdupoint.com/en/rss/rss.aspx?site=latest-news/pakistan/sindh/larkana',
    'https://www.urdupoint.com/en/rss/rss.aspx?site=latest-news/pakistan/sindh/sukkur',
    'https://www.urdupoint.com/en/rss/rss.aspx?site=latest-news/pakistan/sindh/hyderabad'
]

urdu_rss_urls = [
    'https://jang.com.pk/rss/feed/3',
    'https://www.geo.tv/rss/1/20',
    'https://dunya.com.pk/ur/sindh/rss.xml',
    'https://www.express.pk/feed/',
    'https://www.urdupoint.com/daily/livenews/rss.xml',
    'https://www.bbc.com/urdu/institutional/2009/03/090306_rss_feed',
    'https://www.qaumiawaz.com/stories.rss',
    'https://www.qaumiawaz.com/stories.rss?section=news',
    'https://www.urduvoa.com/rssfeeds',
    'https://www.suchtv.pk/urdu/latestnews?format=feed&type=rss',
    'https://tgstat.com/channel/@BOLNetworkOfficial/rss',
    'https://tgstat.com/channel/@ExpressNewsLive/rss',
    'https://tgstat.com/channel/@samaatv/rss',
    'https://tgstat.com/channel/@ARYNewsLive/rss',
    'https://tgstat.com/channel/@GeoNewsOfficial/rss'
]

# Function to clean HTML content
def clean_html_content(html_content):
    cleaned_text = ""
    if html_content:
        if isinstance(html_content, list):
            for item in html_content:
                if isinstance(item, dict) and 'value' in item:
                    cleaned_text += item['value'] + " "
                else:
                    cleaned_text += str(item) + " "
            cleaned_text = cleaned_text.strip()
        elif isinstance(html_content, str):
            cleaned_text = html_content

        if cleaned_text:
            cleaned_text = unicodedata.normalize('NFKC', cleaned_text)
            soup = BeautifulSoup(cleaned_text, 'html.parser')
            text_content = soup.get_text(separator=' ', strip=True)
            text_content = ' '.join(text_content.split())
            return text_content
    return ""

# Function to scrape news from RSS feeds
# We add a 'log_callback' to send messages to the GUI
def scrape_news_from_rss(all_keywords, rss_urls, urdu_rss_urls, log_callback):
    news_items = []
    now = datetime.datetime.now(datetime.timezone.utc)
    english_keywords = [k for k in all_keywords if k not in urdu_keywords]

    total_feeds = len(rss_urls)
    for i, url in enumerate(rss_urls):
        # This is how we send updates to the GUI
        log_callback(f"Processing feed {i+1}/{total_feeds}: {url}")
        
        is_urdu = url in urdu_rss_urls
        try:
            feed = feedparser.parse(url)
            for entry in feed.entries:
                published_time = None
                if hasattr(entry, 'published_parsed'):
                    published_time = datetime.datetime(*entry.published_parsed[:6], tzinfo=datetime.timezone.utc)
                elif hasattr(entry, 'published'):
                    try:
                        published_time = parse_date(entry.published)
                        if published_time.tzinfo is None:
                            published_time = published_time.replace(tzinfo=datetime.timezone.utc)
                    except ValueError:
                        continue
                
                if published_time is None or (now - published_time).total_seconds() <= 24 * 3600:
                    title = entry.get('title', 'No title')
                    content_html = entry.get('content', entry.get('summary', ''))
                    cleaned_content = clean_html_content(content_html)
                    text_to_check = title + ' ' + cleaned_content
                    matches_keywords = False
                    
                    if is_urdu:
                        matches_keywords = any(keyword in text_to_check for keyword in urdu_keywords) or \
                                          any(keyword.lower() in text_to_check.lower() for keyword in english_keywords)
                    else:
                        matches_keywords = any(keyword.lower() in text_to_check.lower() for keyword in english_keywords)
                    
                    general_keywords_check = 'karachi' in text_to_check.lower() or 'sindh' in text_to_check.lower()
                    
                    if (matches_keywords or general_keywords_check) and (title.strip() or cleaned_content.strip()):
                        news_items.append({
                            'source': url,
                            'title': title,
                            'link': entry.get('link', 'No link'),
                            'published': published_time if published_time else datetime.datetime.min.replace(tzinfo=datetime.timezone.utc),
                            'content': cleaned_content,
                            'is_urdu': is_urdu
                        })
        except Exception as e:
            log_callback(f"Error processing RSS feed {url}: {e}")

    return news_items

# Function to create the Word document
def create_word_document(news_data, filename="Karachi_Sindh_News.docx"):
    document = Document()
    document.add_heading('Executive Summary', 0)
    
    if news_data:
        document.add_paragraph("The following news items were published related to Karachi and Sindh in the last 24 hours:")
        for i, news_item in enumerate(news_data):
            lang = "Urdu" if news_item['is_urdu'] else "English"
            document.add_paragraph(f"{i+1}. {news_item['title']} ({news_item['source']} - {lang})")
    else:
        document.add_paragraph("No news found for Karachi and Sindh from the provided RSS feeds in the last 24 hours.")

    document.add_page_break()
    document.add_heading('Karachi and Sindh News (Last 24 Hours)', 0)

    if news_data:
        for news_item in news_data:
            document.add_heading(news_item['title'], level=1)
            document.add_paragraph(f"Source: {news_item['source']}")
            published_str = news_item['published'].strftime('%Y-%m-%d %H:%M:%S UTC') if news_item['published'] > datetime.datetime.min.replace(tzinfo=datetime.timezone.utc) else "N/A"
            document.add_paragraph(f"Published: {published_str}")
            document.add_paragraph(f"Link: {news_item['link']}")
            document.add_paragraph(news_item['content'])
            document.add_paragraph("-" * 50)
    else:
        document.add_paragraph("No detailed news items available.")

    section = document.sections[0]
    if not section.footer.paragraphs:
        section.footer.add_paragraph("Algorithm based news generation - BSG")
    else:
        section.footer.paragraphs[0].text = "Algorithm based news generation - BSG"
    
    # Save the file to the user's Desktop for easy access
    desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop') 
    output_path = os.path.join(desktop, filename)

    document.save(output_path)
    return output_path

# This is the main function that the GUI will call.
def run_scraper(log_callback):
    try:
        log_callback("Starting news scraping...")
        all_rss_urls = english_rss_urls + urdu_rss_urls
        relevant_news = scrape_news_from_rss(keywords, all_rss_urls, urdu_rss_urls, log_callback)
        
        log_callback(f"Finished scraping. Found {len(relevant_news)} relevant news items.")

        if not relevant_news:
            log_callback("No relevant news found. Document will not be created.")
            return

        relevant_news.sort(key=lambda x: (not x['is_urdu'], x['published']), reverse=True)

        log_callback("Creating Word document...")
        
        # Add a timestamp to the filename to avoid overwriting
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"Karachi_Sindh_News_{timestamp}.docx"
        
        output_path = create_word_document(relevant_news, filename)
        
        log_callback("\n" + "="*50)
        log_callback(f"SUCCESS! Word document created.")
        log_callback(f"File saved to: {output_path}")
        log_callback("="*50)
    except Exception as e:
        log_callback(f"\nAN ERROR OCCURRED: {e}")