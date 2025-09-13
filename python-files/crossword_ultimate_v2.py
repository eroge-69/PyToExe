#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
crossword_gui_ultimate_v2.py

Adds:
- Rotational block symmetry (toggle) + post-generation symmetrize
- Full Across/Down clue columns in GUI + DOCX/PDF (NYT-style numbering)
- Number font color + bold/italic/underline
- Manual clue editing + Renumber
- Multi-language: RTL shaping, RTL across layout, Vertical TTB (CJK), localized digits
- Light/Dark theme toggle
- Keyboard shortcuts: Ctrl+S (export), Ctrl+Z (undo), Ctrl+R (regenerate), Ctrl+T (theme)
- Cell background color for open cells
- Puzzle numbering for exports

Base author(s): You & ChatGPT
"""

import random
import string
import math
import json
from typing import List, Tuple, Dict, Optional, Set
import os
import csv
import re
import sys
import copy

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, colorchooser

from PIL import Image, ImageDraw, ImageFont, ImageTk

# Optional RTL shaping
try:
    import arabic_reshaper
    from bidi.algorithm import get_display as bidi_get_display
    HAS_RTL = True
except Exception:
    HAS_RTL = False

# =================== BASIC CONFIG ===================
FILL_ALPHABET = string.ascii_uppercase
DEFAULT_CELL_PX = 48
DEFAULT_MARGIN_CELLS = 1
MIN_PREVIEW = 360
MAX_PREVIEW = 900

WORD_BANK = [
    "PYTHON","ALGORITHM","SEARCH","GRID","PUZZLE","CODE","LOGIC","FUNCTION","VARIABLE","LOOP",
    "APPLE","BANANA","ORANGE","GRAPE","PEAR","MANGO","PEACH","CHERRY","LEMON","PAPAYA",
    "PARIS","CAIRO","TOKYO","LONDON","BERLIN","MADRID","DUBAI","DELHI","NAIROBI","SYDNEY",
    "MOUNTAIN","RIVER","OCEAN","FOREST","DESERT","ISLAND","VALLEY","CANYON","PRAIRIE","GLACIER"
]

# New: digit scripts for Unicode-aware numbering
DIGIT_MAPS = {
    "Latin": "0123456789",
    "Arabic-Indic": "٠١٢٣٤٥٦٧٨٩",
    "Persian": "۰۱۲۳۴۵۶۷۸۹",
}
DIGIT_KEYS = list(DIGIT_MAPS.keys())

# New: grid orientation
ORIENTATIONS = ["LTR (Across →)", "RTL (Across ←)", "Vertical TTB (CJK)"]

# =================== UTILITIES ===================
def shape_if_needed(s: str) -> str:
    if not HAS_RTL: return s
    try:
        reshaped = arabic_reshaper.reshape(s)
        return bidi_get_display(reshaped)
    except Exception:
        return s

def localize_digits(n: int, script: str) -> str:
    ds = DIGIT_MAPS.get(script, DIGIT_MAPS["Latin"])
    return "".join(ds[int(d)] for d in str(n))

def hex_to_rgb_tuple(hex_str: str) -> Tuple[int, int, int]:
    h = hex_str.strip().lstrip("#")
    if len(h) == 3: h = "".join(c*2 for c in h)
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

def hex_to_reportlab_rgb01(hex_str: str) -> Tuple[float, float, float]:
    r, g, b = hex_to_rgb_tuple(hex_str)
    return (r/255.0, g/255.0, b/255.0)

def hex_to_docx_fill(hex_str: str) -> str:
    return hex_str.strip().lstrip("#").upper()

def blend_with_white(rgb01: Tuple[float, float, float], alpha01: float) -> Tuple[float, float, float]:
    r, g, b = rgb01; a = max(0.0, min(1.0, alpha01))
    return (a*r + (1-a)*1.0, a*g + (1-a)*1.0, a*b + (1-a)*1.0)

def auto_palette(n: int) -> List[str]:
    if n <= 0: return []
    cols = []
    for k in range(n):
        h = k / max(1, n)
        s = 0.55; v = 0.95
        r, g, b = hsv_to_rgb(h, s, v)
        cols.append("#%02X%02X%02X" % (int(r*255), int(g*255), int(b*255)))
    return cols

def hsv_to_rgb(h, s, v):
    i = int(h*6); f = h*6 - i
    p = v*(1-s); q = v*(1-f*s); t = v*(1-(1-f)*s)
    i = i % 6
    if i==0: r,g,b=v,t,p
    elif i==1: r,g,b=q,v,p
    elif i==2: r,g,b=p,v,t
    elif i==3: r,g,b=p,q,v
    elif i==4: r,g,b=t,p,v
    else: r,g,b=v,p,q
    return r,g,b

def parse_palette_str(palette: str, keys: List[str]) -> Dict[str, str]:
    if not palette.strip(): return {}
    tokens = [t.strip() for t in palette.split(";") if t.strip()]
    mapping = {}
    hexes: List[str] = []
    for tok in tokens:
        if "=" in tok:
            w, col = tok.split("=", 1)
            mapping[w.strip().upper()] = col.strip()
        else:
            hexes.append(tok)
    if hexes:
        for i, k in enumerate(keys):
            mapping.setdefault(k.strip().upper(), hexes[i % len(hexes)])
    return mapping

def build_cell_color_map(entry_cells: Dict[str, List[Tuple[int,int]]],
                         per_key_colors: Dict[str, str],
                         default_hex: Optional[str] = None) -> Dict[Tuple[int,int], str]:
    cell_map: Dict[Tuple[int,int], str] = {}
    for key, coords in entry_cells.items():
        c = per_key_colors.get(key.upper(), default_hex)
        if not c: continue
        for pos in coords:
            cell_map.setdefault(pos, c)
    return cell_map

# =================== GRID SIZE HEURISTICS (unchanged) ===================
def minimal_size(words: List[str], padding: int = 2) -> int:
    if not words: return 10
    longest = max(len(w) for w in words)
    area_target = sum(len(w) for w in words) * 2
    side = max(longest, int(math.sqrt(area_target)) + padding)
    return max(side, 5)

def compute_size_by_density(words: List[str], density_pct: int, min_side: int = 5, padding: int = 0) -> int:
    clean = [w.replace(" ", "") for w in words if w.strip()]
    if not clean: return min_side
    L = sum(len(w) for w in clean)
    longest = max(len(w) for w in clean)
    d = max(30, min(95, int(round(density_pct))))
    area = max(L, int(math.ceil(L / (d / 100.0))))
    side = max(longest, int(math.ceil(math.sqrt(area))) + padding)
    return max(min_side, side)

def compute_size_by_cpw(words: List[str], cpw: int, min_side: int = 5, padding: int = 0) -> int:
    clean = [w.replace(" ", "") for w in words if w.strip()]
    if not clean: return min_side
    num = len(clean)
    longest = max(len(w) for w in clean)
    cpw = max(4, int(cpw))
    area = max(num, cpw * num)
    side = max(longest, int(math.ceil(math.sqrt(area))) + padding)
    return max(min_side, side)

def compute_size_by_cpl(words: List[str], cpl: float, min_side: int = 5, padding: int = 0) -> int:
    clean = [w.replace(" ", "") for w in words if w.strip()]
    if not clean: return min_side
    L = sum(len(w) for w in clean)
    longest = max(len(w) for w in clean)
    cpl = max(1.2, float(cpl))
    area = max(L, int(math.ceil(L * cpl)))
    side = max(longest, int(math.ceil(math.sqrt(area))) + padding)
    return max(min_side, side)

# =================== CROSSWORD GENERATION (unchanged core) ===================
Placement = Tuple[int, int, str, bool]  # (row, col, dir, reversed)

def new_grid(n: int):
    return [[None for _ in range(n)] for _ in range(n)]

def in_bounds(n: int, r: int, c: int) -> bool:
    return 0 <= r < n and 0 <= c < n

def can_place_word(grid, word: str, r: int, c: int, dir_: str, reversed_ok: bool) -> bool:
    n = len(grid)
    dr, dc = (0,1) if dir_ == 'A' else (1,0)
    L = len(word)
    rr, cc = r + dr*(L-1), c + dc*(L-1)
    if not in_bounds(n, rr, cc): return False
    for i, ch in enumerate(word):
        y, x = r + dr*i, c + dc*i
        cell = grid[y][x]
        if cell is None:
            continue
        if cell != ch:
            return False
    before_r, before_c = r - dr, c - dc
    after_r, after_c = r + dr*L, c + dc*L
    if in_bounds(n, before_r, before_c) and grid[before_r][before_c] not in (None,):
        return False
    if in_bounds(n, after_r, after_c) and grid[after_r][after_c] not in (None,):
        return False
    return True

def overlaps_count(grid, word: str, r: int, c: int, dir_: str) -> int:
    n = len(grid); dr, dc = (0,1) if dir_ == 'A' else (1,0)
    cnt = 0
    for i, ch in enumerate(word):
        y, x = r + dr*i, c + dc*i
        if in_bounds(n, y, x) and grid[y][x] == ch:
            cnt += 1
    return cnt

def place_word(grid, word: str, r: int, c: int, dir_: str) -> List[Tuple[int,int]]:
    n = len(grid); dr, dc = (0,1) if dir_ == 'A' else (1,0)
    written: List[Tuple[int,int]] = []
    for i, ch in enumerate(word):
        y, x = r + dr*i, c + dc*i
        if grid[y][x] is None:
            grid[y][x] = ch
            written.append((y, x))
    return written

def undo_writes(grid, written: List[Tuple[int,int]]):
    for (y, x) in written:
        grid[y][x] = None

def candidate_positions(grid, word: str) -> List[Placement]:
    n = len(grid)
    L = len(word)
    out: List[Placement] = []
    # Across
    for r in range(n):
        for c in range(n - L + 1):
            if can_place_word(grid, word, r, c, 'A', False):
                out.append((r, c, 'A', False))
    # Down
    for c in range(n):
        for r in range(n - L + 1):
            if can_place_word(grid, word, r, c, 'D', False):
                out.append((r, c, 'D', False))
    return out

def build_crossword(words: List[str],
                    size: int,
                    allow_reversed: bool = False,
                    require_connected: bool = True,
                    min_overlap: int = 1,
                    max_backtracks: int = 20000,
                    rng: Optional[random.Random] = None):
    rng = rng or random.Random()
    words = [w.strip().replace(" ", "").upper() for w in words if w.strip()]
    if not words:
        raise ValueError("No valid words were provided.")
    seen = set(); clean = []
    for w in words:
        if w not in seen:
            seen.add(w); clean.append(w)
    words = sorted(clean, key=len, reverse=True)

    grid = new_grid(size)
    placed_entries: List[Dict] = []

    def selection_order(remaining: List[str]) -> List[str]:
        scores = []
        for w in remaining:
            cands = candidate_positions(grid, w)
            scores.append((len(cands), -len(w), rng.random(), w))
        scores.sort()
        return [w for _,__,___, w in scores]

    backtracks = 0

    def rec(left: List[str]) -> bool:
        nonlocal backtracks
        if not left:
            return True
        order = selection_order(left)
        w = order[0]
        rest = [x for x in left if x != w]

        cands = candidate_positions(grid, w)
        if require_connected and placed_entries:
            cands_scored = []
            for (r, c, d, rev) in cands:
                ov = overlaps_count(grid, w, r, c, d)
                cands_scored.append(( -ov, rng.random(), (r,c,d,rev) ))
            cands_scored.sort()
            cands = [t[2] for t in cands_scored]
        else:
            rng.shuffle(cands)

        for (r, c, d, rev) in cands:
            ov = overlaps_count(grid, w, r, c, d)
            if require_connected and placed_entries and ov < max(0, int(min_overlap)):
                continue
            written = place_word(grid, w, r, c, d)
            placed_entries.append(
                {"answer": w, "dir": d, "r": r, "c": c,
                 "coords": [(r + (0 if d=='A' else i), c + (i if d=='A' else 0)) for i in range(len(w))]}
            )
            if rec(rest):
                return True
            placed_entries.pop()
            undo_writes(grid, written)
            backtracks += 1
            if backtracks > max_backtracks:
                return False
        return False

    if not rec(words):
        raise RuntimeError("Could not place all words into the grid with current settings.")

    return grid, placed_entries

# =================== NUMBERING & CLUES (NYT-style) ===================
def compute_numbering_nyt(grid, entries: List[Dict]) -> Tuple[Dict[Tuple[int,int], int], Dict[str, Dict]]:
    """
    NYT-style: scan row-major; a non-block cell gets a number if it starts an Across or Down entry.
    Across start: cell not block and (col==0 or left is block) and right is letter
    Down   start: cell not block and (row==0 or above is block) and below is letter
    A cell number is shared by both an Across and a Down starter if both start there.
    """
    n = len(grid)
    cell_number: Dict[Tuple[int,int], int] = {}
    num = 1
    # mark starts
    for r in range(n):
        for c in range(n):
            if grid[r][c] is None: 
                continue
            start_across = (c == 0 or grid[r][c-1] is None) and (c+1 < n and grid[r][c+1] is not None)
            start_down   = (r == 0 or grid[r-1][c] is None) and (r+1 < n and grid[r+1][c] is not None)
            if start_across or start_down:
                cell_number[(r,c)] = num
                num += 1

    # enrich entries with numbers from their start cell
    enriched: Dict[str, Dict] = {}
    for e in entries:
        key = f'{e["dir"]}:{e["r"]}:{e["c"]}:{e["answer"]}'
        number = cell_number.get((e["r"], e["c"]), None)
        item = dict(e); item["number"] = number
        enriched[key] = item
    return cell_number, enriched

def parse_words_and_clues(raw: str) -> Tuple[List[str], Dict[str, str]]:
    """
    Accepts lines separated by comma/newline.
    Syntax: ANSWER or ANSWER = clue text
    """
    if not raw.strip(): return [], {}
    words = []
    clues = {}
    for chunk in raw.replace(",", "\n").splitlines():
        t = chunk.strip()
        if not t: continue
        if "=" in t:
            ans, cl = t.split("=", 1)
            ans = ans.strip().replace(" ", "").upper()
            cl = cl.strip()
            if ans:
                words.append(ans)
                clues[ans] = cl
        else:
            ans = t.strip().replace(" ", "").upper()
            if ans:
                words.append(ans)
    return words, clues

# =================== SYMMETRY & ENTRY EXTRACTION ===================
def rot_sym_pos(n: int, r: int, c: int) -> Tuple[int,int]:
    return (n-1-r, n-1-c)

def enforce_rotational_symmetry_blocks(grid) -> int:
    """
    Best-effort: if a cell is a block (None), ensure its 180° counterpart is also a block.
    Returns number of letters removed to satisfy symmetry (violations fixed).
    """
    n = len(grid)
    removed = 0
    for r in range(n):
        for c in range(n):
            if grid[r][c] is None:
                rr, cc = rot_sym_pos(n, r, c)
                if grid[rr][cc] is not None:
                    grid[rr][cc] = None
                    removed += 1
    return removed

def extract_entries_from_grid(grid) -> List[Dict]:
    """Rescan the grid to build Across/Down entries with answers from letters."""
    n = len(grid)
    out: List[Dict] = []
    # Across
    for r in range(n):
        c = 0
        while c < n:
            if grid[r][c] is None:
                c += 1; continue
            start = c
            while c < n and grid[r][c] is not None:
                c += 1
            if c - start >= 2:  # classic rule: min len 2
                ans = "".join(grid[r][k] for k in range(start, c))
                out.append({"answer": ans, "dir": "A", "r": r, "c": start,
                            "coords": [(r, k) for k in range(start, c)]})
        # Down
    for c in range(n):
        r = 0
        while r < n:
            if grid[r][c] is None:
                r += 1; continue
            start = r
            while r < n and grid[r][c] is not None:
                r += 1
            if r - start >= 2:
                ans = "".join(grid[k][c] for k in range(start, r))
                out.append({"answer": ans, "dir": "D", "r": start, "c": c,
                            "coords": [(k, c) for k in range(start, r)]})
    return out

# =================== RENDER HELPERS ===================
def draw_rect_styled(draw: ImageDraw.ImageDraw, x0, y0, x1, y1, color, width: int, style: str):
    style = (style or "solid").lower()
    if isinstance(color, str):
        r, g, b = hex_to_rgb_tuple(color); color = (r, g, b, 255)
    w = max(1, int(width))
    if style == "solid":
        draw.rectangle([x0, y0, x1, y1], outline=color, width=w)
        return
    dash_len = 8 if style == "dashed" else 2
    gap_len  = 6 if style == "dashed" else 4
    def dash_line(xa, ya, xb, yb):
        dx, dy = xb - xa, yb - ya
        length = max(abs(dx), abs(dy))
        if length == 0: return
        stepx = dx / length; stepy = dy / length
        pos = 0; on = True; cx, cy = xa, ya
        while pos < length:
            seg = min(dash_len if on else gap_len, length - pos)
            nx = cx + stepx*seg; ny = cy + stepy*seg
            if on: draw.line([cx, cy, nx, ny], fill=color, width=w)
            cx, cy = nx, ny; pos += seg; on = not on
    dash_line(x0, y0, x1, y0); dash_line(x1, y0, x1, y1)
    dash_line(x1, y1, x0, y1); dash_line(x0, y1, x0, y0)

# =================== PNG / SVG / PDF / DOCX (updated) ===================
def render_crossword_png(grid, entries, clues_map,
                         png_path: Optional[str],
                         show_solution: bool,
                         cell_px=DEFAULT_CELL_PX, margin_cells=DEFAULT_MARGIN_CELLS,
                         default_highlight_hex="#FFE696", alpha01: float = 0.7,
                         per_key_palette: Optional[Dict[str,str]] = None,
                         grid_thickness_px: int = 1, grid_style: str = "solid", grid_color_hex: str = "#000000",
                         letter_font_path: Optional[str] = None, letter_font_px: Optional[int] = None, letter_color_hex: str = "#000000",
                         letter_outline_px: int = 0, letter_outline_hex: str = "#000000",
                         number_font_path: Optional[str] = None, number_font_px: int = 12, number_color_hex: str = "#000000",
                         # NEW:
                         number_bold: bool = False, number_italic: bool = False, number_underline: bool = False,
                         open_cell_bg_hex: str = "#FFFFFF",
                         digit_script: str = "Latin",
                         orientation: str = "LTR (Across →)",
                         vertical_text: bool = False,
                         scale_factor: float = 1.0, dpi: int = 300,
                         align_h: str = "center", align_v: str = "middle",
                         bold: bool = False, italic: bool = False, underline: bool = False):
    """
    Draws puzzle (letters hidden) or solution (letters shown), with clue numbers styled.
    Blocks = None -> black squares. Open cells filled with open_cell_bg_hex.
    """
    n = len(grid)
    cell = int(max(1, round(cell_px * max(0.25, scale_factor))))
    border = margin_cells * cell
    W = H = n * cell + border * 2

    img = Image.new("RGBA", (W, H), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)

    # Fonts
    def try_font(path, size):
        if not path: return None
        try: return ImageFont.truetype(path, size)
        except Exception: return None

    letter_font = (try_font(letter_font_path, int(letter_font_px or int(cell*0.55)))
                   or try_font("DejaVuSans.ttf", int(letter_font_px or int(cell*0.55)))
                   or try_font("arial.ttf", int(letter_font_px or int(cell*0.55)))
                   or ImageFont.load_default())
    number_font = (try_font(number_font_path, int(number_font_px or max(8, int(cell*0.28))))
                   or try_font("DejaVuSans.ttf", int(number_font_px or max(8, int(cell*0.28))))
                   or try_font("arial.ttf", int(number_font_px or max(8, int(cell*0.28))))
                   or ImageFont.load_default())

    per_key_palette = per_key_palette or {}

    # NYT numbering
    cell_number, enriched = compute_numbering_nyt(grid, entries)

    key_to_cells: Dict[str, List[Tuple[int,int]]] = {}
    for e in entries:
        key = f'{e["dir"]}:{e["r"]}:{e["c"]}:{e["answer"]}'.upper()
        key_to_cells[key] = e["coords"]

    per_cell_hex = {}
    for key, coords in key_to_cells.items():
        answer = key.split(":")[-1]
        col = per_key_palette.get(key) or per_key_palette.get(answer) or None
        if not col:
            col = None
        if col:
            for pos in coords:
                per_cell_hex.setdefault(pos, col)
    if not per_cell_hex and default_highlight_hex:
        for coords in key_to_cells.values():
            for pos in coords:
                per_cell_hex.setdefault(pos, default_highlight_hex)

    # Colors
    lr, lg, lb = hex_to_rgb_tuple(letter_color_hex); letter_color = (lr, lg, lb, 255)
    orr, org, orb = hex_to_rgb_tuple(letter_outline_hex); letter_outline_color = (orr, org, orb, 255)
    nr, ng, nb = hex_to_rgb_tuple(number_color_hex); number_color = (nr, ng, nb, 255)
    cr, cg, cb = hex_to_rgb_tuple(open_cell_bg_hex); open_bg = (cr, cg, cb, 255)

    # Visual padding
    pad = max(1, int(round(cell * 0.06)))

    # Helper: styled number draw
    def draw_number(x, y, txt):
        # Underline baseline
        l, t, r, b = draw.textbbox((0, 0), txt, font=number_font)
        if number_bold:
            for dx, dy in [(-1,0),(1,0),(0,-1),(0,1)]:
                draw.text((x+dx, y+dy), txt, font=number_font, fill=number_color)
        if number_italic:
            shear = math.tan(math.radians(12.0))
            # Simulate italic by drawing offset progressively
            steps = 3
            for s in range(steps):
                draw.text((x + s*shear, y), txt, font=number_font, fill=number_color)
        draw.text((x, y), txt, font=number_font, fill=number_color)
        if number_underline:
            underline_y = y + (b - t) + 1
            draw.line([(x, underline_y), (x + (r-l), underline_y)], fill=number_color, width=1)

    # Render
    for i in range(n):
        for j in range(n):
            x0 = border + j*cell
            y0 = border + i*cell
            x1, y1 = x0 + cell, y0 + cell

            if grid[i][j] is None:
                # Block
                draw.rectangle([x0, y0, x1, y1], fill=(0,0,0,255))
            else:
                # Open cell background
                draw.rectangle([x0, y0, x1, y1], fill=open_bg)
                # Highlight (solution only)
                if show_solution and (i, j) in per_cell_hex:
                    r, g, b = hex_to_rgb_tuple(per_cell_hex[(i, j)])
                    a = int(max(0.0, min(1.0, alpha01)) * 255)
                    draw.rectangle([x0, y0, x1, y1], fill=(r, g, b, a))

            # Border
            draw_rect_styled(draw, x0, y0, x1, y1, grid_color_hex, grid_thickness_px, grid_style)

            # Numbers
            num = cell_number.get((i, j))
            if num is not None:
                txt = localize_digits(num, digit_script)
                tl = (x0 + pad, y0 + pad)
                draw_number(tl[0], tl[1], txt)

            # Letters (solution only)
            if show_solution and grid[i][j] is not None:
                ch = shape_if_needed(grid[i][j])
                if vertical_text:
                    # crude vertical: draw char centered; CJK looks fine as-is
                    l, t, r, b = draw.textbbox((0,0), ch, font=letter_font)
                    tw, th = (r - l), (b - t)
                    cx, cy = x0 + cell/2.0, y0 + cell/2.0
                    draw.text((int(cx - tw/2 - l), int(cy - th/2 - t)), ch, font=letter_font, fill=letter_color)
                else:
                    # outline + bold/italic/underline per letters (existing)
                    l, t, r, b = draw.textbbox((0, 0), ch, font=letter_font)
                    tw, th = (r - l), (b - t)
                    PAD = max(2, int(letter_outline_px))
                    glyph_w = int(tw + PAD*2 + 4)
                    glyph_h = int(th + PAD*2 + 4)
                    glyph_img = Image.new("RGBA", (glyph_w, glyph_h), (0,0,0,0))
                    gdraw = ImageDraw.Draw(glyph_img)

                    tx_local = PAD + 2 - l
                    ty_local = PAD + 2 - t

                    if letter_outline_px:
                        gdraw.text((tx_local, ty_local), ch, font=letter_font, fill=letter_color,
                                   stroke_width=max(1, int(letter_outline_px)), stroke_fill=letter_outline_color)
                    if bold:
                        for dx, dy in [(-1,0),(1,0),(0,-1),(0,1),(-1,-1),(1,-1),(1,1),(-1,1)]:
                            gdraw.text((tx_local+dx, ty_local+dy), ch, font=letter_font, fill=letter_color)
                    gdraw.text((tx_local, ty_local), ch, font=letter_font, fill=letter_color)
                    if italic:
                        shear = math.tan(math.radians(12.0))
                        extra_w = int(math.ceil(glyph_h * shear))
                        glyph_img = glyph_img.transform(
                            (glyph_w + extra_w, glyph_h),
                            Image.AFFINE, (1, shear, 0, 0, 1, 0),
                            resample=Image.BICUBIC
                        )
                    alpha_ch = glyph_img.split()[3]
                    vis = alpha_ch.getbbox() or (0,0,glyph_img.width,glyph_img.height)
                    vis_l, vis_t, vis_r, vis_b = vis

                    if align_h == "left":
                        dest_x = int(round(x0 + pad - vis_l))
                    elif align_h == "right":
                        dest_x = int(round(x1 - pad - vis_r))
                    else:
                        cx = x0 + cell/2.0
                        dest_x = int(round(cx - (vis_l + vis_r)/2.0))
                    if align_v == "top":
                        dest_y = int(round(y0 + pad - vis_t))
                    elif align_v == "bottom":
                        dest_y = int(round(y1 - pad - vis_b))
                    else:
                        cy = y0 + cell/2.0
                        dest_y = int(round(cy - (vis_t + vis_b)/2.0))

                    img.alpha_composite(glyph_img, dest=(dest_x, dest_y))
                    if underline:
                        underline_y = dest_y + vis_b + max(1, int(letter_outline_px / 2))
                        draw.line([(dest_x + vis_l, underline_y), (dest_x + vis_r, underline_y)],
                                  fill=letter_color, width=max(1, int(letter_outline_px) or 1))

    if png_path:
        try:
            img.save(png_path, dpi=(int(dpi), int(dpi)))
        except Exception:
            img.save(png_path)
    return img

def to_svg_crossword(grid, entries, svg_path: str,
                     cell_px: int = 48, margin_cells: int = 1,
                     default_highlight_hex: str = "#FFE696", alpha01: float = 0.7,
                     per_key_palette: Optional[Dict[str,str]] = None,
                     grid_thickness_px: int = 1, grid_style: str = "solid", grid_color_hex: str = "#000000",
                     letter_font_family: str = "DejaVu Sans", letter_px: int = 28, letter_color_hex: str = "#000000",
                     number_font_family: str = "DejaVu Sans", number_px: int = 12, number_color_hex: str = "#000000",
                     with_solution: bool = False,
                     digit_script: str = "Latin",
                     number_bold=False, number_italic=False, number_underline=False,
                     open_cell_bg_hex: str = "#FFFFFF"):
    n = len(grid)
    cell = cell_px
    border = margin_cells * cell
    W = H = n * cell + border * 2

    cell_number, _ = compute_numbering_nyt(grid, entries)

    per_key_palette = per_key_palette or {}
    key_to_cells: Dict[str, List[Tuple[int,int]]] = {}
    for e in entries:
        key = f'{e["dir"]}:{e["r"]}:{e["c"]}:{e["answer"]}'.upper()
        key_to_cells[key] = e["coords"]

    per_cell_hex = {}
    for key, coords in key_to_cells.items():
        ans = key.split(":")[-1]
        col = per_key_palette.get(key) or per_key_palette.get(ans) or None
        if not col: col = None
        if col:
            for pos in coords:
                per_cell_hex.setdefault(pos, col)
    if not per_cell_hex and default_highlight_hex:
        for coords in key_to_cells.values():
            for pos in coords:
                per_cell_hex.setdefault(pos, default_highlight_hex)

    def svg_header():
        return f'<svg xmlns="http://www.w3.org/2000/svg" width="{W}" height="{H}" viewBox="0 0 {W} {H}">\n'
    def svg_rect(x, y, w, h, fill="none", stroke="#000", stroke_width=1, style="solid", opacity=1.0):
        dash = ""
        s = (style or "solid").lower()
        if s == "dashed": dash = ' stroke-dasharray="8 6"'
        elif s == "dotted": dash = ' stroke-dasharray="2 4"'
        return f'<rect x="{x}" y="{y}" width="{w}" height="{h}" fill="{fill}" fill-opacity="{opacity}" stroke="{stroke}" stroke-width="{stroke_width}"{dash} />\n'
    def anchors(h: str): return {"left":"start","center":"middle","right":"end"}.get(h, "middle")
    def baselines(v: str): return {"top":"text-before-edge","middle":"middle","bottom":"text-after-edge"}.get(v, "middle")
    parts = [svg_header()]

    # Cells
    for i in range(n):
        for j in range(n):
            x = border + j*cell; y = border + i*cell
            if grid[i][j] is None:
                parts.append(svg_rect(x, y, cell, cell, fill="#000000", stroke=grid_color_hex,
                                      stroke_width=max(1,int(grid_thickness_px)), style=grid_style, opacity=1.0))
            else:
                parts.append(svg_rect(x, y, cell, cell, fill=open_cell_bg_hex, stroke="none",
                                      stroke_width=0, style="solid", opacity=1.0))
                if with_solution and (i, j) in per_cell_hex:
                    parts.append(svg_rect(x, y, cell, cell, fill=per_cell_hex[(i, j)], stroke="none",
                                          stroke_width=0, style="solid", opacity=max(0.0, min(1.0, alpha01))))
                parts.append(svg_rect(x, y, cell, cell, fill="none", stroke=grid_color_hex,
                                      stroke_width=max(1,int(grid_thickness_px)), style=grid_style, opacity=1.0))

                # Numbers
                num = cell_number.get((i, j))
                if num is not None:
                    txt = localize_digits(num, digit_script)
                    weight = "bold" if number_bold else "normal"
                    style = "italic" if number_italic else "normal"
                    deco = "text-decoration: underline;" if number_underline else ""
                    parts.append(
                        f'<text x="{x + cell*0.12}" y="{y + cell*0.18}" '
                        f'style="font-weight:{weight}; font-style:{style}; {deco}" '
                        f'text-anchor="start" dominant-baseline="text-before-edge" '
                        f'font-family="{number_font_family}" font-size="{number_px}" fill="{number_color_hex}">{txt}</text>\n'
                    )

                if with_solution:
                    parts.append(
                        f'<text x="{x + cell/2}" y="{y + cell/2}" text-anchor="{anchors("center")}" dominant-baseline="{baselines("middle")}" '
                        f'font-family="{letter_font_family}" font-size="{letter_px}" fill="{letter_color_hex}">{grid[i][j]}</text>\n'
                    )

    parts.append("</svg>")
    with open(svg_path, "w", encoding="utf-8") as f:
        f.writelines(parts)

# (PDF/DOCX updated in Part 2)

# =================== CSV / DIRECTIVES (Bulk) ===================
DIRECTIVE_RE = re.compile(r'^@([a-zA-Z_]+)=(.+)$')
def parse_bulk_line(line: str) -> Tuple[str, Dict[str,str], List[str], Dict[str,str]]:
    parts = [p.strip() for p in line.split(",") if p.strip()]
    if not parts: return "", {}, [], {}
    title = parts[0]; idx = 1
    params: Dict[str, str] = {}
    answers: List[str] = []
    clues: Dict[str, str] = {}
    while idx < len(parts):
        tok = parts[idx]
        if tok.startswith("#") and len(tok) >= 4:
            params["color"] = tok
        else:
            m = DIRECTIVE_RE.match(tok)
            if m:
                key, val = m.group(1).lower(), m.group(2).strip()
                params[key] = val
            else:
                for wtok in parts[idx:]:
                    if "=" in wtok:
                        a, cl = wtok.split("=", 1)
                        a = a.strip().replace(" ", "").upper()
                        answers.append(a); clues[a] = cl.strip()
                    else:
                        a = wtok.strip().replace(" ", "").upper()
                        answers.append(a)
                break
        idx += 1
    return title, params, answers, clues

def coerce_bool(val: str, default: bool) -> bool:
    if val is None: return default
    v = str(val).strip().lower()
    return v in ("1","true","t","yes","y","on")

def coerce_int(val: str, default: Optional[int]) -> Optional[int]:
    try: return int(val)
    except: return default

def coerce_float(val: str, default: Optional[float]) -> Optional[float]:
    try: return float(val)
    except: return default

# =================== THEMES ===================
BUILTIN_THEMES = {
    "Classic": {
        "highlight_mode": "single",
        "highlight_hex": "#FFE696",
        "alpha_pct": 70,
        "grid_line_style": "solid",
        "grid_line_color_hex": "#000000",
        "grid_thickness_px": 1,
        "grid_thickness_pt": 0.75,
        "docx_border_pt": 0.75,
        "letter_color_hex": "#000000",
        "number_color_hex": "#000000",
        "letter_outline_px": 0,
        "letter_outline_pt": 0.0,
        "title_color_hex": "#000000",
        "list_color_hex": "#000000",
        "open_cell_bg_hex": "#FFFFFF",
        "ui_theme": "light",
    },
    "Dark": {
        "highlight_mode": "single",
        "highlight_hex": "#5A7FFF",
        "alpha_pct": 60,
        "grid_line_style": "solid",
        "grid_line_color_hex": "#E0E0E0",
        "grid_thickness_px": 1,
        "grid_thickness_pt": 0.75,
        "docx_border_pt": 0.75,
        "letter_color_hex": "#FFFFFF",
        "number_color_hex": "#FFFFFF",
        "letter_outline_px": 0,
        "letter_outline_pt": 0.0,
        "title_color_hex": "#FFFFFF",
        "list_color_hex": "#FFFFFF",
        "open_cell_bg_hex": "#1F1F1F",
        "ui_theme": "dark",
    }
}
# =================== GUI APP ===================
class CrosswordApp:
    def __init__(self, root):
        self.root = root
        root.title("Crossword Builder (Ultimate v2)")

        # --- state
        self.grid_data = None
        self.entries = None
        self.preview_image = None
        self.clues_map: Dict[str, str] = {}
        self.puzzle_counter = 1
        self.undo_stack: List[Tuple] = []
        self.symmetry_violations = 0

        # sizing & constraints
        self.auto_size = tk.BooleanVar(value=True)
        self.size_var = tk.StringVar(value="13")
        self.auto_size_mode = tk.StringVar(value="density")
        self.density_pct = tk.IntVar(value=60)
        self.cells_per_word = tk.IntVar(value=10)
        self.cells_per_letter = tk.DoubleVar(value=1.7)

        # connectivity / overlaps
        self.require_connected = tk.BooleanVar(value=True)
        self.min_overlap = tk.IntVar(value=1)
        self.allow_reversed = tk.BooleanVar(value=False)

        # symmetry
        self.enforce_symmetry = tk.BooleanVar(value=False)

        # highlights
        self.highlight_mode = tk.StringVar(value="single")
        self.highlight_hex = tk.StringVar(value="#FFE696")
        self.alpha_pct = tk.IntVar(value=70)
        self.palette_text = tk.StringVar(value="")

        # grid lines & cell bg
        self.grid_thickness_px = tk.IntVar(value=1)
        self.grid_thickness_pt = tk.DoubleVar(value=0.75)
        self.docx_border_pt = tk.DoubleVar(value=0.75)
        self.grid_line_style = tk.StringVar(value="solid")
        self.grid_line_color_hex = tk.StringVar(value="#000000")
        self.open_cell_bg_hex = tk.StringVar(value="#FFFFFF")

        # geometry + PNG
        self.cell_px = tk.IntVar(value=48)
        self.margin_cells = tk.IntVar(value=1)
        self.pdf_cell_pt = tk.DoubleVar(value=18.0)
        self.png_scale = tk.DoubleVar(value=1.0)
        self.png_dpi = tk.IntVar(value=300)

        # letters & numbers
        self.letter_font_path = tk.StringVar(value="")
        self.letter_font_px = tk.IntVar(value=28)
        self.letter_font_pt = tk.DoubleVar(value=12.0)
        self.letter_color_hex = tk.StringVar(value="#000000")
        self.letter_outline_px = tk.IntVar(value=0)
        self.letter_outline_pt = tk.DoubleVar(value=0.0)
        self.letter_outline_hex = tk.StringVar(value="#000000")
        self.letter_align_h = tk.StringVar(value="center")
        self.letter_align_v = tk.StringVar(value="middle")
        self.letter_bold = tk.BooleanVar(value=False)
        self.letter_italic = tk.BooleanVar(value=False)
        self.letter_underline = tk.BooleanVar(value=False)

        self.number_font_path = tk.StringVar(value="")
        self.number_font_px = tk.IntVar(value=12)
        self.number_font_pt = tk.DoubleVar(value=8.0)
        self.number_color_hex = tk.StringVar(value="#000000")
        self.number_bold = tk.BooleanVar(value=False)
        self.number_italic = tk.BooleanVar(value=False)
        self.number_underline = tk.BooleanVar(value=False)
        self.digit_script = tk.StringVar(value="Latin")

        # multi-lang / orientation
        self.orientation = tk.StringVar(value=ORIENTATIONS[0])
        self.vertical_text = tk.BooleanVar(value=False)

        # title/list
        self.title_font_path = tk.StringVar(value="")
        self.title_font_pt = tk.DoubleVar(value=18.0)
        self.title_color_hex = tk.StringVar(value="#000000")
        self.list_font_path = tk.StringVar(value="")
        self.list_font_pt = tk.DoubleVar(value=10.0)
        self.list_color_hex = tk.StringVar(value="#000000")

        # export options
        self.also_json = tk.BooleanVar(value=True)
        self.show_solution = tk.BooleanVar(value=False)
        self.seed_var = tk.StringVar(value="")

        # bulk defaults (kept as-is for compatibility; not expanded here)
        self.bulk_size_var = tk.StringVar()
        self.bulk_seed = tk.StringVar(value="")
        self.bulk_auto_mode = tk.StringVar(value="density")
        self.bulk_density_pct = tk.IntVar(value=60)
        self.bulk_cpw = tk.IntVar(value=10)
        self.bulk_cpl = tk.DoubleVar(value=1.7)
        self.bulk_require_connected = tk.BooleanVar(value=True)
        self.bulk_min_overlap = tk.IntVar(value=1)
        self.bulk_allow_reversed = tk.BooleanVar(value=False)
        self.bulk_highlight_mode = tk.StringVar(value="single")
        self.bulk_highlight_hex = tk.StringVar(value="#FFE696")
        self.bulk_alpha_pct = tk.IntVar(value=70)
        self.bulk_palette_text = tk.StringVar(value="#ff9999;#99ff99;#9999ff")
        self.bulk_grid_thickness_px = tk.IntVar(value=1)
        self.bulk_grid_thickness_pt = tk.DoubleVar(value=0.75)
        self.bulk_docx_border_pt = tk.DoubleVar(value=0.75)
        self.bulk_grid_line_style = tk.StringVar(value="solid")
        self.bulk_grid_line_color_hex = tk.StringVar(value="#000000")
        self.bulk_letter_font_path = tk.StringVar(value="")
        self.bulk_letter_font_pt = tk.DoubleVar(value=12.0)
        self.bulk_letter_color_hex = tk.StringVar(value="#000000")
        self.bulk_letter_outline_pt = tk.DoubleVar(value=0.0)
        self.bulk_letter_outline_hex = tk.StringVar(value="#000000")
        self.bulk_number_font_path = tk.StringVar(value="")
        self.bulk_number_font_pt = tk.DoubleVar(value=8.0)
        self.bulk_number_color_hex = tk.StringVar(value="#000000")
        self.bulk_title_font_path = tk.StringVar(value="")
        self.bulk_title_font_pt = tk.DoubleVar(value=18.0)
        self.bulk_title_color_hex = tk.StringVar(value="#000000")
        self.bulk_list_font_path = tk.StringVar(value="")
        self.bulk_list_font_pt = tk.DoubleVar(value=10.0)
        self.bulk_list_color_hex = tk.StringVar(value="#000000")
        self.bulk_png_scale = tk.DoubleVar(value=1.0)
        self.bulk_png_dpi = tk.IntVar(value=300)

        # ----- Layout
        nb = ttk.Notebook(root); nb.grid(row=0, column=0, sticky="nsew")
        self.single_tab = ttk.Frame(nb, padding=10); nb.add(self.single_tab, text="Single")
        self.bulk_tab = ttk.Frame(nb, padding=10); nb.add(self.bulk_tab, text="Bulk")

        self._build_single_tab(self.single_tab)
        self._build_bulk_tab(self.bulk_tab)

        root.columnconfigure(0, weight=1); root.rowconfigure(0, weight=1)
        self.update_auto_widgets_state(); self.update_bulk_auto_widgets_state()
        self.apply_builtin_theme("Classic")
        self.generate_single()

        # Shortcuts
        root.bind("<Control-s>", lambda e: self.export_single_dialog())
        root.bind("<Control-S>", lambda e: self.export_single_dialog())
        root.bind("<Control-r>", lambda e: self.generate_single())
        root.bind("<Control-R>", lambda e: self.generate_single())
        root.bind("<Control-z>", lambda e: self.undo_last())
        root.bind("<Control-Z>", lambda e: self.undo_last())
        root.bind("<Control-t>", lambda e: self.toggle_theme())
        root.bind("<Control-T>", lambda e: self.toggle_theme())

    # ---------- UI BUILDERS ----------
    def _build_single_tab(self, parent):
        parent.columnconfigure(2, weight=1); parent.rowconfigure(0, weight=1)
        left = ttk.Frame(parent); center = ttk.Frame(parent); right = ttk.Frame(parent)
        left.grid(row=0, column=0, sticky="nsew", padx=(0,10))
        center.grid(row=0, column=1, sticky="nsew", padx=(0,10))
        right.grid(row=0, column=2, sticky="nsew")

        # Presets + Theme
        presets = ttk.Frame(left); presets.grid(row=0, column=0, columnspan=20, sticky="ew", pady=(0,6))
        ttk.Label(presets, text="Theme:").pack(side="left")
        self.theme_choice = tk.StringVar(value="Classic")
        ttk.Combobox(presets, textvariable=self.theme_choice, values=list(BUILTIN_THEMES.keys()),
                     state="readonly", width=14).pack(side="left")
        ttk.Button(presets, text="Apply", command=lambda: self.apply_builtin_theme(self.theme_choice.get())).pack(side="left", padx=(6,0))
        ttk.Button(presets, text="Toggle Light/Dark (Ctrl+T)", command=self.toggle_theme).pack(side="left", padx=(10,0))
        ttk.Button(presets, text="Save Preset…", command=self.save_preset_json).pack(side="left", padx=(10,0))
        ttk.Button(presets, text="Load Preset…", command=self.load_preset_json).pack(side="left", padx=(6,0))

        ttk.Label(left, text="Answers (one per line, or ANSWER = clue):").grid(row=1, column=0, columnspan=20, sticky="w")
        self.words_box = tk.Text(left, width=44, height=8)
        self.words_box.grid(row=2, column=0, columnspan=20, sticky="nsew", pady=(2,8))
        self.words_box.insert("1.0", "PYTHON = Programming language\nALGORITHM\nSEARCH\nGRID\nPUZZLE")

        # size row
        row3 = ttk.Frame(left); row3.grid(row=3, column=0, columnspan=20, sticky="ew", pady=(2,4))
        ttk.Checkbutton(row3, text="Auto size", variable=self.auto_size, command=self.update_auto_widgets_state).pack(side="left")
        ttk.Label(row3, text="Grid size:").pack(side="left", padx=(8,2))
        ttk.Entry(row3, textvariable=self.size_var, width=6, state="disabled").pack(side="left")
        ttk.Label(row3, text="Auto mode:").pack(side="left", padx=(12,2))
        cbm = ttk.Combobox(row3, textvariable=self.auto_size_mode, values=["density","cells/word","cells/letter"],
                           state="readonly", width=12)
        cbm.pack(side="left"); cbm.bind("<<ComboboxSelected>>", lambda e: self.update_auto_widgets_state())

        ttk.Label(row3, text="Density%").pack(side="left", padx=(10,2))
        self.spin_density = ttk.Spinbox(row3, from_=30, to=95, textvariable=self.density_pct, width=5); self.spin_density.pack(side="left")
        ttk.Label(row3, text="Cells/word").pack(side="left", padx=(10,2))
        self.spin_cpw = ttk.Spinbox(row3, from_=4, to=100, textvariable=self.cells_per_word, width=6); self.spin_cpw.pack(side="left")
        ttk.Label(row3, text="Cells/letter").pack(side="left", padx=(10,2))
        self.spin_cpl = ttk.Spinbox(row3, from_=1.2, to=5.0, increment=0.1, textvariable=self.cells_per_letter, width=6); self.spin_cpl.pack(side="left")

        # connectivity / seed / symmetry
        row4 = ttk.Frame(left); row4.grid(row=4, column=0, columnspan=20, sticky="ew", pady=(2,4))
        ttk.Checkbutton(row4, text="Require connected", variable=self.require_connected).pack(side="left")
        ttk.Label(row4, text="Min overlaps/word:").pack(side="left", padx=(8,2))
        ttk.Spinbox(row4, from_=0, to=4, textvariable=self.min_overlap, width=4).pack(side="left")
        ttk.Checkbutton(row4, text="Allow reversed placement", variable=self.allow_reversed).pack(side="left", padx=(12,0))
        ttk.Label(row4, text="Seed:").pack(side="left", padx=(12,2))
        ttk.Entry(row4, textvariable=self.seed_var, width=10).pack(side="left")
        ttk.Checkbutton(row4, text="Rotational block symmetry", variable=self.enforce_symmetry).pack(side="left", padx=(12,0))
        ttk.Button(row4, text="Symmetrize now", command=self.symmetrize_now).pack(side="left", padx=(6,0))

        # highlight & transparency
        row5 = ttk.Frame(left); row5.grid(row=5, column=0, columnspan=20, sticky="ew", pady=(4,2))
        ttk.Label(row5, text="Highlight mode:").pack(side="left")
        ttk.Radiobutton(row5, text="Single", variable=self.highlight_mode, value="single",
                        command=self.update_preview).pack(side="left")
        ttk.Radiobutton(row5, text="Per-entry palette", variable=self.highlight_mode, value="palette",
                        command=self.update_preview).pack(side="left", padx=(4,0))
        ttk.Label(row5, text="Transparency %:").pack(side="left", padx=(16,4))
        ttk.Spinbox(row5, from_=0, to=100, textvariable=self.alpha_pct, width=5,
                    command=self.update_preview).pack(side="left")
        row5b = ttk.Frame(left); row5b.grid(row=6, column=0, columnspan=20, sticky="ew", pady=(2,2))
        ttk.Label(row5b, text="Single color:").pack(side="left")
        self.single_color_swatch = tk.Label(row5b, text="  ", background=self.highlight_hex.get(), relief="solid", width=3)
        self.single_color_swatch.pack(side="left", padx=(4,6))
        ttk.Button(row5b, text="Pick…", command=lambda: self.pick_hex(self.highlight_hex, self.single_color_swatch, True)).pack(side="left")
        ttk.Label(row5b, text="Palette (hexes or KEY=#hex; ANSWER=#hex …):").pack(side="left", padx=(16,4))
        ttk.Entry(row5b, textvariable=self.palette_text, width=46).pack(side="left")
        ttk.Button(row5b, text="Auto palette", command=self.set_auto_palette_from_entries).pack(side="left", padx=(6,0))

        # grid lines & bg
        row6 = ttk.Frame(left); row6.grid(row=7, column=0, columnspan=20, sticky="ew", pady=(6,2))
        ttk.Label(row6, text="Grid style:").pack(side="left")
        ttk.Combobox(row6, width=8, state="readonly", values=["solid","dashed","dotted"],
                     textvariable=self.grid_line_style).pack(side="left")
        ttk.Label(row6, text="Line color:").pack(side="left", padx=(12,4))
        self.linecolor_swatch = tk.Label(row6, text="  ", background=self.grid_line_color_hex.get(), relief="solid", width=3)
        self.linecolor_swatch.pack(side="left", padx=(0,6))
        ttk.Button(row6, text="Pick…", command=lambda: self.pick_hex(self.grid_line_color_hex, self.linecolor_swatch, True)).pack(side="left")
        ttk.Label(row6, text="Thickness:").pack(side="left", padx=(12,4))
        ttk.Spinbox(row6, from_=1, to=10, textvariable=self.grid_thickness_px, width=4,
                    command=self.update_preview).pack(side="left")
        ttk.Label(row6, text="Open-cell BG:").pack(side="left", padx=(12,4))
        self.bgcolor_swatch = tk.Label(row6, text="  ", background=self.open_cell_bg_hex.get(), relief="solid", width=3)
        self.bgcolor_swatch.pack(side="left", padx=(0,6))
        ttk.Button(row6, text="Pick…", command=lambda: self.pick_hex(self.open_cell_bg_hex, self.bgcolor_swatch, True)).pack(side="left")

        row7 = ttk.Frame(left); row7.grid(row=8, column=0, columnspan=20, sticky="ew", pady=(2,2))
        ttk.Label(row7, text="Cell px:").pack(side="left")
        ttk.Spinbox(row7, from_=18, to=96, textvariable=self.cell_px, width=6, command=self.update_preview).pack(side="left")
        ttk.Label(row7, text="Margin (cells):").pack(side="left", padx=(8,2))
        ttk.Spinbox(row7, from_=0, to=5, textvariable=self.margin_cells, width=6, command=self.update_preview).pack(side="left")
        ttk.Label(row7, text="PDF cell pt:").pack(side="left", padx=(8,2))
        ttk.Spinbox(row7, from_=10.0, to=36.0, increment=0.5, textvariable=self.pdf_cell_pt, width=7).pack(side="left")
        ttk.Label(row7, text="PNG scale:").pack(side="left", padx=(8,2))
        ttk.Spinbox(row7, from_=0.5, to=4.0, increment=0.1, textvariable=self.png_scale, width=6).pack(side="left")
        ttk.Label(row7, text="PNG DPI:").pack(side="left", padx=(8,2))
        ttk.Spinbox(row7, from_=72, to=600, increment=10, textvariable=self.png_dpi, width=7).pack(side="left")

        # Fonts & Styles (left)
        fonts = ttk.LabelFrame(left, text="Fonts & Styles")
        fonts.grid(row=9, column=0, columnspan=20, sticky="ew", pady=(6,2))
        # Letters
        ttk.Label(fonts, text="Letters font:").grid(row=0, column=0, sticky="e")
        ttk.Entry(fonts, textvariable=self.letter_font_path, width=28).grid(row=0, column=1, sticky="w")
        ttk.Button(fonts, text="Browse", command=lambda: self.browse_font(self.letter_font_path)).grid(row=0, column=2, sticky="w")
        ttk.Label(fonts, text="Size(px/pt):").grid(row=0, column=3, sticky="e")
        ttk.Spinbox(fonts, from_=8, to=72, textvariable=self.letter_font_px, width=5, command=self.update_preview).grid(row=0, column=4, sticky="w")
        ttk.Spinbox(fonts, from_=6.0, to=36.0, increment=0.5, textvariable=self.letter_font_pt, width=6).grid(row=0, column=5, sticky="w")
        ttk.Label(fonts, text="Color:").grid(row=0, column=6, sticky="e")
        lf_sw = tk.Label(fonts, text="  ", background=self.letter_color_hex.get(), relief="solid", width=3); lf_sw.grid(row=0, column=7, sticky="w")
        ttk.Button(fonts, text="Pick…", command=lambda: self.pick_hex(self.letter_color_hex, lf_sw, True)).grid(row=0, column=8, sticky="w")
        ttk.Label(fonts, text="Outline px/pt & color:").grid(row=1, column=3, sticky="e")
        ttk.Spinbox(fonts, from_=0, to=10, textvariable=self.letter_outline_px, width=5, command=self.update_preview).grid(row=1, column=4, sticky="w")
        lo_sw = tk.Label(fonts, text="  ", background=self.letter_outline_hex.get(), relief="solid", width=3); lo_sw.grid(row=1, column=7, sticky="w")
        ttk.Spinbox(fonts, from_=0.0, to=5.0, increment=0.25, textvariable=self.letter_outline_pt, width=6).grid(row=1, column=5, sticky="w")
        ttk.Button(fonts, text="Pick…", command=lambda: self.pick_hex(self.letter_outline_hex, lo_sw, True)).grid(row=1, column=8, sticky="w")
        align_row = ttk.Frame(fonts); align_row.grid(row=2, column=0, columnspan=9, sticky="ew", pady=(4,2))
        ttk.Label(align_row, text="Letter align:").pack(side="left")
        cbh = ttk.Combobox(align_row, width=8, state="readonly", values=["left","center","right"], textvariable=self.letter_align_h)
        cbh.pack(side="left", padx=(4,10)); cbh.bind("<<ComboboxSelected>>", lambda e: self.update_preview())
        cbv = ttk.Combobox(align_row, width=8, state="readonly", values=["top","middle","bottom"], textvariable=self.letter_align_v)
        cbv.pack(side="left"); cbv.bind("<<ComboboxSelected>>", lambda e: self.update_preview())
        ttk.Label(align_row, text="  Styles:").pack(side="left", padx=(12,2))
        ttk.Checkbutton(align_row, text="Bold", variable=self.letter_bold, command=self.update_preview).pack(side="left")
        ttk.Checkbutton(align_row, text="Italic", variable=self.letter_italic, command=self.update_preview).pack(side="left")
        ttk.Checkbutton(align_row, text="Underline", variable=self.letter_underline, command=self.update_preview).pack(side="left")

        # Numbers styles
        numrow = ttk.Frame(fonts); numrow.grid(row=3, column=0, columnspan=9, sticky="ew", pady=(6,2))
        ttk.Label(numrow, text="Numbers font:").pack(side="left")
        ttk.Entry(numrow, textvariable=self.number_font_path, width=22).pack(side="left")
        ttk.Button(numrow, text="Browse", command=lambda: self.browse_font(self.number_font_path)).pack(side="left", padx=(4,8))
        ttk.Label(numrow, text="Size(px/pt):").pack(side="left")
        ttk.Spinbox(numrow, from_=6, to=36, textvariable=self.number_font_px, width=5, command=self.update_preview).pack(side="left")
        ttk.Spinbox(numrow, from_=6.0, to=24.0, increment=0.5, textvariable=self.number_font_pt, width=6).pack(side="left", padx=(4,0))
        ttk.Label(numrow, text="Color:").pack(side="left", padx=(6,2))
        nf_sw = tk.Label(numrow, text="  ", background=self.number_color_hex.get(), relief="solid", width=3); nf_sw.pack(side="left")
        ttk.Button(numrow, text="Pick…", command=lambda: self.pick_hex(self.number_color_hex, nf_sw, True)).pack(side="left", padx=(4,8))
        ttk.Checkbutton(numrow, text="Bold", variable=self.number_bold, command=self.update_preview).pack(side="left")
        ttk.Checkbutton(numrow, text="Italic", variable=self.number_italic, command=self.update_preview).pack(side="left")
        ttk.Checkbutton(numrow, text="Underline", variable=self.number_underline, command=self.update_preview).pack(side="left", padx=(0,8))
        ttk.Label(numrow, text="Digits:").pack(side="left")
        ttk.Combobox(numrow, textvariable=self.digit_script, values=DIGIT_KEYS, width=14, state="readonly").pack(side="left", padx=(4,0))

        # Orientation / language
        langrow = ttk.Frame(left); langrow.grid(row=10, column=0, columnspan=20, sticky="ew", pady=(6,2))
        ttk.Label(langrow, text="Orientation:").pack(side="left")
        ttk.Combobox(langrow, textvariable=self.orientation, values=ORIENTATIONS, width=18, state="readonly").pack(side="left", padx=(4,8))
        ttk.Checkbutton(langrow, text="Vertical text (CJK)", variable=self.vertical_text, command=self.update_preview).pack(side="left")

        # Actions
        btns = ttk.Frame(left); btns.grid(row=11, column=0, columnspan=20, sticky="ew", pady=(8,4))
        ttk.Button(btns, text="Generate / Regenerate (Ctrl+R)", command=self.generate_single).pack(side="left")
        ttk.Checkbutton(btns, text="Show solution", variable=self.show_solution, command=self.update_preview).pack(side="left", padx=(10,0))
        ttk.Button(btns, text="Export… (Ctrl+S)", command=self.export_single_dialog).pack(side="left", padx=6)
        ttk.Checkbutton(btns, text="Also save JSON", variable=self.also_json).pack(side="left", padx=(10,0))
        ttk.Button(btns, text="Undo (Ctrl+Z)", command=self.undo_last).pack(side="left", padx=(10,0))

        self.words_label = ttk.Label(left, text="Entries: —"); self.words_label.grid(row=12, column=0, columnspan=20, sticky="w", pady=(6,0))
        self.sym_label = ttk.Label(left, text="Symmetry: OK"); self.sym_label.grid(row=13, column=0, columnspan=20, sticky="w")

        # Center: Canvas preview
        self.canvas = tk.Canvas(center, width=MIN_PREVIEW, height=MIN_PREVIEW, background="#ffffff")
        self.canvas.grid(row=0, column=0, sticky="nsew")
        center.columnconfigure(0, weight=1); center.rowconfigure(0, weight=1)

        # Right: Clues panel
        right.columnconfigure(0, weight=1); right.rowconfigure(1, weight=1)
        ttk.Label(right, text="Clues (double-click a clue to edit)").grid(row=0, column=0, sticky="w")
        self.clue_nb = ttk.Notebook(right); self.clue_nb.grid(row=1, column=0, sticky="nsew")

        self.across_frame = ttk.Frame(self.clue_nb); self.down_frame = ttk.Frame(self.clue_nb)
        self.clue_nb.add(self.across_frame, text="Across")
        self.clue_nb.add(self.down_frame, text="Down")

        self.across_tree = self._make_clue_tree(self.across_frame)
        self.down_tree = self._make_clue_tree(self.down_frame)

        act = ttk.Frame(right); act.grid(row=2, column=0, sticky="ew", pady=(6,0))
        ttk.Button(act, text="Renumber (NYT)", command=self.refresh_clue_views).pack(side="left")
        ttk.Button(act, text="Auto-fill empty clues with answer", command=self.fill_empty_clues).pack(side="left", padx=(8,0))

        # grow left text box
        left.rowconfigure(2, weight=1)

    def _make_clue_tree(self, parent):
        cols = ("#", "Answer", "Clue")
        tree = ttk.Treeview(parent, columns=cols, show="headings", selectmode="browse")
        for c, w in zip(cols, (50, 120, 320)):
            tree.heading(c, text=c); tree.column(c, width=w, anchor="w")
        tree.grid(row=0, column=0, sticky="nsew")
        parent.columnconfigure(0, weight=1); parent.rowconfigure(0, weight=1)
        tree.bind("<Double-1>", lambda e, t=tree: self._edit_clue_inplace(t, e))
        return tree

    def _edit_clue_inplace(self, tree, event):
        item_id = tree.identify_row(event.y)
        col = tree.identify_column(event.x)
        if not item_id or col != "#3":
            return
        x, y, w, h = tree.bbox(item_id, column=col)
        value = tree.set(item_id, "Clue")
        entry = tk.Entry(tree)
        entry.insert(0, value)
        entry.select_range(0, tk.END)
        entry.place(x=x, y=y, width=w, height=h)
        entry.focus_set()
        def save_edit(e=None):
            newv = entry.get()
            tree.set(item_id, "Clue", newv)
            entry.destroy()
            # sync to map
            ans = tree.set(item_id, "Answer")
            self.clues_map[ans] = newv
        entry.bind("<Return>", save_edit)
        entry.bind("<FocusOut>", lambda e: save_edit())

    def _build_bulk_tab(self, parent):
        # keep original bulk UI minimal (unchanged logic)
        parent.columnconfigure(0, weight=1); parent.rowconfigure(3, weight=1)
        mode_frame = ttk.Frame(parent); mode_frame.grid(row=0, column=0, sticky="ew", pady=(0,8))
        self.bulk_mode = tk.StringVar(value="import")
        ttk.Radiobutton(mode_frame, text="Import many lists", value="import", variable=self.bulk_mode,
                        command=self._refresh_bulk_mode).pack(side="left")
        ttk.Radiobutton(mode_frame, text="Auto-generate lists", value="generate", variable=self.bulk_mode,
                        command=self._refresh_bulk_mode).pack(side="left", padx=12)

        opt = ttk.LabelFrame(parent, text="Defaults for Bulk")
        opt.grid(row=1, column=0, sticky="ew", pady=(0,8))
        ttk.Checkbutton(opt, text="Require connected", variable=self.bulk_require_connected).pack(side="left", padx=(6,6))
        ttk.Label(opt, text="Min overlaps:").pack(side="left"); ttk.Spinbox(opt, from_=0, to=4, textvariable=self.bulk_min_overlap, width=4).pack(side="left")
        ttk.Checkbutton(opt, text="Allow reversed", variable=self.bulk_allow_reversed).pack(side="left", padx=(6,6))
        ttk.Label(opt, text="Grid size (blank=auto):").pack(side="left", padx=(12,4))
        ttk.Entry(opt, textvariable=self.bulk_size_var, width=6).pack(side="left")
        ttk.Label(opt, text="Base seed:").pack(side="left", padx=(12,4))
        ttk.Entry(opt, textvariable=self.bulk_seed, width=10).pack(side="left")
        ttk.Label(opt, text="Auto mode:").pack(side="left", padx=(12,4))
        cbm = ttk.Combobox(opt, textvariable=self.bulk_auto_mode, values=["density","cells/word","cells/letter"], state="readonly", width=12)
        cbm.pack(side="left"); cbm.bind("<<ComboboxSelected>>", lambda e: self.update_bulk_auto_widgets_state())
        ttk.Label(opt, text="Density%:").pack(side="left", padx=(12,4))
        self.bulk_spin_density = ttk.Spinbox(opt, from_=30, to=95, textvariable=self.bulk_density_pct, width=5); self.bulk_spin_density.pack(side="left")
        ttk.Label(opt, text="Cells/word:").pack(side="left", padx=(12,4))
        self.bulk_spin_cpw = ttk.Spinbox(opt, from_=4, to=100, textvariable=self.bulk_cpw, width=6); self.bulk_spin_cpw.pack(side="left")
        ttk.Label(opt, text="Cells/letter:").pack(side="left", padx=(12,4))
        self.bulk_spin_cpl = ttk.Spinbox(opt, from_=1.2, to=5.0, increment=0.1, textvariable=self.bulk_cpl, width=6).pack(side="left")

        # import area
        self.import_frame = ttk.LabelFrame(parent, text="Import Lists")
        self.import_frame.grid(row=3, column=0, sticky="nsew")
        self.import_frame.columnconfigure(0, weight=1); self.import_frame.rowconfigure(1, weight=1)
        ib = ttk.Frame(self.import_frame); ib.grid(row=0, column=0, sticky="ew", pady=(0,6))
        ttk.Button(ib, text="Load CSV…", command=self.load_csv).pack(side="left")
        ttk.Label(ib, text="OR paste rows: Title, @size=15, ANSWER=Clue, ANSWER...").pack(side="left", padx=8)
        self.bulk_text = tk.Text(self.import_frame, height=10); self.bulk_text.grid(row=1, column=0, sticky="nsew")
        self.bulk_text.insert("1.0",
            "Tech,@size=13,@connected=1,@minov=1,ALGORITHM=Step-by-step process,SEARCH,GRID,PUZZLE\n"
            "Fruits,@automode=cells/word,@cpw=9,APPLE=Keeps doctor away,ORANGE,GRAPE,PEAR\n"
            "Cities,@automode=cells/letter,@cpl=1.8,PARIS,CAIRO,LONDON,BERLIN"
        )

        # generate tab placeholder (reusing original implementation)
        self.generate_frame = ttk.LabelFrame(parent, text="Auto-generate Lists")
        self.generate_frame.grid(row=3, column=0, sticky="nsew")
        self.generate_frame.columnconfigure(1, weight=1)
        ttk.Label(self.generate_frame, text="How many crosswords:").grid(row=0, column=0, sticky="e", pady=2, padx=4)
        self.gen_count = tk.StringVar(value="5"); ttk.Entry(self.generate_frame, textvariable=self.gen_count, width=6).grid(row=0, column=1, sticky="w")
        ttk.Label(self.generate_frame, text="Words per puzzle:").grid(row=1, column=0, sticky="e", pady=2, padx=4)
        self.gen_words_per = tk.StringVar(value="10"); ttk.Entry(self.generate_frame, textvariable=self.gen_words_per, width=6).grid(row=1, column=1, sticky="w")
        ttk.Label(self.generate_frame, text="Word length range (min-max):").grid(row=2, column=0, sticky="e", pady=2, padx=4)
        self.gen_len_range = tk.StringVar(value="4-9"); ttk.Entry(self.generate_frame, textvariable=self.gen_len_range, width=8).grid(row=2, column=1, sticky="w")

        actions = ttk.Frame(parent); actions.grid(row=4, column=0, sticky="ew", pady=(8,0))
        ttk.Button(actions, text="Export Bulk…", command=self.export_bulk).pack(side="left")

        self._refresh_bulk_mode()

    # ---------- UI helpers ----------
    def update_auto_widgets_state(self):
        auto = self.auto_size.get()
        mode = self.auto_size_mode.get()
        def set_size_entry_state(frame):
            for w in frame.winfo_children():
                if isinstance(w, ttk.Entry) and w.cget("textvariable") == str(self.size_var):
                    w.config(state="disabled" if auto else "normal")
        set_size_entry_state(self.single_tab)
        self.spin_density.config(state="normal" if (auto and mode=="density") else "disabled")
        self.spin_cpw.config(state="normal" if (auto and mode=="cells/word") else "disabled")
        self.spin_cpl.config(state="normal" if (auto and mode=="cells/letter") else "disabled")

    def update_bulk_auto_widgets_state(self):
        mode = self.bulk_auto_mode.get()
        self.bulk_spin_density.config(state="normal" if mode=="density" else "disabled")
        self.bulk_spin_cpw.config(state="normal" if mode=="cells/word" else "disabled")
        # bulk_spin_cpl is created inline; skip

    def browse_font(self, var: tk.StringVar):
        path = filedialog.askopenfilename(title="Select TTF/OTF font", filetypes=[("Font files","*.ttf *.otf"),("All files","*.*")])
        if path: var.set(path); self.update_preview()

    def pick_hex(self, var: tk.StringVar, swatch: tk.Label, refresh_preview: bool):
        c = colorchooser.askcolor(color=var.get())
        if c and c[1]:
            var.set(c[1]); swatch.config(background=c[1])
            if refresh_preview: self.update_preview()

    def _refresh_bulk_mode(self):
        (self.import_frame if self.bulk_mode.get()=="import" else self.generate_frame).tkraise()

    # ---------- THEME ----------
    def apply_builtin_theme(self, name: str):
        cfg = BUILTIN_THEMES.get(name)
        if not cfg: return
        # colors & line
        self.highlight_mode.set(cfg.get("highlight_mode", "single"))
        self.highlight_hex.set(cfg.get("highlight_hex", "#FFE696"))
        self.alpha_pct.set(cfg.get("alpha_pct", 70))
        self.grid_line_style.set(cfg.get("grid_line_style", "solid"))
        self.grid_line_color_hex.set(cfg.get("grid_line_color_hex", "#000000"))
        self.grid_thickness_px.set(cfg.get("grid_thickness_px", 1))
        self.grid_thickness_pt.set(cfg.get("grid_thickness_pt", 0.75))
        self.docx_border_pt.set(cfg.get("docx_border_pt", 0.75))
        self.letter_color_hex.set(cfg.get("letter_color_hex", "#000000"))
        self.number_color_hex.set(cfg.get("number_color_hex", "#000000"))
        self.letter_outline_px.set(cfg.get("letter_outline_px", 0))
        self.letter_outline_pt.set(cfg.get("letter_outline_pt", 0.0))
        self.title_color_hex.set(cfg.get("title_color_hex", "#000000"))
        self.list_color_hex.set(cfg.get("list_color_hex", "#000000"))
        self.open_cell_bg_hex.set(cfg.get("open_cell_bg_hex", "#FFFFFF"))
        self.theme_choice.set(name)
        # UI background
        if cfg.get("ui_theme") == "dark":
            try:
                self.root.configure(bg="#121212")
                self.canvas.configure(background="#121212")
            except:
                pass
        else:
            try:
                self.root.configure(bg="")
                self.canvas.configure(background="#ffffff")
            except:
                pass
        self.update_preview()

    def toggle_theme(self):
        name = self.theme_choice.get()
        nxt = "Dark" if name != "Dark" else "Classic"
        self.apply_builtin_theme(nxt)

    # ---------- SETTINGS SNAPSHOT ----------
    def snapshot(self):
        if self.grid_data is None or self.entries is None:
            return
        self.undo_stack.append((
            copy.deepcopy(self.grid_data),
            copy.deepcopy(self.entries),
            copy.deepcopy(self.clues_map)
        ))
        if len(self.undo_stack) > 10:
            self.undo_stack.pop(0)

    def undo_last(self):
        if not self.undo_stack:
            messagebox.showinfo("Undo", "Nothing to undo."); return
        self.grid_data, self.entries, self.clues_map = self.undo_stack.pop()
        self.refresh_clue_views()
        self.update_preview()

    # ---------- GENERATION ----------
    def parse_words_clues_single(self) -> Tuple[List[str], Dict[str,str]]:
        raw = self.words_box.get("1.0", "end")
        words, clues = parse_words_and_clues(raw)
        return words, clues

    def set_auto_palette_from_entries(self):
        if not self.entries:
            words, _ = self.parse_words_clues_single()
            keys = [w.upper() for w in words]
        else:
            keys = [f'{e["dir"]}:{e["r"]}:{e["c"]}:{e["answer"]}'.upper() for e in self.entries]
        if not keys:
            messagebox.showinfo("No entries", "Enter answers first, then click Auto palette.")
            return
        cols = auto_palette(len(keys))
        mapping = ";".join(f"{k}={c}" for k, c in zip(keys, cols))
        self.palette_text.set(mapping)
        self.highlight_mode.set("palette")
        self.update_preview()

    def generate_single(self):
        try:
            self.snapshot()
            if self.seed_var.get().strip():
                try: random.seed(int(self.seed_var.get().strip()))
                except: pass

            answers, clues = self.parse_words_clues_single()
            self.clues_map.update(clues)

            if not answers:
                messagebox.showinfo("No answers", "Please enter one or more answers (or ANSWER = clue).")
                return

            if self.auto_size.get():
                mode = self.auto_size_mode.get()
                if mode == "density":
                    size = compute_size_by_density(answers, int(self.density_pct.get()))
                elif mode == "cells/word":
                    size = compute_size_by_cpw(answers, int(self.cells_per_word.get()))
                else:
                    size = compute_size_by_cpl(answers, float(self.cells_per_letter.get()))
            else:
                size = int(self.size_var.get())

            rng = random.Random()
            grid, entries = build_crossword(
                answers, size,
                allow_reversed=self.allow_reversed.get(),
                require_connected=self.require_connected.get(),
                min_overlap=int(self.min_overlap.get()),
                max_backtracks=30000, rng=rng
            )

            # Optional symmetry enforcement (best-effort)
            self.symmetry_violations = 0
            if self.enforce_symmetry.get():
                removed = enforce_rotational_symmetry_blocks(grid)
                if removed:
                    self.symmetry_violations = removed
                    # After removing letters, rebuild entries from grid
                    entries = extract_entries_from_grid(grid)

            self.grid_data = grid
            self.entries = entries
            # update labels / clues
            self.words_label.config(text="Entries: " + ", ".join(sorted({e["answer"] for e in entries})))
            if self.symmetry_violations:
                self.sym_label.config(text=f"Symmetry: fixed {self.symmetry_violations} conflicting cell(s)")
            else:
                self.sym_label.config(text="Symmetry: OK")
            self.refresh_clue_views()
            self.update_preview()
        except RuntimeError as e:
            messagebox.showerror("Cannot fit answers", f"{e}\n\nTry: increase grid size, reduce answers/length, lower Min overlaps, or disable 'Require connected'.")
        except Exception as e:
            messagebox.showerror("Error", repr(e))

    def symmetrize_now(self):
        if not self.grid_data: return
        self.snapshot()
        removed = enforce_rotational_symmetry_blocks(self.grid_data)
        if removed:
            self.symmetry_violations = removed
            self.entries = extract_entries_from_grid(self.grid_data)
            self.words_label.config(text="Entries: " + ", ".join(sorted({e["answer"] for e in self.entries})))
            self.refresh_clue_views()
            self.update_preview()
            messagebox.showinfo("Symmetry", f"Made blocks rotationally symmetric. Removed {removed} conflicting letter cell(s).")
        else:
            messagebox.showinfo("Symmetry", "Blocks were already rotationally symmetric.")

    # ---------- CLUES ----------
    def build_across_down_lists(self):
        if not (self.grid_data and self.entries): return [], []
        cell_number, _ = compute_numbering_nyt(self.grid_data, self.entries)
        across, down = [], []
        for e in self.entries:
            num = cell_number.get((e["r"], e["c"]))
            item = (num, e["answer"], self.clues_map.get(e["answer"], ""))
            (across if e["dir"]=='A' else down).append(item)
        across.sort(key=lambda x: (x[0] or 10**9, x[1]))
        down.sort(key=lambda x: (x[0] or 10**9, x[1]))
        return across, down

    def refresh_clue_views(self):
        a, d = self.build_across_down_lists()
        def fill(tree, data):
            tree.delete(*tree.get_children())
            for num, ans, clue in data:
                tree.insert("", "end", values=(num if num else "", ans, clue))
        fill(self.across_tree, a)
        fill(self.down_tree, d)

    def fill_empty_clues(self):
        for e in self.entries or []:
            self.clues_map.setdefault(e["answer"], e["answer"])
        self.refresh_clue_views()

    # ---------- PREVIEW ----------
    def _current_palette(self) -> Dict[str,str]:
        if self.highlight_mode.get() != "palette": return {}
        if not self.entries: return {}
        keys = [f'{e["dir"]}:{e["r"]}:{e["c"]}:{e["answer"]}'.upper() for e in self.entries]
        mp = parse_palette_str(self.palette_text.get(), keys)
        if not mp and keys:
            cols = auto_palette(len(keys))
            for i, k in enumerate(keys):
                mp[k] = cols[i % len(cols)]
        return mp

    def update_preview(self):
        if not self.grid_data or not self.entries: return
        n = len(self.grid_data)
        base_cell = max(18, min(72, max(MIN_PREVIEW // n, min(MAX_PREVIEW // n, self.cell_px.get()))))
        palette = self._current_palette()
        img = render_crossword_png(
            self.grid_data, self.entries, self.clues_map, png_path=None,
            show_solution=self.show_solution.get(), cell_px=base_cell, margin_cells=self.margin_cells.get(),
            default_highlight_hex=self.highlight_hex.get(), alpha01=max(0,min(100,self.alpha_pct.get()))/100.0,
            per_key_palette=palette,
            grid_thickness_px=self.grid_thickness_px.get(),
            grid_style=self.grid_line_style.get(), grid_color_hex=self.grid_line_color_hex.get(),
            letter_font_path=self.letter_font_path.get() or None, letter_font_px=self.letter_font_px.get(),
            letter_color_hex=self.letter_color_hex.get(),
            letter_outline_px=self.letter_outline_px.get(), letter_outline_hex=self.letter_outline_hex.get(),
            number_font_path=self.number_font_path.get() or None, number_font_px=self.number_font_px.get(),
            number_color_hex=self.number_color_hex.get(),
            number_bold=self.number_bold.get(), number_italic=self.number_italic.get(), number_underline=self.number_underline.get(),
            open_cell_bg_hex=self.open_cell_bg_hex.get(),
            digit_script=self.digit_script.get(),
            orientation=self.orientation.get(), vertical_text=self.vertical_text.get(),
            scale_factor=1.0, dpi=self.png_dpi.get(),
            align_h=self.letter_align_h.get(), align_v=self.letter_align_v.get(),
            bold=self.letter_bold.get(), italic=self.letter_italic.get(), underline=self.letter_underline.get()
        )
        self.preview_image = ImageTk.PhotoImage(img)
        self.canvas.config(width=img.width, height=img.height); self.canvas.delete("all")
        self.canvas.create_image(0, 0, anchor="nw", image=self.preview_image)

    # ---------- EXPORT ----------
    def export_single_dialog(self):
        if not self.grid_data or not self.entries:
            messagebox.showinfo("Nothing to export", "Generate a crossword first."); return
        stub = filedialog.asksaveasfilename(
            title="Save as… (choose a base name; extensions added automatically)",
            defaultextension=".docx",
            filetypes=[("DOCX", "*.docx"), ("All files", "*.*")]
        )
        if not stub: return
        if "." in stub: stub = stub.rsplit(".", 1)[0]
        # number the puzzle
        puzzle_no = self.puzzle_counter
        self.puzzle_counter += 1
        base = f"{stub}_puzzle_#{puzzle_no}"
        base_sol = f"{stub}_solution_#{puzzle_no}"
        try:
            palette = self._current_palette()
            alpha01 = max(0, min(100, self.alpha_pct.get()))/100.0

            # DOCX with clue columns
            to_docx_crossword_with_clues(
                self.grid_data, self.entries, f"{base}.docx",
                self.highlight_hex.get(), palette,
                self.docx_border_pt.get(), self.grid_line_style.get(), self.grid_line_color_hex.get().lstrip("#").upper(),
                os.path.basename(self.letter_font_path.get()) or None, self.letter_font_pt.get(), self.letter_color_hex.get(),
                self.number_font_pt.get(), self.number_color_hex.get(),
                os.path.basename(self.title_font_path.get()) or None, self.title_font_pt.get(), self.title_color_hex.get(),
                os.path.basename(self.list_font_path.get()) or None, self.list_font_pt.get(), self.list_color_hex.get(),
                digit_script=self.digit_script.get(),
                number_bold=self.number_bold.get(), number_italic=self.number_italic.get(), number_underline=self.number_underline.get(),
                open_cell_bg_hex=self.open_cell_bg_hex.get(),
                letter_bold=self.letter_bold.get(), letter_italic=self.letter_italic.get(), letter_underline=self.letter_underline.get(),
                with_solution=False, puzzle_title=f"Crossword — Puzzle #{puzzle_no}", clues_map=self.clues_map
            )
            to_docx_crossword_with_clues(
                self.grid_data, self.entries, f"{base_sol}.docx",
                self.highlight_hex.get(), palette,
                self.docx_border_pt.get(), self.grid_line_style.get(), self.grid_line_color_hex.get().lstrip("#").upper(),
                os.path.basename(self.letter_font_path.get()) or None, self.letter_font_pt.get(), self.letter_color_hex.get(),
                self.number_font_pt.get(), self.number_color_hex.get(),
                os.path.basename(self.title_font_path.get()) or None, self.title_font_pt.get(), self.title_color_hex.get(),
                os.path.basename(self.list_font_path.get()) or None, self.list_font_pt.get(), self.list_color_hex.get(),
                digit_script=self.digit_script.get(),
                number_bold=self.number_bold.get(), number_italic=self.number_italic.get(), number_underline=self.number_underline.get(),
                open_cell_bg_hex=self.open_cell_bg_hex.get(),
                letter_bold=self.letter_bold.get(), letter_italic=self.letter_italic.get(), letter_underline=self.letter_underline.get(),
                with_solution=True, puzzle_title=f"Crossword — Solution #{puzzle_no}", clues_map=self.clues_map
            )

            # PDF
            to_pdf_crossword_with_clues(
                self.grid_data, self.entries, f"{base}.pdf",
                self.highlight_hex.get(), palette, alpha01,
                self.grid_thickness_pt.get(), self.grid_line_style.get(), self.grid_line_color_hex.get(),
                self.letter_font_path.get() or None, self.letter_font_pt.get(), self.letter_color_hex.get(),
                self.letter_outline_pt.get(), self.letter_outline_hex.get(),
                self.number_font_path.get() or None, self.number_font_pt.get(), self.number_color_hex.get(),
                self.title_font_path.get() or None, self.title_font_pt.get(), self.title_color_hex.get(),
                self.list_font_path.get() or None, self.list_font_pt.get(), self.list_color_hex.get(),
                digit_script=self.digit_script.get(),
                number_bold=self.number_bold.get(), number_italic=self.number_italic.get(), number_underline=self.number_underline.get(),
                open_cell_bg_hex=self.open_cell_bg_hex.get(),
                cell_pt=self.pdf_cell_pt.get(), show_solution=False,
                puzzle_title=f"Crossword — Puzzle #{puzzle_no}", clues_map=self.clues_map
            )
            to_pdf_crossword_with_clues(
                self.grid_data, self.entries, f"{base_sol}.pdf",
                self.highlight_hex.get(), palette, alpha01,
                self.grid_thickness_pt.get(), self.grid_line_style.get(), self.grid_line_color_hex.get(),
                self.letter_font_path.get() or None, self.letter_font_pt.get(), self.letter_color_hex.get(),
                self.letter_outline_pt.get(), self.letter_outline_hex.get(),
                self.number_font_path.get() or None, self.number_font_pt.get(), self.number_color_hex.get(),
                self.title_font_path.get() or None, self.title_font_pt.get(), self.title_color_hex.get(),
                self.list_font_path.get() or None, self.list_font_pt.get(), self.list_color_hex.get(),
                digit_script=self.digit_script.get(),
                number_bold=self.number_bold.get(), number_italic=self.number_italic.get(), number_underline=self.number_underline.get(),
                open_cell_bg_hex=self.open_cell_bg_hex.get(),
                cell_pt=self.pdf_cell_pt.get(), show_solution=True,
                puzzle_title=f"Crossword — Solution #{puzzle_no}", clues_map=self.clues_map
            )

            # PNG & SVG
            render_crossword_png(self.grid_data, self.entries, self.clues_map, f"{base}.png", show_solution=False,
                                 cell_px=self.cell_px.get(), margin_cells=self.margin_cells.get(),
                                 default_highlight_hex=self.highlight_hex.get(), alpha01=alpha01,
                                 per_key_palette=palette, grid_thickness_px=self.grid_thickness_px.get(),
                                 grid_style=self.grid_line_style.get(), grid_color_hex=self.grid_line_color_hex.get(),
                                 letter_font_path=self.letter_font_path.get() or None, letter_font_px=self.letter_font_px.get(),
                                 letter_color_hex=self.letter_color_hex.get(), letter_outline_px=self.letter_outline_px.get(),
                                 letter_outline_hex=self.letter_outline_hex.get(),
                                 number_font_path=self.number_font_path.get() or None, number_font_px=self.number_font_px.get(),
                                 number_color_hex=self.number_color_hex.get(),
                                 number_bold=self.number_bold.get(), number_italic=self.number_italic.get(), number_underline=self.number_underline.get(),
                                 open_cell_bg_hex=self.open_cell_bg_hex.get(),
                                 digit_script=self.digit_script.get(),
                                 orientation=self.orientation.get(), vertical_text=self.vertical_text.get(),
                                 scale_factor=float(self.png_scale.get()), dpi=int(self.png_dpi.get()),
                                 align_h=self.letter_align_h.get(), align_v=self.letter_align_v.get(),
                                 bold=self.letter_bold.get(), italic=self.letter_italic.get(), underline=self.letter_underline.get())
            render_crossword_png(self.grid_data, self.entries, self.clues_map, f"{base_sol}.png", show_solution=True,
                                 cell_px=self.cell_px.get(), margin_cells=self.margin_cells.get(),
                                 default_highlight_hex=self.highlight_hex.get(), alpha01=alpha01,
                                 per_key_palette=palette, grid_thickness_px=self.grid_thickness_px.get(),
                                 grid_style=self.grid_line_style.get(), grid_color_hex=self.grid_line_color_hex.get(),
                                 letter_font_path=self.letter_font_path.get() or None, letter_font_px=self.letter_font_px.get(),
                                 letter_color_hex=self.letter_color_hex.get(), letter_outline_px=self.letter_outline_px.get(),
                                 letter_outline_hex=self.letter_outline_hex.get(),
                                 number_font_path=self.number_font_path.get() or None, number_font_px=self.number_font_px.get(),
                                 number_color_hex=self.number_color_hex.get(),
                                 number_bold=self.number_bold.get(), number_italic=self.number_italic.get(), number_underline=self.number_underline.get(),
                                 open_cell_bg_hex=self.open_cell_bg_hex.get(),
                                 digit_script=self.digit_script.get(),
                                 orientation=self.orientation.get(), vertical_text=self.vertical_text.get(),
                                 scale_factor=float(self.png_scale.get()), dpi=int(self.png_dpi.get()),
                                 align_h=self.letter_align_h.get(), align_v=self.letter_align_v.get(),
                                 bold=self.letter_bold.get(), italic=self.letter_italic.get(), underline=self.letter_underline.get())

            to_svg_crossword(self.grid_data, self.entries, f"{base}.svg", cell_px=self.cell_px.get(), margin_cells=self.margin_cells.get(),
                             default_highlight_hex=self.highlight_hex.get(), alpha01=alpha01, per_key_palette=palette,
                             grid_thickness_px=self.grid_thickness_px.get(), grid_style=self.grid_line_style.get(), grid_color_hex=self.grid_line_color_hex.get(),
                             letter_font_family=os.path.basename(self.letter_font_path.get()) or "DejaVu Sans",
                             letter_px=self.letter_font_px.get(), letter_color_hex=self.letter_color_hex.get(),
                             number_font_family=os.path.basename(self.number_font_path.get()) or "DejaVu Sans",
                             number_px=self.number_font_px.get(), number_color_hex=self.number_color_hex.get(),
                             with_solution=False, digit_script=self.digit_script.get(),
                             number_bold=self.number_bold.get(), number_italic=self.number_italic.get(), number_underline=self.number_underline.get(),
                             open_cell_bg_hex=self.open_cell_bg_hex.get())
            to_svg_crossword(self.grid_data, self.entries, f"{base_sol}.svg", cell_px=self.cell_px.get(), margin_cells=self.margin_cells.get(),
                             default_highlight_hex=self.highlight_hex.get(), alpha01=alpha01, per_key_palette=palette,
                             grid_thickness_px=self.grid_thickness_px.get(), grid_style=self.grid_line_style.get(), grid_color_hex=self.grid_line_color_hex.get(),
                             letter_font_family=os.path.basename(self.letter_font_path.get()) or "DejaVu Sans",
                             letter_px=self.letter_font_px.get(), letter_color_hex=self.letter_color_hex.get(),
                             number_font_family=os.path.basename(self.number_font_path.get()) or "DejaVu Sans",
                             number_px=self.number_font_px.get(), number_color_hex=self.number_color_hex.get(),
                             with_solution=True, digit_script=self.digit_script.get(),
                             number_bold=self.number_bold.get(), number_italic=self.number_italic.get(), number_underline=self.number_underline.get(),
                             open_cell_bg_hex=self.open_cell_bg_hex.get())

            if self.also_json.get():
                data = self._build_json_payload(title=f"Crossword #{puzzle_no}")
                with open(f"{base}.json", "w", encoding="utf-8") as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)

            messagebox.showinfo("Exported",
                                f"Saved:\n{base}.docx\n{base_sol}.docx\n"
                                f"{base}.pdf\n{base_sol}.pdf\n{base}.svg\n{base_sol}.svg\n"
                                f"{base}.png\n{base_sol}.png" + ("\n" + f"{base}.json" if self.also_json.get() else ""))
        except Exception as e:
            messagebox.showerror("Export error", repr(e))

    def _build_json_payload(self, title: str) -> Dict:
        cell_number, _ = compute_numbering_nyt(self.grid_data, self.entries)
        across = []; down = []
        for e in self.entries:
            number = cell_number.get((e["r"], e["c"]))
            item = {
                "number": number, "answer": e["answer"], "dir": e["dir"],
                "row": e["r"], "col": e["c"], "coords": e["coords"],
                "clue": self.clues_map.get(e["answer"], "")
            }
            (across if e["dir"]=='A' else down).append(item)
        across.sort(key=lambda x: (x["number"], x["row"], x["col"]))
        down.sort(key=lambda x: (x["number"], x["row"], x["col"]))
        return {
            "title": title,
            "grid_size": len(self.grid_data),
            "grid": ["".join(ch if ch else "#" for ch in self.grid_data[r]) for r in range(len(self.grid_data))],
            "across": across, "down": down,
            "settings": self.current_settings()
        }

    # ---------- BULK (reuse base, minor changes omitted for brevity) ----------
    # (Keep the original bulk export methods if you used them; not repeated to save space)

    # ---------- SETTINGS ----------
    def current_settings(self) -> Dict:
        return {
            "highlight_mode": self.highlight_mode.get(),
            "highlight_hex": self.highlight_hex.get(),
            "alpha_pct": int(self.alpha_pct.get()),
            "palette_text": self.palette_text.get(),
            "grid_thickness_px": int(self.grid_thickness_px.get()),
            "grid_thickness_pt": float(self.grid_thickness_pt.get()),
            "docx_border_pt": float(self.docx_border_pt.get()),
            "grid_line_style": self.grid_line_style.get(),
            "grid_line_color_hex": self.grid_line_color_hex.get(),
            "open_cell_bg_hex": self.open_cell_bg_hex.get(),
            "cell_px": int(self.cell_px.get()),
            "margin_cells": int(self.margin_cells.get()),
            "pdf_cell_pt": float(self.pdf_cell_pt.get()),
            "png_scale": float(self.png_scale.get()),
            "png_dpi": int(self.png_dpi.get()),
            "letter_font_path": self.letter_font_path.get(),
            "letter_font_px": int(self.letter_font_px.get()),
            "letter_font_pt": float(self.letter_font_pt.get()),
            "letter_color_hex": self.letter_color_hex.get(),
            "letter_outline_px": int(self.letter_outline_px.get()),
            "letter_outline_pt": float(self.letter_outline_pt.get()),
            "letter_outline_hex": self.letter_outline_hex.get(),
            "letter_align_h": self.letter_align_h.get(),
            "letter_align_v": self.letter_align_v.get(),
            "letter_bold": bool(self.letter_bold.get()),
            "letter_italic": bool(self.letter_italic.get()),
            "letter_underline": bool(self.letter_underline.get()),
            "number_font_path": self.number_font_path.get(),
            "number_font_px": int(self.number_font_px.get()),
            "number_font_pt": float(self.number_font_pt.get()),
            "number_color_hex": self.number_color_hex.get(),
            "number_bold": bool(self.number_bold.get()),
            "number_italic": bool(self.number_italic.get()),
            "number_underline": bool(self.number_underline.get()),
            "digit_script": self.digit_script.get(),
            "orientation": self.orientation.get(),
            "vertical_text": bool(self.vertical_text.get()),
            "title_font_path": self.title_font_path.get(),
            "title_font_pt": float(self.title_font_pt.get()),
            "title_color_hex": self.title_color_hex.get(),
            "list_font_path": self.list_font_path.get(),
            "list_font_pt": float(self.list_font_pt.get()),
            "list_color_hex": self.list_color_hex.get(),
            "auto_size_mode": self.auto_size_mode.get(),
            "density_pct": int(self.density_pct.get()),
            "cells_per_word": int(self.cells_per_word.get()),
            "cells_per_letter": float(self.cells_per_letter.get()),
            "require_connected": bool(self.require_connected.get()),
            "min_overlap": int(self.min_overlap.get()),
            "allow_reversed": bool(self.allow_reversed.get()),
            "enforce_symmetry": bool(self.enforce_symmetry.get()),
            "puzzle_counter_next": self.puzzle_counter
        }
    # ---------- PRESETS ----------
    def save_preset_json(self):
        """Save current UI settings to a JSON preset."""
        data = self.current_settings()
        path = filedialog.asksaveasfilename(
            title="Save preset as…",
            defaultextension=".json",
            filetypes=[("JSON", "*.json"), ("All files", "*.*")]
        )
        if not path:
            return
        try:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            messagebox.showinfo("Preset saved", f"Saved {os.path.basename(path)}")
        except Exception as e:
            messagebox.showerror("Save failed", repr(e))

    def load_preset_json(self):
        """Load settings from a JSON preset and apply to the UI."""
        path = filedialog.askopenfilename(
            title="Load preset…",
            filetypes=[("JSON", "*.json"), ("All files", "*.*")]
        )
        if not path:
            return
        try:
            with open(path, "r", encoding="utf-8") as f:
                cfg = json.load(f)
        except Exception as e:
            messagebox.showerror("Load failed", repr(e))
            return

        self._apply_settings_dict(cfg)
        # refresh UI & preview
        self.update_auto_widgets_state()
        self.update_bulk_auto_widgets_state()
        self.refresh_clue_views()
        self.update_preview()
        messagebox.showinfo("Preset loaded", f"Loaded {os.path.basename(path)}")

    def _apply_settings_dict(self, cfg: dict):
        """Apply a settings dict produced by current_settings(). Unknown keys are ignored."""
        def sets(k, var): 
            if k in cfg and cfg[k] is not None: var.set(str(cfg[k]))
        def seti(k, var): 
            if k in cfg and cfg[k] is not None: var.set(int(cfg[k]))
        def setf(k, var): 
            if k in cfg and cfg[k] is not None: var.set(float(cfg[k]))
        def setb(k, var): 
            if k in cfg and cfg[k] is not None: var.set(bool(cfg[k]))

        # Colors / highlight / grid
        sets("highlight_mode", self.highlight_mode)
        sets("highlight_hex", self.highlight_hex)
        seti("alpha_pct", self.alpha_pct)
        sets("palette_text", self.palette_text)

        seti("grid_thickness_px", self.grid_thickness_px)
        setf("grid_thickness_pt", self.grid_thickness_pt)
        setf("docx_border_pt", self.docx_border_pt)
        sets("grid_line_style", self.grid_line_style)
        sets("grid_line_color_hex", self.grid_line_color_hex)
        sets("open_cell_bg_hex", self.open_cell_bg_hex)

        # Geometry / export sizes
        seti("cell_px", self.cell_px)
        seti("margin_cells", self.margin_cells)
        setf("pdf_cell_pt", self.pdf_cell_pt)
        setf("png_scale", self.png_scale)
        seti("png_dpi", self.png_dpi)

        # Letters
        sets("letter_font_path", self.letter_font_path)
        seti("letter_font_px", self.letter_font_px)
        setf("letter_font_pt", self.letter_font_pt)
        sets("letter_color_hex", self.letter_color_hex)
        seti("letter_outline_px", self.letter_outline_px)
        setf("letter_outline_pt", self.letter_outline_pt)
        sets("letter_outline_hex", self.letter_outline_hex)
        sets("letter_align_h", self.letter_align_h)
        sets("letter_align_v", self.letter_align_v)
        setb("letter_bold", self.letter_bold)
        setb("letter_italic", self.letter_italic)
        setb("letter_underline", self.letter_underline)

        # Numbers
        sets("number_font_path", self.number_font_path)
        seti("number_font_px", self.number_font_px)
        setf("number_font_pt", self.number_font_pt)
        sets("number_color_hex", self.number_color_hex)
        setb("number_bold", self.number_bold)
        setb("number_italic", self.number_italic)
        setb("number_underline", self.number_underline)
        sets("digit_script", self.digit_script)

        # Orientation / language
        sets("orientation", self.orientation)
        setb("vertical_text", self.vertical_text)

        # Titles / lists
        sets("title_font_path", self.title_font_path)
        setf("title_font_pt", self.title_font_pt)
        sets("title_color_hex", self.title_color_hex)
        sets("list_font_path", self.list_font_path)
        setf("list_font_pt", self.list_font_pt)
        sets("list_color_hex", self.list_color_hex)

        # Generation behavior
        sets("auto_size_mode", self.auto_size_mode)
        seti("density_pct", self.density_pct)
        seti("cells_per_word", self.cells_per_word)
        setf("cells_per_letter", self.cells_per_letter)
        setb("require_connected", self.require_connected)
        seti("min_overlap", self.min_overlap)
        setb("allow_reversed", self.allow_reversed)
        setb("enforce_symmetry", self.enforce_symmetry)

        # Puzzle counter (optional)
        if "puzzle_counter_next" in cfg and cfg["puzzle_counter_next"]:
            try:
                self.puzzle_counter = int(cfg["puzzle_counter_next"])
            except Exception:
                pass
    def load_csv(self):
        """Load 'Title, directives, answers...' rows from a CSV and paste them into the Bulk textbox."""
        path = filedialog.askopenfilename(
            title="Select CSV (Title, directives, answers...)",
            filetypes=[("CSV files", "*.csv"), ("Text files", "*.txt"), ("All files", "*.*")]
        )
        if not path:
            return
        try:
            rows = []
            # utf-8-sig handles BOMs gracefully
            with open(path, newline="", encoding="utf-8-sig") as f:
                reader = csv.reader(f)
                for row in reader:
                    if not row:
                        continue
                    # Collapse the CSV row back into the expected comma-separated line,
                    # trimming empty cells and whitespace.
                    parts = [item.strip() for item in row if item and item.strip()]
                    if parts:
                        rows.append(",".join(parts))

            if not rows:
                messagebox.showinfo("Empty CSV", "No usable rows found in the selected file.")
                return

            if not hasattr(self, "bulk_text") or self.bulk_text is None:
                messagebox.showerror("UI not ready", "Bulk paste area is not available.")
                return

            self.bulk_text.delete("1.0", "end")
            self.bulk_text.insert("1.0", "\n".join(rows))
            messagebox.showinfo("Loaded", f"Loaded {len(rows)} row(s) from {os.path.basename(path)}")

        except Exception as e:
            messagebox.showerror("CSV error", repr(e))
    def export_bulk(self):
        """Export many crosswords (from CSV/paste or auto-generated lists) to DOCX/PDF/PNG/SVG/JSON."""
        # --- choose output folder ---
        out_dir = filedialog.askdirectory(title="Choose output folder for bulk export")
        if not out_dir:
            return

        # --- small helpers / fallbacks ---
        def _safe_filename(name: str) -> str:
            if hasattr(self, "safe_filename"):
                return self.safe_filename(name)
            cleaned = "".join(ch for ch in name if ch.isalnum() or ch in (" ", "_", "-")).strip()
            return (cleaned or "Crossword").replace(" ", "_")

        def _get(var, default):
            try:
                return var.get()
            except Exception:
                return default

        # --- read defaults from UI (with defensive fallbacks) ---
        mode_kind = _get(self.bulk_mode, "import")  # "import" or "generate"

        default_size = None
        try:
            sv = _get(self.bulk_size_var, "").strip()
            default_size = int(sv) if sv else None
        except Exception:
            default_size = None

        # generation sizing
        mode_auto = _get(self.bulk_auto_mode, "density")
        default_density = int(_get(self.bulk_density_pct, 60))
        default_cpw = int(_get(self.bulk_cpw, 10))
        default_cpl = float(_get(self.bulk_cpl, 1.7))

        # placement behavior
        default_connected = bool(_get(self.bulk_require_connected, True))
        default_minov = int(_get(self.bulk_min_overlap, 1))
        allow_rev = bool(_get(self.bulk_allow_reversed, False))

        # styling & colors
        hl_mode = _get(self.bulk_highlight_mode, "single")
        default_color = _get(self.bulk_highlight_hex, "#FFE696")
        alpha01 = max(0.0, min(1.0, float(_get(self.bulk_alpha_pct, 70)) / 100.0))
        default_palette_text = _get(self.bulk_palette_text, "")

        line_style = _get(self.bulk_grid_line_style, "solid")
        line_color = _get(self.bulk_grid_line_color_hex, "#000000")
        thick_px = max(1, int(_get(self.bulk_grid_thickness_px, 1)))
        thick_pt = max(0.25, float(_get(self.bulk_grid_thickness_pt, 0.75)))
        docx_border_pt = max(0.25, float(_get(self.bulk_docx_border_pt, 0.75)))

        # fonts
        let_font_path = _get(self.bulk_letter_font_path, "") or None
        let_pt = float(_get(self.bulk_letter_font_pt, 12.0))
        let_color = _get(self.bulk_letter_color_hex, "#000000")
        let_outline_pt = float(_get(self.bulk_letter_outline_pt, 0.0))
        let_outline_hex = _get(self.bulk_letter_outline_hex, "#000000")

        num_font_path = _get(self.bulk_number_font_path, "") or None
        num_pt = float(_get(self.bulk_number_font_pt, 8.0))
        num_color = _get(self.bulk_number_color_hex, "#000000")

        tit_font_path = _get(self.bulk_title_font_path, "") or None
        tit_pt = float(_get(self.bulk_title_font_pt, 18.0))
        tit_color = _get(self.bulk_title_color_hex, "#000000")
        lis_font_path = _get(self.bulk_list_font_path, "") or None
        lis_pt = float(_get(self.bulk_list_font_pt, 10.0))
        lis_color = _get(self.bulk_list_color_hex, "#000000")

        # raster export
        png_scale = float(_get(self.bulk_png_scale, 1.0))
        png_dpi = int(_get(self.bulk_png_dpi, 300))

        # seed (optional)
        base_seed = _get(self.bulk_seed, "").strip()
        try:
            base_seed_int = int(base_seed) if base_seed else None
        except Exception:
            base_seed_int = None

        # --- collect jobs ---
        if mode_kind == "import":
            if hasattr(self, "_collect_import_lists"):
                jobs = self._collect_import_lists()
            else:
                # fallback: try to read from a text box named bulk_text
                raw = getattr(self, "bulk_text", None)
                if raw:
                    lines = raw.get("1.0", "end").strip().splitlines()
                    jobs = []
                    for i, line in enumerate(lines):
                        if not line.strip():
                            continue
                        title = f"Crossword_{i+1}"
                        # naive parse: everything after first cell is answers (ANSWER or ANSWER=Clue)
                        parts = [p.strip() for p in line.split(",") if p.strip()]
                        if not parts:
                            continue
                        title = parts[0]
                        answers, clues = [], {}
                        for p in parts[1:]:
                            if "=" in p:
                                a, cl = p.split("=", 1)
                                a = a.strip().replace(" ", "").upper()
                                answers.append(a)
                                clues[a] = cl.strip()
                            elif p.startswith("@"):
                                # ignore directives in fallback
                                continue
                            else:
                                answers.append(p.strip().replace(" ", "").upper())
                        if answers:
                            jobs.append((title, {}, answers, clues))
                else:
                    messagebox.showerror("No data", "Nothing to export. Paste rows or load a CSV in Bulk tab.")
                    return
        else:  # auto-generate
            if hasattr(self, "_collect_generated_lists"):
                jobs = self._collect_generated_lists()
            else:
                messagebox.showerror("Unavailable", "Auto generation helper is missing.")
                return

        if not jobs:
            messagebox.showinfo("No lists", "No word lists to process.")
            return

        # --- bulk loop ---
        successes = failures = 0
        fail_msgs = []

        for idx, (title, params, answers, clues) in enumerate(jobs):
            try:
                # per-row overrides
                if base_seed_int is not None:
                    random.seed(base_seed_int + idx)
                if "seed" in params:
                    try:
                        random.seed(int(params["seed"]))
                    except Exception:
                        pass

                connected = bool(params.get("connected", default_connected))
                try:
                    minov = int(params.get("minov", default_minov))
                except Exception:
                    minov = default_minov

                # size / auto sizing
                try:
                    size = int(params.get("size")) if "size" in params else default_size
                except Exception:
                    size = default_size

                row_auto = (params.get("automode") or mode_auto).lower()
                density = int(params.get("density", default_density))
                cpw = int(params.get("cpw", default_cpw))
                cpl = float(params.get("cpl", default_cpl))

                if size is None:
                    if row_auto in ("cells/word", "cpw"):
                        size = compute_size_by_cpw(answers, cpw)
                    elif row_auto in ("cells/letter", "cpl"):
                        size = compute_size_by_cpl(answers, cpl)
                    else:
                        size = compute_size_by_density(answers, density)
                # palette / colors
                if hl_mode == "palette":
                    palette_text = params.get("palette", default_palette_text)
                    per_key_palette = parse_palette_str(palette_text, [a.upper() for a in answers]) if palette_text else {}
                    single_color = default_color
                else:
                    per_key_palette = {}
                    single_color = params.get("color", default_color)

                # --- build crossword ---
                grid, entries = build_crossword(
                    answers, size,
                    allow_reversed=allow_rev,
                    require_connected=connected,
                    min_overlap=minov,
                    max_backtracks=30000,
                    rng=random.Random()
                )

                # --- paths ---
                base = os.path.join(out_dir, _safe_filename(title))

                # --- DOCX (puzzle + solution) ---
                to_docx_crossword(
                    grid, entries, f"{base}.docx",
                    single_color, per_key_palette,
                    docx_border_pt, line_style, line_color.lstrip("#").upper(),
                    os.path.basename(let_font_path or "") or None, let_pt, let_color,
                    num_pt, num_color,
                    os.path.basename(tit_font_path or "") or None, tit_pt, tit_color,
                    os.path.basename(lis_font_path or "") or None, lis_pt, lis_color,
                    with_solution=False
                )
                to_docx_crossword(
                    grid, entries, f"{base}_solution.docx",
                    single_color, per_key_palette,
                    docx_border_pt, line_style, line_color.lstrip("#").upper(),
                    os.path.basename(let_font_path or "") or None, let_pt, let_color,
                    num_pt, num_color,
                    os.path.basename(tit_font_path or "") or None, tit_pt, tit_color,
                    os.path.basename(lis_font_path or "") or None, lis_pt, lis_color,
                    with_solution=True
                )

                # --- PDF (puzzle + solution) ---
                to_pdf_crossword(
                    grid, entries, f"{base}.pdf",
                    single_color, per_key_palette, alpha01,
                    thick_pt, line_style, line_color,
                    let_font_path, let_pt, let_color,
                    let_outline_pt, let_outline_hex,
                    num_font_path, num_pt, num_color,
                    tit_font_path, tit_pt, tit_color,
                    lis_font_path, lis_pt, lis_color,
                    cell_pt=18.0, show_solution=False
                )
                to_pdf_crossword(
                    grid, entries, f"{base}_solution.pdf",
                    single_color, per_key_palette, alpha01,
                    thick_pt, line_style, line_color,
                    let_font_path, let_pt, let_color,
                    let_outline_pt, let_outline_hex,
                    num_font_path, num_pt, num_color,
                    tit_font_path, tit_pt, tit_color,
                    lis_font_path, lis_pt, lis_color,
                    cell_pt=18.0, show_solution=True
                )

                # --- PNG (puzzle + solution) ---
                # Heuristic px from pt for raster
                letter_px = int(max(10, round((let_pt or 12.0) * 1.33)))
                number_px = int(max(8, round((num_pt or 8.0) * 1.33)))
                render_crossword_png(
                    grid, entries, clues, f"{base}_puzzle.png",
                    show_solution=False,
                    cell_px=DEFAULT_CELL_PX, margin_cells=DEFAULT_MARGIN_CELLS,
                    default_highlight_hex=single_color, alpha01=alpha01,
                    per_key_palette=per_key_palette,
                    grid_thickness_px=thick_px, grid_style=line_style, grid_color_hex=line_color,
                    letter_font_path=let_font_path, letter_font_px=letter_px, letter_color_hex=let_color,
                    letter_outline_px=int(round(let_outline_pt*1.33)), letter_outline_hex=let_outline_hex,
                    number_font_path=num_font_path, number_font_px=number_px, number_color_hex=num_color,
                    scale_factor=png_scale, dpi=png_dpi
                )
                render_crossword_png(
                    grid, entries, clues, f"{base}_solution.png",
                    show_solution=True,
                    cell_px=DEFAULT_CELL_PX, margin_cells=DEFAULT_MARGIN_CELLS,
                    default_highlight_hex=single_color, alpha01=alpha01,
                    per_key_palette=per_key_palette,
                    grid_thickness_px=thick_px, grid_style=line_style, grid_color_hex=line_color,
                    letter_font_path=let_font_path, letter_font_px=letter_px, letter_color_hex=let_color,
                    letter_outline_px=int(round(let_outline_pt*1.33)), letter_outline_hex=let_outline_hex,
                    number_font_path=num_font_path, number_font_px=number_px, number_color_hex=num_color,
                    scale_factor=png_scale, dpi=png_dpi
                )

                # --- SVG (puzzle + solution) ---
                to_svg_crossword(
                    grid, entries, f"{base}.svg",
                    cell_px=DEFAULT_CELL_PX, margin_cells=DEFAULT_MARGIN_CELLS,
                    default_highlight_hex=single_color, alpha01=alpha01,
                    per_key_palette=per_key_palette,
                    grid_thickness_px=thick_px, grid_style=line_style, grid_color_hex=line_color,
                    letter_font_family=os.path.basename(let_font_path or "") or "DejaVu Sans",
                    letter_px=letter_px, letter_color_hex=let_color,
                    number_font_family=os.path.basename(num_font_path or "") or "DejaVu Sans",
                    number_px=number_px, number_color_hex=num_color,
                    with_solution=False
                )
                to_svg_crossword(
                    grid, entries, f"{base}_solution.svg",
                    cell_px=DEFAULT_CELL_PX, margin_cells=DEFAULT_MARGIN_CELLS,
                    default_highlight_hex=single_color, alpha01=alpha01,
                    per_key_palette=per_key_palette,
                    grid_thickness_px=thick_px, grid_style=line_style, grid_color_hex=line_color,
                    letter_font_family=os.path.basename(let_font_path or "") or "DejaVu Sans",
                    letter_px=letter_px, letter_color_hex=let_color,
                    number_font_family=os.path.basename(num_font_path or "") or "DejaVu Sans",
                    number_px=number_px, number_color_hex=num_color,
                    with_solution=True
                )

                # --- JSON manifest (with clues & numbering) ---
                cell_number, _ = compute_numbering(grid, entries)
                data = {
                    "title": title,
                    "grid_size": len(grid),
                    "grid": ["".join(ch if ch else "#" for ch in row) for row in grid],  # '#' = block
                    "across": [], "down": [],
                    "settings": {
                        "connected": connected,
                        "min_overlap": int(minov),
                        "allow_reversed": bool(allow_rev),
                        "auto_mode": row_auto,
                        "density": default_density,
                        "cpw": default_cpw,
                        "cpl": default_cpl,
                        "line_style": line_style,
                        "line_color": line_color,
                        "thickness_px": thick_px,
                        "thickness_pt": thick_pt,
                        "docx_border_pt": docx_border_pt
                    }
                }
                for e in entries:
                    item = {
                        "number": cell_number.get((e["r"], e["c"])),
                        "answer": e["answer"],
                        "dir": e["dir"],
                        "row": e["r"], "col": e["c"],
                        "coords": e["coords"],
                        "clue": clues.get(e["answer"], "")
                    }
                    (data["across"] if e["dir"] == 'A' else data["down"]).append(item)

                with open(f"{base}.json", "w", encoding="utf-8") as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)

                successes += 1

            except Exception as e:
                failures += 1
                fail_msgs.append(f"{title}: {repr(e)}")

        # --- summary ---
        msg = f"Exported {successes} crossword(s) to:\n{out_dir}"
        if failures:
            msg += f"\n\n{failures} failed:\n" + "\n".join(fail_msgs[:10])
            if failures > 10:
                msg += "\n…"
        messagebox.showinfo("Bulk export complete", msg)

# -------- DOCX & PDF with clue columns (NYT numbering) --------
def to_docx_crossword_with_clues(grid, entries, docx_path: str,
                      default_highlight_hex: str, per_key_palette: Dict[str,str],
                      border_thickness_pts: float, border_style: str, border_color_hex: str,
                      letter_font_name: Optional[str], letter_pt: float, letter_color_hex: str,
                      number_pt: float, number_color_hex: str,
                      title_font_name: Optional[str], title_pt: float, title_color_hex: str,
                      list_font_name: Optional[str], list_pt: float, list_color_hex: str,
                      digit_script: str = "Latin",
                      number_bold=False, number_italic=False, number_underline=False,
                      open_cell_bg_hex: str = "#FFFFFF",
                      letter_bold=False, letter_italic=False, letter_underline=False,
                      with_solution: bool = False, puzzle_title: str = "Crossword", clues_map: Optional[Dict[str,str]] = None):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        messagebox.showwarning("Missing package", "python-docx not installed. Skipping DOCX export.")
        return
    document = Document()

    # Title
    title = document.add_heading(puzzle_title, 0)
    for run in title.runs:
        if title_font_name: run.font.name = title_font_name
        run.font.size = Pt(title_pt)
        r, g, b = hex_to_rgb_tuple(title_color_hex); run.font.color.rgb = RGBColor(r, g, b)

    # Table grid
    n = len(grid)
    cell_number, _ = compute_numbering_nyt(grid, entries)
    table = document.add_table(rows=n, cols=n)
    table.style = 'Table Grid'

    def set_cell_borders(cell, size_pts: float, color_hex: str = "000000", style: str = "single"):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        style_map = {"solid":"single","single":"single","dashed":"dashed","dotted":"dotted","double":"double"}
        val = style_map.get(style.lower(), "single")
        for edge in ('top','left','bottom','right'):
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), val)
            sz = str(int(round(max(0.25, size_pts) * 8)))
            el.set(qn('w:sz'), sz)
            el.set(qn('w:color'), border_color_hex)
            tcBorders.append(el)
        for child in list(tcPr):
            if child.tag.endswith('tcBorders'):
                tcPr.remove(child)
        tcPr.append(tcBorders)

    for i in range(n):
        row = table.rows[i]
        for j in range(n):
            cell = row.cells[j]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(cell, border_thickness_pts, border_color_hex, border_style)
            run = p.add_run()

            if grid[i][j] is None:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), "000000")
                tcPr.append(shd)
                continue

            # open cell background
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), hex_to_docx_fill(open_cell_bg_hex))
            tcPr.append(shd)

            # number
            num = cell_number.get((i, j))
            if num is not None:
                rnum = p.add_run(localize_digits(num, digit_script) + " ")
                if list_font_name: rnum.font.name = list_font_name
                rnum.font.size = Pt(max(6, number_pt))
                nr, ng, nb = hex_to_rgb_tuple(number_color_hex)
                rnum.font.color.rgb = RGBColor(nr, ng, nb)
                rnum.font.bold = bool(number_bold); rnum.font.italic = bool(number_italic); rnum.font.underline = bool(number_underline)
                rnum.font.superscript = True

            if with_solution:
                rlet = p.add_run(grid[i][j])
                if letter_font_name: rlet.font.name = letter_font_name
                rlet.font.size = Pt(letter_pt)
                lr, lg, lb = hex_to_rgb_tuple(letter_color_hex); rlet.font.color.rgb = RGBColor(lr, lg, lb)
                rlet.font.bold = bool(letter_bold); rlet.font.italic = bool(letter_italic); rlet.font.underline = bool(letter_underline)

    # Clues section
    document.add_paragraph("")  # spacer
    document.add_heading("Clues", level=1)
    across, down = [], []
    for e in entries:
        num = cell_number.get((e["r"], e["c"]))
        clue_text = (clues_map or {}).get(e["answer"], "")
        item = (num, e["answer"], clue_text)
        (across if e["dir"]=='A' else down).append(item)
    across.sort(key=lambda x: (x[0] or 10**9, x[1]))
    down.sort(key=lambda x: (x[0] or 10**9, x[1]))

    def add_list(title, items):
        h = document.add_heading(title, level=2)
        for run in h.runs:
            if list_font_name: run.font.name = list_font_name
            run.font.size = Pt(list_pt)
            r, g, b = hex_to_rgb_tuple(list_color_hex); run.font.color.rgb = RGBColor(r, g, b)
        for num, ans, clue in items:
            p = document.add_paragraph()
            r1 = p.add_run((localize_digits(num, digit_script) if num else "") + " ")
            if list_font_name: r1.font.name = list_font_name
            r1.font.size = Pt(list_pt)
            r1.font.bold = True
            r2 = p.add_run(f"{ans}: ")
            if list_font_name: r2.font.name = list_font_name
            r2.font.size = Pt(list_pt)
            r3 = p.add_run(clue or "")
            if list_font_name: r3.font.name = list_font_name
            r3.font.size = Pt(list_pt)

    add_list("Across", across)
    add_list("Down", down)

    document.save(docx_path)

def to_pdf_crossword_with_clues(grid, entries, pdf_path: str,
                     default_highlight_hex: str, per_key_palette: Dict[str,str], alpha01: float,
                     grid_thickness_pt: float, grid_style: str, grid_color_hex: str,
                     letter_font_path: Optional[str], letter_pt: float, letter_color_hex: str,
                     letter_outline_pt: float, letter_outline_hex: str,
                     number_font_path: Optional[str], number_pt: float, number_color_hex: str,
                     title_font_path: Optional[str], title_pt: float, title_color_hex: str,
                     list_font_path: Optional[str], list_pt: float, list_color_hex: str,
                     digit_script: str = "Latin",
                     number_bold=False, number_italic=False, number_underline=False,
                     open_cell_bg_hex: str = "#FFFFFF",
                     cell_pt: float = 18.0, show_solution: bool = False,
                     puzzle_title: str = "Crossword", clues_map: Optional[Dict[str,str]] = None):
    try:
        from reportlab.lib.pagesizes import letter as LETTER
        from reportlab.pdfgen import canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        messagebox.showwarning("Missing package", "reportlab not installed. Skipping PDF export.")
        return

    n = len(grid)
    c = canvas.Canvas(pdf_path, pagesize=LETTER)
    width, height = LETTER

    def ensure_font(path: Optional[str], fallback: str, bold=False, italic=False):
        if path and os.path.exists(path):
            try:
                font_name = os.path.basename(path)
                if font_name not in pdfmetrics.getRegisteredFontNames():
                    pdfmetrics.registerFont(TTFont(font_name, path))
                return font_name
            except Exception:
                pass
        if fallback == "Helvetica":
            if bold and italic: return "Helvetica-BoldOblique"
            if bold: return "Helvetica-Bold"
            if italic: return "Helvetica-Oblique"
        return fallback

    title_font  = ensure_font(title_font_path, "Helvetica", False, False)
    list_font   = ensure_font(list_font_path, "Helvetica", False, False)
    letter_font = ensure_font(letter_font_path, "Helvetica", False, False)
    number_font = ensure_font(number_font_path, "Helvetica", False, False)

    line_rgb = hex_to_reportlab_rgb01(grid_color_hex)
    letter_rgb = hex_to_reportlab_rgb01(letter_color_hex)
    number_rgb = hex_to_reportlab_rgb01(number_color_hex)
    title_rgb = hex_to_reportlab_rgb01(title_color_hex)
    list_rgb = hex_to_reportlab_rgb01(list_color_hex)
    openrgb = hex_to_reportlab_rgb01(open_cell_bg_hex)
    highrgb = hex_to_reportlab_rgb01(default_highlight_hex)

    cell_number, _ = compute_numbering_nyt(grid, entries)

    # Title
    c.setFont(title_font, title_pt)
    c.setFillColorRGB(*title_rgb)
    c.drawString(72, height - 72, puzzle_title)

    # Grid
    cell = max(10.0, float(cell_pt))
    top = height - 130
    left = 72
    c.setLineWidth(max(0.25, grid_thickness_pt))
    def apply_dash(style: str):
        s = (style or "solid").lower()
        if s == "solid": c.setDash()
        elif s == "dashed": c.setDash(6, 4)
        elif s == "dotted": c.setDash(1, 3)
        else: c.setDash()
    apply_dash(grid_style)
    for i in range(n):
        for j in range(n):
            x = left + j * cell
            y = top - i * cell
            if grid[i][j] is None:
                c.setFillColorRGB(0,0,0)
                c.rect(x, y - cell, cell, cell, stroke=1, fill=1)
                continue
            # open cell bg
            c.setFillColorRGB(*openrgb); c.rect(x, y - cell, cell, cell, stroke=0, fill=1)
            # highlight
            # (skip per-entry palette in PDF for simplicity; could add)
            c.setStrokeColorRGB(*line_rgb); c.rect(x, y - cell, cell, cell, stroke=1, fill=0)

            # number
            num = cell_number.get((i, j))
            if num is not None:
                c.setFont(number_font, number_pt)
                c.setFillColorRGB(*number_rgb)
                txt = localize_digits(num, digit_script)
                c.drawString(x + cell*0.07, y - cell*0.22, txt)

            # letter
            if show_solution and grid[i][j] is not None:
                ch = grid[i][j]
                c.setFont(letter_font, letter_pt)
                c.setFillColorRGB(*letter_rgb)
                w = pdfmetrics.stringWidth(ch, letter_font, letter_pt)
                c.drawString(x + (cell - w)/2.0, y - cell/2.0, ch)

    # Clues
    c.setFont(list_font, list_pt); c.setFillColorRGB(*list_rgb)
    across, down = [], []
    for e in entries:
        num = cell_number.get((e["r"], e["c"]))
        clue = (clues_map or {}).get(e["answer"], "")
        (across if e["dir"]=='A' else down).append((num, e["answer"], clue))
    across.sort(key=lambda x: (x[0] or 10**9, x[1])); down.sort(key=lambda x: (x[0] or 10**9, x[1]))
    y = top - n*cell - 24
    c.drawString(72, y, "Across")
    y -= list_pt + 6
    for num, ans, clue in across:
        line = f'{localize_digits(num, digit_script) if num else ""} {ans}: {clue}'
        c.drawString(72, y, line[:120])
        y -= list_pt + 2
        if y < 72: c.showPage(); y = height - 72; c.setFont(list_font, list_pt); c.setFillColorRGB(*list_rgb)
    y -= 12
    c.drawString(72, y, "Down")
    y -= list_pt + 6
    for num, ans, clue in down:
        line = f'{localize_digits(num, digit_script) if num else ""} {ans}: {clue}'
        c.drawString(72, y, line[:120])
        y -= list_pt + 2
        if y < 72: c.showPage(); y = height - 72; c.setFont(list_font, list_pt); c.setFillColorRGB(*list_rgb)

    c.setDash(); c.showPage(); c.save()

# =================== MAIN ===================
def main():
    random.seed()
    root = tk.Tk()
    try:
        style = ttk.Style()
        if "clam" in style.theme_names(): style.theme_use("clam")
    except Exception:
        pass
    app = CrosswordApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
