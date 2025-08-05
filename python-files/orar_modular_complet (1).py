
import sys
from PyQt5.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QLabel, QPushButton, QLineEdit,
    QTableWidget, QTableWidgetItem, QMessageBox, QComboBox, QDateEdit, QHBoxLayout
)
from PyQt5.QtCore import Qt, QDate
from docx import Document

class OrarPostliceal(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Orar Postliceal")
        self.setGeometry(100, 100, 1200, 700)
        self.layout = QVBoxLayout()
        self.setLayout(self.layout)

        self.clase = {}  # cheie: nume clasă, valoare: (calificare, durata)
        self.modul_dates = {}  # cheie: modul, valoare: (start_date, end_date)
        self.orar = []  # conține tuple (clasa, modul, profesor, materie, ore)

        self.init_ui()

    def init_ui(self):
        self.label_clasa = QLabel("Denumire clasă:")
        self.input_clasa = QLineEdit()
        self.label_calificare = QLabel("Calificare:")
        self.input_calificare = QLineEdit()
        self.label_durata = QLabel("Durată (1 sau 2 ani):")
        self.input_durata = QComboBox()
        self.input_durata.addItems(["1", "2"])
        self.btn_adauga_clasa = QPushButton("Adaugă clasă")
        self.btn_adauga_clasa.clicked.connect(self.adauga_clasa)

        self.clasa_layout = QHBoxLayout()
        self.clasa_layout.addWidget(self.label_clasa)
        self.clasa_layout.addWidget(self.input_clasa)
        self.clasa_layout.addWidget(self.label_calificare)
        self.clasa_layout.addWidget(self.input_calificare)
        self.clasa_layout.addWidget(self.label_durata)
        self.clasa_layout.addWidget(self.input_durata)
        self.clasa_layout.addWidget(self.btn_adauga_clasa)

        self.select_clasa = QComboBox()
        self.select_clasa.currentIndexChanged.connect(self.update_module_list)
        self.select_modul = QComboBox()
        self.select_modul.currentIndexChanged.connect(self.update_interval_label)
        self.label_interval = QLabel("Interval modul: -")

        self.table = QTableWidget()
        self.table.setColumnCount(3)
        self.table.setHorizontalHeaderLabels(["Profesor", "Materie", "Ore"])

        self.btn_export = QPushButton("Exportă în Word")
        self.btn_export.clicked.connect(self.exporta_word)

        self.layout.addLayout(self.clasa_layout)
        self.layout.addWidget(QLabel("Selectează clasă:"))
        self.layout.addWidget(self.select_clasa)
        self.layout.addWidget(QLabel("Selectează modul:"))
        self.layout.addWidget(self.select_modul)
        self.layout.addWidget(self.label_interval)
        self.layout.addWidget(self.table)
        self.layout.addWidget(self.btn_export)

    def adauga_clasa(self):
        nume = self.input_clasa.text().strip()
        calificare = self.input_calificare.text().strip()
        durata = int(self.input_durata.currentText())
        if not nume or not calificare:
            QMessageBox.warning(self, "Eroare", "Completează toate câmpurile pentru clasă!")
            return
        self.clase[nume] = (calificare, durata)
        self.select_clasa.addItem(nume)
        self.input_clasa.clear()
        self.input_calificare.clear()

    def update_module_list(self):
        clasa = self.select_clasa.currentText()
        self.select_modul.clear()
        if clasa:
            durata = self.clase[clasa][1]
            nr_module = durata * 4
            for i in range(1, nr_module + 1):
                self.select_modul.addItem(f"Modul {i}")
        self.update_interval_label()

    def update_interval_label(self):
        modul_index = self.select_modul.currentIndex()
        if modul_index >= 0:
            date_start, date_end = self.get_modul_interval(modul_index)
            self.label_interval.setText(f"Interval modul: {date_start.toString('dd MMM yyyy')} - {date_end.toString('dd MMM yyyy')}")

    def get_modul_interval(self, modul_index):
        baza = [("02.09.2024", "25.10.2024"), ("28.10.2024", "20.12.2024"),
                ("13.01.2025", "07.03.2025"), ("10.03.2025", "30.05.2025")]
        i = modul_index % 4
        start_str, end_str = baza[i]
        start = QDate.fromString(start_str, "dd.MM.yyyy")
        end = QDate.fromString(end_str, "dd.MM.yyyy")
        return (start, end)

    def exporta_word(self):
        document = Document()
        document.add_heading("Orar Postliceal - Modul Curent", level=1)
        clasa = self.select_clasa.currentText()
        modul = self.select_modul.currentText()
        interval = self.label_interval.text().replace("Interval modul: ", "")
        document.add_paragraph(f"Clasa: {clasa}")
        document.add_paragraph(f"{modul} – {interval}")
        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Profesor'
        hdr_cells[1].text = 'Materie'
        hdr_cells[2].text = 'Ore'
        for row in range(self.table.rowCount()):
            cells = table.add_row().cells
            cells[0].text = self.table.item(row, 0).text() if self.table.item(row, 0) else ""
            cells[1].text = self.table.item(row, 1).text() if self.table.item(row, 1) else ""
            cells[2].text = self.table.item(row, 2).text() if self.table.item(row, 2) else ""
        document.save("orar_modul_curent.docx")
        QMessageBox.information(self, "Export", "Fișierul orar_modul_curent.docx a fost salvat.")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = OrarPostliceal()
    window.show()
    sys.exit(app.exec_())
