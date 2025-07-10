import sys
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QVBoxLayout, QPushButton, QTextEdit, QWidget, QScrollArea, QHBoxLayout, QLabel,
    QFileDialog, QDesktopWidget, QToolBar,  QMessageBox,
    QDialog, QTextBrowser,QMenu,QAction,QShortcut,QProgressDialog
)
import webbrowser
from PyQt5.QtGui import QFont, QIcon,QTextCursor, QTextCharFormat, QTextListFormat, QPixmap, QKeySequence
from PyQt5.QtCore import Qt, QSize, QPropertyAnimation, QEasingCurve
import time
import Levenshtein
from docx import Document
from docx.shared import Pt 
import pdfplumber
from PyQt5.QtPrintSupport import QPrinter
import re
import asyncio
import edge_tts
import tempfile
from playsound import playsound
import os

a = time.time()
def load_dictionary(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return set(f.read().splitlines())
    except Exception as e:
        #QMessageBox.critical(None, "خطا", f"خطا در بارگیری دیکشنری: {str(e)}")
        return set()

               # پیدا کردن پیشنهادات
def spell_checker(word, kalamat_dorost, num):
    distances = [(w, Levenshtein.distance(word, w)) for w in kalamat_dorost]
    distances.sort(key=lambda x: x[1])
    return [w[0] for w in distances[:num]]

def add_word_and_refresh(word, dict_path, text_edit, kalamat_dorost, suggestions_layout):
    
    if append_to_dictionary(word, dict_path):  
        kalamat_dorost = load_dictionary(dict_path)  
        check_spelling(text_edit, kalamat_dorost, suggestions_layout, clear_layout)
        QMessageBox.information(None, "موفقیت", f"کلمه '{word}' با موفقیت به دیکشنری اضافه شد")
    else:
        QMessageBox.warning(None, "خطا", "خطا در اضافه کردن کلمه به دیکشنری")

def append_to_dictionary(word, dictionary_path):
    """اضافه کردن کلمه به دیکشنری در صورتی که قبلاً وجود نداشته باشد"""
    try:
        with open(dictionary_path, 'r', encoding='utf-8') as file:
            words_in_dict = file.read().splitlines()

        if word not in words_in_dict:
            with open(dictionary_path, 'a', encoding='utf-8') as file:
                file.write('\n' + word)  
            return True
        else:
            return False  
    except Exception as e:
        print(f"خطا در اضافه کردن کلمه به دیکشنری: {e}")
        return False

            # پاک کردن ویجت‌های سیو در دیکشنری بعد از سیو
def clear_layout(layout):
    while layout.count():
        item = layout.takeAt(0)
        widget = item.widget()
        if widget:
            widget.deleteLater()
def check_spelling(text_edit, kalamat_dorost, suggestions_layout,clear_layout,num):
    content = text_edit.toPlainText()
    clear_layout(suggestions_layout)
    

    words = []
    for word in content.split():

        re.sub(r'[^\w\u0600-\u06FF\u0640]', '', word)
        if not word:
            continue
            
        # بررسی وجود "ها" یا "های" در انتهای کلمه
        if word.endswith('های'):
            main_word = word[:-4]
            suffix = 'های'
            words.append(main_word)
            words.append(suffix)
        elif word.endswith('ها'):
            main_word = word[:-3]
            suffix = 'ها'
            words.append(main_word)
            words.append(suffix)
        #elif word.endswith('ان'):
        #    main_word = word[:-2]
        #    suffix = 'ان'
        #    words.append(main_word)
        #    words.append(suffix)
        
        elif word.endswith('؟'):
            main_word = word[:-1]
            suffix = "؟"
            words.append(main_word)
            words.append(suffix)

        
        elif word.endswith("."):
            main_word = word[:-1]
            suffix = "."
            words.append(main_word)
            words.append(suffix)

        
        elif word.endswith("،"):
            main_word = word[:-1]
            suffix = "،"
            words.append(main_word)
            words.append(suffix)

        
        elif word.endswith('!'):
            main_word = word[:-1]
            suffix = "!"
            words.append(main_word)
            words.append(suffix)

        
        elif word.endswith(":"):
            main_word = word[:-1]
            suffix = ":"
            words.append(main_word)
            words.append(suffix)

        
        elif word.endswith("؛"):
            main_word = word[:-1]
            suffix = "؛"
            words.append(main_word)
            words.append(suffix)

        
        elif word.endswith("$"):
            main_word = word[:-1]
            suffix = "$"
            words.append(main_word)
            words.append(suffix)

        
        elif word.endswith(")"):
            main_word = word[:-1]
            suffix = ")"
            words.append(main_word)
            words.append(suffix)

        
        elif word.endswith("#"):
            main_word = word[:-1]
            suffix = "#"
            words.append(main_word)
            words.append(suffix)

        
        elif word.endswith('؟'):
            main_word = word[:-1]
            suffix = "؟"
            words.append(main_word)
            words.append(suffix)

        elif word.endswith(']'):
            main_word = word[:-1]
            suffix = "]"
            words.append(main_word)
            words.append(suffix)

        elif word.endswith('['):
            main_word = word[:-1]
            suffix = "["
            words.append(main_word)
            words.append(suffix)

        elif word.endswith('{'):
            main_word = word[:-1]
            suffix = "{"
            words.append(main_word)
            words.append(suffix)
        
        elif word.endswith('}'):
            main_word = word[:-1]
            suffix = "}"
            words.append(main_word)
            words.append(suffix)
        
        elif word.endswith('<'):
            main_word = word[:-1]
            suffix = "<"
            words.append(main_word)
            words.append(suffix)
        
        elif word.endswith('>'):
            main_word = word[:-1]
            suffix = ">"
            words.append(main_word)
            words.append(suffix)
        
        elif word.endswith('*'):
            main_word = word[:-1]
            suffix = "*"
            words.append(main_word)
            words.append(suffix)
        
        elif word.endswith('«'):
            main_word = word[:-1]
            suffix = "«"
            words.append(main_word)
            words.append(suffix)
        
        elif word.endswith('»'):
            main_word = word[:-1]
            suffix = "»"
            words.append(main_word)
            words.append(suffix)      

        
        elif word.endswith('"'):
            main_word = word[:-1]
            suffix = '"'
            words.append(main_word)
            words.append(suffix) 

        elif word.endswith("'"):
            main_word = word[:-1]
            suffix = "'"
            words.append(main_word)
            words.append(suffix) 

        else:
            words.append(word)
    
    def is_english(word):
        try:
            word.encode(encoding='utf-8').decode('ascii')
        except UnicodeDecodeError:
            return False
        else:
            return True

    def is_number(word):
        if word.isdigit():
            return True
    
    
        persian_digits = {'۰', '۱', '۲', '۳', '۴', '۵', '۶', '۷', '۸', '۹'}
        if all(c in persian_digits for c in word):
            return True
    
    
        written_numbers = [
            'صفر', 'یک', 'دو', 'سه', 'چهار', 'پنج', 'شش', 'هفت', 'هشت', 'نه',
            'ده', 'یازده', 'دوازده', 'سیزده', 'چهارده', 'پانزده', 'شانزده',
            'هفده', 'هجده', 'نوزده', 'بیست', 'سی', 'چهل', 'پنجاه', 'شصت',
            'هفتاد', 'هشتاد', 'نود', 'صد', 'هزار', 'میلیون', 'میلیارد'
        ]
        if word in written_numbers:
            return True
    
        return False

    def convert_english_to_persian_number(word):
        # تبدیل اعداد انگلیسی به فارسی
        english_to_persian = {
            '0': '۰', '1': '۱', '2': '۲', '3': '۳', '4': '۴',
            '5': '۵', '6': '۶', '7': '۷', '8': '۸', '9': '۹'
        }
        return ''.join([english_to_persian.get(c, c) for c in word])    

        

    for word in words:
        
        if not is_english(word)  and word.strip():
            
            if word not in kalamat_dorost and word not in ['ها', 'های','ان',"؟", ".", "!",
                                                            "،", "؛", ":", "#", "$", "(", ")",
                                                              "[", "]", "{", "}"
                                                           , "'", '"', "«", "»", "*", "@", "%", "^"]:
                suggestions = spell_checker(word, kalamat_dorost,num)
                
                if word.isdigit():
                    label = QLabel(f"⚠️ عدد انگلیسی شناسایی شد: '{word}' (لطفاً از اعداد فارسی استفاده کنید)")
                    suggestions_layout.addWidget(label)
                    #label.setFont(QFont("Cairo", 15))
                    
                    # تبدیل عدد انگلیسی به فارسی و نمایش به عنوان پیشنهاد
                    persian_num = convert_english_to_persian_number(word)
             
                if suggestions:
                    # نمایش پیشنهادات تصحیح
                    label = QLabel(f"❌ '{word}':")
                    label.setFont(QFont("Cairo", 13))
                    button_add= QPushButton("اضافه کردن کلمه به دیکشنری")
                    button_add.setIcon(QIcon("add-ellipse-svgrepo-com.png"))
                    button_add.setStyleSheet('''
                            QPushButton {
                                padding: 7px;
                                background-color: qradialgradient(cx:0.5, cy:0.5, radius:0.5,
                                                  fx:0.5, fy:0.5,
                                                  stop:0 #a94444, stop:1 #d47d7d);
                                color: #000000;
                                border: 5px solid qradialgradient(cx:0.5, cy:0.5, radius:0.5,
                                                  fx:0.5, fy:0.5,
                                                  stop:0 #a94444, stop:1 #d47d7d);
                                border-radius: 7px;
                                font-size: 16px;
                                font-family: "Cairo ExtraBold";
                            }
                            QPushButton:hover {
                                background-color: #d9a3a3;
                                border: 5px solid #d9a3a3;
                                border-radius: 7px;
                            }
                            QPushButton:pressed {
                                background-color: #dec8c8;
                                border: 5px solid #dec8c8;
                                border-radius: 7px;
                            }
                        ''')
                    
                    button_add.clicked.connect(lambda _, w=word: add_word_and_refresh(w, "cleaned_farsifinal_no_slash.dic", text_edit, kalamat_dorost, suggestions_layout))
                    suggestions_layout.addWidget(button_add)
                    suggestions_layout.addWidget(label)

                    for suggestion in suggestions:
                        button = QPushButton(suggestion)
                        button.setStyleSheet('''
                            QPushButton {
                                padding: 7px;
                                background-color: qradialgradient(cx:0.5, cy:0.5, radius:0.5,
                                    fx:0.5, fy:0.5,
                                    stop:0 #5da75a, stop:1 #a2d1a1);;
                                color: #000000;
                                font-size: 14px;
                                font-family: "Cairo";
                                text-shadow: 0px 0px 2px white;
                                border: 5px solid qradialgradient(cx:0.5, cy:0.5, radius:0.5,
                                    fx:0.5, fy:0.5,
                                    stop:0 #5da75a, stop:1 #a2d1a1);;
                                border-radius: 7px;
                            }
                            QPushButton:hover {
                                background-color: #aed4ae;
                                border: 5px solid #aed4ae;
                                border-radius: 7px;
                            }
                            QPushButton:pressed {
                                background-color: #c7d1c7;
                                border: 5px solid #c7d1c7;
                                border-radius: 7px;
                            }
                        ''')
                        button.clicked.connect(lambda _, w=word, s=suggestion: replace_word(text_edit, w, s, suggestions_layout))
                        suggestions_layout.addWidget(button)

def clear_word_suggestions(word, suggestions_layout):
    """حذف تمام ویجت‌های مربوط به پیشنهادات یک کلمه خاص"""
    widgets_to_remove = []
    
    # پیدا کردن تمام ویجت‌های مربوط به این کلمه
    for i in reversed(range(suggestions_layout.count())):
        item = suggestions_layout.itemAt(i)
        widget = item.widget()
        if widget and hasattr(widget, 'linked_word') and widget.linked_word == word:
            widgets_to_remove.append(widget)
    
    # حذف ویجت‌ها
    for widget in widgets_to_remove:
        widget.setParent(None)
        widget.deleteLater()

# جایگزینی کلمه اشتباه با پیشنهاد
def replace_word(text_edit, wrong_word, correct_word,suggestions_layout):
    cursor = text_edit.textCursor()
    cursor.select(QTextCursor.Document)
    text = cursor.selectedText()
    new_text = text.replace(wrong_word, correct_word)
    text_edit.setPlainText(new_text)
        # حذف پیشنهادات مربوط به این کلمه
    clear_word_suggestions(wrong_word, suggestions_layout)

                                   # توابع مربوط به ویرایش متن
def set_bold(bold_button,text_edit):
    geom = bold_button.geometry()

    anim_press = QPropertyAnimation(bold_button, b"geometry")
    anim_press.setDuration(150)
    anim_press.setStartValue(geom)
    anim_press.setEndValue(geom.adjusted(5, 5, -5, -5))
    anim_press.setEasingCurve(QEasingCurve.InQuad)

    anim_release = QPropertyAnimation(bold_button, b"geometry")
    anim_release.setDuration(150)
    anim_release.setStartValue(geom.adjusted(5, 5, -5, -5))
    anim_release.setEndValue(geom)
    anim_release.setEasingCurve(QEasingCurve.OutBounce)

    anim_press.start()
    anim_release.start()

        # ذخیره انیمیشن‌ها برای جلوگیری از پاک شدن توسط Garbage Collector
    bold_button._anim_press = anim_press
    bold_button._anim_release = anim_release

    cursor = text_edit.textCursor()
    format = QTextCharFormat()
    format.setFontWeight(QFont.Bold if not cursor.charFormat().font().bold() else QFont.Normal)
    cursor.mergeCharFormat(format)
    text_edit.setFocus()

def set_italic(italic_button,text_edit):
    geom = italic_button.geometry()

    anim_press = QPropertyAnimation(italic_button, b"geometry")
    anim_press.setDuration(150)
    anim_press.setStartValue(geom)
    anim_press.setEndValue(geom.adjusted(5, 5, -5, -5))
    anim_press.setEasingCurve(QEasingCurve.InQuad)

    anim_release = QPropertyAnimation(italic_button, b"geometry")
    anim_release.setDuration(150)
    anim_release.setStartValue(geom.adjusted(5, 5, -5, -5))
    anim_release.setEndValue(geom)
    anim_release.setEasingCurve(QEasingCurve.OutBounce)

    anim_press.start()
    anim_release.start()

        # ذخیره انیمیشن‌ها برای جلوگیری از پاک شدن توسط Garbage Collector
    italic_button._anim_press = anim_press
    italic_button._anim_release = anim_release

    cursor = text_edit.textCursor()
    format = QTextCharFormat()
    format.setFontItalic(not cursor.charFormat().fontItalic())
    cursor.mergeCharFormat(format)
    text_edit.setFocus()

def set_underline(underline_button,text_edit):
    geom = underline_button.geometry()

    anim_press = QPropertyAnimation(underline_button, b"geometry")
    anim_press.setDuration(150)
    anim_press.setStartValue(geom)
    anim_press.setEndValue(geom.adjusted(5, 5, -5, -5))
    anim_press.setEasingCurve(QEasingCurve.InQuad)

    anim_release = QPropertyAnimation(underline_button, b"geometry")
    anim_release.setDuration(150)
    anim_release.setStartValue(geom.adjusted(5, 5, -5, -5))
    anim_release.setEndValue(geom)
    anim_release.setEasingCurve(QEasingCurve.OutBounce)

    anim_press.start()
    anim_release.start()

        # ذخیره انیمیشن‌ها برای جلوگیری از پاک شدن توسط Garbage Collector
    underline_button._anim_press = anim_press
    underline_button._anim_release = anim_release

    cursor = text_edit.textCursor()
    format = QTextCharFormat()
    format.setFontUnderline(not cursor.charFormat().fontUnderline())
    cursor.mergeCharFormat(format)
    text_edit.setFocus()


def insert_bullet_list(bullet_list_button,text_edit):
    geom = bullet_list_button.geometry()

    anim_press = QPropertyAnimation(bullet_list_button, b"geometry")
    anim_press.setDuration(150)
    anim_press.setStartValue(geom)
    anim_press.setEndValue(geom.adjusted(5, 5, -5, -5))
    anim_press.setEasingCurve(QEasingCurve.InQuad)

    anim_release = QPropertyAnimation(bullet_list_button, b"geometry")
    anim_release.setDuration(150)
    anim_release.setStartValue(geom.adjusted(5, 5, -5, -5))
    anim_release.setEndValue(geom)
    anim_release.setEasingCurve(QEasingCurve.OutBounce)

    anim_press.start()
    anim_release.start()

        # ذخیره انیمیشن‌ها برای جلوگیری از پاک شدن توسط Garbage Collector
    bullet_list_button._anim_press = anim_press
    bullet_list_button._anim_release = anim_release

    cursor = text_edit.textCursor()
    list_format = QTextListFormat()
    list_format.setStyle(QTextListFormat.ListDisc)
    cursor.insertList(list_format)
    text_edit.setFocus()

def insert_numbered_list(numbered_list_button,text_edit):
    geom = numbered_list_button.geometry()

    anim_press = QPropertyAnimation(numbered_list_button, b"geometry")
    anim_press.setDuration(150)
    anim_press.setStartValue(geom)
    anim_press.setEndValue(geom.adjusted(5, 5, -5, -5))
    anim_press.setEasingCurve(QEasingCurve.InQuad)

    anim_release = QPropertyAnimation(numbered_list_button, b"geometry")
    anim_release.setDuration(150)
    anim_release.setStartValue(geom.adjusted(5, 5, -5, -5))
    anim_release.setEndValue(geom)
    anim_release.setEasingCurve(QEasingCurve.OutBounce)

    anim_press.start()
    anim_release.start()

    # ذخیره انیمیشن‌ها برای جلوگیری از پاک شدن توسط Garbage Collector
    numbered_list_button._anim_press = anim_press
    numbered_list_button._anim_release = anim_release

    cursor = text_edit.textCursor()
    list_format = QTextListFormat()
    list_format.setStyle(QTextListFormat.ListDecimal)
    cursor.insertList(list_format)
    text_edit.setFocus()

def save_text(text_edit):

    file_path, _ = QFileDialog.getSaveFileName(None, "ذخیره متن", "", "Text Files (*.txt);;HTML Files (*.html)")
    if file_path:
        with open(file_path, "w", encoding="utf-8") as file:
            if file_path.endswith(".html"):
                file.write(text_edit.toHtml())
            else:
                file.write(text_edit.toPlainText())
            QMessageBox.information(None, "موفقیت", "فایل TXT با موفقیت ذخیره شد!")

def save_as_docx(text_edit):
    """ذخیره متن به عنوان فایل ورد (DOCX)"""
    file_path, _ = QFileDialog.getSaveFileName(
        None, "ذخیره به عنوان فایل ورد", "", "Word Files (*.docx)"
    )
    
    if file_path:
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'B Nazanin'
        font.size = Pt(12)
        text = text_edit.toPlainText()
        doc.add_paragraph(text)
        doc.save(file_path)
        QMessageBox.information(None, "موفقیت", "فایل ورد با موفقیت ذخیره شد!")

def save_as_pdf(text_edit):

    """ذخیره متن به عنوان فایل PDF"""
    file_path, _ = QFileDialog.getSaveFileName(
        None, "ذخیره به عنوان PDF", "", "PDF Files (*.pdf)"
    )
    
    if file_path:
        cursor = text_edit.textCursor()
        text_edit.selectAll()
        text_edit.setFontFamily("B Nazanin")
        text_edit.setFontPointSize(12)
        cursor.clearSelection()
        text_edit.setTextCursor(cursor)

        printer = QPrinter(QPrinter.HighResolution)
        printer.setOutputFormat(QPrinter.PdfFormat)
        printer.setOutputFileName(file_path)
        text_edit.document().print_(printer)
        QMessageBox.information(None, "موفقیت", "فایل PDF با موفقیت ذخیره شد!")
def load_text(text_edit):
    file_path, _ = QFileDialog.getOpenFileName(None, "بارگذاری متن", "", "Text Files (*.txt);;HTML Files (*.html)")
    if file_path:
        with open(file_path, "r", encoding="utf-8") as file:
            if file_path.endswith(".html"):
                text_edit.setHtml(file.read())
            else:
                text_edit.setPlainText(file.read())
            QMessageBox.information(None, "موفقیت", "فایل TXT با موفقیت بارگذاری شد!")

def load_docx(text_edit):
    file_path, _ = QFileDialog.getOpenFileName(None, "بارگذاری فایل", "", "Word Files (*.docx)")
    doc = Document(file_path)
    full_text = ""
    for paragraph in doc.paragraphs:
        full_text += paragraph.text + "\n"
    text_edit.append(full_text)

def load_pdf(text_edit):
    file_path, _ = QFileDialog.getOpenFileName(None, "بارگذاری فایل", "", "Pdf Files (*.pdf)")
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text()
        QMessageBox.information(None, "موفقیت", "فایل PDF با موفقیت بارگذاری شد!")
        text_edit.append(text)
    except Exception as e:
        QMessageBox.information(None, "خطا!", "فایل PDF با موفقیت بارگذاری نشد:(!")
        return None
    
def undo_clicked(button,text_edit):
    geom = button.geometry()

    anim_press = QPropertyAnimation(button, b"geometry")
    anim_press.setDuration(150)
    anim_press.setStartValue(geom)
    anim_press.setEndValue(geom.adjusted(5, 5, -5, -5))
    anim_press.setEasingCurve(QEasingCurve.InQuad)

    anim_release = QPropertyAnimation(button, b"geometry")
    anim_release.setDuration(150)
    anim_release.setStartValue(geom.adjusted(5, 5, -5, -5))
    anim_release.setEndValue(geom)
    anim_release.setEasingCurve(QEasingCurve.OutBounce)

    anim_press.start()
    anim_release.start()

    # ذخیره انیمیشن‌ها برای جلوگیری از پاک شدن توسط Garbage Collector
    button._anim_press = anim_press
    button._anim_release = anim_release

    text_edit.undo()

def redo_clicked(button,text_edit):
    geom = button.geometry()

    anim_press = QPropertyAnimation(button, b"geometry")
    anim_press.setDuration(150)
    anim_press.setStartValue(geom)
    anim_press.setEndValue(geom.adjusted(5, 5, -5, -5))
    anim_press.setEasingCurve(QEasingCurve.InQuad)

    anim_release = QPropertyAnimation(button, b"geometry")
    anim_release.setDuration(150)
    anim_release.setStartValue(geom.adjusted(5, 5, -5, -5))
    anim_release.setEndValue(geom)
    anim_release.setEasingCurve(QEasingCurve.OutBounce)

    anim_press.start()
    anim_release.start()

    # ذخیره انیمیشن‌ها برای جلوگیری از پاک شدن توسط Garbage Collector
    button._anim_press = anim_press
    button._anim_release = anim_release

    text_edit.redo()

def amar(spell_check_button,text_edit, kalamat_dorost, suggestions_layout,data_label,num):
    geom = spell_check_button.geometry()

    anim_press = QPropertyAnimation(spell_check_button, b"geometry")
    anim_press.setDuration(150)
    anim_press.setStartValue(geom)
    anim_press.setEndValue(geom.adjusted(5, 5, -5, -5))
    anim_press.setEasingCurve(QEasingCurve.InQuad)

    anim_release = QPropertyAnimation(spell_check_button, b"geometry")
    anim_release.setDuration(150)
    anim_release.setStartValue(geom.adjusted(5, 5, -5, -5))
    anim_release.setEndValue(geom)
    anim_release.setEasingCurve(QEasingCurve.OutBounce)

    anim_press.start()
    anim_release.start()

    # ذخیره انیمیشن‌ها برای جلوگیری از پاک شدن توسط Garbage Collector
    spell_check_button._anim_press = anim_press
    spell_check_button._anim_release = anim_release

    text = text_edit.toPlainText()
    word_count = len(text.split())
    char_count = len(text)
    
    # اگر تابع check_spelling مشکلی داشت، از مقادیر پیش‌فرض استفاده کن
    try:
        
        words_with_sugs, total_sugs = check_spelling(text_edit, kalamat_dorost, suggestions_layout, clear_layout,num)
    except (TypeError, NameError):  # اگر تابع None برگرداند یا تعریف نشده باشد
        words_with_sugs, total_sugs = 0, 0  # مقادیر پیش‌فرض
    
    data_label.setFont(QFont("Cairo", 13))
    
    label_style = """
QLabel {
    background-color: #ffffff;
    border: 2px solid #e37712;
    border-radius: 10px;
    font-family: "Cairo";
    
    color: #000000;
    padding: 15px;
    margin: 10px;
    line-height: 1.6;
}
"""

    message = f"""
    تعداد کلمات: {word_count} تعداد حروف: {char_count}
    """
    
    data_label.setStyleSheet(label_style)
    
    data_label.setText(message)
    
"""async def speak_and_play(text, sound_button):
    try:
        geom = sound_button.geometry()

        anim_press = QPropertyAnimation(sound_button, b"geometry")
        anim_press.setDuration(150)
        anim_press.setStartValue(geom)
        anim_press.setEndValue(geom.adjusted(5, 5, -5, -5))
        anim_press.setEasingCurve(QEasingCurve.InQuad)

        anim_release = QPropertyAnimation(sound_button, b"geometry")
        anim_release.setDuration(150)
        anim_release.setStartValue(geom.adjusted(5, 5, -5, -5))
        anim_release.setEndValue(geom)
        anim_release.setEasingCurve(QEasingCurve.OutBounce)

        anim_press.start()
        anim_release.start()

        # ذخیره انیمیشن‌ها برای جلوگیری از پاک شدن توسط Garbage Collector
        sound_button._anim_press = anim_press
        sound_button._anim_release = anim_release

        if not text:
            QMessageBox.critical(None, "خطا", "متنی وارد نکردید!")
            return

        # ذخیره موقت در یک فایل
        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp_file:
            tmp_path = tmp_file.name
    
        # تبدیل متن به صدا و ذخیره موقت
        tts = edge_tts.Communicate(text=text, voice="fa-IR-FaridNeural")
        await tts.save(tmp_path)
    
        # پخش صدا
        try:
            playsound(tmp_path)
        finally:
            try:
                os.unlink(tmp_path)
            except:
                pass
                
    except Exception as e:
        QMessageBox.critical(None, "خطا", f"😕مشکلی در پخش صدا به وجود آمده!")

async def save_audio(text, sound_button):
    if not text:
        QMessageBox.critical(None, "خطا", "متنی وارد نکردید!")
        return

    # تبدیل متن به صدا
    tts = edge_tts.Communicate(text=text, voice="fa-IR-FaridNeural")

    # باز کردن دیالوگ ذخیره‌سازی فایل
    options = QFileDialog.Options()
    save_path, _ = QFileDialog.getSaveFileName(None, "ذخیره صدا", "", "MP3 Files (*.mp3);;All Files (*)", options=options)

    if not save_path:
        return  # اگر کاربر از دیالوگ خارج شد

    # ذخیره صدا به مسیر انتخاب شده
    await tts.save(save_path)

    QMessageBox.information(None, "موفقیت", "صوت با موفقیت ذخیره شد!")

def on_play_sound(text_edit,sound_button):
    text = text_edit.toPlainText()  # دریافت متن
    asyncio.run(speak_and_play(text, sound_button))

def on_save_sound(text_edit,sound_button):
    text = text_edit.toPlainText()  # دریافت متن
    asyncio.run(save_audio(text, sound_button))
"""
def add_dict(text_edit):
    reply = QMessageBox.question(
    None,
    "تأیید عملیات",
    "آیا مطمئن هستید؟\n .پاک کردن کلمات جدید مقدور نمی‌باشد",
    QMessageBox.Yes | QMessageBox.No
)
    #print(reply)
    if reply==16384:
        matn = text_edit.toPlainText().split()
        try:
            with open("cleaned_farsifinal_no_slash.dic", "r", encoding="utf-8") as file:
                kalamat = [line.strip() for line in file]
        except FileNotFoundError:
            kalamat = []
    
        # تنظیم پیشرفت‌بار
        progress = QProgressDialog("در حال پردازش کلمات...", "انصراف", 0, len(matn), None)
        progress.setWindowTitle("لطفاً صبر کنید")
        progress.setValue(0)
        progress.show()
    
        # یافتن و ذخیره کلمات جدید
        tedad = 0
        new_words = []
    
        for i, kalame in enumerate(matn):
            progress.setValue(i)  # به‌روزرسانی پیشرفت
        
            if kalame not in kalamat and kalame not in new_words:
                tedad += 1
                new_words.append(kalame)
        
            if progress.wasCanceled():
                break
    
    # ذخیره کلمات جدید در فایل (یک بار در انتها)
        if new_words:
            with open("cleaned_farsifinal_no_slash.dic", "a", encoding="utf-8") as file:
                for word in new_words:
                    file.write(f"{word}\n")  # اضافه کردن کاراکتر جدید خط
    
        progress.close()
    
        # نمایش پیام مناسب
        natije=QMessageBox()
        natije.setWindowTitle("نتیجه")
        natije_style = """
QMessageBox {
    background-color:  qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #a94444, stop:1 #d47d7d);
    border: 2px solid  qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #a94444, stop:1 #d47d7d);
    border-radius: 30px;
}

QMessageBox QLabel {
    color: #000000;
    line-height: 1.6;
}

QMessageBox QPushButton {
    background-color: #a94444;
    color: #000000;
    border-radius: 5px;
    padding: 8px 16px;
}
"""
        natije.setStyleSheet(natije_style)
        if tedad == 0:
            matn="تمام کلمات موجود در متن از قبل در دیکشنری وجود داشتند."
            #QMessageBox.information(
            #    None,
            #    "نتیجه",
            #    "تمام کلمات موجود در متن از قبل در دیکشنری وجود داشتند."
            #)
            natije.setText(matn)
            natije.setFont(QFont("Cairo ExtraBold", 18))
            natije.exec_()
        else:
            matn=f"{tedad} کلمه جدید به دیکشنری اضافه شد.""\n""این تغییرات باعث بهبود و شخصی‌سازی نرم‌افزار می‌شود."
            #QMessageBox.information(
            #    None,
            #    "عملیات موفق",
            #    f"{tedad} کلمه جدید به دیکشنری اضافه شد.\n"
            #    "این تغییرات باعث بهبود و شخصی‌سازی نرم‌افزار می‌شود."
            #)
            natije.setText(matn)
            natije.setFont(QFont("Cairo ExtraBold", 18))
            natije.exec_()
def info_page():
    info_page=QDialog()
    info_page.setWindowTitle("اطلاعات")
    info_desktop = QDesktopWidget()
    info_layout=QVBoxLayout()

    info_page.setStyleSheet("""background: #ffffff ;""")

    icon_button= QPushButton()
    icon_button.setIcon(QIcon("new_logo (1).png"))
    icon_button.setIconSize(QSize(150,150))
    icon_button.setStyleSheet('''QPushButton {
        padding: 7px;
        background-color: #ffffff;;
        color: #000000;
        border: #ffffff ;
        border-radius: 7px;
   
    }''')

    info_layout.addWidget(icon_button)

    # ایجاد QTextBrowser
    text_browser = QTextBrowser()


# تنظیم متن فارسی
    about_text = """
    ویراستو 
    اصلاح ساختاری متون فارسی

    ویژگی‌ها:
    • بررسی املای کلمات
    • پیشنهاد کلمات جایگزین
    • قابل ذخیره سازی متون به صورت txt, word, pdf 📂
    • رابط کاربری ساده و روان 

    توسعه‌دهندگان:
    - بهراد صفری نژاد (طراح سایت + ایده پرداز + مدیر مسئول سایت)
    - مهدی صحافیان (طراح رابط کاربری و گرافیکی نرم افزار و سایت و اکستنشن + مدیر مسئول نرم‌افزار)
    - رامتین ایزدی (پیاده ساز الگوریتم غلط یابی و اصلاح متن + اصلاح مشکلات + مدیر مسئول الگوریتم غلط یابی)
    -  امیر محمد افشار (طراح لوگو + مسئول تیم تحقیقات + نظارت بر محصول + ایده پرداز + مدیر مسئول اکستنشن)
    وب سایت: https://virasto.github.io/
    کانال بله: https://ble.ir/join/6XnKEWFmMG
    
    با تشکر از:
    تیم پژوهشی کامپیوتر حلی 2

    نسخه : 2   
    © تمام حقوق برای توسعه دهنده گان محفوظ است - ۱۴۰۴
    """

    font = QFont("Cairo Light", 18)
    font.setBold(True)  
    text_browser.setFont(font)
    
    text_browser.setStyleSheet("""
    QTextBrowser {
        padding: 12px;
        color: #000000;
        text-align: right;
        line-height: 1.8;
        direction: rtl;
    }
    
    QTextBrowser a {
        color: #2980b9;
        text-decoration: none;
    }
    
    QTextBrowser a:hover {
        color: #e74c3c;
        text-decoration: underline;
    }
    
    QTextBrowser::scroll-bar:vertical {
        width: 14px;
        background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
            stop:0 #a2d1a1, stop:1 #d47d7d);;
    }
    
    QTextBrowser::scroll-bar::handle:vertical {
        background: #95a5a6;
        min-height: 30px;
        border-radius: 7px;
    }
    
    QTextBrowser::selection {
        background-color: #3498db;
        color: white;
    }
    
""")
    text_browser.setText(about_text)
# تنظیم راست‌چین برای متن فارسی
    text_browser.setAlignment(Qt.AlignRight)
# چیدمان
    info_layout.addWidget(text_browser)
    info_page.showMaximized()
    info_page.setLayout(info_layout)
    info_page.exec_()
def open_website():
    webbrowser.open("https://virasto.github.io/")

def main():
    app = QApplication(sys.argv)
    window = QMainWindow()
    window.setWindowTitle("ویراستار و غلط یاب فارسی!")
    desktop = QDesktopWidget()
    screen_geometry = desktop.screenGeometry()



    window_width = screen_geometry.width()
    window_height = screen_geometry.height()
    window.setGeometry(0, 0, window_width, window_height)
    window.setStyleSheet("background-color: #ffffff;")

    central_widget = QWidget()
    window.setCentralWidget(central_widget)

    main_layout = QHBoxLayout()
    central_widget.setLayout(main_layout)

    left_layout = QVBoxLayout()
    left_layout.setSpacing(8)
    main_layout.addLayout(left_layout)

    right_layout = QVBoxLayout()
    right_layout.setSpacing(10)
    main_layout.addLayout(right_layout)

    right_layout.addSpacing(10)  
    toolbar1 = QToolBar("نوار ابزار ۱")
    toolbar1.setMovable(True)
    toolbar1.setLayoutDirection(Qt.RightToLeft)
    toolbar1_layout = QVBoxLayout()
    toolbar1_layout.setAlignment(Qt.AlignRight)
    toolbar1.setLayout(toolbar1_layout)
    toolbar1.setFloatable(True)  # قابلیت جدا شدن از پنجره اصلی (اختیاری)
    toolbar1.setStyleSheet("""
    QToolBar {
        spacing: 6px;
    }
    """)
    right_layout.addWidget(toolbar1)
    #right_layout.addStretch(1)  

    text_edit = QTextEdit()
    text_edit.setFont(QFont("Cairo", 20))
    text_edit.setPlaceholderText("                                                                                                                               ...متن خود را بنویسید")
    text_edit.setLayoutDirection(Qt.RightToLeft)
    text_edit.setStyleSheet("""
    QTextEdit {
        background-color: #FCFCFC;
        border: 5px solid qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #45b9bd, stop:1 #91d9db);
        border-radius: 10px;
        padding: 10px;
    }
    QScrollBar:vertical {
        background-color: #ffffff;
        width: 10px;
        margin: 15px 15px 15px 15px;
    }
    QScrollBar::handle:vertical {
        background-color: qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #45b9bd, stop:1 #91d9db);
        min-height: 20px;
        border-radius: 5px;
    }
    QScrollBar::add-line:vertical,
    QScrollBar::sub-line:vertical {
        background: none;
    }
    QScrollBar::add-page:vertical,
    QScrollBar::sub-page:vertical {
        background: none;
    }
    """)
    right_layout.addWidget(text_edit)

    suggestions_panel = QScrollArea()
    suggestions_panel.setStyleSheet("""
    QScrollArea {
        background-color: #ffffff;
        border: 5px solid qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #5da75a, stop:1 #a2d1a1);
        border-radius: 15px;
        margin-top: 10px;
        margin-bottom: 10px;
    }
    QScrollBar:vertical {
        background-color: #ffffff;
        width: 10px;
        margin: 0px 0px 0px 0px;
    }
    QScrollBar::handle:vertical {
        background-color: qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #5da75a, stop:1 #a2d1a1);
        min-height: 20px;
        border-radius: 5px;
    }
    QScrollBar::add-line:vertical,
    QScrollBar::sub-line:vertical {
        background: none;
    }
    QScrollBar::add-page:vertical,
    QScrollBar::sub-page:vertical {
        background: none;
    }
    QScrollBar:horizontal {
        background-color: #e0e0e0;
        height: 10px;
        margin: 10px 0px 0px 0px;
    }
    QScrollBar::handle:horizontal {
        background-color: #4CAF50;
        min-width: 20px;
        border-radius: 5px;
    }
    QScrollBar::add-line:horizontal,
    QScrollBar::sub-line:horizontal {
        background: none;
    }
    QScrollBar::add-page:horizontal,
    QScrollBar::sub-page:horizontal {
        background: none;
    }
    """)
    suggestions_panel.setWidgetResizable(True)
    suggestions_panel.setFixedWidth(320)
    suggestions_widget = QWidget()
    suggestions_layout = QVBoxLayout(suggestions_widget)
    suggestions_panel.setWidget(suggestions_widget)

    ban = QLabel()
    pixmap = QPixmap("بنر درست.png")
    ban.setPixmap(pixmap.scaled(400, 90, Qt.KeepAspectRatio, Qt.SmoothTransformation))
    #ban.setText("ویراستو")
    #Font=QFont("Cairo SemiBold",20)
    #ban.setFont(Font)
    left_layout.addWidget(ban)
    left_layout.addWidget(suggestions_panel)

    # بارگیری دیکشنری
    dict_path = "cleaned_farsifinal_no_slash.dic"
    kalamat_dorost = load_dictionary(dict_path)
    if not kalamat_dorost:
        QMessageBox.critical(None, "خطا", "دیکشنری بارگیری نشد! لطفاً مسیر فایل دیکشنری را بررسی کنید.")
        sys.exit(1)

    # ایجاد دکمه‌ها
    bold_button = QPushButton()
    bold_button.setIcon(QIcon('bold-svgrepo-com.png'))
    bold_button.setIconSize(QSize(30, 30))
    bold_button.setStyleSheet('''QPushButton {
        padding: 7px;
        background-color: qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #45b9bd, stop:1 #91d9db);
        color: #000000;
        border: 5px solid qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #45b9bd, stop:1 #91d9db);
        border-radius: 7px;
    }
    QPushButton:hover {
        background-color: #c7eaeb;
        border: 5px solid #c7eaeb;
    }
    QPushButton:pressed {
        background-color: #91d9db;
    }''')
    bold_button.clicked.connect(lambda: set_bold(bold_button,text_edit))
    toolbar1.addWidget(bold_button)
    shortcut_bold = QShortcut(QKeySequence("Ctrl+B"), window)
    shortcut_bold.activated.connect(lambda: set_bold(bold_button,text_edit))    

    italic_button = QPushButton()
    italic_button.setIcon(QIcon('italic-svgrepo-com.png'))
    italic_button.setIconSize(QSize(30, 30))
    italic_button.setStyleSheet('''QPushButton {
        padding: 7px;
        background-color: qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #45b9bd, stop:1 #91d9db);
        color: #000000;
        border: 5px solid qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #45b9bd, stop:1 #91d9db);
        border-radius: 7px;
    }
    QPushButton:hover {
        background-color: #c7eaeb;
        border: 5px solid #c7eaeb;
    }
    QPushButton:pressed {
        background-color: #91d9db;
    }''')
    italic_button.clicked.connect(lambda: set_italic(italic_button,text_edit))
    toolbar1.addWidget(italic_button)
    shortcut_italic = QShortcut(QKeySequence("Ctrl+L"), window)
    shortcut_italic.activated.connect(lambda: set_italic(italic_button,text_edit))

    underline_button = QPushButton()
    underline_button.setIcon(QIcon('underline-669-svgrepo-com.png'))
    underline_button.setIconSize(QSize(30, 30))
    underline_button.setStyleSheet('''QPushButton {
        padding: 7px;
        background-color: qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #45b9bd, stop:1 #91d9db);
        color: #000000;
        border: 5px solid qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #45b9bd, stop:1 #91d9db);
        border-radius: 7px;
    }
    QPushButton:hover {
        background-color: #c7eaeb;
        border: 5px solid #c7eaeb;
    }
    QPushButton:pressed {
        background-color: #91d9db;
    }''')
    underline_button.clicked.connect(lambda: set_underline(underline_button,text_edit))
    toolbar1.addWidget(underline_button)
    shortcut_underline = QShortcut(QKeySequence("Ctrl+U"), window)
    shortcut_underline.activated.connect(lambda: set_underline(underline_button,text_edit))

    redo_button = QPushButton()
    redo_button.setIcon(QIcon('undo-right-svgrepo-com.png'))
    redo_button.setIconSize(QSize(30, 30))
    redo_button.setStyleSheet('''QPushButton {
        padding: 7px;
        background-color: qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #45b9bd, stop:1 #91d9db);
        color: #000000;
        border: 5px solid qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #45b9bd, stop:1 #91d9db);
        border-radius: 7px;
    }
    QPushButton:hover {
        background-color: #c7eaeb;
        border: 5px solid #c7eaeb;
    }
    QPushButton:pressed {
        background-color: #91d9db;
    }''')
    redo_button.clicked.connect(lambda: redo_clicked(redo_button,text_edit))
    toolbar1.addWidget(redo_button)
    shortcut_redo = QShortcut(QKeySequence("Ctrl+Y"), window)
    shortcut_redo.activated.connect(lambda: redo_clicked(redo_button,text_edit))

    undo_button = QPushButton()
    undo_button.setIcon(QIcon('undo-left-svgrepo-com.png'))
    undo_button.setIconSize(QSize(30, 30))
    undo_button.setStyleSheet('''QPushButton {
        padding: 7px;
        background-color: qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #45b9bd, stop:1 #91d9db);
        color: #000000;
        border: 5px solid qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #45b9bd, stop:1 #91d9db);
        border-radius: 7px;
    }
    QPushButton:hover {
        background-color: #c7eaeb;
        border: 5px solid #c7eaeb;
    }
    QPushButton:pressed {
        background-color: #91d9db;
    }''')
    undo_button.clicked.connect(lambda: undo_clicked(undo_button,text_edit))
    toolbar1.addWidget(undo_button)
    shortcut_undo = QShortcut(QKeySequence("Ctrl+Z"), window)
    shortcut_undo.activated.connect(lambda: undo_clicked(undo_button,text_edit))

    bullet_list_button = QPushButton()
    bullet_list_button.setIcon(QIcon("bullets-svgrepo-com.png"))
    bullet_list_button.setIconSize(QSize(30, 30))
    bullet_list_button.setStyleSheet('''QPushButton {
        padding: 7px;
        background-color: qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #45b9bd, stop:1 #91d9db);
        color: #000000;
        border: 5px solid qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #45b9bd, stop:1 #91d9db);
        border-radius: 7px;
    }
    QPushButton:hover {
        background-color: #c7eaeb;
        border: 5px solid #c7eaeb;
    }
    QPushButton:pressed {
        background-color: #91d9db;
    }''')
    bullet_list_button.clicked.connect(lambda: insert_bullet_list(bullet_list_button,text_edit))
    toolbar1.addWidget(bullet_list_button)
    shortcut_bullet = QShortcut(QKeySequence("Ctrl+shift+B"), window)
    shortcut_bullet.activated.connect(lambda: insert_bullet_list(bullet_list_button, text_edit))

    numbered_list_button = QPushButton()
    numbered_list_button.setIcon(QIcon("text-number-list-ltr-svgrepo-com.png"))
    numbered_list_button.setIconSize(QSize(30, 30))
    numbered_list_button.setStyleSheet('''QPushButton {
        padding: 7px;
        background-color: qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #45b9bd, stop:1 #91d9db);
        color: #000000;
        border: 5px solid qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #45b9bd, stop:1 #91d9db);
        border-radius: 7px;
    }
    QPushButton:hover {
        padding: 7px;                           
        background-color: #c7eaeb;
        border: 5px solid #c7eaeb;
        border-radius: 7px;
    }
    QPushButton:pressed {
        background-color: #91d9db;
    }''')
    toolbar1.addWidget(numbered_list_button)
    numbered_list_button.clicked.connect(lambda: insert_numbered_list(numbered_list_button,text_edit))
    shortcut_numbered = QShortcut(QKeySequence("Ctrl+N"), window)
    shortcut_numbered.activated.connect(lambda: insert_numbered_list(numbered_list_button, text_edit))

#   دکمه‌ی اصلی منو
    save_menu_button = QPushButton()
    font = QFont("Cairo", 10)
    save_menu_button.setFont(font)
    save_menu_button.setIcon(QIcon("archive-down-minimlistic-svgrepo-com.png")) 
    save_menu_button.setIconSize(QSize(30, 30))
    save_menu_button.setStyleSheet('''QPushButton {
    padding: 7px;
    background-color: qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #a94444, stop:1 #d47d7d);
    color: #000000;
    border: 5px solid  qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #a94444, stop:1 #d47d7d);
    border-radius: 7px;
}
QPushButton:hover {
    background-color: #d47d7d;
}
QPushButton:pressed {
    background-color: #d47d7d;
    }''')
    shortcut_save = QShortcut(QKeySequence("Ctrl+S"), window)
    shortcut_save.activated.connect(lambda: save_menu.exec_(
    save_menu_button.mapToGlobal(save_menu_button.rect().bottomLeft())
    ))

    # ساخت منو
    save_menu = QMenu()
    save_menu.setCursor(Qt.PointingHandCursor)  # تغییر نشانگر به دست اشاره‌گر
    save_menu.setStyleSheet('''
    QMenu {
        background-color: qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #a94444, stop:1 #d47d7d);
        border: 1px solid #ffffff;
        border-radius: 20px;            /* گردی گوشه‌های منو */
    }
    QMenu::item {
        padding: 10px 20px;  /* افزایش فاصله عمودی و افقی */
        font-size: 14px;     /* اندازه فونت */
        height: 20px;        /* ارتفاع هر آیتم */
    }
    QMenu::item:selected {
        background-color: #d9a3a3;  /* رنگ هنگام hover */
    }
''')
    
    # ایجاد اکشن‌های منو با آیکن، فونت و اتصال به توابع ذخیره‌سازی
    action_docx = QAction(QIcon("docx-file-format-symbol-svgrepo-com.png"), " Docx ذخیره به ", save_menu)
    action_docx.setFont(font)
    action_docx.triggered.connect(lambda: save_as_docx(text_edit))

    action_pdf = QAction(QIcon("pdf-file-format-symbol-svgrepo-com (1).png"), " PDF ذخیره به ", save_menu)
    action_pdf.setFont(font)
    action_pdf.triggered.connect(lambda: save_as_pdf(text_edit))

    action_txt = QAction(QIcon("archive-down-minimlistic-svgrepo-com.png"), " TXT ذخیره به ", save_menu)
    action_txt.setFont(font)
    action_txt.triggered.connect(lambda: save_text(text_edit))

# افزودن اکشن‌ها به منو
    save_menu.addAction(action_docx)
    save_menu.addAction(action_pdf)
    save_menu.addAction(action_txt)

# اتصال منو به دکمه
    save_menu_button.setMenu(save_menu)

# افزودن دکمه منو به نوار ابزار
    toolbar1.addWidget(save_menu_button)

    load = QPushButton()
    font = QFont("Cairo",10)  # انتخاب فونت، اندازه و ضخامت
    load.setFont(font)
    load.setIcon(QIcon('archive-up-minimlistic-svgrepo-com.png'))
    load.setIconSize(QSize(30, 30))
    load.setStyleSheet('''QPushButton {
        padding: 7px;
        background-color:  qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #a94444, stop:1 #d47d7d);
        color: #000000;
        border: 5px solid  qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #a94444, stop:1 #d47d7d);
        border-radius: 7px;
    }
    QPushButton:hover {
        background-color: #d47d7d;
    }
    QPushButton:pressed {
        background-color: #d47d7d;
    }''')

# ساخت منو
    load_menu = QMenu()
    load_menu.setCursor(Qt.PointingHandCursor)  # تغییر نشانگر به دست اشاره‌گر
    load_menu.setStyleSheet('''
    QMenu {
        background-color: qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #a94444, stop:1 #d47d7d);
        border: 1px solid #ffffff;
        border-radius: 20px;            /* گردی گوشه‌های منو */
        
    }
    QMenu::item {
        padding: 10px 20px;  /* افزایش فاصله عمودی و افقی */
        font-size: 14px;     /* اندازه فونت */
        height: 20px;        /* ارتفاع هر آیتم */
    }
    QMenu::item:selected {
        background-color: #d9a3a3;  /* رنگ هنگام hover */
    }
''')
    
    shortcut_load = QShortcut(QKeySequence("Ctrl+O"), window)
    shortcut_load.activated.connect(lambda: load_menu.exec_(
    load_menu.mapToGlobal(load_menu.rect().bottomLeft())
    ))

    # ایجاد اکشن‌های منو با آیکن، فونت و اتصال به توابع ذخیره‌سازی

    action_pdf = QAction(QIcon("pdf-file-format-symbol-svgrepo-com (1).png"), " PDF بارگذاری به ", load_menu)
    action_pdf.setFont(font)
    action_pdf.triggered.connect(lambda: load_pdf(text_edit))


    action_docx = QAction(QIcon("docx-file-format-symbol-svgrepo-com.png"), " Docx بارگذاری", load_menu)
    action_docx.setFont(font)
    action_docx.triggered.connect(lambda: load_docx(text_edit))



    action_txt = QAction(QIcon("archive-up-minimlistic-svgrepo-com.png"), " TXT بارگذاری", load_menu)
    action_txt.setFont(font)
    action_txt.triggered.connect(lambda: load_text(text_edit))

# افزودن اکشن‌ها به منو
    load_menu.addAction(action_pdf)
    load_menu.addAction(action_docx)
    load_menu.addAction(action_txt)

# اتصال منو به دکمه
    load.setMenu(load_menu)

# افزودن دکمه منو به نوار ابزار
    toolbar1.addWidget(load)

    add_dic = QPushButton()
    add_dic.setIcon(QIcon("dictionary-add-svgrepo-com.png"))
    add_dic.setIconSize(QSize(30, 30))
    add_dic.setStyleSheet('''QPushButton {
        padding: 7px;
        background-color:  qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #a94444, stop:1 #d47d7d);
        color: #000000;
        border: 5px solid  qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #a94444, stop:1 #d47d7d);
        border-radius: 7px;
    }
    QPushButton:hover {
        padding: 7px;                           
        background-color: qlineargradient(
            x1:0, y1:0, x2:1, y2:1,
            stop:0 #d47d7d, stop:1 #d47d7d
        );
        border: 5px solid #d47d7d;
        border-radius: 7px;
    }
    QPushButton:pressed {
        background-color: #efefef;
    }''')
    toolbar1.addWidget(add_dic)
    add_dic.clicked.connect(lambda: add_dict(text_edit))

    shortcut_dic = QShortcut(QKeySequence("Ctrl+W"), window)
    #shortcut_dic.activated.connect(lambda: insert_bullet_list(bullet_list_button, text_edit))

    spell_check_button = QPushButton("بررسی")
    font = QFont("Cairo",10)  # انتخاب فونت، اندازه و ضخامت
    spell_check_button.setFont(font)#                                                   
    spell_check_button.setIcon(QIcon('spell-check-svgrepo-com (4).png'))
    spell_check_button.setIconSize(QSize(30, 30))
    spell_check_button.setStyleSheet('''QPushButton {
        padding: 7px;
        background-color: qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #5da75a, stop:1 #a2d1a1);
        color: #000000;
        border: 5px solid qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #5da75a, stop:1 #a2d1a1);
        border-radius: 7px;
    }
    QPushButton:hover {
        padding: 7px;
        background-color: #aed4ae;
        border: 5px solid #aed4ae;
        border-radius: 7px;
    }
    QPushButton:pressed {
        background-color: #c7d1c7;
        padding: 7px;
        background-color: #c7d1c7;
        border: 5px solid #c7d1c7;
        border-radius: 7px;
    }''')
    spell_check_button.clicked.connect(lambda: amar(spell_check_button,text_edit, kalamat_dorost, suggestions_layout,data_label,num=7))
    toolbar1.addWidget(spell_check_button)
    shortcut_spell_check = QShortcut(QKeySequence("Ctrl+Q"), window)
    shortcut_spell_check.activated.connect(lambda: amar(spell_check_button,text_edit, kalamat_dorost, suggestions_layout,data_label,num=7))

    eslah_button = QPushButton("اصلاح کل")
    font = QFont("Cairo",10)  # انتخاب فونت، اندازه و ضخامت
    eslah_button.setFont(font)#                                                   
    eslah_button.setIcon(QIcon('spell-check-svgrepo-com (4).png'))
    eslah_button.setIconSize(QSize(24, 24))
    eslah_button.setStyleSheet('''QPushButton {
        padding: 7px;
        background-color: qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #5da75a, stop:1 #a2d1a1);
        color: #000000;
        border: 5px solid qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #5da75a, stop:1 #a2d1a1);;
        border-radius: 7px;
    }
    QPushButton:hover {
        padding: 7px;
        background-color: #aed4ae;
        border: 5px solid #aed4ae;
        border-radius: 7px;
    }
    QPushButton:pressed {
        background-color: #c7d1c7;
        padding: 7px;
        background-color: #c7d1c7;
        border: 5px solid #c7d1c7;
        border-radius: 7px;
    }''')
    eslah_button.clicked.connect(lambda: amar(eslah_button,text_edit, kalamat_dorost, suggestions_layout,data_label,num=1))
    #toolbar1.addWidget(eslah_button)
    shortcut_spell_check = QShortcut(QKeySequence("Ctrl+shift+Q"), window)
    
    shortcut_spell_check.activated.connect(lambda: amar(eslah_button,text_edit, kalamat_dorost, suggestions_layout,data_label,num=1))

    """sound_button=QPushButton()
    font = QFont("Cairo",10)  # انتخاب فونت، اندازه و ضخامت
    sound_button.setFont(font)
    sound_button.setIcon(QIcon("sound-svgrepo-com.png"))
    sound_button.setIconSize(QSize(30, 30))
    sound_button.setStyleSheet('''QPushButton {
        padding: 7px;
        background-color: qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #2366b3, stop:1 #6c75f5);
        color: #000000;
        border: 5px solid qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #2366b3, stop:1 #6c75f5);
        border-radius: 7px;
    }
    QPushButton:hover {
        padding: 7px;
        background-color: #6c75f5;
        border: 5px solid #6c75f5;
        border-radius: 7px;
    }
    QPushButton:pressed {
        background-color: #424ef5;
        padding: 7px;
        background-color: #424ef5;
        border: 5px solid #424ef5;
        border-radius: 7px;
    }''')
    sound_button.clicked.connect(lambda: asyncio.run(speak_and_play(text_edit.toPlainText(),sound_button)))
    toolbar1.addWidget(sound_button)
    shortcut_voice = QShortcut(QKeySequence("Ctrl+shift+V"), window)
    shortcut_voice.activated.connect(lambda: asyncio.run(speak_and_play(text_edit.toPlainText(),sound_button)))

    shortcut_save_voice=QShortcut(QKeySequence("shift+V"), window)
    shortcut_save_voice.activated.connect(lambda: asyncio.run(save_audio(text_edit.toPlainText(),sound_button)))

    sound_menu=QMenu()
    sound_menu.setCursor(Qt.PointingHandCursor)  # تغییر نشانگر به دست اشاره‌گر
    sound_menu.setStyleSheet('''
    QMenu {
        background-color: qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #2366b3, stop:1 #6c75f5);
        border: 1px solid #ffffff;
        border-radius: 20px;            /* گردی گوشه‌های منو */
    }
    QMenu::item {
        padding: 10px 20px;  /* افزایش فاصله عمودی و افقی */
        font-size: 14px;     /* اندازه فونت */
        height: 20px;        /* ارتفاع هر آیتم */
    }
    QMenu::item:selected {
        background-color: #6c75f5;  /* رنگ هنگام hover */
    }
''')
    
            # ایجاد اکشن‌های منو با آیکن، فونت و اتصال به توابع ذخیره‌سازی
    playsound = QAction(QIcon("sound-svgrepo-com.png"), "پخش صدا", sound_menu)
    playsound.setFont(font)
    playsound.triggered.connect(lambda: on_play_sound(text_edit,sound_button))

    savesound = QAction(QIcon('download.png'), "ذخیره فایل صوتی", sound_menu)
    savesound.setFont(font)
    savesound.triggered.connect(lambda: on_save_sound(text_edit,sound_button)) 

    
    sound_menu.addAction(playsound)
    sound_menu.addAction(savesound)

    sound_button.setMenu(sound_menu)"""

    data_label=QLabel(":آمار")
    #databar.addWidget(data_label)
    data_label.setText("متنی را وارد نکردید!")
    label_style = """
QLabel {
    background-color: #e2b891;
    border: 2px solid #e37712;
    border-radius: 10px;
    font-family: "Cairo";
    
    color: #000000;
    padding: 15px;
    margin: 10px;
    line-height: 1.6;
}
"""
    pixmap = QPixmap('data-monitor-svgrepo-com.png') 
    pixmap = pixmap.scaled(30, 30, Qt.KeepAspectRatio)  
    #data_label.setPixmap(pixmap)
    data_label.setStyleSheet(label_style)
    data_label.setFixedHeight(100)
    data_label.setFont(QFont("Cairo", 13))
    right_layout.addWidget(data_label)

    info_button = QPushButton()
    info_button.setIcon(QIcon('info-circle-svgrepo-com (2).png'))
    info_button.setIconSize(QSize(30, 30))
    font = QFont("Cairo",10)  # انتخاب فونت، اندازه و ضخامت
    info_button.setFont(font)
    info_button.setStyleSheet('''
    QPushButton {
        padding: 7px;
        background: qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #b77a08, stop:1 #f39c12);
        color: #000000;
        border: 5px solid qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #b77a08, stop:1 #f39c12);
        border-radius: 7px;                  
        
        
    }
    QPushButton:hover {
        background: #f39c12
    }
    QPushButton:pressed {
        background: #f39c12
    }
''')
    info_button.clicked.connect(lambda: info_page(info_button))
    toolbar1.addWidget(info_button)
    shortcut_info = QShortcut(QKeySequence("Ctrl+I"), window)
    shortcut_info.activated.connect(lambda: info_menu.exec_(
    info_button.mapToGlobal(info_button.rect().bottomLeft())
    ))

    info_menu=QMenu()
    info_menu.setCursor(Qt.PointingHandCursor)  # تغییر نشانگر به دست اشاره‌گر
    info_menu.setStyleSheet('''
    QMenu {
        background-color: qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #b77a08, stop:1 #f39c12);
        border: 1px solid #ffffff;
        border-radius: 20px;            /* گردی گوشه‌های منو */
    }
    QMenu::item {
        padding: 10px 20px;  /* افزایش فاصله عمودی و افقی */
        font-size: 14px;     /* اندازه فونت */
        height: 20px;        /* ارتفاع هر آیتم */
    }
    QMenu::item:selected {
        background-color: qradialgradient(cx:0.5, cy:0.5, radius:0.5,
            fx:0.5, fy:0.5,
            stop:0 #a2d1a1, stop:1 #91d9fe);  /* رنگ هنگام hover */
    }
''')
    
        # ایجاد اکشن‌های منو با آیکن، فونت و اتصال به توابع ذخیره‌سازی
    info = QAction(QIcon("info-circle-svgrepo-com (2).png"), " درباره‌ی ما", info_menu)
    info.setFont(font)
    info.triggered.connect(lambda: info_page())

    web_addres = QAction(QIcon('virs.png'), " وب سایت", info_menu)
    web_addres.setFont(font)
    web_addres.triggered.connect(lambda: webbrowser.open("https://virasto.github.io/")) 

    bale_addres = QAction(QIcon('bale-column-text-black.png'), " کانال بله", info_menu)
    bale_addres.setFont(font)
    bale_addres.triggered.connect(lambda: webbrowser.open("https://ble.ir/join/6XnKEWFmMG"))

# افزودن اکشن‌ها به منو
    info_menu.addAction(info)
    info_menu.addAction(web_addres)
    info_menu.addAction(bale_addres)

    # اتصال منو به دکمه
    info_button.setMenu(info_menu)

    window.showMaximized()
    window.show()
    #print(time.time() - a)

    app.setWindowIcon(QIcon('virs.png'))
    
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()