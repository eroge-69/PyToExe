import sys
import os
import re
import subprocess
import platform
from datetime import datetime
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QTabWidget, QWidget, QVBoxLayout, QHBoxLayout,
    QFormLayout, QLineEdit, QDateEdit, QComboBox, QPushButton, QTableView,
    QLabel, QHeaderView, QAbstractItemView, QMessageBox, QFileDialog,
    QDialog, QDialogButtonBox, QGridLayout, QInputDialog, QListWidget,
    QListWidgetItem, QMenu, QTextEdit, QSplitter, QCompleter, QSizePolicy,
    QStyledItemDelegate, QStyle
)
from PyQt5.QtCore import Qt, QDate, QSortFilterProxyModel, QRegExp, QStringListModel, QEvent, QSize
# Düzeltme: QImage doğrudan kullanılmak üzere eklendi
from PyQt5.QtGui import QStandardItemModel, QStandardItem, QKeySequence, QPixmap, QImage, QPainter # QPainter added

import openpyxl
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn

import tempfile
import time
import json # JSON module added
import sqlite3 # SQLite module added

# Win32print is only available on Windows.
# Alternative methods or control mechanisms can be added for other platforms.
if sys.platform.startswith("win"):
    try:
        import win32print
        import win32api
        import win32con
    except ImportError:
        QMessageBox.warning(None, "Module Error",
                            "To use Windows printer features, you need to install the 'pywin32' library.\n"
                            "Command: pip install pywin32")
        win32print = None
        win32api = None
        win32con = None

# PDF and Image processing libraries
try:
    import fitz  # PyMuPDF (no longer directly used for PDF preview, but kept for potential other uses if any)
    from PIL import Image
    from PIL.ImageQt import ImageQt # For converting PIL Image to QImage (no longer directly used for PDF preview, but kept for potential other uses if any)
    
    # New import for pdf2image method
    from pdf2image import convert_from_path
except ImportError:
    QMessageBox.warning(None, "Module Error",
                        "To use PDF features, you need to install 'PyMuPDF', 'Pillow', and 'pdf2image' libraries.\n"
                        "Also, ensure 'Poppler' is installed and its bin directory is in your system's PATH.\n"
                        "Commands: pip install PyMuPDF Pillow pdf2image")
    fitz = None
    Image = None
    ImageQt = None
    convert_from_path = None # Set pdf2image function to None as well



class ButtonDelegate(QStyledItemDelegate):
    """
    A delegate that paints an 'Open' icon in the specified column and handles clicks.
    The icon is drawn centered in the cell; clicks open the PDF management dialog.
    """
    def __init__(self, parent=None):
        super().__init__(parent)

    def paint(self, painter, option, index):
        """
        Paints an 'open folder' icon for the PDF column.
        """
        if index.column() == 10:  # PDF column
            try:
                style = QApplication.style()
                icon = style.standardIcon(QStyle.SP_DirOpenIcon)
                # Choose an icon size that fits the row height
                size = max(16, option.rect.height() - 8)
                pix = icon.pixmap(QSize(size, size))
                x = option.rect.x() + (option.rect.width() - pix.width()) // 2
                y = option.rect.y() + (option.rect.height() - pix.height()) // 2
                painter.save()
                painter.drawPixmap(x, y, pix)
                painter.restore()
            except Exception:
                # Fallback: draw default representation
                super().paint(painter, option, index)
        else:
            super().paint(painter, option, index)

    def editorEvent(self, event, model, option, index):
        """
        Handle mouse release events on the icon cell and open the PDF management dialog.
        """
        if event.type() == QEvent.MouseButtonRelease and index.column() == 10:
            try:
                # Retrieve the record id from column 0 of the same row
                record_id = index.model().data(index.sibling(index.row(), 0), Qt.DisplayRole)
                # Try to convert to int when possible
                try:
                    record_id_int = int(record_id)
                except Exception:
                    record_id_int = record_id
                # Call parent method to open PDF management dialog
                self.parent().open_pdf_management_dialog(record_id_int)
            except Exception:
                pass
            return True
        return super().editorEvent(event, model, option, index)

    def createEditor(self, parent, option, index):
        """
        No editor is needed for the icon cell.
        """
        return None
class PrinterSelectionDialog(QDialog):
    """
    Dialog window for printer selection.
    Yazıcı seçimi için diyalog penceresi.
    """
    def __init__(self, parent=None, initial_printer=None):
        super().__init__(parent)
        self.setWindowTitle("Select Printer") # Yazıcı Seç
        self.setFixedSize(400, 300)

        layout = QVBoxLayout()

        self.printer_list_widget = QListWidget()
        self.load_printers()
        layout.addWidget(self.printer_list_widget)

        buttons = QDialogButtonBox(
            QDialogButtonBox.Ok | QDialogButtonBox.Cancel
        )
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

        self.setLayout(layout)
        self.selected_printer = None

        # Select the initial printer if provided
        # Sağlanan başlangıç yazıcısını seç
        if initial_printer:
            items = self.printer_list_widget.findItems(initial_printer, Qt.MatchExactly)
            if items:
                self.printer_list_widget.setCurrentItem(items[0])
                self.selected_printer = initial_printer

    def load_printers(self):
        """
        Lists available printers on the system.
        Sistemdeki mevcut yazıcıları listeler.
        """
        self.printer_list_widget.clear()
        if sys.platform.startswith("win") and win32print:
            try:
                printers = [printer[2] for printer in win32print.EnumPrinters(win32print.PRINTER_ENUM_LOCAL)]
                for printer in printers:
                    self.printer_list_widget.addItem(printer)
            except Exception as e:
                QMessageBox.warning(self, "Printer Loading Error", f"An error occurred while loading printers: {str(e)}") # Yazıcı Yükleme Hatası, Yazıcılar yüklenirken bir hata oluştu
        elif sys.platform.startswith("darwin"): # macOS
            try:
                # Use 'lpstat -p' command to list printers on macOS
                # macOS'ta yazıcıları listelemek için 'lpstat -p' komutunu kullan
                result = subprocess.run(["lpstat", "-p"], capture_output=True, text=True, check=True)
                # Parse output to get printer names
                # Çıktıyı yazıcı adlarını almak için ayrıştır
                printers = []
                for line in result.stdout.splitlines():
                    match = re.search(r"printer (\S+) is idle", line)
                    if match:
                        printers.append(match.group(1))
                for printer in printers:
                    self.printer_list_widget.addItem(printer)
            except Exception as e:
                QMessageBox.warning(self, "Printer Loading Error", f"An error occurred while loading macOS printers: {str(e)}\n" # Yazıcı Yükleme Hatası, macOS yazıcıları yüklenirken bir hata oluştu
                                                                     "Please ensure 'lpstat' command is working on your system.") # Lütfen 'lpstat' komutunun sisteminizde çalıştığından emin olun
        else: # Linux
            try:
                # Use 'lpstat -p' command to list printers on Linux
                # Linux'ta yazıcıları listelemek için 'lpstat -p' komutunu kullan
                result = subprocess.run(["lpstat", "-p"], capture_output=True, text=True, check=True)
                printers = []
                for line in result.stdout.splitlines():
                    match = re.search(r"printer (\S+) is idle", line)
                    if match:
                        printers.append(match.group(1))
                for printer in printers:
                    self.printer_list_widget.addItem(printer)
            except Exception as e:
                QMessageBox.warning(self, "Printer Loading Error", f"An error occurred while loading Linux printers: {str(e)}\n" # Yazıcı Yükleme Hatası, Linux yazıcıları yüklenirken bir hata oluştu
                                                                     "Please ensure 'lpstat' command is working on your system.") # Lütfen 'lpstat' komutunun sisteminizde çalıştığından emin olun
        
        if self.printer_list_widget.count() == 0:
            self.printer_list_widget.addItem("No printer found.") # Yazıcı bulunamadı
            self.printer_list_widget.setEnabled(False)


    def accept(self):
        selected_items = self.printer_list_widget.selectedItems()
        if selected_items:
            self.selected_printer = selected_items[0].text()
        super().accept()

    def get_selected_printer(self):
        return self.selected_printer


class UserDialog(QDialog):
    """
    Authorized person add/edit dialog window.
    Yetkili kişi ekleme/düzenleme diyalog penceresi.
    """
    def __init__(self, user=None, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Authorized Person Management") # Yetkili Kişi Yönetimi
        self.setFixedSize(400, 200)

        layout = QGridLayout()

        # Form elements
        # Form elemanları
        layout.addWidget(QLabel("Organization Name:"), 0, 0) # Kuruluş Adı:
        self.organization = QLineEdit()
        layout.addWidget(self.organization, 0, 1)

        layout.addWidget(QLabel("Title:"), 1, 0) # Ünvan:
        self.title = QLineEdit()
        layout.addWidget(self.title, 1, 1)

        layout.addWidget(QLabel("Full Name:"), 2, 0) # Tam Adı:
        self.fullname = QLineEdit()
        layout.addWidget(self.fullname, 2, 1)

        # Buttons
        # Butonlar
        buttons = QDialogButtonBox(
            QDialogButtonBox.Ok | QDialogButtonBox.Cancel
        )
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons, 3, 0, 1, 2)

        # Load data if opened for editing
        # Düzenleme için açıldıysa verileri yükle
        if user:
            self.organization.setText(user.get('organization', ''))
            self.title.setText(user.get('title', ''))
            self.fullname.setText(user.get('fullname', ''))

        self.setLayout(layout)

    def get_data(self):
        """
        Returns data from the dialog window as a dictionary.
        Diyalog penceresindeki verileri sözlük olarak döndürür.
        """
        return {
            'organization': self.organization.text(),
            'title': self.title.text(),
            'fullname': self.fullname.text()
        }

class PreviewDialog(QDialog):
    """
    Dialog window for previewing and optionally editing record details.
    Kayıt detaylarını önizlemek ve isteğe bağlı olarak düzenlemek için diyalog penceresi.
    """
    def __init__(self, record, db_manager, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Record Preview and Edit") # Kayıt Önizleme ve Düzenle
        self.setGeometry(150, 150, 700, 700) # Increased size for better preview and edit form
                                           # Daha iyi önizleme ve düzenleme formu için boyut artırıldı
        
        self.original_record = record # Keep original for reference
                                     # Referans için orijinali sakla
        self.edited_record = dict(record) # Work on a copy of the record
                                         # Kaydın bir kopyası üzerinde çalış
        self.db = db_manager
        
        self.save_and_print_requested = False
        self.save_changes_requested = False

        main_layout = QVBoxLayout()
        splitter = QSplitter(Qt.Vertical)

        # --- Top Section: Read-only Preview ---
        # Üst Bölüm: Salt Okunur Önizleme
        preview_widget = QWidget()
        preview_layout = QVBoxLayout(preview_widget)
        preview_layout.addWidget(QLabel("Document Preview:")) # Belge Önizlemesi:
        
        self.preview_text_edit = QTextEdit()
        self.preview_text_edit.setReadOnly(True)
        self.preview_text_edit.setFontPointSize(10)
        preview_layout.addWidget(self.preview_text_edit)
        
        splitter.addWidget(preview_widget)

        # --- Bottom Section: Editable Form ---
        # Alt Bölüm: Düzenlenebilir Form
        edit_form_widget = QWidget()
        edit_form_layout = QVBoxLayout(edit_form_widget)
        edit_form_layout.addWidget(QLabel("Edit Record Details:")) # Kayıt Detaylarını Düzenle:

        form_fields_layout = QFormLayout()

        self.surname_edit = QLineEdit()
        self.surname_edit.setText(self.edited_record.get('surname', ''))
        form_fields_layout.addRow("Surname:", self.surname_edit) # Soyadı:

        self.given_names_edit = QLineEdit()
        self.given_names_edit.setText(self.edited_record.get('given_names', ''))
        form_fields_layout.addRow("Given Names:", self.given_names_edit) # Adı:

        self.sex_combo = QComboBox()
        self.sex_combo.addItems(["ERKEK/MALE", "AÝAL/FEMALE"])
        self.sex_combo.setCurrentText(self.edited_record.get('sex', 'ERKEK/MALE'))
        form_fields_layout.addRow("Sex:", self.sex_combo) # Cinsiyet:

        self.place_of_birth_edit = QLineEdit()
        self.place_of_birth_edit.setText(self.edited_record.get('place_of_birth', 'TKM'))
        form_fields_layout.addRow("Place of Birth:", self.place_of_birth_edit) # Doğum Yeri:

        self.date_of_birth_edit = QDateEdit()
        self.date_of_birth_edit.setDisplayFormat("dd.MM.yyyy")
        self.date_of_birth_edit.setCalendarPopup(True)
        # Parse date string to QDate object
        try:
            self.date_of_birth_edit.setDate(QDate.fromString(self.edited_record.get('date_of_birth', '01.01.2000'), "dd.MM.yyyy"))
        except ValueError:
            self.date_of_birth_edit.setDate(QDate(2000, 1, 1))
        form_fields_layout.addRow("Date of Birth:", self.date_of_birth_edit) # Doğum Tarihi:

        self.date_of_issue_edit = QDateEdit()
        self.date_of_issue_edit.setDisplayFormat("dd.MM.yyyy")
        self.date_of_issue_edit.setCalendarPopup(True)
        try:
            self.date_of_issue_edit.setDate(QDate.fromString(self.edited_record.get('date_of_issue', QDate.currentDate().toString("dd.MM.yyyy")), "dd.MM.yyyy"))
        except ValueError:
            self.date_of_issue_edit.setDate(QDate.currentDate())
        form_fields_layout.addRow("Date of Issue:", self.date_of_issue_edit) # Veriliş Tarihi:

        self.date_of_expiry_edit = QDateEdit()
        self.date_of_expiry_edit.setDisplayFormat("dd.MM.yyyy")
        self.date_of_expiry_edit.setCalendarPopup(True)
        try:
            self.date_of_expiry_edit.setDate(QDate.fromString(self.edited_record.get('date_of_expiry', ''), "dd.MM.yyyy"))
        except ValueError:
            self.date_of_expiry_edit.setDate(QDate.currentDate().addMonths(1).addDays(-1)) # Default if invalid
                                                                                               # Geçersizse varsayılan

        form_fields_layout.addRow("Date of Expiry:", self.date_of_expiry_edit) # Son Geçerlilik Tarihi:

        self.telefon_edit = QLineEdit()
        # Remove "+90 5" prefix for editing in mask
        # Maskede düzenleme için "+90 5" önekini kaldır
        phone_display = self.edited_record.get('telefon', '')
        if phone_display.startswith("+90 5"):
            self.telefon_edit.setText(phone_display[5:].strip())
        else:
            self.telefon_edit.setText(phone_display)
        self.telefon_edit.setInputMask("000 000 00 00;_")
        form_fields_layout.addRow("Phone (+90 5):", self.telefon_edit) # Telefon (+90 5):

        self.lp_number_edit = QLineEdit()
        # Remove "A№" prefix for editing
        # Düzenleme için "A№" önekini kaldır
        lp_display = self.edited_record.get('lp_number', '')
        if lp_display.startswith("A№"):
            self.lp_number_edit.setText(lp_display[2:].strip())
        else:
            self.lp_number_edit.setText(lp_display)
        form_fields_layout.addRow("LP Number (A№):", self.lp_number_edit) # LP Numarası (A№):

        # Authorized Person selection for editing
        # Düzenleme için Yetkili Kişi seçimi
        self.auth_person_combo = QComboBox()
        self._populate_auth_person_combo()
        form_fields_layout.addRow("Authorized Person:", self.auth_person_combo) # Yetkili Kişi:

        edit_form_layout.addLayout(form_fields_layout)

        # Buttons for editing section
        # Düzenleme bölümü için butonlar
        edit_buttons_layout = QHBoxLayout()
        self.update_preview_btn = QPushButton("Update Preview") # Önizlemeyi Güncelle
        self.update_preview_btn.clicked.connect(self._update_dialog_preview)
        edit_buttons_layout.addWidget(self.update_preview_btn)

        self.save_changes_btn = QPushButton("Save Changes") # Değişiklikleri Kaydet
        self.save_changes_btn.clicked.connect(self._save_changes_only)
        edit_buttons_layout.addWidget(self.save_changes_btn)

        self.save_print_btn = QPushButton("Save & Print") # Kaydet ve Yazdır
        self.save_print_btn.clicked.connect(self._save_changes_and_print)
        edit_buttons_layout.addWidget(self.save_print_btn)
        
        self.cancel_btn = QPushButton("Cancel") # İptal
        self.cancel_btn.clicked.connect(self.reject)
        edit_buttons_layout.addWidget(self.cancel_btn)

        edit_form_layout.addLayout(edit_buttons_layout)
        
        splitter.addWidget(edit_form_widget)
        splitter.setSizes([400, 300]) # Adjusted sizes for preview and edit sections
                                    # Önizleme ve düzenleme bölümleri için ayarlanmış boyutlar

        main_layout.addWidget(splitter)
        self.setLayout(main_layout)

        # Initial preview update
        # Başlangıç önizleme güncellemesi
        self._update_dialog_preview()


    def _populate_auth_person_combo(self):
        """Populates the authorized person combo box in the dialog."""
        """Diyalogdaki yetkili kişi açılır kutusunu doldurur."""
        self.auth_person_combo.clear()
        users = self.db.get_users()
        
        current_auth_person_id_in_record = self.edited_record.get('auth_person', {}).get('id')
        
        selected_index = -1
        for i, user in enumerate(users):
            display_text = f"{user['organization']} - {user['title']} - {user['fullname']}"
            self.auth_person_combo.addItem(display_text, user) # Store the whole user dictionary as data
            if user.get('id') == current_auth_person_id_in_record:
                selected_index = i
        
        if selected_index != -1:
            self.auth_person_combo.setCurrentIndex(selected_index)
        elif users: # If no match, but there are users, select the first one
                    # Eşleşme yoksa ancak kullanıcılar varsa, ilkini seç
            self.auth_person_combo.setCurrentIndex(0)
            self.edited_record['auth_person'] = users[0] # Update edited record with default selected
        else: # No users at all
              # Hiç kullanıcı yok
            self.edited_record['auth_person'] = {}


    def get_modified_record_data(self):
        """
        Collects data from editable fields and returns a dictionary.
        Düzenlenebilir alanlardan verileri toplar ve bir sözlük döndürür.
        """
        # Collect current data from the editable fields
        # Düzenlenebilir alanlardan mevcut verileri topla
        modified_data = {
            'id': self.edited_record.get('id'), # Keep original ID
            'surname': self.surname_edit.text().strip(),
            'given_names': self.given_names_edit.text().strip(),
            'sex': self.sex_combo.currentText(),
            'place_of_birth': self.place_of_birth_edit.text().strip(),
            'date_of_birth': self.date_of_birth_edit.date().toString("dd.MM.yyyy"),
            'date_of_issue': self.date_of_issue_edit.date().toString("dd.MM.yyyy"),
            'date_of_expiry': self.date_of_expiry_edit.date().toString("dd.MM.yyyy"),
            'telefon': "+90 5" + self.telefon_edit.text().strip(),
            'lp_number': "A№" + self.lp_number_edit.text().strip(),
            'auth_person': self.auth_person_combo.currentData() # Get the full user dict
        }
        return modified_data


    def _update_dialog_preview(self):
        """
        Updates the read-only preview text based on the current editable field values.
        Salt okunur önizleme metnini mevcut düzenlenebilir alan değerlerine göre günceller.
        """
        # Temporarily update self.edited_record with current form values for preview generation
        # Önizleme oluşturma için self.edited_record'u mevcut form değerleriyle geçici olarak güncelle
        self.edited_record.update(self.get_modified_record_data())

        # Now generate the preview text using the (temporarily) updated edited_record
        # Şimdi (geçici olarak) güncellenmiş edited_record kullanarak önizleme metnini oluştur
        
        # This function mimics the generate_docx structure but creates plain text
        # Bu fonksiyon generate_docx yapısını taklit eder ancak düz metin oluşturur
        surname = self.edited_record['surname'].upper()
        given_names = self.edited_record['given_names'].upper()
        date_of_birth = self.edited_record['date_of_birth']
        date_of_expiry = self.edited_record['date_of_expiry']
        place_of_birth = self.edited_record['place_of_birth']
        sex = self.edited_record['sex'].upper()
        date_of_issue = self.edited_record['date_of_issue']
        
        auth = self.edited_record.get('auth_person', {}) or {'organization':'', 'title':'', 'fullname':''}
        organization = auth.get('organization','')
        title = auth.get('title','')
        fullname = auth.get('fullname','')

        preview_lines = ["", ""]
        indent = " " * 18 
        preview_lines.append(f"{indent}{surname}")
        preview_lines.append(f"{indent}{given_names}")

        base_line_start_preview_char = 18
        conceptual_tab_stop_preview_char = 45 
        shifted_tab_stop_preview_char = conceptual_tab_stop_preview_char - 4 

        left_text_nationality = "TURKMENISTAN"
        spacing_nationality = " " * (shifted_tab_stop_preview_char - len(left_text_nationality) - base_line_start_preview_char)
        preview_lines.append(f"{indent}{left_text_nationality}{spacing_nationality}{date_of_birth}")

        left_text_place_sex = place_of_birth
        spacing_place_sex = " " * (shifted_tab_stop_preview_char - len(left_text_place_sex) - base_line_start_preview_char)
        preview_lines.append(f"{indent}{left_text_place_sex}{spacing_place_sex}{sex}")

        left_text_dates = date_of_issue
        spacing_dates = " " * (shifted_tab_stop_preview_char - len(left_text_dates) - base_line_start_preview_char)
        preview_lines.append(f"{indent}{left_text_dates}{spacing_dates}{date_of_expiry}")

        preview_lines.append("") 

        preview_lines.append(f"{indent}{organization}")
        preview_lines.append("") 

        left_text_auth = title
        right_text_auth = fullname
        spacing_auth = " " * (conceptual_tab_stop_preview_char - len(left_text_auth) - base_line_start_preview_char) 
        preview_lines.append(f"{indent}{left_text_auth}{spacing_auth}{right_text_auth}")

        self.preview_text_edit.setPlainText("\n".join(preview_lines))


    def _save_changes_only(self):
        """Saves the edited changes back to the database without printing."""
        """Düzenlenen değişiklikleri yazdırmadan veritabanına kaydeder."""
        updated_data = self.get_modified_record_data()
        self.db.update_record(updated_data) # Call the new update_record method
        self.save_changes_requested = True
        QMessageBox.information(self, "Success", "Changes saved successfully!") # Başarı, Değişiklikler başarıyla kaydedildi!
        self.accept() # Close dialog

    def _save_changes_and_print(self):
        """Saves the edited changes and triggers a print."""
        """Düzenlenen değişiklikleri kaydeder ve yazdırma işlemini başlatır."""
        updated_data = self.get_modified_record_data()
        self.db.update_record(updated_data) # Call the new update_record method
        self.save_changes_requested = True
        self.save_and_print_requested = True
        QMessageBox.information(self, "Success", "Changes saved successfully! Document will now print.") # Başarı, Değişiklikler başarıyla kaydedildi! Belge şimdi yazdırılacak.
        self.accept() # Close dialog

class PdfManagementDialog(QDialog):
    """
    Dialog window for managing PDFs associated with a record.
    Bir kayıtla ilişkili PDF'leri yönetmek için diyalog penceresi.
    """
    def __init__(self, record_id, db_manager, parent=None):
        super().__init__(parent)
        self.record_id = record_id
        self.db = db_manager
        self.setWindowTitle(f"PDF Management for Record ID: {record_id}") # Kayıt Kimliği için PDF Yönetimi:
        self.setGeometry(200, 200, 900, 700)

        main_layout = QVBoxLayout()
        
        # Splitter for PDF list and preview area
        # PDF listesi ve önizleme alanı için ayırıcı
        splitter = QSplitter(Qt.Horizontal)
        
        # Left side: PDF list and controls
        # Sol taraf: PDF listesi ve kontroller
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        
        left_layout.addWidget(QLabel("Associated PDFs:")) # İlişkili PDF'ler:
        self.pdf_list_widget = QListWidget()
        self.pdf_list_widget.currentItemChanged.connect(self.preview_selected_pdf)
        left_layout.addWidget(self.pdf_list_widget)

        button_layout = QHBoxLayout()
        self.add_pdf_btn = QPushButton("PDF Ekle") # Add PDF
        self.add_pdf_btn.clicked.connect(self.add_pdf)
        button_layout.addWidget(self.add_pdf_btn)

        self.delete_pdf_btn = QPushButton("PDF Sil") # Delete PDF
        self.delete_pdf_btn.clicked.connect(self.delete_pdf)
        button_layout.addWidget(self.delete_pdf_btn)

        self.download_pdf_btn = QPushButton("PDF İndir") # Download PDF
        self.download_pdf_btn.clicked.connect(self.download_pdf)
        button_layout.addWidget(self.download_pdf_btn)
        
        left_layout.addLayout(button_layout)
        splitter.addWidget(left_widget)

        # Right side: PDF preview
        # Sağ taraf: PDF önizlemesi
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.addWidget(QLabel("PDF Preview:")) # PDF Önizlemesi:
        
        self.pdf_preview_label = QLabel("No PDF selected for preview.") # Önizleme için PDF seçilmedi.
        self.pdf_preview_label.setAlignment(Qt.AlignCenter)
        self.pdf_preview_label.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.pdf_preview_label.setScaledContents(True) # Scale pixmap to fit label size
                                                      # Pixmap'i etiket boyutuna sığacak şekilde ölçeklendir
        right_layout.addWidget(self.pdf_preview_label)
        
        splitter.addWidget(right_widget)
        
        splitter.setSizes([300, 600]) # Initial sizes for list and preview
                                    # Liste ve önizleme için başlangıç boyutları

        main_layout.addWidget(splitter)
        self.setLayout(main_layout)

        self.load_pdfs()

    def load_pdfs(self):
        """Loads PDF files associated with the record from the database and populates the list widget."""
        """Kayda bağlı PDF dosyalarını veritabanından yükler ve liste widget'ını doldurur."""
        self.pdf_list_widget.clear()
        pdfs = self.db.get_pdf_files(self.record_id)
        for pdf in pdfs:
            item = QListWidgetItem(pdf['filename'])
            item.setData(Qt.UserRole, pdf) # Store full PDF data (id, filename, filepath) in item
            self.pdf_list_widget.addItem(item)
        if not pdfs:
            self.pdf_preview_label.setText("No PDF selected for preview.")

    def add_pdf(self):
        """Allows user to select a PDF file, copies it to app data, and saves its info to the database."""
        """Kullanıcının bir PDF dosyası seçmesine, uygulama verilerine kopyalamasına ve bilgilerini veritabanına kaydetmesine olanak tanır."""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "PDF Seç", "", "PDF Files (*.pdf)" # Select PDF, PDF Dosyaları
        )

        if not file_path:
            return

        # Ensure PDF processing libraries are available
        if not convert_from_path: # Check for pdf2image specific function
            QMessageBox.warning(self, "Missing Modules", "PDF functionality requires 'pdf2image' library and 'Poppler' installed. Please install them.")
            return

        try:
            # Create app data directory for PDFs if it doesn't exist
            pdf_data_dir = self.db._get_pdf_data_path()
            os.makedirs(pdf_data_dir, exist_ok=True)

            original_filename = os.path.basename(file_path)
            # Create a unique filename to avoid conflicts
            unique_filename = f"{self.record_id}_{int(time.time())}_{original_filename}"
            destination_path = os.path.join(pdf_data_dir, unique_filename)

            # Copy the file
            with open(file_path, 'rb') as src, open(destination_path, 'wb') as dest:
                dest.write(src.read())

            # Save to database
            self.db.add_pdf_file(self.record_id, original_filename, destination_path)
            QMessageBox.information(self, "Başarılı", f"PDF '{original_filename}' başarıyla eklendi.") # Success, PDF successfully added.
            self.load_pdfs()
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"PDF eklenirken bir hata oluştu: {str(e)}") # Error, An error occurred while adding PDF:

    def delete_pdf(self):
        """Deletes the selected PDF from the database and file system."""
        """Seçilen PDF'yi veritabanından ve dosya sisteminden siler."""
        selected_item = self.pdf_list_widget.currentItem()
        if not selected_item:
            QMessageBox.warning(self, "Uyarı", "Lütfen silmek için bir PDF seçin.") # Warning, Please select a PDF to delete.
            return

        pdf_data = selected_item.data(Qt.UserRole)
        pdf_id = pdf_data['id']
        pdf_filepath = pdf_data['filepath']
        pdf_filename = pdf_data['filename']

        reply = QMessageBox.question(
            self, "Silme Onayı", # Confirm Deletion
            f"PDF '{pdf_filename}' silinecektir. Devam etmek istiyor musunuz?", # PDF will be deleted. Do you want to continue?
            QMessageBox.Yes | QMessageBox.No
        )

        if reply == QMessageBox.Yes:
            try:
                if self.db.delete_pdf_file(pdf_id):
                    if os.path.exists(pdf_filepath):
                        os.remove(pdf_filepath)
                    QMessageBox.information(self, "Başarılı", f"PDF '{pdf_filename}' başarıyla silindi.") # Success, PDF successfully deleted.
                    self.load_pdfs()
                    self.pdf_preview_label.clear() # Clear preview after deletion
                    self.pdf_preview_label.setText("No PDF selected for preview.")
                else:
                    QMessageBox.warning(self, "Hata", "PDF silinemedi.") # Error, PDF could not be deleted.
            except Exception as e:
                QMessageBox.critical(self, "Hata", f"PDF silinirken bir hata oluştu: {str(e)}") # Error, An error occurred while deleting PDF:

    def download_pdf(self):
        """Allows user to download the selected PDF file."""
        """Kullanıcının seçilen PDF dosyasını indirmesine olanak tanır."""
        selected_item = self.pdf_list_widget.currentItem()
        if not selected_item:
            QMessageBox.warning(self, "Uyarı", "Lütfen indirmek için bir PDF seçin.") # Warning, Please select a PDF to download.
            return

        pdf_data = selected_item.data(Qt.UserRole)
        pdf_filepath = pdf_data['filepath']
        pdf_filename = pdf_data['filename']

        if not os.path.exists(pdf_filepath):
            QMessageBox.warning(self, "Hata", "Dosya bulunamadı. Lütfen yöneticinizle iletişime geçin.") # Error, File not found. Please contact your administrator.
            return

        save_path, _ = QFileDialog.getSaveFileName(
            self, "PDF İndir", pdf_filename, "PDF Files (*.pdf)" # Download PDF, PDF Dosyaları
        )

        if not save_path:
            return

        try:
            with open(pdf_filepath, 'rb') as src, open(save_path, 'wb') as dest:
                dest.write(src.read())
            QMessageBox.information(self, "Başarılı", f"PDF '{pdf_filename}' başarıyla indirildi.") # Success, PDF successfully downloaded.
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"PDF indirilirken bir hata oluştu: {str(e)}") # Error, An error occurred while downloading PDF:

    def preview_selected_pdf(self, current_item, previous_item):
        """
        Previews the selected PDF file by converting its first page to an image using pdf2image.
        Seçilen PDF dosyasını ilk sayfasını pdf2image kullanarak resme dönüştürerek önizler.
        """
        if not current_item:
            self.pdf_preview_label.clear()
            self.pdf_preview_label.setText("No PDF selected for preview.")
            return

        pdf_data = current_item.data(Qt.UserRole)
        pdf_filepath = pdf_data['filepath']

        if not os.path.exists(pdf_filepath):
            self.pdf_preview_label.setText("Dosya bulunamadı.") # File not found.
            QMessageBox.warning(self, "Hata", "Önizlenecek dosya bulunamadı. Silinmiş olabilir.") # Error, File to preview not found. It might have been deleted.
            return
        
        # Ensure pdf2image library is available
        if not convert_from_path:
            self.pdf_preview_label.setText("Gerekli modüller eksik (pdf2image/Poppler).") # Required modules are missing.
            QMessageBox.warning(self, "Missing Modules", "PDF önizleme işlevi için 'pdf2image' kütüphanesi ve 'Poppler' kurulu olmalıdır. Lütfen bunları yükleyin.")
            return

        try:
            # PDF'yi görüntüye dönüştür
            # Using dpi=200 for better quality image conversion
            images = convert_from_path(pdf_filepath, first_page=1, last_page=1, dpi=200)
            if not images:
                raise ValueError("PDF'den görüntü oluşturulamadı.")

            # Geçici bir dosyaya kaydet
            temp_img_path = tempfile.mktemp(suffix='.jpg')
            images[0].save(temp_img_path, 'JPEG')

            # QPixmap ile yükle
            pixmap = QPixmap(temp_img_path)
            if pixmap.isNull():
                raise ValueError("Görüntü yüklenemedi.")

            scaled_pixmap = pixmap.scaled(
                self.pdf_preview_label.size(),
                Qt.KeepAspectRatio,
                Qt.SmoothTransformation
            )
            self.pdf_preview_label.setPixmap(scaled_pixmap)
            self.pdf_preview_label.setText("") # Clear "No PDF selected" text

            # Geçici dosyayı sil
            os.remove(temp_img_path)

        except Exception as e:
            self.pdf_preview_label.setText(f"Önizleme hatası: {str(e)}")
            QMessageBox.critical(self, "Hata", f"PDF önizleme hatası: {str(e)}")


class DatabaseManager:
    """
    Class to manage application data and authorized person information.
    Uses SQLite database.
    Uygulama verilerini ve yetkili kişi bilgilerini yönetmek için sınıf.
    SQLite veritabanı kullanır.
    """
    def __init__(self):
        # Local application data path for user data
        # Kullanıcı verileri için yerel uygulama veri yolu
        app_data_path = os.path.join(os.environ['APPDATA'] if sys.platform == 'win32' else os.path.expanduser('~/.local/share'), 'LP_App')
        os.makedirs(app_data_path, exist_ok=True)
        self.db_path = os.path.join(app_data_path, 'lp_data.db')
        self.conn = None
        self.connect_db()
        self.create_tables()
        self._recalculate_next_id() # Call to set next_id initially
                                    # Başlangıçta next_id'yi ayarlamak için çağrı
        
    def connect_db(self):
        """Connects to the database."""
        """Veritabanına bağlanır."""
        try:
            self.conn = sqlite3.connect(self.db_path, check_same_thread=False)
            self.conn.row_factory = sqlite3.Row # For accessing by column names
                                                # Sütun adlarına göre erişim için
        except sqlite3.Error as e:
            QMessageBox.critical(None, "Database Error", f"Could not connect to database: {e}") # Veritabanı Hatası, Veritabanına bağlanılamadı
            sys.exit(1)

    def create_tables(self):
        """Creates necessary tables."""
        """Gerekli tabloları oluşturur."""
        cursor = self.conn.cursor()
        # Records table
        # Kayıtlar tablosu
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS records (
                id INTEGER PRIMARY KEY, -- AUTOINCREMENT removed
                surname TEXT,
                given_names TEXT,
                sex TEXT,
                place_of_birth TEXT,
                date_of_birth TEXT,
                date_of_issue TEXT,
                date_of_expiry TEXT,
                telefon TEXT,
                lp_number TEXT,
                auth_person_json TEXT -- Store authorized person info as JSON
                                    -- Yetkili kişi bilgilerini JSON olarak sakla
            )
        """)
        # Users (Authorized Persons) table
        # Kullanıcılar (Yetkili Kişiler) tablosu
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                organization TEXT,
                title TEXT,
                fullname TEXT
            )
        """)
        # Settings table for application preferences (e.g., default printer, current authorized person)
        # Uygulama tercihleri için ayarlar tablosu (örn. varsayılan yazıcı, mevcut yetkili kişi)
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS settings (
                setting_name TEXT PRIMARY KEY,
                setting_value TEXT
            )
        """)
        # New: PDF Files table
        # Yeni: PDF Dosyaları tablosu
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS pdf_files (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                record_id INTEGER,
                filename TEXT,
                filepath TEXT,
                FOREIGN KEY(record_id) REFERENCES records(id) ON DELETE CASCADE
            )
        """)
        self.conn.commit()
        self._add_default_user_if_not_exists()

    def _add_default_user_if_not_exists(self):
        """Adds a default authorized person if not exists."""
        """Varsayılan bir yetkili kişi yoksa ekler."""
        cursor = self.conn.cursor()
        cursor.execute("SELECT COUNT(*) FROM users WHERE organization = ? AND title = ? AND fullname = ?",
                       ('Consulate General of Turkmenistan in Istanbul', 'Consul General', 'M. Ovezov'))
        if cursor.fetchone()[0] == 0:
            cursor.execute("INSERT INTO users (organization, title, fullname) VALUES (?, ?, ?)",
                           ('Consulate General of Turkmenistan in Istanbul', 'Consul General', 'M. Ovezov'))
            self.conn.commit()

    def _recalculate_next_id(self):
        """
        Recalculates the next suitable ID based on existing records.
        Ensures new IDs are always unique and increment from the highest existing ID.
        Mevcut kayıtlara göre bir sonraki uygun kimliği yeniden hesaplar.
        Yeni kimliklerin her zaman benzersiz olmasını ve mevcut en yüksek kimlikten artmasını sağlar.
        """
        cursor = self.conn.cursor()
        cursor.execute("SELECT id FROM records ORDER BY id ASC")
        existing_ids = [row[0] for row in cursor.fetchall()]
        
        # Find the smallest positive integer not in existing_ids
        # Mevcut kimliklerde olmayan en küçük pozitif tam sayıyı bul
        next_id = 1
        for id_val in existing_ids:
            if id_val == next_id:
                next_id += 1
            elif id_val > next_id:
                break # Found a gap, next_id is the smallest available
                      # Bir boşluk bulundu, next_id en küçük kullanılabilir olanıdır
        self.next_id = next_id


    def add_record(self, record):
        """
        Adds a new record and assigns a unique ID.
        Yeni bir kayıt ekler ve benzersiz bir kimlik atar.
        """
        cursor = self.conn.cursor()
        auth_person_json = json.dumps(record.get('auth_person', {}))
        
        # Use the pre-calculated next_id
        # Önceden hesaplanmış next_id'yi kullan
        record_id = self.next_id
        
        cursor.execute("""
            INSERT INTO records (id, surname, given_names, sex, place_of_birth, date_of_birth,
                                 date_of_issue, date_of_expiry, telefon, lp_number, auth_person_json)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (record_id, record['surname'], record['given_names'], record['sex'],
              record['place_of_birth'], record['date_of_birth'],
              record['date_of_issue'], record['date_of_expiry'],
              record['telefon'], record['lp_number'], auth_person_json))
        self.conn.commit()
        self._recalculate_next_id() # Recalculate after a new record is added
                                    # Yeni bir kayıt eklendikten sonra yeniden hesapla
        return record_id

    def get_records(self):
        """
        Returns all records.
        Tüm kayıtları döndürür.
        """
        cursor = self.conn.cursor()
        cursor.execute("SELECT * FROM records ORDER BY id DESC") # To get the newest at the top
                                                                # En yeniyi en üste almak için
        records = []
        for row in cursor.fetchall():
            record = dict(row)
            if 'auth_person_json' in record and record['auth_person_json']:
                try:
                    record['auth_person'] = json.loads(record['auth_person_json'])
                except json.JSONDecodeError:
                    record['auth_person'] = {} # Empty dictionary in case of invalid JSON
                                                # Geçersiz JSON durumunda boş sözlük
            else:
                record['auth_person'] = {}
            records.append(record)
        return records

    def update_record(self, record_data):
        """
        Updates an existing record in the database.
        Veritabanındaki mevcut bir kaydı günceller.
        """
        cursor = self.conn.cursor()
        auth_person_json = json.dumps(record_data.get('auth_person', {}))
        
        cursor.execute("""
            UPDATE records SET 
                surname = ?, given_names = ?, sex = ?, place_of_birth = ?, 
                date_of_birth = ?, date_of_issue = ?, date_of_expiry = ?, 
                telefon = ?, lp_number = ?, auth_person_json = ?
            WHERE id = ?
        """, (record_data['surname'], record_data['given_names'], record_data['sex'],
              record_data['place_of_birth'], record_data['date_of_birth'],
              record_data['date_of_issue'], record_data['date_of_expiry'],
              record_data['telefon'], record_data['lp_number'], 
              auth_person_json, record_data['id']))
        self.conn.commit()
        return cursor.rowcount > 0


    def delete_record(self, record_id):
        """
        Deletes a record by its ID.
        Kimliğine göre bir kaydı siler.
        """
        cursor = self.conn.cursor()
        cursor.execute("DELETE FROM records WHERE id = ?", (record_id,))
        self.conn.commit()
        # After deletion, remove any potential gaps in IDs by re-assigning them
        # Silme işleminden sonra, kimliklerdeki olası boşlukları yeniden atayarak kaldır
        self._reorganize_ids()
        self._recalculate_next_id() # Recalculate after a record is deleted
                                    # Bir kayıt silindikten sonra yeniden hesapla
        return cursor.rowcount > 0

    def _reorganize_ids(self):
        """
        Reorganizes IDs to ensure they are sequential without gaps.
        Kimlikleri boşluksuz ardışık olmalarını sağlamak için yeniden düzenler.
        """
        cursor = self.conn.cursor()
        cursor.execute("SELECT id FROM records ORDER BY id ASC")
        existing_ids = [row[0] for row in cursor.fetchall()]

        new_id = 1
        for old_id in existing_ids:
            if old_id != new_id:
                cursor.execute("UPDATE records SET id = ? WHERE id = ?", (new_id, old_id))
            new_id += 1
        self.conn.commit()
        # Also, make sure to compact the database to reclaim space from deleted rows
        # Ayrıca, silinen satırlardan alanı geri kazanmak için veritabanını sıkıştırdığınızdan emin olun
        cursor.execute("VACUUM")
        self.conn.commit()


    def add_user(self, user):
        """
        Adds a new authorized person.
        Yeni bir yetkili kişi ekler.
        """
        cursor = self.conn.cursor()
        cursor.execute("INSERT INTO users (organization, title, fullname) VALUES (?, ?, ?)",
                       (user['organization'], user['title'], user['fullname']))
        self.conn.commit()
        return cursor.lastrowid

    def get_users(self):
        """
        Returns all authorized persons.
        Tüm yetkili kişileri döndürür.
        """
        cursor = self.conn.cursor()
        cursor.execute("SELECT * FROM users")
        return [dict(row) for row in cursor.fetchall()]

    def delete_user(self, user_id):
        """
        Deletes the authorized person with the specified ID.
        Belirtilen kimliğe sahip yetkili kişiyi siler.
        """
        cursor = self.conn.cursor()
        cursor.execute("DELETE FROM users WHERE id = ?", (user_id,))
        self.conn.commit()
        return cursor.rowcount > 0

    def update_user(self, user_id, user_data):
        """
        Updates the authorized person with the specified ID.
        Belirtilen kimliğe sahip yetkili kişiyi günceller.
        """
        cursor = self.conn.cursor()
        cursor.execute("UPDATE users SET organization = ?, title = ?, fullname = ? WHERE id = ?",
                       (user_data['organization'], user_data['title'], user_data['fullname'], user_id))
        self.conn.commit()
        return cursor.rowcount > 0

    def save_setting(self, name, value):
        """
        Saves an application setting.
        Bir uygulama ayarını kaydeder.
        """
        cursor = self.conn.cursor()
        cursor.execute("INSERT OR REPLACE INTO settings (setting_name, setting_value) VALUES (?, ?)",
                       (name, value))
        self.conn.commit()

    def get_setting(self, name):
        """
        Retrieves an application setting.
        Bir uygulama ayarını alır.
        """
        cursor = self.conn.cursor()
        cursor.execute("SELECT setting_value FROM settings WHERE setting_name = ?", (name,))
        result = cursor.fetchone()
        return result[0] if result else None

    # New PDF management methods
    # Yeni PDF yönetim yöntemleri
    def _get_pdf_data_path(self):
        """Returns the path to the directory where PDF files are stored."""
        """PDF dosyalarının saklandığı dizinin yolunu döndürür."""
        app_data_path = os.path.join(os.environ['APPDATA'] if sys.platform == 'win32' else os.path.expanduser('~/.local/share'), 'LP_App')
        pdf_data_dir = os.path.join(app_data_path, 'pdfs')
        os.makedirs(pdf_data_dir, exist_ok=True)
        return pdf_data_dir

    def add_pdf_file(self, record_id, filename, filepath):
        """Adds a new PDF file entry to the database."""
        """Veritabanına yeni bir PDF dosyası girişi ekler."""
        cursor = self.conn.cursor()
        cursor.execute("INSERT INTO pdf_files (record_id, filename, filepath) VALUES (?, ?, ?)",
                       (record_id, filename, filepath))
        self.conn.commit()
        return cursor.lastrowid

    def get_pdf_files(self, record_id):
        """Returns all PDF files associated with a given record_id."""
        """Belirli bir record_id ile ilişkili tüm PDF dosyalarını döndürür."""
        cursor = self.conn.cursor()
        cursor.execute("SELECT * FROM pdf_files WHERE record_id = ?", (record_id,))
        return [dict(row) for row in cursor.fetchall()]

    def delete_pdf_file(self, pdf_id):
        """Deletes a PDF file entry from the database."""
        """Veritabanından bir PDF dosyası girişini siler."""
        cursor = self.conn.cursor()
        cursor.execute("DELETE FROM pdf_files WHERE id = ?", (pdf_id,))
        self.conn.commit()
        return cursor.rowcount > 0


class MainWindow(QMainWindow):
    """
    Main application window.
    Ana uygulama penceresi.
    """
    # List of ISO 3166-1 alpha-3 country codes for Place of Birth
    # Doğum Yeri için ISO 3166-1 alpha-3 ülke kodları listesi
    COUNTRY_CODES = [
        "ABW", "AFG", "AGO", "AIA", "ALA", "ALB", "AND", "ARE", "ARG", "ARM", "ASM", "ATA", "ATF", "ATG",
        "AUS", "AUT", "AZE", "BDI", "BEL", "BEN", "BES", "BFA", "BGD", "BGR", "BHR", "BHS", "BIH", "BLM",
        "BLR", "BLZ", "BMU", "BOL", "BRA", "BRB", "BRN", "BTN", "BVT", "BWA", "CAF", "CAN", "CCK", "CHE",
        "CHL", "CHN", "CIV", "CMR", "COD", "COG", "COK", "COK", "COL", "COM", "CPV", "CRI", "CUB", "CUW", "CXR",
        "CYM", "CYP", "CZE", "DEU", "DJI", "DMA", "DNK", "DOM", "DZA", "ECU", "EGY", "ERI", "ESH", "ESP",
        "EST", "ETH", "FIN", "FJI", "FLK", "FRA", "FRO", "FSM", "GAB", "GBR", "GEO", "GGY", "GHA", "GIB",
        "GIN", "GLP", "GMB", "GNB", "GNQ", "GRC", "GRD", "GRL", "GTM", "GUF", "GUM", "GUY", "HKG", "HMD",
        "HND", "HRV", "HTI", "HUN", "IDN", "IMN", "IND", "IOT", "IRL", "IRN", "IRQ", "ISL", "ISR", "ITA",
        "JAM", "JEY", "JOR", "JPN", "KAZ", "KEN", "KGZ", "KHM", "KIR", "KNA", "KOR", "KWT", "LAO", "LBN",
        "LBR", "LBY", "LCA", "LIE", "LKA", "LSO", "LTU", "LUX", "LVA", "MAC", "MAF", "MAR", "MCO", "MDA",
        "MDG", "MDV", "MEX", "MHL", "MKD", "MLI", "MLT", "MMR", "MNE", "MNG", "MNP", "MOZ", "MRT", "MSR",
        "MTQ", "MUS", "MWI", "MYS", "MYT", "NAM", "NCL", "NER", "NFK", "NGA", "NIC", "NIU", "NLD", "NOR",
        "NPL", "NRU", "NZL", "OMN", "PAK", "PAN", "PCN", "PER", "PHL", "PLW", "PNG", "POL", "PRI", "PRK",
        "PRT", "PRY", "PSE", "PYF", "QAT", "REU", "ROU", "RUS", "RWA", "SAU", "SLB", "SCN", "SDN", "SEN",
        "SGP", "SGS", "SHN", "SJM", "SLB", "SLE", "SLV", "SMR", "SOM", "SPM", "SRB", "SSD", "STP", "SUR",
        "SVK", "SVN", "SWE", "SWZ", "SXM", "SYC", "SYR", "TCA", "TCD", "TGO", "THA", "TJK", "TKL", "TKM",
        "TLS", "TON", "TTO", "TUN", "TUR", "TUV", "TWN", "TZA", "UGA", "UKR", "UMI", "URY", "USA", "UZB",
        "VAT", "VCT", "VEN", "VGB", "VIR", "VNM", "VUT", "WLF", "WSM", "YEM", "ZAF", "ZMB", "ZWE"
    ]


    def __init__(self):
        super().__init__()
        self.setWindowTitle("Certificate For Return To Turkmenistan") # Belge Yönetim Sistemi
        self.setGeometry(100, 100, 1000, 600)

        self.db = DatabaseManager()
        # "Empty" date value used for filtering
        # Filtreleme için kullanılan "boş" tarih değeri
        self._filter_min_date = QDate(2000, 1, 1)
        
        # Load the selected printer from settings
        # Ayarlardan seçilen yazıcıyı yükle
        self.selected_printer = self.db.get_setting('default_printer') 
        
        # Load the current authorized person from settings
        # Ayarlardan mevcut yetkili kişiyi yükle
        self.current_authorized_person_id = self.db.get_setting('current_auth_person_id')
        self.current_authorized_person_data = None # Will be populated after users are loaded

        # Flag indicating if date_of_expiry is being set automatically
        # Son kullanma tarihinin otomatik olarak ayarlanıp ayarlanmadığını gösteren bayrak
        self._setting_expiry_date_automatically = False 
        
        self.init_ui()

    def init_ui(self):
        """
        Initializes the user interface and creates tabs.
        Kullanıcı arayüzünü başlatır ve sekmeler oluşturur.
        """
        # Main tabs
        # Ana sekmeler
        self.tabs = QTabWidget()
        self.setCentralWidget(self.tabs)

        # Tab 1: Data Entry
        # Sekme 1: Veri Girişi
        self.tab1 = QWidget()
        self.tabs.addTab(self.tab1, "Data Entry") # Veri Girişi
        self.init_tab1()

        # Tab 2: Database
        # Sekme 2: Veritabanı
        self.tab2 = QWidget()
        self.tabs.addTab(self.tab2, "Database") # Veritabanı
        self.init_tab2()

        # Menu bar - Authorized Person Management and Printer Settings
        # Menü çubuğu - Yetkili Kişi Yönetimi ve Yazıcı Ayarları
        menubar = self.menuBar()
        user_menu = menubar.addMenu("Authorized Person Management") # Yetkili Kişi Yönetimi

        add_user_action = user_menu.addAction("Add Authorized Person") # Yetkili Kişi Ekle
        add_user_action.triggered.connect(self.add_user)

        manage_users_action = user_menu.addAction("Manage Authorized Persons") # Yetkili Kişileri Yönet
        manage_users_action.triggered.connect(self.manage_users)

        # New: Current Authorized Person
        # Yeni: Mevcut Yetkili Kişi
        user_menu.addSeparator()
        current_auth_person_action = user_menu.addAction("Set Current Authorized Person") # Mevcut Yetkili Kişiyi Ayarla
        current_auth_person_action.triggered.connect(self.set_current_authorized_person)

        printer_menu = menubar.addMenu("Printer Settings") # Yazıcı Ayarları
        select_printer_action = printer_menu.addAction("Select Printer") # Yazıcı Seç
        select_printer_action.triggered.connect(self.select_printer)
        
        # Status bar label to show selected printer
        # Seçilen yazıcıyı göstermek için durum çubuğu etiketi
        self.status_label = QLabel("Selected Printer: None") # Seçilen Yazıcı: Yok
        self.statusBar().addWidget(self.status_label)
        self.update_printer_status()


    def init_tab1(self):
        """
        Initializes the Data Entry tab and arranges form elements.
        Veri Girişi sekmesini başlatır ve form elemanlarını düzenler.
        """
        # Main layout with splitter for form and preview
        # Form ve önizleme için ayırıcı ile ana düzen
        splitter = QSplitter(Qt.Vertical)
        
        # Top widget for form
        # Form için üst widget
        form_widget = QWidget()
        form_layout = QVBoxLayout(form_widget)

        # Form fields
        # Form alanları
        entry_form_layout = QFormLayout()

        self.surname = QLineEdit()
        entry_form_layout.addRow("Surname:", self.surname) # Soyadı:

        self.given_names = QLineEdit()
        entry_form_layout.addRow("Given Names:", self.given_names) # Adı:

        self.sex = QComboBox()
        self.sex.addItems(["ERKEK/MALE", "AÝAL/FEMALE"]) # Gender options updated to include Turkmen and English
        entry_form_layout.addRow("Sex:", self.sex) # Cinsiyet:

        # Place of Birth with QCompleter for country codes
        # Ülke kodları için QCompleter ile Doğum Yeri
        self.place_of_birth = QLineEdit("TKM")
        self.country_completer = QCompleter(self.COUNTRY_CODES, self)
        self.country_completer.setCaseSensitivity(Qt.CaseInsensitive)
        self.country_completer.setFilterMode(Qt.MatchStartsWith)
        self.country_completer.setCompletionMode(QCompleter.PopupCompletion)
        self.place_of_birth.setCompleter(self.country_completer)
        entry_form_layout.addRow("Place of Birth:", self.place_of_birth) # Doğum Yeri:


        self.date_of_birth = QDateEdit()
        self.date_of_birth.setDisplayFormat("dd.MM.yyyy")
        self.date_of_birth.setCalendarPopup(True)
        entry_form_layout.addRow("Date of Birth:", self.date_of_birth) # Doğum Tarihi:

        self.date_of_issue = QDateEdit(QDate.currentDate())
        self.date_of_issue.setDisplayFormat("dd.MM.yyyy")
        self.date_of_issue.setCalendarPopup(True)
        # Connect dateChanged signal of date_of_issue to a slot that updates date_of_expiry
        # date_of_issue'nin dateChanged sinyalini, son kullanma tarihini güncelleyen bir yuvaya bağla
        self.date_of_issue.dateChanged.connect(self._update_date_of_expiry_auto)
        entry_form_layout.addRow("Date of Issue:", self.date_of_issue) # Veriliş Tarihi:

        self.date_of_expiry = QDateEdit()
        self.date_of_expiry.setDisplayFormat("dd.MM.yyyy")
        self.date_of_expiry.setCalendarPopup(True)
        # Disconnect automatic update when user manually changes date_of_expiry
        # Kullanıcı date_of_expiry'yi manuel olarak değiştirdiğinde otomatik güncellemeyi ayır
        self.date_of_expiry.dateChanged.connect(self._handle_expiry_date_manual_change)
        entry_form_layout.addRow("Date of Expiry:", self.date_of_expiry) # Son Geçerlilik Tarihi:

        # Initialize expiry date based on current issue date
        # Mevcut veriliş tarihine göre son kullanma tarihini başlat
        self._update_date_of_expiry_auto(self.date_of_issue.date())


        # Phone field with "+90 5" prefix
        # "+90 5" önekli telefon alanı
        phone_layout = QHBoxLayout()
        phone_prefix = QLabel("+90 5")
        phone_prefix.setFixedWidth(50)  # Adjust width as needed
                                        # Gerektiği gibi genişliği ayarla
        self.telefon = QLineEdit()
        self.telefon.setInputMask("000 000 00 00;_") # Mask for remaining digits
                                                    # Kalan rakamlar için maske
        phone_layout.addWidget(phone_prefix)
        phone_layout.addWidget(self.telefon)
        entry_form_layout.addRow("Phone:", phone_layout) # Telefon:

        # LP Number field with "A№" prefix
        # "A№" önekli LP Numarası alanı
        lp_layout = QHBoxLayout()
        self.lp_prefix = QLabel("A№") # Fixed 'A' before '№'
                                    # '№'den önce sabit 'A'
        self.lp_prefix.setFixedWidth(30) # Adjusted width for "A№"
                                        # "A№" için ayarlanmış genişlik
        self.lp_prefix.setAlignment(Qt.AlignCenter)
        lp_layout.addWidget(self.lp_prefix) # LP prefix label added
        self.lp_number = QLineEdit()
        lp_layout.addWidget(self.lp_number)
        entry_form_layout.addRow("LP Number:", lp_layout) # LP Numarası:

        # Authorized person selection - now this is just for manual override/view
        # Yetkili kişi seçimi - şimdi bu sadece manuel geçersiz kılma/görüntüleme içindir
        self.authorized_person = QComboBox()
        self.update_auth_person_combo() # Load authorized persons
                                        # Yetkili kişileri yükle
        entry_form_layout.addRow("Authorized Person:", self.authorized_person) # Yetkili Kişi:

        form_layout.addLayout(entry_form_layout)

        # Buttons
        # Butonlar
        btn_layout = QHBoxLayout()

        self.save_btn = QPushButton("Save to Database") # Veritabanına Kaydet
        self.save_btn.setShortcut(QKeySequence("Ctrl+S"))
        self.save_btn.clicked.connect(self.save_to_database)
        btn_layout.addWidget(self.save_btn)

        self.print_btn = QPushButton("Print Directly and Save") # Doğrudan Yazdır ve Kaydet
        self.print_btn.setShortcut(QKeySequence("Ctrl+P"))
        self.print_btn.clicked.connect(self.print_and_save)
        btn_layout.addWidget(self.print_btn)

        self.preview_btn = QPushButton("Preview") # Önizleme
        self.preview_btn.setShortcut(QKeySequence("Ctrl+R"))
        self.preview_btn.clicked.connect(self.preview_document)
        btn_layout.addWidget(self.preview_btn)

        self.clear_btn = QPushButton("Clear Form") # Formu Temizle
        self.clear_btn.clicked.connect(self.clear_form)
        btn_layout.addWidget(self.clear_btn)

        form_layout.addLayout(btn_layout)

        # Bottom widget for preview
        # Önizleme için alt widget
        preview_widget = QWidget()
        preview_layout = QVBoxLayout(preview_widget)
        
        preview_layout.addWidget(QLabel("Document Preview:")) # Belge Önizlemesi:
        
        self.preview_text = QTextEdit()
        self.preview_text.setReadOnly(True)
        preview_layout.addWidget(self.preview_text)
        
        # Add widgets to splitter
        # Bölücüye widget'ları ekle
        splitter.addWidget(form_widget)
        splitter.addWidget(preview_widget)
        
        # Set initial sizes (form larger than preview) - Adjusted for larger preview
        # Başlangıç boyutlarını ayarla (form önizlemeden daha büyük) - Daha büyük önizleme için ayarlandı
        splitter.setSizes([350, 400]) # Made preview area larger
                                    # Önizleme alanını büyüt

        # Set splitter as main layout
        # Bölücüyü ana düzen olarak ayarla
        tab1_layout = QVBoxLayout(self.tab1)
        tab1_layout.addWidget(splitter)

    def _update_date_of_expiry_auto(self, issue_date: QDate):
        """
        Sets the Expiry Date automatically to one month and one day before the issue date.
        This function is called when date_of_issue changes.
        Bu fonksiyon, date_of_issue değiştiğinde son kullanma tarihini otomatik olarak veriliş tarihinden bir ay bir gün öncesine ayarlar.
        """
        # Set flag to true to prevent _handle_expiry_date_manual_change from triggering
        # _handle_expiry_date_manual_change'nin tetiklenmesini önlemek için bayrağı true yap
        self._setting_expiry_date_automatically = False 
        
        # Calculate expiry date: add one month, then subtract one day
        # Son kullanma tarihini hesapla: bir ay ekle, sonra bir gün çıkar
        # QDate.addMonths handles month transitions correctly
        # QDate.addMonths ay geçişlerini doğru şekilde yönetir
        expiry_date = issue_date.addMonths(1)
        expiry_date = expiry_date.addDays(-1)
        self.date_of_expiry.setDate(expiry_date)
        
        # Reset flag
        # Bayrağı sıfırla
        self._setting_expiry_date_automatically = False

    def _handle_expiry_date_manual_change(self):
        """
        Manages manual changes to Expiry Date.
        If the change is manual, disconnects the signal from date_of_issue.
        Son Kullanma Tarihindeki manuel değişiklikleri yönetir. Değişiklik manuel ise, date_of_issue'den sinyali keser.
        """
        # If the change is not automatic, disconnect the signal from date_of_issue
        # Değişiklik otomatik değilse, date_of_issue'den sinyali kes
        if not self._setting_expiry_date_automatically:
            try:
                self.date_of_issue.dateChanged.disconnect(self._update_date_of_expiry_auto)
            except TypeError:
                # Signal already disconnected, ignore
                # Sinyal zaten bağlantısı kesilmiş, yoksay
                pass

    def init_tab2(self):
        """
        Initializes the Database tab and arranges the table and control elements.
        Veritabanı sekmesini başlatır ve tablo ile kontrol elemanlarını düzenler.
        """
        layout = QVBoxLayout()

        # Control buttons
        # Kontrol butonları
        control_layout = QHBoxLayout()

        self.refresh_btn = QPushButton("Refresh") # Yenile
        self.refresh_btn.clicked.connect(self.load_data)
        control_layout.addWidget(self.refresh_btn)

        self.export_btn = QPushButton("Export to Excel") # Excel'e Aktar
        self.export_btn.clicked.connect(self.export_to_excel_selected_or_all)
        control_layout.addWidget(self.export_btn)

        self.import_btn = QPushButton("Import from Excel") # Excel'den İçe Aktar
        self.import_btn.clicked.connect(self.import_from_excel)
        control_layout.addWidget(self.import_btn)

        self.print_selected_btn = QPushButton("Print Selected") # Seçileni Yazdır
        self.print_selected_btn.clicked.connect(self.print_selected)
        control_layout.addWidget(self.print_selected_btn)

        self.delete_selected_btn = QPushButton("Delete Selected") # Seçileni Sil
        self.delete_selected_btn.clicked.connect(self.delete_selected_records)
        control_layout.addWidget(self.delete_selected_btn)

        self.select_all_btn = QPushButton("Select All") # Tümünü Seç
        self.select_all_btn.clicked.connect(self.select_all)
        control_layout.addWidget(self.select_all_btn)

        # "Sort by Newest" button added
        # "En Yeniye Göre Sırala" butonu eklendi
        self.sort_latest_btn = QPushButton("Sort by Newest") # En Yeniye Göre Sırala
        self.sort_latest_btn.clicked.connect(self.sort_by_latest_records)
        control_layout.addWidget(self.sort_latest_btn)

        # "Sort by Oldest" button added
        self.sort_oldest_btn = QPushButton("Sort by Oldest")
        self.sort_oldest_btn.clicked.connect(self.sort_by_oldest_records)
        control_layout.addWidget(self.sort_oldest_btn)

        # New: Preview button for database entries
        # Yeni: Veritabanı girişleri için önizleme butonu
        self.preview_db_btn = QPushButton("Preview Record") # Kaydı Önizle
        self.preview_db_btn.clicked.connect(self.preview_selected_record)
        control_layout.addWidget(self.preview_db_btn)

        layout.addLayout(control_layout)

        # Filtering - using a special value for initial empty display
        # Filtreleme - başlangıçta boş görüntü için özel bir değer kullanma
        filter_layout = QHBoxLayout()

        # Advanced Live Search Input - shortened
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Search (Surname, Name, LP#, Phone)") # Ara (Soyadı, Adı, LP Numarası, Telefon) - shortened
        self.search_input.textChanged.connect(self.live_search) # Connect for live search
        # Adjusting stretch factor for the search input to make it shorter
        filter_layout.addWidget(self.search_input, 2) # Giving it a stretch of 2

        filter_layout.addWidget(QLabel("Filter by Date:")) # Tarihe Göre Filtrele:

        self.filter_date = QDateEdit()
        self.filter_date.setDisplayFormat("dd.MM.yyyy")
        self.filter_date.setCalendarPopup(True)
        # Set to current date but do not filter initially
        # Mevcut tarihe ayarla ama başlangıçta filtreleme yapma
        self.filter_date.setDate(QDate.currentDate())
        filter_layout.addWidget(self.filter_date, 1) # Giving it a stretch of 1

        self.filter_field = QComboBox()
        # First option is empty, user will choose
        # İlk seçenek boş, kullanıcı seçecek
        self.filter_field.addItem("Select Field") # Alan Seç
        self.filter_field.addItems([
            "Date of Birth", "Date of Issue", "Date of Expiry" # Doğum Tarihi, Veriliş Tarihi, Son Geçerlilik Tarihi
        ])
        self.filter_field.setCurrentIndex(0)
        filter_layout.addWidget(self.filter_field, 1) # Giving it a stretch of 1

        self.apply_filter_btn = QPushButton("Apply Date Filter") # Tarih Filtresi Uygula
        self.apply_filter_btn.clicked.connect(self.apply_date_filter)
        filter_layout.addWidget(self.apply_filter_btn)

        self.clear_filter_btn = QPushButton("Clear All Filters") # Tüm Filtreleri Temizle
        self.clear_filter_btn.clicked.connect(self.clear_all_filters)
        filter_layout.addWidget(self.clear_filter_btn)

        layout.addLayout(filter_layout)

        # Table
        # Tablo
        self.table = QTableView()
        self.table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table.setSelectionMode(QAbstractItemView.ExtendedSelection) # Allow multi-selection
                                                                        # Çoklu seçime izin ver
        self.model = QStandardItemModel()
        self.proxy_model = QSortFilterProxyModel()
        self.proxy_model.setSourceModel(self.model)
        self.table.setModel(self.proxy_model)

        # Set the custom delegate for the "PDF" column
        # "PDF" sütunu için özel delegeyi ayarla
        self.table.setItemDelegateForColumn(10, ButtonDelegate(self)) # Index 10 for "PDF" column

        # Connect selection model's selectionChanged signal
        self.table.selectionModel().selectionChanged.connect(self.update_selected_count_label)

        # Column headers
        # Sütun başlıkları
        # Changed: Added a new column for PDF
        self.model.setHorizontalHeaderLabels([
            "ID", "Surname", "Given Names", "Sex", "Place of Birth", "Date of Birth",
            "Date of Issue", "Date of Expiry", "Phone", "A №", "PDF" 
        ])

        # Table resizing
        # Tablo yeniden boyutlandırma
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.table.verticalHeader().setVisible(False)

        # Right-click context menu
        # Sağ tıklama bağlam menüsü
        self.table.setContextMenuPolicy(Qt.CustomContextMenu)
        self.table.customContextMenuRequested.connect(self.open_table_context_menu)

        layout.addWidget(self.table)

        # Info label
        # Bilgi etiketi
        self.info_label = QLabel("Total Records: 0") # Toplam Kayıt: 0
        layout.addWidget(self.info_label)

        # New label for selected record count
        self.selected_info_label = QLabel("Selected Records: 0")
        layout.addWidget(self.selected_info_label)


        self.tab2.setLayout(layout)
        self.load_data() # Load data
                         # Verileri yükle

    def open_table_context_menu(self, position):
        """
        Opens the right-click context menu for the table.
        Tablo için sağ tıklama bağlam menüsü.
        """
        menu = QMenu()
        
        print_action = menu.addAction("Print Selected") # Seçileni Yazdır
        delete_action = menu.addAction("Delete Selected") # Seçileni Sil
        export_excel_action = menu.addAction("Export Selected to Excel") # Seçileni Excel'e Aktar
        preview_action = menu.addAction("Preview Record") # Kaydı Önizle
        
        # New: PDF Management action in context menu
        # Yeni: Bağlam menüsünde PDF Yönetimi eylemi
        pdf_action = menu.addAction("Manage PDF") # PDF Yönet
        # No action to connect to (will be handled by exec_)

        action = menu.exec_(self.table.viewport().mapToGlobal(position))

        if action == print_action:
            self.print_selected()
        elif action == delete_action:
            self.delete_selected_records()
        elif action == export_excel_action:
            self.export_to_excel_selected_or_all(selected_only=True)
        elif action == preview_action:
            self.preview_selected_record()
        elif action == pdf_action:
            self.open_pdf_management_dialog_from_context()

    def update_auth_person_combo(self):
        """
        Updates the authorized person dropdown list and sets the current authorized person.
        Yetkili kişi açılır listesini günceller ve mevcut yetkili kişiyi ayarlar.
        """
        self.authorized_person.clear()
        users = self.db.get_users()
        
        # Populate the combo box
        # Açılır kutuyu doldur
        for user in users:
            display_text = f"{user['organization']} - {user['title']} - {user['fullname']}"
            self.authorized_person.addItem(display_text, user) # Store the whole user dictionary as data
                                                                # Tüm kullanıcı sözlüğünü veri olarak sakla

        # Set the current authorized person from settings if available
        # Ayarlardan mevcut yetkili kişiyi ayarla (varsa)
        current_auth_person_id_from_settings = self.db.get_setting('current_auth_person_id')
        
        found_current_user = False
        if current_auth_person_id_from_settings:
            try:
                # Find the user by ID
                # Kimliğe göre kullanıcıyı bul
                for user in users:
                    if str(user['id']) == current_auth_person_id_from_settings:
                        self.current_authorized_person_data = user
                        found_current_user = True
                        break
                
                if found_current_user:
                    # If user is found, set it as current in the combo box
                    # Kullanıcı bulunursa, açılır kutuda mevcut olarak ayarla
                    for i in range(self.authorized_person.count()):
                        if self.authorized_person.itemData(i)['id'] == self.current_authorized_person_data['id']:
                            self.authorized_person.setCurrentIndex(i)
                            break
                else:
                    # If current_authorized_person_id_from_settings is not found in the database, clear it
                    # current_authorized_person_id_from_settings veritabanında bulunamazsa, temizle
                    self.db.save_setting('current_auth_person_id', '')
                    self.current_authorized_person_data = None # Reset
            except Exception:
                # Handle potential errors if parsing fails or data is corrupt
                # Ayrıştırma başarısız olursa veya veriler bozuksa olası hataları ele al
                self.current_authorized_person_data = None
                self.db.save_setting('current_auth_person_id', '') # Clear invalid setting
        
        # If no current authorized person is set in settings, or if the stored ID was invalid/deleted,
        # default to the first one if available.
        # Ayarlarda mevcut yetkili kişi ayarlanmamışsa veya saklanan kimlik geçersiz/silinmişse,
        # varsa ilkine varsayılan olarak ayarla.
        if not self.current_authorized_person_data and users:
            self.current_authorized_person_data = users[0]
            self.authorized_person.setCurrentIndex(0) # Select the first item
            self.db.save_setting('current_auth_person_id', str(users[0]['id']))
        elif not users:
            self.current_authorized_person_data = None
            self.db.save_setting('current_auth_person_id', '') # No users, clear setting


    def set_current_authorized_person(self):
        """
        Allows the user to select and set a 'current' authorized person.
        Kullanıcının bir 'mevcut' yetkili kişi seçmesine ve ayarlamasına olanak tanır.
        """
        users = self.db.get_users()
        if not users:
            QMessageBox.information(self, "Information", "No authorized persons registered to select.") # Bilgi, Seçmek için kayıtlı yetkili kişi yok.
            return

        user_names = [
            f"{user['id']}. {user['organization']} - {user['title']} - {user['fullname']}"
            for user in users
        ]
        
        # Pre-select the currently set authorized person in the dialog
        # Diyalogda şu anda ayarlı yetkili kişiyi önceden seç
        current_idx = -1
        if self.current_authorized_person_data:
            for i, user in enumerate(users):
                if user['id'] == self.current_authorized_person_data['id']:
                    current_idx = i
                    break

        selected_user_display, ok = QInputDialog.getItem(
            self, "Set Current Authorized Person", "Select the current authorized person:", # Mevcut Yetkili Kişiyi Ayarla, Mevcut yetkili kişiyi seçin:
            user_names, current_idx, False
        )

        if ok and selected_user_display:
            # Extract ID from the display string
            # Görüntü dizesinden Kimliği çıkar
            selected_id_str = selected_user_display.split('.')[0].strip()
            # Find the actual user object
            # Gerçek kullanıcı nesnesini bul
            selected_user = next((u for u in users if f"{u['id']}" == selected_id_str), None)

            if selected_user:
                self.current_authorized_person_data = selected_user
                self.db.save_setting('current_auth_person_id', str(selected_user['id']))
                # Update the combo box in Data Entry tab to reflect the change
                # Veri Girişi sekmesindeki açılır kutuyu değişikliği yansıtacak şekilde güncelle
                self.update_auth_person_combo()
                QMessageBox.information(self, "Success", f"Current authorized person set to: {selected_user['fullname']}") # Başarı, Mevcut yetkili kişi şuna ayarlandı:
            else:
                QMessageBox.warning(self, "Error", "Selected authorized person not found in database.") # Hata, Seçilen yetkili kişi veritabanında bulunamadı.
        

    def format_date(self, date):
        """
        Returns date in 'dd.MM.yyyy' format.
        Tarihi 'dd.MM.yyyy' formatında döndürür.
        """
        return date.toString("dd.MM.yyyy")

    def clear_form(self):
        """
        Clears the fields in the entry form.
        Giriş formundaki alanları temizler.
        """
        # Clear all fields except date fields
        # Tarih alanları dışındaki tüm alanları temizle
        self.surname.clear()
        self.given_names.clear()
        self.sex.setCurrentIndex(0)
        self.place_of_birth.setText("TKM") # Reset to default TKM
        # Clear only the editable part of the phone number
        # Telefon numarasının yalnızca düzenlenebilir kısmını temizle
        self.telefon.clear() 
        self.lp_number.clear()

        # Reconnect date_of_issue signal if it was disconnected
        # date_of_issue sinyali bağlantısı kesildiyse yeniden bağla
        try:
            self.date_of_issue.dateChanged.disconnect(self._update_date_of_expiry_auto)
        except TypeError:
            pass # Signal already disconnected or never connected
                # Sinyal zaten bağlantısı kesilmiş veya hiç bağlanmamış

        self.date_of_issue.dateChanged.connect(self._update_date_of_expiry_auto)

        # Reset issue date to current date
        # Veriliş tarihini bugünün tarihine sıfırla
        self.date_of_issue.setDate(QDate.currentDate())
        # Reset birth and expiry dates to default values
        # Doğum ve son kullanma tarihlerini varsayılan değerlere sıfırla
        self.date_of_birth.setDate(QDate(2000, 1, 1)) # Or a starting date you define
                                                    # Veya tanımladığınız bir başlangıç tarihi
        
        # Automatically set expiry date based on the newly reset issue date
        # Yeni sıfırlanan veriliş tarihine göre son kullanma tarihini otomatik olarak ayarla
        self._update_date_of_expiry_auto(self.date_of_issue.date())


        # Update authorized person combo box
        # Yetkili kişi açılır kutusunu güncelle
        self.update_auth_person_combo() # This will now reset to the 'Current Authorized Person'
                                        # Bu şimdi 'Mevcut Yetkili Kişi'ye sıfırlanacak
        # Clear preview
        # Önizlemeyi temizle
        self.preview_text.clear()

        # Set focus back to the first input field (surname)
        # Odağı ilk giriş alanına (soyadı) geri ayarla
        self.surname.setFocus()

    def save_to_database(self):
        """
        Collects and validates data, saves to database without printing.
        Verileri toplar ve doğrular, yazdırmadan veritabanına kaydeder.
        """
        # Collect data
        # Verileri topla
        record = {
            'surname': self.surname.text().strip(),
            'given_names': self.given_names.text().strip(),
            'sex': self.sex.currentText(), # This will now be 'ERKEK/MALE' or 'AÝAL/FEMALE'
            'place_of_birth': self.place_of_birth.text().strip(),
            'date_of_birth': self.format_date(self.date_of_birth.date()),
            'date_of_issue': self.format_date(self.date_of_issue.date()),
            'date_of_expiry': self.format_date(self.date_of_expiry.date()),
            'telefon': "+90 5" + self.telefon.text().strip(), # Prepend fixed part
                                                            # Sabit kısmı öne ekle
            'lp_number': "A№" + self.lp_number.text().strip(), # Changed to include '№'
                                                            # '№' içerecek şekilde değiştirildi
            'auth_person': self.current_authorized_person_data # Use the current authorized person
        }

        # Validation
        # Doğrulama
        if not all([record['surname'], record['given_names'], record['lp_number']]):
            QMessageBox.warning(self, "Missing Information", "Please fill in all required fields!") # Eksik Bilgi, Lütfen tüm gerekli alanları doldurun!
            return
        if not self.current_authorized_person_data:
            QMessageBox.warning(self, "Missing Information", "Please select a Current Authorized Person under 'Authorized Person Management' menu.") # Eksik Bilgi, Lütfen 'Yetkili Kişi Yönetimi' menüsü altında bir Mevcut Yetkili Kişi seçin.
            return

        # Save to database
        # Veritabanına kaydet
        self.db.add_record(record)
        self.load_data() # Refresh database table
                         # Veritabanı tablosunu yenile
        # QMessageBox.information(self, "Success", "Record successfully saved to database!") # Message removed
        self.clear_form() # Clear form and set focus
                          # Formu temizle ve odağı ayarla

    def generate_preview_text(self, record):
        """
        Generates a plain text preview of the document matching the Word output.
        Belgenin Word çıktısına uygun düz metin önizlemesini oluşturur.
        """
        # This function creates the content for the preview_text QTextEdit
        # and closely mimics the structure and spacing of the generated Word document.
        # Bu fonksiyon, preview_text QTextEdit için içeriği oluşturur
        # ve oluşturulan Word belgesinin yapısını ve boşluklarını yakından taklit eder.
        
        # Word document parameters for alignment and spacing
        # Hizalama ve boşluk için Word belgesi parametreleri
        # These are approximations based on typical font sizes and indents
        # Bunlar, tipik yazı tipi boyutları ve girintilere dayalı yaklaşıklıklardır
        # and are used to simulate the visual layout of the Word document.
        # ve Word belgesinin görsel düzenini simüle etmek için kullanılır.
        
        surname = record['surname'].upper()
        given_names = record['given_names'].upper()
        date_of_birth = record['date_of_birth']
        place_of_birth = record['place_of_birth']
        sex = record['sex'].upper() # This will use the value from the QComboBox directly
        date_of_issue = record['date_of_issue']
        date_of_expiry = record['date_of_expiry']
        
        # Use the provided record's auth_person data, not the current_authorized_person_data
        # Sağlanan kaydın auth_person verilerini kullan, current_authorized_person_data'yı değil
        auth = record.get('auth_person', {}) or {'organization':'', 'title':'', 'fullname':''}
        organization = auth.get('organization','')
        title = auth.get('title','')
        fullname = auth.get('fullname','')

        # --- Simulate Word Document Layout ---
        # Word Belgesi Düzenini Simüle Et
        # Adjust spacing and tabs to visually match the DOCX output.
        # DOCX çıktısıyla görsel olarak eşleşmesi için boşlukları ve sekmeleri ayarla.
        # This will require careful use of spaces and newlines.
        # Bu, boşlukların ve yeni satırların dikkatli kullanılmasını gerektirecektir.
        
        # Initial top spacing
        # Başlangıçtaki üst boşluk
        preview_lines = ["", ""] # Two empty lines for top margin simulation
                                # Üst kenar boşluğu simülasyonu için iki boş satır

        # Surname and Given Names - Left aligned with first line indent
        # Soyadı ve Adı - İlk satır girintili sola hizalı
        # Using string formatting to simulate indent
        # Girintiyi simüle etmek için dize biçimlendirmesi kullanma
        indent = " " * 18 # Approximate number of characters for Cm(4.2) indent with Arial 12pt
                        # Cm(4.2) girintisi için yaklaşık karakter sayısı Arial 12pt ile
        preview_lines.append(f"{indent}{surname}")
        preview_lines.append(f"{indent}{given_names}")

        base_line_start_preview_char = 18
        conceptual_tab_stop_preview_char = 45 
        shifted_tab_stop_preview_char = conceptual_tab_stop_preview_char - 4 

        left_text_nationality = "TURKMENISTAN"
        spacing_nationality = " " * (shifted_tab_stop_preview_char - len(left_text_nationality) - base_line_start_preview_char)
        preview_lines.append(f"{indent}{left_text_nationality}{spacing_nationality}{date_of_birth}")

        left_text_place_sex = place_of_birth
        spacing_place_sex = " " * (shifted_tab_stop_preview_char - len(left_text_place_sex) - base_line_start_preview_char)
        preview_lines.append(f"{indent}{left_text_place_sex}{spacing_place_sex}{sex}")

        left_text_dates = date_of_issue
        spacing_dates = " " * (shifted_tab_stop_preview_char - len(left_text_dates) - base_line_start_preview_char)
        preview_lines.append(f"{indent}{left_text_dates}{spacing_dates}{date_of_expiry}")

        preview_lines.append("") 

        preview_lines.append(f"{indent}{organization}")
        preview_lines.append("") 

        left_text_auth = title
        right_text_auth = fullname
        spacing_auth = " " * (conceptual_tab_stop_preview_char - len(left_text_auth) - base_line_start_preview_char) 
        preview_lines.append(f"{indent}{left_text_auth}{spacing_auth}{right_text_auth}")

        return "\n".join(preview_lines)


    def preview_document(self):
        """
        Generates a preview of the document and displays it in the preview area.
        Belgenin önizlemesini oluşturur ve önizleme alanında gösterir.
        """
        # Collect data
        # Verileri topla
        record = {
            'surname': self.surname.text().strip(),
            'given_names': self.given_names.text().strip(),
            'sex': self.sex.currentText(), # This will now be 'ERKEK/MALE' or 'AÝAL/FEMALE'
            'place_of_birth': self.place_of_birth.text().strip(),
            'date_of_birth': self.format_date(self.date_of_birth.date()),
            'date_of_issue': self.format_date(self.date_of_issue.date()),
            'date_of_expiry': self.format_date(self.date_of_expiry.date()),
            'telefon': "+90 5" + self.telefon.text().strip(), # Prepend fixed part
                                                            # Sabit kısmı öne ekle
            'lp_number': "A№" + self.lp_number.text().strip(), # Changed to include '№'
                                                            # '№' içerecek şekilde değiştirildi
            'auth_person': self.current_authorized_person_data # Use the current authorized person for preview
        }

        # Validation
        # Doğrulama
        if not all([record['surname'], record['given_names'], record['lp_number']]):
            QMessageBox.warning(self, "Missing Information", "Please fill in all required fields!") # Eksik Bilgi, Lütfen tüm gerekli alanları doldurun!
            return
        if not self.current_authorized_person_data:
            QMessageBox.warning(self, "Missing Information", "Please select a Current Authorized Person under 'Authorized Person Management' menu.") # Eksik Bilgi, Lütfen 'Yetkili Kişi Yönetimi' menüsü altında bir Mevcut Yetkili Kişi seçin.
            return

        # Generate preview text
        # Önizleme metni oluştur
        preview_text = self.generate_preview_text(record)
        self.preview_text.setPlainText(preview_text)

    def print_and_save(self):
        """
        Collects and validates data, saves to database, generates document, and prints directly.
        Verileri toplar ve doğrular, veritabanına kaydeder, belge oluşturur ve doğrudan yazdırır.
        """
        # Collect data
        # Verileri topla
        record = {
            'surname': self.surname.text().strip(),
            'given_names': self.given_names.text().strip(),
            'sex': self.sex.currentText(), # This will now be 'ERKEK/MALE' or 'AÝAL/FEMALE'
            'place_of_birth': self.place_of_birth.text().strip(),
            'date_of_birth': self.format_date(self.date_of_birth.date()),
            'date_of_issue': self.format_date(self.date_of_issue.date()),
            'date_of_expiry': self.format_date(self.date_of_expiry.date()),
            'telefon': "+90 5" + self.telefon.text().strip(), # Prepend fixed part
                                                            # Sabit kısmı öne ekle
            'lp_number': "A№" + self.lp_number.text().strip(), # Changed to include '№'
                                                            # '№' içerecek şekilde değiştirildi
            'auth_person': self.current_authorized_person_data # Use the current authorized person for saving and printing
        }

        # Validation
        # Doğrulama
        if not all([record['surname'], record['given_names'], record['lp_number']]):
            QMessageBox.warning(self, "Missing Information", "Please fill in all required fields!") # Eksik Bilgi, Lütfen tüm gerekli alanları doldurun!
            return
        if not self.current_authorized_person_data:
            QMessageBox.warning(self, "Missing Information", "Please select a Current Authorized Person under 'Authorized Person Management' menu.") # Eksik Bilgi, Lütfen 'Yetkili Kişi Yönetimi' menüsü altında bir Mevcut Yetkili Kişi seçin.
            return

        # Generate document
        # Belge oluştur
        filepath = self.generate_docx(record, open_after_save=False) # Do not open for preview
                                                                    # Önizleme için açma
        if filepath:
            # Print document directly
            # Belgeyi doğrudan yazdır
            self.print_file(filepath, self.selected_printer)

            # Try to remove the temporary Word file from user's disk.
            # Geçici Word dosyasını kullanıcının diskinden kaldırmayı dene.
            try:
                # Short pause to allow print job to read the file (helps on Windows)
                # Yazdırma işinin dosyayı okumasına izin vermek için kısa bir duraklama (Windows'ta yardımcı olur)
                time.sleep(2) # Increased sleep from 1 second to 5 seconds
                            # Uyku süresi 1 saniyeden 5 saniyeye çıkarıldı
                if os.path.exists(filepath):
                    os.remove(filepath)
            except Exception:
                pass

            # Save to database after successful printing
            # Başarılı yazdırmadan sonra veritabanına kaydet
            self.db.add_record(record)
            self.load_data() # Refresh database table
                             # Veritabanı tablosunu yenile
            # QMessageBox.information(self, "Success", "Record successfully saved to database!") # Message removed
        else:
            QMessageBox.warning(self, "Error", "Document could not be generated.") # Hata, Belge oluşturulamadı.

        # Clear form and set focus
        # Formu temizle ve odağı ayarla
        self.clear_form()

    
    def generate_docx(self, record, open_after_save=False):
        """
        Generates a document matching the layout in lp-word.pdf (A4).
        Only layout/margins/indents/tab stops adjusted to match the sample.
        lp-word.pdf (A4) dosyasındaki düzenle eşleşen bir belge oluşturur.
        Yalnızca düzen/kenar boşlukları/girintiler/sekme durakları örnekle eşleşecek şekilde ayarlandı.
        """
        doc = Document()

        # Page margins - Adjusted to match lp-word.pdf example on A4
        # Sayfa kenar boşlukları - A4'teki lp-word.pdf örneğiyle eşleşecek şekilde ayarlandı
        for section in doc.sections:
            section.left_margin = Cm(1.0)   # left margin (approx)
                                            # sol kenar boşluğu (yaklaşık)
            section.right_margin = Cm(2.5)  # right margin (approx)
                                            # sağ kenar boşluğu (yaklaşık)
            section.top_margin = Cm(0.3)    # top margin (adjusted)
                                            # üst kenar boşluğu (ayarlı)
            section.bottom_margin = Cm(1.0) # bottom margin (approx)
                                            # alt kenar boşluğu (yaklaşık)

        # Style settings
        # Stil ayarları
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'  # Font changed to Times New Roman
                                    # Yazı tipi Times New Roman olarak değiştirildi
        font.size = Pt(11)             # Font size changed to 11
                                    # Yazı tipi boyutu 11 olarak değiştirildi

        # Helper: add paragraph with first line indent
        # Yardımcı: ilk satır girintili paragraf ekle
        def add_paragraph_with_indent(doc_obj, text, font_size=Pt(11), bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, # Font size changed to 11
                                      space_before=0, space_after=0, left_indent=0, first_line_indent=0):
            p = doc_obj.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_before = space_before
            p.paragraph_format.space_after = space_after
            p.paragraph_format.left_indent = left_indent
            p.paragraph_format.first_line_indent = first_line_indent
            run = p.add_run(text)
            run.bold = bold
            run.font.size = font_size
            run.font.name = 'Times New Roman' # Font changed to Times New Roman
            try:
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman') # Font changed to Times New Roman
            except Exception:
                pass
            return p

        # Helper: add paragraph with right aligned tab stop (left and right text)
        # Yardımcı: sağa hizalı sekme durağı olan paragraf ekle (sol ve sağ metin)
        def add_tabbed_paragraph(doc_obj, left_text, right_text, font_size=Pt(11), bold=True, # Font size changed to 11
                                 space_before=0, space_after=0, left_indent=0, tab_stop_pos_relative=None):
            p = doc_obj.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = space_before
            p.paragraph_format.space_after = space_after
            p.paragraph_format.left_indent = left_indent

            if tab_stop_pos_relative is not None:
                try:
                    tab_stops = p.paragraph_format.tab_stops
                    tab_stops.clear_all() # Clear existing tab stops to ensure only ours is applied
                                          # Mevcut sekme duraklarını temizle, böylece sadece bizimki uygulanır
                    tab_stops.add_tab_stop(tab_stop_pos_relative, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
                except Exception:
                    # Some python-docx versions might not expose tab_stops; quietly ignore.
                    # Bazı python-docx sürümleri sekme duraklarını göstermeyebilir; sessizce yoksay.
                    pass

            run_left = p.add_run(left_text)
            run_left.bold = bold
            run_left.font.size = font_size
            run_left.font.name = 'Times New Roman' # Font changed to Times New Roman
            try:
                run_left._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman') # Font changed to Times New Roman
            except Exception:
                pass

            p.add_run('	') # This is the tab character
            
            run_right = p.add_run(right_text)
            run_right.bold = bold
            run_right.font.size = font_size
            run_right.font.name = 'Times New Roman' # Font changed to Times New Roman
            try:
                run_right._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman') # Font changed to Times New Roman
            except Exception:
                pass

            return p
            
        # Yeni yardımcı fonksiyon: Sağ sütunu biraz sola kaydırmak için ayarlanmış sekme duraklı paragraf ekler.
        # New helper function: Adds a paragraph with an adjusted tab stop to shift the right column slightly left.
        def add_right_shifted_tabbed_paragraph(doc_obj, left_text, right_text, font_size=Pt(11), bold=True,
                                               space_before=0, space_after=0, left_indent=0,
                                               base_tab_stop_pos_relative=Cm(11), shift_amount=Cm(1.0)):
            p = doc_obj.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = space_before
            p.paragraph_format.space_after = space_after
            p.paragraph_format.left_indent = left_indent

            # Sağ sütunu sola doğru kaydırmak için yeni, ayarlanmış sekme durağı konumu hesapla.
            # Calculate the new, adjusted tab stop position to shift the right column left.
            shifted_tab_stop_pos = base_tab_stop_pos_relative - shift_amount

            try:
                tab_stops = p.paragraph_format.tab_stops
                tab_stops.clear_all() # Mevcut sekme duraklarını temizle, böylece sadece bu uygulanır.
                                      # Clear existing to ensure only this one applies.
                tab_stops.add_tab_stop(shifted_tab_stop_pos, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
            except Exception:
                pass

            run_left = p.add_run(left_text)
            run_left.bold = bold
            run_left.font.size = font_size
            run_left.font.name = 'Times New Roman'
            try:
                run_left._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            except Exception:
                pass

            p.add_run('	') # Bu bir sekme karakteridir.
                            # This is the tab character.

            run_right = p.add_run(right_text)
            run_right.bold = bold
            run_right.font.size = font_size
            run_right.font.name = 'Times New Roman'
            try:
                run_right._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            except Exception:
                pass

            return p

        # Top spacing to mimic lp-word.pdf layout
        # lp-word.pdf düzenini taklit etmek için üst boşluk
        add_paragraph_with_indent(doc, "", space_after=Pt(5))
        add_paragraph_with_indent(doc, "", space_after=Pt(3))

        # Surname and Given Names with a first line indent matching the sample
        # Örnekle eşleşen ilk satır girintili Soyadı ve Adı
        # First line indent set to 5.2 cm to align with lp-word.pdf layout.
        # İlk satır girintisi lp-word.pdf düzeniyle hizalamak için 5.2 cm'ye ayarlandı.
        add_paragraph_with_indent(doc, record['surname'].upper(), font_size=Pt(11), bold=True, # Font size changed to 11
                                  left_indent=Cm(0), first_line_indent=Cm(3.7), space_after=Pt(8))
        add_paragraph_with_indent(doc, record['given_names'].upper(), font_size=Pt(11), bold=True, # Font size changed to 11
                                  left_indent=Cm(0), first_line_indent=Cm(3.7), space_after=Pt(10))

        # Tab stop position for the main block, set to 11 cm relative to the paragraph's left indent.
        # Ana blok için sekme durağı konumu, paragrafın sol girintisine göre 11 cm olarak ayarlandı.
        # Bu değer, 'Title' ve 'Full Name' gibi diğer alanlar için de kullanılacaktır.
        # This value will also be used for other fields like 'Title' and 'Full Name'.
        tab_stop_relative_pos = Cm(11)

        # Nationality and Date of Birth (left and right)
        # Sağ sütunu sola doğru kaydırmak için yeni eklenen fonksiyonu kullan.
        # Use the newly added function to shift the right column left.
        add_right_shifted_tabbed_paragraph(doc, "TURKMENISTAN", record['date_of_birth'], font_size=Pt(11), bold=True,
                                           left_indent=Cm(3.7), base_tab_stop_pos_relative=tab_stop_relative_pos,
                                           shift_amount=Cm(1.0), space_after=Pt(12.5)) # Sağ sütunu 1 cm sola kaydır
                                                                                       # Shift right column 1 cm left

        # Place of Birth and Sex
        # Sağ sütunu sola doğru kaydırmak için yeni eklenen fonksiyonu kullan.
        # Use the newly added function to shift the right column left.
        add_right_shifted_tabbed_paragraph(doc, record['place_of_birth'], record['sex'].upper(), font_size=Pt(11), bold=True,
                                           left_indent=Cm(3.7), base_tab_stop_pos_relative=tab_stop_relative_pos,
                                           shift_amount=Cm(1.0), space_after=Pt(12.5)) # Sağ sütunu 1 cm sola kaydır
                                                                                       # Shift right column 1 cm left

        # Date of Issue and Date of Expiry
        # Sağ sütunu sola doğru kaydırmak için yeni eklenen fonksiyonu kullan.
        # Use the newly added function to shift the right column left.
        add_right_shifted_tabbed_paragraph(doc, record['date_of_issue'], record['date_of_expiry'], font_size=Pt(11), bold=True,
                                           left_indent=Cm(3.7), base_tab_stop_pos_relative=tab_stop_relative_pos,
                                           shift_amount=Cm(1.0), space_after=Pt(9.5)) # Sağ sütunu 1 cm sola kaydır
                                                                                     # Shift right column 1 cm left

        # Authorized / institution block
        # Yetkili / kurum bloğu
        # Use the auth_person data from the record, which should be the *current* authorized person when generated
        # Kayıttaki auth_person verilerini kullan, bu oluşturulduğunda *mevcut* yetkili kişi olmalıdır
        auth = record.get('auth_person', {}) or {'organization':'', 'title':'', 'fullname':''}

        # Organization name - slightly smaller font to match sample spacing
        # Kuruluş adı - örnek boşluğa uyacak şekilde biraz daha küçük yazı tipi
        add_paragraph_with_indent(doc, auth.get('organization',''), font_size=Pt(11), bold=True, # Font size changed to 11
                                  left_indent=Cm(0), first_line_indent=Cm(3.6), space_before=Pt(0), space_after=Pt(7.0))

        # Title and Full Name - left and right aligned using the original tab stop (11 cm)
        # Ünvan ve Tam Adı - orijinal sekme durağını (11 cm) kullanarak sola ve sağa hizalı
        add_tabbed_paragraph(doc, auth.get('title',''), auth.get('fullname',''), font_size=Pt(11), bold=True, # Font size changed to 11
                             left_indent=Cm(2.8), tab_stop_pos_relative=tab_stop_relative_pos, space_after=Pt(7.0))

        # Save to a temporary file and return path
        # Geçici bir dosyaya kaydet ve yolu döndür
        safe_name = re.sub(r'[^A-Za-z0-9_\-\.]', '_', f"LP_{record['surname']}_{record['given_names']}.docx")
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx', prefix='LP_')
        filename = tmp.name
        tmp.close()
        doc.save(filename)

        return filename


    def open_file(self, filepath):
        """
        Opens the file with its default application based on the platform.
        Dosyayı platforma göre varsayılan uygulamasıyla açar.
        """
        try:
            if sys.platform.startswith("win"):
                os.startfile(filepath)
            elif sys.platform.startswith("darwin"):  # macOS
                subprocess.run(["open", filepath], check=True)
            else:  # Linux
                subprocess.run(["xdg-open", filepath], check=True)
        except Exception as e:
            QMessageBox.warning(self, "File Open Error", f"Could not open file:\n{str(e)}") # Dosya Açma Hatası, Dosya açılamadı

    def select_printer(self):
        """
        Allows the user to select a printer from the system.
        Kullanıcının sistemden bir yazıcı seçmesine olanak tanır.
        """
        # Pass the current selected printer to the dialog for initial selection
        # Başlangıç seçimi için mevcut seçilen yazıcıyı diyaloga ilet
        dialog = PrinterSelectionDialog(self, initial_printer=self.selected_printer)
        if dialog.exec_() == QDialog.Accepted:
            selected = dialog.get_selected_printer()
            if selected and selected != "No printer found.": # Yazıcı bulunamadı
                self.selected_printer = selected
                # Save the selected printer to settings
                # Seçilen yazıcıyı ayarlara kaydet
                self.db.save_setting('default_printer', self.selected_printer)
                QMessageBox.information(self, "Printer Selection", f"Printer '{self.selected_printer}' selected.") # Yazıcı Seçimi, Yazıcı seçildi
                self.update_printer_status()
            else:
                self.selected_printer = None
                # Clear the default printer setting if nothing is selected
                # Hiçbir şey seçilmezse varsayılan yazıcı ayarını temizle
                self.db.save_setting('default_printer', '') 
                QMessageBox.warning(self, "Printer Selection", "No printer selected or no printer found.") # Yazıcı Seçimi, Yazıcı seçilmedi veya yazıcı bulunamadı.
                self.update_printer_status()
        
    def update_printer_status(self):
        """
        Updates the printer information in the status bar.
        Durum çubuğundaki yazıcı bilgilerini günceller.
        """
        if self.selected_printer:
            self.status_label.setText(f"Selected Printer: {self.selected_printer}") # Seçilen Yazıcı:
        else:
            self.status_label.setText("Selected Printer: None (Default will be used)") # Seçilen Yazıcı: Yok (Varsayılan kullanılacak)

    def print_file(self, filepath, printer_name=None):
        """
        Sends the .docx file to the default printer on Windows or selected/default on Unix-like systems.
        .docx dosyasını Windows'ta varsayılan yazıcıya veya Unix benzeri sistemlerde seçilen/varsayılan yazıcıya gönderir.
        """
        try:
            if sys.platform.startswith("win"):
                # On Windows, for .docx files, the most reliable way to print is to allow ShellExecute to
                # Windows'ta, .docx dosyaları için en güvenilir yazdırma yolu, ShellExecute'in
                # handle it via the 'print' verb of the associated application. This will print to the system's default printer.
                # ilgili uygulamanın 'print' fiili aracılığıyla onu işlemesine izin vermektir. Bu, sistemin varsayılan yazıcısına yazdıracaktır.
                if win32api: # Ensure win32api of pywin32 is available
                            # pywin32'nin win32api'sinin mevcut olduğundan emin olun
                    try:
                        # If a specific printer is provided, try the 'printto' verb which allows us to specify the target printer.
                        # Belirli bir yazıcı sağlanmışsa, hedef yazıcıyı belirtmemizi sağlayan 'printto' fiilini deneyin.
                        if printer_name:
                            # Use the printto verb and pass the desired printer as a parameter
                            # printto fiilini kullanın ve istenen yazıcıyı bir parametre olarak geçirin
                            try:
                                win32api.ShellExecute(0, "printto", filepath, f'\"{printer_name}\"', os.path.dirname(filepath), 0)
                                # QMessageBox.information(self, "Print Sent", f"Document sent to printer: {printer_name}") # Message removed
                            except Exception as inner_err:
                                # Fallback to simple print if 'printto' fails
                                # 'printto' başarısız olursa basit yazdırmaya geri dön
                                win32api.ShellExecute(0, "print", filepath, None, os.path.dirname(filepath), 0)
                                # QMessageBox.information(self, "Print Sent", "Document sent to DEFAULT Windows printer (fallback).") # Message removed
                        else:
                            # No specific printer requested, use default 'print' verb
                            # Belirli bir yazıcı istenmedi, varsayılan 'print' fiilini kullan
                            win32api.ShellExecute(0, "print", filepath, None, os.path.dirname(filepath), 0)
                            # QMessageBox.information(self, "Print Sent", "Document sent to DEFAULT Windows printer.") # Message removed
                    except Exception as shell_error:
                        QMessageBox.warning(self, "Printing Error", f"Could not send document to printer via ShellExecute:\n{str(shell_error)}\n" # Yazdırma Hatası, Belge ShellExecute aracılığıyla yazıcıya gönderilemedi
                                                                 "Please ensure Microsoft Word (or a compatible application) is installed "
                                                                 "and set as the default program for .docx files, "
                                                                 "and your default printer is correctly configured.") # Lütfen Microsoft Word'ün (veya uyumlu bir uygulamanın) kurulu olduğundan ve .docx dosyaları için varsayılan program olarak ayarlandığından ve varsayılan yazıcınızın doğru yapılandırıldığından ve varsayılan yazıcınızın doğru yapılandırıldığından emin olun.
                else:
                    # Fallback if win32api is not imported (e.g., pywin32 not installed)
                    # win32api içe aktarılmazsa geri dönüş (örn. pywin32 kurulu değil)
                    try:
                        os.startfile(filepath, "print")
                        # QMessageBox.information(self, "Print Sent", "Document sent to DEFAULT Windows printer.") # Message removed
                    except Exception as e:
                        QMessageBox.warning(self, "Printing Error", f"Could not send document to printer:\n{str(e)}") # Yazdırma Hatası, Belge yazıcıya gönderilemedi

            else:
                # Unix-like systems (Linux/macOS) can usually handle direct printing to specified printers
                # Unix benzeri sistemler (Linux/macOS) genellikle belirtilen yazıcılara doğrudan yazdırma işlemini halledebilir.
                # using lp/lpr commands.
                # lp/lpr komutlarını kullanarak.
                print_successful = False
                command = ["lp"]
                if printer_name and printer_name != "No printer found.": # Yazıcı bulunamadı
                    command.extend(["-d", printer_name])
                command.append(filepath)

                try:
                    subprocess.run(command, check=True)
                    print_successful = True
                except FileNotFoundError:
                    # Fallback to lpr if lp is not found
                    # lp bulunamazsa lpr'ye geri dön
                    command = ["lpr"]
                    if printer_name and printer_name != "No printer found.": # Yazıcı bulunamadı
                        command.extend(["-P", printer_name]) # lpr uses -P for printer name
                                                            # lpr yazıcı adı için -P kullanır
                    command.append(filepath)
                    try:
                        subprocess.run(command, check=True)
                        print_successful = True
                    except FileNotFoundError:
                        QMessageBox.warning(self, "Printing Error", "Neither 'lp' nor 'lpr' command found on your system.") # Yazdırma Hatası, Sisteminizde 'lp' veya 'lpr' komutu bulunamadı.
                    except Exception as e:
                        QMessageBox.warning(self, "Printing Error", f"An error occurred with 'lpr' command:\n{str(e)}") # Yazdırma Hatası, 'lpr' komutuyla ilgili bir hata oluştu
                except Exception as e:
                    QMessageBox.warning(self, "Printing Error", f"An error occurred with 'lp' command:\n{str(e)}") # Yazdırma Hatası, 'lp' komutuyla ilgili bir hata oluştu

                if print_successful:
                    pass # QMessageBox.information(self, "Print Successful", "Document sent to printer!") # Message removed
                else:
                    QMessageBox.warning(self, "Printing Error", "Could not send document to printer on Unix-like system.") # Yazdırma Hatası, Unix benzeri sistemde belge yazıcıya gönderilemedi.

        except Exception as e:
            QMessageBox.warning(self, "Printing Error", f"An unexpected error occurred during printing:\n{str(e)}") # Yazdırma Hatası, Yazdırma sırasında beklenmeyen bir hata oluştu

    def load_data(self):
        """
        Loads records from the database into the table.
        Veritabanından kayıtları tabloya yükler.
        """
        # Block signals to prevent intermediate updates of the view
        self.model.blockSignals(True)
        # Clear the model
        self.model.removeRows(0, self.model.rowCount())
        
        records = self.db.get_records()
        # Set the row count in advance to avoid repeated calls to insertRow
        self.model.setRowCount(len(records))
        
        # If there are no records, unblock signals and update the label
        if not records:
            self.model.blockSignals(False)
            self.info_label.setText(f"Total Records: 0")
            self.update_selected_count_label()
            return
        
        # Populate the model
        for i, record in enumerate(records):
            # Set each column
            self.model.setItem(i, 0, QStandardItem(str(record['id'])))
            self.model.setItem(i, 1, QStandardItem(record['surname']))
            self.model.setItem(i, 2, QStandardItem(record['given_names']))
            self.model.setItem(i, 3, QStandardItem(record['sex']))
            self.model.setItem(i, 4, QStandardItem(record['place_of_birth']))
            self.model.setItem(i, 5, QStandardItem(record['date_of_birth']))
            self.model.setItem(i, 6, QStandardItem(record['date_of_issue']))
            self.model.setItem(i, 7, QStandardItem(record['date_of_expiry']))
            self.model.setItem(i, 8, QStandardItem(record['telefon']))
            self.model.setItem(i, 9, QStandardItem(record['lp_number']))

            # Set an empty QStandardItem for the PDF column. The delegate will draw the button.
            # PDF sütunu için boş bir QStandardItem ayarla. Delege düğmeyi çizecek.
            self.model.setItem(i, 10, QStandardItem("")) 
        
        # Unblock signals and notify the view of the changes
        self.model.blockSignals(False)
        # Emit dataChanged for the entire model
        self.model.layoutChanged.emit()
        
        self.info_label.setText(f"Total Records: {len(records)}")
        self.update_selected_count_label()

    def live_search(self):
        """
        Performs live filtering based on text input.
        Metin girişine göre canlı filtreleme yapar.
        """
        search_text = self.search_input.text().strip()
        if search_text:
            # Set the filter using a regular expression to search across multiple columns
            # We will search in Surname (column 1), Given Names (column 2), LP Number (column 9), Phone (column 8)
            # Soyadı (sütun 1), Adı (sütun 2), LP Numarası (sütun 9), Telefon (sütun 8)
            
            # The pattern is case-insensitive
            # Desen büyük/küçük harfe duyarsızdır
            regex = QRegExp(search_text, Qt.CaseInsensitive, QRegExp.RegExp)
            
            # Create a combined filter for the desired columns
            # İstenen sütunlar için birleştirilmiş filtre oluştur
            self.proxy_model.setFilterCaseSensitivity(Qt.CaseInsensitive)
            
            # We need to manually filter by checking each column if setFilterKeyColumn(-1) is not desired,
            # Or if we want to combine multiple specific columns.
            # Here, we will explicitly set the filter function to search in specific columns.
            # setFilterKeyColumn(-1) istenmiyorsa veya birden çok belirli sütunu birleştirmek istiyorsak,
            # her sütunu kontrol ederek manuel olarak filtrelememiz gerekir.
            # Burada, belirli sütunlarda arama yapmak için filtre fonksiyonunu açıkça ayarlayacağız.
            self.proxy_model.setFilterKeyColumn(-1) # Reset to search all columns
            
            self.proxy_model.setFilterRegExp(regex) # Apply the regex to all filterable columns
            
            # Note: For more complex "OR" logic across specific columns,
            # you might need to subclass QSortFilterProxyModel and override filterAcceptsRow.
            # But for simple regex search across all columns or the main ones, setFilterRegExp works well.

        else:
            # Clear the filter if the search input is empty
            # Arama girişi boşsa filtreyi temizle
            self.proxy_model.setFilterRegExp(QRegExp(""))
            self.proxy_model.setFilterKeyColumn(-1) # Ensure no specific column filter is left
                                                    # Belirli bir sütun filtresi kalmadığından emin olun

        # Update filtered record count
        # Filtrelenmiş kayıt sayısını güncelle
        filtered_count = self.proxy_model.rowCount()
        self.info_label.setText(f"Filtered Records: {filtered_count} (Total: {len(self.db.get_records())})") # Filtrelenmiş Kayıtlar: (Toplam: )
        self.update_selected_count_label() # Update selected count after filter

    def apply_date_filter(self):
        """
        Applies date filter to the table.
        Tabloya tarih filtresi uygular.
        """
        # Do not apply filter if user has not selected a field
        # Kullanıcı bir alan seçmediyse filtreyi uygulama
        if self.filter_field.currentIndex() == 0:
            QMessageBox.information(self, "Filter Missing", "Please select a field for date filtering first.") # Filtre Eksik, Lütfen önce tarih filtrelemesi için bir alan seçin.
            return

        filter_date = self.filter_date.date().toString("dd.MM.yyyy")
        filter_field_index = self.filter_field.currentIndex() - 1  # Because 0 is "Select Field"
                                                                    # Çünkü 0 "Alan Seç"

        # Map field indices (5: Date of Birth, 6: Date of Issue, 7: Date of Expiry)
        # Alan dizinlerini eşle (5: Doğum Tarihi, 6: Veriliş Tarihi, 7: Son Geçerlilik Tarihi)
        column_index = 5 + filter_field_index

        self.proxy_model.setFilterKeyColumn(column_index)
        self.proxy_model.setFilterFixedString(filter_date)

        # Update filtered record count
        # Filtrelenmiş kayıt sayısını güncelle
        filtered_count = self.proxy_model.rowCount()
        self.info_label.setText(f"Filtered Records: {filtered_count} (Total: {len(self.db.get_records())})") # Filtrelenmiş Kayıtlar: (Toplam: )
        self.update_selected_count_label() # Update selected count after filter

    def clear_all_filters(self):
        """
        Clears all applied filters (live search and date filter).
        Uygulanan tüm filtreleri (canlı arama ve tarih filtresi) temizler.
        """
        self.proxy_model.setFilterFixedString("") # Clear date filter
        self.search_input.clear() # Clear live search input and trigger live_search to clear its filter
        self.filter_field.setCurrentIndex(0)
        self.filter_date.setDate(QDate.currentDate()) # Reset date to current, but not _filter_min_date
        self.info_label.setText(f"Total Records: {len(self.db.get_records())}") # Toplam Kayıt:
        self.update_selected_count_label() # Update selected count after filter clear

    def select_all(self):
        """
        Selects all rows in the table.
        Tablodaki tüm satırları seçer.
        """
        self.table.selectAll()
        self.update_selected_count_label() # Update selected count after selecting all

    def sort_by_latest_records(self):
        """
        Sorts records by ID descending (newest first), then by insertion time.
        Kayıtları önce ID'ye göre azalan (en yeni), sonra eklenme zamanına göre sıralar.
        """
        # First sort by ID descending (primary)
        self.proxy_model.sort(0, Qt.DescendingOrder)
        
        # Then sort by database insertion time (secondary)
        # Get the actual record creation order from the database
        records = self.db.get_records()
        creation_order = {record['id']: idx for idx, record in enumerate(records)}
        
        # Create a custom sorting function
        def custom_sort(row):
            record_id = int(self.model.item(row, 0).text())
            return creation_order.get(record_id, 0)
        
        # Apply custom sort
        self.proxy_model.setSortRole(Qt.UserRole)
        for row in range(self.model.rowCount()):
            self.model.setData(self.model.index(row, 0), custom_sort(row), Qt.UserRole)
        self.proxy_model.sort(0, Qt.DescendingOrder)

    def sort_by_oldest_records(self):
        """
        Sorts records by ID ascending (oldest first), then by insertion time.
        Kayıtları önce ID'ye göre artan (en eski), sonra eklenme zamanına göre sıralar.
        """
        # First sort by ID ascending (primary)
        self.proxy_model.sort(0, Qt.AscendingOrder)
        
        # Then sort by database insertion time (secondary)
        # Get the actual record creation order from the database
        records = self.db.get_records()
        creation_order = {record['id']: idx for idx, record in enumerate(records)}
        
        # Create a custom sorting function
        def custom_sort(row):
            record_id = int(self.model.item(row, 0).text())
            return creation_order.get(record_id, 0)
        
        # Apply custom sort
        self.proxy_model.setSortRole(Qt.UserRole)
        for row in range(self.model.rowCount()):
            self.model.setData(self.model.index(row, 0), custom_sort(row), Qt.UserRole)
        self.proxy_model.sort(0, Qt.AscendingOrder)

    def update_selected_count_label(self):
        """
        Updates the label showing the number of selected records in the table.
        Tabloda seçilen kayıt sayısını gösteren etiketi günceller.
        """
        selected_rows = self.table.selectionModel().selectedRows()
        self.selected_info_label.setText(f"Selected Records: {len(selected_rows)}")

    def preview_selected_record(self):
        """
        Opens a dialog to preview and edit the details of the selected record from the database table.
        Veritabanı tablosundan seçilen kaydın detaylarını önizlemek ve düzenlemek için bir diyalog açar.
        """
        selected_rows = self.table.selectionModel().selectedRows()
        if not selected_rows:
            QMessageBox.warning(self, "Warning", "Please select a record to preview!") # Uyarı, Lütfen önizlemek için bir kayıt seçin!
            return

        if len(selected_rows) > 1:
            QMessageBox.warning(self, "Warning", "Please select only one record to preview!") # Uyarı, Lütfen önizlemek için yalnızca bir kayıt seçin!
            return

        proxy_row = selected_rows[0]
        source_row = self.proxy_model.mapToSource(proxy_row).row()
        record_to_edit = self.db.get_records()[source_row]

        dialog = PreviewDialog(record_to_edit, self.db, self)
        if dialog.exec_() == QDialog.Accepted:
            # Check if save or print actions were requested from the dialog
            # Diyalogdan kaydetme veya yazdırma işlemleri istenip istenmediğini kontrol et
            if dialog.save_changes_requested:
                # Data is already updated in DB via dialog's internal save methods
                # Veriler zaten diyaloğun dahili kaydetme yöntemleriyle VT'de güncellendi
                self.load_data() # Refresh main table to show updated record
                                 # Güncellenen kaydı göstermek için ana tabloyu yenile
                if dialog.save_and_print_requested:
                    # Generate and print the document using the *modified* record data
                    # *Değiştirilmiş* kayıt verilerini kullanarak belgeyi oluştur ve yazdır
                    filepath = self.generate_docx(dialog.edited_record, open_after_save=False)
                    if filepath:
                        self.print_file(filepath, self.selected_printer)
                        try:
                            time.sleep(2)
                            if os.path.exists(filepath):
                                os.remove(filepath)
                        except Exception:
                            pass
        
    def open_pdf_management_dialog(self, record_id):
        """Opens the PDF management dialog for a specific record ID."""
        """Belirli bir kayıt kimliği için PDF yönetim diyalogunu açar."""
        dialog = PdfManagementDialog(record_id, self.db, self)
        dialog.exec_()

    def open_pdf_management_dialog_from_context(self):
        """
        Opens the PDF management dialog for the selected record from the context menu.
        Bağlam menüsünden seçilen kayıt için PDF yönetim diyalogunu açar.
        """
        selected_rows = self.table.selectionModel().selectedRows()
        if not selected_rows:
            QMessageBox.warning(self, "Uyarı", "Lütfen PDF'leri yönetmek için bir kayıt seçin!") # Warning, Please select a record to manage PDFs!
            return
        if len(selected_rows) > 1:
            QMessageBox.warning(self, "Uyarı", "Lütfen PDF'leri yönetmek için yalnızca bir kayıt seçin!") # Warning, Please select only one record to manage PDFs!
            return
        
        proxy_row = selected_rows[0]
        source_row = self.proxy_model.mapToSource(proxy_row).row()
        record_id = self.db.get_records()[source_row]['id']
        self.open_pdf_management_dialog(record_id)


    def print_selected(self):
        """
        Generates document and prints directly for selected records.
        Seçilen kayıtlar için belge oluşturur ve doğrudan yazdırır.
        """
        selected_rows = self.table.selectionModel().selectedRows()
        if not selected_rows:
            QMessageBox.warning(self, "Warning", "Please select records to print!") # Uyarı, Lütfen yazdırmak için kayıt seçin!
            return

        printed_count = 0
        for proxy_row in selected_rows:
            source_row = self.proxy_model.mapToSource(proxy_row).row()
            record = self.db.get_records()[source_row]
            
            # Generate document for each record, do not open for preview
            # Her kayıt için belge oluştur, önizleme için açma
            filepath = self.generate_docx(record, open_after_save=False)

            if filepath:
                # Print document directly
                # Belgeyi doğrudan yazdır
                self.print_file(filepath, self.selected_printer)
                # Try to remove the temporary Word file
                # Geçici Word dosyasını kaldırmayı dene
                try:
                    time.sleep(2) # Increased sleep from 1 second to 5 seconds
                                # Uyku süresi 1 saniyeden 5 saniyeye çıkarıldı
                    if os.path.exists(filepath):
                        os.remove(filepath)
                except Exception:
                    pass
                printed_count += 1
            else:
                QMessageBox.warning(self, "Error", f"Document for ID {record.get('id', 'N/A')} could not be generated. Printing cancelled.") # Hata, Kimlik {record.get('id', 'N/A')} için belge oluşturulamadı. Yazdırma iptal edildi.


        if printed_count > 0:
            QMessageBox.information(self, "Completed", f"{printed_count} documents printed!") # Tamamlandı, belge yazdırıldı!
        else:
            QMessageBox.information(self, "Information", "No documents were printed.") # Bilgi, Hiçbir belge yazdırılmadı.

    def delete_selected_records(self):
        """
        Deletes selected records from the database and updates the table.
        Seçilen kayıtları veritabanından siler ve tabloyu günceller.
        """
        selected_rows = self.table.selectionModel().selectedRows()
        if not selected_rows:
            QMessageBox.warning(self, "Warning", "Please select records to delete!") # Uyarı, Lütfen silmek için kayıt seçin!
            return

        reply = QMessageBox.question(
            self, "Confirm Deletion", # Silme Onayı
            f"{len(selected_rows)} selected records will be deleted. Do you want to continue?", # seçilen kayıt silinecektir. Devam etmek istiyor musunuz?
            QMessageBox.Yes | QMessageBox.No 
        )

        if reply == QMessageBox.Yes: 
            # Collect IDs to delete
            # Silinecek kimlikleri topla
            record_ids_to_delete = []
            # Get the current records list from DB only once to avoid re-fetching after each deletion
            # Her silme işleminden sonra yeniden getirmeyi önlemek için mevcut kayıtlar listesini VT'den yalnızca bir kez al
            all_current_records = self.db.get_records()
            
            for proxy_row in selected_rows:
                source_row = self.proxy_model.mapToSource(proxy_row).row()
                # Ensure source_row is within bounds of all_current_records
                if 0 <= source_row < len(all_current_records):
                    record_id = all_current_records[source_row]['id']
                    record_ids_to_delete.append(record_id)
                else:
                    QMessageBox.warning(self, "Error", "Selected record is out of sync with database. Please refresh and try again.")
                    return # Exit to prevent further errors

            deleted_count = 0
            for record_id in record_ids_to_delete:
                # Also delete associated PDF files
                pdfs_to_delete = self.db.get_pdf_files(record_id)
                for pdf in pdfs_to_delete:
                    try:
                        if os.path.exists(pdf['filepath']):
                            os.remove(pdf['filepath'])
                        self.db.delete_pdf_file(pdf['id']) # Delete from DB after file removal
                    except Exception as e:
                        print(f"Error deleting PDF file {pdf['filepath']}: {e}") # You might want a QMessageBox here for severe errors
                
                if self.db.delete_record(record_id):
                    deleted_count += 1
            
            self.load_data() # Refresh table after all deletions
                             # Tüm silme işlemlerinden sonra tabloyu yenile
            QMessageBox.information(self, "Success", f"{deleted_count} records successfully deleted!") # Başarı, kayıt başarıyla silindi!
        else:
            QMessageBox.information(self, "Information", "Deletion cancelled.") # Bilgi, Silme işlemi iptal edildi.

    def export_to_excel_selected_or_all(self, selected_only=False):
        """
        Exports data to an Excel file. Can export selected records or all records.
        Verileri bir Excel dosyasına aktarır. Seçilen kayıtları veya tüm kayıtları aktarabilir.
        """
        records_to_export = []
        if selected_only:
            selected_rows = self.table.selectionModel().selectedRows()
            if not selected_rows:
                QMessageBox.warning(self, "Warning", "Please select records to export to Excel!") # Uyarı, Lütfen Excel'e aktarmak için kayıt seçin!
                return
            all_current_records = self.db.get_records() # Get all records once
            for proxy_row in selected_rows:
                source_row = self.proxy_model.mapToSource(proxy_row).row()
                if 0 <= source_row < len(all_current_records):
                    records_to_export.append(all_current_records[source_row])
                else:
                    QMessageBox.warning(self, "Error", "Selected record is out of sync. Please refresh and try again.")
                    return
        else:
            records_to_export = self.db.get_records()
            if not records_to_export:
                QMessageBox.warning(self, "Warning", "No records to export!") # Uyarı, Aktarılacak kayıt yok!
                return


        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export to Excel", "", "Excel Files (*.xlsx)" # Excel'e Aktar, Excel Dosyaları
        )

        if not file_path:
            return

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "LP_Records"

        # Headers
        # Başlıklar
        headers = [
            "ID", "Surname", "Given Names", "Sex", "Place of Birth", "Date of Birth",
            "Date of Issue", "Date of Expiry", "Phone", "LP №" 
        ]
        ws.append(headers)

        # Data
        # Veriler
        for record in records_to_export:
            # Retrieve the correct authorized person data from the record itself
            # Kayıttan doğru yetkili kişi verilerini al
            auth_person_info = record.get('auth_person', {})
            
            ws.append([
                record.get('id', ''),
                record.get('surname', ''),
                record.get('given_names', ''),
                record.get('sex', ''),
                record.get('place_of_birth', ''),
                record.get('date_of_birth', ''),
                record.get('date_of_issue', ''),
                record.get('date_of_expiry', ''),
                record.get('telefon', ''),
                record.get('lp_number', '')
            ])

        wb.save(file_path)
        QMessageBox.information(self, "Success", f"Data successfully exported:\n{file_path}") # Başarı, Veriler başarıyla aktarıldı:

    def import_from_excel(self):
        """
        Imports data from an Excel file into the database.
        Bir Excel dosyasından veritabanına veri aktarır.
        """
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Import from Excel", "", "Excel Files (*.xlsx)" # Excel'den İçe Aktar, Excel Dosyaları
        )

        if not file_path:
            return

        try:
            wb = openpyxl.load_workbook(file_path)
            ws = wb.active

            # Skip header row (assumption: first row is header)
            # Başlık satırını atla (varsayım: ilk satır başlık)
            for row in ws.iter_rows(min_row=2, values_only=True):
                if not any(row):  # Skip empty rows
                                # Boş satırları atla
                    continue

                # Fix date format (dd.mm.yyyy)
                # Tarih biçimini düzelt (gg.aa.yyyy)
                def fix_date(date_data):
                    if isinstance(date_data, datetime):
                        return date_data.strftime("%d.%m.%Y")
                    if isinstance(date_data, str):
                        # Try "YYYY-MM-DD" or "DD.MM.YYYY" format
                        # "YYYY-AA-GG" veya "GG.AA-YYYY" biçimini dene
                        try:
                            # Parse YYYY-MM-DD format
                            # YYYY-MM-DD biçimini ayrıştır
                            if re.match(r"\d{4}-\d{2}-\d{2}", date_data):
                                return datetime.strptime(date_data, "%Y-%m-%d").strftime("%d.%m.%Y")
                            # Return DD.MM.YYYY format directly
                            # GG.AA.YYYY biçimini doğrudan döndür
                            elif re.match(r"\d{2}\.\d{2}\.\d{4}", date_data):
                                return date_data
                        except ValueError:
                            pass # Do nothing on error, return as is
                                # Hata durumunda hiçbir şey yapma, olduğu gibi döndür
                    return str(date_data) if date_data is not None else ""


                # Ensure correct column mapping for data read from Excel
                # Excel'den okunan veriler için doğru sütun eşlemesini sağla
                # We assume the column order in the Excel file is the same as expected
                # Excel dosyasındaki sütun sırasının beklendiği gibi olduğunu varsayıyoruz
                # Get LP number from the correct column
                # LP numarasını doğru sütundan al
                record = {
                    'surname': str(row[1]) if row[1] else "", # Start from 1 as first column is ID
                                                            # İlk sütun Kimlik olduğu için 1'den başla
                    'given_names': str(row[2]) if row[2] else "",
                    'sex': str(row[3]) if row[3] else "",
                    'place_of_birth': str(row[4]) if row[4] else "TKM",
                    'date_of_birth': fix_date(row[5]),
                    'date_of_issue': fix_date(row[6]),
                    'date_of_expiry': fix_date(row[7]),
                    'telefon': str(row[8]) if row[8] else "",
                    'lp_number': str(row[9]) if row[9] else "", # Use correct column index
                                                                # Doğru sütun dizinini kullan
                    # When importing, use the current authorized person to associate with new records
                    # İçe aktarırken, yeni kayıtlarla ilişkilendirmek için mevcut yetkili kişiyi kullan
                    'auth_person': self.current_authorized_person_data if self.current_authorized_person_data else (self.db.get_users()[0] if self.db.get_users() else {})
                }

                self.db.add_record(record)

            self.load_data()
            QMessageBox.information(self, "Success", "Data successfully imported!") # Başarı, Veriler başarıyla aktarıldı!
        except Exception as e:
            QMessageBox.critical(self, "Error", f"An error occurred while reading the Excel file:\n{str(e)}") # Hata, Excel dosyası okunurken bir hata oluştu:

    def add_user(self):
        """
        Opens a dialog window to add a new authorized person.
        Yeni bir yetkili kişi eklemek için bir diyalog penceresi açar.
        """
        dialog = UserDialog()
        if dialog.exec_() == QDialog.Accepted:
            user_data = dialog.get_data()
            new_user_id = self.db.add_user(user_data)
            self.update_auth_person_combo() # Refresh the combo box and set the current person
            
            # If this is the first user, set them as the current authorized person
            # Bu ilk kullanıcıysa, onları mevcut yetkili kişi olarak ayarla
            if not self.current_authorized_person_data:
                self.current_authorized_person_data = self.db.get_users()[-1] # Get the newly added user
                self.db.save_setting('current_auth_person_id', str(self.current_authorized_person_data['id']))
                self.update_auth_person_combo() # Re-update to ensure combo box selection is correct
            QMessageBox.information(self, "Success", "Authorized person successfully added!") # Başarı, Yetkili kişi başarıyla eklendi!

    def manage_users(self):
        """
        Opens a dialog window to manage authorized persons (Add/Edit/Delete).
        Yetkili kişileri yönetmek için bir diyalog penceresi açar (Ekle/Düzenle/Sil).
        """
        users = self.db.get_users()
        if not users:
            QMessageBox.information(self, "Information", "No authorized persons registered to select.") # Bilgi, Seçmek için kayıtlı yetkili kişi yok.
            return

        user_names = [
            f"{user['id']}. {user['organization']} - {user['title']} - {user['fullname']}"
            for user in users
        ]
        
        selected_user_display, ok = QInputDialog.getItem(
            self, "Authorized Person Management", "Select an authorized person to manage:", # Yetkili Kişi Yönetimi, Yönetmek için bir yetkili kişi seçin:
            user_names, 0, False
        )

        if ok and selected_user_display:
            selected_id = int(selected_user_display.split('.')[0])
            selected_user = next((u for u in users if u['id'] == selected_id), None)

            if not selected_user:
                QMessageBox.warning(self, "Error", "Selected authorized person not found.") # Hata, Seçilen yetkili kişi bulunamadı.
                return

            msg = QMessageBox()
            msg.setWindowTitle("Authorized Person Management") # Yetkili Kişi Yönetimi
            msg.setText(f"Selected Authorized Person:\n{selected_user_display}") # Seçilen Yetkili Kişi:

            # Buttons
            # Butonlar
            edit_btn = msg.addButton("Edit", QMessageBox.ActionRole) # Düzenle
            del_btn = msg.addButton("Delete", QMessageBox.ActionRole) # Sil
            close_btn = msg.addButton("Close", QMessageBox.RejectRole) # Kapat

            msg.exec_()

            if msg.clickedButton() == edit_btn:
                self.edit_user(selected_user)
            elif msg.clickedButton() == del_btn:
                self.delete_user(selected_user)
        else:
            # If user presses "Cancel" or doesn't select an authorized person,
            # Kullanıcı "İptal"e basarsa veya yetkili kişi seçmezse,
            # offer option to add a new one or do nothing
            # yeni bir tane ekleme veya hiçbir şey yapmama seçeneği sun
            msg = QMessageBox()
            msg.setWindowTitle("Authorized Person Management") # Yetkili Kişi Yönetimi
            msg.setText("No authorized person selected.") # Yetkili kişi seçilmedi.
            add_new_btn = msg.addButton("Add New", QMessageBox.ActionRole) # Yeni Ekle
            close_btn = msg.addButton("Close", QMessageBox.RejectRole) # Kapat
            msg.exec_()
            if msg.clickedButton() == add_new_btn:
                self.add_user()

    def edit_user(self, user_to_edit):
        """
        Opens a dialog window to edit an existing authorized person.
        Mevcut bir yetkili kişiyi düzenlemek için bir diyalog penceresi açar.
        """
        dialog = UserDialog(user_to_edit)
        if dialog.exec_() == QDialog.Accepted:
            updated_data = dialog.get_data()
            if self.db.update_user(user_to_edit['id'], updated_data):
                self.update_auth_person_combo() # Refresh the combo box and set the current person
                # If the edited user was the current authorized person, update the stored data
                # Düzenlenen kullanıcı mevcut yetkili kişi ise, saklanan verileri güncelle
                if self.current_authorized_person_data and self.current_authorized_person_data['id'] == user_to_edit['id']:
                    self.current_authorized_person_data.update(updated_data)
                QMessageBox.information(self, "Success", "Authorized person successfully updated!") # Başarı, Yetkili kişi başarıyla güncellendi!
            else:
                QMessageBox.warning(self, "Error", "Authorized person could not be updated!") # Hata, Yetkili kişi güncellenemedi!

    def delete_user(self, user_to_delete):
        """
        Opens a dialog window to delete an existing authorized person.
        Mevcut bir yetkili kişiyi silmek için bir diyalog penceresi açar.
        """
        # Prevent deletion if it's the last authorized person
        # Son yetkili kişiyse silmeyi önle
        if len(self.db.get_users()) <= 1:
            QMessageBox.warning(self, "Deletion Error", "Cannot delete the last authorized person. At least one authorized person must exist.") # Silme Hatası, Son yetkili kişi silinemez. En az bir yetkili kişi bulunmalıdır.
            return

        reply = QMessageBox.question(
            self, "Confirm Deletion", # Silme Onayı
            f"Authorized person '{user_to_delete['fullname']}' will be deleted. Do you want to continue?", # Yetkili kişi '{user_to_delete['fullname']}' silinecektir. Devam etmek istiyor musunuz?
            QMessageBox.Yes | QMessageBox.No 
        )

        if reply == QMessageBox.Yes: 
            if self.db.delete_user(user_to_delete['id']):
                # If the deleted user was the current authorized person, reset to a new current
                # Silinen kullanıcı mevcut yetkili kişi ise, yeni bir mevcut kişiye sıfırla
                if self.current_authorized_person_data and self.current_authorized_person_data['id'] == user_to_delete['id']:
                    remaining_users = self.db.get_users()
                    if remaining_users:
                        self.current_authorized_person_data = remaining_users[0]
                        self.db.save_setting('current_auth_person_id', str(remaining_users[0]['id']))
                    else:
                        self.current_authorized_person_data = None
                        self.db.save_setting('current_auth_person_id', '')
                
                self.update_auth_person_combo() # Refresh the combo box and set the current person
                QMessageBox.information(self, "Success", "Authorized person successfully deleted!") # Başarı, Yetkili kişi başarıyla silindi!
            else:
                QMessageBox.warning(self, "Error", "Authorized person could not be deleted!") # Hata, Yetkili kişi silinemedi!


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())
