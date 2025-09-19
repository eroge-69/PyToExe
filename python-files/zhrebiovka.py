# -*- coding: utf-8 -*-
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn
from docx.oxml import parse_xml
import os
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from datetime import datetime
import tempfile
import shutil
import re

# === ФУНКЦИИ ДЛЯ РАБОТЫ С ДАННЫМИ ===
def load_data(file_path):
    """Загружает данные из Excel файла"""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Файл {file_path} не найден!")
    
    df = pd.read_excel(file_path)
    
    # Оставляем только нужные колонки
    required_columns = ['Фамилия и имя спортсмена', 'Пол', 'Возраст', 'Квал', 'Вес']
    for col in required_columns:
        if col not in df.columns:
            raise ValueError(f"Отсутствует обязательная колонка: {col}")
    
    # Преобразуем вес в числовой формат
    def convert_weight(weight):
        if pd.isna(weight):
            return 0.0
        try:
            if isinstance(weight, (int, float)):
                return float(weight)
            
            weight_str = str(weight).strip().replace(',', '.')
            weight_str = re.sub(r'[^\d.]', '', weight_str)
            return float(weight_str)
        except (ValueError, TypeError):
            return 0.0
    
    df['Вес'] = df['Вес'].apply(convert_weight)
    df = df[df['Вес'] > 0]
    
    return df[required_columns]

def determine_age_group(age):
    """Определяет возрастную группу для тхэквондо"""
    try:
        age = int(age)
        if 6 <= age <= 9:
            return "6-9 лет"
        elif 10 <= age <= 12:
            return "10-12 лет"
        elif 13 <= age <= 18:
            return "13-18 лет"
    except (ValueError, TypeError):
        pass
    return "Другое"

def get_qualification_value(qual):
    """Возвращает числовое значение квалификации для сравнения"""
    if pd.isna(qual) or str(qual).strip().lower() in ['без квал.', 'nan', '']:
        return 0
    
    qual_str = str(qual).lower().strip()
    
    qual_values = {
        'без квал.': 0,
        '10 гуп': 1, '9 гуп': 2, '8 гуп': 3, '7 гуп': 4,
        '6 гуп': 5, '5 гуп': 6, '4 гуп': 7, '3 гуп': 8, '2 гуп': 9, '1 гуп': 10,
        '1 дан': 11, '2 дан': 12, '3 дан': 13
    }
    
    return qual_values.get(qual_str, 999)

def get_weight_diff_limit(age_group):
    """Возвращает максимальную разницу в весе для возрастной категории"""
    if age_group == "6-9 лет":
        return 2.0
    elif age_group == "10-12 лет":
        return 3.0
    elif age_group == "13-18 лет":
        return 4.0
    return 2.0

def create_pairs(participants, age_group):
    """Создает пары спортсменов с ограничениями для возрастной категории"""
    if not participants:
        return [], []
    
    max_weight_diff = get_weight_diff_limit(age_group)
    max_qual_diff = 4
    
    # Создаем копию списка для работы
    available_participants = participants.copy()
    pairs = []
    
    # Продолжаем пока есть участники для pairing
    while len(available_participants) >= 2:
        best_pair = None
        best_score = float('inf')
        
        # Ищем наилучшую пару среди всех возможных комбинаций
        for i in range(len(available_participants)):
            for j in range(i + 1, len(available_participants)):
                p1 = available_participants[i]
                p2 = available_participants[j]
                
                weight_diff = abs(p1['Вес'] - p2['Вес'])
                qual_diff = abs(get_qualification_value(p1['Квал']) - get_qualification_value(p2['Квал']))
                
                # Проверяем ограничения
                if weight_diff <= max_weight_diff and qual_diff <= max_qual_diff:
                    # Вычисляем общий score (чем меньше - тем лучше)
                    # Даем больший вес разнице в весе (70%) и меньший квалификации (30%)
                    score = (weight_diff * 0.7) + (qual_diff * 0.3)
                    
                    if score < best_score:
                        best_score = score
                        best_pair = (i, j)
        
        # Если нашли подходящую пару
        if best_pair:
            i, j = best_pair
            # Убедимся что индексы в правильном порядке (больший индекс сначала)
            if j < i:
                i, j = j, i
            pairs.append([available_participants[i], available_participants[j]])
            # Удаляем участников из доступных (сначала больший индекс!)
            available_participants.pop(j)
            available_participants.pop(i)
        else:
            # Если не нашли больше подходящих пар, выходим
            break
    
    # Оставшиеся участники - одиночные
    single_participants = available_participants
    
    return pairs, single_participants

def format_participant_info(participant):
    """Форматирует информацию о спортсмене в две строки"""
    qual = ''
    if pd.notna(participant['Квал']):
        qual = str(participant['Квал']).strip()
        if qual.lower() in ['nan', 'без квал.', '']:
            qual = 'без квал.'
    else:
        qual = 'без квал.'
    
    # Возвращаем две строки: ФИО и данные в скобках
    return (f"{participant['Фамилия и имя спортсмена']}",
            f"({int(participant['Возраст'])} лет, {qual}, {participant['Вес']} кг)")

def set_cell_borders(cell, border_type='all'):
    """Устанавливает границы для ячейки"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is not None:
        tcPr.remove(tcBorders)
    
    if border_type == 'all':
        border_xml = '''
        <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        </w:tcBorders>
        '''
    else:
        border_xml = '''
        <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:val="nil"/>
            <w:left w:val="nil"/>
            <w:bottom w:val="nil"/>
            <w:right w:val="nil"/>
        </w:tcBorders>
        '''
    
    borders_element = parse_xml(border_xml)
    tcPr.append(borders_element)

def create_age_group_table(doc, age_group, participants):
    """Создает таблицу для одной возрастной группы"""
    if not participants:
        return
    
    # Разделяем участников по полу
    boys = [p for p in participants if p['Пол'] == 'м']
    girls = [p for p in participants if p['Пол'] == 'ж']
    
    # Создаем пары отдельно для мальчиков и девочек
    boys_pairs, boys_singles = create_pairs(boys, age_group)
    girls_pairs, girls_singles = create_pairs(girls, age_group)
    
    total_rows = len(boys_pairs) + len(boys_singles) + len(girls_pairs) + len(girls_singles)
    if total_rows == 0:
        return
    
    # Заголовок возрастной группы
    title = doc.add_paragraph(f"Возрастная категория: {age_group}")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)
    title.runs[0].font.name = 'Times New Roman'
    doc.add_paragraph()
    
    # Создаем таблицу
    table = doc.add_table(rows=total_rows + 1, cols=3)
    table.style = 'Table Grid'
    
    table.columns[0].width = Cm(6.0)
    table.columns[1].width = Cm(3.0)
    table.columns[2].width = Cm(6.0)
    
    # Заголовки таблицы
    headers = ['ФИО спортсмена', 'Результат', 'ФИО спортсмена']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
        set_cell_borders(cell, 'all')
    
    row_idx = 1
    
    # Мальчики - пары
    for pair in boys_pairs:
        cell = table.cell(row_idx, 0)
        name_line, info_line = format_participant_info(pair[0])
        cell.text = name_line + "\n" + info_line
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
        set_cell_borders(cell, 'all')
        
        cell = table.cell(row_idx, 1)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
        set_cell_borders(cell, 'all')
        
        cell = table.cell(row_idx, 2)
        name_line, info_line = format_participant_info(pair[1])
        cell.text = name_line + "\n" + info_line
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
        set_cell_borders(cell, 'all')
        row_idx += 1
    
    # Мальчики - одиночные
    for single in boys_singles:
        cell = table.cell(row_idx, 0)
        name_line, info_line = format_participant_info(single)
        cell.text = name_line + "\n" + info_line
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
        set_cell_borders(cell, 'all')
        
        cell = table.cell(row_idx, 1)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
        set_cell_borders(cell, 'all')
        
        cell = table.cell(row_idx, 2)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
        set_cell_borders(cell, 'all')
        row_idx += 1
    
    # Девочки - пары
    for pair in girls_pairs:
        cell = table.cell(row_idx, 0)
        name_line, info_line = format_participant_info(pair[0])
        cell.text = name_line + "\n" + info_line
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
        set_cell_borders(cell, 'all')
        
        cell = table.cell(row_idx, 1)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
        set_cell_borders(cell, 'all')
        
        cell = table.cell(row_idx, 2)
        name_line, info_line = format_participant_info(pair[1])
        cell.text = name_line + "\n" + info_line
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
        set_cell_borders(cell, 'all')
        row_idx += 1
    
    # Девочки - одиночные
    for single in girls_singles:
        cell = table.cell(row_idx, 0)
        name_line, info_line = format_participant_info(single)
        cell.text = name_line + "\n" + info_line
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
        set_cell_borders(cell, 'all')
        
        cell = table.cell(row_idx, 1)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
        set_cell_borders(cell, 'all')
        
        cell = table.cell(row_idx, 2)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
        set_cell_borders(cell, 'all')
        row_idx += 1
    
    doc.add_page_break()

def remove_empty_pages(doc):
    """Удаляет пустые страницы (первую и последнюю) из документа"""
    if len(doc.sections) > 0:
        # Удаляем последнюю страницу
        if len(doc.paragraphs) > 0:
            last_paragraph = doc.paragraphs[-1]
            if not last_paragraph.text.strip():
                p = last_paragraph._element
                p.getparent().remove(p)
        
        # Удаляем первую страницу если она пустая
        if len(doc.paragraphs) > 0:
            first_paragraph = doc.paragraphs[0]
            if not first_paragraph.text.strip():
                p = first_paragraph._element
                p.getparent().remove(p)

class LotteryApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Жеребьевка спаррингов по тхэквондо")
        self.root.geometry("600x400")
        self.root.resizable(False, False)
        
        self.excel_file = None
        self.template_file = None
        self.age_groups_data = {}
        
        self.setup_ui()
    
    def setup_ui(self):
        title_label = ttk.Label(self.root, text="Жеребьевка спаррингов по тхэквондо", font=("Arial", 16, "bold"))
        title_label.pack(pady=20)
        
        self.load_excel_btn = ttk.Button(self.root, text="Загрузить Excel файл с данными", command=self.load_excel_file)
        self.load_excel_btn.pack(pady=10)
        
        self.excel_status = ttk.Label(self.root, text="Excel файл: не загружен", font=("Arial", 9))
        self.excel_status.pack(pady=5)
        
        self.load_template_btn = ttk.Button(self.root, text="Загрузить шаблон Word", command=self.load_template_file)
        self.load_template_btn.pack(pady=10)
        
        self.template_status = ttk.Label(self.root, text="Шаблон Word: не загружен", font=("Arial", 9))
        self.template_status.pack(pady=5)
        
        self.process_btn = ttk.Button(self.root, text="Запустить жеребьевку", command=self.process_lottery, state='disabled')
        self.process_btn.pack(pady=15)
        
        self.status_label = ttk.Label(self.root, text="Ожидание загрузки файлов", font=("Arial", 10))
        self.status_label.pack(pady=5)
        
        self.download_btn = ttk.Button(self.root, text="Скачать результат", command=self.download_file, state='disabled')
        self.download_btn.pack(pady=10)
        
        self.result_label = ttk.Label(self.root, text="", font=("Arial", 9))
        self.result_label.pack(pady=5)
    
    def load_excel_file(self):
        file_path = filedialog.askopenfilename(title="Выберите Excel файл с данными", filetypes=[("Excel files", "*.xlsx *.xls")])
        if file_path:
            self.excel_file = file_path
            self.excel_status.config(text=f"Excel файл: {os.path.basename(file_path)}")
            self.check_files_loaded()
    
    def load_template_file(self):
        file_path = filedialog.askopenfilename(title="Выберите шаблон Word", filetypes=[("Word documents", "*.docx")])
        if file_path:
            self.template_file = file_path
            self.template_status.config(text=f"Шаблон Word: {os.path.basename(file_path)}")
            self.check_files_loaded()
    
    def check_files_loaded(self):
        if self.excel_file and self.template_file:
            self.process_btn.config(state='normal')
            self.status_label.config(text="Файлы загружены. Нажмите 'Запустить жеребьевку'")
        else:
            self.process_btn.config(state='disabled')
    
    def process_lottery(self):
        try:
            df = load_data(self.excel_file)
            if df.empty:
                messagebox.showerror("Ошибка", "Файл не содержит данных")
                return
            
            df['Возрастная группа'] = df['Возраст'].apply(determine_age_group)
            df = df[df['Возрастная группа'] != "Другое"]
            
            self.age_groups_data = {}
            for age_group in ["6-9 лет", "10-12 лет", "13-18 лет"]:
                age_group_participants = df[df['Возрастная группа'] == age_group].to_dict('records')
                if age_group_participants:
                    self.age_groups_data[age_group] = age_group_participants
            
            if not self.age_groups_data:
                messagebox.showerror("Ошибка", "Не создано ни одной возрастной группы")
                return
            
            self.download_btn.config(state='normal')
            
            total_pairs = 0
            total_singles = 0
            for age_group, participants in self.age_groups_data.items():
                boys = [p for p in participants if p['Пол'] == 'м']
                girls = [p for p in participants if p['Пол'] == 'ж']
                boys_pairs, boys_singles = create_pairs(boys, age_group)
                girls_pairs, girls_singles = create_pairs(girls, age_group)
                total_pairs += len(boys_pairs) + len(girls_pairs)
                total_singles += len(boys_singles) + len(girls_singles)
            
            self.result_label.config(text=f"Создано {len(self.age_groups_data)} возрастных групп, {total_pairs} пар, {total_singles} одиночных участников")
            self.status_label.config(text="Жеребьевка завершена. Нажмите 'Скачать результат'")
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Произошла ошибка: {str(e)}")
    
    def download_file(self):
        if not self.age_groups_data:
            messagebox.showerror("Ошибка", "Нет данных для сохранения")
            return
        
        output_path = filedialog.asksaveasfilename(
            title="Сохранить результат",
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx")],
            initialfile=f"жеребьевка_тхэквондо_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        )
        
        if output_path:
            try:
                shutil.copy2(self.template_file, output_path)
                final_doc = Document(output_path)
                style = final_doc.styles['Normal']
                style.font.name = 'Times New Roman'
                style.font.size = Pt(12)
                
                # Добавляем разрыв страницы
                final_doc.add_page_break()
                
                # Добавляем таблицы для каждой возрастной группы
                for age_group, participants in self.age_groups_data.items():
                    create_age_group_table(final_doc, age_group, participants)
                
                # Удаляем пустые страницы
                remove_empty_pages(final_doc)
                
                # Сохраняем результат
                final_doc.save(output_path)
                messagebox.showinfo("Успех", f"Файл успешно сохранен:\n{output_path}")
                
            except Exception as e:
                messagebox.showerror("Ошибка", f"Ошибка при сохранении: {str(e)}")

def main():
    root = tk.Tk()
    app = LotteryApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()