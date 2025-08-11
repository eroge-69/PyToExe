from docx import Document
import os
from docx2pdf import convert
from PyPDF2 import PdfMerger
from PyQt5.uic import loadUi
from PyQt5.QtWidgets import QApplication, QTableWidgetItem, QMessageBox

def play():
    debut=w.debut.text()
    fin=w.fin.text()
    vehicule=w.vehicule.currentText()
    if debut=="" or fin=="":
        QMessageBox.critical(w,"Erreur","Veuillez saisir début et fin")
        return False
    if not debut.isdecimal() or int(debut)<0:
        QMessageBox.critical(w,"Erreur début","Veuillez saisir un entier positif")
        return False
        
    if not fin.isdecimal() or int(fin)<=int(debut):
        QMessageBox.critical(w,"Erreur fin","Veuillez saisir un entier > debut")
        return False
    
    if (generer(vehicule, int(debut), int(fin))):
        QMessageBox.information(w,"OK","Bons de sortie générés ")

def generer(vehicule, start_num, end_num ):  

    # Dossier temporaire
    os.makedirs("temp_docs", exist_ok=True)
    os.makedirs("temp_pdfs", exist_ok=True)
    # Charger le modèle
    template_path = "Bon_de_sortie.docx"
    template = Document(template_path)
    output_docx = "Bons.docx"
    old_number = "xxxxx"  # Texte à remplacer
    old_vehicule= "VVVVV"
    for num in range(start_num, end_num + 1):
        # Dupliquer le document
        doc = Document(template_path)
        # Remplacer le numéro dans tout le texte
        for para in doc.paragraphs:
            if old_number in para.text:
                para.text = para.text.replace(old_number, f"{num}")
            elif old_vehicule in para.text:
                para.text = para.text.replace(old_vehicule, f"{vehicule}")
        doc_path = f"temp_docs/bon_{num}.docx"
        doc.save(doc_path)
        # Convertir en PDF
        pdf_path = f"temp_pdfs/bon_{num}.pdf"
        convert(doc_path, pdf_path)

    # Fusionner tous les PDF
    merger = PdfMerger()
    for num in range(start_num, end_num + 1):
        merger.append(f"temp_pdfs/bon_{num}.pdf")
    merger.write("Bons_de_sortie.pdf")
    merger.close()
    return True
# 
# 
#     
#     # Ajouter le contenu du modèle modifié au document final
#     for element in doc.element.body:
#         final_doc.element.body.append(element)
#     
#     # Saut de page sauf après le dernier
#     if num != end_num:
#         final_doc.add_page_break()
# 
#     
# # Sauvegarder avec un nom unique
# final_doc.save(output_docx)
#



app = QApplication([])
w = loadUi ("interface_bon.ui")
w.show()
w.bt.clicked.connect ( play )

app.exec_()