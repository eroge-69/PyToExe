
import wmi
import psutil
from docx import Document
import socket
from pathlib import Path

doc = Document()
doc.add_heading('System Hardware Report', 0)

c = wmi.WMI()

# Motherboard
for board in c.Win32_BaseBoard():
    doc.add_heading('Motherboard:', level=1)
    doc.add_paragraph(f"Model: {board.Product}")
    doc.add_paragraph(f"Manufacturer: {board.Manufacturer}")

# CPU
for processor in c.Win32_Processor():
    doc.add_heading('CPU:', level=1)
    doc.add_paragraph(f"Name: {processor.Name}")
    doc.add_paragraph(f"Cores: {processor.NumberOfCores}")
    doc.add_paragraph(f"Threads: {processor.NumberOfLogicalProcessors}")

# RAM
doc.add_heading('RAM:', level=1)
mem = psutil.virtual_memory()
doc.add_paragraph(f"Total: {round(mem.total / (1024**3), 2)} GB")

for ram in c.Win32_PhysicalMemory():
    mem_type = ram.MemoryType
    mem_type_str = {
        20: "DDR",
        21: "DDR2",
        24: "DDR3",
        26: "DDR4",
        30: "DDR5"
    }.get(mem_type, f"Unknown ({mem_type})")
    doc.add_paragraph(f"Type: {mem_type_str}")


# Hard Disk
doc.add_heading('Hard Disk:', level=1)
for disk in c.Win32_DiskDrive():
    doc.add_paragraph(f"Model: {disk.Model}")
    doc.add_paragraph(f"Size: {round(int(disk.Size) / (1024**3), 2)} GB")

# Monitors
doc.add_heading('Monitors:', level=1)
monitors = c.Win32_DesktopMonitor()
if monitors:
    for m in monitors:
        doc.add_paragraph(f"Name: {m.Name}")
else:
    doc.add_paragraph("No monitor detected.")

# Printers
doc.add_heading('Printers/Scanners:', level=1)
printers = c.Win32_Printer()
if printers:
    for p in printers:
        doc.add_paragraph(f"Name: {p.Name} | Default: {'Yes' if p.Default else 'No'}")
else:
    doc.add_paragraph("No printer or scanner connected.")

# Antivirus
doc.add_heading('Antivirus:', level=1)
try:
    found_av = False
    for av in c.Win32_Product():
        if "antivirus" in av.Name.lower():
            doc.add_paragraph(f"Antivirus: {av.Name}")
            found_av = True
    if not found_av:
        doc.add_paragraph("No antivirus software detected.")
except:
    doc.add_paragraph("Unable to detect antivirus software.")

# Save to Desktop
desktop = Path.home() / "Desktop"
output_file = desktop / "shenasnameh_sakht_afzari.docx"
doc.save(output_file)

print(f"Saved: {output_file}")
