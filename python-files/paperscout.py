import os
import re
import fitz  # PyMuPDF
import tkinter as tk
from tkinter import filedialog
from ttkbootstrap import Style
from ttkbootstrap import ttk
from tkinter.scrolledtext import ScrolledText
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import threading
import queue
from concurrent.futures import ThreadPoolExecutor

class PaperScoutApp:
    def __init__(self, root):
        self.root = root
        self.root.title("PaperScout - PDF Search Tool")
        self.root.geometry("1200x700")
        
        # Initialize with default theme
        self.style = Style("flatly")
        self.current_theme = "flatly"
        
        self.folder_path = tk.StringVar()
        self.search_term = tk.StringVar()
        self.case_sensitive = tk.BooleanVar()
        self.whole_word = tk.BooleanVar()
        self.results = []
        
        # Threading-related attributes
        self.search_thread = None
        self.stop_search_flag = False
        self.result_queue = queue.Queue()
        self.search_active = False
        
        # Thread pool for parallel PDF processing
        self.executor = ThreadPoolExecutor(max_workers=4)

        self.build_ui()
        self.setup_thread_monitor()

    def build_ui(self):
        # Top frame with controls
        top_frame = ttk.Frame(self.root, padding=10)
        top_frame.pack(fill="x", padx=10, pady=(10, 5))

        # Theme selector
        ttk.Label(top_frame, text="Theme:").pack(side="left", padx=(0, 5))
        self.theme_var = tk.StringVar(value=self.current_theme)
        theme_menu = ttk.OptionMenu(
            top_frame,
            self.theme_var,
            self.current_theme,
            *sorted(self.style.theme_names()),
            command=self.change_theme
        )
        theme_menu.pack(side="left", padx=(0, 20))
        self.theme_var.trace("w", lambda *_: self.change_theme(self.theme_var.get()))

        # Search controls
        ttk.Label(top_frame, text="Folder:").pack(side="left")
        ttk.Entry(top_frame, textvariable=self.folder_path, width=40).pack(side="left", padx=5)
        ttk.Button(top_frame, text="Select Folder", command=self.select_folder).pack(side="left", padx=5)

        ttk.Entry(top_frame, textvariable=self.search_term, width=30).pack(side="left", padx=5)
        ttk.Button(top_frame, text="Search", command=self.start_search).pack(side="left", padx=5)
        self.stop_button = ttk.Button(top_frame, text="Stop", command=self.stop_search, state="disabled")
        self.stop_button.pack(side="left", padx=5)

        ttk.Checkbutton(top_frame, text="Case Sensitive", variable=self.case_sensitive).pack(side="left", padx=5)
        ttk.Checkbutton(top_frame, text="Whole Word", variable=self.whole_word).pack(side="left")

        # Main content area
        self.build_content_area()

    def build_content_area(self):
        # Content frame (separated for theme changes)
        if hasattr(self, 'content_frame'):
            self.content_frame.destroy()
            
        self.content_frame = ttk.Frame(self.root)
        self.content_frame.pack(fill="both", expand=True, padx=10, pady=5)

        # Results table with scrollbars
        table_frame = ttk.Frame(self.content_frame)
        table_frame.pack(fill="both", expand=True, pady=(0, 10))
        
        tree_scroll_y = ttk.Scrollbar(table_frame)
        tree_scroll_y.pack(side="right", fill="y")
        
        tree_scroll_x = ttk.Scrollbar(table_frame, orient="horizontal")
        tree_scroll_x.pack(side="bottom", fill="x")

        self.results_table = ttk.Treeview(
            table_frame,
            columns=("Document", "Page", "Paragraph"),
            show="headings",
            yscrollcommand=tree_scroll_y.set,
            xscrollcommand=tree_scroll_x.set
        )
        self.results_table.pack(fill="both", expand=True)
        
        tree_scroll_y.config(command=self.results_table.yview)
        tree_scroll_x.config(command=self.results_table.xview)

        self.results_table.heading("Document", text="Document")
        self.results_table.heading("Page", text="Page")
        self.results_table.heading("Paragraph", text="Paragraph")
        self.results_table.column("Document", width=200, stretch=False)
        self.results_table.column("Page", width=50, stretch=False)
        self.results_table.column("Paragraph", width=600, stretch=True)
        self.results_table.bind("<<TreeviewSelect>>", self.show_paragraph)

        # Paragraph viewer
        viewer_frame = ttk.Frame(self.content_frame)
        viewer_frame.pack(fill="both", expand=True, pady=(0, 10))
        
        ttk.Label(viewer_frame, text="Selected Paragraph:").pack(anchor="w")
        self.paragraph_viewer = ScrolledText(
            viewer_frame, 
            wrap="word", 
            height=10, 
            font=("Segoe UI", 11),
            padx=5,
            pady=5
        )
        self.paragraph_viewer.pack(fill="both", expand=True)
        self.paragraph_viewer.configure(state="disabled")
        self.paragraph_viewer.tag_configure("highlight", background="yellow", foreground="black")

        # Bottom buttons
        if hasattr(self, 'bottom_frame'):
            self.bottom_frame.destroy()
            
        self.bottom_frame = ttk.Frame(self.root, padding=10)
        self.bottom_frame.pack(fill="x", padx=10, pady=(5, 10))

        ttk.Button(self.bottom_frame, text="Export to Word", command=self.export_to_word).pack(side="right", padx=5)
        ttk.Button(self.bottom_frame, text="Export to PDF", command=self.export_to_pdf).pack(side="right")

        # Status area
        if hasattr(self, 'status_frame'):
            self.status_frame.destroy()
            
        self.status_frame = ttk.Frame(self.root, padding=0)
        self.status_frame.pack(fill="x", padx=10, pady=(0, 10))

        self.progress = ttk.Progressbar(self.status_frame, mode="determinate")
        self.progress.pack(fill="x", pady=(0, 5))
        self.status_label = ttk.Label(self.status_frame, text="Ready")
        self.status_label.pack(anchor="w")

    def change_theme(self, theme_name):
        """Safe theme changing without recursion"""
        if theme_name == self.current_theme:
            return
            
        self.current_theme = theme_name
        self.style = Style(theme_name)
        
        # Store current state
        current_search = self.search_term.get()
        current_folder = self.folder_path.get()
        current_results = self.results.copy()
        current_selection = self.results_table.selection()
        current_scroll = self.paragraph_viewer.yview()
        
        # Rebuild content area only
        self.build_content_area()
        
        # Restore state
        self.search_term.set(current_search)
        self.folder_path.set(current_folder)
        self.results = current_results
        
        # Restore results if any
        for result in current_results:
            self.results_table.insert("", "end", values=result)
            
        # Restore selection and scroll
        if current_selection:
            self.results_table.selection_set(current_selection)
            self.results_table.focus(current_selection)
        self.paragraph_viewer.yview_moveto(current_scroll[0])
        
        self.status_label.config(text=f"Theme changed to {theme_name}")

    # [Keep all your existing methods unchanged]
    def select_folder(self):
        folder = filedialog.askdirectory()
        if folder:
            self.folder_path.set(folder)

    def start_search(self):
        if self.search_active:
            return
            
        folder = self.folder_path.get()
        term = self.search_term.get()
        
        if not os.path.isdir(folder) or not term:
            self.status_label.config(text="Please select a folder and enter a search term.")
            return

        # Clear previous results
        self.results_table.delete(*self.results_table.get_children())
        self.paragraph_viewer.configure(state="normal")
        self.paragraph_viewer.delete("1.0", tk.END)
        self.paragraph_viewer.configure(state="disabled")
        self.results = []
        
        # Reset progress bar
        self.progress["value"] = 0
        
        # Get list of PDF files
        pdf_files = [f for f in os.listdir(folder) if f.lower().endswith(".pdf")]
        if not pdf_files:
            self.status_label.config(text="No PDF files found in the selected folder.")
            return
            
        self.progress["maximum"] = len(pdf_files)
        
        # Set up search parameters
        flags = 0 if self.case_sensitive.get() else re.IGNORECASE
        pattern = r"\b" + re.escape(term) + r"\b" if self.whole_word.get() else re.escape(term)
        self.regex = re.compile(pattern, flags)
        self.current_search_term = term  # Store for highlighting
        
        # Update UI
        self.status_label.config(text=f"Searching in {len(pdf_files)} files...")
        self.stop_button.config(state="normal")
        self.search_active = True
        self.stop_search_flag = False
        
        # Start the search in a separate thread
        self.search_thread = threading.Thread(
            target=self.perform_search,
            args=(folder, pdf_files),
            daemon=True
        )
        self.search_thread.start()

    def perform_search(self, folder, pdf_files):
        """Perform the search across multiple PDF files using thread pool"""
        try:
            futures = []
            for file in pdf_files:
                if self.stop_search_flag:
                    break
                    
                path = os.path.join(folder, file)
                future = self.executor.submit(self.search_in_pdf, path)
                future.add_done_callback(self.update_progress)
                futures.append(future)
                
            # Wait for all tasks to complete unless stopped
            for future in futures:
                if self.stop_search_flag:
                    break
                future.result()
                
        except Exception as e:
            self.result_queue.put(("error", str(e)))
        finally:
            # Ensure progress bar reaches 100% when done
            if not self.stop_search_flag:
                self.progress["value"] = self.progress["maximum"]
            self.result_queue.put(("complete", None))

    def search_in_pdf(self, filepath):
        """Search a single PDF file for matches"""
        if self.stop_search_flag:
            return
            
        filename = os.path.basename(filepath)
        results = []
        
        try:
            doc = fitz.open(filepath)
            for page_num in range(len(doc)):
                if self.stop_search_flag:
                    break
                    
                page = doc.load_page(page_num)
                text = page.get_text()
                paragraphs = text.split("\n\n")
                
                for paragraph in paragraphs:
                    if self.regex.search(paragraph):
                        snippet = paragraph.strip().replace("\n", " ")
                        results.append((filename, page_num + 1, snippet))
            
            if results:
                self.result_queue.put(("results", results))
                
        except Exception as e:
            self.result_queue.put(("error", f"Error processing {filename}: {str(e)}"))

    def update_progress(self, future):
        """Callback to update progress bar"""
        if not self.stop_search_flag:
            # Increment progress only if not at max already
            if self.progress["value"] < self.progress["maximum"]:
                self.progress["value"] += 1
                self.root.update_idletasks()

    def stop_search(self):
        """Stop the current search operation"""
        self.stop_search_flag = True
        self.search_active = False
        self.status_label.config(text="Search stopped by user.")
        self.stop_button.config(state="disabled")

    def setup_thread_monitor(self):
        """Check the result queue periodically and update UI"""
        try:
            while True:
                try:
                    msg_type, data = self.result_queue.get_nowait()
                    
                    if msg_type == "results":
                        for result in data:
                            self.results.append(result)
                            self.results_table.insert("", "end", values=result)
                            
                    elif msg_type == "error":
                        self.status_label.config(text=f"Error: {data}")
                        
                    elif msg_type == "complete":
                        self.search_active = False
                        self.stop_button.config(state="disabled")
                        if not self.stop_search_flag:
                            self.status_label.config(
                                text=f"Search complete. {len(self.results)} result(s) found."
                            )
                            
                except queue.Empty:
                    pass
                    
                self.root.after(100, self.setup_thread_monitor)
                return
                
        except Exception as e:
            print(f"Error in thread monitor: {e}")
            self.root.after(100, self.setup_thread_monitor)

    def show_paragraph(self, event):
        selected = self.results_table.selection()
        if not selected:
            return
            
        values = self.results_table.item(selected[0])["values"]
        paragraph = values[2]

        self.paragraph_viewer.configure(state="normal")
        self.paragraph_viewer.delete("1.0", tk.END)
        
        # Insert the paragraph text
        self.paragraph_viewer.insert(tk.END, paragraph)
        
        # Highlight all occurrences of the search term
        if hasattr(self, 'current_search_term') and self.current_search_term:
            self.highlight_search_terms(self.current_search_term)
            
        self.paragraph_viewer.configure(state="disabled")

    def highlight_search_terms(self, term):
        """Highlight all occurrences of the search term in the paragraph viewer"""
        start_pos = "1.0"
        flags = 0 if self.case_sensitive.get() else "nocase"
        pattern = r"\b" + re.escape(term) + r"\b" if self.whole_word.get() else re.escape(term)
        
        while True:
            pos = self.paragraph_viewer.search(pattern, start_pos, stopindex=tk.END, regexp=True, nocase=(flags != 0))
            if not pos:
                break
            end_pos = f"{pos}+{len(term)}c"
            self.paragraph_viewer.tag_add("highlight", pos, end_pos)
            start_pos = end_pos

    def export_to_word(self):
        if not self.results:
            self.status_label.config(text="No results to export.")
            return

        doc = Document()
        doc.add_heading("Search Results", 0)
        for result in self.results:
            doc.add_paragraph(f"Document: {result[0]}")
            doc.add_paragraph(f"Page: {result[1]}")
            doc.add_paragraph(result[2])
            doc.add_paragraph("-" * 50)

        filepath = filedialog.asksaveasfilename(defaultextension=".docx")
        if filepath:
            doc.save(filepath)
            self.status_label.config(text=f"Exported to Word: {filepath}")

    def export_to_pdf(self):
        if not self.results:
            self.status_label.config(text="No results to export.")
            return

        filepath = filedialog.asksaveasfilename(defaultextension=".pdf")
        if not filepath:
            return

        doc = SimpleDocTemplate(filepath)
        styles = getSampleStyleSheet()
        flowables = [Paragraph("Search Results", styles["Heading1"])]

        for result in self.results:
            flowables.append(Paragraph(f"<b>Document:</b> {result[0]}", styles["Normal"]))
            flowables.append(Paragraph(f"<b>Page:</b> {result[1]}", styles["Normal"]))
            flowables.append(Paragraph(result[2], styles["Normal"]))
            flowables.append(Spacer(1, 12))

        doc.build(flowables)
        self.status_label.config(text=f"Exported to PDF: {filepath}")

    def __del__(self):
        """Clean up thread pool when the app is closed"""
        if hasattr(self, 'executor'):
            self.executor.shutdown(wait=False)

if __name__ == "__main__":
    root = tk.Tk()
    app = PaperScoutApp(root)
    root.mainloop()