import tkinter as tk
from tkinter import ttk, messagebox, Menu, Scrollbar, Canvas, Toplevel
import requests
from io import BytesIO
from PIL import Image, ImageTk
from docx import Document
from docx.shared import Inches, Mm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import re
import pickle
from datetime import datetime
import time
import win32com.client
import psutil

DOC_FILE = "YouTube_Summary.docx"
DATA_FILE = "video_data.pkl"
ICON_PATH = "logo.ico"

if os.path.exists(DATA_FILE):
    with open(DATA_FILE, "rb") as f:
        video_entries = pickle.load(f)
else:
    video_entries = []

def extract_video_id(url):
    pattern = r"(?:v=|\/)([0-9A-Za-z_-]{11})"
    match = re.search(pattern, url)
    return match.group(1) if match else None

def download_thumbnail_image(video_id):
    urls = [
        f"https://img.youtube.com/vi/{video_id}/maxresdefault.jpg",
        f"https://img.youtube.com/vi/{video_id}/hqdefault.jpg"
    ]
    for url in urls:
        response = requests.get(url)
        if response.status_code == 200:
            return BytesIO(response.content)
    raise Exception("Failed to download thumbnail.")

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def close_word_if_open():
    for proc in psutil.process_iter():
        try:
            if proc.name().lower() == "winword.exe":
                proc.kill()
                time.sleep(2)
                return True
        except (psutil.NoSuchProcess, psutil.AccessDenied):
            continue
    return False

def reopen_word_file():
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = True
    word.Documents.Open(os.path.abspath(DOC_FILE))

def regenerate_word(entries):
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(297)
    section.page_height = Mm(210)

    try:
        document.add_picture("logo.png", width=Inches(2))
    except:
        pass

    topic_groups = {}
    for entry in entries:
        topic_groups.setdefault(entry['topic'], []).append(entry)

    for topic in sorted(topic_groups.keys()):
        p = document.add_paragraph(topic)
        p.style = "Heading 1"
        if p.runs:
            p.runs[0].font.size = Pt(20)

        videos = topic_groups[topic]
        cols = 3
        rows = (len(videos) + cols - 1) // cols
        table = document.add_table(rows=rows, cols=cols)

        for idx, video in enumerate(videos):
            row_idx = idx // cols
            col_idx = idx % cols
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            temp_img = "temp.jpg"
            Image.open(video['thumbnail']).save(temp_img)
            run = cell.paragraphs[0].add_run()
            run.add_picture(temp_img, width=Inches(2.5))
            p_link = cell.add_paragraph()
            add_hyperlink(p_link, video['link'], "Watch Video")
            cell.add_paragraph(video['description'])
            os.remove(temp_img)

    try:
        document.save(DOC_FILE)
    except PermissionError:
        closed = close_word_if_open()
        if closed:
            time.sleep(2)
            document.save(DOC_FILE)
            reopen_word_file()
        else:
            show_messagebox("Error", "Could not close Word. Please close it manually and retry.", box_type="error")

def show_messagebox(title, message, icon=ICON_PATH, box_type="info"):
    try:
        root.iconbitmap(icon)
    except:
        pass
    if box_type == "info":
        messagebox.showinfo(title, message)
    elif box_type == "warning":
        messagebox.showwarning(title, message)
    elif box_type == "error":
        messagebox.showerror(title, message)

def on_add_topic():
    new_topic = new_topic_entry.get().strip()
    if new_topic and new_topic not in topic_cb["values"]:
        topic_cb["values"] = list(topic_cb["values"]) + [new_topic]
        topic_cb.set(new_topic)
        new_topic_entry.delete(0, tk.END)
    else:
        show_messagebox("Invalid Topic", "Enter a new unique topic.", box_type="warning")

def on_submit():
    url = url_entry.get().strip()
    topic = topic_cb.get().strip()
    description = desc_text.get("1.0", tk.END).strip()
    if not url or not topic or not description:
        show_messagebox("Missing Data", "Please fill all fields.", box_type="warning")
        return
    video_id = extract_video_id(url)
    if not video_id:
        show_messagebox("Invalid URL", "Cannot extract video ID.", box_type="error")
        return
    for entry in video_entries:
        if extract_video_id(entry["link"]) == video_id:
            show_messagebox("Duplicate Video", f"The entered video already exists in '{entry['topic']}'.")
            return
    try:
        thumb_data = download_thumbnail_image(video_id)
        video_entries.append({
            "topic": topic,
            "description": description,
            "link": url,
            "thumbnail": BytesIO(thumb_data.read()),
            "date_added": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        })
        with open(DATA_FILE, "wb") as f:
            pickle.dump(video_entries, f)
        regenerate_word(video_entries)
        show_messagebox("Success", f"Added video under topic '{topic}'.")
        url_entry.delete(0, tk.END)
        desc_text.delete("1.0", tk.END)
    except Exception as e:
        show_messagebox("Error", str(e), box_type="error")

def open_manage_videos():
    win = Toplevel(root)
    try:
        manage_win.iconbitmap(ICON_PATH)
    except:
        pass
    win.title("Manage Videos")
    win.geometry("650x500")
    canvas = Canvas(win)
    scrollbar = Scrollbar(win, orient="vertical", command=canvas.yview)
    scroll_frame = tk.Frame(canvas)
    exit_btn = tk.Button(win, text="Exit", command=win.destroy)
    exit_btn.pack(padx=10,pady=10)
    scroll_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
    canvas.create_window((0, 0), window=scroll_frame, anchor="nw")
    canvas.configure(yscrollcommand=scrollbar.set)
    canvas.pack(side="left", fill="both", expand=True)
    scrollbar.pack(side="right", fill="y")
    for idx, entry in enumerate(video_entries):
        frame = tk.Frame(scroll_frame, bd=1, relief="solid", padx=10, pady=5)
        frame.pack(pady=5, fill="x", expand=True)
        try:
            thumb_data = entry.get("thumbnail")
            if not thumb_data:
                video_id = extract_video_id(entry["link"])
                thumb_data = download_thumbnail_image(video_id)
                entry["thumbnail"] = BytesIO(thumb_data.read())  # Cache it
            with open(DATA_FILE, "wb") as f:
                pickle.dump(video_entries, f)
                img = Image.open(entry["thumbnail"])
        except Exception as e:
            img = Image.new("RGB", (120, 90), color="gray")  # Fallback image
            print(f"Failed to load thumbnail for video: {e}")

        img.thumbnail((120, 90))
        img_tk = ImageTk.PhotoImage(img)
        label = tk.Label(frame, image=img_tk)
        label.image = img_tk
        label.grid(row=0, column=0, rowspan=3)
        tk.Label(frame, text=f"Topic: {entry['topic']}").grid(row=0, column=1, sticky="w")
        tk.Label(frame, text=f"Added: {entry['date_added']}").grid(row=1, column=1, sticky="w")
        tk.Button(frame, text="Delete", command=lambda i=idx: confirm_delete_video(i, win)).grid(row=2, column=2, rowspan=1)
        tk.Button(frame, text="Edit", command=lambda i=idx: edit_video_entry(i, win)).grid(row=0, column=2, rowspan=2)
       

def confirm_delete_video(index, parent_window):
    countdown = 6

    def update_label():
        nonlocal countdown
        if countdown > 0:
            label.config(text=f"Confirm delete video? {countdown}s")
            countdown -= 1
            confirm_win.after(1000, update_label)
        else:
            confirm_btn.config(state="normal")

    def do_delete():
        delete_video(index, parent_window)
        confirm_win.destroy()

    confirm_win = Toplevel(root)
    confirm_win.title("Delete Confirmation")
    confirm_win.geometry("300x120")
    try:
        confirm_win.iconbitmap(ICON_PATH)
    except:
        pass

    label = tk.Label(confirm_win, text="Confirm delete video? 6s", font=("Arial", 11))
    label.pack(pady=15)

    confirm_btn = tk.Button(confirm_win, text="Confirm", state="disabled", command=do_delete)
    confirm_btn.pack(pady=5)

    cancel_btn = tk.Button(confirm_win, text="Cancel", command=confirm_win.destroy)
    cancel_btn.pack()

    update_label()

def edit_video_entry(index, parent_window):
    video = video_entries[index]

    def save_edit():
        new_topic = topic_entry.get().strip()
        new_desc = desc_box.get("1.0", tk.END).strip()
        if not new_topic or not new_desc:
            show_messagebox("Invalid", "Topic and Description cannot be empty.", box_type="warning")
            return
        video["topic"] = new_topic
        video["description"] = new_desc
        with open(DATA_FILE, "wb") as f:
            pickle.dump(video_entries, f)
        regenerate_word(video_entries)
        edit_win.destroy()
        parent_window.destroy()
        open_manage_videos()

    edit_win = Toplevel(root)
    edit_win.title("Edit Video Entry")
    edit_win.geometry("400x250")
    try:
        edit_win.iconbitmap(ICON_PATH)
    except:
        pass

    tk.Label(edit_win, text="Edit Topic:").pack(pady=5)
    topic_entry = ttk.Combobox(edit_win, values=sorted(set(e["topic"] for e in video_entries)), state="normal")
    topic_entry.set(video["topic"])
    topic_entry.pack(pady=5)

    tk.Label(edit_win, text="Edit Description:").pack(pady=5)
    desc_box = tk.Text(edit_win, width=40, height=5)
    desc_box.insert("1.0", video["description"])
    desc_box.pack(pady=5)

    tk.Button(edit_win, text="Save", bg="green", fg="white", command=save_edit).pack(pady=10)
    tk.Button(edit_win, text="Cancel", command=edit_win.destroy).pack()



def delete_video(index, window):
    del video_entries[index]
    with open(DATA_FILE, "wb") as f:
        pickle.dump(video_entries, f)
    regenerate_word(video_entries)
    window.destroy()
    open_manage_videos()

def edit_topics():
    win = Toplevel(root)
    try:
        edit_win.iconbitmap(ICON_PATH)
    except:
        pass
    win.title("Edit Topics")
    win.geometry("400x200")
    tk.Label(win, text="Select Topic to Rename:").pack(pady=5)
    existing = sorted(set(e['topic'] for e in video_entries))
    selected_topic = ttk.Combobox(win, values=existing, state="readonly")
    selected_topic.pack(pady=5)
    tk.Label(win, text="New Name:").pack()
    new_name_entry = tk.Entry(win)
    new_name_entry.pack(pady=5)
    def rename():
        old = selected_topic.get().strip()
        new = new_name_entry.get().strip()
        if not old or not new:
            return
        for entry in video_entries:
            if entry['topic'] == old:
                entry['topic'] = new
        with open(DATA_FILE, "wb") as f:
            pickle.dump(video_entries, f)
        regenerate_word(video_entries)
        # Updates the topic dropdown in the main UI
        all_topics = sorted(set(e['topic'] for e in video_entries))
        topic_cb["values"] = all_topics

        win.destroy()
    tk.Button(win, text="Rename", command=rename).pack(pady=10)
    tk.Button(win, text="Exit", command=win.destroy).pack(pady=10)

def Ver():
    win = Toplevel(root)
    win.title("Version")
    win.geometry("150x80")
    tk.Label(win, text="Version 1.2").pack(pady=5)
root = tk.Tk()
root.title("YouTube Video to Word Organizer")
root.geometry("500x350")
try:
    root.iconbitmap(ICON_PATH)
except:
    pass
menubar = Menu(root)
options_menu = Menu(menubar, tearoff=0)
options_menu.add_command(label="Manage Videos", command=open_manage_videos)
options_menu.add_command(label="Edit Topics", command=edit_topics)
options_menu.add_command(label="Version", command=Ver)
menubar.add_cascade(label="Options", menu=options_menu)
root.config(menu=menubar)
try:
    logo_img = ImageTk.PhotoImage(Image.open("logo.png").resize((80, 80)))
    tk.Label(root, image=logo_img).grid(row=0, column=0, columnspan=3, pady=10)
except:
    pass
tk.Label(root, text="YouTube Link:").grid(row=1, column=0, padx=10, pady=10, sticky="e")
url_entry = tk.Entry(root, width=50)
url_entry.grid(row=1, column=1, columnspan=2)
tk.Label(root, text="Topic:").grid(row=2, column=0, padx=10, pady=10, sticky="e")
topic_cb = ttk.Combobox(root, values=["Tutorial", "Education", "Vlog"], state="readonly", width=47)
topic_cb.grid(row=2, column=1, columnspan=2)
topic_cb.set("Tutorial")
tk.Label(root, text="Add Topic:").grid(row=3, column=0, padx=10, pady=5, sticky="e")
new_topic_entry = tk.Entry(root, width=30)
new_topic_entry.grid(row=3, column=1, padx=5, pady=5, sticky="w")
tk.Button(root, text="Add", command=on_add_topic).grid(row=3, column=2, padx=5, pady=5, sticky="w")
tk.Label(root, text="Description:").grid(row=4, column=0, padx=10, pady=10, sticky="ne")
desc_text = tk.Text(root, width=38, height=5)
desc_text.grid(row=4, column=1, columnspan=2, padx=10)
tk.Button(root, text="Submit to Word", bg="green", fg="white", width=25, command=on_submit).grid(row=5, column=0, columnspan=3, pady=20)
root.mainloop()
