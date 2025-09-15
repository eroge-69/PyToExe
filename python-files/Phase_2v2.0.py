# -*- coding: utf-8 -*-
"""
Phase 2 - Integrated FID Analysis Tool with Modern GUI
@author: z041763 Jawahar Bharanitharan
Version: 2.0
Date: 29/08/2025
Update: Integrated Phase 1 features with manual code and dico preparation using CustomTkinter
"""

import pandas as pd
import tkinter as tk
from tkinter import filedialog, messagebox, Checkbutton, IntVar
import os
import shutil
from docx import Document
import re
import json
import time
import win32com.client as win32
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from docxcompose.composer import Composer
import zipfile, tempfile
from lxml import etree
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
import customtkinter as ctk
import threading
from pathlib import Path
import warnings
import contextlib
import atexit
import sys
import logging
from datetime import datetime

# Suppress warnings (existing code)
warnings.filterwarnings("ignore", category=UserWarning, module="openpyxl")
warnings.filterwarnings("ignore", message="Data Validation extension is not supported")

def ignore_runtime_error(exctype, value, traceback):
    if exctype == RuntimeError and "main thread is not in main loop" in str(value):
        return
    sys.__excepthook__(exctype, value, traceback)

sys.excepthook = ignore_runtime_error


# Set CustomTkinter appearance (will be overridden by settings)
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

# Global variables for GUI
current_step = 0
total_steps = 10

# ==== Settings for style format ====
NEW_FONT = "Times New Roman"
NEW_FONT_SIZE_PT = 14
NEW_COLOR_RGB = "000000"
HEADING_LEVELS = 6
# ===================================

NAMESPACE = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

class LogManager:
    """Handles all logging functionality for the application"""
    
    def __init__(self):
        self.log_dir = "Log"
        self.session_id = time.strftime("%Y%m%d_%H%M%S")
        self.log_filename = f"{self.session_id}.log"
        self.log_filepath = os.path.join(self.log_dir, self.log_filename)
        
        # Create log directory if it doesn't exist
        self.setup_log_directory()
        
        # Setup file logger
        self.setup_logger()
        
        # Log session start
        self.log_info("=" * 60)
        self.log_info(f"Phase 2 Tool Session Started - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        self.log_info("=" * 60)
    
    def setup_log_directory(self):
        """Create log directory if it doesn't exist"""
        try:
            if not os.path.exists(self.log_dir):
                os.makedirs(self.log_dir)
                print(f"Created log directory: {self.log_dir}")
        except Exception as e:
            print(f"Error creating log directory: {e}")
    
    def setup_logger(self):
        """Setup file logger configuration"""
        try:
            # Create logger
            self.logger = logging.getLogger('Phase2Tool')
            self.logger.setLevel(logging.DEBUG)
            
            # Clear any existing handlers
            for handler in self.logger.handlers[:]:
                self.logger.removeHandler(handler)
            
            # Create file handler
            file_handler = logging.FileHandler(self.log_filepath, encoding='utf-8')
            file_handler.setLevel(logging.DEBUG)
            
            # Create formatter
            formatter = logging.Formatter(
                '[%(asctime)s] %(levelname)s - %(message)s',
                datefmt='%Y-%m-%d %H:%M:%S'
            )
            file_handler.setFormatter(formatter)
            
            # Add handler to logger
            self.logger.addHandler(file_handler)
            
        except Exception as e:
            print(f"Error setting up logger: {e}")
    
    def log_info(self, message):
        """Log info level message"""
        try:
            self.logger.info(message)
        except:
            # Fallback to direct file write
            self.write_to_file("INFO", message)
    
    def log_error(self, message):
        """Log error level message"""
        try:
            self.logger.error(message)
        except:
            self.write_to_file("ERROR", message)
    
    def log_warning(self, message):
        """Log warning level message"""
        try:
            self.logger.warning(message)
        except:
            self.write_to_file("WARNING", message)
    
    def log_debug(self, message):
        """Log debug level message"""
        try:
            self.logger.debug(message)
        except:
            self.write_to_file("DEBUG", message)
    
    def write_to_file(self, level, message):
        """Fallback method to write directly to file"""
        try:
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            with open(self.log_filepath, 'a', encoding='utf-8') as f:
                f.write(f"[{timestamp}] {level} - {message}\n")
        except Exception as e:
            print(f"Critical: Could not write to log file: {e}")
    
    def log_function_start(self, function_name):
        """Log function start"""
        self.log_info(f"Started function: {function_name}")
    
    def log_function_end(self, function_name, success=True):
        """Log function end with status"""
        status = "SUCCESS" if success else "FAILED"
        self.log_info(f"Finished function: {function_name} - {status}")
    
    def log_file_operation(self, operation, filepath, success=True):
        """Log file operations"""
        status = "SUCCESS" if success else "FAILED"
        self.log_info(f"File {operation}: {filepath} - {status}")
    
    def log_user_action(self, action):
        """Log user actions"""
        self.log_info(f"User Action: {action}")
    
    def log_configuration(self, config_data):
        """Log configuration changes"""
        self.log_info("Configuration updated:")
        for key, value in config_data.items():
            # Don't log sensitive information in full paths
            if isinstance(value, str) and len(value) > 50:
                logged_value = f"...{value[-30:]}"
            else:
                logged_value = value
            self.log_info(f"  {key}: {logged_value}")
    
    def log_session_end(self):
        """Log session end"""
        self.log_info("=" * 60)
        self.log_info(f"Phase 2 Tool Session Ended - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        self.log_info("=" * 60)
    
    def get_log_filepath(self):
        """Get current log file path"""
        return self.log_filepath
    
    def get_log_stats(self):
        """Get log file statistics"""
        try:
            if os.path.exists(self.log_filepath):
                size = os.path.getsize(self.log_filepath)
                with open(self.log_filepath, 'r', encoding='utf-8') as f:
                    lines = sum(1 for _ in f)
                return {
                    'filepath': self.log_filepath,
                    'size_bytes': size,
                    'size_kb': round(size / 1024, 2),
                    'lines': lines,
                    'session_id': self.session_id
                }
        except:
            pass
        return None


class SettingsDialog:
    """Settings dialog for configuration management and theme selection"""
    
    def __init__(self, parent, config_file="Config.json"):
        self.parent = parent
        self.config_file = config_file
        self.config_data = {}
        self.changes_made = False
        
        # Log settings dialog opening
        self.parent.log_manager.log_user_action("Opened Settings Dialog")
       
        # Load existing config
        self.load_config()
        
        # Create settings window
        self.create_settings_window()
        
    def load_config(self):
        """Load configuration from JSON file"""
        try:
            if os.path.exists(self.config_file):
                with open(self.config_file, 'r') as file:
                    self.config_data = json.load(file)
                self.parent.log_manager.log_file_operation("loaded", self.config_file, True)
            else:
                # Create default config
                self.config_data = {
                    "Model": "",
                    "REMS": "",
                    "AEMS": "", 
                    "Zid": "",
                    "AEMS_SUIVI_PATH": "",
                    "appearance_mode": "dark",
                    "color_theme": "blue"
                }
                self.save_config()
                self.parent.log_manager.log_info("Created default configuration file")
        except Exception as e:
            self.parent.log_manager.log_error(f"Error loading config: {str(e)}")
            messagebox.showerror("Config Error", f"Error loading config: {str(e)}")
            self.config_data = {}
    
    def save_config(self):
        """Save configuration to JSON file"""
        try:
            with open(self.config_file, 'w') as file:
                json.dump(self.config_data, file, indent=4)
            self.parent.log_manager.log_file_operation("saved", self.config_file, True)
            self.parent.log_manager.log_configuration(self.config_data)
            return True
        except Exception as e:
            self.parent.log_manager.log_error(f"Error saving config: {str(e)}")
            messagebox.showerror("Save Error", f"Error saving config: {str(e)}")
            return False
    
    def create_settings_window(self):
        """Create the settings dialog window"""
        self.parent.log_manager.log_info("Opened Settings")
        self.dialog = ctk.CTkToplevel(self.parent.app)
        self.dialog.title("Settings - Configuration & Themes")
        self.dialog.geometry("900x900")
        self.dialog.resizable(True, True)
        
        # Make dialog modal
        self.dialog.transient(self.parent.app)
        self.dialog.grab_set()
        
        # Center the dialog
        self.center_window()
        
        # Create main scrollable frame
        main_frame = ctk.CTkScrollableFrame(self.dialog)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Title
        title_label = ctk.CTkLabel(main_frame, text="Application Settings", 
                                 font=ctk.CTkFont(size=20, weight="bold"))
        title_label.pack(pady=(0, 20))
        
        # Theme Settings Section
        self.create_theme_section(main_frame)
        
        # Configuration Section
        self.create_config_section(main_frame)
        
        # Buttons
        self.create_buttons(main_frame)
        
        # Handle window close
        self.dialog.protocol("WM_DELETE_WINDOW", self.on_close)
        
    def center_window(self):
        """Center the dialog window"""
        self.dialog.update_idletasks()
        x = (self.dialog.winfo_screenwidth() // 2) - (600 // 2)
        y = (self.dialog.winfo_screenheight() // 2) - (700 // 2)
        self.dialog.geometry(f"700x700+{x}+{y}")
    
    def create_theme_section(self, parent):
        """Create theme selection section"""
        # Theme Section Frame
        theme_frame = ctk.CTkFrame(parent)
        theme_frame.pack(fill="x", pady=(0, 20))
        
        # Section Title
        ctk.CTkLabel(theme_frame, text="Appearance & Theme Settings", 
                    font=ctk.CTkFont(size=16, weight="bold")).pack(pady=10)
        
        # Appearance Mode
        appearance_frame = ctk.CTkFrame(theme_frame)
        appearance_frame.pack(fill="x", padx=10, pady=5)
        
        ctk.CTkLabel(appearance_frame, text="Appearance Mode:", 
                    font=ctk.CTkFont(weight="bold")).pack(anchor="w", padx=10, pady=5)
        
        # Get current appearance mode
        current_appearance = self.config_data.get("appearance_mode", "dark")
        self.appearance_var = ctk.StringVar(value=current_appearance)
        
        appearance_options = ctk.CTkFrame(appearance_frame)
        appearance_options.pack(fill="x", padx=10, pady=5)
        
        ctk.CTkRadioButton(appearance_options, text="System (Auto)", 
                          variable=self.appearance_var, value="system").pack(side="left", padx=10)
        ctk.CTkRadioButton(appearance_options, text="Light Mode", 
                          variable=self.appearance_var, value="light").pack(side="left", padx=10)
        ctk.CTkRadioButton(appearance_options, text="Dark Mode", 
                          variable=self.appearance_var, value="dark").pack(side="left", padx=10)
        
        # Color Theme
        color_frame = ctk.CTkFrame(theme_frame)
        color_frame.pack(fill="x", padx=10, pady=5)
        
        ctk.CTkLabel(color_frame, text="Color Theme:", 
                    font=ctk.CTkFont(weight="bold")).pack(anchor="w", padx=10, pady=5)
        
        # Get current color theme
        current_theme = self.config_data.get("color_theme", "blue")
        self.theme_var = ctk.StringVar(value=current_theme)
        
        color_options = ctk.CTkFrame(color_frame)
        color_options.pack(fill="x", padx=10, pady=5)
        
        ctk.CTkRadioButton(color_options, text="Blue", 
                          variable=self.theme_var, value="blue").pack(side="left", padx=10)
        ctk.CTkRadioButton(color_options, text="Dark Blue", 
                          variable=self.theme_var, value="dark-blue").pack(side="left", padx=10)
        ctk.CTkRadioButton(color_options, text="Green", 
                          variable=self.theme_var, value="green").pack(side="left", padx=10)
        
        # Preview button
        preview_btn = ctk.CTkButton(theme_frame, text="Preview Theme", 
                                   command=self.preview_theme)
        preview_btn.pack(pady=10)
        
    def create_config_section(self, parent):
        """Create configuration section"""
        # Config Section Frame
        config_frame = ctk.CTkFrame(parent)
        config_frame.pack(fill="both", expand=True)
        
        # Section Title
        ctk.CTkLabel(config_frame, text="Path Configuration", 
                    font=ctk.CTkFont(size=16, weight="bold")).pack(pady=10)
        
        # Create input fields for each config item
        self.config_entries = {}
        
        config_items = [
            ("Model", "Model Template Path", "Path to the Word document template file"),
            ("REMS", "REMS Path", "Path to REMS repair methods folder"),
            ("AEMS", "AEMS Path", "Path to AEMS repair methods folder"),
            ("AEMS_SUIVI_PATH", "AEMS SUIVI PATH", "Path to AEMS repair methods SUIVI Sheet"),
            ("Zid", "User ID (Zid)", "Your user ID for OneDrive path")
        ]
        
        for key, label, tooltip in config_items:
            self.create_config_entry(config_frame, key, label, tooltip)
    
    def create_config_entry(self, parent, key, label, tooltip):
        """Create a configuration entry field"""
        # Frame for this config item
        item_frame = ctk.CTkFrame(parent)
        item_frame.pack(fill="x", padx=10, pady=5)
        
        # Label
        label_widget = ctk.CTkLabel(item_frame, text=f"{label}:", 
                                   font=ctk.CTkFont(weight="bold"))
        label_widget.pack(anchor="w", padx=10, pady=(10, 0))
        
        # Tooltip/description
        if tooltip:
            desc_label = ctk.CTkLabel(item_frame, text=tooltip, 
                                     font=ctk.CTkFont(size=11),
                                     text_color="gray")
            desc_label.pack(anchor="w", padx=10)
        
        # Input frame
        input_frame = ctk.CTkFrame(item_frame)
        input_frame.pack(fill="x", padx=10, pady=(5, 10))
        input_frame.grid_columnconfigure(0, weight=1)
        
        # Entry field
        current_value = self.config_data.get(key, "")
        entry_var = ctk.StringVar(value=current_value)
        entry = ctk.CTkEntry(input_frame, textvariable=entry_var, 
                            placeholder_text=f"Enter {label.lower()}...")
        entry.grid(row=0, column=0, sticky="ew", padx=(0, 5))
        
        # Browse button (for path fields)
        if key != "Zid":
            browse_type = "folder" if key in ["REMS", "AEMS"] else "file"
            browse_btn = ctk.CTkButton(input_frame, text="Browse", width=80,
                                     command=lambda: self.browse_path(entry_var, browse_type, key))
            browse_btn.grid(row=0, column=1)
        
        # Store reference
        self.config_entries[key] = entry_var
    
    def browse_path(self, entry_var, browse_type, config_key):
        """Browse for file or folder path"""
        if browse_type == "folder":
            path = filedialog.askdirectory(title=f"Select {config_key} folder")
        else:
            if config_key == "Model":
                path = filedialog.askopenfilename(
                    title="Select Model template file",
                    filetypes=[("Word documents", "*.docx"), ("All files", "*.*")]
                )
            else:
                path = filedialog.askopenfilename(title=f"Select {config_key} file")
        
        if path:
            entry_var.set(path)
    
    def create_buttons(self, parent):
        """Create dialog buttons"""
        button_frame = ctk.CTkFrame(parent)
        button_frame.pack(fill="x", pady=20)
        
        # Button container
        btn_container = ctk.CTkFrame(button_frame)
        btn_container.pack()
        
        # Buttons
        ctk.CTkButton(btn_container, text="Save & Apply", 
                     command=self.save_and_apply,
                     font=ctk.CTkFont(weight="bold")).pack(side="left", padx=10)
        
        ctk.CTkButton(btn_container, text="Apply Theme Only", 
                     command=self.apply_theme_only).pack(side="left", padx=10)
        
        ctk.CTkButton(btn_container, text="Reset to Defaults", 
                     command=self.reset_defaults).pack(side="left", padx=10)
        
        ctk.CTkButton(btn_container, text="Cancel", 
                     command=self.on_close).pack(side="left", padx=10)
    
    def preview_theme(self):
        """Preview selected theme without saving"""
        appearance = self.appearance_var.get()
        color_theme = self.theme_var.get()
        
        try:
            ctk.set_appearance_mode(appearance)
            ctk.set_default_color_theme(color_theme)
            
            # Show preview message
            messagebox.showinfo("Theme Preview", 
                              f"Theme preview applied!\n"
                              f"Appearance: {appearance.title()}\n"
                              f"Color Theme: {color_theme.title()}\n\n"
                              f"Note: New widgets will show the color theme changes.")
            
        except Exception as e:
            messagebox.showerror("Preview Error", f"Error applying theme preview: {str(e)}")
    
    def apply_theme_only(self):
        """Apply theme changes without saving config paths"""
        try:
            appearance = self.appearance_var.get()
            color_theme = self.theme_var.get()
            
            # Apply theme
            ctk.set_appearance_mode(appearance)
            ctk.set_default_color_theme(color_theme)
            
            # Save only theme settings
            self.config_data["appearance_mode"] = appearance
            self.config_data["color_theme"] = color_theme
            
            if self.save_config():
                messagebox.showinfo("Theme Applied", 
                                  f"Theme settings applied and saved!\n"
                                  f"Appearance: {appearance.title()}\n"
                                  f"Color Theme: {color_theme.title()}")
                self.changes_made = True
            
        except Exception as e:
            messagebox.showerror("Error", f"Error applying theme: {str(e)}")
    
    def save_and_apply(self):
        """Save all configuration and apply theme"""
        try:
            self.parent.log_manager.log_user_action("Saving configuration and applying theme")
            # Validate paths
            if not self.validate_paths():
                return
            
            # Update config data
            for key, entry_var in self.config_entries.items():
                old_value = self.config_data.get(key, "")
                new_value = entry_var.get().strip()
                if old_value != new_value:
                    self.parent.log_manager.log_info(f"Config changed - {key}: '{old_value}' -> '{new_value}'")
                self.config_data[key] = new_value
            
            # Update theme settings
            self.config_data["appearance_mode"] = self.appearance_var.get()
            self.config_data["color_theme"] = self.theme_var.get()
            
            # Apply theme
            ctk.set_appearance_mode(self.config_data["appearance_mode"])
            ctk.set_default_color_theme(self.config_data["color_theme"])
            
            # Save config
            if self.save_config():
                messagebox.showinfo("Settings Saved", 
                                  "All settings have been saved and applied successfully!")
                self.changes_made = True
                self.parent.log_manager.log_user_action("Settings saved and applied successfully")
                self.dialog.destroy()
            
        except Exception as e:
            self.parent.log_manager.log_error(f"Error saving settings: {str(e)}")
            messagebox.showerror("Save Error", f"Error saving settings: {str(e)}")
    
    def validate_paths(self):
        """Validate entered paths"""
        errors = []
        
        # Check Model path
        model_path = self.config_entries["Model"].get().strip()
        if model_path and not os.path.exists(model_path):
            errors.append(f"Model template file not found: {model_path}")
        
        # Check REMS path  
        rems_path = self.config_entries["REMS"].get().strip()
        if rems_path and not os.path.exists(rems_path):
            errors.append(f"REMS folder not found: {rems_path}")
        
        # Check AEMS path
        aems_path = self.config_entries["AEMS"].get().strip()
        if aems_path and not os.path.exists(aems_path):
            errors.append(f"AEMS folder not found: {aems_path}")
        
        # Check Zid
        zid = self.config_entries["Zid"].get().strip()
        if not zid:
            errors.append("User ID (Zid) is required")
        
        if errors:
            messagebox.showerror("Validation Error", 
                               "Please fix the following issues:\n\n" + "\n".join(errors))
            self.parent.log_manager.log_error("Validation Error", 
                               "Please fix the following issues:\n\n" + "\n".join(errors))
            return False
        
        return True
    
    def reset_defaults(self):
        """Reset to default configuration"""
        if messagebox.askyesno("Reset Settings", 
                             "Are you sure you want to reset all settings to defaults?\n"
                             "This will clear all current configuration."):
            
            # Reset config data
            self.config_data = {
                "Model": "",
                "REMS": "",
                "AEMS": "",
                "Zid": "",
                "appearance_mode": "dark",
                "color_theme": "blue"
            }
            
            # Reset GUI elements
            for key, entry_var in self.config_entries.items():
                entry_var.set("")
            
            self.appearance_var.set("dark")
            self.theme_var.set("blue")
            
            # Apply default theme
            ctk.set_appearance_mode("dark")
            ctk.set_default_color_theme("blue")
            
            self.parent.log_manager.log_user_action("Settings have been reset to defaults.")
            messagebox.showinfo("Reset Complete", "Settings have been reset to defaults.")
    
    def on_close(self):
        """Handle dialog close"""
        self.dialog.grab_release()
        self.dialog.destroy()
        
        
class Phase2Tool:
    def __init__(self):
        try:
            # Initialize log manager first
            self.log_manager = LogManager()
            self.log_manager.log_info("Initializing Phase 2 Tool...")
            
            # Load theme settings from config before creating GUI
            self.load_theme_settings()
            
            self.app = ctk.CTk()
            self.app.title("DED DFS OneStep Tool V2.0")
            self.app.geometry("1200x800")
            self.app.minsize(1000, 700)
            self.log_manager.log_info("Main application window created")
            
            # Initialize variables after main window
            self.setup_variables()
            self.setup_gui()
            
            # Proper cleanup
            self.app.protocol("WM_DELETE_WINDOW", self.on_closing)
            atexit.register(self.cleanup)
            
        except Exception as e:
            if hasattr(self, 'log_manager'):
                self.log_manager.log_error(f"Error initializing application: {e}")
            print(f"Error initializing application: {e}")
            
    def load_theme_settings(self):
        """Load theme settings from config file"""
        try:
            if os.path.exists("Config.json"):
                with open("Config.json", 'r') as file:
                    config = json.load(file)
                
                appearance = config.get("appearance_mode", "dark")
                color_theme = config.get("color_theme", "blue")
                
                ctk.set_appearance_mode(appearance)
                ctk.set_default_color_theme(color_theme)
                self.log_manager.log_info(f"Theme loaded - Appearance: {appearance}, Color: {color_theme}")
        except Exception as e:
            # Use defaults if config loading fails
            ctk.set_appearance_mode("dark")
            ctk.set_default_color_theme("blue")  
            self.log_manager.log_warning(f"Could not load theme settings, using defaults: {e}")
            
    def log_message(self, message):
        """Add message to log area and log file"""
        timestamp = time.strftime('%H:%M:%S')
        display_message = f"[{timestamp}] {message}"
        
        # Add to GUI log
        self.log_text.insert("end", f"{display_message}\n")
        self.log_text.see("end")
        self.app.update()
        
        # Add to file log
        self.log_manager.log_info(message)
    
    def log_error_message(self, message):
        """Log error message to both GUI and file"""
        timestamp = time.strftime('%H:%M:%S')
        display_message = f"[{timestamp}] ERROR: {message}"
        
        # Add to GUI log
        self.log_text.insert("end", f"{display_message}\n")
        self.log_text.see("end")
        self.app.update()
        
        # Add to file log
        self.log_manager.log_error(message)
            
    def setup_variables(self):
        """Initialize all tkinter variables"""
        with contextlib.suppress(Exception):
            self.FID_path_var = ctk.StringVar()
            self.DFS_path_var = ctk.StringVar()
            self.PFID_path_var = ctk.StringVar()
            self.radio_var = ctk.StringVar(value="EV")
            self.activity_var = ctk.StringVar(value="THERMAL")
            self.progress_var = ctk.DoubleVar()
            self.status_var = ctk.StringVar(value="Ready")
            
    def setup_gui(self):
        # Configure appearance and window
        self.app.title("DED DFS OneStep Tool V2.0")
        self.app.geometry("1200x800")
        self.app.minsize(1000, 700)
        
        # Header
        header_frame = ctk.CTkFrame(self.app)
        header_frame.pack(fill="x", padx=10, pady=10)
        
        title_label = ctk.CTkLabel(header_frame, text="DED DFS - ONEstep tool", font=ctk.CTkFont(size=24, weight="bold"))
        title_label.pack(side="left", padx=20, pady=10)
        
        # Button frame for settings and log info
        button_frame = ctk.CTkFrame(header_frame)
        button_frame.pack(side="right", padx=20, pady=10)
        
        # Log info button
        log_info_btn = ctk.CTkButton(button_frame, text="📋 Log Info", command=self.show_log_info,
                                     width=90, height=32,
                                     font=ctk.CTkFont(size=12))
        log_info_btn.pack(side="left", padx=(0, 5))
        
        # Settings button
        settings_btn = ctk.CTkButton(button_frame, text="⚙️ Settings", command=self.open_settings,
                                     width=90, height=32,
                                     font=ctk.CTkFont(size=12))
        settings_btn.pack(side="left")
        
        # Main content frame (holds left and right)
        main_frame = ctk.CTkFrame(self.app)
        main_frame.pack(fill="both", expand=True, padx=10, pady=(0, 10))
        
        # Left panel (Scrollable)
        left_panel = ctk.CTkScrollableFrame(main_frame, width=400)
        left_panel.pack(side="left", fill="y", padx=(10, 5), pady=10)
        
        # Right panel
        right_panel = ctk.CTkFrame(main_frame)
        right_panel.pack(side="left", fill="both", expand=True, padx=(5, 10), pady=10)
        
        # Setup panel sections using packed layout
        self.setup_left_panel(left_panel)
        self.setup_right_panel(right_panel)
        
        # Status bar
        self.setup_status_bar()
    
        self.log_manager.log_info("GUI setup completed")
        
    def show_log_info(self):
        """Show information about current log file"""
        try:
            stats = self.log_manager.get_log_stats()
            if stats:
                message = (f"Current Log Session Information:\n\n"
                          f"Session ID: {stats['session_id']}\n"
                          f"Log File: {stats['filepath']}\n"
                          f"File Size: {stats['size_kb']} KB\n"
                          f"Total Lines: {stats['lines']}\n\n"
                          f"Log files are stored in the 'Log' folder.\n"
                          f"Each session creates a new log file with timestamp.")
                
                messagebox.showinfo("Log Information", message)
                self.log_manager.log_user_action("Viewed log information")
            else:
                messagebox.showwarning("Log Info", "Could not retrieve log information.")
        except Exception as e:
            self.log_manager.log_error(f"Error showing log info: {str(e)}")
            messagebox.showerror("Error", f"Error retrieving log info: {str(e)}")
        
    def open_settings(self):
        """Open settings dialog"""
        try:
            settings_dialog = SettingsDialog(self)
        except Exception as e:
            self.log_error_message(f"Could not open settings: {str(e)}")
            messagebox.showerror("Settings Error", f"Could not open settings: {str(e)}")   
            
    def setup_left_panel(self, parent):
        # Input Files Section
        input_section = ctk.CTkFrame(parent)
        input_section.pack(fill="x", padx=10, pady=10)
    
        ctk.CTkLabel(input_section, text="Input Files", font=ctk.CTkFont(size=16, weight="bold")).pack(pady=5)
    
        # FID Path
        fid_frame = ctk.CTkFrame(input_section)
        fid_frame.pack(fill="x", padx=10, pady=5)
        ctk.CTkLabel(fid_frame, text="FID File (.xlsm):").pack(anchor="w", padx=5)
    
        fid_input_frame = ctk.CTkFrame(fid_frame)
        fid_input_frame.pack(fill="x", padx=5, pady=5)
        
        self.fid_entry = ctk.CTkEntry(fid_input_frame, textvariable=self.FID_path_var, placeholder_text="Select FID file...")
        self.fid_entry.pack(side="left", fill="x", expand=True, padx=(0,5))
        ctk.CTkButton(fid_input_frame, text="Browse", command=self.browse_FID, width=80).pack(side="left")
    
        # DFS Path
        dfs_frame = ctk.CTkFrame(input_section)
        dfs_frame.pack(fill="x", padx=10, pady=5)
        ctk.CTkLabel(dfs_frame, text="DED/DFS File (.docx):").pack(anchor="w", padx=5)
    
        dfs_input_frame = ctk.CTkFrame(dfs_frame)
        dfs_input_frame.pack(fill="x", padx=5, pady=5)
    
        self.dfs_entry = ctk.CTkEntry(dfs_input_frame, textvariable=self.DFS_path_var, placeholder_text="Select DED/DFS file...")
        self.dfs_entry.pack(side="left", fill="x", expand=True, padx=(0,5))
        ctk.CTkButton(dfs_input_frame, text="Browse", command=self.browse_DFS, width=80).pack(side="left")
    
        # Previous FID Path
        pfid_frame = ctk.CTkFrame(input_section)
        pfid_frame.pack(fill="x", padx=10, pady=5)
        ctk.CTkLabel(pfid_frame, text="Previous FID (for comparison):").pack(anchor="w", padx=5)
    
        pfid_input_frame = ctk.CTkFrame(pfid_frame)
        pfid_input_frame.pack(fill="x", padx=5, pady=5)
    
        self.pfid_entry = ctk.CTkEntry(pfid_input_frame, textvariable=self.PFID_path_var, placeholder_text="Select previous FID file...")
        self.pfid_entry.pack(side="left", fill="x", expand=True, padx=(0,5))
        ctk.CTkButton(pfid_input_frame, text="Browse", command=self.browse_PFID, width=80).pack(side="left")
    
        # Configuration Section
        config_section = ctk.CTkFrame(parent)
        config_section.pack(fill="x", padx=10, pady=10)
    
        ctk.CTkLabel(config_section, text="Configuration", font=ctk.CTkFont(size=16, weight="bold")).pack(pady=5)
    
        # Engine Type
        engine_frame = ctk.CTkFrame(config_section)
        engine_frame.pack(fill="x", padx=10, pady=5)
        ctk.CTkLabel(engine_frame, text="Engine Type:", font=ctk.CTkFont(weight="bold")).pack(anchor="w", padx=5)
        radio_frame = ctk.CTkFrame(engine_frame)
        radio_frame.pack(fill="x", padx=5, pady=5)
        ctk.CTkRadioButton(radio_frame, text="EV", variable=self.radio_var, value="EV").pack(side="left", padx=10)
        ctk.CTkRadioButton(radio_frame, text="Hybrid", variable=self.radio_var, value="Hybrid").pack(side="left", padx=10)
        ctk.CTkRadioButton(radio_frame, text="ICE", variable=self.radio_var, value="ICE").pack(side="left", padx=10)
    
        # Activity Type
        activity_frame = ctk.CTkFrame(config_section)
        activity_frame.pack(fill="x", padx=10, pady=5)
        ctk.CTkLabel(activity_frame, text="Activity Type:", font=ctk.CTkFont(weight="bold")).pack(anchor="w", padx=5)
        activity_radio_frame = ctk.CTkFrame(activity_frame)
        activity_radio_frame.pack(fill="x", padx=5, pady=5)
        ctk.CTkRadioButton(activity_radio_frame, text="THERMAL", variable=self.activity_var, value="THERMAL").pack(side="left", padx=10)
        ctk.CTkRadioButton(activity_radio_frame, text="EV", variable=self.activity_var, value="EV").pack(side="left", padx=10)

    def setup_right_panel(self, parent):
        # ==== INDIVIDUAL ACTIONS SECTION ====
        actions_section = ctk.CTkFrame(parent)
        actions_section.pack(fill="x", padx=10, pady=10)
    
        ctk.CTkLabel(actions_section, text="Individual Actions", font=ctk.CTkFont(size=20, weight="bold")).pack(pady=(10, 2))
        
        # PHASE 1
        ctk.CTkLabel(actions_section, text="Phase 1 Functions", font=ctk.CTkFont(size=16, weight="bold")).pack(pady=(10,2))
        phase1_frame = ctk.CTkFrame(actions_section)
        phase1_frame.pack(fill="x", padx=10, pady=(0, 5), expand=True)
        for btn_text, btn_cmd in [
            ("Extract DTC/FT", self.extract_dtc_ft),
            ("Generate 6th Chapter", self.gentabfiche),
            ("Fetch Repair Methods", self.fetch_rm),
            ("FID Comparison", self.fid_comp),
        ]:
            ctk.CTkButton(phase1_frame, text=btn_text, command=btn_cmd).pack(
                side="left", padx=8, pady=7, expand=True, fill="x"
            )
    
        # PHASE 2
        ctk.CTkLabel(actions_section, text="Phase 2 Functions", font=ctk.CTkFont(size=16, weight="bold")).pack(pady=(12,2))
        phase2_frame = ctk.CTkFrame(actions_section)
        phase2_frame.pack(fill="x", padx=10, pady=(0, 10), expand=True)
        for btn_text, btn_cmd in [
            ("RTF to DOCX & Merge", self.RTF_merge),
            ("Prepare DICO", self.prepare_dico),
            ("Replace DICO in Word", self.replace_dico),
            ("Generate Report", self.generate_report),
        ]:
            ctk.CTkButton(phase2_frame, text=btn_text, command=btn_cmd).pack(
                side="left", padx=8, pady=7, expand=True, fill="x"
            )
    
        # ==== WORKFLOW SECTION ====
        workflow_section = ctk.CTkFrame(parent)
        workflow_section.pack(fill="both", expand=True, padx=10, pady=10)
        ctk.CTkLabel(workflow_section, text="Complete Workflow", font=ctk.CTkFont(size=18, weight="bold")).pack(pady=10)
    
        # Progress bar
        self.progress_bar = ctk.CTkProgressBar(workflow_section)
        self.progress_bar.pack(fill="x", padx=20, pady=10)
        self.progress_bar.set(0)
    
        # Progress label
        self.progress_label = ctk.CTkLabel(workflow_section, text="Ready to start workflow")
        self.progress_label.pack(pady=5)
    
        # Workflow button
        self.workflow_button = ctk.CTkButton(
            workflow_section, text="One STEP",
            command=self.One_STEP,
            height=40, font=ctk.CTkFont(size=15, weight="bold")
        )
        self.workflow_button.pack(pady=20, padx=20, fill="x")
    
        # Log area
        ctk.CTkLabel(workflow_section, text="Process Log", font=ctk.CTkFont(size=14, weight="bold")).pack(anchor="w", padx=20, pady=(20, 5))
        self.log_text = ctk.CTkTextbox(workflow_section, height=150)
        self.log_text.pack(fill="both", expand=True, padx=20, pady=(0, 20))

        
    def setup_status_bar(self):
        status_frame = ctk.CTkFrame(self.app, height=30)
        status_frame.pack(fill="x", padx=10, pady=(0, 10))
    
        self.status_label = ctk.CTkLabel(status_frame, textvariable=self.status_var)
        self.status_label.pack(side="left", padx=10, pady=5)
    
        # Log file info (right side)
        log_info = f"Log: {self.log_manager.session_id}.log"
        log_label = ctk.CTkLabel(status_frame, text=log_info,
                                 font=ctk.CTkFont(size=10),
                                 text_color="gray")
        log_label.pack(side="right", padx=10, pady=5)

    def update_progress(self, value, message=""):
        """Update progress bar and message"""
        self.progress_bar.set(value)
        if message:
            self.progress_label.configure(text=message)
        self.app.update()
        
    # File browsing functions
    def browse_FID(self):
        file_path = filedialog.askopenfilename(
            title='Select FID file',
            filetypes=[("Excel files", "*.xlsm"), ("All files", "*.*")]
        )
        if file_path:
            self.FID_path_var.set(file_path)
            self.log_message(f"FID file selected: {os.path.basename(file_path)}")
            
    def browse_DFS(self):
        file_path = filedialog.askopenfilename(
            title='Select DED/DFS file',
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")]
        )
        if file_path:
            self.DFS_path_var.set(file_path)
            self.log_message(f"DED/DFS file selected: {os.path.basename(file_path)}")
            
    def browse_PFID(self):
        file_path = filedialog.askopenfilename(
            title='Select Previous FID file',
            filetypes=[("Excel files", "*.xlsm"), ("All files", "*.*")]
        )
        if file_path:
            self.PFID_path_var.set(file_path)
            self.log_message(f"Previous FID file selected: {os.path.basename(file_path)}")
    
    # Phase 1 Functions (from original code)
    def find_cell(self, table, element):
        """Function to find index of an element in a table"""
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                if element == cell.text:
                    return (row_idx, cell_idx)
        return None
    
    def extract_dtc_ft(self):
        """Function to Extract DTC FT from DUD/SOD"""
        try:
            self.log_message("Starting DTC/FT extraction...")
            DFS = self.DFS_path_var.get()
            if not DFS:
                messagebox.showwarning("Error", "Please select DED/DFS file")
                return
                
            DTC = []
            FT = []
            doc = Document(DFS)
            
            # Process tables
            for table in doc.tables:
                try:
                    if table.cell(0, 0).text == "Diagnosis confirmation":
                        cell_pos = self.find_cell(table, "DTC")
                        if cell_pos:
                            row_idx, cell_idx = cell_pos
                            dtc = table.cell(row_idx, cell_idx + 1).text
                            fault_types = table.cell(row_idx, cell_idx + 3).text.split("\n")
                            for ft in fault_types:
                                if len(ft) > 1:
                                    FT.append(ft)
                                    DTC.append(dtc)
                except:
                    continue
                    
            # Process paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text
                if "see " in text.lower() and "DTC " in text:
                    dtc_matches = re.findall(r'[P-U]?[0-9A-F]{4}h?|[P-U]?\d{4}h?', text)
                    description = text.lower().split("see", 1)[1].strip()
                    for dtc in dtc_matches:
                        DTC.append(dtc)
                        FT.append(description)
                        
            # Save results
            df = pd.DataFrame({'DTC': DTC, 'FT': FT})
            output_path = DFS + '.xlsx'
            df.to_excel(output_path, index=False)
            
            self.log_message(f"DTC/FT extraction completed. Found {len(DTC)} entries. Saved to: {output_path}")
            messagebox.showinfo("Success", f"DTC/FT extraction completed!\nFound {len(DTC)} entries.")
            
        except Exception as e:
            self.log_message(f"Error in DTC/FT extraction: {str(e)}")
            messagebox.showerror("Error", f"Error in DTC/FT extraction: {str(e)}")
    
    def gentabfiche(self):
        """Function to create DTC Resource table from FID"""
        try:
            self.log_message("Starting 6th chapter generation...")
            FID = self.FID_path_var.get()
            if not FID:
                messagebox.showwarning("Error", "Please select FID file")
                return
                
            # Check for config file
            if not os.path.exists('Config.json'):
                messagebox.showerror("Error", "Config.json file not found. Please create configuration file.")
                return
                
            with open('Config.json', 'r') as file:
                config = json.load(file)
            
            Model_path = config.get("Model")
            if not Model_path or not os.path.exists(Model_path):
                messagebox.showerror("Error", "Model template file not found. Please check Config.json")
                return
                
            # Read FID data
            Vxx_fmc_Base = pd.read_excel(FID, sheet_name="Vxx_fmc_Base", skiprows=2)
            Vxx_fmc = Vxx_fmc_Base.loc[Vxx_fmc_Base['État cible du diagnostic'].isin(["ACTIVE", "AUTO-ACTIVE"])]
            
            Filename = os.path.basename(FID)
            New = Document(Model_path)
            DTC_list = list(set(Vxx_fmc["Code HMB"]))
            DTC_list.sort()
            
            self.log_message(f"Processing {len(DTC_list)} DTCs...")
            
            count = 50
            for i, dtc in enumerate(DTC_list):
                self.update_progress((i + 1) / len(DTC_list), f"Processing DTC {i+1}/{len(DTC_list)}: {dtc}")
                
                Vxx = Vxx_fmc.loc[Vxx_fmc['Code HMB'] == dtc]
                Desc_HMB = str(Vxx["Description HMB"].iloc[0])
                LB = list(Vxx["Code LB"])
                Desc_LB = list(Vxx["Description LB"])
                DC_Bit_4 = list(Vxx["Description des conditions d'activation"])
                DC_Bit_0 = list(Vxx["Description des seuils de détection"])
                
                Model = Document(Model_path)
                dtc = str(dtc)
                
                # Update paragraphs
                for paragraph in Model.paragraphs:
                    if '@Filename@' in paragraph.text:
                        paragraph.text = paragraph.text.replace("@Filename@", f"{Filename}_{dtc}")
                    if '@01@' in paragraph.text:
                        paragraph.text = paragraph.text.replace("@01@", dtc)
                    if '@02@' in paragraph.text:
                        paragraph.text = paragraph.text.replace("@02@", Desc_HMB)
                        
                # Update tables
                for table in Model.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if '@01@' in paragraph.text:
                                    paragraph.text = paragraph.text.replace("@01@", dtc)
                                elif '@02@' in paragraph.text:
                                    paragraph.text = paragraph.text.replace("@02@", Desc_HMB)
                                else:
                                    for j in range(len(LB)):
                                        j_str = str(j)
                                        if f"@03{j_str}@" in paragraph.text:
                                            paragraph.text = paragraph.text.replace(f"@03{j_str}@", str(LB[j]))
                                        elif f"@04{j_str}@" in paragraph.text:
                                            paragraph.text = paragraph.text.replace(f"@04{j_str}@", str(Desc_LB[j]))
                                        elif f"@05{j_str}@" in paragraph.text:
                                            paragraph.text = paragraph.text.replace(f"@05{j_str}@", str(DC_Bit_4[j]))
                                        elif f"@06{j_str}@" in paragraph.text:
                                            paragraph.text = paragraph.text.replace(f"@06{j_str}@", str(DC_Bit_0[j]))
                                        elif "@090@" in paragraph.text:
                                            FT = Vxx.loc[Vxx['Code LB'] == LB[j]]
                                            Snapshot = ""
                                            try:
                                                Snapshot = f"{FT['DID[0]'].iloc[0]} - {FT['Description'].iloc[0]}\n"
                                                for x in range(1, count):
                                                    try:
                                                        DID = str(FT[f"DID[{x}]"].iloc[0])
                                                        DID_Desc = str(FT[f"Description.{x}"].iloc[0])
                                                        if DID not in ["No specific Snapshot", "", "nan"]:
                                                            Snapshot += f"{DID} - {DID_Desc}\n"
                                                    except KeyError:
                                                        break
                                            except:
                                                Snapshot = "-"
                                            paragraph.text = paragraph.text.replace("@090@", Snapshot)
                
                # Clean up unfilled rows and cells
                for table in Model.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if '@03' in paragraph.text:
                                    try:
                                        table_element = table._tbl
                                        table_element.remove(row._tr)
                                        paragraph.text = ""
                                    except:
                                        pass
                
                # Clean unfilled cells
                for table in Model.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if '@' in paragraph.text:
                                    paragraph.text = ""
                                if 'nan' in paragraph.text:
                                    paragraph.text = paragraph.text.replace("nan", "-")
                                if "- -" in paragraph.text:
                                    paragraph.text = paragraph.text.replace("- -", "-")
                
                # Append to new document
                for element in Model.element.body:
                    New.element.body.append(element)
                    
            # Save document
            output_path = f"{FID}.docx"
            New.save(output_path)
            
            self.update_progress(1.0, "6th chapter generation completed!")
            self.log_message(f"6th chapter generated and saved to: {output_path}")
            messagebox.showinfo("Success", "6th chapter (Diagnostic resources table) created successfully!")
            
        except Exception as e:
            self.log_message(f"Error in 6th chapter generation: {str(e)}")
            messagebox.showerror("Error", f"Error in 6th chapter generation: {str(e)}")
    
    def get_user_function(self, item_list):
        """Function to create a GUI with list of functions as checkbox"""
        selected_items = []
        
        def on_confirm():
            nonlocal selected_items
            selected_items = [item for item, var in zip(items, var_list) if var.get()]
            popup.destroy()
            
        popup = ctk.CTkToplevel(self.app)
        popup.title("Select DED/DFS Function")
        popup.geometry("400x500")
        popup.grab_set()
        
        # Main frame
        main_frame = ctk.CTkFrame(popup)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        ctk.CTkLabel(main_frame, text="Select Functions", font=ctk.CTkFont(size=16, weight="bold")).pack(pady=10)
        
        # Scrollable frame for checkboxes
        scroll_frame = ctk.CTkScrollableFrame(main_frame)
        scroll_frame.pack(fill="both", expand=True, pady=10)
        
        items = [item for item in item_list if pd.notna(item)]
        var_list = []
        
        # Select all checkbox
        select_all_var = ctk.BooleanVar()
        
        def select_all():
            if select_all_var.get():
                for var in var_list:
                    var.set(True)
            else:
                for var in var_list:
                    var.set(False)
        
        select_all_cb = ctk.CTkCheckBox(scroll_frame, text="Select All", variable=select_all_var, command=select_all)
        select_all_cb.pack(anchor="w", padx=10, pady=5)
        
        # Individual checkboxes
        for item in items:
            var = ctk.BooleanVar()
            checkbox = ctk.CTkCheckBox(scroll_frame, text=str(item), variable=var)
            checkbox.pack(anchor="w", padx=20, pady=2)
            var_list.append(var)
        
        # Confirm button
        ctk.CTkButton(main_frame, text="Confirm", command=on_confirm).pack(pady=10)
        
        popup.wait_window()
        return selected_items
    
    def read_active_dtc(self, path):
        """Function to read active DTC from FID"""
        App = pd.read_excel(path, sheet_name="Applicabilité", header=8)
        App = App.iloc[:, :15]
        App = App[["Nom du défaut", "DFS", "Device", "Fault Type", "Etat Cible\n(Activé/Inhibé)"]]
        App = App.loc[App['Etat Cible\n(Activé/Inhibé)'].isin(["ACTIVE", "AUTO-ACTIVE"])]
        
        App.columns = ["ID", "FUNCTION", "DTC", "FT", "Etat Cible"]
        App["Dem_Eventid"] = App["ID"].str.replace("Vxx_fmc", "Dem_EventId")
        return App
    
    def delete_files_in_folder(self, folder_path):
        """Deletes all files in the specified folder"""
        if not os.path.exists(folder_path):
            return
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)
            try:
                if os.path.isfile(file_path):
                    os.remove(file_path)
            except Exception as e:
                self.log_message(f"Error deleting {file_path}: {e}")
    
    def copy_file(self, source_folder, destination_folder, base_filename):
        """Copy files with .rtf and .xls extensions"""
        extensions = ['.xls', '.rtf']
        os.makedirs(destination_folder, exist_ok=True)
        result = []
        
        for ext in extensions:
            base_filename = base_filename.strip()
            filename = base_filename + ext
            source_path = os.path.join(source_folder, filename)
            destination_path = os.path.join(destination_folder, filename)
            
            if os.path.isfile(source_path):
                try:
                    shutil.copy2(source_path, destination_path)
                    if ext == ".rtf":
                        result.append(f"File '{filename}' copied successfully to {destination_folder}.")
                except Exception as e:
                    if ext == ".rtf":
                        result.append(f"Error copying '{filename}': {e}")
            else:
                if ext == ".rtf":
                    result.append(f"File '{filename}' not found in source folder.({source_folder})")
        return result
    
    def fetch_rm(self):
        """Fetch Repair Methods function"""
        try:
            self.log_message("Starting repair methods fetching...")
            
            # Check configuration
            if not os.path.exists('Config.json'):
                messagebox.showerror("Error", "Config.json file not found")
                return
                
            with open('Config.json', 'r') as file:
                config = json.load(file)
                
            REMS = config.get("REMS")
            AEMS = config.get("AEMS")
            #Zid = config.get("Zid")
            Suivi = config.get("AEMS_SUIVI_PATH")
            
            if not all([REMS, AEMS, Suivi]):
                messagebox.showerror("Error", "Configuration incomplete. Check REMS, AEMS, and Suivi in Config.json")
                return
                
            #Suivi = f"C:/Users/{Zid}/OneDrive - Alliance/0-CHARGE_APV/Suivi Méthodes Diagnostic AEMS.xlsm"
            
            
            FID = self.FID_path_var.get()
            Eng_Type = self.radio_var.get()
            Activity = self.activity_var.get()
            
            if not FID:
                messagebox.showwarning("Error", "Please select FID file")
                return
                
            # Read data
            Active_DTC = self.read_active_dtc(FID)
            Master = pd.read_excel(Suivi, sheet_name="DTCs HEVC_ECM")
            
            # Process Master data
            Master = Master.iloc[:, :18]
            Master.columns = ['blank', 'Dem_EventId_XXX (AEMS)', 'Equivalent REMS (Vxx_fmc_XXX)', 'Strategie (CDC)', 'Titre DTC',
                            'Function', 'Numéro de Composant Schématique:', 'Device', 'Fault type',
                            'DTC (Pcode)', 'Ft', 'Perimeter DEA-MWA6', 'Responsabilité', "Engine",
                            'Répartition DED/DFS', "Pilotes", "DMFI/IS/ARAS0", "Nom de la Fiche"]
                            
            Master = Master.drop(['blank', 'Equivalent REMS (Vxx_fmc_XXX)', 'Strategie (CDC)',
                                'Function', 'Numéro de Composant Schématique:', 'Device', 'Fault type',
                                'Perimeter DEA-MWA6', 'Responsabilité', "Engine",
                                "Pilotes", "DMFI/IS/ARAS0"], axis=1)
                                
            Master.columns = ["Dem_Eventid", "Title DTC", "DTC (Pcode)", "FT(Pcode)", "Function(Master)", "Fiche"]
            
            # Merge data
            Master = pd.merge(Active_DTC, Master, on='Dem_Eventid', how='left')
            
            funct_nan = Master[Master["Function(Master)"].isnull()]
            funct = self.get_user_function(Master["Function(Master)"].unique())
            
            if not funct:
                self.log_message("No functions selected")
                return
                
            filtered_data = Master[Master["Function(Master)"].isin(funct)]
            filtered_data = filtered_data.reset_index(drop=True)
            fiche_nan = filtered_data[filtered_data["Fiche"].isnull()]
            
            # Process fiche names based on engine type
            for i in range(len(filtered_data["Fiche"])):
                if len(str(filtered_data["Fiche"].iloc[i])) > 50:
                    lst = str(filtered_data["Fiche"].iloc[i]).replace(",", "\n").split("\n")
                    lst = [s.strip() for s in lst]
                    fichename = ""
                    
                    # Fiche priority logic
                    priorities = {
                        "Hybrid": ["h", "e", "x"],
                        "EV": ["e", "h", "x"],
                        "ICE": ["g", "x", "d"]
                    }
                    
                    engine_priorities = priorities.get(Eng_Type, ["x", "h", "e"])
                    
                    for suffix in engine_priorities:
                        for j in lst:
                            if j.endswith(suffix):
                                fichename = j
                                break
                        if fichename:
                            break
                            
                    if not fichename and lst:
                        fichename = lst[0][:51] if len(lst[0]) > 50 else "No relevant fiche found"
                        
                    filtered_data.loc[i, "Fiche"] = fichename
            
            # Store inputs
            Fun = '-'.join(str(x) for x in funct)
            input_data = [['FID: ', FID], ['Engine Type: ', Eng_Type], ['Function: ', Fun]]
            Inputs = pd.DataFrame(input_data, columns=['Inputs', 'File'])
            
            # Clear and prepare Fiches folder
            folder_path = "Fiches"
            self.delete_files_in_folder(folder_path)
            os.makedirs(folder_path, exist_ok=True)
            
            df = filtered_data
            Folders = list(set(list(df["Function(Master)"])))
            
            # Process each function folder
            file_status_data = []
            
            
            if Activity.upper() == "THERMAL":
                # All files go into Fiches/DUD_HECM
                dud_folder = os.path.join(folder_path, "DUD_HECM")
                # Clear and recreate
                self.delete_files_in_folder(dud_folder)
                if os.path.exists(dud_folder):
                    try:
                        os.rmdir(dud_folder)
                    except Exception:
                        pass
                os.makedirs(dud_folder, exist_ok=True)
    
                for j in Folders:
                    self.log_message(f"Processing function: {j}")
                    filtered_df = df[df["Function(Master)"] == j]
                    fiches = list(set(list(filtered_df["Fiche"])))
                    for i in fiches:
                        if i == "No relevant fiche found" or str(i) == "nan":
                            file_status_data.append({'Function': j, 'Fiche': i,
                                                     'Status': i, 'Path': i})
                            continue
                        fiche_path = ""
                        split_up = i.split("_")
                        if len(split_up) > 1:
                            if len(split_up[1]) == 2:  # REMS
                                path = f"{split_up[1]}/{split_up[1]}_{split_up[2]}/{split_up[1]}_{split_up[2]}_{split_up[3]}"
                                fiche_path = os.path.join(REMS, path.upper())
                            elif len(split_up[1]) == 4:  # AEMS
                                path = f"{split_up[1][:2]}/{split_up[1]}/{split_up[1]}_{split_up[2]}"
                                fiche_path = os.path.join(AEMS, path.upper())
                        # All files copied to dud_folder
                        status = self.copy_file(fiche_path, dud_folder, i)
                        for S in status:
                            file_status_data.append({'Function': j, 'Fiche': i, 'Status': S, 'Path': fiche_path})
    
            else:  # Activity is EV or anything else - original logic
                for j in Folders:
                    self.log_message(f"Processing function: {j}")
                    function_folder = os.path.join(folder_path, str(j))
                    if os.path.exists(function_folder):
                        self.delete_files_in_folder(function_folder)
                        os.rmdir(function_folder)
                    os.makedirs(function_folder, exist_ok=True)
                    filtered_df = df[df["Function(Master)"] == j]
                    fiches = list(set(list(filtered_df["Fiche"])))
                    for i in fiches:
                        if i == "No relevant fiche found" or str(i) == "nan":
                            file_status_data.append({'Function': j, 'Fiche': i,
                                                     'Status': i, 'Path': i})
                            continue
                        fiche_path = ""
                        split_up = i.split("_")
                        if len(split_up) > 1:
                            if len(split_up[1]) == 2:  # REMS
                                path = f"{split_up[1]}/{split_up[1]}_{split_up[2]}/{split_up[1]}_{split_up[2]}_{split_up[3]}"
                                fiche_path = os.path.join(REMS, path.upper())
                            elif len(split_up[1]) == 4:  # AEMS
                                path = f"{split_up[1][:2]}/{split_up[1]}/{split_up[1]}_{split_up[2]}"
                                fiche_path = os.path.join(AEMS, path.upper())
                        # Copy to function-specific folder as before
                        status = self.copy_file(fiche_path, function_folder, i)
                        for S in status:
                            file_status_data.append({'Function': j, 'Fiche': i, 'Status': S, 'Path': fiche_path})
                
            file_status = pd.DataFrame(file_status_data)
            
            # Save all data to Excel
            output_path = f"{FID}_RM_Fetch.xlsx"
            with pd.ExcelWriter(output_path) as writer:
                Inputs.to_excel(writer, sheet_name='Inputs', index=False)
                Active_DTC.to_excel(writer, sheet_name='Current_FID', index=False)
                Master.to_excel(writer, sheet_name='Master', index=False)
                filtered_data.to_excel(writer, sheet_name='DED-DFS', index=False)
                funct_nan.to_excel(writer, sheet_name='Function_missing', index=False)
                fiche_nan.to_excel(writer, sheet_name='Fiche_missing', index=False)
                file_status.to_excel(writer, sheet_name='Fetching_Status', index=False)
            
            self.log_message(f"Repair methods fetching completed. Results saved to: {output_path}")
            messagebox.showinfo("Success", "Repair methods fetching completed successfully!")
            
        except Exception as e:
            self.log_message(f"Error in repair methods fetching: {str(e)}")
            messagebox.showerror("Error", f"Error in repair methods fetching: {str(e)}")
    
    def fid_comp(self):
        """FID Comparison function"""
        try:
            self.log_message("Starting FID comparison...")
            
            FID = self.FID_path_var.get()
            pFID = self.PFID_path_var.get()
            
            if not FID or not pFID:
                messagebox.showwarning("Error", "Please select both current and previous FID files")
                return
                
            vxx_id_new = self.read_active_dtc(FID)
            vxx_id_old = self.read_active_dtc(pFID)
            
            # Find new and removed DTCs
            new_dtc = vxx_id_new[~vxx_id_new.isin(vxx_id_old.to_dict(orient='list')).all(axis=1)]
            inhibe_dtc = vxx_id_old[~vxx_id_old.isin(vxx_id_new.to_dict(orient='list')).all(axis=1)]
            
            # Save comparison results
            output_path = f"{FID}_Analysis.xlsx"
            with pd.ExcelWriter(output_path) as writer:
                new_dtc.to_excel(writer, sheet_name='New DTC', index=False)
                inhibe_dtc.to_excel(writer, sheet_name='Inhibe DTC', index=False)
            
            self.log_message(f"FID comparison completed. Found {len(new_dtc)} new DTCs and {len(inhibe_dtc)} removed DTCs.")
            self.log_message(f"Results saved to: {output_path}")
            messagebox.showinfo("Success", f"FID comparison completed!\nNew DTCs: {len(new_dtc)}\nRemoved DTCs: {len(inhibe_dtc)}")
            
        except Exception as e:
            self.log_message(f"Error in FID comparison: {str(e)}")
            messagebox.showerror("Error", f"Error in FID comparison: {str(e)}")
    
    # Phase 2 Functions
    def update_heading_styles(self, styles_xml_path):
        """Update heading styles in Word document"""
        tree = etree.parse(styles_xml_path)
        root = tree.getroot()
        changed = 0
        
        for style in root.findall("w:style", NAMESPACE):
            style_id = style.get("{%s}styleId" % NAMESPACE["w"])
            if not style_id or not style_id.startswith("Heading"):
                continue
                
            # Find or create w:rPr (run properties)
            rpr = style.find("w:rPr", NAMESPACE)
            if rpr is None:
                rpr = etree.SubElement(style, "{%s}rPr" % NAMESPACE["w"])
            
            # Set font name
            rfonts = rpr.find("w:rFonts", NAMESPACE)
            if rfonts is None:
                rfonts = etree.SubElement(rpr, "{%s}rFonts" % NAMESPACE["w"])
            rfonts.set("{%s}ascii" % NAMESPACE["w"], NEW_FONT)
            rfonts.set("{%s}hAnsi" % NAMESPACE["w"], NEW_FONT)
            rfonts.set("{%s}cs" % NAMESPACE["w"], NEW_FONT)
            
            # Set font size
            size_elem = rpr.find("w:sz", NAMESPACE)
            if size_elem is None:
                size_elem = etree.SubElement(rpr, "{%s}sz" % NAMESPACE["w"])
            size_elem.set("{%s}val" % NAMESPACE["w"], str(NEW_FONT_SIZE_PT * 2))
            
            # Set color
            color_elem = rpr.find("w:color", NAMESPACE)
            if color_elem is None:
                color_elem = etree.SubElement(rpr, "{%s}color" % NAMESPACE["w"])
            color_elem.set("{%s}val" % NAMESPACE["w"], NEW_COLOR_RGB)
            
            changed += 1
            
        tree.write(styles_xml_path, xml_declaration=True, encoding="utf-8")
        return changed
    
    def get_heading_styles(self, styles_xml_path):
        """Get heading styles from document"""
        import re
        tree = etree.parse(styles_xml_path)
        heading_styles = []
        
        for style in tree.findall('w:style', NAMESPACE):
            style_id = style.get('{%s}styleId' % NAMESPACE['w'])
            if style_id and re.fullmatch(r"Heading[1-9]\d*", style_id):
                heading_styles.append(style_id)
                
        heading_styles.sort(key=lambda s: int(s.replace("Heading", "")))
        return heading_styles
    
    def set_headings_numbering(self, numbering_xml_path, heading_styles):
        """Set unified numbering for headings"""
        tree = etree.parse(numbering_xml_path)
        root = tree.getroot()
        
        # Remove existing abstractNum with heading styles
        for absnum in root.findall('w:abstractNum', NAMESPACE):
            for lvl in absnum.findall('w:lvl', NAMESPACE):
                pstyle = lvl.find('w:pStyle', NAMESPACE)
                if pstyle is not None and pstyle.get('{%s}val' % NAMESPACE['w']).startswith("Heading"):
                    root.remove(absnum)
                    break
        
        # Create new abstractNum
        absnum = etree.SubElement(root, '{%s}abstractNum' % NAMESPACE['w'])
        absnum.set('{%s}abstractNumId' % NAMESPACE['w'], '100')
        mlt = etree.SubElement(absnum, '{%s}multiLevelType' % NAMESPACE['w'])
        mlt.set('{%s}val' % NAMESPACE['w'], 'multilevel')
        
        for i, style in enumerate(heading_styles):
            lvl = etree.SubElement(absnum, '{%s}lvl' % NAMESPACE['w'])
            lvl.set('{%s}ilvl' % NAMESPACE['w'], str(i))
            etree.SubElement(lvl, '{%s}start' % NAMESPACE['w']).set('{%s}val' % NAMESPACE['w'], '1')
            etree.SubElement(lvl, '{%s}numFmt' % NAMESPACE['w']).set('{%s}val' % NAMESPACE['w'], 'decimal')
            dots = ".".join([f"%{j+1}" for j in range(i+1)])
            etree.SubElement(lvl, '{%s}lvlText' % NAMESPACE['w']).set('{%s}val' % NAMESPACE['w'], dots)
            etree.SubElement(lvl, '{%s}lvlJc' % NAMESPACE['w']).set('{%s}val' % NAMESPACE['w'], 'left')
            pstyle = etree.SubElement(lvl, '{%s}pStyle' % NAMESPACE['w'])
            pstyle.set('{%s}val' % NAMESPACE['w'], style)
            ind = etree.SubElement(lvl, '{%s}ind' % NAMESPACE['w'])
            ind.set('{%s}left' % NAMESPACE['w'], str(i*360))
            ind.set('{%s}hanging' % NAMESPACE['w'], '0')
        
        # Add num instance
        num = etree.SubElement(root, '{%s}num' % NAMESPACE['w'])
        num.set('{%s}numId' % NAMESPACE['w'], '200')
        absnumid = etree.SubElement(num, '{%s}abstractNumId' % NAMESPACE['w'])
        absnumid.set('{%s}val' % NAMESPACE['w'], '100')
        
        tree.write(numbering_xml_path, xml_declaration=True, encoding='utf-8')
    
    def rebase_heading_numbering(self, docx_dir, heading_styles):
        """Rebase heading numbering in document"""
        docxml = os.path.join(docx_dir, "word", "document.xml")
        doc_tree = etree.parse(docxml)
        doc_root = doc_tree.getroot()
        
        for p in doc_root.findall('.//w:p', NAMESPACE):
            ppr = p.find('w:pPr', NAMESPACE)
            if ppr is not None:
                pstyle = ppr.find('w:pStyle', NAMESPACE)
                if pstyle is not None and pstyle.get('{%s}val' % NAMESPACE["w"], "") in heading_styles:
                    # Remove old numPr
                    numpr = ppr.find('w:numPr', NAMESPACE)
                    if numpr is not None:
                        ppr.remove(numpr)
                        
                    # Create new numPr
                    this_style = pstyle.get('{%s}val' % NAMESPACE["w"])
                    ilvl_level = str(int(this_style.replace('Heading', '')) - 1)
                    numpr = etree.Element('{%s}numPr' % NAMESPACE["w"])
                    ilvl = etree.Element('{%s}ilvl' % NAMESPACE["w"])
                    ilvl.set('{%s}val' % NAMESPACE["w"], ilvl_level)
                    numid = etree.Element('{%s}numId' % NAMESPACE["w"])
                    numid.set('{%s}val' % NAMESPACE["w"], '200')
                    numpr.append(ilvl)
                    numpr.append(numid)
                    ppr.append(numpr)
        
        doc_tree.write(docxml, xml_declaration=True, encoding='utf-8')
    
    def patch_docx(self, fname, outname):
        """Apply styling patches to DOCX file"""
        with tempfile.TemporaryDirectory() as tmpdir:
            # Unzip
            with zipfile.ZipFile(fname, 'r') as zip_ref:
                zip_ref.extractall(tmpdir)
            
            styles_xml = os.path.join(tmpdir, 'word', 'styles.xml')
            numbering_xml = os.path.join(tmpdir, 'word', 'numbering.xml')
            
            # Update styles and numbering
            changed = self.update_heading_styles(styles_xml)
            heading_styles = self.get_heading_styles(styles_xml)
            self.set_headings_numbering(numbering_xml, heading_styles)
            self.rebase_heading_numbering(tmpdir, heading_styles)
            
            # Rezip
            with zipfile.ZipFile(outname, "w", zipfile.ZIP_DEFLATED) as docx_file:
                for root, _, files in os.walk(tmpdir):
                    for file in files:
                        abs_filename = os.path.join(root, file)
                        rel_filename = os.path.relpath(abs_filename, tmpdir)
                        docx_file.write(abs_filename, rel_filename)
        
        return changed
    
    def remove_content_before_first_heading(self, doc, filename):
        """Remove content before first heading"""
        body = doc.element.body
        found_heading = False
        to_remove = []
        
        for child in body.iterchildren():
            if isinstance(child, CT_P):
                para = Paragraph(child, doc)
                if para.style and para.style.name and "Heading" in para.style.name:
                    found_heading = True
                    break
            if not found_heading:
                to_remove.append(child)
        
        for elem in to_remove:
            body.remove(elem)
        
        if to_remove:
            self.log_message(f"Removed {len(to_remove)} elements from '{os.path.basename(filename)}'")
        
        return doc
    
    def rtf_to_docx(self, input_rtf):
        """Convert RTF to DOCX using Word COM"""
        output_docx = input_rtf.replace(".rtf", ".docx")
        
        try:
            input_rtf = os.path.abspath(input_rtf)
            output_docx = os.path.abspath(output_docx)
            
            if os.path.exists(output_docx):
                self.log_message(f"{output_docx} already exists, skipping conversion.")
                return output_docx
            
            # Kill any existing Word processes
            os.system("taskkill /im WINWORD.EXE /f >nul 2>&1")
            time.sleep(0.5)
            
            word = win32.Dispatch("Word.Application")
            word.DisplayAlerts = 0
            word.Visible = False
            
            doc = word.Documents.Open(input_rtf)
            doc.SaveAs(output_docx, FileFormat=16)
            doc.Close(SaveChanges=0)
            word.Quit()
            time.sleep(0.5)
            
            self.log_message(f"Converted {os.path.basename(input_rtf)} to DOCX")
            return output_docx
            
        except Exception as e:
            self.log_message(f"Error converting {input_rtf}: {e}")
            try:
                word.Quit()
            except:
                pass
            return None
    
    def insert_styled_paragraph_at_start(self, doc, text):
        """Insert styled paragraph at document start"""
        p = doc.add_paragraph()
        run = p.add_run(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(136, 136, 136)
        
        # Move to start
        body = doc._body._element
        body.remove(p._element)
        body.insert(0, p._element)
    
    def add_styled_paragraph_at_end(self, doc, text):
        """Add styled paragraph at document end"""
        p = doc.add_paragraph()
        run = p.add_run(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(136, 136, 136)
    
    def merge_documents(self, docx_file_paths, output_filename='combined_from_list.docx'):
        """Merge multiple DOCX files"""
        sorted_file_paths = sorted(docx_file_paths)
        if not sorted_file_paths:
            self.log_message("No DOCX files to merge")
            return None
        
        self.log_message(f"Combining {len(sorted_file_paths)} files in alphabetical order")
        
        master_document_path = sorted_file_paths[0]
        try:
            master_doc_obj = Document(master_document_path)
            processed_master_doc = self.remove_content_before_first_heading(master_doc_obj, master_document_path)
            
            # Insert begin marker
            self.insert_styled_paragraph_at_start(processed_master_doc, 
                                                f"______Begin of {os.path.basename(master_document_path)}______")
            composer = Composer(processed_master_doc)
            
        except Exception as e:
            self.log_message(f"Error processing master document '{os.path.basename(master_document_path)}': {e}")
            return None
        
        # Process remaining documents
        if len(sorted_file_paths) > 1:
            for doc_path in sorted_file_paths[1:]:
                try:
                    doc_to_append = Document(doc_path)
                    processed_doc = self.remove_content_before_first_heading(doc_to_append, doc_path)
                    
                    # Insert begin marker
                    self.insert_styled_paragraph_at_start(processed_doc, 
                                                        f"______Begin of {os.path.basename(doc_path)}______")
                    composer.append(processed_doc)
                    
                    # Add end marker
                    self.add_styled_paragraph_at_end(composer.doc, 
                                                   f"______End of {os.path.basename(doc_path)}______")
                    
                except Exception as e:
                    self.log_message(f"Could not append {os.path.basename(doc_path)}: {e}")
                    continue
        
        # Add end marker for master document
        self.add_styled_paragraph_at_end(composer.doc, 
                                       f"______End of {os.path.basename(master_document_path)}______")
        
        try:
            output_path = os.path.join(os.getcwd(), output_filename)
            composer.save(output_path)
            self.log_message(f"Successfully combined {len(sorted_file_paths)} documents into '{output_filename}'")
            return output_path
        except Exception as e:
            self.log_message(f"Error saving final document: {e}")
            return None
    
    def RTF_merge(self):
        """RTF to DOCX conversion and merging for all Fiches subfolders, output by subfolder name."""
        try:
            self.log_message("Starting batch RTF to DOCX conversion and merging for all Fiches subfolders...")
    
            base_folder = "Fiches"
            if not os.path.exists(base_folder):
                messagebox.showerror("Error", "Fiches folder not found. Please run previous steps first.")
                return
    
            subfolders = [f for f in os.listdir(base_folder)
                          if os.path.isdir(os.path.join(base_folder, f))]
    
            if not subfolders:
                messagebox.showinfo("No Folders", "No subfolders found in Fiches.")
                return
    
            any_merged = False
    
            for subfolder in subfolders:
                folder_path = os.path.join(base_folder, subfolder)
                # Find RTF files in the subfolder
                rtf_files = [os.path.join(folder_path, f) for f in os.listdir(folder_path) if f.lower().endswith(".rtf")]
    
                if not rtf_files:
                    self.log_message(f"[{subfolder}]: No .rtf files found, skipping.")
                    continue
    
                self.log_message(f"[{subfolder}]: Found {len(rtf_files)} RTF files to process")
    
                # Convert RTF files to DOCX
                converted_files = []
                for i, rtf_file in enumerate(rtf_files):
                    self.update_progress((i + 1) / (len(rtf_files) + 2), f"[{subfolder}] Converting {i+1}/{len(rtf_files)}: {os.path.basename(rtf_file)}")
                    docx_file = self.rtf_to_docx(rtf_file)
                    if docx_file:
                        converted_files.append(docx_file)
                if not converted_files:
                    self.log_message(f"[{subfolder}]: No RTF files could be converted to DOCX, skipping.")
                    continue
    
                # Merge documents
                self.update_progress(0.90, f"[{subfolder}] Merging documents...")
                output_filename = os.path.join(folder_path, f"{subfolder}.docx")
                final_output_path = self.merge_documents(converted_files, output_filename)
    
                if not final_output_path:
                    self.log_message(f"[{subfolder}]: Document merging failed.")
                    continue
    
                # Apply styling
                self.update_progress(0.95, f"[{subfolder}] Applying styling...")
                outname = os.path.splitext(output_filename)[0] + "_styled.docx"
                changed = self.patch_docx(output_filename, outname)
    
                self.log_message(f"[{subfolder}] RTF merging completed. Output saved as '{outname}'")
                any_merged = True
    
            self.update_progress(1.0, "All RTF merging operations completed!")
            if any_merged:
                messagebox.showinfo("Success", "RTF conversion & merging completed!\nEach folder in Fiches containing RTF files has a merged and styled DOCX output named after the folder.\n\nFeatures Applied:\n• Content before first heading removed\n• Documents merged alphabetically\n• Styling applied")
            else:
                messagebox.showinfo("Notice", "No .rtf files found in any Fiches subfolder. No merged outputs created.")
    
        except Exception as e:
            self.log_message(f"Error in RTF merging: {str(e)}")
            messagebox.showerror("Error", f"Error in RTF merging: {str(e)}")

    
    def list_xls_files(self, folder_path):
        """List XLS files in folder"""
        xls_files = []
        try:
            xls_files = [file for file in os.listdir(folder_path) if file.endswith('.xls')]
        except FileNotFoundError:
            self.log_message(f"Folder '{folder_path}' does not exist.")
        except Exception as e:
            self.log_message(f"Error listing files: {e}")
        return xls_files
    
    def prepare_dico(self):
        """Prepare DICO sheet for each subfolder in Fiches, saving Dico.xlsx to each subfolder."""
        try:
            self.log_message("Starting DICO preparation (per subfolder)...")
            base_folder = "Fiches"
            if not os.path.exists(base_folder):
                messagebox.showerror("Error", "Fiches folder not found. Please run 'Fetch Repair Methods' first.")
                return
    
            subfolders = [f for f in os.listdir(base_folder)
                          if os.path.isdir(os.path.join(base_folder, f))]
    
            if not subfolders:
                messagebox.showinfo("No Folders", "No subfolders found in Fiches.")
                return
    
            for subfolder in subfolders:
                folder_path = os.path.join(base_folder, subfolder)
                xls_files = [file for file in os.listdir(folder_path) if file.lower().endswith('.xls')]
                if not xls_files:
                    self.log_message(f"No XLS files in {subfolder}, skipping.")
                    continue
    
                self.log_message(f"[{subfolder}] Found {len(xls_files)} XLS files to process")
                dico = pd.DataFrame()
                processed = 0
    
                for xl in xls_files:
                    xl_path = os.path.join(folder_path, xl)
                    self.update_progress((processed + 1) / len(xls_files),
                                        f"[{subfolder}] Processing {processed+1}/{len(xls_files)}: {xl}")
                    try:
                        df = pd.read_excel(xl_path, sheet_name="Diagnosis resources", skiprows=4)
                        df = df.dropna(axis=1, how='all')
                        if "DID" not in df.columns:
                            self.log_message(f"[{subfolder}] Skipping {xl} - no DID column found")
                            processed += 1
                            continue
                        df.dropna(subset=["DID"], inplace=True)
                        df["Fiche"] = xl_path
                        dico = pd.concat([dico, df], ignore_index=True)
                    except Exception as e:
                        self.log_message(f"[{subfolder}] Skipping {xl} - error: {e}")
                        processed += 1
                        continue
                    processed += 1
    
                if dico.empty:
                    self.log_message(f"[{subfolder}] No valid DICO data after processing, skipping saving.")
                    continue
    
                # Filter out routines and IO control
                try:
                    dico = dico[~dico["DID"].str.contains("Routines|IO Control", na=False)]
                except Exception:
                    pass
    
                columns_to_check = [col for col in dico.columns if col not in ['DID', 'Fiche']]
                for col in columns_to_check:
                    dico[col] = dico[col].apply(lambda x: None if isinstance(x, str) and len(x) > 6 else x)
    
                def summarize_row(row):
                    unique_values = row.dropna().unique()
                    return unique_values[0] if len(unique_values) == 1 else None
    
                dico['Value'] = dico[columns_to_check].apply(summarize_row, axis=1)
                dico.loc[dico['Value'].notna(), columns_to_check] = None
                dico = dico.dropna(axis=1, how='all')
                
                if 'DID' in dico.columns:
                    dico['DID'] = dico['DID'].astype(str).str.strip().str.replace(" ", "")
                if 'Value' in dico.columns:
                    dico['Value'] = dico['Value'].astype(str).str.strip().str.replace(" ", "")
    
                # Save DICO to the same subfolder
                output_path = os.path.join(folder_path, "Dico.xlsx")
                dico.to_excel(output_path, index=False)
                self.log_message(f"[{subfolder}] DICO preparation completed. {len(dico)} entries saved to {output_path}")
    
            self.update_progress(1.0, "All DICO preparations completed!")
            messagebox.showinfo("Success", "DICO preparation completed for all subfolders.\n"
                                "Each folder (with valid XLS) has its own Dico.xlsx.")
    
        except Exception as e:
            self.log_message(f"Error in DICO preparation: {str(e)}")
            messagebox.showerror("Error", f"Error in DICO preparation: {str(e)}")

    
    def replace_dico(self):
    
        self.log_message("Starting DICO placeholder replacement in RTF merge outputs...")
    
        base_folder = "Fiches"
        if not os.path.exists(base_folder):
            messagebox.showerror("Error", "Fiches folder not found.")
            return
    
        subfolders = [f for f in os.listdir(base_folder) if os.path.isdir(os.path.join(base_folder, f))]
        any_processed = False
    
        for subfolder in subfolders:
            folder_path = os.path.join(base_folder, subfolder)
            dico_path = os.path.join(folder_path, "Dico.xlsx")
            docx_name_base = os.path.join(folder_path, subfolder)
            docx_path = docx_name_base + "_styled.docx"
            if not os.path.exists(docx_path):
                docx_path = docx_name_base + ".docx"
            if not (os.path.exists(dico_path) and os.path.exists(docx_path)):
                self.log_message(f"[{subfolder}]: Missing DICO.xlsx or DOCX, skipping.")
                continue
    
            try:
                dico_df = pd.read_excel(dico_path)
                doc = Document(docx_path)
            except Exception as e:
                self.log_message(f"[{subfolder}]: Error opening files: {e}")
                continue
    
            # Build robust replacement dictionary (strip, remove all spaces, @-enforce)
            rep_dict = {}
            for _, row in dico_df.iterrows():
                if pd.isna(row.get("DID")) or pd.isna(row.get("Value")):
                    continue
                # Clean up DID and Value
                did = str(row["DID"]).strip().replace(" ", "")
                val = str(row["Value"]).strip().replace(" ", "")
                if not did or not val or val.lower() in ["nan", "none", ""]:
                    continue
                # Enforce @ placeholder
                if not did.startswith("@"):
                    did = "@" + did
                if not did.endswith("@"):
                    did = did + "@"
                rep_dict[did] = val
    
            replacements_made = 0
    
            def replace_in_paragraph(paragraph, rep_dict):
                nonlocal replacements_made
                runs = paragraph.runs
                if not runs:
                    return
                combined = ''.join(run.text for run in runs)
                new_text = combined
                # Replace by length (longest first)
                for did in sorted(rep_dict, key=len, reverse=True):
                    if did in new_text:
                        cnt = new_text.count(did)
                        if cnt:
                            replacements_made += cnt
                        new_text = new_text.replace(did, str(rep_dict[did]))
                if new_text != combined and runs:
                    for run in runs:
                        run.text = ""
                    runs[0].text = new_text
    
            def replace_in_table(table, rep_dict):
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            replace_in_paragraph(para, rep_dict)
                        # Recursively check nested tables
                        for nested_table in cell.tables:
                            replace_in_table(nested_table, rep_dict)
    
            # Process all document paragraphs (outside tables)
            for para in doc.paragraphs:
                replace_in_paragraph(para, rep_dict)
    
            # Process all tables (including nested)
            for table in doc.tables:
                replace_in_table(table, rep_dict)
    
            # Save updated DOCX
            updated_path = os.path.join(folder_path, f"{subfolder}_DICO_Updated.docx")
            try:
                doc.save(updated_path)
                self.log_message(f"[{subfolder}]: DICO replacement done. {replacements_made} replacements. Saved: {updated_path}")
                any_processed = True
            except Exception as e:
                self.log_message(f"[{subfolder}]: Error saving updated DOCX: {e}")
    
        if any_processed:
            messagebox.showinfo("Success", "DICO placeholder replacement done for all valid Fiches subfolders.")
        else:
            messagebox.showinfo("Notice", "No valid RTF-Merge+DICO pairs found; no replacements done.")

    
    def generate_report(self):
        """Generate comprehensive report for all activities"""
        try:
            self.log_message("Generating comprehensive report...")
            
            # Collect all generated files and create summary
            FID = self.FID_path_var.get()
            if not FID:
                messagebox.showwarning("Error", "Please select FID file first")
                return
            
            report_data = []
            
            # Check Phase 1 outputs
            phase1_files = [
                (f"{FID}_RM_Fetch.xlsx", "Repair Methods Fetch Results"),
                (f"{FID}_Analysis.xlsx", "FID Comparison Analysis"),
                (f"{FID}.docx", "6th Chapter (Diagnostic Resources)"),
                (f"{self.DFS_path_var.get()}.xlsx", "DTC/FT Extraction Results")
            ]
            
            for file_path, description in phase1_files:
                if file_path and os.path.exists(file_path):
                    size = os.path.getsize(file_path)
                    modified = time.ctime(os.path.getmtime(file_path))
                    report_data.append({
                        'Phase': 'Phase 1',
                        'File': os.path.basename(file_path),
                        'Description': description,
                        'Size (KB)': round(size/1024, 2),
                        'Last Modified': modified,
                        'Status': 'Available'
                    })
                else:
                    report_data.append({
                        'Phase': 'Phase 1',
                        'File': os.path.basename(file_path) if file_path else 'N/A',
                        'Description': description,
                        'Size (KB)': 0,
                        'Last Modified': 'N/A',
                        'Status': 'Not Generated'
                    })
            
            # Check Phase 2 outputs
            phase2_files = [
                ("Bibant_Output_styled.docx", "RTF Merged Document"),
                ("Fiches//Dico.xlsx", "DICO Sheet"),
                ("*_DICO_Updated.docx", "DICO Updated Documents")
            ]
            
            for file_pattern, description in phase2_files:
                if "*" in file_pattern:
                    # Find files matching pattern
                    found_files = []
                    for root, dirs, files in os.walk("."):
                        for file in files:
                            if file.endswith("_DICO_Updated.docx"):
                                found_files.append(os.path.join(root, file))
                    
                    if found_files:
                        for file_path in found_files:
                            size = os.path.getsize(file_path)
                            modified = time.ctime(os.path.getmtime(file_path))
                            report_data.append({
                                'Phase': 'Phase 2',
                                'File': os.path.basename(file_path),
                                'Description': description,
                                'Size (KB)': round(size/1024, 2),
                                'Last Modified': modified,
                                'Status': 'Available'
                            })
                    else:
                        report_data.append({
                            'Phase': 'Phase 2',
                            'File': 'No DICO updated files',
                            'Description': description,
                            'Size (KB)': 0,
                            'Last Modified': 'N/A',
                            'Status': 'Not Generated'
                        })
                else:
                    if os.path.exists(file_pattern):
                        size = os.path.getsize(file_pattern)
                        modified = time.ctime(os.path.getmtime(file_pattern))
                        report_data.append({
                            'Phase': 'Phase 2',
                            'File': os.path.basename(file_pattern),
                            'Description': description,
                            'Size (KB)': round(size/1024, 2),
                            'Last Modified': modified,
                            'Status': 'Available'
                        })
                    else:
                        report_data.append({
                            'Phase': 'Phase 2',
                            'File': os.path.basename(file_pattern),
                            'Description': description,
                            'Size (KB)': 0,
                            'Last Modified': 'N/A',
                            'Status': 'Not Generated'
                        })
            
            # Create report DataFrame
            report_df = pd.DataFrame(report_data)
            
            # Add summary information
            summary_data = [
                ['Report Generated', time.strftime('%Y-%m-%d %H:%M:%S')],
                ['FID File', os.path.basename(FID) if FID else 'N/A'],
                ['DED/DFS File', os.path.basename(self.DFS_path_var.get()) if self.DFS_path_var.get() else 'N/A'],
                ['Engine Type', self.radio_var.get()],
                ['Activity Type', self.activity_var.get()],
                ['Total Files Generated', len([r for r in report_data if r['Status'] == 'Available'])],
                ['Phase 1 Files', len([r for r in report_data if r['Phase'] == 'Phase 1' and r['Status'] == 'Available'])],
                ['Phase 2 Files', len([r for r in report_data if r['Phase'] == 'Phase 2' and r['Status'] == 'Available'])]
            ]
            
            summary_df = pd.DataFrame(summary_data, columns=['Parameter', 'Value'])
            
            # Save comprehensive report
            report_filename = f"Comprehensive_Report_{time.strftime('%Y%m%d_%H%M%S')}.xlsx"
            
            with pd.ExcelWriter(report_filename) as writer:
                summary_df.to_excel(writer, sheet_name='Summary', index=False)
                report_df.to_excel(writer, sheet_name='File_Status', index=False)
            
            self.log_message(f"Comprehensive report generated: {report_filename}")
            
            # Show summary in message box
            available_files = len([r for r in report_data if r['Status'] == 'Available'])
            total_files = len(report_data)
            
            messagebox.showinfo("Report Generated", 
                              f"Comprehensive report generated successfully!\n\n"
                              f"Files Available: {available_files}/{total_files}\n"
                              f"Report saved as: {report_filename}")
            
        except Exception as e:
            self.log_message(f"Error generating report: {str(e)}")
            messagebox.showerror("Error", f"Error generating report: {str(e)}")
    
    def One_STEP(self):
        """Start the complete workflow in a separate thread"""
        def workflow_thread():
            try:
                # Disable workflow button
                self.workflow_button.configure(state="disabled", text="Running Workflow...")
                
                self.log_message("Starting complete workflow...")
                self.update_progress(0, "Initializing workflow...")
                
                # Validate inputs
                FID = self.FID_path_var.get()
                if not FID:
                    messagebox.showerror("Error", "Please select FID file before starting workflow")
                    return
                
                # Step 1: Extract DTC/FT (if DFS is provided)
                if self.DFS_path_var.get():
                    self.update_progress(0.1, "Step 1: Extracting DTC/FT...")
                    self.extract_dtc_ft()
                
                # Step 2: Generate 6th Chapter
                self.update_progress(0.2, "Step 2: Generating 6th Chapter...")
                self.gentabfiche()
                
                # Step 3: Fetch Repair Methods
                self.update_progress(0.4, "Step 3: Fetching Repair Methods...")
                self.fetch_rm()
                
                # Step 4: Prepare DICO
                self.update_progress(0.6, "Step 4: Preparing DICO...")
                self.prepare_dico()
                
                # Step 5: RTF to DOCX Merge (if RTF files are available)
                self.update_progress(0.7, "Step 5: Processing RTF files...")
                self.RTF_merge()
                
                # Step 6: DICO replacement
                self.update_progress(0.8, "Step 5: Processing RTF files...")
                self.replace_dico()

                # Step 7: Generate Report
                self.update_progress(0.9, "Step 6: Generating comprehensive report...")
                self.generate_report()
                
                # Complete
                self.update_progress(1.0, "Workflow completed successfully!")
                self.log_message("Complete workflow finished successfully!")
                
                messagebox.showinfo("Workflow Complete", 
                                  "Complete workflow finished successfully!\n\n"
                                  "All Phase 1 and Phase 2 processes have been executed.\n"
                                  "Check the log for detailed information.")
                
            except Exception as e:
                self.log_message(f"Error in complete workflow: {str(e)}")
                messagebox.showerror("Workflow Error", f"Error in complete workflow: {str(e)}")
            
            finally:
                # Re-enable workflow button
                self.workflow_button.configure(state="normal", text="Start Complete Workflow")
        
        # Start workflow in separate thread
        thread = threading.Thread(target=workflow_thread)
        thread.daemon = True
        thread.start()
        
    def on_closing(self):
        """Handle proper application shutdown (safe and clean)"""
        try:
            # Optionally log session end or perform custom cleanup
            if hasattr(self, 'log_manager'):
                self.log_manager.log_session_end()
    
            # Call your resource cleanup routine if any
            self.cleanup()
    
            # Properly destroy the GUI
            if hasattr(self, 'app'):
                self.app.quit()
                self.app.destroy()
        except Exception as e:
            print(f"Error during shutdown: {e}")
    
    def cleanup(self):
        """Clean up resources"""
        with contextlib.suppress(Exception):
            pass
    
    def run(self):
        """Run the application with error handling"""
        try:
            self.app.mainloop()
        except KeyboardInterrupt:
            self.on_closing()
        except Exception as e:
            print(f"Application error: {e}")
            self.on_closing()

# Main execution
if __name__ == "__main__":
    # Create and run the application
    app = Phase2Tool()
    app.run()
