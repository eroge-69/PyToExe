from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox

def ekstrak_rancangan(tabel):
    data = []
    for row in tabel.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        data.append(cells)
    return data

def buat_perubahan(data_rancangan):
    perubahan = []
    catatan = []
    for row in data_rancangan:
        pasal = row[0]
        isi_lama = row[1]
        isi_baru = isi_lama.replace("warga negara", "setiap orang")  # Contoh sesuai asas kesetaraan
        alasan = "Disesuaikan dengan asas kesetaraan (Lampiran II UU 12/2011)"

        perubahan.append([pasal, isi_baru, alasan])
        catatan.append([pasal, f"Frasa 'warga negara' diganti 'setiap orang' sesuai asas kesetaraan"])
    return perubahan, catatan

def isi_tabel(tabel, data):
    while len(tabel.rows) > 1:
        tabel._tbl.remove(tabel.rows[1]._tr)
    for baris in data:
        row = tabel.add_row()
        for i, value in enumerate(baris):
            row.cells[i].text = value

def proses_dokumen(filepath):
    try:
        doc = Document(filepath)
        tabel1 = doc.tables[0]
        tabel2 = doc.tables[1]
        tabel3 = doc.tables[2]

        data_rancangan = ekstrak_rancangan(tabel1)
        data_perubahan, data_catatan = buat_perubahan(data_rancangan)

        isi_tabel(tabel2, data_perubahan)
        isi_tabel(tabel3, data_catatan)

        output_path = filepath.replace(".docx", "_hasil.docx")
        doc.save(output_path)
        messagebox.showinfo("Berhasil", f"Dokumen berhasil diproses:\n{output_path}")
    except Exception as e:
        messagebox.showerror("Gagal", f"Terjadi kesalahan:\n{e}")

def pilih_file():
    filepath = filedialog.askopenfilename(filetypes=[("Word Document", "*.docx")])
    if filepath:
        proses_dokumen(filepath)

root = tk.Tk()
root.title("Pengisi Tabel Perubahan Raperbup")

tk.Label(root, text="Masukkan file Raperbup (.docx) dengan 3 tabel:").pack(pady=10)
tk.Button(root, text="Pilih File", command=pilih_file).pack(pady=10)

root.mainloop()
