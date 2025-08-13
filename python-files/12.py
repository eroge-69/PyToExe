from docx import Document

def process_docx(input_file, output_file):
    # Загружаем документ
    doc = Document(input_file)
    
    # Предполагаем, что таблица первая в документе
    table = doc.tables[0]
    
    # Создаем новый документ для результатов
    new_doc = Document()
    new_table = new_doc.add_table(rows=1, cols=2)
    
    # Обрабатываем строки, начиная с третьей (индекс 2)
    for row in table.rows[2:]:
        # Проверяем, есть ли в пятом столбце "АО"
        if len(row.cells) >= 5 and "АО" in row.cells[4].text:
            # Копируем первый и второй столбцы
            col1 = row.cells[0].text
            col2 = row.cells[1].text
            
            # Добавляем строку в новую таблицу
            new_row = new_table.add_row()
            new_row.cells[0].text = col1
            new_row.cells[1].text = col2
    
    # Сохраняем новый документ
    new_doc.save(output_file)
    print(f"Обработанный документ сохранен как {output_file}")

# Использование
input_file = "13.08.docx"  # Укажите путь к исходному файлу
output_file = "out.docx"  # Укажите путь для сохранения результата
process_docx(input_file, output_file)