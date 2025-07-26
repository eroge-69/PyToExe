import tkinter as tk
from tkinter import messagebox, filedialog
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os
import re

template_file = None
entries = {}
output_name_entry = None

def select_template():
    global template_file, entries
    for widget in form_frame.winfo_children():
        widget.destroy()
    entries.clear()

    file_path = filedialog.askopenfilename(
        title="Select Word Template",
        filetypes=[("Word Documents", "*.docx")]
    )
    if file_path:
        template_file = file_path
        template_label.config(text=os.path.basename(file_path))
        fields = extract_placeholders(file_path)
        if not fields:
            messagebox.showwarning("No Fields", "No placeholders like {{field}} found in the document.")
            return
        for i, field in enumerate(fields):
            tk.Label(form_frame, text=field + ":").grid(row=i, column=0, padx=10, pady=5, sticky="ne")
            text_widget = tk.Text(form_frame, width=40, height=2, wrap="word")
            text_widget.grid(row=i, column=1, padx=10, pady=5)
            text_widget.bind("<KeyRelease>", lambda e, t=text_widget: auto_resize(t))
            entries[field] = text_widget

        # Output file name entry
        i += 1
        tk.Label(form_frame, text="Output File Name (no extension):").grid(row=i, column=0, padx=10, pady=10, sticky="e")
        global output_name_entry
        output_name_entry = tk.Entry(form_frame, width=40)
        output_name_entry.grid(row=i, column=1, padx=10, pady=10)
        output_name_entry.insert(0, "output")

def auto_resize(text_widget):
    content = text_widget.get("1.0", "end-1c")
    lines = content.count("\n") + 1
    text_widget.configure(height=min(max(lines, 2), 10))

def extract_placeholders(docx_path):
    doc = Document(docx_path)
    pattern = r"{{(.*?)}}"
    placeholders = []
    seen = set()

    # Search in paragraphs
    for para in doc.paragraphs:
        found = re.findall(pattern, para.text)
        for f in found:
            if f not in seen:
                placeholders.append(f)
                seen.add(f)

    # Search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    found = re.findall(pattern, para.text)
                    for f in found:
                        if f not in seen:
                            placeholders.append(f)
                            seen.add(f)

    return placeholders

def fill_template():
    if not template_file or not os.path.exists(template_file):
        messagebox.showerror("Error", "No valid template file selected.")
        return

    data = { field: entry.get("1.0", "end-1c").strip() for field, entry in entries.items() }
    output_name = output_name_entry.get().strip() or "output"
    if not output_name.lower().endswith(".docx"):
        output_name += ".docx"

    doc = Document(template_file)

    # Replace in paragraphs
    for para in doc.paragraphs:
        for key, val in data.items():
            para.text = para.text.replace(f"{{{{{key}}}}}", val)
        for run in para.runs:
            run.font.name = "Noto Serif Malayalam"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "Noto Serif Malayalam")
            run.font.size = Pt(12)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, val in data.items():
                        para.text = para.text.replace(f"{{{{{key}}}}}", val)
                    for run in para.runs:
                        run.font.name = "Noto Serif Malayalam"
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), "Noto Serif Malayalam")
                        run.font.size = Pt(12)

    doc.save(output_name)
    messagebox.showinfo("Success", f"Document saved as '{output_name}'")

# GUI setup
root = tk.Tk()
root.title("DOCX Auto-Filler (Dynamic Final Version)")

tk.Button(root, text="Select Template", command=select_template, bg="blue", fg="white").pack(pady=10)
template_label = tk.Label(root, text="No file selected", fg="gray")
template_label.pack()

form_frame = tk.Frame(root)
form_frame.pack(padx=20, pady=10)

tk.Button(root, text="Generate Document", command=fill_template, bg="green", fg="white").pack(pady=10)

root.mainloop()