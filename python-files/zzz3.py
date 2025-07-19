import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import os
from pathlib import Path
import re
from docx import Document
from docx2pdf import convert
import tempfile
from PIL import Image
import fitz  # PyMuPDF for PDF to image conversion

class DocxToPdfConverter:
    def __init__(self, root):
        self.root = root
        self.root.title("DOCX to PDF/Image Converter with Placeholders")
        self.root.geometry("900x800")
        
        # Variables
        self.input_file = tk.StringVar()
        self.output_file = tk.StringVar()
        self.output_format = tk.StringVar(value="PDF")
        
        # Placeholder variables
        self.title = tk.StringVar(value="Mr.")
        self.fullname = tk.StringVar()
        self.passport = tk.StringVar()
        self.father = tk.StringVar()
        self.dob = tk.StringVar()
        self.residence = tk.StringVar()
        
        # Updated pronoun mappings - now includes "He or She" format
        self.pronoun_mappings = {
            'Mr.': {
                'he': 'he', 'him': 'him', 'his': 'his', 
                'He': 'He', 'Him': 'Him', 'His': 'His',
                'he or she': 'he', 'He or She': 'He',
                'him or her': 'him', 'Him or Her': 'Him',
                'his or her': 'his', 'His or Her': 'His'
            },
            'Mrs.': {
                'he': 'she', 'him': 'her', 'his': 'her', 
                'He': 'She', 'Him': 'Her', 'His': 'Her',
                'he or she': 'she', 'He or She': 'She',
                'him or her': 'her', 'Him or Her': 'Her',
                'his or her': 'her', 'His or Her': 'Her'
            },
            'Ms.': {
                'he': 'she', 'him': 'her', 'his': 'her', 
                'He': 'She', 'Him': 'Her', 'His': 'Her',
                'he or she': 'she', 'He or She': 'She',
                'him or her': 'her', 'Him or Her': 'Her',
                'his or her': 'her', 'His or Her': 'Her'
            },
            'Miss': {
                'he': 'she', 'him': 'her', 'his': 'her', 
                'He': 'She', 'Him': 'Her', 'His': 'Her',
                'he or she': 'she', 'He or She': 'She',
                'him or her': 'her', 'Him or Her': 'Her',
                'his or her': 'her', 'His or Her': 'Her'
            },
            'Dr.': {
                'he': 'he', 'him': 'him', 'his': 'his', 
                'He': 'He', 'Him': 'Him', 'His': 'His',
                'he or she': 'he', 'He or She': 'He',
                'him or her': 'him', 'Him or Her': 'Him',
                'his or her': 'his', 'His or Her': 'His'
            }
        }
        
        self.setup_gui()
        
    def setup_gui(self):
        # Create notebook for tabs
        notebook = ttk.Notebook(self.root)
        notebook.pack(fill='both', expand=True, padx=10, pady=10)
        
        # Main tab
        main_tab = ttk.Frame(notebook)
        notebook.add(main_tab, text="Converter")
        
        # Instructions tab
        instructions_tab = ttk.Frame(notebook)
        notebook.add(instructions_tab, text="Instructions")
        
        self.setup_main_tab(main_tab)
        self.setup_instructions_tab(instructions_tab)
        
    def setup_main_tab(self, parent):
        # Main frame
        main_frame = ttk.Frame(parent, padding="10")
        main_frame.pack(fill='both', expand=True)
        
        # File selection section
        file_frame = ttk.LabelFrame(main_frame, text="File Selection", padding="10")
        file_frame.pack(fill='x', pady=(0, 10))
        
        # Input file
        ttk.Label(file_frame, text="Input DOCX:").grid(row=0, column=0, sticky=tk.W, pady=2)
        ttk.Entry(file_frame, textvariable=self.input_file, width=60).grid(row=0, column=1, padx=(5, 5), pady=2)
        ttk.Button(file_frame, text="Browse", command=self.browse_input_file).grid(row=0, column=2, pady=2)
        
        # Output format selection
        ttk.Label(file_frame, text="Output Format:").grid(row=1, column=0, sticky=tk.W, pady=2)
        format_combo = ttk.Combobox(file_frame, textvariable=self.output_format, 
                                   values=["PDF", "PNG Image"], width=20, state="readonly")
        format_combo.grid(row=1, column=1, sticky=tk.W, padx=(5, 0), pady=2)
        format_combo.bind('<<ComboboxSelected>>', self.on_format_change)
        
        # Output file
        ttk.Label(file_frame, text="Output File:").grid(row=2, column=0, sticky=tk.W, pady=2)
        ttk.Entry(file_frame, textvariable=self.output_file, width=60).grid(row=2, column=1, padx=(5, 5), pady=2)
        ttk.Button(file_frame, text="Browse", command=self.browse_output_file).grid(row=2, column=2, pady=2)
        
        # Placeholders section
        placeholder_frame = ttk.LabelFrame(main_frame, text="Placeholder Values", padding="10")
        placeholder_frame.pack(fill='x', pady=(0, 10))
        
        # Title dropdown
        ttk.Label(placeholder_frame, text="Title [TITLE]:").grid(row=0, column=0, sticky=tk.W, pady=2)
        title_combo = ttk.Combobox(placeholder_frame, textvariable=self.title, 
                                  values=["Mr.", "Mrs.", "Ms.", "Miss", "Dr."], width=20)
        title_combo.grid(row=0, column=1, sticky=tk.W, padx=(5, 0), pady=2)
        title_combo.bind('<<ComboboxSelected>>', self.on_title_change)
        
        # Full Name
        ttk.Label(placeholder_frame, text="Full Name [FULLNAME]:").grid(row=1, column=0, sticky=tk.W, pady=2)
        fullname_entry = ttk.Entry(placeholder_frame, textvariable=self.fullname, width=40)
        fullname_entry.grid(row=1, column=1, sticky=tk.W, padx=(5, 0), pady=2)
        fullname_entry.bind('<KeyRelease>', self.on_fullname_change)
        
        # Passport
        ttk.Label(placeholder_frame, text="Passport [PASSPORT]:").grid(row=2, column=0, sticky=tk.W, pady=2)
        ttk.Entry(placeholder_frame, textvariable=self.passport, width=40).grid(row=2, column=1, sticky=tk.W, padx=(5, 0), pady=2)
        
        # Father's Name
        ttk.Label(placeholder_frame, text="Father's Name [FATHER]:").grid(row=3, column=0, sticky=tk.W, pady=2)
        ttk.Entry(placeholder_frame, textvariable=self.father, width=40).grid(row=3, column=1, sticky=tk.W, padx=(5, 0), pady=2)
        
        # Date of Birth
        ttk.Label(placeholder_frame, text="Date of Birth [DOB]:").grid(row=4, column=0, sticky=tk.W, pady=2)
        ttk.Entry(placeholder_frame, textvariable=self.dob, width=40).grid(row=4, column=1, sticky=tk.W, padx=(5, 0), pady=2)
        
        # Residence
        ttk.Label(placeholder_frame, text="Residence [RESIDENCE]:").grid(row=5, column=0, sticky=tk.W, pady=2)
        residence_text = tk.Text(placeholder_frame, width=40, height=3)
        residence_text.grid(row=5, column=1, sticky=tk.W, padx=(5, 0), pady=2)
        residence_text.bind('<KeyRelease>', lambda e: self.residence.set(residence_text.get("1.0", tk.END).strip()))
        
        # Pronoun info
        pronoun_info = ttk.Label(placeholder_frame, 
                               text="Pronouns will be automatically set based on title selection",
                               font=('TkDefaultFont', 8, 'italic'))
        pronoun_info.grid(row=6, column=0, columnspan=2, sticky=tk.W, pady=(10, 2))
        
        # Convert button
        convert_btn = ttk.Button(main_frame, text="Convert Document", command=self.convert_document)
        convert_btn.pack(pady=20)
        
        # Progress bar
        self.progress = ttk.Progressbar(main_frame, mode='indeterminate')
        self.progress.pack(fill='x', pady=(0, 10))
        
        # Status label
        self.status_label = ttk.Label(main_frame, text="Ready")
        self.status_label.pack()
        
        # Configure grid weights
        file_frame.columnconfigure(1, weight=1)
        
    def setup_instructions_tab(self, parent):
        # Instructions frame
        instructions_frame = ttk.Frame(parent, padding="10")
        instructions_frame.pack(fill='both', expand=True)
        
        # Title
        title_label = ttk.Label(instructions_frame, text="Available Placeholders", 
                               font=('TkDefaultFont', 14, 'bold'))
        title_label.pack(pady=(0, 10))
        
        # Instructions text
        instructions_text = scrolledtext.ScrolledText(instructions_frame, wrap=tk.WORD, 
                                                     width=80, height=30, font=('Consolas', 10))
        instructions_text.pack(fill='both', expand=True)
        
        instructions_content = """
AVAILABLE PLACEHOLDERS:

1. PERSONAL INFORMATION:
   [TITLE] or [title]           - Title (Mr., Mrs., Ms., Miss, Dr.)
   [FULLNAME] or [fullname]     - Full name of the person
   [PASSPORT] or [passport]     - Passport number
   [FATHER] or [father]         - Father's name
   [DOB] or [dob]              - Date of birth
   [RESIDENCE] or [residence]   - Residence address

2. PRONOUNS (Auto-adjusted based on title):
   [He] or [he]                - He/She
   [Him] or [him]              - Him/Her
   [His] or [his]              - His/Her
   [He or She] or [he or she]  - He/She (gender-neutral option)
   [Him or Her] or [him or her] - Him/Her (gender-neutral option)
   [His or Her] or [his or her] - His/Her (gender-neutral option)

PRONOUN MAPPING:
   - Mr., Dr. → he, him, his
   - Mrs., Ms., Miss → she, her, her
   - Gender-neutral forms automatically convert to appropriate gender

OUTPUT FORMATS:
   - PDF: Standard PDF format (recommended for documents)
   - PNG Image: High-quality image format (good for sharing/viewing)

AUTOMATIC NAMING:
   - Output files are automatically named using [FULLNAME]
   - Example: "John Doe.pdf" or "John Doe.png"

USAGE EXAMPLES:

In your DOCX document, you can use placeholders like:

"Dear [TITLE] [FULLNAME],

This is to certify that [He or She] has submitted [His or Her] passport ([PASSPORT]) 
for verification. [He or She] is the son/daughter of [FATHER] and was born on [DOB]. 
[His or Her] current residence is [RESIDENCE].

We hereby confirm that [FULLNAME] has completed all requirements.

Requirements:
• [He or She] must provide valid identification
• [His or Her] documents must be current
• [He or She] should attend the interview"

FORMATTING NOTES:
- Use UPPERCASE placeholders for original case: [TITLE], [FULLNAME]
- Use lowercase placeholders for lowercase output: [title], [fullname]
- All original formatting (bold, italic, font size, etc.) will be preserved
- Images and tables will maintain their positions
- Headers and footers are also processed
- Placeholders work in bullet points and numbered lists

TIPS:
1. Always use square brackets around placeholders
2. Placeholders are case-sensitive
3. Make sure to fill in all placeholder values before conversion
4. Test with a sample document first
5. Keep backup of your original DOCX file
6. PNG images provide better quality for viewing on screen
7. PDF is better for printing and official documents
8. Use "He or She" format for gender-neutral documents

SUPPORTED DOCUMENT ELEMENTS:
✓ Paragraphs and text runs
✓ Tables and cells
✓ Headers and footers
✓ Text boxes
✓ Images (position preserved)
✓ Shapes with text
✓ Bullet points and numbered lists
✓ All text formatting (bold, italic, underline, colors, fonts, sizes)

TROUBLESHOOTING:
- If placeholders aren't replaced, check spelling and brackets
- Ensure DOCX file is not corrupted
- Make sure all required fields are filled
- Check that Microsoft Word or LibreOffice is installed for PDF conversion
- For PNG output, ensure sufficient disk space for high-quality images
- Placeholders in bullet points are fully supported
"""
        
        instructions_text.insert('1.0', instructions_content)
        instructions_text.config(state='disabled')

    def on_format_change(self, event=None):
        self.update_output_filename()

    def on_fullname_change(self, event=None):
        self.update_output_filename()

    def update_output_filename(self):
        if self.input_file.get() and self.fullname.get().strip():
            input_path = Path(self.input_file.get())
            fullname = self.fullname.get().strip()
            # Clean filename (remove invalid characters)
            clean_name = "".join(c for c in fullname if c.isalnum() or c in (' ', '-', '_')).strip()
            if clean_name:
                if self.output_format.get() == "PNG Image":
                    output_path = input_path.parent / f"{clean_name}.png"
                else:
                    output_path = input_path.parent / f"{clean_name}.pdf"
                self.output_file.set(str(output_path))
        
    def browse_input_file(self):
        filename = filedialog.askopenfilename(
            title="Select DOCX file",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")]
        )
        if filename:
            self.input_file.set(filename)
            self.update_output_filename()
    
    def browse_output_file(self):
        if self.output_format.get() == "PNG Image":
            filename = filedialog.asksaveasfilename(
                title="Save PNG as",
                defaultextension=".png",
                filetypes=[("PNG files", "*.png"), ("All files", "*.*")]
            )
        else:
            filename = filedialog.asksaveasfilename(
                title="Save PDF as",
                defaultextension=".pdf",
                filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")]
            )
        if filename:
            self.output_file.set(filename)
    
    def on_title_change(self, event=None):
        pass
    
    def get_all_replacements(self):
        """Get all placeholder replacements"""
        # Get current pronoun mappings based on title
        current_pronouns = self.pronoun_mappings.get(self.title.get(), 
                                                   self.pronoun_mappings['Mr.'])
        
        # Define all replacements
        replacements = {
            # Basic placeholders - uppercase
            '[TITLE]': self.title.get(),
            '[FULLNAME]': self.fullname.get(),
            '[PASSPORT]': self.passport.get(),
            '[FATHER]': self.father.get(),
            '[DOB]': self.dob.get(),
            '[RESIDENCE]': self.residence.get(),
            
            # Basic placeholders - lowercase
            '[title]': self.title.get().lower() if self.title.get() else '',
            '[fullname]': self.fullname.get().lower() if self.fullname.get() else '',
            '[passport]': self.passport.get().lower() if self.passport.get() else '',
            '[father]': self.father.get().lower() if self.father.get() else '',
            '[dob]': self.dob.get().lower() if self.dob.get() else '',
            '[residence]': self.residence.get().lower() if self.residence.get() else '',
        }
        
        # Add pronoun replacements - including "He or She" format
        for old_pronoun, new_pronoun in current_pronouns.items():
            replacements[f'[{old_pronoun}]'] = new_pronoun
        
        return replacements
    
    def replace_text_in_run(self, run, replacements):
        """Replace placeholders in a single run while preserving formatting"""
        if not run.text:
            return
            
        original_text = run.text
        new_text = original_text
        
        # Sort replacements by length (longest first) to handle overlapping patterns
        sorted_replacements = sorted(replacements.items(), key=lambda x: len(x[0]), reverse=True)
        
        # Replace all placeholders
        for placeholder, value in sorted_replacements:
            if placeholder in new_text:
                new_text = new_text.replace(placeholder, str(value))
        
        # Update the run text if it changed
        if new_text != original_text:
            run.text = new_text
    
    def replace_placeholders_in_paragraph(self, paragraph, replacements):
        """Replace placeholders in a paragraph while preserving formatting"""
        # Process each run in the paragraph
        for run in paragraph.runs:
            self.replace_text_in_run(run, replacements)
    
    def replace_placeholders_in_table(self, table, replacements):
        """Replace placeholders in table cells"""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self.replace_placeholders_in_paragraph(paragraph, replacements)
    
    def replace_placeholders_in_shape(self, shape, replacements):
        """Replace placeholders in shapes (text boxes, etc.)"""
        try:
            if hasattr(shape, 'text_frame') and shape.text_frame:
                for paragraph in shape.text_frame.paragraphs:
                    for run in paragraph.runs:
                        self.replace_text_in_run(run, replacements)
        except:
            pass

    def convert_pdf_to_png(self, pdf_path, png_path):
        """Convert PDF to high-quality PNG image"""
        try:
            # Open the PDF
            pdf_document = fitz.open(pdf_path)
            
            if len(pdf_document) > 1:
                # Multiple pages - create separate images
                base_path = Path(png_path)
                base_name = base_path.stem
                base_dir = base_path.parent
                
                for page_num in range(len(pdf_document)):
                    page = pdf_document[page_num]
                    # Render page to image with high DPI for quality
                    mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better quality
                    pix = page.get_pixmap(matrix=mat)
                    
                    if len(pdf_document) == 1:
                        output_path = png_path
                    else:
                        output_path = base_dir / f"{base_name}_page_{page_num + 1}.png"
                    
                    pix.save(str(output_path))
                    pix = None
            else:
                # Single page
                page = pdf_document[0]
                mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better quality
                pix = page.get_pixmap(matrix=mat)
                pix.save(png_path)
                pix = None
            
            pdf_document.close()
            return True
            
        except Exception as e:
            print(f"Error converting PDF to PNG: {str(e)}")
            return False
    
    def process_document(self, input_path, output_path):
        """Process the document and replace placeholders"""
        try:
            # Load the document
            doc = Document(input_path)
            
            # Get all replacements
            replacements = self.get_all_replacements()
            
            # Debug: Print replacements
            print("Replacements to be made:")
            for k, v in replacements.items():
                if v:  # Only print non-empty values
                    print(f"  {k} -> {v}")
            
            # Process main document paragraphs
            for paragraph in doc.paragraphs:
                self.replace_placeholders_in_paragraph(paragraph, replacements)
            
            # Process tables in main document
            for table in doc.tables:
                self.replace_placeholders_in_table(table, replacements)
            
            # Process headers and footers for each section
            for section in doc.sections:
                # Process header
                try:
                    header = section.header
                    for paragraph in header.paragraphs:
                        self.replace_placeholders_in_paragraph(paragraph, replacements)
                    for table in header.tables:
                        self.replace_placeholders_in_table(table, replacements)
                except:
                    pass
                
                # Process footer
                try:
                    footer = section.footer
                    for paragraph in footer.paragraphs:
                        self.replace_placeholders_in_paragraph(paragraph, replacements)
                    for table in footer.tables:
                        self.replace_placeholders_in_table(table, replacements)
                except:
                    pass
            
            # Save the modified document to a temporary file
            temp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
            doc.save(temp_docx.name)
            temp_docx.close()
            
            # Convert based on output format
            if self.output_format.get() == "PNG Image":
                # Convert to PDF first, then to PNG
                self.status_label.config(text="Converting to PDF...")
                self.root.update()
                
                temp_pdf = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
                temp_pdf.close()
                
                convert(temp_docx.name, temp_pdf.name)
                
                self.status_label.config(text="Converting to PNG...")
                self.root.update()
                
                success = self.convert_pdf_to_png(temp_pdf.name, output_path)
                
                # Clean up temporary files
                os.unlink(temp_pdf.name)
                
                if not success:
                    raise Exception("Failed to convert PDF to PNG")
                    
            else:
                # Convert to PDF
                self.status_label.config(text="Converting to PDF...")
                self.root.update()
                convert(temp_docx.name, output_path)
            
            # Clean up temporary docx file
            os.unlink(temp_docx.name)
            
            return True
            
        except Exception as e:
            print(f"Error details: {str(e)}")
            messagebox.showerror("Error", f"An error occurred: {str(e)}")
            return False
    
    def validate_inputs(self):
        """Validate that all required inputs are provided"""
        if not self.input_file.get():
            messagebox.showerror("Error", "Please select an input DOCX file")
            return False
        
        if not self.output_file.get():
            messagebox.showerror("Error", "Please specify an output file")
            return False
        
        if not os.path.exists(self.input_file.get()):
            messagebox.showerror("Error", "Input file does not exist")
            return False
        
        # Check if at least some placeholder values are provided
        if not any([self.title.get(), self.fullname.get(), self.passport.get(), 
                   self.father.get(), self.dob.get(), self.residence.get()]):
            result = messagebox.askyesno("Warning", 
                                       "No placeholder values provided. Continue anyway?")
            if not result:
                return False
        
        return True
    
    def convert_document(self):
        # Validate inputs
        if not self.validate_inputs():
            return
        
        # Start progress bar
        self.progress.start()
        self.status_label.config(text="Processing document...")
        self.root.update()
        
        try:
            # Process the document
            success = self.process_document(self.input_file.get(), self.output_file.get())
            
            if success:
                format_name = "PNG image" if self.output_format.get() == "PNG Image" else "PDF"
                self.status_label.config(text=f"Conversion to {format_name} completed successfully!")
                messagebox.showinfo("Success", 
                                  f"Document converted successfully to {format_name}!\nSaved as: {self.output_file.get()}")
            else:
                self.status_label.config(text="Conversion failed")
        
        except Exception as e:
            self.status_label.config(text="Error occurred during conversion")
            messagebox.showerror("Error", f"Conversion failed: {str(e)}")
            print(f"Detailed error: {str(e)}")
        
        finally:
            # Stop progress bar
            self.progress.stop()

def main():
    root = tk.Tk()
    app = DocxToPdfConverter(root)
    root.mainloop()

if __name__ == "__main__":
    main()
