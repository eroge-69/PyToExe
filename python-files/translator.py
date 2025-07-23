
import os
from docx import Document
from deep_translator import GoogleTranslator

def translate_docx(pdf_or_docx_path, target_lang="ko"):
    ext = os.path.splitext(pdf_or_docx_path)[1].lower()
    if ext == ".docx":
        doc = Document(pdf_or_docx_path)
        for p in doc.paragraphs:
            if p.text.strip():
                p.text = GoogleTranslator(source="auto", target=target_lang).translate(p.text)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        cell.text = GoogleTranslator(source="auto", target=target_lang).translate(cell.text)
        out = pdf_or_docx_path.replace(ext, f"_translated_{target_lang}.docx")
        doc.save(out)
        print("Saved:", out)
    elif ext == ".pdf":
        from pdf2docx import Converter
        temp_docx = pdf_or_docx_path.replace(".pdf", ".docx")
        Converter(pdf_or_docx_path).convert(temp_docx)
        translate_docx(temp_docx, target_lang)
    else:
        print("지원되지 않는 파일 형식:", ext)

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("사용법: python translator.py <파일경로>")
    else:
        translate_docx(sys.argv[1])
