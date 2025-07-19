
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox
import os

replacements = {
    "கொ": "கொ", "கோ": "கோ", "கௌ": "கெள",
    "ஙொ": "ஙொ", "ஙோ": "ஙோ", "ஙௌ": "ஙெள",
    "சொ": "சொ", "சோ": "சோ", "சௌ": "செள",
    "ஞொ": "ஞொ", "ஞோ": "ஞோ", "ஞௌ": "ஞெள",
    "டொ": "டொ", "டோ": "டோ", "டௌ": "டெள",
    "ணொ": "ணொ", "ணோ": "ணோ", "ணௌ": "ணெள",
    "தொ": "தொ", "தோ": "தோ", "தௌ": "தெள",
    "நொ": "நெப", "நோ": "நோ", "நௌ": "நெள",
    "பொ": "பொ", "போ": "போ", "பௌ": "பெள",
    "மொ": "மொ", "மோ": "மோ", "மௌ": "மெள",
    "யொ": "யொ", "யோ": "யோ", "யௌ": "யெள",
    "ரொ": "ரொ", "ரோ": "ரோ", "ரௌ": "ரெள",
    "லொ": "லொ", "லோ": "லோ", "லௌ": "லெள",
    "வொ": "வொ", "வோ": "வோ", "வௌ": "வெள",
    "ழொ": "ழொ", "ழோ": "ழோ", "ழௌ": "ழெள",
    "ளொ": "ளொ", "ளோ": "ளோ", "ளௌ": "ளெள",
    "றொ": "றொ", "றோ": "றோ", "றௌ": "றெள",
    "னொ": "னொ", "னோ": "னோ", "னௌ": "னெள",
    "ஹொ": "ஹொ", "ஹோ": "ஹோ", "ஜொ": "ஜொ", "ஜோ": "ஜோ"
}

def replace_words_in_docx(input_path, output_path):
    doc = Document(input_path)
    for para in doc.paragraphs:
        for old, new in replacements.items():
            if old in para.text:
                para.text = para.text.replace(old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old, new in replacements.items():
                    if old in cell.text:
                        cell.text = cell.text.replace(old, new)
    doc.save(output_path)

def select_and_replace():
    input_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
    if not input_path:
        return
    folder = os.path.dirname(input_path)
    output_path = os.path.join(folder, "output.docx")
    replace_words_in_docx(input_path, output_path)
    messagebox.showinfo("Done", "File saved as: " + output_path)

root = tk.Tk()
root.title("Tamil Word Replacer")
root.geometry("300x150")

label = tk.Label(root, text="Replace Words in .docx", font=("Arial", 12))
label.pack(pady=10)

btn = tk.Button(root, text="Select DOCX and Convert", command=select_and_replace)
btn.pack(pady=20)

root.mainloop()
