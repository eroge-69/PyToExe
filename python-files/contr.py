import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from tkinter.constants import *
import mysql.connector
import datetime
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
from ttkbootstrap.tableview import Tableview
from ttkbootstrap.validation import add_regex_validation
from ttkbootstrap.widgets import DateEntry
import zipfile
from io import BytesIO
import os
import winsound
import threading
import tempfile
import hashlib
from plyer import notification
import time
import logging
from docx.enum.text import WD_BREAK


class ContratApplication:
    def __init__(self, root):
        self.root = root
        self.root.title("Gestion des Contrats - Imbert Mnif")
        self.root.geometry("1600x900")
        self.style = ttk.Style(theme='flatly')
        self.style.configure("Treeview", rowheight=30, font=('Segoe UI', 10))
        self.style.configure("Treeview.Heading", font=('Segoe UI', 11, 'bold'))

        # Initialisation des variables
        self.last_contract_check = None
        self.alert_timer = None
        self.check_interval = 4 * 60 * 1000  # 4 minutes
        self.contract_hash = {}
        self.notified_contracts = set()
        self.sound_file = os.path.normpath(r"D:\UIAlert_Notification lasolisa 4 (ID 2066)_LS.wav")
        self.sound_enabled = True
        self.undo_stack = []
        self.current_employee = None
        self.logo_path = r"D:\imbertlogo.png"

        # Define validators
        self.validators = {
            "matricule": r'^\w+$',
            "cin": r'^\d{8}$',
            "date_naissance": r'^\d{2}/\d{2}/\d{4}$',
            "date_cin": r'^\d{2}/\d{2}/\d{4}$',
            "email": r'^[^@]+@[^@]+\.[^@]+$',
            "telephone": r'^\+?\d{10,12}$',
            "salaire": r'^\d+(\.\d{1,2})?$',
            "prime": r'^\d+(\.\d{1,2})?$',
            "date_debut": r'^\d{2}/\d{2}/\d{4}$',
            "date_fin": r'^\d{2}/\d{2}/\d{4}$'
        }

        # Configuration de la base de données
        self.db_config = {
            "host": "192.168.1.210",
            "user": "omar",
            "password": "1234",
            "database": "rh",
            "charset": "utf8mb4",
            "collation": "utf8mb4_unicode_ci"
        }

        self.initialize_database()
        self.initialize_ui()
        self.load_data()
        self.background_alert_service()

    def save_and_generate(self):
        try:
            # Récupération des données
            employee_data = {
                'matricule': self.get_widget_value(self.entries['matricule']),
                'nom': self.get_widget_value(self.entries['nom']),
                'prenom': self.get_widget_value(self.entries['prenom']),
                'genre': self.variables["genre"].get(),
                'date_naissance': self.get_widget_value(self.entries['date_naissance']) or None,
                'lieu_naissance': self.get_widget_value(self.entries['lieu_naissance']) or None,
                'adresse': self.get_widget_value(self.entries['adresse']) or None,
                'ville': self.get_widget_value(self.entries['ville']) or None,
                'cin': self.get_widget_value(self.entries['cin']) or None,
                'date_cin': self.get_widget_value(self.entries['date_cin']) or None,
                'lieu_cin': self.get_widget_value(self.entries['lieu_cin']) or None,
                'poste': self.get_widget_value(self.entries['poste']) or None,
                'email': self.get_widget_value(self.entries['email']) or None,
                'telephone': self.get_widget_value(self.entries['telephone']) or None
            }

            contract_data = {
                'matricule': employee_data['matricule'],
                'type_contrat': self.variables["contract_type"].get(),
                'date_debut': self.get_widget_value(self.contract_entries['date_debut']),
                'date_fin': self.get_widget_value(self.contract_entries['date_fin']) if self.variables[
                                                                                            "contract_type"].get() == "CDD" else None,
                'salaire_base': self.get_widget_value(self.contract_entries['salaire']),
                'prime': self.get_widget_value(self.contract_entries['prime']),
                'salary_type': self.variables["salary_type"].get()
            }

            # Validation des champs obligatoires
            required_fields = {
                'matricule': employee_data['matricule'],
                'nom': employee_data['nom'],
                'prenom': employee_data['prenom'],
                'date_debut': contract_data['date_debut'],
                'salaire_base': contract_data['salaire_base'],
                'prime': contract_data['prime']
            }

            for field, value in required_fields.items():
                if not value or value.isspace():
                    Messagebox.show_error(f"Le champ {field} est obligatoire.", "Erreur de validation")
                    return

            # Validation des formats
            for field, value in employee_data.items():
                if field in self.validators and value and value.strip() and not re.match(self.validators[field], value):
                    Messagebox.show_error(f"Format invalide pour {field}.", "Erreur de validation")
                    return

            for field, value in contract_data.items():
                if field in self.validators and value and value.strip() and not re.match(self.validators[field],
                                                                                         str(value)):
                    Messagebox.show_error(f"Format invalide pour {field}.", "Erreur de validation")
                    return

            # Validation des valeurs numériques
            try:
                contract_data['salaire_base'] = float(contract_data['salaire_base'])
                contract_data['prime'] = float(contract_data['prime'])
                if contract_data['salaire_base'] <= 0 or contract_data['prime'] < 0:
                    Messagebox.show_error("Le salaire doit être positif et la prime non négative.",
                                          "Erreur de validation")
                    return
            except ValueError:
                Messagebox.show_error("Le salaire et la prime doivent être des nombres valides.",
                                      "Erreur de validation")
                return

            # Validation des dates
            try:
                date_debut = datetime.datetime.strptime(contract_data['date_debut'], "%d/%m/%Y")
                if contract_data['type_contrat'] == "CDD" and contract_data['date_fin']:
                    date_fin = datetime.datetime.strptime(contract_data['date_fin'], "%d/%m/%Y")
                    if date_fin <= date_debut:
                        Messagebox.show_error("La date de fin doit être postérieure à la date de début.",
                                              "Erreur de validation")
                        return
                    duration_days = (date_fin - date_debut).days
                    if duration_days < 30:
                        def show_contract_termination_message():
                            Messagebox.show_warning(
                                f"Le contrat de {employee_data['prenom']} {employee_data['nom']} (Matricule: {employee_data['matricule']}) "
                                "est arrêté car sa durée est inférieure à 30 jours.",
                                "Arrêt du contrat"
                            )
                            try:
                                with self.conn.cursor() as cursor:
                                    cursor.execute(
                                        "UPDATE contrats SET statut = %s WHERE matricule = %s AND date_debut = %s",
                                        ("arrêté", employee_data['matricule'], contract_data['date_debut'])
                                    )
                                    self.conn.commit()
                                logging.info(
                                    f"Contrat arrêté pour {employee_data['matricule']} (durée: {duration_days} jours)")
                            except mysql.connector.Error as e:
                                logging.error(f"Erreur lors de la mise à jour du statut du contrat: {str(e)}")
                                Messagebox.show_error(f"Erreur base de données: {str(e)}", "Erreur")

                        self.root.after(300000, show_contract_termination_message)
            except ValueError:
                Messagebox.show_error("Les dates doivent être au format JJ/MM/AAAA.", "Erreur de validation")
                return

            # Vérification de l'unicité du matricule
            with self.conn.cursor() as cursor:
                cursor.execute("SELECT matricule FROM employes WHERE matricule = %s", (employee_data['matricule'],))
                if cursor.fetchone():
                    Messagebox.show_error("Ce matricule existe déjà.", "Erreur de validation")
                    return

            # Génération du texte du contrat
            contract_text = self.generate_contract_text(employee_data, contract_data)

            if not contract_text.strip():
                Messagebox.show_error("Erreur: Le texte du contrat est vide.", "Erreur de génération")
                return

            # Création du document Word
            doc = self.create_contract_doc(employee_data['matricule'], contract_text.strip())
            contract_data['texte_contrat'] = contract_text.strip()

            # Affichage du contrat
            self.contract_text.delete(1.0, tk.END)
            self.contract_text.insert(tk.END, contract_text.strip())

            # Enregistrement dans la base de données
            with self.conn.cursor() as cursor:
                # Insertion de l'employé
                cursor.execute('''
                               INSERT INTO employes (matricule, nom, prenom, genre, date_naissance, lieu_naissance,
                                                     adresse, ville, cin, date_cin, lieu_cin, poste, email, telephone)
                               VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                               ''', (
                                   employee_data['matricule'], employee_data['nom'], employee_data['prenom'],
                                   employee_data['genre'], employee_data['date_naissance'],
                                   employee_data['lieu_naissance'],
                                   employee_data['adresse'], employee_data['ville'], employee_data['cin'],
                                   employee_data['date_cin'], employee_data['lieu_cin'], employee_data['poste'],
                                   employee_data['email'], employee_data['telephone']
                               ))

                # Insertion du contrat
                cursor.execute('''
                               INSERT INTO contrats (matricule, type_contrat, date_creation, date_debut, date_fin,
                                                     salaire_base, prime, salary_type, texte_contrat, statut)
                               VALUES (%s, %s, NOW(), %s, %s, %s, %s, %s, %s, %s)
                               ''', (
                                   contract_data['matricule'], contract_data['type_contrat'],
                                   contract_data['date_debut'], contract_data['date_fin'],
                                   contract_data['salaire_base'], contract_data['prime'],
                                   contract_data['salary_type'], contract_data['texte_contrat'],
                                   'actif' if duration_days >= 30 else 'en attente'
                               ))

                self.conn.commit()

            # Mise à jour de l'interface
            self.current_employee = {
                'matricule': employee_data['matricule'],
                'nom': employee_data['nom'],
                'prenom': employee_data['prenom'],
                'genre': employee_data['genre'],
                'date_naissance': employee_data['date_naissance'],
                'lieu_naissance': employee_data['lieu_naissance'],
                'adresse': employee_data['adresse'],
                'ville': employee_data['ville'],
                'cin': employee_data['cin'],
                'date_cin': employee_data['date_cin'],
                'lieu_cin': employee_data['lieu_cin'],
                'poste': employee_data['poste'],
                'email': employee_data['email'],
                'telephone': employee_data['telephone'],
                'texte_contrat': contract_text.strip()
            }

            self.load_data()
            self.clear_form()
            self.status_var.set("Employé et contrat enregistrés avec succès.")
            Messagebox.show_info("Enregistrement et génération réussis.", "Succès")

        except mysql.connector.Error as e:
            self.conn.rollback()
            Messagebox.show_error(f"Erreur de base de données: {str(e)}", "Erreur")
            logging.error(f"Erreur SQL dans save_and_generate: {str(e)}")
        except Exception as e:
            self.conn.rollback()
            Messagebox.show_error(f"Erreur inattendue: {str(e)}", "Erreur")
            logging.error(f"Erreur inattendue dans save_and_generate: {str(e)}")

    def initialize_database(self):
        try:
            self.conn = mysql.connector.connect(**self.db_config)
            self.create_database()
        except mysql.connector.Error as e:
            Messagebox.show_error(f"Erreur de connexion: {str(e)}", "Erreur")
            self.root.quit()
            return



    def initialize_ui(self):
        # Variables
        self.variables = {
            "genre": tk.StringVar(value="féminin"),
            "matricule": tk.StringVar(),
            "contract_type": tk.StringVar(value="CDD"),
            "salary_type": tk.StringVar(value="hourly")
        }

        # Configuration de l'interface
        main_panel = ttk.Frame(self.root)
        main_panel.pack(fill=BOTH, expand=True, padx=15, pady=15)

        self.notebook = ttk.Notebook(main_panel, bootstyle=PRIMARY)
        self.notebook.pack(fill=BOTH, expand=True)

        # Onglets
        self.create_employee_tab()
        self.create_search_tab()
        self.create_contract_tab()
        self.create_list_tab()

        # Barre de statut
        status_frame = ttk.Frame(main_panel, bootstyle=INFO)
        status_frame.pack(fill=X, pady=(10, 0))

        self.status_var = tk.StringVar(value="Prêt")
        ttk.Label(status_frame, textvariable=self.status_var, bootstyle=(INFO, INVERSE),
                  font=('Segoe UI', 10)).pack(side=LEFT, padx=10)

        # Boutons d'aide
        ttk.Button(status_frame, text="Aide", command=self.show_help,
                   bootstyle=(INFO, OUTLINE)).pack(side=RIGHT, padx=5)
        ttk.Button(status_frame, text="À propos", command=self.show_about,
                   bootstyle=(INFO, OUTLINE)).pack(side=RIGHT, padx=5)

        # Option pour les alertes sonores
        self.sound_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(
            status_frame,
            text="Alertes sonores",
            variable=self.sound_var,
            command=self.toggle_sound,
            bootstyle="round-toggle"
        ).pack(side=RIGHT, padx=5)

    def toggle_sound(self):
        self.sound_enabled = self.sound_var.get()

    def create_employee_tab(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Nouvel Employé")

        # Configuration du défilement
        canvas = tk.Canvas(frame, highlightthickness=0)
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview, bootstyle=PRIMARY)
        scrollable_frame = ttk.Frame(canvas)

        scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side=LEFT, fill=BOTH, expand=True, padx=10, pady=10)
        scrollbar.pack(side=RIGHT, fill=Y)

        # Champs du formulaire
        fields = [
            ("Matricule*", "matricule", r'^\w+$', ttk.Entry),
            ("Nom*", "nom", None, ttk.Entry),
            ("Prénom*", "prenom", None, ttk.Entry),
            ("Date Naissance (JJ/MM/AAAA)", "date_naissance", r'^\d{2}/\d{2}/\d{4}$', DateEntry),
            ("Lieu Naissance", "lieu_naissance", None, ttk.Entry),
            ("Adresse", "adresse", None, ttk.Entry),
            ("Ville", "ville", None, ttk.Entry),
            ("CIN", "cin", r'^\d{8}$', ttk.Entry),
            ("Date CIN (JJ/MM/AAAA)", "date_cin", r'^\d{2}/\d{2}/\d{4}$', DateEntry),
            ("Lieu CIN", "lieu_cin", None, ttk.Entry),
            ("Poste", "poste", None, ttk.Entry),
            ("Email", "email", r'^[^@]+@[^@]+\.[^@]+$', ttk.Entry),
            ("Téléphone", "telephone", r'^\+?\d{10,12}$', ttk.Entry)
        ]

        self.entries = {}
        form_frame = ttk.LabelFrame(scrollable_frame, text="Informations Employé", bootstyle=PRIMARY)
        form_frame.pack(fill=BOTH, padx=10, pady=10, expand=True)

        for i, (label, field, regex, widget_type) in enumerate(fields):
            ttk.Label(form_frame, text=label, font=('Segoe UI', 10)).grid(row=i, column=0, padx=5, pady=5, sticky=E)
            entry = widget_type(form_frame, bootstyle="primary") if widget_type != DateEntry else DateEntry(form_frame,
                                                                                                            bootstyle="primary",
                                                                                                            dateformat="%d/%m/%Y")
            if widget_type == DateEntry:
                entry.entry.configure(justify="center")
            entry.grid(row=i, column=1, padx=5, pady=5, sticky=EW)
            if regex and widget_type == ttk.Entry:
                add_regex_validation(entry, regex)
                entry.bind("<KeyRelease>", lambda e, f=field: self.validate_field(e.widget, f))
            self.entries[field] = entry

        # Genre
        ttk.Label(form_frame, text="Genre*", font=('Segoe UI', 10)).grid(row=len(fields), column=0, sticky=E, padx=5,
                                                                         pady=5)
        genre_frame = ttk.Frame(form_frame)
        genre_frame.grid(row=len(fields), column=1, sticky=W)
        ttk.Radiobutton(genre_frame, text="Féminin", variable=self.variables["genre"], value="féminin",
                        bootstyle="primary-toolbutton").pack(side=LEFT, padx=5)
        ttk.Radiobutton(genre_frame, text="Masculin", variable=self.variables["genre"], value="masculin",
                        bootstyle="primary-toolbutton").pack(side=LEFT, padx=5)

        # Valeurs par défaut
        self.entries['ville'].insert(0, "المحرس")
        self.entries['lieu_cin'].insert(0, "تونس")

        # Section Contrat
        contract_frame = ttk.LabelFrame(scrollable_frame, text="Détails du Contrat", bootstyle=PRIMARY)
        contract_frame.pack(fill=BOTH, padx=10, pady=10, expand=True)

        ttk.Label(contract_frame, text="Type de Contrat*", font=('Segoe UI', 10)).grid(row=0, column=0, padx=5, pady=5,
                                                                                       sticky=E)
        ttk.Radiobutton(contract_frame, text="CDD", variable=self.variables["contract_type"], value="CDD",
                        bootstyle="primary-toolbutton", command=self.toggle_date_fin).grid(row=0, column=1, sticky=W)
        ttk.Radiobutton(contract_frame, text="CDI", variable=self.variables["contract_type"], value="CDI",
                        bootstyle="primary-toolbutton", command=self.toggle_date_fin).grid(row=0, column=2, sticky=W)

        contract_fields = [
            ("Date Début (JJ/MM/AAAA)*", "date_debut", r'^\d{2}/\d{2}/\d{4}$', DateEntry),
            ("Date Fin (JJ/MM/AAAA)", "date_fin", r'^\d{2}/\d{2}/\d{4}$', DateEntry),
            ("Salaire Base*", "salaire", r'^\d+(\.\d{1,2})?$', ttk.Entry),
            ("Prime*", "prime", r'^\d+(\.\d{1,2})?$', ttk.Entry)
        ]

        self.contract_entries = {}
        for i, (label, field, regex, widget_type) in enumerate(contract_fields):
            ttk.Label(contract_frame, text=label, font=('Segoe UI', 10)).grid(row=i + 1, column=0, padx=5, pady=5,
                                                                              sticky=E)
            entry = widget_type(contract_frame, bootstyle="primary") if widget_type != DateEntry else DateEntry(
                contract_frame, bootstyle="primary", dateformat="%d/%m/%Y")
            if widget_type == DateEntry:
                entry.entry.configure(justify="center")
                if field == "date_fin" and self.variables["contract_type"].get() == "CDI":
                    entry.entry.config(state=DISABLED)
            entry.grid(row=i + 1, column=1, padx=5, pady=5, sticky=EW)
            if regex and widget_type == ttk.Entry:
                add_regex_validation(entry, regex)
                entry.bind("<KeyRelease>", lambda e, f=field: self.validate_field(e.widget, f))
            self.contract_entries[field] = entry

        # Valeurs par défaut contrat
        self.contract_entries['date_debut'].entry.delete(0, tk.END)
        self.contract_entries['date_debut'].entry.insert(0, datetime.datetime.now().strftime("%d/%m/%Y"))
        self.contract_entries['salaire'].insert(0, "2500.00")
        self.contract_entries['prime'].insert(0, "500.00")

        # Type de salaire
        ttk.Label(contract_frame, text="Type de Salaire*", font=('Segoe UI', 10)).grid(row=len(contract_fields) + 1,
                                                                                       column=0, padx=5, pady=5,
                                                                                       sticky=E)
        salary_type_frame = ttk.Frame(contract_frame)
        salary_type_frame.grid(row=len(contract_fields) + 1, column=1, sticky=W)
        ttk.Radiobutton(salary_type_frame, text="Par Heure", variable=self.variables["salary_type"], value="hourly",
                        bootstyle="primary-toolbutton").pack(side=LEFT, padx=5)
        ttk.Radiobutton(salary_type_frame, text="Par Mois", variable=self.variables["salary_type"], value="monthly",
                        bootstyle="primary-toolbutton").pack(side=LEFT, padx=5)

        # Boutons
        button_frame = ttk.Frame(scrollable_frame)
        button_frame.pack(fill=X, pady=10)
        ttk.Button(button_frame, text="Enregistrer et Générer", command=self.save_and_generate, bootstyle=SUCCESS).pack(
            side=LEFT, padx=5)
        ttk.Button(button_frame, text="Réinitialiser", command=self.clear_form, bootstyle=WARNING).pack(side=LEFT,
                                                                                                        padx=5)

        scrollable_frame.columnconfigure(1, weight=1)

        # Activation du défilement avec la molette
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

        canvas.bind_all("<MouseWheel>", on_mousewheel)

    def create_search_tab(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Recherche")

        search_frame = ttk.Frame(frame)
        search_frame.pack(fill=X, padx=10, pady=10)
        ttk.Label(search_frame, text="Matricule:", font=('Segoe UI', 10)).pack(side=LEFT)
        self.search_combo = ttk.Combobox(search_frame, textvariable=self.variables["matricule"], font=('Segoe UI', 10))
        self.search_combo.pack(side=LEFT, padx=5, expand=True, fill=X)
        ttk.Button(search_frame, text="Rechercher", command=self.search_employee, bootstyle=INFO).pack(side=LEFT,
                                                                                                       padx=5)

        info_frame = ttk.LabelFrame(frame, text="Informations Employé", bootstyle=PRIMARY)
        info_frame.pack(fill=BOTH, expand=True, padx=10, pady=5)
        self.info_text = tk.Text(info_frame, wrap=WORD, height=12, font=('Segoe UI', 10))
        scrollbar = ttk.Scrollbar(info_frame, command=self.info_text.yview, bootstyle=PRIMARY)
        self.info_text.config(yscrollcommand=scrollbar.set)
        self.info_text.pack(side=LEFT, fill=BOTH, expand=True)
        scrollbar.pack(side=RIGHT, fill=Y)

        button_frame = ttk.Frame(frame)
        button_frame.pack(fill=X, padx=10, pady=10)
        self.view_contract_btn = ttk.Button(button_frame, text="Voir Contrat", command=self.show_last_contract,
                                            bootstyle=(PRIMARY, OUTLINE), state=DISABLED)
        self.view_contract_btn.pack(side=LEFT, padx=5)
        self.edit_btn = ttk.Button(button_frame, text="Modifier Employé",
                                   command=lambda: self.edit_employee(self.current_employee['matricule']),
                                   bootstyle=(WARNING, OUTLINE), state=DISABLED)
        self.edit_btn.pack(side=LEFT, padx=5)
        self.delete_btn = ttk.Button(button_frame, text="Supprimer Employé",
                                     command=lambda: self.delete_employee(self.current_employee['matricule']),
                                     bootstyle=(DANGER, OUTLINE), state=DISABLED)
        self.delete_btn.pack(side=LEFT, padx=5)

    def create_contract_tab(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Contrat")

        self.contract_text = tk.Text(frame, wrap=WORD, font=('Arial', 11))
        self.contract_text.pack(side=LEFT, fill=BOTH, expand=True)
        scrollbar = ttk.Scrollbar(frame, command=self.contract_text.yview, bootstyle=PRIMARY)
        self.contract_text.config(yscrollcommand=scrollbar.set)
        scrollbar.pack(side=RIGHT, fill=Y)

        self.contract_text.tag_configure("rtl", justify="right")
        self.contract_text.insert(tk.END, "", "rtl")

        button_frame = ttk.Frame(frame)
        button_frame.pack(fill=X, padx=10, pady=10)
        ttk.Button(button_frame, text="Exporter Word", command=self.export_word, bootstyle=SUCCESS).pack(side=LEFT,
                                                                                                         padx=5)
        ttk.Button(button_frame, text="Copier", command=self.copy_contract, bootstyle=INFO).pack(side=LEFT, padx=5)
        ttk.Button(button_frame, text="Imprimer", command=self.print_contract, bootstyle=PRIMARY).pack(side=LEFT,
                                                                                                       padx=5)

    def create_list_tab(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Liste Employés")

        summary_frame = ttk.Frame(frame, bootstyle=INFO)
        summary_frame.pack(fill=X, padx=10, pady=5)

        stats_frame = ttk.Frame(summary_frame)
        stats_frame.pack(side=LEFT, fill=X, expand=True)

        self.total_label = ttk.Label(stats_frame, text="Total: 0", font=('Segoe UI', 10))
        self.total_label.pack(side=LEFT, padx=10)
        self.cdd_label = ttk.Label(stats_frame, text="CDD: 0", font=('Segoe UI', 10))
        self.cdd_label.pack(side=LEFT, padx=10)
        self.cdi_label = ttk.Label(stats_frame, text="CDI: 0", font=('Segoe UI', 10))
        self.cdi_label.pack(side=LEFT, padx=10)
        self.salary_label = ttk.Label(stats_frame, text="Salaire Moyen: 0.00 TND", font=('Segoe UI', 10))
        self.salary_label.pack(side=LEFT, padx=10)

        refresh_btn = ttk.Button(summary_frame, text="🔄 Actualiser", command=self.load_employee_table,
                                 bootstyle=(INFO, OUTLINE))
        refresh_btn.pack(side=RIGHT, padx=5)

        filter_frame = ttk.Frame(frame)
        filter_frame.pack(fill=X, padx=10, pady=5)
        ttk.Label(filter_frame, text="Filtrer par:", font=('Segoe UI', 10)).pack(side=LEFT)
        self.filter_var = tk.StringVar()
        self.filter_combo = ttk.Combobox(filter_frame, textvariable=self.filter_var,
                                         values=["Nom", "Matricule", "Type Contrat"],
                                         font=('Segoe UI', 10))
        self.filter_combo.pack(side=LEFT, padx=5)
        self.filter_entry = ttk.Entry(filter_frame, font=('Segoe UI', 10))
        self.filter_entry.pack(side=LEFT, padx=5, expand=True, fill=X)
        ttk.Button(filter_frame, text="Filtrer", command=self.apply_filter, bootstyle=INFO).pack(side=LEFT, padx=5)
        ttk.Button(filter_frame, text="Réinitialiser", command=self.reset_filter, bootstyle=WARNING).pack(side=LEFT,
                                                                                                          padx=5)
        ttk.Button(filter_frame, text="Exporter Tous", command=self.export_all_contracts, bootstyle=SUCCESS).pack(
            side=LEFT, padx=5)
        ttk.Button(filter_frame, text="Annuler", command=self.undo_action, bootstyle=(WARNING, OUTLINE)).pack(side=LEFT,
                                                                                                              padx=5)

        self.employee_table = Tableview(
            frame, coldata=self.get_column_definitions(), rowdata=[], paginated=True, searchable=True,
            bootstyle=PRIMARY, autoalign=True, stripecolor=('lightblue', None), pagesize=20
        )
        self.employee_table.pack(fill=BOTH, expand=True, padx=10, pady=10)

        self.employee_table.view.bind("<Double-1>", self.edit_cell)
        self.context_menu = tk.Menu(self.root, tearoff=0, font=('Segoe UI', 10))
        self.context_menu.add_command(label="Modifier", command=self.context_menu_modify)
        self.context_menu.add_command(label="Supprimer", command=self.context_menu_delete)
        self.context_menu.add_command(label="Voir Contrat", command=self.context_menu_view_contract)
        self.context_menu.add_command(label="Exporter Contrat", command=self.context_menu_export_contract)
        self.employee_table.view.bind("<Button-3>", self.show_context_menu)

    def get_column_definitions(self):
        return [
            {"text": "Matricule", "stretch": False, "width": 100},
            {"text": "Nom", "stretch": True, "width": 150},
            {"text": "Prénom", "stretch": True, "width": 150},
            {"text": "Genre", "stretch": False, "width": 80},
            {"text": "Date Naissance", "stretch": True, "width": 120},
            {"text": "Lieu Naissance", "stretch": True, "width": 150},
            {"text": "Adresse", "stretch": True, "width": 200},
            {"text": "Ville", "stretch": True, "width": 100},
            {"text": "CIN", "stretch": True, "width": 100},
            {"text": "Date CIN", "stretch": True, "width": 120},
            {"text": "Lieu CIN", "stretch": True, "width": 150},
            {"text": "Poste", "stretch": True, "width": 150},
            {"text": "Email", "stretch": True, "width": 200},
            {"text": "Téléphone", "stretch": True, "width": 120},
            {"text": "Type Contrat", "stretch": True, "width": 100},
            {"text": "Date Début", "stretch": True, "width": 120},
            {"text": "Date Fin", "stretch": True, "width": 120},
            {"text": "Salaire Base", "stretch": True, "width": 100},
            {"text": "Prime", "stretch": True, "width": 100},
            {"text": "Type Salaire", "stretch": True, "width": 100},
        ]

    def toggle_date_fin(self):
        state = DISABLED if self.variables["contract_type"].get() == "CDI" else NORMAL
        self.contract_entries['date_fin'].entry.config(state=state)

    def load_data(self):
        self.load_matricules()
        self.load_employee_table()

    def load_matricules(self):
        with self.conn.cursor() as cursor:
            cursor.execute("SELECT matricule FROM employes ORDER BY matricule")
            self.search_combo['values'] = [row[0] for row in cursor.fetchall()]

    def load_employee_table(self):
        with self.conn.cursor() as cursor:
            cursor.execute('''
                           SELECT e.matricule,
                                  e.nom,
                                  e.prenom,
                                  e.genre,
                                  e.date_naissance,
                                  e.lieu_naissance,
                                  e.adresse,
                                  e.ville,
                                  e.cin,
                                  e.date_cin,
                                  e.lieu_cin,
                                  e.poste,
                                  e.email,
                                  e.telephone,
                                  c.type_contrat,
                                  c.date_debut,
                                  c.date_fin,
                                  c.salaire_base,
                                  c.prime,
                                  c.salary_type
                           FROM employes e
                                    LEFT JOIN contrats c ON e.matricule = c.matricule
                               AND c.id = (SELECT MAX(id) FROM contrats WHERE matricule = e.matricule)
                           ''')
            self.update_table_data(cursor.fetchall())
            self.status_var.set(f"{len(self.employee_table.get_rows())} employés chargés")

    def update_table_data(self, rows):
        self.employee_table.delete_rows()
        today = datetime.datetime.now().date()

        for row in rows:
            matricule = row[0]
            action_frame = ttk.Frame(self.employee_table.view)

            warning = ""
            if row[16]:  # Si date_fin existe
                try:
                    end_date = datetime.datetime.strptime(row[16], "%d/%m/%Y").date()
                    days_left = (end_date - today).days
                    if 0 <= days_left <= 30:
                        warning = "⚠️ "
                except ValueError:
                    pass

            ttk.Button(action_frame, text="Modifier",
                       command=lambda m=matricule: self.edit_employee(m),
                       bootstyle=(PRIMARY, OUTLINE), width=8).pack(side=LEFT, padx=2)
            ttk.Button(action_frame, text="Supprimer",
                       command=lambda m=matricule: self.delete_employee(m),
                       bootstyle=(DANGER, OUTLINE), width=8).pack(side=LEFT, padx=2)
            ttk.Button(action_frame, text="Contrat",
                       command=lambda m=matricule: self.view_contract_from_table(m),
                       bootstyle=(INFO, OUTLINE), width=8).pack(side=LEFT, padx=2)

            display_row = list(row[:14])
            display_row[1] = warning + display_row[1]

            display_row += [
                row[14] or "N/A",
                row[15] or "N/A",
                row[16] or "N/A",
                str(row[17]) if row[17] is not None else "N/A",
                str(row[18]) if row[18] is not None else "N/A",
                row[19] or "N/A",
                action_frame
            ]
            self.employee_table.insert_row(values=display_row)

        self.update_summary()

    def update_summary(self):
        try:
            with self.conn.cursor() as cursor:
                cursor.execute("SELECT COUNT(DISTINCT matricule) FROM employes")
                total_employees = cursor.fetchone()[0]

                cursor.execute('''
                               SELECT COUNT(CASE WHEN c.type_contrat = 'CDD' THEN 1 END) AS cdd_count,
                                      COUNT(CASE WHEN c.type_contrat = 'CDI' THEN 1 END) AS cdi_count,
                                      AVG(c.salaire_base)                                AS avg_salary
                               FROM (SELECT matricule, MAX(id) AS max_id
                                     FROM contrats
                                     GROUP BY matricule) AS latest
                                        JOIN contrats c ON c.id = latest.max_id
                               ''')
                stats = cursor.fetchone()
                cdd_count = stats[0] or 0
                cdi_count = stats[1] or 0
                avg_salary = stats[2] or 0

                today = datetime.datetime.now().date()
                cursor.execute('''
                               SELECT COUNT(*)
                               FROM (SELECT c.matricule, MAX(c.id) AS max_id
                                     FROM contrats c
                                     WHERE c.type_contrat = 'CDD'
                                     GROUP BY c.matricule) AS latest
                                        JOIN contrats c ON c.id = latest.max_id
                               WHERE c.date_fin IS NOT NULL
                                 AND DATEDIFF(STR_TO_DATE(c.date_fin, '%d/%m/%Y'), %s) <= 30
                                 AND DATEDIFF(STR_TO_DATE(c.date_fin, '%d/%m/%Y'), %s) >= 0
                               ''', (today, today))
                expiring_soon = cursor.fetchone()[0] or 0

                if expiring_soon > 0:
                    self.play_alert_sound()
                    cursor.execute('''
                                   SELECT e.matricule,
                                          e.nom,
                                          e.prenom,
                                          c.date_fin,
                                          DATEDIFF(STR_TO_DATE(c.date_fin, '%d/%m/%Y'), %s) AS jours_restants
                                   FROM employes e
                                            JOIN contrats c ON e.matricule = c.matricule
                                   WHERE c.type_contrat = 'CDD'
                                     AND c.date_fin IS NOT NULL
                                     AND DATEDIFF(STR_TO_DATE(c.date_fin, '%d/%m/%Y'), %s) BETWEEN 0 AND 30
                                   ORDER BY jours_restants
                                   ''', (today, today))
                    expiring_contracts = cursor.fetchall()

                    message = "⚠️ ALERTE : Contrats CDD expirant bientôt ⚠️\n\n"
                    for contract in expiring_contracts:
                        matricule, nom, prenom, date_fin, jours_restants = contract
                        message += f"• {nom} {prenom} (Matricule: {matricule}) - "
                        message += f"Expire le {date_fin} (dans {jours_restants} jours)\n"

                    Messagebox.show_warning(message, "Alerte Contrats", parent=self.root)
                    self.stop_alert_sound()
                    self.status_var.set(f"⚠ {expiring_soon} contrats expirent dans ≤30 jours")
                else:
                    self.status_var.set("Aucun contrat expirant bientôt")

            self.total_label.config(text=f"Employés: {total_employees}")
            self.cdd_label.config(text=f"CDD Actifs: {cdd_count}")
            self.cdi_label.config(text=f"CDI Actifs: {cdi_count}")
            self.salary_label.config(text=f"Salaire Moyen: {avg_salary:.2f} TND")

        except mysql.connector.Error as e:
            print(f"Erreur lors de la mise à jour du résumé: {e}")
            self.status_var.set("Erreur lors de la mise à jour des statistiques")

    def apply_filter(self):
        filter_type = self.filter_var.get()
        filter_value = self.filter_entry.get().strip().lower()
        if not filter_type or not filter_value:
            self.load_employee_table()
            return

        query = '''
                SELECT e.matricule, \
                       e.nom, \
                       e.prenom, \
                       e.genre, \
                       e.date_naissance, \
                       e.lieu_naissance,
                       e.adresse, \
                       e.ville, \
                       e.cin, \
                       e.date_cin, \
                       e.lieu_cin, \
                       e.poste, \
                       e.email, \
                       e.telephone,
                       c.type_contrat, \
                       c.date_debut, \
                       c.date_fin, \
                       c.salaire_base, \
                       c.prime, \
                       c.salary_type
                FROM employes e
                         LEFT JOIN contrats c ON e.matricule = c.matricule
                    AND c.id = (SELECT MAX(id) FROM contrats WHERE matricule = e.matricule)
                WHERE {} \
                '''
        params = (f"%{filter_value}%",) if filter_type != "Type Contrat" else (filter_value.upper(),)
        condition = {
            "Nom": "LOWER(e.nom) LIKE %s",
            "Matricule": "e.matricule LIKE %s",
            "Type Contrat": "c.type_contrat = %s"
        }.get(filter_type, "1=1")

        with self.conn.cursor() as cursor:
            cursor.execute(query.format(condition), params)
            self.update_table_data(cursor.fetchall())
            self.status_var.set(f"{len(self.employee_table.get_rows())} employés trouvés")

    def reset_filter(self):
        self.filter_var.set("")
        self.filter_entry.delete(0, tk.END)
        self.load_employee_table()

    def show_context_menu(self, event):
        row_id = self.employee_table.view.identify_row(event.y)
        if row_id:
            self.employee_table.view.selection_set(row_id)
            self.selected_matricule = self.employee_table.get_row(row_id).values[0]
            self.context_menu.post(event.x_root, event.y_root)

    def context_menu_modify(self):
        self.edit_employee(self.selected_matricule)

    def context_menu_delete(self):
        self.delete_employee(self.selected_matricule)

    def context_menu_view_contract(self):
        self.current_employee = {'matricule': self.selected_matricule}
        self.show_last_contract()

    def context_menu_export_contract(self):
        self.current_employee = {'matricule': self.selected_matricule}
        self.show_last_contract()
        self.export_word()

    def clear_form(self):
        for entry in self.entries.values():
            if isinstance(entry, ttk.Entry):
                entry.delete(0, tk.END)
            elif isinstance(entry, DateEntry):
                entry.entry.delete(0, tk.END)

        for entry in self.contract_entries.values():
            if isinstance(entry, ttk.Entry):
                entry.delete(0, tk.END)
            elif isinstance(entry, DateEntry):
                entry.entry.delete(0, tk.END)

        self.entries['ville'].insert(0, "المحرس")
        self.entries['lieu_cin'].insert(0, "تونس")
        self.contract_entries['date_debut'].entry.insert(0, datetime.datetime.now().strftime("%d/%m/%Y"))
        self.contract_entries['salaire'].insert(0, "2500.00")
        self.contract_entries['prime'].insert(0, "500.00")
        self.variables["genre"].set("féminin")
        self.variables["contract_type"].set("CDD")
        self.variables["salary_type"].set("hourly")
        self.status_var.set("Formulaire réinitialisé")

    def search_employee(self):
        matricule = self.variables["matricule"].get()
        if not matricule:
            Messagebox.show_warning("Veuillez saisir un matricule", "Attention")
            return

        with self.conn.cursor() as cursor:
            cursor.execute('''
                           SELECT e.*, c.type_contrat, c.date_debut, c.date_fin, c.salaire_base, c.prime, c.salary_type
                           FROM employes e
                                    LEFT JOIN contrats c ON e.matricule = c.matricule
                               AND c.id = (SELECT MAX(id) FROM contrats WHERE matricule = e.matricule)
                           WHERE e.matricule = %s
                           ''', (matricule,))
            employee = cursor.fetchone()

        if employee:
            self.current_employee = {
                'matricule': employee[0], 'nom': employee[1], 'prenom': employee[2], 'genre': employee[3],
                'date_naissance': employee[4], 'lieu_naissance': employee[5], 'adresse': employee[6],
                'ville': employee[7], 'cin': employee[8], 'date_cin': employee[9], 'lieu_cin': employee[10],
                'poste': employee[11], 'email': employee[12], 'telephone': employee[13],
                'type_contrat': employee[14] or "N/A", 'date_debut': employee[15] or "N/A",
                'date_fin': employee[16] or "N/A",
                'salaire_base': str(employee[17]) if employee[17] is not None else "N/A",
                'prime': str(employee[18]) if employee[18] is not None else "N/A", 'salary_type': employee[19] or "N/A"
            }
            info = f"""Matricule: {employee[0]}
Nom: {employee[1]} {employee[2]}
Genre: {employee[3]}
Date Naissance: {employee[4] or 'N/A'} à {employee[5] or 'N/A'}
Adresse: {employee[6] or 'N/A'}, {employee[7] or 'N/A'}
CIN: {employee[8] or 'N/A'} (délivré le {employee[9] or 'N/A'} à {employee[10] or 'N/A'})
Poste: {employee[11] or 'N/A'}
Email: {employee[12] or 'N/A'}
Téléphone: {employee[13] or 'N/A'}
Type Contrat: {employee[14] or 'N/A'}
Date Début: {employee[15] or 'N/A'}
Date Fin: {employee[16] or 'N/A'}
Salaire Base: {employee[17] or 'N/A'} TND
Prime: {employee[18] or 'N/A'} TND
Type Salaire: {employee[19] or 'N/A'}"""
            self.info_text.delete(1.0, tk.END)
            self.info_text.insert(tk.END, info)
            self.view_contract_btn.config(state=NORMAL)
            self.edit_btn.config(state=NORMAL)
            self.delete_btn.config(state=NORMAL)
            self.status_var.set(f"Employé trouvé: {employee[1]} {employee[2]}")
        else:
            self.clear_search()

    def clear_search(self):
        self.info_text.delete(1.0, tk.END)
        self.current_employee = None
        self.view_contract_btn.config(state=DISABLED)
        self.edit_btn.config(state=DISABLED)
        self.delete_btn.config(state=DISABLED)
        Messagebox.show_info("Aucun employé trouvé", "Information")
        self.status_var.set("Aucun résultat")

    def show_last_contract(self):
        if not self.current_employee:
            Messagebox.show_warning("Veuillez rechercher un employé", "Attention")
            return

        with self.conn.cursor() as cursor:
            cursor.execute("""
                           SELECT texte_contrat, type_contrat, date_debut, date_fin, salaire_base, prime, salary_type
                           FROM contrats
                           WHERE matricule = %s
                           ORDER BY id DESC LIMIT 1
                           """, (self.current_employee['matricule'],))
            contract = cursor.fetchone()

        if contract:
            self.contract_text.delete(1.0, tk.END)
            self.contract_text.insert(tk.END, contract[0])
            self.variables["contract_type"].set(contract[1])
            self.variables["salary_type"].set(contract[6])

            self.contract_entries['date_debut'].entry.delete(0, tk.END)
            self.contract_entries['date_debut'].entry.insert(0, contract[2] or "")

            self.contract_entries['date_fin'].entry.delete(0, tk.END)
            self.contract_entries['date_fin'].entry.insert(0, contract[3] or "")

            self.contract_entries['salaire'].delete(0, tk.END)
            self.contract_entries['salaire'].insert(0, str(contract[4]) if contract[4] is not None else "2500.00")

            self.contract_entries['prime'].delete(0, tk.END)
            self.contract_entries['prime'].insert(0, str(contract[5]) if contract[5] is not None else "500.00")

            self.notebook.select(2)
            self.status_var.set("Dernier contrat affiché")
        else:
            Messagebox.show_info("Aucun contrat trouvé", "Information")
            self.status_var.set("Aucun contrat")

    def view_contract_from_table(self, matricule):
        self.current_employee = {'matricule': matricule}
        self.show_last_contract()

    def export_word(self):
        if not self.current_employee or not isinstance(self.current_employee, dict):
            Messagebox.show_error("Aucun employé sélectionné", "Erreur")
            return

        try:
            doc = self.create_contract_doc(self.current_employee['matricule'],
                                           self.contract_text.get(1.0, tk.END).strip())

            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Document Word", "*.docx")],
                initialfile=f"Contrat_{self.current_employee['matricule']}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )
            if not file_path:
                return

            doc.save(file_path)
            Messagebox.show_info(f"Contrat exporté avec succès sous {file_path}", "Succès")
            self.status_var.set("Contrat exporté avec succès.")
        except Exception as e:
            Messagebox.show_error(f"Erreur lors de l'exportation: {str(e)}", "Erreur")
            logging.error(f"Erreur dans export_word: {str(e)}")

    def export_all_contracts(self):
        with self.conn.cursor() as cursor:
            cursor.execute("""
                           SELECT matricule, texte_contrat
                           FROM contrats
                           WHERE id IN (SELECT MAX(id) FROM contrats GROUP BY matricule)
                           """)
            contracts = cursor.fetchall()

        if not contracts:
            Messagebox.show_warning("Aucun contrat à exporter", "Attention")
            return

        filepath = filedialog.asksaveasfilename(
            defaultextension=".zip",
            filetypes=[("Archive ZIP", "*.zip")],
            initialfile=f"Contrats_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
            title="Enregistrer tous les contrats"
        )
        if filepath:
            try:
                with zipfile.ZipFile(filepath, 'w', zipfile.ZIP_DEFLATED) as zipf:
                    for matricule, texte in contracts:
                        doc = self.create_contract_doc(matricule, texte)
                        doc_io = BytesIO()
                        doc.save(doc_io)
                        doc_io.seek(0)
                        zipf.writestr(f"Contrat_{matricule}.docx", doc_io.read())
                self.status_var.set(f"Exportation réussie: {filepath}")
                Messagebox.show_info(f"Tous les contrats exportés: {filepath}", "Succès")
            except Exception as e:
                Messagebox.show_error(f"Erreur d'exportation: {str(e)}", "Erreur")

    def copy_contract(self):
        self.root.clipboard_clear()
        self.root.clipboard_append(self.contract_text.get(1.0, tk.END))
        self.status_var.set("Texte du contrat copié")

    def print_contract(self):
        if not self.contract_text.get(1.0, tk.END).strip():
            Messagebox.show_warning("Aucun contrat à imprimer", "Attention")
            return

        try:
            temp_dir = tempfile.gettempdir()
            temp_file = os.path.join(temp_dir, f"contrat_temp_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")

            doc = self.create_contract_doc(
                self.current_employee['matricule'],
                self.contract_text.get(1.0, tk.END).strip()
            )
            doc.save(temp_file)

            try:
                import win32print
                import win32api

                printer_name = win32print.GetDefaultPrinter()
                win32api.ShellExecute(
                    0,
                    "print",
                    temp_file,
                    f'/d:"{printer_name}"',
                    temp_dir,
                    0
                )
                self.status_var.set(f"Contrat envoyé à l'imprimante {printer_name}")

            except ImportError:
                if os.name == 'posix':
                    import subprocess
                    subprocess.run(['lpr', temp_file])
                    self.status_var.set("Contrat envoyé à l'imprimante par défaut")
                else:
                    os.startfile(temp_file, "print")
                    self.status_var.set("Ouverture du contrat pour impression")

            try:
                if os.name == 'nt':
                    os.startfile(temp_file)
                else:
                    import subprocess
                    subprocess.run(['xdg-open', temp_file])
                Messagebox.show_info(
                    f"Le contrat a été envoyé à l'imprimante et ouvert pour aperçu.\n"
                    f"Fichier temporaire: {temp_file}",
                    "Impression"
                )
            except Exception as preview_error:
                Messagebox.show_warning(
                    f"Contrat imprimé, mais erreur lors de l'ouverture de l'aperçu:\n{str(preview_error)}\n"
                    f"Fichier temporaire: {temp_file}",
                    "Avertissement"
                )
        except Exception as e:
            Messagebox.show_error(
                f"Erreur lors de l'impression:\n{str(e)}",
                "Erreur d'impression"
            )

    def edit_cell(self, event):
        row_id = self.employee_table.view.identify_row(event.y)
        column = self.employee_table.view.identify_column(event.x)
        if not row_id or not column:
            return

        col_idx = int(column.replace("#", "")) - 1
        col_name = self.get_column_definitions()[col_idx]["text"]
        if col_name in ["Matricule", "Actions"]:
            return

        row_data = self.employee_table.get_row(row_id).values
        matricule = row_data[0]
        current_value = row_data[col_idx]

        entry = ttk.Entry(self.employee_table.view, bootstyle="primary", font=('Segoe UI', 10))
        entry.insert(0, current_value)
        entry.place(x=event.x_root - self.employee_table.view.winfo_rootx(),
                    y=event.y_root - self.employee_table.view.winfo_rooty(),
                    anchor="nw")

        def validate_input(value):
            validators = {
                "Date Naissance": r'^\d{2}/\d{2}/\d{4}$', "Date CIN": r'^\d{2}/\d{2}/\d{4}$',
                "Date Début": r'^\d{2}/\d{2}/\d{4}$', "Date Fin": r'^\d{2}/\d{2}/\d{4}$',
                "Email": r'^[^@]+@[^@]+\.[^@]+$', "Genre": r'^(féminin|masculin)$',
                "Type Contrat": r'^(CDD|CDI)$', "Salaire Base": r'^\d+(\.\d{1,2})?$',
                "Prime": r'^\d+(\.\d{1,2})?$', "Type Salaire": r'^(hourly|monthly)$'
            }
            return bool(re.match(validators.get(col_name, r'.*'), value)) and \
                (float(value) > 0 if col_name in ["Salaire Base", "Prime"] and value else True)

        def save_edit(event=None):
            new_value = entry.get()
            if not validate_input(new_value):
                Messagebox.show_error(f"Valeur invalide pour {col_name}", "Erreur")
                entry.destroy()
                return

            try:
                with self.conn.cursor() as cursor:
                    sql_field = {
                        "Nom": "nom", "Prénom": "prenom", "Genre": "genre",
                        "Date Naissance": "date_naissance", "Lieu Naissance": "lieu_naissance",
                        "Adresse": "adresse", "Ville": "ville", "CIN": "cin",
                        "Date CIN": "date_cin", "Lieu CIN": "lieu_cin", "Poste": "poste",
                        "Email": "email", "Téléphone": "telephone", "Type Contrat": "type_contrat",
                        "Date Début": "date_debut", "Date Fin": "date_fin",
                        "Salaire Base": "salaire_base", "Prime": "prime",
                        "Type Salaire": "salary_type"
                    }.get(col_name)

                    if sql_field:
                        if col_name in ["Type Contrat", "Date Début", "Date Fin",
                                        "Salaire Base", "Prime", "Type Salaire"]:
                            cursor.execute("""
                                           SELECT id
                                           FROM contrats
                                           WHERE matricule = %s
                                           ORDER BY id DESC LIMIT 1
                                           """, (matricule,))
                            contract_id = cursor.fetchone()
                            if contract_id:
                                cursor.execute(f"""
                                    UPDATE contrats 
                                    SET {sql_field} = %s 
                                    WHERE id = %s
                                """, (float(new_value) if col_name in ["Salaire Base", "Prime"] else new_value,
                                      contract_id[0]))
                                self.undo_stack.append(("contract_update", matricule, contract_id[0],
                                                        sql_field, current_value))
                            else:
                                Messagebox.show_warning("Aucun contrat trouvé", "Attention")
                                entry.destroy()
                                return
                        else:
                            cursor.execute(f"""
                                UPDATE employes 
                                SET {sql_field} = %s 
                                WHERE matricule = %s
                            """, (new_value, matricule))
                            self.undo_stack.append(("employee_update", matricule, sql_field, current_value))

                        self.conn.commit()
                        self.load_employee_table()
                        self.status_var.set(f"Champ {col_name} mis à jour pour {matricule}")
                        if self.current_employee and self.current_employee['matricule'] == matricule:
                            self.search_employee()
            except Exception as e:
                Messagebox.show_error(f"Erreur de mise à jour: {str(e)}", "Erreur")
            finally:
                entry.destroy()

        entry.bind("<Return>", save_edit)
        entry.bind("<FocusOut>", save_edit)
        entry.focus_set()

    def validate_field(self, widget, field):
        value = widget.get()
        validators = {
            "matricule": lambda v: bool(re.match(r'^\w+$', v)) if v else False,
            "cin": lambda v: bool(re.match(r'^\d{8}$', v)) if v else True,
            "date_naissance": lambda v: self.is_valid_date(v) if v else True,
            "date_cin": lambda v: self.is_valid_date(v) if v else True,
            "email": lambda v: bool(re.match(r'^[^@]+@[^@]+\.[^@]+$', v)) if v else True,
            "telephone": lambda v: bool(re.match(r'^\+?\d{10,12}$', v)) if v else True,
            "salaire": lambda v: bool(re.match(r'^\d+(\.\d{1,2})?$', v)) and float(v) > 0 if v else False,
            "prime": lambda v: bool(re.match(r'^\d+(\.\d{1,2})?$', v)) and float(v) >= 0 if v else False,
            "date_debut": lambda v: self.is_valid_date(v) if v else False,
            "date_fin": lambda v: self.is_valid_date(v) if v else True
        }
        widget.configure(bootstyle="danger" if not validators.get(field, lambda x: True)(value) else "primary")

    def is_valid_date(self, date_str):
        try:
            if not date_str or not re.match(r'^\d{2}/\d{2}/\d{4}$', date_str.strip()):
                return False
            datetime.datetime.strptime(date_str.strip(), "%d/%m/%Y")
            return True
        except ValueError:
            return False


    def create_contract_doc(self, matricule, texte):
        doc = Document()
        section = doc.sections[0]
        section.left_margin = section.right_margin = Inches(8 / 25.4)  # 8mm
        section.top_margin = section.bottom_margin = Inches(5 / 25.4)  # 5mm
        section.is_right_to_left = True

        # Tableau principal
        main_table = doc.add_table(rows=1, cols=3)
        main_table.columns[0].width = Inches(15 / 25.4)  # 15mm
        main_table.columns[1].width = Inches(80 / 25.4)  # 80mm
        main_table.columns[2].width = Inches(40 / 25.4)  # 40mm
        main_table.rows[0].height = Inches(15 / 25.4)  # 15mm
        main_table.style = 'Table Grid'

        # Cellule Logo
        logo_cell = main_table.cell(0, 0)
        try:
            logo_para = logo_cell.add_paragraph()
            logo_run = logo_para.add_run()
            logo_run.add_picture(self.logo_path, width=Inches(12 / 25.4), height=Inches(12 / 25.4))
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"Erreur lors du chargement du logo: {e}")
            logo_cell.text = ""

        # Cellule Titre
        title_cell = main_table.cell(0, 1)
        title_para = title_cell.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        french_title = title_para.add_run("FORMULAIRE\n")
        french_title.bold = True
        french_title.font.name = "Helvetica"
        french_title.font.size = Pt(8)
        french_title.font.rtl = True

        arabic_title = title_para.add_run(
            f"عقد شغل لمدة {'غير محددة' if self.variables['contract_type'].get() == 'CDI' else 'محدودة'}\n")
        arabic_title.bold = True
        arabic_title.font.name = "Arial"
        arabic_title.font.size = Pt(8)
        arabic_title.font.rtl = True

        # Tableau d'informations
        info_cell = main_table.cell(0, 2)
        info_table = info_cell.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        info_table.columns[0].width = info_table.columns[1].width = Inches(15 / 25.4)

        info_data = [
            ("Réf.", f"FO-RH-0{4 if self.variables['contract_type'].get() == 'CDI' else 3}"),
            ("Date", datetime.datetime.now().strftime('%d/%m/%Y')),
            ("Version", "01"),
            ("Page", "1/1")
        ]

        for row_idx, (label, value) in enumerate(info_data):
            label_cell = info_table.cell(row_idx, 0)
            label_para = label_cell.add_paragraph(label)
            label_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            label_run = label_para.runs[0]
            label_run.font.name = "Helvetica"
            label_run.font.size = Pt(7)
            label_run.font.color.rgb = RGBColor(100, 100, 100)
            label_run.font.rtl = True

            value_cell = info_table.cell(row_idx, 1)
            value_para = value_cell.add_paragraph(value)
            value_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            value_run = value_para.runs[0]
            value_run.font.name = "Helvetica"
            value_run.font.size = Pt(7)
            value_run.font.rtl = True

        # Matricule
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        matricule_para = doc.add_paragraph()
        matricule_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        matricule_run = matricule_para.add_run(f"Matricule: {matricule}")
        matricule_run.bold = True
        matricule_run.font.name = "Arial"
        matricule_run.font.size = Pt(9)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

        # Ajout du texte du contrat
        for paragraph in texte.strip().split('\n'):
            if paragraph.strip():
                body_para = doc.add_paragraph()
                body_para.paragraph_format.space_after = Pt(2)
                body_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                body_run = body_para.add_run(paragraph.strip())
                body_run.font.name = "Arial"
                body_run.font.size = Pt(9)
                body_run.font.rtl = True

        return doc

    def edit_employee(self, matricule):
        top = ttk.Toplevel(self.root)
        top.title(f"Modifier Employé {matricule}")
        top.geometry("900x900")

        with self.conn.cursor() as cursor:
            cursor.execute("SELECT * FROM employes WHERE matricule = %s", (matricule,))
            employee = cursor.fetchone()
            cursor.execute(
                "SELECT type_contrat, date_debut, date_fin, salaire_base, prime, salary_type FROM contrats WHERE matricule = %s ORDER BY id DESC LIMIT 1",
                (matricule,))
            contract = cursor.fetchone()

        if not employee:
            Messagebox.show_error("Employé non trouvé", "Erreur")
            top.destroy()
            return

        notebook = ttk.Notebook(top, bootstyle=PRIMARY)
        notebook.pack(fill=BOTH, expand=True, padx=10, pady=10)

        # Frame pour les informations employé
        employee_frame = ttk.Frame(notebook)
        notebook.add(employee_frame, text="Détails Employé")

        # Frame pour les détails du contrat
        contract_frame = ttk.Frame(notebook)
        notebook.add(contract_frame, text="Détails Contrat")

        # Variables pour les radiobuttons
        genre_var = tk.StringVar(value=employee[3])
        contract_type_var = tk.StringVar(value=contract[0] if contract else "CDD")
        salary_type_var = tk.StringVar(value=contract[5] if contract else "hourly")

        # Dictionnaires pour stocker les références aux widgets
        entries = {}
        contract_entries = {}

        # Champs employé
        fields = [
            ("Matricule", "matricule", employee[0], True, ttk.Entry),
            ("Nom*", "nom", employee[1], False, ttk.Entry),
            ("Prénom*", "prenom", employee[2], False, ttk.Entry),
        ]

        for i, (label, field, value, disabled, widget_type) in enumerate(fields):
            ttk.Label(employee_frame, text=label, font=('Segoe UI', 10)).grid(row=i, column=0, padx=5, pady=5, sticky=E)
            entry = widget_type(employee_frame, bootstyle="primary")
            entry.insert(0, value)
            if disabled:
                entry.config(state='disabled')
            entry.grid(row=i, column=1, padx=5, pady=5, sticky=EW)
            entries[field] = entry

        # Champs dates (utilisant DateEntry)
        date_fields = [
            ("Date Naissance (JJ/MM/AAAA)", "date_naissance", employee[4], False, employee_frame),
            ("Date CIN (JJ/MM/AAAA)", "date_cin", employee[9], False, employee_frame),
            ("Date Début (JJ/MM/AAAA)*", "date_debut", contract[1] if contract else "", False, contract_frame),
            ("Date Fin (JJ/MM/AAAA)", "date_fin", contract[2] if contract else "", contract_type_var.get() == "CDI",
             contract_frame),
        ]

        for i, (label, field, value, disabled, frame) in enumerate(date_fields):
            ttk.Label(frame, text=label, font=('Segoe UI', 10)).grid(row=i, column=0, padx=5, pady=5, sticky=E)
            entry = DateEntry(frame, bootstyle="primary", dateformat="%d/%m/%Y")
            if value:
                entry.entry.delete(0, tk.END)
                entry.entry.insert(0, value)
            if disabled:
                entry.entry.config(state='disabled')
            entry.grid(row=i, column=1, padx=5, pady=5, sticky=EW)

            if frame == employee_frame:
                entries[field] = entry
            else:
                contract_entries[field] = entry

        # Autres champs employé
        other_fields = [
            ("Lieu Naissance", "lieu_naissance", employee[5], False, ttk.Entry),
            ("Adresse", "adresse", employee[6], False, ttk.Entry),
            ("Ville", "ville", employee[7], False, ttk.Entry),
            ("CIN", "cin", employee[8], False, ttk.Entry),
            ("Lieu CIN", "lieu_cin", employee[10], False, ttk.Entry),
            ("Poste", "poste", employee[11], False, ttk.Entry),
            ("Email", "email", employee[12], False, ttk.Entry),
            ("Téléphone", "telephone", employee[13], False, ttk.Entry),
        ]

        start_row = len(date_fields) if frame == employee_frame else 0
        for i, (label, field, value, disabled, widget_type) in enumerate(other_fields, start=start_row):
            ttk.Label(employee_frame, text=label, font=('Segoe UI', 10)).grid(row=i, column=0, padx=5, pady=5, sticky=E)
            entry = widget_type(employee_frame, bootstyle="primary")
            entry.insert(0, value or "")
            entry.grid(row=i, column=1, padx=5, pady=5, sticky=EW)
            entries[field] = entry

        # Genre
        ttk.Label(employee_frame, text="Genre*", font=('Segoe UI', 10)).grid(row=len(other_fields) + start_row,
                                                                             column=0, sticky=E, padx=5, pady=5)
        genre_frame = ttk.Frame(employee_frame)
        genre_frame.grid(row=len(other_fields) + start_row, column=1, sticky=W)
        ttk.Radiobutton(genre_frame, text="Féminin", variable=genre_var, value="féminin",
                        bootstyle="primary-toolbutton").pack(side=LEFT, padx=5)
        ttk.Radiobutton(genre_frame, text="Masculin", variable=genre_var, value="masculin",
                        bootstyle="primary-toolbutton").pack(side=LEFT, padx=5)

        # Type de contrat
        ttk.Label(contract_frame, text="Type de Contrat*", font=('Segoe UI', 10)).grid(row=len(date_fields), column=0,
                                                                                       padx=5, pady=5, sticky=E)
        contract_type_frame = ttk.Frame(contract_frame)
        contract_type_frame.grid(row=len(date_fields), column=1, sticky=W)
        ttk.Radiobutton(contract_type_frame, text="CDD", variable=contract_type_var, value="CDD",
                        bootstyle="primary-toolbutton",
                        command=lambda: contract_entries['date_fin'].entry.config(state=NORMAL)).pack(side=LEFT, padx=5)
        ttk.Radiobutton(contract_type_frame, text="CDI", variable=contract_type_var, value="CDI",
                        bootstyle="primary-toolbutton",
                        command=lambda: contract_entries['date_fin'].entry.config(state=DISABLED)).pack(side=LEFT,
                                                                                                        padx=5)

        # Salaire et prime
        salary_fields = [
            ("Salaire Base*", "salaire", str(contract[3]) if contract else "2500.00", False, ttk.Entry),
            ("Prime*", "prime", str(contract[4]) if contract else "500.00", False, ttk.Entry),
        ]

        for i, (label, field, value, disabled, widget_type) in enumerate(salary_fields, start=len(date_fields) + 1):
            ttk.Label(contract_frame, text=label, font=('Segoe UI', 10)).grid(row=i, column=0, padx=5, pady=5, sticky=E)
            entry = widget_type(contract_frame, bootstyle="primary")
            entry.insert(0, value)
            entry.grid(row=i, column=1, padx=5, pady=5, sticky=EW)
            contract_entries[field] = entry

        # Type de salaire
        ttk.Label(contract_frame, text="Type de Salaire*", font=('Segoe UI', 10)).grid(row=len(date_fields) + 3,
                                                                                       column=0, padx=5, pady=5,
                                                                                       sticky=E)
        salary_type_frame = ttk.Frame(contract_frame)
        salary_type_frame.grid(row=len(date_fields) + 3, column=1, sticky=W)
        ttk.Radiobutton(salary_type_frame, text="Par Heure", variable=salary_type_var, value="hourly",
                        bootstyle="primary-toolbutton").pack(side=LEFT, padx=5)
        ttk.Radiobutton(salary_type_frame, text="Par Mois", variable=salary_type_var, value="monthly",
                        bootstyle="primary-toolbutton").pack(side=LEFT, padx=5)

        # Boutons
        button_frame = ttk.Frame(top)
        button_frame.pack(fill=X, pady=10)
        ttk.Button(button_frame, text="Enregistrer",
                   command=lambda: self.save_employee_and_contract_changes(
                       matricule, entries, contract_entries, genre_var, contract_type_var, salary_type_var, top),
                   bootstyle=SUCCESS).pack(side=LEFT, padx=5)
        ttk.Button(button_frame, text="Annuler", command=top.destroy, bootstyle=WARNING).pack(side=LEFT, padx=5)

        employee_frame.columnconfigure(1, weight=1)
        contract_frame.columnconfigure(1, weight=1)

    def save_employee_and_contract_changes(self, matricule, entries, contract_entries, genre_var, contract_type_var,
                                           salary_type_var, top):
        try:
            # Validation des champs obligatoires
            required_employee = ['nom', 'prenom']
            required_contract = ['date_debut', 'salaire', 'prime']

            for field in required_employee:
                value = self.get_widget_value(entries[field])
                if not value:
                    Messagebox.show_error(f"Le champ {field} est obligatoire", "Erreur")
                    return

            for field in required_contract:
                value = self.get_widget_value(contract_entries[field])
                if not value:
                    Messagebox.show_error(f"Le champ {field} est obligatoire", "Erreur")
                    return

            # Validation des dates
            date_debut = self.get_widget_value(contract_entries['date_debut'])
            if not self.is_valid_date(date_debut):
                Messagebox.show_error("Format de date invalide pour la date de début (JJ/MM/AAAA)", "Erreur")
                return

            if contract_type_var.get() == "CDD":
                date_fin = self.get_widget_value(contract_entries['date_fin'])
                if not self.is_valid_date(date_fin):
                    Messagebox.show_error("Format de date invalide pour la date de fin (JJ/MM/AAAA)", "Erreur")
                    return

                if datetime.datetime.strptime(date_fin, "%d/%m/%Y") <= datetime.datetime.strptime(date_debut,
                                                                                                  "%d/%m/%Y"):
                    Messagebox.show_error("La date de fin doit être après la date de début", "Erreur")
                    return

            # Validation des valeurs numériques
            try:
                salaire = float(self.get_widget_value(contract_entries['salaire']))
                prime = float(self.get_widget_value(contract_entries['prime']))
                if salaire <= 0 or prime < 0:
                    Messagebox.show_error("Le salaire doit être positif et la prime non négative", "Erreur")
                    return
            except ValueError:
                Messagebox.show_error("Le salaire et la prime doivent être des nombres valides", "Erreur")
                return

            with self.conn.cursor() as cursor:
                # Mise à jour employé
                cursor.execute('''
                               UPDATE employes
                               SET nom            = %s,
                                   prenom         = %s,
                                   genre          = %s,
                                   date_naissance = %s,
                                   lieu_naissance = %s,
                                   adresse        = %s,
                                   ville          = %s,
                                   cin            = %s,
                                   date_cin       = %s,
                                   lieu_cin       = %s,
                                   poste          = %s,
                                   email          = %s,
                                   telephone      = %s
                               WHERE matricule = %s
                               ''', (
                    self.get_widget_value(entries['nom']),
                    self.get_widget_value(entries['prenom']),
                    genre_var.get(),
                    self.get_widget_value(entries['date_naissance']) or None,
                    self.get_widget_value(entries['lieu_naissance']) or None,
                    self.get_widget_value(entries['adresse']) or None,
                    self.get_widget_value(entries['ville']) or None,
                    self.get_widget_value(entries['cin']) or None,
                    self.get_widget_value(entries['date_cin']) or None,
                    self.get_widget_value(entries['lieu_cin']) or None,
                    self.get_widget_value(entries['poste']) or None,
                    self.get_widget_value(entries['email']) or None,
                    self.get_widget_value(entries['telephone']) or None,
                    matricule
                               ))

                # Mise à jour contrat
                cursor.execute('''
                               UPDATE contrats
                               SET type_contrat = %s,
                                   date_debut   = %s,
                                   date_fin     = %s,
                                   salaire_base = %s,
                                   prime        = %s,
                                   salary_type  = %s
                               WHERE matricule = %s ORDER BY id DESC LIMIT 1
                               ''', (
                                   contract_type_var.get(),
                                   date_debut,
                                   self.get_widget_value(
                                       contract_entries['date_fin']) if contract_type_var.get() == "CDD" else None,
                                   salaire,
                                   prime,
                                   salary_type_var.get(),
                                   matricule
                               ))

                self.conn.commit()
                Messagebox.show_info("Mise à jour réussie", "Succès")
                self.load_data()
                top.destroy()

        except Exception as e:
            Messagebox.show_error(f"Erreur lors de la sauvegarde: {str(e)}", "Erreur")

    def delete_employee(self, matricule):
        if not Messagebox.yesno(f"Confirmer la suppression de l'employé {matricule}?", "Confirmation"):
            return

        try:
            with self.conn.cursor() as cursor:
                cursor.execute("SELECT * FROM employes WHERE matricule = %s", (matricule,))
                employee_data = cursor.fetchone()
                cursor.execute("SELECT * FROM contrats WHERE matricule = %s ORDER BY id DESC LIMIT 1", (matricule,))
                contract_data = cursor.fetchone()

                if not employee_data:
                    Messagebox.show_error("Employé non trouvé", "Erreur")
                    return

                cursor.execute("DELETE FROM contrats WHERE matricule = %s", (matricule,))
                cursor.execute("DELETE FROM employes WHERE matricule = %s", (matricule,))
                self.conn.commit()

                self.undo_stack.append((
                    "employee_delete", matricule,
                    {
                        'matricule': employee_data[0], 'nom': employee_data[1], 'prenom': employee_data[2],
                        'genre': employee_data[3], 'date_naissance': employee_data[4],
                        'lieu_naissance': employee_data[5], 'adresse': employee_data[6],
                        'ville': employee_data[7], 'cin': employee_data[8], 'date_cin': employee_data[9],
                        'lieu_cin': employee_data[10], 'poste': employee_data[11],
                        'email': employee_data[12], 'telephone': employee_data[13]
                    },
                    {
                        'id': contract_data[0], 'type_contrat': contract_data[2], 'date_creation': contract_data[3],
                        'date_debut': contract_data[4], 'date_fin': contract_data[5], 'salaire_base': contract_data[6],
                        'prime': contract_data[7], 'salary_type': contract_data[8], 'texte_contrat': contract_data[9]
                    } if contract_data else None
                ))

                self.load_data()
                self.clear_search()
                self.status_var.set(f"Employé {matricule} supprimé")
        except mysql.connector.Error as e:
            Messagebox.show_error(f"Erreur de suppression: {str(e)}", "Erreur")

    def undo_action(self):
        if not self.undo_stack:
            Messagebox.show_info("Aucune action à annuler", "Information")
            return

        action_type, matricule, employee_data, contract_data = self.undo_stack.pop()
        try:
            with self.conn.cursor() as cursor:
                if action_type == "employee_delete":
                    cursor.execute('''
                                   INSERT INTO employes (matricule, nom, prenom, genre, date_naissance, lieu_naissance,
                                                         adresse, ville, cin, date_cin, lieu_cin, poste, email,
                                                         telephone)
                                   VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                                   ''', (
                                       employee_data['matricule'], employee_data['nom'], employee_data['prenom'],
                                       employee_data['genre'], employee_data['date_naissance'],
                                       employee_data['lieu_naissance'],
                                       employee_data['adresse'], employee_data['ville'], employee_data['cin'],
                                       employee_data['date_cin'], employee_data['lieu_cin'], employee_data['poste'],
                                       employee_data['email'], employee_data['telephone']
                                   ))
                    if contract_data:
                        cursor.execute('''
                                       INSERT INTO contrats (id, matricule, type_contrat, date_creation, date_debut,
                                                             date_fin, salaire_base, prime, salary_type, texte_contrat)
                                       VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                                       ''', (
                                           contract_data['id'], matricule, contract_data['type_contrat'],
                                           contract_data['date_creation'], contract_data['date_debut'],
                                           contract_data['date_fin'], contract_data['salaire_base'],
                                           contract_data['prime'], contract_data['salary_type'],
                                           contract_data['texte_contrat']
                                       ))
                    self.status_var.set(f"Suppression de {matricule} annulée")
                elif action_type == "employee_update":
                    cursor.execute(f"""
                        UPDATE employes 
                        SET {employee_data} = %s 
                        WHERE matricule = %s
                    """, (contract_data, matricule))
                    self.status_var.set(f"Mise à jour de {employee_data} pour {matricule} annulée")
                elif action_type == "contract_update":
                    cursor.execute(f"""
                        UPDATE contrats 
                        SET {employee_data} = %s 
                        WHERE id = %s
                    """, (contract_data, matricule))
                    self.status_var.set(f"Mise à jour du contrat pour {matricule} annulée")

                self.conn.commit()
                self.load_data()
                if self.current_employee and self.current_employee['matricule'] == matricule:
                    self.search_employee()
        except mysql.connector.Error as e:
            Messagebox.show_error(f"Erreur lors de l'annulation: {str(e)}", "Erreur")

    def show_help(self):
        Messagebox.show_info(
            """Application de Gestion des Contrats
Version: 1.0
Fonctionnalités:
- Ajouter, modifier, supprimer des employés
- Générer des contrats CDD/CDI en arabe
- Exporter les contrats en Word ou ZIP
- Rechercher et filtrer les employés
- Modifier les données directement dans le tableau
- Annuler la dernière action (suppression ou modification)

Pour plus d'aide, contactez le support technique.""",
            "Aide"
        )

    def show_about(self):
        Messagebox.show_info(
            """Application de Gestion des Contrats
Développée par: Imbert Mnif
Version: 1.0
© Imbert Mnif. Tous droits réservés.""",
            "À propos"
        )

    def get_widget_value(self, widget):
        if isinstance(widget, (ttk.Entry, ttk.Combobox)):
            return widget.get().strip()
        elif isinstance(widget, DateEntry):
            return widget.entry.get().strip()
        elif hasattr(widget, 'get'):
            return widget.get().strip()
        return ""

    def play_alert_sound(self):
        if self.sound_enabled and hasattr(self, 'sound_file') and self.sound_file and os.path.exists(self.sound_file):
            try:
                def play_loop():
                    while not hasattr(self, 'alert_stopped') or not self.alert_stopped:
                        winsound.PlaySound(self.sound_file, winsound.SND_FILENAME | winsound.SND_ASYNC)
                        time.sleep(2)

                self.alert_thread = threading.Thread(target=play_loop)
                self.alert_thread.daemon = True
                self.alert_thread.start()
            except Exception as e:
                print(f"Erreur de lecture du son d'alerte: {e}")
                try:
                    def play_system_loop():
                        while not hasattr(self, 'alert_stopped') or not self.alert_stopped:
                            winsound.PlaySound("SystemExclamation", winsound.SND_ALIAS | winsound.SND_ASYNC)
                            time.sleep(2)

                    self.alert_thread = threading.Thread(target=play_system_loop)
                    self.alert_thread.daemon = True
                    self.alert_thread.start()
                except:
                    pass

    def stop_alert_sound(self):
        try:
            self.alert_stopped = True
            winsound.PlaySound(None, 0)
            if hasattr(self, 'alert_thread'):
                self.alert_thread.join(timeout=0.1)
        except Exception as e:
            print(f"Erreur lors de l'arrêt du son: {e}")

    def stop_alert_timer(self):
        if self.alert_timer:
            self.root.after_cancel(self.alert_timer)

    def background_alert_service(self):
        try:
            if not hasattr(self, 'conn') or not self.conn.is_connected():
                self.conn = mysql.connector.connect(**self.db_config)

            today = datetime.datetime.now().date()
            with self.conn.cursor() as cursor:
                cursor.execute('''
                               SELECT e.matricule,
                                      e.nom,
                                      e.prenom,
                                      c.date_fin,
                                      DATEDIFF(STR_TO_DATE(c.date_fin, '%d/%m/%Y'), %s) AS jours_restants
                               FROM employes e
                                        JOIN contrats c ON e.matricule = c.matricule
                               WHERE c.type_contrat = 'CDD'
                                 AND c.date_fin IS NOT NULL
                                 AND DATEDIFF(STR_TO_DATE(c.date_fin, '%d/%m/%Y'), %s) BETWEEN 0 AND 30
                               ORDER BY jours_restants
                               ''', (today, today))
                expiring_contracts = cursor.fetchall()

                new_alerts = []
                for contract in expiring_contracts:
                    matricule, nom, prenom, date_fin, jours_restants = contract
                    contract_key = f"{matricule}_{date_fin.strip().lower()}"

                    if contract_key not in self.notified_contracts:
                        new_alerts.append(contract)
                        self.notified_contracts.add(contract_key)

                if new_alerts:
                    message = "⚠️ ALERTE : Contrats CDD expirant bientôt ⚠️\n\n"
                    for contract in new_alerts:
                        matricule, nom, prenom, date_fin, jours_restants = contract
                        message += f"• {nom} {prenom} (Matricule: {matricule}) - "
                        message += f"Expire le {date_fin} (dans {jours_restants} jours)\n"

                    self.play_alert_sound()
                    Messagebox.show_warning(message, "Alerte Contrats", parent=self.root)
                    self.stop_alert_sound()
                    self.status_var.set(f"⚠ {len(new_alerts)} nouveaux contrats expirent dans ≤30 jours")

                self.alert_timer = self.root.after(self.check_interval, self.background_alert_service)

        except Exception as e:
            print(f"Erreur dans le service d'alerte en arrière-plan: {e}")
            self.alert_timer = self.root.after(self.check_interval, self.background_alert_service)

    def __del__(self):
        self.stop_alert_timer()
        self.stop_alert_sound()
        if hasattr(self, 'conn') and self.conn.is_connected():
            self.conn.close()


    def generate_contract_text(self, employee_data, contract_data):
        try:
            # Determine contract type in Arabic
            contract_type = "غير محددة" if contract_data['type_contrat'] == "CDI" else "محدودة"

            # Format date fin (only for CDD)
            date_fin_text = f" إلى {contract_data['date_fin']}" if contract_data['type_contrat'] == "CDD" and contract_data[
                'date_fin'] else ""

            # Format salary type in Arabic
            salary_type_text = "بالساعة" if contract_data['salary_type'] == "hourly" else "شهري"

            # Build the contract text
            contract_text = f"""عقد عمل {contract_type}
    بين الشركة: Imbert Mnif
    والموظف: {employee_data['prenom']} {employee_data['nom']}
    الرقم التعريفي: {employee_data['matricule']}
    المنصب: {employee_data['poste'] or 'غير محدد'}
    تاريخ الميلاد: {employee_data['date_naissance'] or 'غير محدد'}
    مكان الميلاد: {employee_data['lieu_naissance'] or 'غير محدد'}
    العنوان: {employee_data['adresse'] or 'غير محدد'}، {employee_data['ville'] or 'غير محدد'}
    رقم بطاقة التعريف: {employee_data['cin'] or 'غير محدد'}
    تاريخ الإصدار: {employee_data['date_cin'] or 'غير محدد'}
    مكان الإصدار: {employee_data['lieu_cin'] or 'غير محدد'}
    البريد الإلكتروني: {employee_data['email'] or 'غير محدد'}
    الهاتف: {employee_data['telephone'] or 'غير محدد'}
    الراتب الأساسي: {contract_data['salaire_base']} دينار تونسي
    المنحة: {contract_data['prime']} دينار تونسي
    نوع الأجرة: {salary_type_text}
    تاريخ البدء: {contract_data['date_debut']}{date_fin_text}
    """
            return contract_text.strip()
        except Exception as e:
            logging.error(f"Erreur dans generate_contract_text: {str(e)}")
            raise Exception(f"Erreur lors de la génération du texte du contrat: {str(e)}")

if __name__ == "__main__":
    app = ContratApplication(ttk.Window())
    app.root.mainloop()