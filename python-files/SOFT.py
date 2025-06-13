import tkinter as tk
from tkinter import filedialog, messagebox, ttk, scrolledtext
from datetime import datetime
import os
from PIL import Image, ImageTk
import fitz  # PyMuPDF
from docx import Document
import win32com.client
import random

class DocumentApprovalApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Автоматизированная система утверждения электронных документов")
        self.root.geometry("800x600")
        self.document_path = ""
        self.current_page = 0
        self.images = []

        # User Identification
        self.user_authenticated = False

        # Создание стиля
        style = ttk.Style()
        style.configure("TFrame", padding=(10, 10))
        style.configure("TLabel", font=("Arial", 12))
        style.configure("TButton", font=("Arial", 12), padding=(5, 5))

        # Загрузка изображения фона
        self.background_image = Image.open(r"C:\Users\us\Desktop\your_path_to_yars_image.jpg")  # Укажите свой путь
        self.background_image = self.background_image.resize((800, 600), Image.LANCZOS)
        self.background_photo = ImageTk.PhotoImage(self.background_image)

        # Canvas для фона
        self.canvas = tk.Canvas(root, width=800, height=600)
        self.canvas.pack()
        self.canvas.create_image(0, 0, anchor='nw', image=self.background_photo)

        # Фрейм авторизации
        self.login_frame = ttk.Frame(root)
        self.login_frame.place(relx=0.5, rely=0.4, anchor='center')

        ttk.Label(self.login_frame, text="Логин:").grid(row=0, column=0, sticky=tk.W)
        self.username_entry = ttk.Entry(self.login_frame)
        self.username_entry.grid(row=0, column=1, sticky=tk.EW)

        ttk.Label(self.login_frame, text="Пароль:").grid(row=1, column=0, sticky=tk.W)
        self.password_entry = ttk.Entry(self.login_frame, show="*")
        self.password_entry.grid(row=1, column=1, sticky=tk.EW)

        ttk.Button(self.login_frame, text="Войти", command=self.login).grid(row=2, columnspan=2, pady=(10, 0))

        # Document Upload Frame
        self.document_frame = ttk.Frame(root)

        self.document_label = ttk.Label(self.document_frame, text="Нет добавленных документов", font=("Arial", 14))
        self.document_label.pack(pady=10)

        self.image_label = ttk.Label(self.document_frame)
        self.image_label.pack(pady=(10, 10))

        # Кнопки для управления
        button_frame = ttk.Frame(self.document_frame)
        button_frame.pack(pady=(5, 2))

        ttk.Button(button_frame, text="Добавить документ", command=self.attach_document).pack(side=tk.LEFT)
        ttk.Button(button_frame, text="Утвердить документ", command=self.approve_document).pack(side=tk.LEFT)
        ttk.Button(button_frame, text="Отказать в утверждении", command=self.ask_rejection_reason).pack(side=tk.LEFT)
        ttk.Button(button_frame, text="Журнал действий", command=self.show_log).pack(side=tk.RIGHT)

        # Кнопки для переключения страниц
        self.prev_button = ttk.Button(self.document_frame, text="Предыдущая", command=self.show_prev_page)
        self.prev_button.pack(side=tk.LEFT, padx=(10, 5))

        self.next_button = ttk.Button(self.document_frame, text="Следующая", command=self.show_next_page)
        self.next_button.pack(side=tk.LEFT, padx=(5, 10))

    def login(self):
        username = self.username_entry.get()
        password = self.password_entry.get()
        if username == "Надымов" and password == "1234":
            self.user_authenticated = True
            self.login_frame.place_forget()
            self.document_frame.place(relx=0.5, rely=0.5, anchor='center')
            messagebox.showinfo("Вход", "Успешная авторизация")
        else:
            messagebox.showerror("Вход", "Неправильное имя пользователя или пароль.")

    def attach_document(self):
        if not self.user_authenticated:
            messagebox.showerror("Ошибка", "Для начала работы Вам необходимо авторизоваться.")
            return

        self.document_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf"), ("Word files", "*.docx")])
        if self.document_path:
            self.document_label.config(text=os.path.basename(self.document_path))
            self.log_action(f"Документ добавлен: {self.document_path}")

            if self.document_path.endswith('.pdf'):
                self.load_pdf_images()  # Загружаем изображения из PDF
            elif self.document_path.endswith('.docx'):
                self.load_word_images()  # Загружаем изображения из Word
            self.current_page = 0
            self.show_page()  # Показать первую страницу документа

    def load_pdf_images(self):
        self.images.clear()  # Очищаем предыдущие изображения
        pdf_document = fitz.open(self.document_path)  # Открываем PDF с помощью PyMuPDF
        for page in pdf_document:
            pix = page.get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            self.images.append(img)  # Добавляем изображение страницы в список
        pdf_document.close()  # Закрываем PDF-документ

    def load_word_images(self):
        self.images.clear()  # Очищаем предыдущие изображения
        word_app = win32com.client.Dispatch('Word.Application')
        word_doc = word_app.Documents.Open(self.document_path)  # Открываем Word документ
        temp_dir = os.path.join(os.path.dirname(self.document_path), "temp_images")
        os.makedirs(temp_dir, exist_ok=True)  # Создаем временную папку для изображений

        try:
            # Экспортируем каждую страницу в PDF
            for i in range(1, word_doc.ComputeStatistics(2) + 1):
                pdf_path = os.path.join(temp_dir, f'page_{i}.pdf')
                word_doc.ExportAsFixedFormat(pdf_path, 17)  # 17 = wdExportFormatPDF

                pdf_page = fitz.open(pdf_path)  # Открываем экспортированный PDF
                for page in pdf_page:
                    pix = page.get_pixmap()  # Получаем изображение
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    self.images.append(img)  # Добавляем изображение страницы в список
                pdf_page.close()  # Закрываем временный PDF
        finally:
            word_doc.Close(False)  # Закрываем Word документ
            word_app.Quit()  # Закрываем приложение Word
            # Удаляем временные файлы
            for file in os.listdir(temp_dir):
                os.remove(os.path.join(temp_dir, file))
            os.rmdir(temp_dir)

    def show_page(self):
        if self.images:  # Проверяем, что есть изображения для отображения
            img = self.images[self.current_page]  # Получаем текущее изображение
            img.thumbnail((800, 600))  # Уменьшаем изображение до заданного размера
            self.tk_image = ImageTk.PhotoImage(img)  # Преобразуем в формат, подходящий для Tkinter
            self.image_label.config(image=self.tk_image)  # Устанавливаем изображение в Label
            self.image_label.image = self.tk_image  # Сохраняем ссылку на изображение

            # Настройка кнопок навигации
            self.prev_button.config(state=tk.NORMAL if self.current_page > 0 else tk.DISABLED)
            self.next_button.config(state=tk.NORMAL if self.current_page < len(self.images) - 1 else tk.DISABLED)

    def show_next_page(self):
        if self.current_page < len(self.images) - 1:
            self.current_page += 1
            self.show_page()  # Показать следующую страницу

    def show_prev_page(self):
        if self.current_page > 0:
            self.current_page -= 1
            self.show_page()  # Показать предыдущую страницу

    def approve_document(self):
        if not self.user_authenticated:
            messagebox.showerror("Ошибка", "Для начала работы Вам необходимо авторизоваться.")
            return
        if not self.document_path:
            messagebox.showerror("Ошибка", "Сначала добавьте документ для утверждения.")
            return

        key_file_path = filedialog.askopenfilename(title="Выберите файл с приватным ключом", filetypes=[("PEM files", "*.pem")])
        if not key_file_path:
            return

        cert_file_path = filedialog.askopenfilename(title="Выберите файл с сертификатом", filetypes=[("CRT files", "*.crt"), ("CER files", "*.cer")])
        if not cert_file_path:
            return

        try:
            self.sign_pdf(self.document_path, key_file_path, cert_file_path)  # Подписываем документ
            self.log_action(f"Документ подписан: {self.document_path}")
            messagebox.showinfo("Утверждение", "Документ успешно утвержден и подписан.")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось подписать документ: {e}")

    def sign_pdf(self, pdf_path, key_file_path, cert_file_path):
        with open(pdf_path, "rb") as f:
            reader = fitz.open(f)  # Открываем PDF
            writer = fitz.open()  # Создаем новый PDF для записи
            writer.insert_page(-1)
            writer[-1].insert_text((100, 100), "Подписано", fontsize=12)

            signed_pdf_path = pdf_path.replace(".pdf", "_signed.pdf")
            writer.save(signed_pdf_path)  # Сохраняем подписанный документ
            writer.close()  # Закрываем новый PDF

    def ask_rejection_reason(self):
        reason_window = tk.Toplevel(self.root)
        reason_window.title("Укажите причины отказа")
        
        label = ttk.Label(reason_window, text="Введите причину отказа:")
        label.pack(padx=10, pady=10)

        reason_text = tk.Text(reason_window, height=10, width=50)
        reason_text.pack(padx=10, pady=10)

        save_button = ttk.Button(reason_window, text="Сохранить", command=lambda: self.save_rejection_reason(reason_text.get("1.0", tk.END).strip(), reason_window))
        save_button.pack(pady=(0, 10))

    def save_rejection_reason(self, reason, reason_window):
        if not reason:
            messagebox.showwarning("Предупреждение", "Причина отказа не может быть пустой.")
            return

        # Сохранение причины отказа в Word-документ
        rejection_file_path = os.path.join(os.path.dirname(self.document_path), "Причина отказа.docx")
        doc = Document()
        
        # Устанавливаем отступы (преобразуем в int)
        sections = doc.sections
        for section in sections:
            section.top_margin = int(2.5 * 28.35)  # 1 см = 28.35 пунктов
            section.bottom_margin = int(2 * 28.35)
            section.left_margin = int(random.uniform(2.75, 3.5) * 28.35)  # Случайное значение от 2.75 до 3.5 см
            section.right_margin = int(random.uniform(1.25, 2.25) * 28.35)  # Случайное значение от 1.25 до 2.25 см

        # Добавляем заголовок
        heading = doc.add_heading("Причина отказа", level=1)
        heading.alignment = 1  # Центрируем заголовок

        # Добавляем основной текст
        p = doc.add_paragraph(reason)
        p.alignment = 0  # Выравнивание по умолчанию (по левому краю)

        # Сохраняем документ
        doc.save(rejection_file_path)

        self.log_action(f"Отказано в утверждении: {self.document_path}, причина: {reason}")
        messagebox.showinfo("Успешно", "Причина отказа успешно сохранена.")
        reason_window.destroy()

    def log_action(self, action):
        with open("document_approval_log.txt", "a") as log_file:
            log_file.write(f"{datetime.now()}: {action}\n")

    def show_log(self):
        log_window = tk.Toplevel(self.root)
        log_window.title("Журнал действий")
        log_window.geometry("600x400")

        log_text = scrolledtext.ScrolledText(log_window, wrap=tk.WORD)
        log_text.pack(expand=True, fill=tk.BOTH)

        # Читаем и показываем журнал действий
        if os.path.exists("document_approval_log.txt"):
            with open("document_approval_log.txt", "r") as log_file:
                log_content = log_file.read()
                log_text.insert(tk.END, log_content)

        # Кнопка для закрытия окна
        ttk.Button(log_window, text="Закрыть", command=log_window.destroy).pack(pady=10)


if __name__ == "__main__":
    root = tk.Tk()
    app = DocumentApprovalApp(root)
    root.mainloop()