# scramble_test_generator.py
import sys
import os
import random
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import zipfile

hebrew_bullets = ["א", "ב", "ג", "ד", "ה", "ו", "ז", "ח", "ט", "י"]

def set_rtl(paragraph):
    p = paragraph._element
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    p_pr = p.get_or_add_pPr()
    p_pr.append(bidi)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

def parse_docx(file_path):
    doc = Document(file_path)
    elements = []
    current_text = ""
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if text.startswith("[") and text.endswith("]"):
            tag = text.strip("[]").lower()
            elements.append((tag, current_text.strip()))
            current_text = ""
        else:
            if current_text:
                current_text += "\n"
            current_text += text
    title, intro, questions = "", "", []
    q_buffer = {}
    for tag, content in elements:
        if tag == "title":
            title = content
        elif tag == "intro" or tag == "info":
            intro += content + "\n"
        elif tag.startswith("q"):
            if q_buffer:
                questions.append(q_buffer)
            q_buffer = {"question": content, "answers": []}
        elif tag == "a":
            if q_buffer:
                q_buffer["answers"].append(content)
    if q_buffer:
        questions.append(q_buffer)
    return title, intro.strip(), questions

def generate_versions(file_path, count, output_zip):
    title, intro, questions = parse_docx(file_path)
    output_paths = []

    for n in range(1, count + 1):
        doc = Document()

        p_title = doc.add_paragraph(title)
        set_rtl(p_title)
        doc.add_paragraph("[Title]")

        p_intro = doc.add_paragraph(intro)
        set_rtl(p_intro)
        doc.add_paragraph("[Intro]")

        scrambled = questions.copy()
        random.shuffle(scrambled)
        for i, q in enumerate(scrambled, 1):
            doc.add_paragraph("")
            doc.add_paragraph("")
            p_num = doc.add_paragraph()
            run_num = p_num.add_run(f"שאלה מספר {i}")
            run_num.bold = True
            run_num.underline = True
            set_rtl(p_num)
            p_q = doc.add_paragraph(q['question'])
            set_rtl(p_q)
            answers = q['answers'].copy()
            random.shuffle(answers)
            for j, a in enumerate(answers):
                bullet = hebrew_bullets[j % len(hebrew_bullets)]
                p_a = doc.add_paragraph(f"{bullet}. {a}")
                set_rtl(p_a)

        out_file = f"scrambled_version_{n}.docx"
        doc.save(out_file)
        output_paths.append(out_file)

    with zipfile.ZipFile(output_zip, 'w') as zipf:
        for file in output_paths:
            zipf.write(file)
            os.remove(file)

if __name__ == "__main__":
    docx_path = sys.argv[1]
    version_count = int(sys.argv[2])
    output_zip = sys.argv[3]
    generate_versions(docx_path, version_count, output_zip)
