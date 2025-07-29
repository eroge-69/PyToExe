import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document
import pandas as pd
import re
import sys
import os

def parse_word_to_excel(input_word, output_excel):
    """Основная функция парсинга"""
    try:
        doc = Document(input_word)
        cables = []
        
        # Шаблоны для поиска данных
        patterns = [
            re.compile(r"(?P<type>\w[\w-]+)\s+(?P<section>[\dхx.]+)\s+(?P<length>\d+)"), # "ВВГ-нг 3х2.5 100"
            re.compile(r"(?P<type>\w[\w-]+)[\s,]+сеч[.\s]*(?P<section>[\dхx.]+)[\s,]+длин[.\s]*(?P<length>\d+)") # "Тип: ВВГ-нг, сеч. 3х2.5, длина 100"
        ]
        
        # Поиск в тексте
        for para in doc.paragraphs:
            for pattern in patterns:
                match = pattern.search(para.text.replace(",", " "))
                if match:
                    cables.append({
                        'Тип кабеля': match.group('type').strip(),
                        'Сечение': match.group('section').strip().replace("x", "х"),
                        'Длина (м)': int(match.group('length'))
                    })
                    break
        
        # Поиск в таблицах
        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) >= 3:
                    cables.append({
                        'Тип кабеля': row.cells[0].text.strip(),
                        'Сечение': row.cells[1].text.strip().replace("x", "х"),
                        'Длина (м)': int(float(row.cells[2].text.strip()))
                    })
        
        if not cables:
            raise ValueError("Данные о кабелях не найдены!")
        
        # Создаем DataFrame
        df = pd.DataFrame(cables)
        total = df['Длина (м)'].sum()
        
        # Сохраняем в Excel
        with pd.ExcelWriter(output_excel) as writer:
            df.to_excel(writer, sheet_name='Данные', index=False)
            
            # Добавляем итоги
            summary = df.groupby(['Тип кабеля', 'Сечение'])['Длина (м)'].sum().reset_index()
            summary.to_excel(writer, sheet_name='Итоги', index=False)
            
            # Форматирование
            workbook = writer.book
            header_format = workbook.add_format({
                'bold': True, 
                'align': 'center',
                'valign': 'vcenter',
                'border': 1
            })
            
            for sheet in writer.sheets.values():
                sheet.set_column('A:C', 25)
                sheet.freeze_panes(1, 0)
                for col_num, value in enumerate(df.columns.values):
                    sheet.write(0, col_num, value, header_format)
            
            writer.sheets['Итоги'].write(len(summary)+1, 0, 'ВСЕГО', header_format)
            writer.sheets['Итоги'].write(len(summary)+1, 2, total, header_format)
        
        return True, f"Файл успешно сохранен:\n{output_excel}"
    
    except Exception as e:
        return False, f"Ошибка: {str(e)}"

class CableParserApp:
    """Графический интерфейс"""
    def __init__(self):
        self.window = tk.Tk()
        self.window.title("Парсер кабельной документации v1.0")
        self.window.geometry("600x300")
        
        # Стиль
        style = ttk.Style()
        style.configure('TButton', font=('Arial', 10))
        style.configure('TLabel', font=('Arial', 10))
        
        # Заголовок
        ttk.Label(self.window, text="Конвертер Word → Excel", font=('Arial', 14, 'bold'))\
            .pack(pady=10)
        
        # Поля ввода
        frame = ttk.Frame(self.window)
        frame.pack(pady=10, padx=20, fill=tk.X)
        
        ttk.Label(frame, text="Word-документ:").grid(row=0, column=0, sticky='w')
        self.word_entry = ttk.Entry(frame, width=50)
        self.word_entry.grid(row=1, column=0, padx=5)
        ttk.Button(frame, text="Выбрать", command=self.browse_word).grid(row=1, column=1)
        
        ttk.Label(frame, text="Excel-файл:").grid(row=2, column=0, sticky='w', pady=(10,0))
        self.excel_entry = ttk.Entry(frame, width=50)
        self.excel_entry.grid(row=3, column=0, padx=5)
        ttk.Button(frame, text="Выбрать", command=self.browse_excel).grid(row=3, column=1)
        
        # Кнопка обработки
        ttk.Button(self.window, text="Конвертировать", command=self.process).pack(pady=20)
        
        # Статус
        self.status = ttk.Label(self.window, text="", foreground='gray')
        self.status.pack()
        
        self.window.mainloop()
    
    def browse_word(self):
        filename = filedialog.askopenfilename(
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
        )
        if filename:
            self.word_entry.delete(0, tk.END)
            self.word_entry.insert(0, filename)
            # Автозаполнение имени для Excel
            if not self.excel_entry.get():
                base = os.path.splitext(filename)[0]
                self.excel_entry.insert(0, f"{base}_результат.xlsx")
    
    def browse_excel(self):
        filename = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")]
        )
        if filename:
            self.excel_entry.delete(0, tk.END)
            self.excel_entry.insert(0, filename)
    
    def process(self):
        word_file = self.word_entry.get()
        excel_file = self.excel_entry.get()
        
        if not word_file:
            messagebox.showerror("Ошибка", "Выберите Word-файл!")
            return
        
        if not excel_file:
            messagebox.showerror("Ошибка", "Укажите файл для сохранения!")
            return
        
        self.status.config(text="Обработка...", foreground='blue')
        self.window.update()
        
        success, msg = parse_word_to_excel(word_file, excel_file)
        
        if success:
            self.status.config(text=msg, foreground='green')
            messagebox.showinfo("Готово", msg)
        else:
            self.status.config(text=msg, foreground='red')
            messagebox.showerror("Ошибка", msg)

if __name__ == "__main__":
    # Проверка аргументов командной строки
    if len(sys.argv) > 1:
        word_file = sys.argv[1]
        excel_file = sys.argv[2] if len(sys.argv) > 2 else f"{os.path.splitext(word_file)[0]}_результат.xlsx"
        parse_word_to_excel(word_file, excel_file)
    else:
        CableParserApp()