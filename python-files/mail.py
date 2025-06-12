import os
import json
import pandas as pd
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from mailmerge import MailMerge
from datetime import datetime
import shutil
from docx import Document
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from copy import deepcopy
from docx.enum.section import WD_SECTION
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import sys
import tempfile

# 在程序啟動時初始化 python-docx 環境
def initialize_docx_environment():
    """
    初始化 python-docx 環境，避免使用內置默認模板
    這可以解決 "Failed to extract docx\\templates\\default.docx: failed to open target file!" 錯誤
    通過設置環境變量，確保所有模板都從外部讀取
    """
    try:
        # 設置本地模板目錄
        local_templates_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'templates')
        os.makedirs(local_templates_dir, exist_ok=True)
        
        # 設置環境變量，告訴 python-docx 使用指定的模板目錄
        os.environ['DOCX_TEMPLATE_PATH'] = local_templates_dir
        print(f"已設置環境變量 DOCX_TEMPLATE_PATH={local_templates_dir}")
        
    except Exception as e:
        print(f"初始化 python-docx 環境時出錯: {e}")
        # 不中斷程序執行，繼續嘗試運行

# 強制設置一個可寫入的模板目錄（EXE執行時就是當前資料夾下的 templates）
os.environ['DOCX_TEMPLATE_PATH'] = os.path.join(os.getcwd(), 'templates')
os.makedirs(os.environ['DOCX_TEMPLATE_PATH'], exist_ok=True)

# 自定義函數，用於創建 Document 對象，避免使用默認模板
def create_document(file_path=None):
    """
    創建一個新的 Document 對象，避免使用默認模板
    如果提供了文件路徑，則使用該路徑創建文檔
    如果沒有提供路徑，則提示用戶選擇模板
    
    :param file_path: 可選的文件路徑，如果提供則從該文件創建文檔
    :return: Document 對象
    """
    try:
        if file_path and os.path.exists(file_path):
            # 如果提供了有效的文件路徑，使用該路徑創建文檔
            try:
                return Document(file_path)
            except Exception as e:
                print(f"使用提供的文件路徑創建文檔時出錯: {e}")
                # 如果創建失敗，提示用戶選擇模板
                from tkinter import messagebox, filedialog
                messagebox.showwarning("警告", "無法使用指定的模板文件，請選擇其他模板文件")
                new_file_path = filedialog.askopenfilename(
                    title="選擇Word模板文件",
                    filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
                )
                if new_file_path:
                    return Document(new_file_path)
                else:
                    raise Exception("未選擇有效的模板文件")
        else:
            # 如果未提供文件路徑，提示用戶選擇模板
            from tkinter import messagebox, filedialog
            messagebox.showwarning("警告", "請選擇Word模板文件\n\n所有模板必須從外部讀取，系統不提供默認模板")
            new_file_path = filedialog.askopenfilename(
                title="選擇Word模板文件",
                filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
            )
            if new_file_path:
                return Document(new_file_path)
            else:
                raise Exception("未選擇有效的模板文件")
    except Exception as e:
        print(f"創建文檔時出錯: {e}")
        raise e

# 在程序啟動時調用初始化函數
initialize_docx_environment()

def get_application_path():
    """獲取應用程序路徑，處理 PyInstaller 打包情況"""
    if getattr(sys, 'frozen', False):
        # 如果是 PyInstaller 打包的可執行文件
        return os.path.dirname(sys.executable)
    else:
        # 如果是普通 Python 腳本
        return os.path.dirname(os.path.abspath(__file__))

def resource_path(relative_path):
    """獲取資源路徑，處理 PyInstaller 打包情況"""
    base_path = get_application_path()
    return os.path.join(base_path, relative_path)

class MailMergeApp:
    def __init__(self, root):
        self.root = root
        self.root.title("程序選擇")
        self.root.geometry("400x300")
        self.root.configure(bg="#f0f0f0")

        # 初始化共用變量
        self.status_var = tk.StringVar()
        self.status_var.set("就緒")
        
        # 初始化配置文件路徑
        self.config_file = os.path.join(get_application_path(), 'config.json')
        
        # 創建模板目錄
        self.templates_dir = os.path.join(get_application_path(), 'templates')
        os.makedirs(self.templates_dir, exist_ok=True)
        
        # 創建輸出目錄
        self.output_dir = os.path.join(get_application_path(), 'output')
        os.makedirs(self.output_dir, exist_ok=True)
        
        # 設置 python-docx 環境變量，指向模板目錄
        os.environ['DOCX_TEMPLATE_PATH'] = self.templates_dir
        
        # 清理主窗口
        for widget in self.root.winfo_children():
            widget.destroy()

        self.create_selection_ui()

    def load_config(self, mode):
        """統一的配置載入函數"""
        try:
            if not os.path.exists(self.config_file):
                return {}
                
            with open(self.config_file, 'r', encoding='utf-8') as f:
                config = json.load(f)
                return config.get(mode, {})
        except Exception as e:
            print(f"載入{mode}模式配置錯誤: {e}")
            return {}

    def save_config(self, mode, config_data):
        """統一的配置保存函數"""
        try:
            # 讀取現有配置
            current_config = {}
            if os.path.exists(self.config_file):
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    current_config = json.load(f)
            
            # 更新指定模式的配置
            current_config[mode] = config_data
            
            # 保存更新後的配置
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(current_config, f, ensure_ascii=False, indent=4)
        except Exception as e:
            messagebox.showerror("錯誤", f"保存{mode}模式配置失敗: {e}")

    def update_preview_table(self, tree, df):
        """統一的預覽表格更新函數"""
        for i in tree.get_children():
            tree.delete(i)
        if df is None or df.empty:
            return
            
        tree["columns"] = list(df.columns)
        for col in df.columns:
            tree.heading(col, text=col)
            tree.column(col, width=100)
            
        for index, row in df.iterrows():
            tree.insert("", "end", values=list(row))

    def populate_filter_listbox(self, listbox, data, column_name, date_format=False):
        """統一的篩選列表框填充函數"""
        if data is None or data.empty or column_name not in data.columns:
            return
            
        listbox.delete(0, tk.END)
        values = []
        
        for val in data[column_name]:
            if pd.notna(val):
                try:
                    if date_format:
                        val_str = str(val).strip()
                        if val_str and val_str.lower() != 'nan':
                            values.append(val_str)
                    else:
                        formatted_val = str(val).strip()
                        if formatted_val and formatted_val.lower() != 'nan':
                            values.append(formatted_val)
                except Exception as e:
                    print(f"處理值時出錯: {val}, 錯誤: {e}")
                    continue
        
        # 去重
        values = list(set(values))
        
        # 如果是日期，尝试按日期对象排序
        if date_format:
            try:
                date_objects = []
                for date_str in values:
                    try:
                        if '/' in date_str:
                            parts = date_str.split('/')
                            if len(parts) == 3:
                                day = int(parts[0])
                                month = int(parts[1])
                                year = int(parts[2])
                                date_obj = datetime(year, month, day)
                                date_objects.append((date_obj, date_str))
                                continue
                        
                        date_obj = pd.to_datetime(date_str)
                        date_objects.append((date_obj, date_str))
                    except:
                        date_objects.append((date_str, date_str))
                
                # 排序
                date_items = [item for item in date_objects if isinstance(item[0], datetime)]
                string_items = [item for item in date_objects if not isinstance(item[0], datetime)]
                date_items.sort(key=lambda x: x[0], reverse=True)
                
                # 使用原始字符串格式
                sorted_values = [item[1] for item in date_items + string_items]
            except Exception as e:
                print(f"日期排序错误: {e}")
                sorted_values = sorted(values, reverse=True)
        else:
            sorted_values = sorted(values)
                    
        for val in sorted_values:
            listbox.insert(tk.END, val)

    def apply_filter(self, data, filter_conditions):
        """統一的篩選函數"""
        if data is None or data.empty:
            return data
            
        filtered_df = data.copy()
        for column, values in filter_conditions.items():
            if values:
                if any(keyword in column.lower() for keyword in ['日期', 'date']):
                    # 日期篩選 - 使用原始的日期字符串进行比对，不转换格式
                    date_mask = pd.Series([False] * len(filtered_df))
                    for i, val in enumerate(filtered_df[column]):
                        if pd.notna(val):
                            try:
                                # 使用原始字符串进行比较
                                val_str = str(val).strip()
                                if val_str in values:
                                    date_mask.iloc[i] = True
                            except Exception as e:
                                print(f"筛选日期时出错: {val}, 错误: {e}")
                    filtered_df = filtered_df[date_mask]
                else:
                    # 一般篩選
                    filtered_df = filtered_df[filtered_df[column].astype(str).isin(values)]
                    
        return filtered_df

    def clear_filter(self, listboxes):
        """統一的篩選清除函數"""
        for listbox in listboxes:
            if hasattr(self, listbox):
                getattr(self, listbox).selection_clear(0, tk.END)

    def reload_data(self, mode):
        """統一的數據重新載入函數"""
        excel_file = getattr(self, f'{mode}_excel_file').get()
        if not excel_file:
            messagebox.showwarning("警告", "請先選擇Excel文件")
            return
            
        try:
            load_func = getattr(self, f'load_{mode}_excel_data')
            load_func()
            self.status_var.set(f"已重新載入Excel數據，共 {len(getattr(self, f'{mode}_excel_data'))} 筆記錄")
            messagebox.showinfo("成功", "Excel數據已重新載入")
        except Exception as e:
            messagebox.showerror("錯誤", f"重新載入Excel數據失敗: {e}")
            self.status_var.set("重新載入Excel數據失敗")

    def create_selection_ui(self):
        """創建初始的模式選擇界面"""
        selection_frame = ttk.Frame(self.root, padding=20)
        selection_frame.pack(expand=True, fill=tk.BOTH)

        self.root.title("自動轉換系統")
        self.root.geometry("400x400")  # 調整窗口大小

        # 添加標題標籤
        title_label = ttk.Label(
            selection_frame, 
            text="請選擇操作模式", 
            font=('Microsoft YaHei', 16, 'bold')
        )
        title_label.pack(pady=(0, 30))

        # 設置按鈕樣式
        style = ttk.Style()
        style.configure(
            'Mode.TButton',
            font=('Microsoft YaHei', 12),
            padding=15
        )

        # 創建按鈕容器
        button_frame = ttk.Frame(selection_frame)
        button_frame.pack(pady=10)

        # 創建三個按鈕，垂直排列
        buttons = [
            ("入職文件", self.start_onboarding_mode),
            ("早操記錄", self.start_morning_briefing_mode),
            ("安全訓練", self.start_safety_training_mode)
        ]

        for text, command in buttons:
            btn = ttk.Button(
                button_frame,
                text=text,
                command=command,
                style='Mode.TButton',
                width=30
            )
            btn.pack(pady=15)

    def create_return_button(self, parent_frame):
        """創建返回按鈕"""
        # 創建返回按鈕容器（置於底部）
        return_frame = ttk.Frame(parent_frame)
        return_frame.pack(side=tk.BOTTOM, fill=tk.X, pady=10)
        
        # 創建返回按鈕
        return_btn = ttk.Button(
            return_frame,
            text="返回主菜單",
            command=self.return_to_main,
            style='Return.TButton'
        )
        return_btn.pack(side=tk.LEFT, padx=10)

    def return_to_main(self):
        """返回主菜單"""
        # 清理當前界面
        for widget in self.root.winfo_children():
            widget.destroy()
        
        # 重新創建選擇界面
        self.create_selection_ui()

    def start_onboarding_mode(self):
        """啟動入職模式界面"""
        # 清理選擇界面
        for widget in self.root.winfo_children():
            widget.destroy()
        
        self.root.title("入職文件列印") # 設置新標題
        self.root.geometry("900x700") # 根據入職界面調整大小，可能需要更大

        # 初始化入職模式所需的變量
        self.onboarding_excel_file = tk.StringVar()
        self.onboarding_template_file = tk.StringVar()
        self.onboarding_output_path = tk.StringVar()
        self.onboarding_excel_data = None # 用於存儲加載的Excel數據
        self.status_var = tk.StringVar() # 添加狀態欄變量
        self.status_var.set("就緒") # 設置初始狀態
        
        # 入職模式的配置文件
        self.onboarding_config_file = resource_path('onboarding_config.json')
        
        # 加載入職模式的配置
        self.load_onboarding_config()

        # 創建界面
        self.create_onboarding_ui()

    def create_onboarding_ui(self):
        """創建入職模式的GUI界面"""
        main_frame = ttk.Frame(self.root, padding=10)
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 按鈕區域
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=(0, 10))

        # 定義入職模式的按鈕及其命令 (命令需要新實現或調整)
        onboarding_buttons = [
            ("輸入Excel", self.select_onboarding_excel), # 需要新方法
            ("選擇模板", self.select_onboarding_template), # 需要新方法
            ("輸出位置選擇", self.select_onboarding_output), # 需要新方法
            ("套用篩選", self.apply_onboarding_filter), # 需要新方法
            ("清除篩選", self.clear_onboarding_filter), # 需要新方法
            ("重新載入", self.reload_onboarding_data), # 新增重新載入按鈕
            ("生成Word", self.generate_onboarding_word),  # 需要新方法
            ("返回主菜單", self.return_to_main)  # 添加返回按鈕
        ]

        for i, (btn_text, btn_command) in enumerate(onboarding_buttons):
            btn = ttk.Button(btn_frame, text=btn_text, command=btn_command)
            btn.grid(row=0, column=i, padx=5)

        # 設置區域
        settings_frame = ttk.LabelFrame(main_frame, text="設置", padding=10)
        settings_frame.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(settings_frame, text="Excel文件:").grid(row=0, column=0, sticky=tk.W, pady=5)
        ttk.Entry(settings_frame, textvariable=self.onboarding_excel_file, width=50).grid(row=0, column=1, sticky=tk.W+tk.E, padx=5)
        
        ttk.Label(settings_frame, text="模板文件:").grid(row=1, column=0, sticky=tk.W, pady=5)
        ttk.Entry(settings_frame, textvariable=self.onboarding_template_file, width=50).grid(row=1, column=1, sticky=tk.W+tk.E, padx=5)
        
        ttk.Label(settings_frame, text="輸出位置:").grid(row=2, column=0, sticky=tk.W, pady=5)
        ttk.Entry(settings_frame, textvariable=self.onboarding_output_path, width=50).grid(row=2, column=1, sticky=tk.W+tk.E, padx=5)

        # 郵件合併欄位區域
        onboarding_filter_frame = ttk.LabelFrame(main_frame, text="篩選條件", padding=10)
        onboarding_filter_frame.pack(fill=tk.X, pady=(0, 10))

        # ID 多選列表框
        ttk.Label(onboarding_filter_frame, text="ID：").grid(row=0, column=0, sticky=tk.NW, pady=5)
        id_list_frame = ttk.Frame(onboarding_filter_frame)
        id_list_frame.grid(row=0, column=1, sticky=tk.NSEW, pady=5)
        self.onboarding_id_listbox = tk.Listbox(id_list_frame, selectmode=tk.MULTIPLE, exportselection=False, height=6, width=25)
        id_yscroll = ttk.Scrollbar(id_list_frame, orient=tk.VERTICAL, command=self.onboarding_id_listbox.yview)
        self.onboarding_id_listbox.configure(yscrollcommand=id_yscroll.set)
        id_xscroll = ttk.Scrollbar(id_list_frame, orient=tk.HORIZONTAL, command=self.onboarding_id_listbox.xview)
        self.onboarding_id_listbox.configure(xscrollcommand=id_xscroll.set)
        self.onboarding_id_listbox.grid(row=0, column=0, sticky=tk.NSEW)
        id_yscroll.grid(row=0, column=1, sticky=tk.NS)
        id_xscroll.grid(row=1, column=0, sticky=tk.EW)
        id_list_frame.grid_rowconfigure(0, weight=1)
        id_list_frame.grid_columnconfigure(0, weight=1)

        # 入職日期多選列表框
        ttk.Label(onboarding_filter_frame, text="入職日期：").grid(row=0, column=2, sticky=tk.NW, pady=5, padx=(10,0))
        date_list_frame = ttk.Frame(onboarding_filter_frame)
        date_list_frame.grid(row=0, column=3, sticky=tk.NSEW, pady=5)
        self.onboarding_date_listbox = tk.Listbox(date_list_frame, selectmode=tk.MULTIPLE, exportselection=False, height=6, width=25)
        date_yscroll = ttk.Scrollbar(date_list_frame, orient=tk.VERTICAL, command=self.onboarding_date_listbox.yview)
        self.onboarding_date_listbox.configure(yscrollcommand=date_yscroll.set)
        date_xscroll = ttk.Scrollbar(date_list_frame, orient=tk.HORIZONTAL, command=self.onboarding_date_listbox.xview)
        self.onboarding_date_listbox.configure(xscrollcommand=date_xscroll.set)
        self.onboarding_date_listbox.grid(row=0, column=0, sticky=tk.NSEW)
        date_yscroll.grid(row=0, column=1, sticky=tk.NS)
        date_xscroll.grid(row=1, column=0, sticky=tk.EW)
        date_list_frame.grid_rowconfigure(0, weight=1)
        date_list_frame.grid_columnconfigure(0, weight=1)

        # 分判多選列表框
        ttk.Label(onboarding_filter_frame, text="分判：").grid(row=0, column=4, sticky=tk.NW, pady=5, padx=(10,0))
        subcontractor_list_frame = ttk.Frame(onboarding_filter_frame)
        subcontractor_list_frame.grid(row=0, column=5, sticky=tk.NSEW, pady=5)
        self.onboarding_subcontractor_listbox = tk.Listbox(subcontractor_list_frame, selectmode=tk.MULTIPLE, exportselection=False, height=6, width=25)
        sub_yscroll = ttk.Scrollbar(subcontractor_list_frame, orient=tk.VERTICAL, command=self.onboarding_subcontractor_listbox.yview)
        self.onboarding_subcontractor_listbox.configure(yscrollcommand=sub_yscroll.set)
        sub_xscroll = ttk.Scrollbar(subcontractor_list_frame, orient=tk.HORIZONTAL, command=self.onboarding_subcontractor_listbox.xview)
        self.onboarding_subcontractor_listbox.configure(xscrollcommand=sub_xscroll.set)
        self.onboarding_subcontractor_listbox.grid(row=0, column=0, sticky=tk.NSEW)
        sub_yscroll.grid(row=0, column=1, sticky=tk.NS)
        sub_xscroll.grid(row=1, column=0, sticky=tk.EW)
        subcontractor_list_frame.grid_rowconfigure(0, weight=1)
        subcontractor_list_frame.grid_columnconfigure(0, weight=1)

        # 配置列寬度
        onboarding_filter_frame.grid_columnconfigure(1, weight=1)  # ID列表框
        onboarding_filter_frame.grid_columnconfigure(3, weight=1)  # 日期列表框
        onboarding_filter_frame.grid_columnconfigure(5, weight=1)  # 分判列表框

        # 預覽區域
        preview_frame = ttk.LabelFrame(main_frame, text="預覽", padding=10)
        preview_frame.pack(fill=tk.BOTH, expand=True)

        self.onboarding_preview_table = ttk.Treeview(preview_frame, show='headings')
        preview_yscroll = ttk.Scrollbar(preview_frame, orient=tk.VERTICAL, command=self.onboarding_preview_table.yview)
        preview_xscroll = ttk.Scrollbar(preview_frame, orient=tk.HORIZONTAL, command=self.onboarding_preview_table.xview)
        self.onboarding_preview_table.configure(yscrollcommand=preview_yscroll.set, xscrollcommand=preview_xscroll.set)
        
        preview_yscroll.pack(side=tk.RIGHT, fill=tk.Y)
        preview_xscroll.pack(side=tk.BOTTOM, fill=tk.X)
        self.onboarding_preview_table.pack(fill=tk.BOTH, expand=True)

        # 添加狀態欄
        status_bar = ttk.Label(self.root, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W)
        status_bar.pack(side=tk.BOTTOM, fill=tk.X)

    def load_onboarding_config(self):
        """載入入職模式的配置"""
        config = self.load_config('onboarding')
        self.onboarding_excel_file.set(config.get('excel_file', ''))
        self.onboarding_template_file.set(config.get('template_file', ''))
        self.onboarding_output_path.set(config.get('output_path', ''))

    def save_onboarding_config(self):
        """保存入職模式的配置"""
        config = {
            'excel_file': self.onboarding_excel_file.get(),
            'template_file': self.onboarding_template_file.get(),
            'output_path': self.onboarding_output_path.get()
        }
        self.save_config('onboarding', config)

    def select_onboarding_excel(self):
        """選擇入職Excel文件"""
        file_path = filedialog.askopenfilename(
            title="選擇入職Excel文件",
            filetypes=[("Excel files", "*.xlsx *.xls"), ("All files", "*.*")],
            initialdir=os.path.dirname(self.onboarding_excel_file.get()) if self.onboarding_excel_file.get() else None
        )
        if file_path:
            self.onboarding_excel_file.set(file_path)
            self.load_onboarding_excel_data()
            self.save_onboarding_config()
            self.status_var.set(f"已選擇Excel文件: {os.path.basename(file_path)}")

    def load_onboarding_excel_data(self):
        if not self.onboarding_excel_file.get():
            messagebox.showerror("錯誤", "請先選擇Excel文件")
            self.onboarding_excel_data = None # Ensure data is None if no file
            self.filtered_onboarding_data = None # Initialize filtered data to None
            return
        try:
            self.onboarding_excel_data = pd.read_excel(self.onboarding_excel_file.get(), engine='openpyxl')
            # 更新預覽表
            self.update_onboarding_preview_table(self.onboarding_excel_data)
            # 更新ID、入職日期、分判的列表框選項
            self.populate_onboarding_listboxes()
            # If read_excel was successful, onboarding_excel_data is a DataFrame (possibly empty)
            self.filtered_onboarding_data = self.onboarding_excel_data.copy()
        except Exception as e:
            messagebox.showerror("錯誤", f"讀取Excel文件失敗: {e}")
            self.onboarding_excel_data = None
            self.filtered_onboarding_data = None # Initialize filtered data to None on error

    def update_onboarding_preview_table(self, df):
        # 清除舊數據
        for i in self.onboarding_preview_table.get_children():
            self.onboarding_preview_table.delete(i)
        if df is None or df.empty:
            return
        # 設置列
        self.onboarding_preview_table["columns"] = list(df.columns)
        for col in df.columns:
            self.onboarding_preview_table.heading(col, text=col)
            self.onboarding_preview_table.column(col, width=100, anchor='w') # 可調整寬度
        # 插入數據
        for index, row in df.iterrows():
            self.onboarding_preview_table.insert("", "end", values=list(row))

    def populate_onboarding_listboxes(self):
        if self.onboarding_excel_data is None or self.onboarding_excel_data.empty:
            return
        # 1. 填充ID列表框（A列，索引0）
        id_col = self.onboarding_excel_data.iloc[:, 0].astype(str)
        ids = sorted(set([i for i in id_col if i.lower() != 'nan' and i.strip() != '']))
        self.onboarding_id_listbox.delete(0, tk.END)
        for i in ids:
            self.onboarding_id_listbox.insert(tk.END, i)
        
        # 移除对不存在控件的引用
        # 2. 填充"職安證號碼"列表框 - 界面中不存在此控件
        # if '職安証號碼' in self.onboarding_excel_data.columns:
        #     cic_numbers = sorted(self.onboarding_excel_data['職安証號碼'].dropna().unique())
        #     for cic in cic_numbers:
        #         self.onboarding_cic_listbox.insert(tk.END, cic)
        
        # 3. 填充"姓名"列表框 - 界面中不存在此控件
        # if '姓名' in self.onboarding_excel_data.columns:
        #     names = sorted(self.onboarding_excel_data['姓名'].dropna().unique())
        #     for name in names:
        #         self.onboarding_name_listbox.insert(tk.END, name)
        
        # 4. 填充"入職日期"列表框（只顯示yyyy-mm-dd）
        if '入職日期' in self.onboarding_excel_data.columns:
            dates = []
            for val in self.onboarding_excel_data['入職日期'].dropna():
                if pd.notna(val) and str(val).lower() != 'nan' and str(val).strip() != '':
                    try:
                        # 轉為datetime再格式化
                        date_obj = pd.to_datetime(val)
                        date_str = date_obj.strftime('%Y-%m-%d')
                        dates.append(date_str)
                    except Exception as e:
                        print(f"處理入職日期時出錯: {val}, 錯誤: {e}")
                        # 無法轉換則保留原始字符串
                        dates.append(str(val))
            # 去重並排序
            sorted_dates = sorted(set(dates), reverse=True)
            self.onboarding_date_listbox.delete(0, tk.END)
            for date_str in sorted_dates:
                self.onboarding_date_listbox.insert(tk.END, date_str)
        
        # 5. 填充"分判"列表框
        if '分判' in self.onboarding_excel_data.columns:
            subcontractors = sorted(self.onboarding_excel_data['分判'].astype(str).unique())
            for sub in subcontractors:
                if sub.lower() != 'nan' and sub.strip() != '':
                    self.onboarding_subcontractor_listbox.insert(tk.END, sub)

    def select_onboarding_template(self):
        """選擇入職Word模板文件"""
        file_path = filedialog.askopenfilename(
            title="選擇入職Word模板文件",
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")],
            initialdir=os.path.dirname(self.onboarding_template_file.get()) if self.onboarding_template_file.get() else None
        )
        if file_path:
            self.onboarding_template_file.set(file_path)
            self.save_onboarding_config()
            self.status_var.set(f"已選擇模板: {os.path.basename(file_path)}")
            
            # 嘗試讀取模板中的欄位信息
            try:
                with MailMerge(file_path) as document:
                    fields = document.get_merge_fields()
                    if fields:
                        field_list = ", ".join(fields)
                        messagebox.showinfo("模板欄位", f"模板包含以下合併欄位:\n{field_list}")
            except Exception as e:
                messagebox.showwarning("警告", f"讀取模板欄位時出錯: {e}")

    def select_onboarding_output(self):
        """選擇入職文件輸出位置"""
        dir_path = filedialog.askdirectory(
            title="選擇入職文件輸出目錄",
            initialdir=self.onboarding_output_path.get() if self.onboarding_output_path.get() else None
        )
        if dir_path:
            self.onboarding_output_path.set(dir_path)
            self.save_onboarding_config()
            self.status_var.set(f"已選擇輸出目錄: {dir_path}")

    def apply_onboarding_filter(self):
        if self.onboarding_excel_data is None or self.onboarding_excel_data.empty:
            messagebox.showinfo("提示", "請先載入Excel數據")
            return

        filtered_df = self.onboarding_excel_data.copy()
        active_filters = False

        # 先根據ID篩選（A列，索引0）
        selected_id_indices = self.onboarding_id_listbox.curselection()
        if selected_id_indices:
            active_filters = True
            selected_ids = [self.onboarding_id_listbox.get(i) for i in selected_id_indices]
            id_col = self.onboarding_excel_data.iloc[:, 0].astype(str)
            filtered_df = filtered_df[id_col.isin(selected_ids)]

        # 移除对不存在的 CIC 列表框和姓名列表框的引用
        # 获取CIC号码筛选条件
        # selected_cic_indices = self.onboarding_cic_listbox.curselection()
        # if selected_cic_indices:
        #     active_filters = True
        #     selected_cics = [self.onboarding_cic_listbox.get(i) for i in selected_cic_indices]
        #     filtered_df = filtered_df[filtered_df['職安証號碼'].astype(str).isin(selected_cics)]

        # 获取姓名筛选条件
        # selected_names_indices = self.onboarding_name_listbox.curselection()
        # if selected_names_indices:
        #     active_filters = True
        #     selected_names = [self.onboarding_name_listbox.get(i) for i in selected_names_indices]
        #     filtered_df = filtered_df[filtered_df['姓名'].astype(str).isin(selected_names)]

        # 獲取並應用"入職日期"篩選條件
        selected_dates_indices = self.onboarding_date_listbox.curselection()
        if selected_dates_indices:
            active_filters = True
            selected_dates = [self.onboarding_date_listbox.get(i) for i in selected_dates_indices]
            try:
                date_mask = pd.Series([False] * len(filtered_df))
                for i, val in enumerate(filtered_df['入職日期']):
                    if pd.notna(val):
                        try:
                            val_str = pd.to_datetime(val).strftime('%Y-%m-%d')
                            if val_str in selected_dates:
                                date_mask.iloc[i] = True
                        except Exception as e:
                            print(f"入職日期篩選出錯: {val}, 錯誤: {e}")
                filtered_df = filtered_df[date_mask]
            except Exception as e:
                print(f"篩選入職日期時發生錯誤: {e}")
                messagebox.showwarning("篩選警告", f"篩選入職日期時發生錯誤，該篩選可能未生效。\n錯誤詳情: {e}")

        # 獲取並應用"分判"篩選條件
        selected_subcontractors_indices = self.onboarding_subcontractor_listbox.curselection()
        if selected_subcontractors_indices:
            active_filters = True
            selected_subcontractors = [self.onboarding_subcontractor_listbox.get(i) for i in selected_subcontractors_indices]
            filtered_df = filtered_df[filtered_df['分判'].astype(str).isin(selected_subcontractors)]

        # 更新預覽表格（修正控件名稱）
        self.update_onboarding_preview_table(filtered_df)
        self.onboarding_filtered_data = filtered_df  # 保存篩選後的資料
        
        # 更新狀態欄（修正控件名稱）
        record_count = len(filtered_df) if filtered_df is not None else 0
        self.status_var.set(f"已篩選 {record_count} 筆記錄")

        return filtered_df

    def clear_onboarding_filter(self):
        # 清除ID列表框的選擇
        if hasattr(self, 'onboarding_id_listbox'):
            self.onboarding_id_listbox.selection_clear(0, tk.END)
        
        # 清除入職日期列表框的選擇
        if hasattr(self, 'onboarding_date_listbox'):
            self.onboarding_date_listbox.selection_clear(0, tk.END)
        
        # 清除分判列表框的選擇
        if hasattr(self, 'onboarding_subcontractor_listbox'):
            self.onboarding_subcontractor_listbox.selection_clear(0, tk.END)
        
        # 如果已加載Excel數據，則更新預覽表以顯示所有數據
        if self.onboarding_excel_data is not None and not self.onboarding_excel_data.empty:
            self.update_onboarding_preview_table(self.onboarding_excel_data)
        else:
            # 如果沒有數據，可以選擇清空預覽表或顯示提示
            # 這裡假設 update_onboarding_preview_table 能夠處理空的 DataFrame
            # 或者您可以直接清空 Treeview
            if hasattr(self, 'onboarding_preview_table'):
                 for i in self.onboarding_preview_table.get_children():
                    self.onboarding_preview_table.delete(i)
            # messagebox.showinfo("提示", "沒有數據可顯示或Excel文件尚未加載。")
        
        # 可以選擇性地添加一個消息提示用戶篩選已清除
        # messagebox.showinfo("提示", "篩選已清除，顯示所有數據。")

    def generate_onboarding_word(self):
        """生成入職Word文件"""
        if self.onboarding_excel_data is None:
            messagebox.showwarning("警告", "請先載入Excel文件")
            return

        template_path = self.onboarding_template_file.get()
        if not template_path or not os.path.exists(template_path):
            # 如果模板路径为空或不存在，提示用户选择模板
            messagebox.showwarning("警告", "請選擇入職Word模板文件\n\n所有模板必須從外部讀取，系統不提供默認模板")
            self.select_onboarding_template()
            return
            
        # 获取当前选择的模板路径（可能已经被select_template更新）
        template_path = self.onboarding_template_file.get()
        if not template_path or not os.path.exists(template_path):
            # 用户取消了模板选择
            messagebox.showinfo("提示", "已取消選擇模板，無法生成Word文件")
            return

        output_dir = self.onboarding_output_path.get()
        if not output_dir:
            # 如果未指定輸出目錄，則默認在程序目錄下創建 output_onboarding 文件夾
            output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output_onboarding")
            self.onboarding_output_path.set(output_dir)
            self.save_onboarding_config()
        
        os.makedirs(output_dir, exist_ok=True)

        # 獲取預覽表中的數據 (假設是篩選後的或全部數據)
        # 我們需要從 self.onboarding_preview_table 獲取數據，或者直接使用篩選後的 DataFrame
        # 為了簡化，我們假設 apply_onboarding_filter 後，filtered_df 是最新的篩選結果
        # 如果沒有篩選，則使用 self.onboarding_excel_data
        
        # 重新獲取當前預覽表中的數據
        current_preview_df = None
        if self.onboarding_preview_table.get_children():
            cols = self.onboarding_preview_table["columns"]
            data_rows = []
            for item in self.onboarding_preview_table.get_children():
                data_rows.append(self.onboarding_preview_table.item(item)["values"])
            if data_rows:
                current_preview_df = pd.DataFrame(data_rows, columns=cols)
        
        if current_preview_df is None or current_preview_df.empty:
            # 如果預覽表是空的，但原始數據存在，則使用原始數據
            if self.onboarding_excel_data is not None and not self.onboarding_excel_data.empty:
                current_preview_df = self.onboarding_excel_data
            else:
                messagebox.showinfo("提示", "沒有數據可供生成Word文件")
                return

        # Word模板占位符与Excel列名的对应关系
        # 注意：Excel列名需要与您实际文件中的列名完全一致
        merge_fields_map = {
            'name': '工人姓名',
            'trade': '工種',
            'Green_Card': '平安咭號編號',
            'Green_Date': '平安咭有效日期',
            'date': '入職日期',
            'cic_card': '工人註冊證編號',
            'cic_date': '工人註冊證有效日期',
            'phone': '聯絡電話號碼',
            'subcon': '分判',
            'salary': '日薪',
            'overtime': '超時薪金方法計算',
            'E_contact': '緊急聯絡人',
            'E_phone': '緊急聯絡人電話',
            'workyear': '從事上述工種經驗(多少年)',
            'level': '註冊電業人員級別 (如有)',
            'reason': '未能操作上述工作的原因',
            'work_process': '工人負責工序 ',  # 新增此映射
        }

        # 检查是否有分判商列
        if '分判' not in current_preview_df.columns:
            messagebox.showwarning("警告", "Excel數據中缺少'分判'列，這可能會影響文件生成")
        
        # 按分判商分组生成文件
        subcontractors = current_preview_df['分判'].unique() if '分判' in current_preview_df.columns else ['未知分判商']
        
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        
        total_files = 0
        for subcon in subcontractors:
            if pd.isna(subcon):
                subcon_name = "未知分判商"
            else:
                subcon_name = str(subcon)
            
            # 获取该分判商的所有记录
            if '分判' in current_preview_df.columns:
                subcon_df = current_preview_df[current_preview_df['分判'] == subcon]
            else:
                subcon_df = current_preview_df
            
            if subcon_df.empty:
                continue
            
            # 为每个记录生成一个Word文件
            for index, row in subcon_df.iterrows():
                try:
                    # 创建一个字典，用于映射Word模板中的占位符
                    merge_fields = {}
                    for word_field, excel_col in merge_fields_map.items():
                        if excel_col in row:
                            # 处理日期格式
                            if excel_col in ['入職日期', '平安咭有效日期', '工人註冊證有效日期'] and pd.notna(row[excel_col]):
                                try:
                                    date_value = pd.to_datetime(row[excel_col])
                                    merge_fields[word_field] = date_value.strftime('%Y-%m-%d')
                                except:
                                    merge_fields[word_field] = str(row[excel_col])
                            else:
                                merge_fields[word_field] = str(row[excel_col]) if pd.notna(row[excel_col]) else ""
                    
                    # 获取工人姓名作为文件名的一部分
                    worker_name = row['工人姓名'] if '工人姓名' in row and pd.notna(row['工人姓名']) else f"工人{index}"
                    
                    # 创建输出文件路径
                    output_filename = os.path.join(output_dir, f"入職文件- {subcon_name}_{worker_name}_{timestamp}.docx")
                    
                    # 使用 docx-mailmerge 进行邮件合并
                    document = MailMerge(template_path)
                    document.merge(**merge_fields)
                    document.write(output_filename)
                    
                    total_files += 1
                    self.status_var.set(f"正在生成Word文件 {total_files}/{len(subcon_df)}")
                    
                except Exception as e:
                    messagebox.showerror("錯誤", f"生成 {worker_name} 的Word文件時出錯: {e}")
        
        if total_files > 0:
            messagebox.showinfo("成功", f"已成功生成 {total_files} 個Word文件")
            # 保存配置
            self.save_onboarding_config()
        else:
            messagebox.showwarning("警告", "沒有生成任何Word文件")
        
        self.status_var.set(f"已生成 {total_files} 個Word文件")

    def start_morning_briefing_mode(self):
        """啟動早操（早操合并工具）模式界面"""
        # 清理選擇界面
        for widget in self.root.winfo_children():
            widget.destroy()

        self.root.title("早操合并工具")
        self.root.geometry("900x600")

        # 初始化變量
        self.excel_file = tk.StringVar()
        self.output_path = tk.StringVar()
        self.template_file_path = tk.StringVar()
        # 移除默认模板设置，改为空字符串
        self.template_file_path.set("")
        self.excel_data = None
        
        self.company_filter = tk.StringVar()
        self.subject_var = tk.StringVar()
        self.trainer_var = tk.StringVar()
        self.time_var = tk.StringVar()
        self.place_var = tk.StringVar()
        
        self.company_combo = None
        
        # 加載早操模式的配置
        self.load_morning_config()

        # 創建界面
        self.create_ui()

    def load_config(self, mode):
        """統一的配置載入函數"""
        try:
            if not os.path.exists(self.config_file):
                return {}
                
            with open(self.config_file, 'r', encoding='utf-8') as f:
                config = json.load(f)
                return config.get(mode, {})
        except Exception as e:
            print(f"載入{mode}模式配置錯誤: {e}")
            return {}

    def save_config(self, mode, config_data):
        """統一的配置保存函數"""
        try:
            # 讀取現有配置
            current_config = {}
            if os.path.exists(self.config_file):
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    current_config = json.load(f)
            
            # 更新指定模式的配置
            current_config[mode] = config_data
            
            # 保存更新後的配置
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(current_config, f, ensure_ascii=False, indent=4)
        except Exception as e:
            messagebox.showerror("錯誤", f"保存{mode}模式配置失敗: {e}")

    def create_ui(self):
        """創建早操模式的GUI界面"""
        main_frame = ttk.Frame(self.root, padding=10)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 按鈕區域
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=(0, 10))
        
        # 功能按鈕
        buttons = [
            ("輸入Excel", self.select_excel),
            ("選擇模板", self.select_template),
            ("輸出位置選擇", self.select_output),
            ("套用篩選", self.apply_filter),
            ("清除篩選", self.clear_filter),
            ("重新載入", self.reload_morning_data),
            ("生成Word", self.generate_word),
            ("返回主菜單", self.return_to_main)
        ]
        
        for i, (btn_text, btn_command) in enumerate(buttons):
            btn = ttk.Button(btn_frame, text=btn_text, command=btn_command)
            btn.grid(row=0, column=i, padx=5)
        
        # 設置區域
        settings_frame = ttk.LabelFrame(main_frame, text="設置", padding=10)
        settings_frame.pack(fill=tk.X, pady=(0, 10))
        
        # Excel文件路徑
        ttk.Label(settings_frame, text="Excel文件:").grid(row=0, column=0, sticky=tk.W, pady=5)
        ttk.Entry(settings_frame, textvariable=self.excel_file, width=50).grid(row=0, column=1, sticky=tk.W+tk.E, padx=5)
        
        # 模板文件路徑
        ttk.Label(settings_frame, text="模板文件:").grid(row=1, column=0, sticky=tk.W, pady=5)
        ttk.Entry(settings_frame, textvariable=self.template_file_path, width=50).grid(row=1, column=1, sticky=tk.W+tk.E, padx=5)
        
        # 輸出位置
        ttk.Label(settings_frame, text="輸出位置:").grid(row=2, column=0, sticky=tk.W, pady=5)
        ttk.Entry(settings_frame, textvariable=self.output_path, width=50).grid(row=2, column=1, sticky=tk.W+tk.E, padx=5)
        
        # 篩選區域
        filter_frame = ttk.LabelFrame(main_frame, text="篩選條件", padding=10)
        filter_frame.pack(fill=tk.X, pady=(0, 10))
        
        # 第一行
        ttk.Label(filter_frame, text="工程項目及\n項目編號：").grid(row=0, column=0, sticky=tk.W, pady=5)
        ttk.Label(filter_frame, text="HE202202\nDesign and Construction of a District Court Building at\nCaroline Hill Road").grid(row=0, column=1, sticky=tk.W, pady=5)
        
        # 公司名稱多選 - 包含滾動條
        ttk.Label(filter_frame, text="公司名稱：").grid(row=0, column=2, sticky=tk.NW, pady=5)
        
        company_list_frame = ttk.Frame(filter_frame)
        company_list_frame.grid(row=0, column=3, sticky=tk.NSEW, pady=5)
        
        self.company_listbox = tk.Listbox(company_list_frame, selectmode=tk.MULTIPLE, exportselection=False, height=6, width=30)
        
        company_yscroll = ttk.Scrollbar(company_list_frame, orient=tk.VERTICAL, command=self.company_listbox.yview)
        self.company_listbox.configure(yscrollcommand=company_yscroll.set)
        
        company_xscroll = ttk.Scrollbar(company_list_frame, orient=tk.HORIZONTAL, command=self.company_listbox.xview)
        self.company_listbox.configure(xscrollcommand=company_xscroll.set)
        
        self.company_listbox.grid(row=0, column=0, sticky=tk.NSEW)
        company_yscroll.grid(row=0, column=1, sticky=tk.NS)
        company_xscroll.grid(row=1, column=0, sticky=tk.EW)
        
        company_list_frame.grid_rowconfigure(0, weight=1)
        company_list_frame.grid_columnconfigure(0, weight=1)

        # 日期多選列表框
        ttk.Label(filter_frame, text="日期：").grid(row=1, column=2, sticky=tk.NW, pady=5)
        date_list_frame = ttk.Frame(filter_frame)
        date_list_frame.grid(row=1, column=3, sticky=tk.NSEW, pady=5)
        
        self.date_listbox = tk.Listbox(date_list_frame, selectmode=tk.MULTIPLE, exportselection=False, height=6, width=30)
        
        date_yscroll = ttk.Scrollbar(date_list_frame, orient=tk.VERTICAL, command=self.date_listbox.yview)
        self.date_listbox.configure(yscrollcommand=date_yscroll.set)
        
        date_xscroll = ttk.Scrollbar(date_list_frame, orient=tk.HORIZONTAL, command=self.date_listbox.xview)
        self.date_listbox.configure(xscrollcommand=date_xscroll.set)
        
        self.date_listbox.grid(row=0, column=0, sticky=tk.NSEW)
        date_yscroll.grid(row=0, column=1, sticky=tk.NS)
        date_xscroll.grid(row=1, column=0, sticky=tk.EW)
        
        date_list_frame.grid_rowconfigure(0, weight=1)
        date_list_frame.grid_columnconfigure(0, weight=1)
        
        # 添加合併選項的checkbox
        self.merge_companies_var = tk.BooleanVar()
        self.merge_companies_var.set(False)  # 預設不合併
        merge_checkbox = ttk.Checkbutton(filter_frame, text="合併多個公司到同一個Word檔案", variable=self.merge_companies_var)
        merge_checkbox.grid(row=1, column=0, columnspan=2, sticky=tk.W, pady=5)
        
        # 課程題目、講授者等其他UI元素
        ttk.Label(filter_frame, text="課程題目：").grid(row=2, column=0, sticky=tk.W, pady=5)
        self.subject_combo = ttk.Combobox(filter_frame, textvariable=self.subject_var, width=20)
        self.subject_combo['values'] = getattr(self, 'subject_options', [])
        self.subject_combo.grid(row=2, column=1, sticky=tk.W, pady=5)
        
        ttk.Label(filter_frame, text="講授者：").grid(row=2, column=2, sticky=tk.W, pady=5)
        self.trainer_combo = ttk.Combobox(filter_frame, textvariable=self.trainer_var, width=20)
        self.trainer_combo['values'] = getattr(self, 'trainer_options', [])
        self.trainer_combo.grid(row=2, column=3, sticky=tk.W, pady=5)
        
        ttk.Label(filter_frame, text="時間：").grid(row=3, column=0, sticky=tk.W, pady=5)
        time_entry = ttk.Entry(filter_frame, textvariable=self.time_var, width=20)
        time_entry.grid(row=3, column=1, sticky=tk.W, pady=5)
        
        ttk.Label(filter_frame, text="地點：").grid(row=3, column=2, sticky=tk.W, pady=5)
        place_entry = ttk.Entry(filter_frame, textvariable=self.place_var, width=20)
        place_entry.grid(row=3, column=3, sticky=tk.W, pady=5)
        
        # 預覽區域
        preview_frame = ttk.LabelFrame(main_frame, text="預覽區", padding=10)
        preview_frame.pack(fill=tk.BOTH, expand=True)
        
        # 建立預覽表格
        columns = ("CWRNo", "Company", "Attendee Chinese Name", "Attendance Date")
        self.preview_tree = ttk.Treeview(preview_frame, columns=columns, show="headings")
        
        for col in columns:
            self.preview_tree.heading(col, text=col)
            self.preview_tree.column(col, width=100)
        
        # 添加預覽表格的滾動條
        scrollbar = ttk.Scrollbar(preview_frame, orient=tk.VERTICAL, command=self.preview_tree.yview)
        self.preview_tree.configure(yscrollcommand=scrollbar.set)
        
        self.preview_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # 添加狀態欄
        self.status_var = tk.StringVar()
        self.status_var.set("就緒")
        status_bar = ttk.Label(self.root, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W)
        status_bar.pack(side=tk.BOTTOM, fill=tk.X)

    def select_excel(self):
        """選擇Excel文件"""
        file_path = filedialog.askopenfilename(
            title="選擇Excel文件",
            filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")]
        )
        if file_path:
            self.excel_file.set(file_path)
            self.load_excel_data()
            self.save_morning_config()
    
    def select_template(self):
        """選擇Word模板文件"""
        file_path = filedialog.askopenfilename(
            title="選擇Word模板文件",
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")],
            initialdir=os.path.dirname(self.template_file_path.get()) if self.template_file_path.get() else None
        )
        if file_path:
            self.template_file_path.set(file_path)
            self.status_var.set(f"已選擇模板: {os.path.basename(file_path)}")
            self.save_morning_config()
            
            # 嘗試讀取模板中的欄位信息
            try:
                with MailMerge(file_path) as document:
                    fields = document.get_merge_fields()
                    if fields:
                        field_list = ", ".join(fields)
                        messagebox.showinfo("模板欄位", f"模板包含以下合併欄位:\n{field_list}")
            except Exception as e:
                messagebox.showwarning("警告", f"讀取模板欄位時出錯: {e}")
    
    def select_output(self):
        """選擇輸出位置"""
        dir_path = filedialog.askdirectory(title="選擇輸出目錄")
        if dir_path:
            self.output_path.set(dir_path)
            self.save_morning_config()
    
    def load_excel_data(self):
        """載入Excel數據"""
        try:
            excel_file = self.excel_file.get()
            if not excel_file:
                return
            
            self.excel_data = pd.read_excel(excel_file)
            print(f"成功載入Excel文件，共 {len(self.excel_data)} 行數據")
            
            # 清空現有的表格數據
            for item in self.preview_tree.get_children():
                self.preview_tree.delete(item)
            
            try:
                company_col = self.excel_data.iloc[:, 21]  # V列是索引21
                date_col = self.excel_data.iloc[:, 6]      # G列是索引6（修正）
                cwrno_col = self.excel_data.iloc[:, 8]     # I列是索引8
                name_col = self.excel_data.iloc[:, 11]     # L列是索引11
                
                print(f"成功讀取特定欄位")
                print(f"公司欄位示例值: {company_col.iloc[0] if len(company_col) > 0 else 'N/A'}")
                print(f"日期欄位示例值: {date_col.iloc[0] if len(date_col) > 0 else 'N/A'}")
                
                company_col = company_col.astype(str)
                date_col = date_col.astype(str)
                cwrno_col = cwrno_col.astype(str)
                name_col = name_col.astype(str)
                
                companies = sorted(company_col.dropna().unique().tolist())
                companies = [c for c in companies if c.lower() != 'nan' and c.strip() != '']
                
                print(f"找到 {len(companies)} 個公司選項")
                
                # 日期下拉選單（G列）
                date_values = date_col.dropna().unique().tolist()
                dates = [d for d in date_values if d.lower() != 'nan' and d.strip() != '']
                print(f"找到 {len(dates)} 個日期選項")
                if dates and len(dates) > 0:
                    print(f"日期排序結果 (前5項): {dates[:5]}")
                
                if companies and len(companies) > 0:
                    self.company_listbox.delete(0, tk.END)
                    for c in companies:
                        self.company_listbox.insert(tk.END, c)
                if dates and len(dates) > 0:
                    self.date_listbox.delete(0, tk.END)
                    for d in dates:
                        self.date_listbox.insert(tk.END, d)
            except Exception as e:
                print(f"處理Excel欄位時出錯: {e}")
                messagebox.showwarning("警告", f"讀取Excel欄位時出現問題，請檢查Excel文件格式: {e}")
                return
            
            # 顯示數據在預覽表格中
            if not self.excel_data.empty:
                record_count = 0
                for i in range(len(self.excel_data)):
                    try:
                        self.preview_tree.insert("", "end", values=(
                            cwrno_col.iloc[i],      # CWRNo (I列)
                            company_col.iloc[i],    # Company (V列)
                            name_col.iloc[i],       # Attendee Chinese Name (L列)
                            date_col.iloc[i]        # Attendance Date (G列)
                        ))
                        record_count += 1
                    except Exception as e:
                        print(f"插入預覽表格第 {i} 行時出錯: {e}")
                        continue
                print(f"已插入 {record_count} 筆記錄到預覽表格")
            self.status_var.set(f"已載入 {len(self.excel_data)} 筆記錄")
        except Exception as e:
            print(f"載入Excel數據失敗: {e}")
            messagebox.showerror("錯誤", f"載入Excel數據失敗: {e}")
            self.status_var.set("載入Excel數據失敗")

    def apply_filter(self):
        if self.excel_data is None:
            messagebox.showwarning("警告", "請先載入Excel文件")
            return
        selected_companies = [self.company_listbox.get(i) for i in self.company_listbox.curselection()]
        selected_dates = [self.date_listbox.get(i) for i in self.date_listbox.curselection()]
        print(f"篩選條件 - 公司: {selected_companies}, 日期: {selected_dates}")
        for item in self.preview_tree.get_children():
            self.preview_tree.delete(item)
        try:
            company_col = self.excel_data.iloc[:, 21].astype(str)
            date_col = self.excel_data.iloc[:, 6].astype(str)  # G列是索引6（修正）
            cwrno_col = self.excel_data.iloc[:, 8].astype(str)
            name_col = self.excel_data.iloc[:, 11].astype(str)
            
            mask = pd.Series([True] * len(self.excel_data))
            if selected_companies:
                mask = mask & company_col.isin(selected_companies)
            if selected_dates:
                date_mask = date_col.isin(selected_dates)
                mask = mask & date_mask
            filtered_indices = mask[mask].index
            record_count = 0
            for i in filtered_indices:
                try:
                    self.preview_tree.insert("", "end", values=(
                        cwrno_col.iloc[i],
                        company_col.iloc[i],
                        name_col.iloc[i],
                        date_col.iloc[i]
                    ))
                    record_count += 1
                except Exception as e:
                    print(f"插入篩選記錄 {i} 時出錯: {e}")
                    continue
            self.status_var.set(f"已篩選 {len(filtered_indices)} 筆記錄")
        except Exception as e:
            print(f"篩選數據時出錯: {e}")
            messagebox.showerror("錯誤", f"篩選數據時出錯: {e}")

    def clear_filter(self):
        # 清除公司列表框的選擇
        if hasattr(self, 'company_listbox'):
            self.company_listbox.selection_clear(0, tk.END)
        
        # 清除日期列表框的選擇
        if hasattr(self, 'date_listbox'):
            self.date_listbox.selection_clear(0, tk.END)
        
        # 清除分判列表框的選擇
        if hasattr(self, 'company_combo'):
            self.company_combo.set('') # 假設 company_combo 是 Combobox

        # 如果已加載Excel數據，則更新預覽表以顯示所有數據
        if self.excel_data is not None and not self.excel_data.empty:
            self.update_preview_table(self.preview_tree, self.excel_data) # 注意這裡是 update_preview_table
        else:
            # 如果沒有數據，清空預覽表
            if hasattr(self, 'preview_tree'): # 早操模式的預覽表名是 preview_tree
                 for i in self.preview_tree.get_children():
                    self.preview_tree.delete(i)
            # messagebox.showinfo("提示", "沒有數據可顯示或Excel文件尚未加載。")

        # messagebox.showinfo("提示", "篩選已清除。")

    def fill_attendance_table_with_company_date(self, template_path: str, output_path: str, data: pd.DataFrame, company_name: str, specific_date: str):
        """
        將特定公司和日期的資料依序填入Word出席表格，支援多頁且每頁保留標題，並替換佔位符。
        :param template_path: Word模板路徑
        :param output_path: 輸出Word路徑
        :param data: pandas DataFrame，欄位需有 'cic', 'chi', 't_date'
        :param company_name: 當前處理的公司名稱，如果為"多個公司"，則從data的company欄位讀取
        :param specific_date: 當前處理的日期
        """
        from copy import deepcopy
        import shutil
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
        import os

        # 首先創建一個臨時文件來保存我們的結果，避免直接修改原模板
        temp_output = output_path + ".temp"
        shutil.copy2(template_path, temp_output)
        
        try:
            # 打開臨時文件進行編輯，使用自定義的create_document函數
            doc = create_document(temp_output)
            
            # 查找目標表格（8列16行）
            table = None
            for t in doc.tables:
                if len(t.rows) >= 16 and len(t.columns) == 8:
                    table = t
                    break
            
            if not table:
                # 如果找不到特定表格，嘗試使用第一個表格或記錄錯誤
                if doc.tables:
                    print("警告: 找不到8欄16行的出席表格，將使用文件中的第一個表格。請檢查模板！")
                    table = doc.tables[0]
                else:
                    raise Exception("模板中沒有找到任何表格，請檢查模板！")

            # 處理資料
            # 確保 'cic', 'chi', 't_date' 欄位存在，若不存在則使用空值或預設值
            records_data = []
            for _, row in data.iterrows():
                cic = str(row['cic']) if 'cic' in row else ''
                chi = str(row['chi']) if 'chi' in row else ''
                # 移除添加公司名稱到姓名的代碼，即使是多個公司模式也不添加
                # t_date 來自 specific_date，因為數據已經按公司和日期分組
                records_data.append([cic, chi, specific_date])
            
            total_pages = (len(records_data) + 29) // 30 # 每頁30筆記錄

            # 儲存原始表格的XML元素和表頭
            orig_tbl_xml = deepcopy(table._element)
            header_cells_content = []
            if len(table.rows) > 0:
                for col_idx in range(len(table.columns)):
                    header_cells_content.append(table.cell(0, col_idx).text)
            else: # 如果表格是空的，創建一個預設的表頭結構
                header_cells_content = [f"Header {i+1}" for i in range(len(table.columns))]

            for page_num in range(total_pages):
                current_table_to_fill = None
                if page_num == 0:
                    current_table_to_fill = table
                else:
                    # 插入分頁符
                    p = doc.add_paragraph()
                    run = p.add_run()
                    run.add_break(WD_BREAK.PAGE)
                    # 複製原始表格
                    new_tbl_element = deepcopy(orig_tbl_xml)
                    doc._body._element.append(new_tbl_element)
                    current_table_to_fill = doc.tables[-1] # 新複製的表格是最後一個
                
                # 填充表頭 (確保即使是新複製的表格也有表頭)
                if len(current_table_to_fill.rows) > 0 and len(header_cells_content) == len(current_table_to_fill.columns):
                    for col_idx, header_text in enumerate(header_cells_content):
                        cell = current_table_to_fill.cell(0, col_idx)
                        cell.text = header_text
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 嚴格取本頁的30筆資料
                page_records = records_data[page_num*30:(page_num+1)*30]
                # 如果記錄不足30筆，用空記錄補齊
                while len(page_records) < 30:
                    page_records.append(['', '', ''])

                # 填充數據行 (從第二行開始，索引為1)
                # 確保表格至少有16行來填充數據 (1行表頭 + 15行數據)
                while len(current_table_to_fill.rows) < 16:
                    current_table_to_fill.add_row()

                for i in range(15): # 每頁15行數據 * 2列 = 30筆記錄
                    data_row_idx = i + 1 # 表格中的數據行索引
                    
                    # 左側 (前4欄: 編號, CIC, 姓名, 日期)
                    current_table_to_fill.cell(data_row_idx, 0).text = f"{page_num*30 + i + 1}."
                    current_table_to_fill.cell(data_row_idx, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    current_table_to_fill.cell(data_row_idx, 1).text = str(page_records[i][0]).replace('\n', '').replace('\r', '').strip()
                    current_table_to_fill.cell(data_row_idx, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    current_table_to_fill.cell(data_row_idx, 2).text = str(page_records[i][1])
                    current_table_to_fill.cell(data_row_idx, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    current_table_to_fill.cell(data_row_idx, 3).text = str(page_records[i][2]) # 日期
                    current_table_to_fill.cell(data_row_idx, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 右側 (後4欄: 編號, CIC, 姓名, 日期)
                    current_table_to_fill.cell(data_row_idx, 4).text = f"{page_num*30 + i + 16}."
                    current_table_to_fill.cell(data_row_idx, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    current_table_to_fill.cell(data_row_idx, 5).text = str(page_records[i+15][0]).replace('\n', '').replace('\r', '').strip()
                    current_table_to_fill.cell(data_row_idx, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    current_table_to_fill.cell(data_row_idx, 6).text = str(page_records[i+15][1])
                    current_table_to_fill.cell(data_row_idx, 6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    current_table_to_fill.cell(data_row_idx, 7).text = str(page_records[i+15][2]) # 日期
                    current_table_to_fill.cell(data_row_idx, 7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 替換佔位符的函數
            def replace_placeholders_in_doc(document, mapping):
                # 替換正文段落
                for p in document.paragraphs:
                    for key, value in mapping.items():
                        if key in p.text:
                            inline = p.runs
                            for item_run in inline:
                                if key in item_run.text:
                                    item_run.text = item_run.text.replace(key, value)
                # 替換正文表格
                for tbl in document.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            for p_in_cell in cell.paragraphs:
                                for key, value in mapping.items():
                                    if key in p_in_cell.text:
                                        inline = p_in_cell.runs
                                        for item_run in inline:
                                            if key in item_run.text:
                                                item_run.text = item_run.text.replace(key, value)
                # 替換所有 section 的 header/footer
                for section in document.sections:
                    for header_footer in [section.header, section.footer]:
                        if header_footer is None: continue
                        for p in header_footer.paragraphs:
                            for key, value in mapping.items():
                                if key in p.text:
                                    inline = p.runs
                                    for item_run in inline:
                                        if key in item_run.text:
                                            item_run.text = item_run.text.replace(key, value)
                        for tbl in header_footer.tables:
                            for row in tbl.rows:
                                for cell in row.cells:
                                    for p_in_cell in cell.paragraphs:
                                        for key, value in mapping.items():
                                            if key in p_in_cell.text:
                                                inline = p_in_cell.runs
                                                for item_run in inline:
                                                    if key in item_run.text:
                                                        item_run.text = item_run.text.replace(key, value)

            # 替換所有欄位，包括公司名稱和日期
            company_display_name = company_name
            if company_name == "多個公司":
                # 如果是多個公司模式，獲取所有不重複的公司名稱
                if 'company' in data.columns:
                    unique_companies = data['company'].unique()
                    company_display_name = ", ".join(unique_companies)
                
            placeholders_map = {
                '«subject»': self.subject_var.get(),
                '«Subject»': self.subject_var.get(),
                '«trainer»': self.trainer_var.get(),
                '«Trainer»': self.trainer_var.get(),
                '«time»': self.time_var.get(),
                '«Time»': self.time_var.get(),
                '«place»': self.place_var.get(),
                '«Place»': self.place_var.get(),
                '«company name»': company_display_name, # 使用處理後的公司名稱
                '«date»': specific_date      # 使用傳入的特定日期
            }
            replace_placeholders_in_doc(doc, placeholders_map)
            
            doc.save(output_path)
            
            print(f"已成功為公司 '{company_display_name}' 和日期 '{specific_date}' 產生出席表：{output_path}")
        except Exception as e:
            print(f"處理文件時出錯: {e}")
            raise e
        finally:
            # 清理臨時文件
            try:
                if os.path.exists(temp_output):
                    os.remove(temp_output)
            except Exception as e_remove:
                print(f"清理臨時文件 {temp_output} 失敗: {e_remove}")
                pass # 即使刪除失敗也繼續

    def generate_word(self):
        """生成Word文件"""
        if self.excel_data is None:
            messagebox.showwarning("警告", "請先載入Excel文件")
            return
            
        template_path = self.template_file_path.get()
        if not template_path or not os.path.exists(template_path):
            # 如果模板路径为空或不存在，提示用户选择模板
            messagebox.showwarning("警告", "請選擇Word模板文件\n\n所有模板必須從外部讀取，系統不提供默認模板")
            self.select_template()
            return
            
        # 获取当前选择的模板路径（可能已经被select_template更新）
        template_path = self.template_file_path.get()
        if not template_path or not os.path.exists(template_path):
            # 用户取消了模板选择
            messagebox.showinfo("提示", "已取消選擇模板，無法生成Word文件")
            return

        output_dir = self.output_path.get()
        if not output_dir:
            output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")
        os.makedirs(output_dir, exist_ok=True)
        selected_companies = [self.company_listbox.get(i) for i in self.company_listbox.curselection()]
        selected_dates = [self.date_listbox.get(i) for i in self.date_listbox.curselection()]
        if not selected_companies or not selected_dates:
            messagebox.showwarning("警告", "請選擇公司與日期")
            return
        
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        total_count = 0
        
        # 檢查是否需要合併多個公司到同一個Word檔案
        merge_companies = self.merge_companies_var.get()
        
        if merge_companies and len(selected_companies) > 1:
            # 合併多個公司到同一個Word檔案
            all_data = []
            for date in selected_dates:
                # 收集所有選定日期和公司的數據
                for company in selected_companies:
                    for item in self.preview_tree.get_children():
                        values = self.preview_tree.item(item)["values"]
                        if values and len(values) >= 4:
                            if str(values[1]) == company and str(values[3]) == date:
                                all_data.append({
                                    "cic": str(values[0]),
                                    "chi": str(values[2]),
                                    "t_date": str(values[3]),
                                    "company": str(values[1])  # 保存公司名稱
                                })
                
                if not all_data:
                    continue
                    
                df = pd.DataFrame(all_data)
                # 使用所有公司名稱作為檔案名的一部分
                company_names = "多公司"
                output_file = os.path.join(output_dir, f"合併文件_{company_names}_{date}_{timestamp}.docx")
                
                try:
                    # 修改fill_attendance_table_with_company_date方法的調用，傳遞所有公司名稱
                    self.fill_attendance_table_with_company_date(template_path, output_file, df, "多個公司", date)
                    total_count += 1
                except Exception as e:
                    messagebox.showerror("錯誤", f"{company_names} {date} 產生失敗: {e}")
                
                # 清空數據，準備下一個日期
                all_data = []
        else:
            # 原有邏輯：每個公司一個Word檔案
            for company in selected_companies:
                for date in selected_dates:
                    group_data = []
                    for item in self.preview_tree.get_children():
                        values = self.preview_tree.item(item)["values"]
                        if values and len(values) >= 4:
                            if str(values[1]) == company and str(values[3]) == date:
                                group_data.append({
                                    "cic": str(values[0]),
                                    "chi": str(values[2]),
                                    "t_date": str(values[3])
                                })
                    if not group_data:
                        continue
                    df = pd.DataFrame(group_data)
                    output_file = os.path.join(output_dir, f"合併文件_{company}_{date}_{timestamp}.docx")
                    try:
                        self.fill_attendance_table_with_company_date(template_path, output_file, df, company, date)
                        total_count += 1
                    except Exception as e:
                        messagebox.showerror("錯誤", f"{company} {date} 產生失敗: {e}")
        
        # 保存當前的課程名稱、講授者、時間和地點等信息到配置文件
        # 添加課程名稱和講授者到選項列表中（如果不存在）
        current_subject = self.subject_var.get()
        current_trainer = self.trainer_var.get()
        
        if current_subject and current_subject not in self.subject_options:
            self.subject_options.append(current_subject)
            # 更新下拉選單選項
            if hasattr(self, 'subject_combo'):
                self.subject_combo['values'] = self.subject_options
                
        if current_trainer and current_trainer not in self.trainer_options:
            self.trainer_options.append(current_trainer)
            # 更新下拉選單選項
            if hasattr(self, 'trainer_combo'):
                self.trainer_combo['values'] = self.trainer_options
        
        # 保存配置
        self.save_morning_config()
        
        if total_count > 0:
            messagebox.showinfo("成功", "所有公司與日期的合併文件已產生！")
            self.status_var.set("已生成所有合併文件")
        else:
            messagebox.showwarning("警告", "沒有任何符合條件的資料產生文件")

    def reload_onboarding_data(self):
        """重新載入Excel數據並更新界面"""
        if not self.onboarding_excel_file.get():
            messagebox.showwarning("警告", "請先選擇Excel文件")
            return
            
        try:
            # 重新載入Excel數據
            self.load_onboarding_excel_data()
            
            # 更新狀態欄
            self.status_var.set(f"已重新載入Excel數據，共 {len(self.onboarding_excel_data)} 筆記錄")
            
            # 顯示成功消息
            messagebox.showinfo("成功", "Excel數據已重新載入")
            
        except Exception as e:
            messagebox.showerror("錯誤", f"重新載入Excel數據失敗: {e}")
            self.status_var.set("重新載入Excel數據失敗")

    def reload_morning_data(self):
        """重新載入早操Excel數據並更新界面"""
        if not self.excel_file.get():
            messagebox.showwarning("警告", "請先選擇Excel文件")
            return
            
        try:
            # 重新載入Excel數據
            self.load_excel_data()
            
            # 更新狀態欄
            self.status_var.set(f"已重新載入Excel數據，共 {len(self.excel_data)} 筆記錄")
            
            # 顯示成功消息
            messagebox.showinfo("成功", "Excel數據已重新載入")
            
        except Exception as e:
            messagebox.showerror("錯誤", f"重新載入Excel數據失敗: {e}")
            self.status_var.set("重新載入Excel數據失敗")

    def start_safety_training_mode(self):
        """啟動安全訓練模式界面"""
        # 清理選擇界面
        for widget in self.root.winfo_children():
            widget.destroy()

        self.root.title("安全訓練文件列印")
        self.root.geometry("900x700")

        # 初始化安全訓練模式所需的變量
        self.safety_excel_file = tk.StringVar()
        self.safety_template_file = tk.StringVar()
        self.safety_output_path = tk.StringVar()
        self.safety_excel_data = None
        self.status_var = tk.StringVar()
        self.status_var.set("就緒")
        # 先初始化 today/tittle 變數
        self.safety_today_var = tk.StringVar()
        self.safety_tittle_var = tk.StringVar()
        self.safety_today_var.set(datetime.now().strftime('%Y-%m-%d'))
        # 安全訓練模式的配置文件
        self.safety_config_file = resource_path('safety_training_config.json')
        # 加載安全訓練模式的配置
        self.load_safety_config()
        # 創建主框架
        main_frame = ttk.Frame(self.root, padding=10)
        main_frame.pack(fill=tk.BOTH, expand=True)
        self.create_safety_ui()
        # 添加狀態欄
        status_bar = ttk.Label(self.root, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W)
        status_bar.pack(side=tk.BOTTOM, fill=tk.X)

    def load_safety_config(self):
        """載入安全訓練模式的配置"""
        config = self.load_config('safety')
        self.safety_excel_file.set(config.get('excel_file', ''))
        self.safety_template_file.set(config.get('template_file', ''))
        self.safety_output_path.set(config.get('output_path', ''))
        self.safety_today_var.set(config.get('today', datetime.now().strftime('%Y-%m-%d')))
        self.safety_tittle_var.set(config.get('tittle', ''))
        self.safety_tittle_options = config.get('tittle_options', [])

    def save_safety_config(self):
        """保存安全訓練模式的配置"""
        config = {
            'excel_file': self.safety_excel_file.get(),
            'template_file': self.safety_template_file.get(),
            'output_path': self.safety_output_path.get(),
            'today': self.safety_today_var.get(),
            'tittle': self.safety_tittle_var.get(),
            'tittle_options': getattr(self, 'safety_tittle_options', [])
        }
        self.save_config('safety', config)

    def create_safety_ui(self):
        """創建安全訓練模式的GUI界面"""
        main_frame = ttk.Frame(self.root, padding=10)
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 按鈕區域
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=(0, 10))

        # 定義安全訓練模式的按鈕及其命令
        safety_buttons = [
            ("輸入Excel", self.select_safety_excel),
            ("選擇模板", self.select_safety_template),
            ("輸出位置選擇", self.select_safety_output),
            ("套用篩選", self.apply_safety_filter),
            ("清除篩選", self.clear_safety_filter),
            ("重新載入", self.reload_safety_data),
            ("生成Word", self.generate_safety_word),
            ("返回主菜單", self.return_to_main)
        ]

        for i, (btn_text, btn_command) in enumerate(safety_buttons):
            btn = ttk.Button(btn_frame, text=btn_text, command=btn_command)
            btn.grid(row=0, column=i, padx=5)

        # 設置區域
        settings_frame = ttk.LabelFrame(main_frame, text="設置", padding=10)
        settings_frame.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(settings_frame, text="Excel文件:").grid(row=0, column=0, sticky=tk.W, pady=5)
        ttk.Entry(settings_frame, textvariable=self.safety_excel_file, width=50).grid(row=0, column=1, sticky=tk.W+tk.E, padx=5)
        
        ttk.Label(settings_frame, text="模板文件:").grid(row=1, column=0, sticky=tk.W, pady=5)
        ttk.Entry(settings_frame, textvariable=self.safety_template_file, width=50).grid(row=1, column=1, sticky=tk.W+tk.E, padx=5)
        
        ttk.Label(settings_frame, text="輸出位置:").grid(row=2, column=0, sticky=tk.W, pady=5)
        ttk.Entry(settings_frame, textvariable=self.safety_output_path, width=50).grid(row=2, column=1, sticky=tk.W+tk.E, padx=5)

        # 篩選區域
        filter_frame = ttk.LabelFrame(main_frame, text="篩選條件", padding=10)
        filter_frame.pack(fill=tk.X, pady=(0, 10))

        # 日期多選列表框
        ttk.Label(filter_frame, text="日期：").grid(row=0, column=0, sticky=tk.NW, pady=5)
        date_list_frame = ttk.Frame(filter_frame)
        date_list_frame.grid(row=0, column=1, sticky=tk.NSEW, pady=5)
        
        self.safety_date_listbox = tk.Listbox(date_list_frame, selectmode=tk.MULTIPLE, exportselection=False, height=6, width=30)
        date_yscroll = ttk.Scrollbar(date_list_frame, orient=tk.VERTICAL, command=self.safety_date_listbox.yview)
        self.safety_date_listbox.configure(yscrollcommand=date_yscroll.set)
        date_xscroll = ttk.Scrollbar(date_list_frame, orient=tk.HORIZONTAL, command=self.safety_date_listbox.xview)
        self.safety_date_listbox.configure(xscrollcommand=date_xscroll.set)
        
        self.safety_date_listbox.grid(row=0, column=0, sticky=tk.NSEW)
        date_yscroll.grid(row=0, column=1, sticky=tk.NS)
        date_xscroll.grid(row=1, column=0, sticky=tk.EW)
        
        date_list_frame.grid_rowconfigure(0, weight=1)
        date_list_frame.grid_columnconfigure(0, weight=1)

        # 分判多選列表框
        ttk.Label(filter_frame, text="分判：").grid(row=0, column=2, sticky=tk.NW, pady=5, padx=(10,0))
        subcon_list_frame = ttk.Frame(filter_frame)
        subcon_list_frame.grid(row=0, column=3, sticky=tk.NSEW, pady=5)
        
        self.safety_subcon_listbox = tk.Listbox(subcon_list_frame, selectmode=tk.MULTIPLE, exportselection=False, height=6, width=30)
        subcon_yscroll = ttk.Scrollbar(subcon_list_frame, orient=tk.VERTICAL, command=self.safety_subcon_listbox.yview)
        self.safety_subcon_listbox.configure(yscrollcommand=subcon_yscroll.set)
        subcon_xscroll = ttk.Scrollbar(subcon_list_frame, orient=tk.HORIZONTAL, command=self.safety_subcon_listbox.xview)
        self.safety_subcon_listbox.configure(xscrollcommand=subcon_xscroll.set)
        
        self.safety_subcon_listbox.grid(row=0, column=0, sticky=tk.NSEW)
        subcon_yscroll.grid(row=0, column=1, sticky=tk.NS)
        subcon_xscroll.grid(row=1, column=0, sticky=tk.EW)
        
        subcon_list_frame.grid_rowconfigure(0, weight=1)
        subcon_list_frame.grid_columnconfigure(0, weight=1)

        # 預覽區域
        preview_frame = ttk.LabelFrame(main_frame, text="預覽", padding=10)
        preview_frame.pack(fill=tk.BOTH, expand=True)

        self.safety_preview_table = ttk.Treeview(preview_frame, show='headings')
        preview_yscroll = ttk.Scrollbar(preview_frame, orient=tk.VERTICAL, command=self.safety_preview_table.yview)
        preview_xscroll = ttk.Scrollbar(preview_frame, orient=tk.HORIZONTAL, command=self.safety_preview_table.xview)
        self.safety_preview_table.configure(yscrollcommand=preview_yscroll.set, xscrollcommand=preview_xscroll.set)
        
        preview_yscroll.pack(side=tk.RIGHT, fill=tk.Y)
        preview_xscroll.pack(side=tk.BOTTOM, fill=tk.X)
        self.safety_preview_table.pack(fill=tk.BOTH, expand=True)

        # 2. UI新增欄位（在 create_safety_ui 的 filter_frame 區域下方插入）
        ttk.Label(filter_frame, text="今日日期：").grid(row=1, column=0, sticky=tk.W, pady=5)
        ttk.Entry(filter_frame, textvariable=self.safety_today_var, width=20).grid(row=1, column=1, sticky=tk.W, pady=5)
        ttk.Label(filter_frame, text="課程名稱：").grid(row=1, column=2, sticky=tk.W, pady=5)
        self.safety_tittle_combo = ttk.Combobox(filter_frame, textvariable=self.safety_tittle_var, width=20)
        self.safety_tittle_combo['values'] = getattr(self, 'safety_tittle_options', [])
        self.safety_tittle_combo.grid(row=1, column=3, sticky=tk.W, pady=5)

    def load_safety_excel_data(self):
        """載入安全訓練Excel數據"""
        try:
            if not self.safety_excel_file.get():
                return
            self.safety_excel_data = pd.read_excel(self.safety_excel_file.get())
            
            # 更新預覽表格
            self.update_safety_preview_table(self.safety_excel_data)
            
            # 更新篩選列表框
            self.update_safety_filter_listboxes()
            
        except Exception as e:
            messagebox.showerror("錯誤", f"載入Excel數據失敗: {e}")

    def update_safety_filter_listboxes(self):
        """更新安全訓練篩選列表框的選項（日期只用K列原始內容）"""
        if self.safety_excel_data is None or self.safety_excel_data.empty:
            return

        # 清空現有選項
        self.safety_date_listbox.delete(0, tk.END)
        self.safety_subcon_listbox.delete(0, tk.END)

        try:
            # 僅用K列（第10列，索引10）作為日期來源
            date_col = self.safety_excel_data.iloc[:, 10].astype(str)
            dates = [d for d in date_col if d.lower() != 'nan' and d.strip() != '']
            # 保持原始順序且去重
            seen = set()
            unique_dates = []
            for d in dates:
                if d not in seen:
                    unique_dates.append(d)
                    seen.add(d)
            for d in unique_dates:
                self.safety_date_listbox.insert(tk.END, d)

            # 分判欄位維持原有邏輯
            subcon_columns = [col for col in self.safety_excel_data.columns if any(keyword in col for keyword in ['分判', '分判商', 'subcon', 'Subcon'])]
            if subcon_columns:
                subcon_column = subcon_columns[0]  # 使用找到的第一個分判欄位
                subcons = []
                for subcon in self.safety_excel_data[subcon_column]:
                    if pd.notna(subcon):
                        subcon_str = str(subcon).strip()
                        if subcon_str and subcon_str.lower() != 'nan':
                            subcons.append(subcon_str)
                unique_subcons = sorted(set(subcons))
                for subcon in unique_subcons:
                    self.safety_subcon_listbox.insert(tk.END, subcon)

        except Exception as e:
            print(f"更新篩選列表框時出錯: {e}")
            messagebox.showerror("錯誤", f"更新篩選選項時出錯: {e}")

    def apply_safety_filter(self):
        """應用安全訓練數據篩選（日期只用K列原始內容）"""
        if self.safety_excel_data is None or self.safety_excel_data.empty:
            messagebox.showinfo("提示", "請先載入安全訓練Excel數據")
            return
        filtered_df = self.safety_excel_data.copy()
        try:
            # 獲取並應用日期篩選條件
            selected_dates_indices = self.safety_date_listbox.curselection()
            selected_dates = [self.safety_date_listbox.get(i) for i in selected_dates_indices]
            # 應用日期篩選（只用K列）
            if selected_dates:
                date_col = filtered_df.iloc[:, 10].astype(str)
                date_mask = date_col.isin(selected_dates)
                filtered_df = filtered_df[date_mask]
            # 分判篩選維持原有邏輯
            selected_subcons_indices = self.safety_subcon_listbox.curselection()
            selected_subcons = [self.safety_subcon_listbox.get(i) for i in selected_subcons_indices]
            if selected_subcons:
                subcon_columns = [col for col in filtered_df.columns if any(keyword in col for keyword in ['分判', '分判商', 'subcon', 'Subcon'])]
                if subcon_columns:
                    subcon_column = subcon_columns[0]
                    subcon_mask = filtered_df[subcon_column].astype(str).apply(lambda x: x.strip() in selected_subcons)
                    filtered_df = filtered_df[subcon_mask]
            self.update_preview_table(self.safety_preview_table, filtered_df)
            self.safety_filtered_data = filtered_df
            record_count = len(filtered_df) if filtered_df is not None else 0
            self.status_var.set(f"已篩選 {record_count} 筆記錄")
        except Exception as e:
            print(f"篩選安全訓練資料時出錯: {e}")
            messagebox.showerror("錯誤", f"篩選資料時出錯: {e}")

    def clear_safety_filter(self):
        """清除安全訓練數據篩選"""
        try:
            # 清除列表框的選擇
            self.safety_date_listbox.selection_clear(0, tk.END)
            self.safety_subcon_listbox.selection_clear(0, tk.END)
            
            # 如果已加載數據，顯示所有數據
            if self.safety_excel_data is not None and not self.safety_excel_data.empty:
                self.update_safety_preview_table(self.safety_preview_table, self.safety_excel_data)
                self.status_var.set("已清除所有篩選")
            else:
                # 如果沒有數據，清空預覽表
                for i in self.safety_preview_table.get_children():
                    self.safety_preview_table.delete(i)
        except Exception as e:
            print(f"清除篩選時出錯: {e}")
            messagebox.showerror("錯誤", f"清除篩選時出錯: {e}")

    def fill_safety_table_with_records(self, template_path: str, output_path: str, data: pd.DataFrame, subcon_name: str):
        """
        將特定分判公司的資料填入Word表格，每頁15行（包含標題行）、5列（包含標題列）
        :param template_path: Word模板路徑
        :param output_path: 輸出Word路徑
        :param data: pandas DataFrame，包含該分判公司的所有記錄
        :param subcon_name: 分判公司名稱
        """
        print(f"開始處理分判公司: {subcon_name}")
        print(f"數據記錄數: {len(data)}")

        # 創建臨時文件
        temp_output = output_path + ".temp"
        shutil.copy2(template_path, temp_output)
        
        try:
            # 打開臨時文件，使用自定義的create_document函數
            doc = create_document(temp_output)
            
            # 查找目標表格
            if not doc.tables:
                raise Exception("模板中沒有找到表格")
            
            table = doc.tables[0]
            print(f"找到表格，行數: {len(table.rows)}")
            
            # 檢查表格結構
            if len(table.rows) < 15:
                raise Exception(f"表格必須有15行（包含標題行），當前行數: {len(table.rows)}")
            
            first_row = table.rows[0]
            if not first_row.cells:
                raise Exception("表格第一行沒有單元格")
                
            print(f"表格第一行單元格數: {len(first_row.cells)}")
            
            if len(first_row.cells) < 5:
                raise Exception(f"表格必須有5列（包含標題列），當前列數: {len(first_row.cells)}")
            
            # 保存原始表格的XML元素和標題行
            orig_tbl_xml = deepcopy(table._element)
            header_row = deepcopy(table.rows[0]._element)
            first_column = [cell.text for cell in [row.cells[0] for row in table.rows]]
            
            # 計算需要的頁數（每頁14行數據，1行標題）
            total_records = len(data)
            records_per_page = 14  # 每頁14行數據（第1行是標題）
            total_pages = (total_records + records_per_page - 1) // records_per_page
            print(f"總頁數: {total_pages}")
            
            # 處理每一頁
            for page_num in range(total_pages):
                print(f"處理第 {page_num + 1} 頁")
                current_table = table if page_num == 0 else doc.tables[-1]
                
                if page_num > 0:
                    # 插入分頁符
                    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                    # 複製原始表格
                    new_tbl_element = deepcopy(orig_tbl_xml)
                    doc._body._element.append(new_tbl_element)
                    current_table = doc.tables[-1]
                
                # 獲取當前頁的記錄
                start_idx = page_num * records_per_page
                end_idx = min((page_num + 1) * records_per_page, total_records)
                page_records = data.iloc[start_idx:end_idx]
                
                # 確保表格有15行
                while len(current_table.rows) < 15:
                    current_table.add_row()
                
                # 保持第一行（標題行）不變
                # 填充數據（從第2行開始，索引1）
                for i, (_, record) in enumerate(page_records.iterrows(), start=1):
                    try:
                        row = current_table.rows[i]
                        
                        # 第一列：序號（自動生成）
                        row.cells[0].text = f"{start_idx + i}."
                        
                        # 第二列：工人姓名
                        name_val = record.get('工人姓名', '')
                        row.cells[1].text = str(name_val) if pd.notna(name_val) else ''
                        
                        # 第三列：分判
                        row.cells[2].text = str(subcon_name)
                        
                        # 第四列：平安咭號碼
                        card_val = record.get('平安咭號編號', '')
                        row.cells[3].text = str(card_val) if pd.notna(card_val) else ''

                        # 第五列：簽名欄（空白）
                        row.cells[4].text = ''
                        
                        # 設置數據單元格居中對齊（除了第一列）
                        for cell in row.cells:
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                    except Exception as e:
                        print(f"處理第 {i} 行時出錯: {e}")
                        continue
                
                # 清空剩餘的行（保持第一列不變）
                for i in range(len(page_records) + 1, 15):
                    try:
                        row = current_table.rows[i]
                        # 保持第一列（標題列）不變
                        row.cells[0].text = ''  # 清空序號
                        # 清空其他列
                        for cell in row.cells[1:]:
                            cell.text = ''
                    except Exception as e:
                        print(f"清空第 {i} 行時出錯: {e}")
            
            # 替換文檔中的合併欄位
            def replace_merge_fields(paragraph):
                for run in paragraph.runs:
                    if '«date»' in run.text:
                        run.text = run.text.replace('«date»', datetime.now().strftime('%Y-%m-%d'))
                    if '«subcon»' in run.text:
                        run.text = run.text.replace('«subcon»', subcon_name)
                    if '«Green_Card»' in run.text and len(data) > 0:
                        card_val = data.iloc[0].get('平安咭號編號', '')
                        run.text = run.text.replace('«Green_Card»', str(card_val) if pd.notna(card_val) else '')
                    if '«total»' in run.text:
                        run.text = run.text.replace('«total»', str(total_records))
                    if '«today»' in run.text:
                        run.text = run.text.replace('«today»', self.safety_today_var.get())
                    if '«tittle»' in run.text:
                        run.text = run.text.replace('«tittle»', self.safety_tittle_var.get())

            # 替換正文中的合併欄位
            for paragraph in doc.paragraphs:
                replace_merge_fields(paragraph)

            # 替換表格中的合併欄位
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_merge_fields(paragraph)

            # 替換頁眉頁腳中的合併欄位
            for section in doc.sections:
                # 處理頁眉
                if section.header:
                    for paragraph in section.header.paragraphs:
                        replace_merge_fields(paragraph)
                    for table in section.header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    replace_merge_fields(paragraph)
                
                # 處理頁腳
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        replace_merge_fields(paragraph)
                    for table in section.footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    replace_merge_fields(paragraph)
            
            # 保存文件
            doc.save(output_path)
            print(f"文件已保存: {output_path}")
            
        except Exception as e:
            print(f"處理文件時出錯: {str(e)}")
            raise e
            
        finally:
            # 清理臨時文件
            try:
                if os.path.exists(temp_output):
                    os.remove(temp_output)
            except Exception as e:
                print(f"清理臨時文件時出錯: {e}")
                pass

    def generate_safety_word(self):
        """生成安全訓練Word文件"""
        if self.safety_excel_data is None or self.safety_excel_data.empty:
            messagebox.showwarning("警告", "請先載入Excel文件")
            return

        template_path = self.safety_template_file.get()
        if not template_path or not os.path.exists(template_path):
            # 如果模板路径为空或不存在，提示用户选择模板
            messagebox.showwarning("警告", "請選擇安全訓練Word模板文件\n\n所有模板必須從外部讀取，系統不提供默認模板")
            self.select_safety_template()
            return
            
        # 获取当前选择的模板路径（可能已经被select_template更新）
        template_path = self.safety_template_file.get()
        if not template_path or not os.path.exists(template_path):
            # 用户取消了模板选择
            messagebox.showinfo("提示", "已取消選擇模板，無法生成Word文件")
            return

        output_dir = self.safety_output_path.get()
        if not output_dir:
            output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output_safety")
            self.safety_output_path.set(output_dir)
            self.save_safety_config()
        
        os.makedirs(output_dir, exist_ok=True)

        # 獲取預覽表中的數據
        filtered_data = None
        if hasattr(self, 'safety_filtered_data') and self.safety_filtered_data is not None and not self.safety_filtered_data.empty:
            filtered_data = self.safety_filtered_data
        else:
            filtered_data = self.safety_excel_data

        # 檢查是否有分包商列
        if '分判' not in filtered_data.columns:
            messagebox.showwarning("警告", "Excel數據中缺少'分判'列，這可能會影響文件生成")
        
        # 按分包商分組生成文件
        subcontractors = filtered_data['分判'].unique() if '分判' in filtered_data.columns else ['未知分包商']
        
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        
        total_files = 0
        for subcon in subcontractors:
            if pd.isna(subcon):
                subcon_name = "未知分包商"
            else:
                subcon_name = str(subcon)
            
            # 獲取該分包商的所有記錄
            if '分判' in filtered_data.columns:
                subcon_df = filtered_data[filtered_data['分判'] == subcon]
            else:
                subcon_df = filtered_data
            
            if subcon_df.empty:
                continue
            
            # 創建輸出文件路徑
            output_filename = os.path.join(output_dir, f"{subcon_name}入職安全訓練_{timestamp}.docx")
            
            try:
                # 使用fill_safety_table_with_records方法填充表格
                self.fill_safety_table_with_records(template_path, output_filename, subcon_df, str(subcon_name))
                total_files += 1
                self.status_var.set(f"已生成 {subcon_name} 的安全訓練文件")
            except Exception as e:
                messagebox.showerror("錯誤", f"生成 {subcon_name} 的安全訓練文件時出錯: {e}")
        
        if total_files > 0:
            messagebox.showinfo("成功", f"已成功生成 {total_files} 個安全訓練Word文件")
            # 保存配置
            self.save_safety_config()
        else:
            messagebox.showwarning("警告", "沒有生成任何安全訓練Word文件")
        
        self.status_var.set(f"已生成 {total_files} 個安全訓練Word文件")

    def select_safety_excel(self):
        """選擇安全訓練Excel文件"""
        file_path = filedialog.askopenfilename(
            title="選擇安全訓練Excel文件",
            filetypes=[("Excel files", "*.xlsx *.xls"), ("All files", "*.*")]
        )
        if file_path:
            self.safety_excel_file.set(file_path)
            self.load_safety_excel_data()
            self.save_safety_config()
            self.status_var.set(f"已選擇Excel文件: {os.path.basename(file_path)}")

    def select_safety_template(self):
        """選擇安全訓練Word模板文件"""
        file_path = filedialog.askopenfilename(
            title="選擇安全訓練Word模板文件",
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")],
            initialdir=os.path.dirname(self.safety_template_file.get()) if self.safety_template_file.get() else None
        )
        if file_path:
            self.safety_template_file.set(file_path)
            self.save_safety_config()
            self.status_var.set(f"已選擇模板: {os.path.basename(file_path)}")
            
            # 嘗試讀取模板中的欄位信息
            try:
                with MailMerge(file_path) as document:
                    fields = document.get_merge_fields()
                    if fields:
                        field_list = ", ".join(fields)
                        messagebox.showinfo("模板欄位", f"模板包含以下合併欄位:\n{field_list}")
            except Exception as e:
                messagebox.showwarning("警告", f"讀取模板欄位時出錯: {e}")

    def select_safety_output(self):
        """選擇安全訓練文件輸出位置"""
        dir_path = filedialog.askdirectory(title="選擇安全訓練文件輸出目錄")
        if dir_path:
            self.safety_output_path.set(dir_path)
            self.save_safety_config()

    def update_safety_preview_table(self, df):
        """更新安全訓練預覽表格"""
        for i in self.safety_preview_table.get_children():
            self.safety_preview_table.delete(i)
        if df is None or df.empty:
            return
        self.safety_preview_table["columns"] = list(df.columns)
        for col in df.columns:
            self.safety_preview_table.heading(col, text=col)
            self.safety_preview_table.column(col, width=100)
        for index, row in df.iterrows():
            self.safety_preview_table.insert("", "end", values=list(row))

    def reload_safety_data(self):
        """重新載入安全訓練Excel數據"""
        if not self.safety_excel_file.get():
            messagebox.showwarning("警告", "請先選擇Excel文件")
            return
        try:
            self.load_safety_excel_data()
            self.status_var.set("已重新載入Excel數據")
            messagebox.showinfo("成功", "Excel數據已重新載入")
        except Exception as e:
            messagebox.showerror("錯誤", f"重新載入Excel數據失敗: {e}")
            self.status_var.set("重新載入Excel數據失敗")

    def load_morning_config(self):
        """載入早操模式的配置"""
        config = self.load_config('morning')
        self.excel_file.set(config.get('excel_file', ''))
        self.output_path.set(config.get('output_path', ''))
        self.template_file_path.set(config.get('template_file', ''))
        self.subject_var.set(config.get('subject', ''))
        self.trainer_var.set(config.get('trainer', ''))
        self.time_var.set(config.get('time', ''))
        self.place_var.set(config.get('place', ''))
        self.subject_options = config.get('subject_options', [])
        self.trainer_options = config.get('trainer_options', [])
        
        # 載入合併選項的狀態
        if hasattr(self, 'merge_companies_var'):
            self.merge_companies_var.set(config.get('merge_companies', False))

    def save_morning_config(self):
        """保存早操模式的配置"""
        config = {
            'excel_file': self.excel_file.get(),
            'output_path': self.output_path.get(),
            'template_file': self.template_file_path.get(),
            'subject': self.subject_var.get(),
            'trainer': self.trainer_var.get(),
            'time': self.time_var.get(),
            'place': self.place_var.get(),
            'subject_options': self.subject_options,
            'trainer_options': self.trainer_options,
            'merge_companies': self.merge_companies_var.get() if hasattr(self, 'merge_companies_var') else False
        }
        self.save_config('morning', config)

    def generate_attendance_word(self):
        """根據日期和公司生成出席記錄Word文件"""
        if self.excel_data is None:
            messagebox.showwarning("警告", "請先載入Excel文件")
            return
            
        template_path = self.template_file_path.get()
        if not template_path or not os.path.exists(template_path):
            # 如果模板路径为空或不存在，提示用户选择模板
            messagebox.showwarning("警告", "請選擇有效的Word模板文件")
            self.select_template()
            return
            
        # 获取当前选择的模板路径（可能已经被select_template更新）
        template_path = self.template_file_path.get()
        if not template_path or not os.path.exists(template_path):
            # 用户取消了模板选择
            return

    def update_filter_results(self):
        """根據選擇的篩選條件更新結果預覽"""
        try:
            # 獲取所選的公司
            selected_companies = []
            for i in self.company_listbox.curselection():
                selected_companies.append(self.company_listbox.get(i))
            
            # 獲取所選的日期
            selected_dates = []
            for i in self.date_listbox.curselection():
                selected_dates.append(self.date_listbox.get(i))
            
            # 清空預覽樹
            for i in self.preview_tree.get_children():
                self.preview_tree.delete(i)
            
            # 如果沒有選擇任何篩選條件，顯示所有數據
            if not selected_companies and not selected_dates:
                self.status_var.set("未選擇篩選條件，顯示所有記錄")
                return
            
            # 根據選擇的篩選條件更新結果
            if not hasattr(self, 'excel_data') or self.excel_data is None:
                self.status_var.set("未載入Excel數據")
                return
            
            try:
                company_col = self.excel_data.iloc[:, 21].astype(str)
                date_col = self.excel_data.iloc[:, 10].astype(str)  # K列是索引10
                cwrno_col = self.excel_data.iloc[:, 8].astype(str)
                name_col = self.excel_data.iloc[:, 11].astype(str)
                
                mask = pd.Series([True] * len(self.excel_data))
                if selected_companies:
                    mask = mask & company_col.isin(selected_companies)
                
                if selected_dates:
                    # 直接使用字符串比較，因為listbox中的日期已經與Excel中的格式統一
                    date_mask = pd.Series([False] * len(self.excel_data))
                    for i, val in enumerate(date_col):
                        if pd.notna(val):
                            val_str = str(val).strip()
                            if val_str in selected_dates:
                                date_mask.iloc[i] = True
                    mask = mask & date_mask
                    
                filtered_indices = mask[mask].index
                record_count = 0
                for i in filtered_indices:
                    try:
                        self.preview_tree.insert("", "end", values=(
                            cwrno_col.iloc[i],
                            company_col.iloc[i],
                            name_col.iloc[i],
                            date_col.iloc[i]
                        ))
                        record_count += 1
                    except Exception as e:
                        print(f"插入篩選記錄 {i} 時出錯: {e}")
                        continue
                self.status_var.set(f"已篩選 {len(filtered_indices)} 筆記錄")
            except Exception as e:
                print(f"篩選數據時出錯: {e}")
                messagebox.showerror("錯誤", f"篩選數據時出錯: {e}")
        except Exception as e:
            print(f"更新篩選結果時出錯: {e}")
            messagebox.showerror("錯誤", f"更新篩選結果時出錯: {e}")

if __name__ == "__main__":
    root = tk.Tk()
    app = MailMergeApp(root)
    root.mainloop()